#!/usr/bin/env python3
"""
Generate formatted DOCX for Calvince Odhiambo BBM 415 CAT
Moi University - Management of Financial Institutions
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_run(run, bold=False, italic=False, size=12, color=(0, 0, 0)):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(*color)


def set_para_spacing(para, before=0, after=6, spacing=1.5):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if spacing == 1.5:
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    elif spacing == 1.0:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    else:
        pf.line_spacing = Pt(spacing * 12)


def cover_line(doc, text, size=12, bold=False, after=6, align=WD_ALIGN_PARAGRAPH.CENTER):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run(run, bold=bold, size=size)
    para.alignment = align
    set_para_spacing(para, before=0, after=after, spacing=1.0)
    return para


def body_para(doc, text, bold=False, italic=False, indent=0, before=0, after=6,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run(run, bold=bold, italic=italic, size=12)
    para.alignment = align
    set_para_spacing(para, before=before, after=after, spacing=1.5)
    if indent:
        para.paragraph_format.left_indent = Inches(indent)
    return para


def question_heading(doc, text, before=14, after=6):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run(run, bold=True, size=12)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(para, before=before, after=after, spacing=1.5)
    return para


def question_text(doc, text, before=4, after=6):
    """The actual question statement — bold + italic."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run(run, bold=True, italic=True, size=12)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(para, before=before, after=after, spacing=1.5)
    return para


def sub_heading(doc, text, before=10, after=4):
    """Sub-point heading like 'Commercial Banks', 'Duration Gap Analysis'."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run(run, bold=True, size=12)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(para, before=before, after=after, spacing=1.5)
    return para


def bullet_para(doc, text, bold_prefix=None, before=0, after=4):
    """Bullet / numbered point. bold_prefix is the label part (e.g. 'i) ')."""
    para = doc.add_paragraph()
    if bold_prefix:
        r1 = para.add_run(bold_prefix)
        set_run(r1, bold=True, size=12)
        r2 = para.add_run(text)
        set_run(r2, size=12)
    else:
        run = para.add_run(text)
        set_run(run, size=12)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(para, before=before, after=after, spacing=1.5)
    para.paragraph_format.left_indent = Inches(0.35)
    return para


def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)


def generate_cat_docx(output='Calvince_BBM_415_CAT.docx'):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.0)

    # ======================== COVER PAGE ========================
    # Logo
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(logo_para, before=12, after=10, spacing=1.0)
    try:
        logo_para.add_run().add_picture(
            'assets/moi_uni_logo/moi_logo.png', width=Inches(1.4))
    except Exception:
        pass

    cover_line(doc, 'MOI UNIVERSITY', size=16, bold=True, after=4)
    cover_line(doc, 'ANNEX CAMPUS', size=14, bold=True, after=4)
    cover_line(doc, 'SCHOOL OF BUSINESS AND ECONOMICS', size=13, bold=True, after=24)

    # Details table on cover
    details = [
        ('COURSE CODE',    'BBM 415'),
        ('COURSE TITLE',   'MANAGEMENT OF FINANCIAL INSTITUTIONS'),
        ('NAME',           'ODHIAMBO CALVINCE'),
        ('REG NUMBER',     'BBM/1483/23'),
        ('TASK',           'CAT'),
        ('DATE',           '3RD APRIL 2026'),
        ('LECTURER',       'DR. JAPHET KOGEI'),
    ]
    for label, value in details:
        para = doc.add_paragraph()
        r1 = para.add_run(f'{label:<16}: ')
        set_run(r1, bold=True, size=12)
        r2 = para.add_run(value)
        set_run(r2, size=12)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_spacing(para, before=0, after=8, spacing=1.0)

    add_page_break(doc)

    # ======================== QUESTION ONE ========================
    question_heading(doc, 'QUESTION ONE')
    body_para(doc, 'Financial institutions occupy a unique position in the economy due to their '
              'special characteristics and operational frameworks that distinguish them from '
              'non-financial enterprises.')

    question_text(doc, 'a) Discuss the nature of financial institutions and explain why they are '
                  'considered unique in terms of their operations and risk exposure. (5 Marks)')

    body_para(doc, 'Financial institutions are organizations that provide financial services for '
              'their clients or members. They differ significantly from non-financial enterprises '
              'like manufacturing or retail firms because their primary inventory is money and '
              'financial contracts rather than physical goods.')
    body_para(doc, 'The fundamental nature of a financial institution is to act as a financial '
              'intermediary. They bridge the gap between "surplus units" (Savers/Investors) and '
              '"deficit units" (borrowers), ensuring the efficient allocation of capital within '
              'the economy.')

    sub_heading(doc, 'Reasons Why Financial Institutions are Unique', before=10, after=4)

    bullet_para(doc,
                'They borrow short term and lend long term. Depositors can withdraw funds on '
                'demand, but loans may be extended for years. This maturity mismatch creates '
                'significant liquidity risk.',
                bold_prefix='i) Maturity Transformation:  ')
    bullet_para(doc,
                'Financial institutions operate with high levels of debt relative to equity. '
                'Banks, for example, fund a significant portion of their operations using '
                'deposits and borrowed funds, making them more vulnerable to small shocks or '
                'adverse economic events.',
                bold_prefix='ii) High Leverage:  ')
    bullet_para(doc,
                'Financial institutions transform risky assets into less risky claims for '
                'savers. For example, a bank may pool many risky individual loans but provide '
                'depositors with relatively safe and liquid claims on those assets.',
                bold_prefix='iii) Risk Transformation:  ')

    sub_heading(doc, 'Unique Risk Exposure', before=10, after=4)
    body_para(doc, 'Because of their specific operations, financial institutions face risks that '
              'are far more volatile than those of non-financial firms:')
    bullet_para(doc,
                'Because they deal in long-term assets and short-term liabilities, sudden '
                'changes in market interest rates can drastically impact their profit margins.',
                bold_prefix='Interest Rate Risk:  ')
    bullet_para(doc,
                'The primary risk is that borrowers will fail to repay their loans, which '
                'directly affects the financial institution\'s solvency and capital base.',
                bold_prefix='Credit/Default Risk:  ')
    bullet_para(doc,
                'A sudden surge in depositor withdrawals can leave a bank unable to meet its '
                'short-term obligations, even if it is technically solvent.',
                bold_prefix='Liquidity Risk:  ')

    question_text(doc,
                  'b) With reference to at least two different types of depository institutions '
                  '(banks, insurance companies, or security firms), explain how their operational '
                  'characteristics expose them to different types of risks. (5 Marks)',
                  before=12)

    body_para(doc, 'The way a financial institution operates — what it takes in as liabilities '
              'and what it invests in as assets — determines the specific risks it faces. Two '
              'key types of financial institutions are examined below.')

    sub_heading(doc, 'i. Commercial Banks', before=10, after=4)
    body_para(doc, 'Banks take in deposits (savings and checking accounts) that customers can '
              'withdraw at any time, then use those deposits to issue long-term loans (mortgages '
              'and business loans). This business model exposes them to the following risks:')
    bullet_para(doc,
                'Because liabilities are on demand but assets are locked in for years, a bank '
                'faces risk if too many depositors want their money back simultaneously (a '
                '"bank run").',
                bold_prefix='Liquidity Risk:  ')
    bullet_para(doc,
                'The risk that borrowers will not repay their loans, leading to a loss of the '
                'bank\'s principal and interest income.',
                bold_prefix='Credit/Default Risk:  ')
    bullet_para(doc,
                'If market interest rates rise, the bank must pay more to retain its depositors, '
                'but the income from its older, fixed-rate loans stays the same, squeezing '
                'profit margins.',
                bold_prefix='Interest Rate Risk:  ')

    sub_heading(doc, 'ii. Insurance Companies', before=10, after=4)
    body_para(doc, 'Insurance companies collect premiums in exchange for a promise to pay out if '
              'a specific event occurs (death, accident, fire). They invest these premiums in '
              'long-term, safe securities (such as government bonds) to ensure they have funds '
              'for future claims. Their resulting risks include:')
    bullet_para(doc,
                'The risk that actual claims paid out are much higher than the premiums '
                'collected — for example, due to a natural disaster or an unexpected spike '
                'in mortality rates.',
                bold_prefix='Underwriting Risk:  ')
    bullet_para(doc,
                'Since they hold long-duration assets, insurance companies are highly sensitive '
                'to market crashes or inflation that devalues their long-term bond holdings, '
                'reducing their ability to meet future claim obligations.',
                bold_prefix='Investment Risk:  ')

    add_page_break(doc)

    # ======================== QUESTION TWO ========================
    question_heading(doc, 'QUESTION TWO')

    question_text(doc,
                  'a) Define interest rate risk and explain the two primary sources of interest '
                  'rate risk exposure in a bank\'s balance sheet. (5 Marks)')

    body_para(doc, 'Interest rate risk is the potential for a bank\'s financial condition — '
              'specifically its earnings (Net Interest Income) and the economic value of its '
              'equity — to be adversely affected by movements in market interest rates. It is '
              'one of the most significant financial risks facing depository institutions '
              'such as banks.')

    sub_heading(doc, 'Two Primary Sources of Interest Rate Risk', before=10, after=4)

    bullet_para(doc,
                'This is the most common form of interest rate risk. It arises from timing '
                'differences in the maturity (for fixed-rate instruments) and repricing '
                '(for floating-rate instruments) of bank assets, liabilities, and off-balance '
                'sheet positions. For example, if a bank funds a long-term, fixed-rate mortgage '
                'with short-term certificates of deposit (a liability), it faces risk: if '
                'interest rates rise, the bank must pay more to depositors to retain funding, '
                'while the income from the mortgage remains fixed, squeezing profit margins.',
                bold_prefix='1. Repricing Risk (Maturity Mismatch):  ')
    bullet_para(doc,
                'This occurs when the interest rates on different instruments change by '
                'different amounts or at different times, even if they have similar repricing '
                'frequencies. For example, a bank might have a loan priced based on the prime '
                'rate but funded by deposits priced based on LIBOR. Even if both rates move in '
                'the same direction, if the cost of LIBOR-based deposits rises faster than the '
                'income from prime-based loans, the bank\'s net interest margin will decline.',
                bold_prefix='2. Basis Risk:  ')

    question_text(doc,
                  'b) Discuss the concept of duration as a measure of interest rate sensitivity, '
                  'and illustrate with a practical example how duration gap analysis helps a bank '
                  'manage its interest rate risk exposure. (5 Marks)',
                  before=12)

    body_para(doc, 'Duration is a comprehensive measure of the timing of the cash flows of a '
              'financial instrument. Unlike simple maturity, it accounts for the size and timing '
              'of all interest and principal payments, providing a more precise measure of '
              'interest rate sensitivity.')

    sub_heading(doc, 'Key Properties of Duration', before=10, after=4)
    bullet_para(doc,
                'Duration measures how much the price (or market value) of a financial asset '
                'or liability changes when interest rates change.',
                bold_prefix='Sensitivity:  ')
    bullet_para(doc,
                'There is an inverse relationship between interest rates and the value of '
                'fixed-income instruments. If interest rates rise, the market value of the '
                'instrument falls.',
                bold_prefix='The Inverse Relationship:  ')
    bullet_para(doc,
                'A duration of 5 years implies that for every 1% increase in interest rates, '
                'the value of the instrument will decrease by approximately 5%.',
                bold_prefix='The Rule of Thumb:  ')

    sub_heading(doc, 'Duration Gap Analysis', before=10, after=4)
    body_para(doc, 'The Duration Gap (DGAP) measures the sensitivity of a bank\'s net worth to '
              'changes in interest rates and is calculated as:')
    body_para(doc, 'DGAP = DA \u2212 (L/A \u00d7 DL)',
              bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=4)

    for line in [
        'DA = Average duration of assets',
        'DL = Average duration of liabilities',
        'L/A = Ratio of total liabilities to total assets',
    ]:
        bullet_para(doc, line)

    sub_heading(doc, 'Practical Example', before=10, after=4)
    body_para(doc, 'Imagine a bank with the following balance sheet:')
    for line in [
        'Assets (A): Ksh 100 million in long-term loans with a duration (DA) of 6 years.',
        'Liabilities (L): Ksh 90 million in short-term deposits with a duration (DL) of 1 year.',
        'Equity: Ksh 10 million.',
    ]:
        bullet_para(doc, line)

    body_para(doc, 'DGAP = 6 \u2212 (90/100 \u00d7 1) = 6 \u2212 0.9 = 5.1 years',
              bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=4)

    body_para(doc, 'This large positive DGAP means the bank\'s assets have a much longer '
              'duration than its liabilities. If interest rates rise by 1%:')
    bullet_para(doc,
                'The value of the Ksh 100 million assets would drop by roughly 6% '
                '(\u2212Ksh 6 million).')
    bullet_para(doc,
                'The value of the Ksh 90 million liabilities would drop by only 1% '
                '(\u2212Ksh 0.9 million).')
    bullet_para(doc,
                'The bank\'s equity would therefore decline by Ksh 5.1 million '
                '(Ksh 6m loss on assets offset by only Ksh 0.9m gain on liabilities), '
                'wiping out over half of the bank\'s capital.')

    body_para(doc, 'This analysis shows that to manage interest rate risk, the bank should '
              'seek to reduce its DGAP — either by shortening the duration of its assets, '
              'extending the duration of its liabilities, or using interest rate derivatives '
              'such as swaps.')

    add_page_break(doc)

    # ======================== QUESTION THREE ========================
    question_heading(doc, 'QUESTION THREE')
    body_para(doc, 'Credit risk remains a fundamental concern for all financial institutions '
              'engaged in lending activities.')

    question_text(doc,
                  'a) Explain what credit risk is and discuss the key components that determine '
                  'a borrower\'s creditworthiness when financial institutions evaluate loan '
                  'applications. (5 Marks)')

    body_para(doc, 'Credit risk is the risk that a borrower or counterparty fails to meet '
              'contractual obligations (principal and/or interest), leading to default losses '
              'for the lender. Essentially, it is the possibility of a loss resulting from a '
              'borrower\'s failure to repay a loan or meet contractual obligations.')

    sub_heading(doc, 'Key Components of Creditworthiness: The 5 Cs of Credit',
                before=10, after=4)
    body_para(doc, 'When financial institutions evaluate loan applications, they typically use '
              'a framework known as the 5 Cs of Credit to determine a borrower\'s '
              'creditworthiness:')

    for label, text in [
        ('1) Character:  ',
         'This refers to the borrower\'s reputation or track record for repaying debts. '
         'Lenders examine credit reports and credit scores to see how reliably the applicant '
         'has handled credit in the past. It reflects the borrower\'s honesty and reliability.'),
        ('2) Capital:  ',
         'This represents the borrower\'s own investment in the project or their overall net '
         'worth. A large down payment on a home or significant savings indicate the borrower '
         'has "skin in the game," which reduces the risk for the lender.'),
        ('3) Collateral:  ',
         'Collateral consists of assets (such as a house, car, or equipment) that the borrower '
         'pledges as security for the loan. If the borrower defaults, the lender can seize the '
         'collateral to recover their losses.'),
        ('4) Capacity:  ',
         'This measures the borrower\'s ability to repay the loan. Lenders compare the '
         'borrower\'s income against their recurring debts (the debt-to-income or DTI ratio), '
         'examining cash flow, job stability, and other income sources.'),
        ('5) Conditions:  ',
         'This refers to external factors that could affect the borrower\'s ability to pay, '
         'including the purpose of the loan (e.g., business expansion), the state of the '
         'broader economy, and prevailing interest rate conditions.'),
    ]:
        bullet_para(doc, text, bold_prefix=label)

    question_text(doc,
                  'b) Discuss the various methods that financial institutions can employ to '
                  'manage and mitigate credit risk exposure in their lending portfolios. '
                  'Illustrate your answer with at least two practical examples. (5 Marks)',
                  before=12)

    body_para(doc, 'Credit risk is the possibility of a loss resulting from a borrower\'s '
              'failure to repay a loan or meet contractual obligations. To mitigate this risk, '
              'financial institutions employ several strategies:')

    bullet_para(doc,
                'Before lending, institutions evaluate a borrower\'s creditworthiness using '
                'the 5 Cs of Credit framework — Character, Capacity, Capital, Collateral, and '
                'Conditions. This ensures only creditworthy borrowers receive loans.',
                bold_prefix='i) Credit Risk Assessment (5 Cs of Credit):  ')
    bullet_para(doc,
                'Financial institutions spread loans across different industries, geographic '
                'regions, and types of borrowers. A downturn in one sector (like agriculture) '
                'does not then collapse the entire lending portfolio.',
                bold_prefix='ii) Portfolio Diversification:  ')
    bullet_para(doc,
                'Institutions charge higher interest rates to riskier borrowers. This "risk '
                'premium" compensates the bank for the higher probability of default.',
                bold_prefix='iii) Risk-Based Pricing:  ')
    bullet_para(doc,
                'Requiring borrowers to pledge assets as security reduces potential losses in '
                'the event of default.',
                bold_prefix='iv) Collateralization:  ')

    sub_heading(doc, 'Practical Examples', before=10, after=4)
    bullet_para(doc,
                'When a bank issues a home loan (mortgage), the house itself serves as '
                'collateral. If the homeowner stops making payments, the bank has the legal '
                'right to foreclose on the property. By selling the house, the bank recovers '
                'the outstanding loan amount, significantly reducing its credit risk exposure.',
                bold_prefix='Example 1 \u2014 Mortgage Lending (Collateralization):  ')
    bullet_para(doc,
                'A commercial bank may lend to hundreds of different small businesses ranging '
                'from local restaurants to technology startups. By diversifying its portfolio, '
                'the bank ensures that if the restaurant industry struggles, the success of '
                'tech firms can offset those losses. Additionally, the bank might include a '
                'covenant requiring businesses to provide quarterly financial statements, '
                'enabling close monitoring of financial health.',
                bold_prefix='Example 2 \u2014 Small Business Loans (Diversification '
                             'and Covenants):  ')

    add_page_break(doc)

    # ======================== QUESTION FOUR ========================
    question_heading(doc, 'QUESTION FOUR')
    body_para(doc, 'Capital adequacy and deposit insurance are two critical regulatory '
              'mechanisms designed to protect the stability of financial institutions and '
              'safeguard depositors.')

    question_text(doc,
                  'a) Explain the concept of capital adequacy ratios and discuss why regulatory '
                  'authorities impose minimum capital requirements on financial institutions. '
                  '(5 Marks)')

    body_para(doc, 'The Capital Adequacy Ratio (CAR), also known as the Capital-to-Risk '
              'Weighted Assets Ratio (CRAR), is a measurement of a bank\'s available capital '
              'expressed as a percentage of its risk-weighted credit exposure. It is used to '
              'protect depositors and promote the stability and efficiency of financial systems '
              'around the world.')

    body_para(doc, 'The Basel III framework, for instance, requires banks to maintain a minimum '
              'CAR of 8%, with Tier 1 capital (core capital such as equity and retained '
              'earnings) forming the most critical component. In Kenya, the Central Bank of '
              'Kenya mandates a minimum CAR of 14.5% for commercial banks.')

    sub_heading(doc, 'Reasons Why Regulatory Authorities Impose Minimum Capital Requirements',
                before=10, after=4)

    for label, text in [
        ('i) Protection of Depositors:  ',
         'By ensuring a bank has enough of its own "skin in the game," regulators protect '
         'depositors\' money. If a bank fails, its capital is used to absorb losses and '
         'pay off liabilities before depositors lose their savings.'),
        ('ii) Restricting Excessive Risk-Taking:  ',
         'When banks are required to hold more capital against riskier assets, it creates a '
         'financial disincentive for them to take on excessive high-risk debt, encouraging '
         'more prudent lending and investment practices.'),
        ('iii) Promoting Financial Stability (Systemic Risk Mitigation):  ',
         'The failure of one large bank can trigger a "domino effect" across the entire '
         'economy. Minimum capital requirements reduce the likelihood of individual bank '
         'failures, thereby protecting the broader financial system.'),
        ('iv) Maintaining Public Confidence:  ',
         'The banking system relies fundamentally on trust. Knowing that banks are required '
         'to maintain a certain level of financial strength prevents bank runs — where many '
         'customers simultaneously withdraw funds out of fear of insolvency.'),
    ]:
        bullet_para(doc, text, bold_prefix=label)

    question_text(doc,
                  'b) Discuss the role of deposit insurance as a liability guarantee mechanism, '
                  'and explain how it protects depositors while also creating moral hazard '
                  'issues that regulators must address. (5 Marks)',
                  before=12)

    body_para(doc, 'Deposit insurance is a protective measure implemented in many countries to '
              'safeguard bank depositors, in full or in part, from losses caused by a bank\'s '
              'inability to pay its debts when due. It acts as a "safety net" for the banking '
              'system, underpinning public confidence in financial institutions.')

    sub_heading(doc, 'Role as a Liability Guarantee Mechanism', before=10, after=4)
    bullet_para(doc,
                'Deposit insurance prevents a localized bank failure from spiraling into a '
                'systemic financial crisis by maintaining public trust in the banking sector '
                'as a whole.',
                bold_prefix='Systemic Protection:  ')
    bullet_para(doc,
                'By guaranteeing that deposits are safe up to a prescribed limit, deposit '
                'insurance prevents bank runs where many customers withdraw funds '
                'simultaneously due to fear of bank insolvency.',
                bold_prefix='Confidence and Stability:  ')

    sub_heading(doc, 'How Deposit Insurance Protects Depositors', before=10, after=4)
    bullet_para(doc,
                'If a bank fails, the deposit insurance agency — for example, the FDIC in '
                'the United States or the Kenya Deposit Insurance Corporation (KDIC) in Kenya '
                '— pays out the insured amount directly to depositors.',
                bold_prefix='Direct Reimbursement:  ')
    bullet_para(doc,
                'Small, unsophisticated depositors are not required to constantly monitor the '
                'bank\'s financial health, as their funds are guaranteed up to a certain limit. '
                'In Kenya, the KDIC currently protects deposits up to Ksh 500,000 per '
                'depositor per institution.',
                bold_prefix='Peace of Mind:  ')

    sub_heading(doc, 'Moral Hazard: The Key Challenge', before=10, after=4)
    body_para(doc, 'While beneficial, deposit insurance creates a moral hazard problem — a '
              'situation where parties take on more risk because they do not bear the full '
              'consequences of that risk:')
    bullet_para(doc,
                'Because depositors are protected and therefore less likely to withdraw funds '
                'in response to risky bank behaviour, bank management may be tempted to engage '
                'in higher-risk lending or investment strategies to seek higher profits, '
                'knowing the "safety net" exists if they fail.',
                bold_prefix='For Banks:  ')
    bullet_para(doc,
                'Insured depositors have little incentive to monitor or discipline bank '
                'behaviour, removing a key check on excessive risk-taking.',
                bold_prefix='For Depositors:  ')

    sub_heading(doc, 'Regulatory Responses to Moral Hazard', before=10, after=4)
    body_para(doc, 'To counter the moral hazard created by deposit insurance, regulators '
              'implement the following strategies:')
    for item in [
        'Prudential Supervision — Regular audits and examinations of banks by regulators '
        'to detect and address excessive risk-taking.',
        'Risk-Based Premiums — Banks that take on more risk pay higher deposit insurance '
        'premiums, aligning the cost of insurance with actual risk levels.',
        'Coverage Limits — Restricting the guarantee to a maximum amount (e.g., Ksh 500,000 '
        'in Kenya) preserves some depositor incentive to monitor bank health for amounts '
        'above the limit.',
        'Capital Requirements — Mandatory capital ratios (such as CAR) ensure banks '
        'maintain their own buffer against losses, reducing reliance on the insurance fund.',
    ]:
        bullet_para(doc, item, bold_prefix='\u2022  ')

    doc.save(output)
    print(f'DOCX saved: {output}')
    return output


if __name__ == '__main__':
    generate_cat_docx('Calvince_BBM_415_CAT.docx')
