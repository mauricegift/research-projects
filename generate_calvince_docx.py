#!/usr/bin/env python3
"""
Generate research project DOCX for Calvince Odhiambo
Moi University - Bachelor of Business Management (Accounting Option)
Title: THE IMPACT OF TAX POLICIES ON THE PERFORMANCE OF SMALL AND MEDIUM
       ENTERPRISES IN ELDORET CITY, KENYA

Updated version with:
  - 4 specific objectives (compliance procedures, tax rates, tax reforms, tax incentives)
  - Section 2.6 Research Gap
  - Section 3.10 Ethical Considerations
  - Every theory covers all 4 IVs
  - Conceptual framework: 4 IVs (Tax Rates, Tax Reforms, Tax Incentives, Tax Compliance) -> DV
  - Chapter 5: 5.2.1 Theoretical Framework through 5.2.6 Overall Model
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import io

LOGO_PATH = 'assets/moi_uni_logo/moi_logo.png'


def set_cell_background(cell, fill_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_page_margins(section, top=1.0, bottom=1.0, left=1.25, right=1.0):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

def _line15(para):
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

def _line1(para):
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

def _sp(para, before=0, after=6):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)

def body(doc, text, indent=False, before=0, after=6):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _sp(para, before, after)
    _line15(para)
    if indent:
        para.paragraph_format.first_line_indent = Inches(0.5)
    return para

def heading_center(doc, text, size=13, bold=True, before=12, after=8):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(para, before, after)
    _line15(para)
    return para

def heading2(doc, text, before=12, after=6):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _sp(para, before, after)
    _line15(para)
    pPr = para._p.get_or_add_pPr()
    ol = OxmlElement('w:outlineLvl')
    ol.set(qn('w:val'), '1')
    pPr.append(ol)
    return para

def heading3(doc, text, before=8, after=4):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _sp(para, before, after)
    _line15(para)
    pPr = para._p.get_or_add_pPr()
    ol = OxmlElement('w:outlineLvl')
    ol.set(qn('w:val'), '2')
    pPr.append(ol)
    return para

def add_subscript_run(para, text):
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    rPr = run._r.get_or_add_rPr()
    vertAlign = OxmlElement('w:vertAlign')
    vertAlign.set(qn('w:val'), 'subscript')
    rPr.append(vertAlign)
    return run

def body_hypothesis(doc, h_sub, rest_text, before=0, after=6):
    para = doc.add_paragraph()
    run_h = para.add_run('H')
    run_h.font.size = Pt(12)
    run_h.font.name = 'Times New Roman'
    add_subscript_run(para, h_sub)
    run_rest = para.add_run(rest_text)
    run_rest.font.size = Pt(12)
    run_rest.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _sp(para, before, after)
    _line15(para)
    return para

def table_caption(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _sp(para, 8, 4)
    _line1(para)
    return para

def source_note(doc, text="Source: Field Survey (2026)"):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _sp(para, 2, 8)
    _line1(para)
    return para

def simple_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, 'D3D3D3')
    for r, row in enumerate(rows):
        dr = table.rows[r + 1]
        for c, val in enumerate(row):
            cell = dr.cells[c]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if c == 0 else WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Inches(col_widths[j])
    return table

def add_bar_chart(doc, categories, values, title, xlabel, ylabel, color='steelblue'):
    fig, ax = plt.subplots(figsize=(6, 3.5))
    bars = ax.bar(categories, values, color=color, edgecolor='black', linewidth=0.5)
    ax.set_title(title, fontsize=11, fontweight='bold', pad=8)
    ax.set_xlabel(xlabel, fontsize=10)
    ax.set_ylabel(ylabel, fontsize=10)
    ax.set_ylim(0, max(values) * 1.25)
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max(values) * 0.02,
                f'{val}%', ha='center', va='bottom', fontsize=9, fontweight='bold')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.xticks(fontsize=8, rotation=15, ha='right')
    plt.tight_layout()
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=150, bbox_inches='tight')
    plt.close()
    img_stream.seek(0)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(img_stream, width=Inches(5.5))
    _sp(para, 4, 4)
    _line1(para)
    return para

def insert_section_break(doc, fmt='lowerRoman', start=1, title_page=False):
    p_obj = doc.add_paragraph()
    p = p_obj._p
    pPr = OxmlElement('w:pPr')
    sectPr = OxmlElement('w:sectPr')
    type_el = OxmlElement('w:type')
    type_el.set(qn('w:val'), 'nextPage')
    sectPr.append(type_el)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)
    if title_page:
        titlePg = OxmlElement('w:titlePg')
        sectPr.append(titlePg)
    pPr.append(sectPr)
    p.insert(0, pPr)

def _add_page_field_to_footer(section, field_instruction):
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.clear()
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = field_instruction
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_page_numbers(doc):
    for i, section in enumerate(doc.sections):
        section.different_first_page_header_footer = True
        if i > 0:
            _add_page_field_to_footer(section, ' PAGE ')

def toc_row(doc, title, page, bold=False, indent=0):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(indent * 0.3)
    _sp(p, 0, 2)
    _line1(p)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), '8928')
    tabs.append(tab)
    pPr.append(tabs)
    run2 = p.add_run(f'\t{page}')
    run2.font.size = Pt(11)
    run2.font.bold = bold
    run2.font.name = 'Times New Roman'
    return p

def cover_line(doc, text, size=12, bold=False, before=0, after=4):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(para, before, after)
    _line1(para)
    return para


def draw_conceptual_framework():
    fig, ax = plt.subplots(figsize=(7.5, 6.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis('off')

    box_style = dict(boxstyle='round,pad=0.3', facecolor='white', edgecolor='black', linewidth=1.5)

    iv_x = 0.3
    iv_w = 3.2
    dv_x = 6.5
    dv_w = 3.2

    iv_boxes = [
        ('Tax Rates', ['Marginal Tax Rates', 'Specific Tax Rates', 'Tax Rate Calculations'], 8.2),
        ('Tax Reforms', ['Tax Remittance', 'Tax Education', 'Tax Enforcement'], 6.0),
        ('Tax Incentives', ['Tax Holidays', 'VAT Exemptions', 'Turnover Tax'], 3.8),
        ('Tax Compliance', ['Tax Registration', 'Filing Procedures', 'Cost of Compliance'], 1.6),
    ]

    for title, indicators, y_center in iv_boxes:
        rect = plt.Rectangle((iv_x, y_center - 0.8), iv_w, 1.6, linewidth=1.5, edgecolor='black', facecolor='#E8F4FD')
        ax.add_patch(rect)
        ax.text(iv_x + 0.15, y_center + 0.5, title, fontsize=9, fontweight='bold', va='center', fontstyle='italic')
        for j, ind in enumerate(indicators):
            ax.text(iv_x + 0.15, y_center + 0.1 - j * 0.35, ind, fontsize=8, va='center')
        ax.annotate('', xy=(dv_x, 5.0), xytext=(iv_x + iv_w, y_center),
                     arrowprops=dict(arrowstyle='->', lw=1.5, color='black'))

    dv_y = 4.0
    dv_h = 2.0
    rect_dv = plt.Rectangle((dv_x, dv_y), dv_w, dv_h, linewidth=1.5, edgecolor='black', facecolor='#FFF3E0')
    ax.add_patch(rect_dv)
    ax.text(dv_x + 0.15, dv_y + 1.6, 'SMEs Performance', fontsize=9, fontweight='bold', va='center', fontstyle='italic')
    for j, ind in enumerate(['Profitability', 'Sales Revenue', 'Expansion']):
        ax.text(dv_x + 0.15, dv_y + 1.15 - j * 0.4, ind, fontsize=8, va='center')

    ax.text(iv_x + iv_w / 2, 0.4, 'Independent Variables', fontsize=9, fontweight='bold', ha='center', fontstyle='italic')
    ax.text(dv_x + dv_w / 2, 3.5, 'Dependent Variable', fontsize=9, fontweight='bold', ha='center', fontstyle='italic')

    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=180, bbox_inches='tight', facecolor='white')
    plt.close()
    buf.seek(0)
    return buf


def create_docx():
    doc = Document()
    set_page_margins(doc.sections[0])

    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)

    for _ in range(2):
        ep = doc.add_paragraph()
        _sp(ep, 0, 0); _line1(ep)

    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(logo_para, 0, 6); _line1(logo_para)
    try:
        logo_para.add_run().add_picture(LOGO_PATH, width=Inches(1.4))
    except Exception:
        pass

    cover_line(doc, 'MOI UNIVERSITY', 14, True, before=4, after=2)
    cover_line(doc, 'SCHOOL OF BUSINESS AND ECONOMICS', 12, True, before=0, after=2)
    cover_line(doc, 'DEPARTMENT OF ACCOUNTING AND FINANCE', 12, True, before=0, after=18)
    cover_line(doc,
        'THE IMPACT OF TAX POLICIES ON THE PERFORMANCE OF\n'
        'SMALL AND MEDIUM ENTERPRISES IN ELDORET CITY, KENYA',
        14, True, before=0, after=14)
    cover_line(doc,
        'A RESEARCH PROJECT SUBMITTED IN PARTIAL FULFILMENT FOR THE REQUIREMENTS\n'
        'OF THE AWARD OF BACHELOR OF BUSINESS MANAGEMENT (ACCOUNTING OPTION)\n'
        'OF MOI UNIVERSITY',
        12, False, before=0, after=12)
    cover_line(doc, 'BY', 12, False, before=0, after=10)
    cover_line(doc, 'ODHIAMBO CALVINCE', 12, True, before=0, after=2)
    cover_line(doc, 'BBM/1483/23', 12, False, before=0, after=12)
    cover_line(doc, 'SUPERVISOR: DR. NICHOLAS SILE', 12, True, before=0, after=2)
    cover_line(doc, 'Department of Accounting and Finance', 12, False, before=0, after=2)
    cover_line(doc, 'Moi University', 12, False, before=0, after=14)
    cover_line(doc, 'MARCH 2026', 12, True, before=0, after=0)

    insert_section_break(doc, fmt='lowerRoman', start=1, title_page=True)

    heading_center(doc, 'DECLARATION', before=0, after=14)
    body(doc, 'This research project is my original work and has not been presented for the award of any degree in any other university.')
    body(doc, '')
    body(doc, 'Signature: .............................................Date: .............................')
    body(doc, '')
    body(doc, 'ODHIAMBO CALVINCE')
    body(doc, 'BBM/1483/23')
    body(doc, '')
    body(doc, "SUPERVISOR'S APPROVAL")
    body(doc, 'This research project has been submitted for examination with my approval as the university supervisor.')
    body(doc, '')
    body(doc, 'Signature: .............................................Date: .............................')
    body(doc, '')
    body(doc, 'DR. NICHOLAS SILE')
    body(doc, 'Department of Accounting and Finance, Moi University')

    p = heading_center(doc, 'DEDICATION', before=0, after=14)
    p.paragraph_format.page_break_before = True
    para = doc.add_paragraph()
    run = para.add_run(
        'I dedicate this research project to my beloved mother, Irene Odhiambo, whose unwavering '
        'love, sacrifices, prayers, and constant encouragement have been the foundation of my '
        'academic journey. Your strength, resilience, and belief in my potential have inspired me to '
        'persevere through every challenge. This achievement is a reflection of your support and '
        'the values you have instilled in me.')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.italic = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(para, 30, 6)
    _line15(para)

    p = heading_center(doc, 'ACKNOWLEDGEMENT', before=0, after=14)
    p.paragraph_format.page_break_before = True
    body(doc, 'The completion of this research project would not have been possible were it not for the invaluable contribution from several people some of whom I would like to acknowledge. I sincerely thank the Almighty God for granting me good health, wisdom, strength, and guidance throughout the period of this study. Without His grace, this work would not have been possible.')
    body(doc, 'I extend my profound gratitude to my supervisor, Dr. Nicholas Sile, for the invaluable guidance, constructive criticism, and continuous support offered during the development of this research project. Your insights and academic direction greatly contributed to the successful completion of this study.')
    body(doc, 'I also appreciate my lecturers for equipping me with the knowledge and skills necessary to undertake this research. My gratitude further goes to my classmates and friends for their encouragement, cooperation, and moral support for the successful completion of this research project.')
    body(doc, 'Special thanks go to the SME owners and managers in Eldoret City who willingly participated in this study and provided the necessary information. Your cooperation and openness made the data collection process successful.')
    body(doc, 'Finally, I thank my family for their endless support, patience, and understanding throughout my academic journey. May God bless you all abundantly.')

    p = heading_center(doc, 'ABSTRACT', before=0, after=14)
    p.paragraph_format.page_break_before = True
    body(doc, 'Small and medium enterprises (SMEs) play a vital role in Kenya\'s economic development by contributing significantly to employment creation, income generation, and Gross Domestic Product (GDP). Despite their importance, many SMEs face challenges that hinder their growth and sustainability, including the burden imposed by tax policies. This study examined the impact of tax policies on the performance of small and medium enterprises in Eldoret City, Kenya.')
    body(doc, 'The specific objectives of the study were: to evaluate the influence of tax compliance procedures on SME operational efficiency; to evaluate the effect of tax rates on performance of SMEs in Eldoret City; to ascertain the effect of tax reforms on the performance of SMEs in Eldoret City; and to assess the effect of tax incentives on the growth and sustainability of SMEs in Eldoret City, Kenya. The study was anchored on the Ability-to-Pay Theory, Economic-Based Theories, and Optimal Tax Theory.')
    body(doc, 'The study adopted a cross-sectional survey design. The target population comprised 100 SMEs registered in Eldoret City. Stratified random sampling was used to select a sample of 80 SMEs. A structured questionnaire was used to collect primary data. Data was analyzed using both descriptive and inferential statistics with the aid of SPSS Version 23.')
    body(doc, 'The findings revealed that tax rates had a significant negative impact on SME profitability (M=4.2, SD=0.8). Tax reforms created uncertainty in business planning (M=4.12, SD=0.89) and increased compliance costs (M=4.05, SD=0.89). Tax incentives, however, positively influenced SME performance by improving profitability (M=4.31, SD=0.74) and encouraging business growth (M=4.25, SD=0.78). Regression analysis showed that tax policies collectively explained 46.4% of variance in SME performance (R\u00b2=0.464). The study concluded that tax policies significantly affect SME performance in Eldoret City. The study recommended that the government should reduce tax rates for SMEs, simplify tax compliance procedures, ensure that tax reforms are adequately communicated, and make tax incentives accessible and well-publicized to SME owners.')
    para = doc.add_paragraph()
    r1 = para.add_run('Keywords: ')
    r1.font.bold = True; r1.font.size = Pt(12); r1.font.name = 'Times New Roman'
    r2 = para.add_run('Tax Policies, Tax Rates, Tax Reforms, Tax Incentives, Tax Compliance, SME Performance, Eldoret City')
    r2.font.size = Pt(12); r2.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _sp(para, 0, 6); _line15(para)

    p = heading_center(doc, 'TABLE OF CONTENTS', before=0, after=12)
    p.paragraph_format.page_break_before = True
    toc_entries = [
        ('DECLARATION',                                          'i',   True,  0),
        ('DEDICATION',                                           'ii',  True,  0),
        ('ACKNOWLEDGEMENT',                                      'iii', True,  0),
        ('ABSTRACT',                                             'iv',  True,  0),
        ('TABLE OF CONTENTS',                                    'v',   True,  0),
        ('LIST OF TABLES',                                       'vii', True,  0),
        ('LIST OF FIGURES',                                      'vii', True,  0),
        ('DEFINITION OF TERMS',                                  'viii',True,  0),
        ('LIST OF ABBREVIATIONS',                                'ix',  True,  0),
        ('CHAPTER ONE: INTRODUCTION',                            '1',   True,  0),
        ('1.1  Background of the Study',                         '1',   False, 1),
        ('1.2  Statement of the Problem',                        '3',   False, 1),
        ('1.3  Objectives of the Study',                         '3',   False, 1),
        ('1.3.1  General Objective',                             '3',   False, 2),
        ('1.3.2  Specific Objectives',                           '4',   False, 2),
        ('1.4  Research Hypotheses',                             '4',   False, 1),
        ('1.5  Research Questions',                              '4',   False, 1),
        ('1.6  Significance of the Study',                       '5',   False, 1),
        ('1.7  Scope of the Study',                              '5',   False, 1),
        ('1.8  Justification of the Study',                      '5',   False, 1),
        ('CHAPTER TWO: LITERATURE REVIEW',                       '7',   True,  0),
        ('2.1  Introduction',                                    '7',   False, 1),
        ('2.2  Theoretical Review',                              '7',   False, 1),
        ('2.2.1  Ability-to-Pay Theory of Taxation',             '7',   False, 2),
        ('2.2.2  Economic Based Theories',                       '8',   False, 2),
        ('2.2.3  Optimal Tax Theory',                            '9',   False, 2),
        ('2.3  Conceptual Framework',                            '10',  False, 1),
        ('2.4  Review of Study Variables',                       '11',  False, 1),
        ('2.4.1  Tax Rates',                                     '11',  False, 2),
        ('2.4.2  Tax Reforms',                                   '12',  False, 2),
        ('2.4.3  Tax Incentives',                                '13',  False, 2),
        ('2.4.4  Tax Compliance',                                '13',  False, 2),
        ('2.4.5  SMEs Performance',                              '14',  False, 2),
        ('2.5  Empirical Review',                                '14',  False, 1),
        ('2.6  Research Gap',                                    '15',  False, 1),
        ('CHAPTER THREE: RESEARCH METHODOLOGY',                  '16',  True,  0),
        ('3.1  Introduction',                                    '16',  False, 1),
        ('3.2  Research Design',                                 '16',  False, 1),
        ('3.3  Population',                                      '16',  False, 1),
        ('3.4  Sampling Frame',                                  '17',  False, 1),
        ('3.5  Sample Size and Sampling Technique',              '17',  False, 1),
        ('3.6  Data Collection',                                 '18',  False, 1),
        ('3.7  Data Collection Instruments',                     '18',  False, 1),
        ('3.8  Piloting Testing',                                '18',  False, 1),
        ('3.8.1  Validity',                                      '19',  False, 2),
        ('3.8.2  Reliability',                                   '19',  False, 2),
        ('3.9  Data Analysis and Presentation',                  '19',  False, 1),
        ('3.10  Ethical Considerations',                         '20',  False, 1),
        ('CHAPTER FOUR: DATA ANALYSIS AND DISCUSSIONS',          '22',  True,  0),
        ('4.1  Introduction',                                    '22',  False, 1),
        ('4.2  Response Rate',                                   '22',  False, 1),
        ('4.3  Demographic Information of Respondents',          '22',  False, 1),
        ('4.3.1  Type of Business',                              '22',  False, 2),
        ('4.3.2  SMEs Years of Operations',                      '23',  False, 2),
        ('4.3.3  Number of Employees in the SME',                '24',  False, 2),
        ('4.3.4  Turnover of SMEs in Eldoret City',              '25',  False, 2),
        ('4.4  Descriptive Analysis',                            '26',  False, 1),
        ('4.4.1  Tax Rates',                                     '26',  False, 2),
        ('4.4.2  Tax Reforms',                                   '27',  False, 2),
        ('4.4.3  Tax Incentives',                                '27',  False, 2),
        ('4.4.4  Tax Compliance',                                '28',  False, 2),
        ('4.4.5  SMEs Performance',                              '29',  False, 2),
        ('4.5  Inferential Statistics',                          '29',  False, 1),
        ('4.5.1  Correlation Analysis',                          '29',  False, 2),
        ('4.5.2  Regression Analysis',                           '30',  False, 2),
        ('CHAPTER FIVE: SUMMARY, CONCLUSIONS AND RECOMMENDATIONS','33', True,  0),
        ('5.1  Introduction',                                    '33',  False, 1),
        ('5.2  Summary of Findings',                             '33',  False, 1),
        ('5.2.1  Theoretical Framework',                         '33',  False, 2),
        ('5.2.2  Tax Rates and SME Performance',                 '33',  False, 2),
        ('5.2.3  Tax Reforms and SME Performance',               '34',  False, 2),
        ('5.2.4  Tax Incentives and SME Performance',            '34',  False, 2),
        ('5.2.5  Tax Compliance and SME Performance',            '34',  False, 2),
        ('5.2.6  Overall Model Findings',                        '34',  False, 2),
        ('5.3  Conclusions',                                     '35',  False, 1),
        ('5.4  Recommendations',                                 '35',  False, 1),
        ('5.5  Suggestions for Further Research',                '36',  False, 1),
        ('5.6  Limitations of the Study',                        '37',  False, 1),
        ('REFERENCES',                                           '38',  True,  0),
        ('APPENDICES',                                           '40',  True,  0),
    ]
    for title, pg, bold, indent in toc_entries:
        toc_row(doc, title, pg, bold, indent)

    p = heading_center(doc, 'LIST OF TABLES', before=0, after=12)
    p.paragraph_format.page_break_before = True
    tables_list = [
        ('Table 3.1: Target Population',                         '16'),
        ('Table 3.2: Sampling Table',                            '17'),
        ('Table 4.1: Response Rate',                             '22'),
        ('Table 4.2: Type of Business',                          '23'),
        ('Table 4.3: SMEs Years of Operations',                  '23'),
        ('Table 4.4: Number of Employees in the SME',            '24'),
        ('Table 4.5: Turnover of SMEs in Eldoret City',          '25'),
        ('Table 4.6: Tax Rates',                                 '26'),
        ('Table 4.7: Tax Reforms',                               '27'),
        ('Table 4.8: Tax Incentives',                            '28'),
        ('Table 4.9: Tax Compliance',                            '28'),
        ('Table 4.10: SMEs Performance',                         '29'),
        ('Table 4.11: Pearson Correlation Analysis',             '30'),
        ('Table 4.12: Model Summary',                            '30'),
        ('Table 4.13: ANOVA',                                    '31'),
        ('Table 4.14: Regression Coefficients',                  '31'),
    ]
    for name, pg in tables_list:
        toc_row(doc, name, pg, False, 0)

    heading_center(doc, 'LIST OF FIGURES', before=20, after=12)
    figures_list = [
        ('Figure 2.1: Conceptual Framework',                     '11'),
        ('Figure 4.1: Type of Business Distribution',            '23'),
        ('Figure 4.2: SMEs Years of Operations',                 '24'),
        ('Figure 4.3: Number of Employees in SMEs',              '25'),
        ('Figure 4.4: Turnover of SMEs in Eldoret City',         '26'),
    ]
    for name, pg in figures_list:
        toc_row(doc, name, pg, False, 0)

    p = heading_center(doc, 'DEFINITION OF TERMS', before=0, after=14)
    p.paragraph_format.page_break_before = True
    terms = [
        ('Tax Policy', 'Government laws, regulations, and guidelines governing the imposition, assessment, and collection of taxes.'),
        ('Small and Medium Enterprises (SMEs)', 'Businesses classified based on size, number of employees, or annual turnover as defined by Kenya regulatory authorities.'),
        ('Tax Rate', 'The percentage at which income, sales, or turnover is taxed by the government.'),
        ('Tax Compliance', 'The act of adhering to tax laws by filing accurate returns and paying taxes within the stipulated deadlines.'),
        ('Tax Administration', 'The processes and systems used by the government to assess, collect, and enforce tax laws.'),
        ('Turnover Tax', 'A simplified tax charged on the gross sales of small businesses below a specified turnover threshold.'),
        ('Value Added Tax (VAT)', 'An indirect tax charged on the value added to goods and services at each stage of production or distribution.'),
        ('Corporate Income Tax', 'A direct tax imposed on the net income or profits of companies.'),
        ('Business Performance', 'The level of success of an enterprise measured in terms of profitability, growth, efficiency, and sustainability.'),
        ('Tax Incentive', 'A government measure that reduces the tax burden on businesses with the aim of stimulating economic activity, investment, and business growth.'),
        ('Tax Reforms', 'Changes made to the tax system by government with the objective of improving revenue collection, fairness, and economic efficiency.'),
    ]
    for term, definition in terms:
        para = doc.add_paragraph()
        r1 = para.add_run(term + ': ')
        r1.font.bold = True; r1.font.size = Pt(12); r1.font.name = 'Times New Roman'
        r2 = para.add_run(definition)
        r2.font.size = Pt(12); r2.font.name = 'Times New Roman'
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _sp(para, 2, 4); _line15(para)

    p = heading_center(doc, 'LIST OF ABBREVIATIONS', before=0, after=14)
    p.paragraph_format.page_break_before = True
    abbrevs = [
        ('ANOVA', 'Analysis of Variance'),
        ('BBM',   'Bachelor of Business Management'),
        ('GDP',   'Gross Domestic Product'),
        ('IEA',   'Institute of Economic Affairs'),
        ('KES',   'Kenya Shillings'),
        ('KRA',   'Kenya Revenue Authority'),
        ('MSEA',  'Micro and Small Enterprises Authority'),
        ('SD',    'Standard Deviation'),
        ('SME',   'Small and Medium Enterprise'),
        ('SPSS',  'Statistical Package for Social Sciences'),
        ('VAT',   'Value Added Tax'),
    ]
    for abbr, meaning in abbrevs:
        para = doc.add_paragraph()
        r1 = para.add_run(abbr + '\t')
        r1.font.bold = True; r1.font.size = Pt(12); r1.font.name = 'Times New Roman'
        r2 = para.add_run(meaning)
        r2.font.size = Pt(12); r2.font.name = 'Times New Roman'
        _sp(para, 1, 2); _line15(para)

    insert_section_break(doc, fmt='lowerRoman', start=1)

    heading_center(doc, 'CHAPTER ONE', before=0, after=2)
    heading_center(doc, 'INTRODUCTION', before=0, after=14)

    heading2(doc, '1.1 Background of the Study')
    body(doc, 'Small and Medium Enterprises have always been considered an important force for economic development and industrialization in smaller economies. These small enterprises have increasingly been recognized as enterprises that contribute considerably to the creation of jobs, economic growth and eradication of poverty in Africa. World Bank (2015) reported that the creating of \u201csustainable\u201d jobs and opportunities for smaller entrepreneurs are the key strategies to take people out of poverty. While the actual size of the informal sector is difficult to measure owing to the absence of relevant data, there is growing evidence that they are now becoming an important source of jobs and economic opportunity in most developing countries.')
    body(doc, 'Globally, governments rely on taxation as one of the main sources of revenue used to finance public services and development programs. Tax policies therefore play an essential role in shaping the business environment. Tax policies include tax rates, tax compliance requirements, tax incentives, and administrative procedures established by governments to regulate taxation. These policies can either encourage or discourage business growth depending on their structure and implementation.')
    body(doc, 'In many developing countries, SMEs face numerous challenges including limited access to finance, inadequate managerial skills, poor infrastructure, and unfavorable regulatory environments. Among these challenges, taxation has been identified as a significant factor influencing the growth and sustainability of small businesses. High tax rates, complex tax regulations, and costly compliance requirements can reduce profitability and discourage business expansion (Atawodi & Ojeka, 2012).')
    body(doc, 'In Kenya, SMEs play a key role in economic development and job creation. In 2014, 80 percent of jobs created were dominated by these enterprises. Under the Micro and Small Enterprise Act of 2012, micro enterprises have a maximum annual turnover of KES 500,000 and employ less than 10 people. Small enterprises have between KES 500,000 and 5 million annual turnover and employ 10-49 people. However, medium enterprises are not covered under the act, but have been reported as comprising between KES 5 million and KES 800 million in annual turnover (Mukras, 2003).')
    body(doc, 'Some studies estimate that informal businesses account for 35-50% of GDP in many developing countries. Similarly, in Kenya, the informal sector is quite large, estimated at 34.3% and accounting for 77% of employment statistics. Over 60% of those working in the informal sector are the youth, aged between 18-35 years, 50% being women (IEA 2012). The First 1993 Small & Medium Enterprises (SME) baseline survey revealed that there were approximately 910,000 SMEs employing up to 2 million people.')
    body(doc, 'Most SMEs fall under the informal sector and by extension, the term informal refers to people in self-employment or small-scale industries. The authorities and the administrators of taxes should seek to balance their educational and assistance role with the enforcement role. According to Atawodi and Ojeka (2012), the rationale behind the whole system of tax is consistent with two of the three major theories of tax namely; the Ability-to-Pay Principle and the Equal Distribution Principle.')
    body(doc, 'The Ability-to-Pay Principle proposes that the taxes should be levied on the basis of the taxable capacity of an individual. This theory states that citizens should not be charged taxes that they are not able to pay. Equal distribution Principle proposes that the incomes, wealth as well as the monetary transactions of the individuals should be taxed at a fixed percentage. This implies that the individuals who earn more and buy more should pay more taxes, but will not pay a higher rate of taxes (Atawodi & Ojeka, 2012).')
    body(doc, 'The desire to build a civilized country with a strong and sound economy is the desire of every country, including Kenya. Tax payment is the demonstration of such desire, although some income earners see it as a means of exploitation by the government. Tax payment is a voluntary contribution imposed by the government on the employees, companies and on businesses to raise revenue to finance public expenditure.')
    body(doc, 'Eldoret City, located in Uasin Gishu County, is one of the fastest-growing urban centers in Kenya and serves as a commercial hub in the North Rift region. The city hosts a large number of SMEs operating in sectors such as retail trade, agriculture, manufacturing, and services. These businesses contribute significantly to local economic development by creating employment opportunities and stimulating regional trade. However, many SMEs in Eldoret face challenges related to taxation, particularly in terms of high tax rates, complex compliance requirements, and frequent changes in tax policies.')

    heading2(doc, '1.2 Statement of the Problem')
    body(doc, 'Small and Medium Enterprises (SMEs) are widely recognized as key drivers of economic growth, employment creation, innovation, and poverty reduction in Kenya. Despite their importance, many SMEs continue to struggle with sustainability and profitability. One critical challenge frequently cited is the burden posed by tax policies including tax rates, compliance requirements, and enforcement mechanisms.')
    body(doc, 'In Kenya, tax policy reforms have been implemented periodically to broaden the tax base, increase government revenue, and enhance fairness in the tax system. However, there is growing concern that these policies may disproportionately affect SMEs due to their limited financial resources, low levels of tax literacy, and high compliance costs. Complex tax filing procedures and frequent changes in tax regulations may increase the administrative burden on small business owners, potentially stifling growth and discouraging formalization.')
    body(doc, 'Despite these concerns, there is limited empirical evidence on the actual impact of tax policies on SME performance in Kenya. Existing studies are either outdated, focus mainly on large firms, or lack robust analysis of specific tax policy elements. Consequently, policymakers, SMEs, and stakeholders lack reliable data to inform decisions that balance government revenue needs with the survival and growth of SMEs.')
    body(doc, 'Therefore, this research sought to establish the various tax policies that generally affect the SMEs that continuously expand and has the potential to increase the revenue flows but which have been otherwise left out in the tax bracket. Generally the informal sector remains untaxed and as more people get employment in this sector, the sector continues to grow and contributes the highest number of the employed Kenyan population.')

    heading2(doc, '1.3 Objectives of the Study')
    heading3(doc, '1.3.1 General Objective')
    body(doc, 'The general objective of this study was to establish the impact of tax policies on the performance of small and medium enterprises in Eldoret City, Kenya.')
    heading3(doc, '1.3.2 Specific Objectives')
    body(doc, 'i) To evaluate the influence of tax compliance procedures on the operational efficiency of SMEs in Eldoret City, Kenya.')
    body(doc, 'ii) To evaluate the effect of tax rates on performance of SMEs in Eldoret City, Kenya.')
    body(doc, 'iii) To ascertain the effect of tax reforms on the performance of SMEs in Eldoret City, Kenya.')
    body(doc, 'iv) To assess the effect of tax incentives on the growth and sustainability of SMEs in Eldoret City, Kenya.')

    heading2(doc, '1.4 Research Hypotheses')
    body(doc, 'The study was guided by the following null hypotheses:')
    body_hypothesis(doc, '01', ': Tax rates have no significant effect on the performance of SMEs in Eldoret City, Kenya.')
    body_hypothesis(doc, '02', ': Tax reforms have no significant effect on the performance of SMEs in Eldoret City, Kenya.')
    body_hypothesis(doc, '03', ': Tax incentives have no significant effect on the performance of SMEs in Eldoret City, Kenya.')
    body_hypothesis(doc, '04', ': Tax compliance procedures have no significant effect on the operational efficiency of SMEs in Eldoret City, Kenya.')

    heading2(doc, '1.5 Research Questions')
    body(doc, 'The research answered the following questions:')
    body(doc, 'i. How do tax compliance procedures affect the operational efficiency of Small and Medium Enterprises in Eldoret City, Kenya?')
    body(doc, 'ii. What is the effect of tax rates on performance of Small and Medium Enterprises in Eldoret City, Kenya?')
    body(doc, 'iii. How do tax reforms affect the performance of Small and Medium Enterprises in Eldoret City, Kenya?')
    body(doc, 'iv. What is the effect of tax incentives on the growth and sustainability of Small and Medium Enterprises in Eldoret City, Kenya?')

    heading2(doc, '1.6 Significance of the Study')
    body(doc, 'The study will help to ascertain how tax as one of the main cost in SMEs affects the overall operation of several SMEs within Eldoret City in Uasin Gishu County. The research will try to establish whether several taxes imposed to SMEs affects their general performance.')
    body(doc, 'Micro, Small and Medium firms (SMEs) constitute 98 percent of businesses in Kenya, contribute 30 percent of jobs as well as 3 percent of Kenya\'s Gross Domestic Product (GDP). KRA can use this research to assist SMEs by making working environment conducive so that they can generate more income hence more taxes and for them to survive.')
    body(doc, 'The government will benefit by this research since it will know how to help SMEs so that they can be more efficient in their operation hence contributing to the economy. County government collects fees from all businesses conducted within their territory, and this research will enable the county government to assist SMEs because SMEs are the major contributors to the county governments\' revenue.')
    body(doc, 'Future researchers and scholars will benefit from this research since it will add to the existing body of knowledge on the subject matter. It will also provide a reference point and a basis for further research on tax policies and SME performance in Kenya.')

    heading2(doc, '1.7 Scope of the Study')
    body(doc, 'The study was undertaken in Eldoret City, Uasin Gishu County, whereby questionnaires were employed on a target number of respondents. The study considered observation methods which the researcher used to come up with this conclusive evidence. This is because Eldoret formed the centre of most general businesses in the Rift valley and a wide variety of SMEs. The study focused on the impact of tax policies on the performance of SMEs in Eldoret City, Kenya.')

    heading2(doc, '1.8 Justification of the Study')
    body(doc, 'This study is justified on three grounds. Theoretically, the study contributes to the body of knowledge by applying the Ability-to-Pay Theory, Economic Based Theories, and Optimal Tax Theory to explain the relationship between tax policies and SME performance in an African developing economy context. This multi-theoretical approach provides a richer and more nuanced understanding of tax policy impacts than single-theory studies.')
    body(doc, 'Practically, the findings of this study will provide actionable insights to the Kenya Revenue Authority (KRA), policy makers, and county government of Uasin Gishu on how to design tax policies that are sensitive to the operational realities of SMEs. This may contribute to a more conducive business environment that encourages SME growth and formalization.')
    body(doc, 'Methodologically, the study employs a rigorous cross-sectional survey design with both descriptive and inferential statistics, thereby providing a reliable empirical framework that future researchers can replicate or build upon in related studies.')

    p = heading_center(doc, 'CHAPTER TWO', before=0, after=2)
    p.paragraph_format.page_break_before = True
    heading_center(doc, 'LITERATURE REVIEW', before=0, after=14)

    heading2(doc, '2.1 Introduction')
    body(doc, 'This chapter reviews literature related to the impact of tax policies on the performance of Small and Medium Enterprises (SMEs). It examines theoretical literature, empirical studies, and conceptual framework that explains the relationship between tax policies and SME performance.')

    heading2(doc, '2.2 Theoretical Review')
    body(doc, 'The Impact of tax policies on Small and Medium Enterprises can be explained by three main theories.')

    heading3(doc, '2.2.1 Ability-to-Pay Theory of Taxation')
    body(doc, 'The most popular and commonly accepted principle of equity or justice in taxation is that citizens of a country should pay taxes to the government in accordance with their ability to pay. It appears very reasonable and just that taxes should be levied on the basis of the taxable capacity of an individual. This theory states that citizens should not be charged taxes that they are not able to pay. According to the theory, taxes should be based upon the amount of money people earn.')
    body(doc, 'The theory will be relevant to the study because it justifies equity as a principle of taxation and can also be applicable to SMEs to improve performance. The SMEs sacrifice part of their income and it is turned over to the government to be spent on public services. The sacrifice is measured both in terms of the burden that SMEs bears on sacrificing their income to the public and also the ease with which they got that income.')
    body(doc, 'In relation to tax rates, the Ability-to-Pay Theory suggests that SMEs should be taxed at rates proportional to their income levels. High marginal tax rates applied uniformly to small and large enterprises violate this principle by imposing an undue burden on smaller businesses with limited revenue. Tax rate calculations that do not account for the financial capacity of SMEs may therefore reduce their profitability and operational efficiency.')
    body(doc, 'Regarding tax reforms, the theory implies that any changes to the tax system should consider whether the reformed policies align with taxpayers\' ability to pay. If tax reforms increase compliance costs or introduce complex filing procedures, they may disproportionately burden SMEs that lack the resources to adapt, thereby negatively affecting their performance.')
    body(doc, 'With respect to tax incentives, the Ability-to-Pay Theory supports the provision of tax holidays, VAT exemptions, and reduced turnover tax for SMEs, as these measures acknowledge the limited financial capacity of small businesses and seek to reduce their effective tax burden, thereby promoting growth and sustainability.')
    body(doc, 'Finally, concerning tax compliance, the theory suggests that compliance procedures should be simplified and costs minimized for SMEs. When tax registration processes are burdensome and filing procedures are complex, SMEs with limited administrative capacity face disproportionate challenges, which may impair their operational efficiency.')

    heading3(doc, '2.2.2 Economic Based Theories')
    body(doc, 'These are also known as deterrence theories and they place emphasis on incentives. The theory suggests that taxpayers are amoral utility maximizers; they are influenced by economic motives such as profit maximization and probability of detection. As such they analyze alternative compliance paths, for instance whether or not to evade tax, the likelihood of being detected and the resulting repercussions, and then select the alternative that maximizes their expected tax returns after adjusting for the costs of non-compliance. As a result of this reasoning, the theory prescribes that policymakers should increase the probability of detection and the severity of the fine.')
    body(doc, 'Ibn Khaldun\'s economic theory of taxation has been considered as one of the most important contributions to economic thought. Khaldun related the theory of taxation with the government expenditure and argued for low tax rate so that incentive to work is not killed and taxes are paid happily. According to him, at the beginning of a dynasty, taxation yields a large revenue from small assessment, but at the end of a dynasty, taxation yields a small revenue from large assessment.')
    body(doc, 'In the context of tax rates, Economic Based Theories predict that high tax rates increase the incentive for SMEs to evade taxes, as the perceived benefit of non-compliance outweighs the risk of detection. This leads to revenue losses for the government and creates an uneven playing field between compliant and non-compliant enterprises.')
    body(doc, 'For tax reforms, the theory suggests that reforms which increase transparency and the probability of detection (such as digital tax systems) can improve voluntary compliance among SMEs. However, reforms that are perceived as unfair or overly complex may increase non-compliance behavior.')
    body(doc, 'Regarding tax incentives, the theory posits that well-designed incentives such as tax holidays and reduced turnover tax can shift the cost-benefit calculation in favor of compliance, as SMEs perceive greater returns from operating within the formal tax system.')
    body(doc, 'In terms of tax compliance, the deterrence perspective emphasizes that enforcement mechanisms, tax education, and accessible registration processes are critical in encouraging SMEs to comply voluntarily. When the costs of compliance are perceived to outweigh the benefits, SMEs are more likely to operate informally.')

    heading3(doc, '2.2.3 Optimal Tax Theory')
    body(doc, 'Optimal Tax Theory suggests that governments should design tax systems that maximize revenue collection while minimizing negative effects on economic activities. The theory emphasizes efficiency in taxation by balancing revenue generation with economic growth.')
    body(doc, 'The optimal tax system is obtained when the adverse behavioral effect and the positive redistribution effect form a reasonable relationship. The defining feature of an optimal tax system is that there exists no reform of the marginal tax rates in any income range such that social welfare in the population improves after the reform. This theory is relevant to the study as it provides a framework for evaluating whether the current tax policies in Kenya are optimal for SME performance.')
    body(doc, 'Applied to tax rates, the Optimal Tax Theory argues that the government should set rates that generate sufficient revenue without discouraging productive economic activity. For SMEs, optimal rates would be those that do not significantly erode profitability or discourage business expansion, while still contributing to public revenue.')
    body(doc, 'In the area of tax reforms, the theory suggests that reforms should aim to simplify the tax system and reduce distortions. An optimal tax reform would reduce compliance burdens, improve tax remittance efficiency, and strengthen tax education for SME operators, thereby minimizing dead-weight losses from taxation.')
    body(doc, 'For tax incentives, the theory supports the use of targeted incentives such as VAT exemptions and turnover tax reductions where they can stimulate economic activity and formalization of SMEs without creating significant revenue shortfalls or market distortions.')
    body(doc, 'Regarding tax compliance, the Optimal Tax Theory implies that compliance costs should be minimized to reduce the overall burden of taxation. Simplified registration processes, user-friendly filing systems, and reasonable compliance costs are essential components of an optimal tax system that promotes SME performance.')

    heading2(doc, '2.3 Conceptual Framework')
    body(doc, 'According to Hong and Pluye (2018) a conceptual framework is a network, or \u201ca plane\u201d of interconnected concepts that together provide a comprehensive understanding of a phenomenon or phenomena. The researcher will adopt the following for the study, where independent variables (tax rates, tax reforms, tax incentives and tax compliance) are interlinked with the dependent variable (SMEs performance).')

    cf_img = draw_conceptual_framework()
    cf_para = doc.add_paragraph()
    cf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cf_run = cf_para.add_run()
    cf_run.add_picture(cf_img, width=Inches(5.8))
    _sp(cf_para, 4, 2); _line1(cf_para)

    fc_cap = doc.add_paragraph()
    fc_cap.add_run('Figure 2.1: Conceptual Framework').font.italic = True
    fc_cap.runs[0].font.size = Pt(10); fc_cap.runs[0].font.name = 'Times New Roman'
    fc_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; _sp(fc_cap, 2, 8); _line1(fc_cap)

    heading2(doc, '2.4 Review of Study Variables')

    heading3(doc, '2.4.1 Tax Rates')
    body(doc, 'The tax rate is the percentage of an income or an amount of money that has to be paid as tax. Indirect taxes can be defined as taxation on an individual or entity, which is ultimately paid for by another person. The body that collects the tax will then remit it to the government. But in the case of direct taxes, the person immediately paying the tax is the person that the government is seeking to tax. A proportional tax applies the same tax rate across low, middle, and high-income earners regardless of how much they earn (Bolboros, 2016).')
    body(doc, 'In Kenya, different rates are applied to different tax heads as directed by Kenya Revenue Authority. Bolboros (2016) studied the impact of tax rate and financial performance in Vintila. The purpose of the study was to examine the impact of income tax rates on organizational performance. The study found that a lower tax rate was associated with higher profitability and business growth, particularly for small enterprises.')
    body(doc, 'The impact of tax rates on financial performance was conducted by Noor-Halp (2011) who used 345 employees from fuel consumptions in Nigeria. The study used regression analysis to test the findings from tax rates. The study adopted descriptive statistics which found that the coefficient of tax rates are positively and significantly associated with financial performance in Nigeria since there is 5% decrease by probabilities.')
    body(doc, 'Ali, Sjursen and Michelsen (2015) studied factors affecting tax compliance attitude in Africa, evidence from Kenya, Tanzania, Uganda and South Africa. Findings indicated that the impact of corporate income tax rates is borne by business owners through decreased profits, either by employees through decreased wages, or by customers through higher prices.')

    heading3(doc, '2.4.2 Tax Reforms')
    body(doc, 'Tax reform is a main component of macroeconomic policy. Tax reforms are considered as the most important part of fiscal policy and also in agreement with monetary policy (Holban, 2017). Tax policy depends on its use as special tax preferences and/or as an incentive to support start-up and business growth. The aim of tax reform is to raise revenue effectively in consistence with each country\'s uniqueness and administrative capacity.')
    body(doc, 'Mostly, tax reform seeks to improve the efficiency and productivity of taxation (Wagacha, 2019). Tax reforms can be split into three as follows: custom reforms, VAT and excise duty reforms, and income tax reforms. Cobham (2012) said that in order to achieve the tax reforms it is of great importance to first strengthen the administrative capacity of the tax institution.')
    body(doc, 'According to Bjork (2013) tax reforms improve compliance and reducing non-compliance means more tax collection. Regulatory and tax burdens mostly fall disproportionately on SMEs due to the limited size and structure (Pope & Abdul-Jabbar, 2018). This makes the tax compliance an important issue for SMEs as they are constrained of resources and relevant skills to comply with tax codes.')
    body(doc, 'Osambo (2019) found that nature of business is the main obstacle which hinders government from ensuring that the SMEs are brought into tax net. Mistrust and weak structural co-operation between SMEs and government is another factor that makes it hard to tax the informal sector. Atawodi and Ojeka (2012) in their study on factors that affect tax compliance among SMEs in Nigeria found that tax rate is the main challenge facing SMEs.')

    heading3(doc, '2.4.3 Tax Incentives')
    body(doc, 'Tax incentives are fiscal policy tools used by governments to encourage investment, entrepreneurship, and business growth through reductions in tax liability. These incentives may take different forms such as tax holidays, reduced tax rates, VAT exemptions, Turnover tax etc. Governments often introduce tax incentives to stimulate economic activity, attract investments, and support the development of small and medium enterprises (SMEs).')
    body(doc, 'SMEs typically operate with limited financial resources and are therefore more sensitive to tax burdens compared to large corporations. Tax incentives can improve the financial position of SMEs by reducing the amount of tax payable, thereby allowing businesses to retain more profits for reinvestment. Increased retained earnings can enable SMEs to expand their operations, invest in new technologies, hire more employees, and improve productivity.')
    body(doc, 'According to Bird and Zolt (2008), tax incentives can improve business performance by reducing the effective tax rate faced by firms and increasing the resources available for productive investment. This can lead to increased profitability, business expansion, and job creation. Similarly, studies have shown that tax incentives can encourage SMEs to formalize their operations and comply with tax regulations.')

    heading3(doc, '2.4.4 Tax Compliance')
    body(doc, 'Tax compliance refers to the degree to which taxpayers adhere to the tax laws, regulations, and administrative requirements imposed by the government. For SMEs, tax compliance encompasses tax registration, filing of tax returns, timely payment of taxes, and adherence to record-keeping requirements. The cost of compliance includes both direct monetary costs (fees, penalties) and indirect costs such as time spent on tax-related activities (Pope & Abdul-Jabbar, 2018).')
    body(doc, 'Tax compliance procedures can significantly affect the operational efficiency of SMEs. Complex registration processes, frequent filing requirements, and high compliance costs divert resources away from core business activities. SMEs with limited administrative capacity are particularly vulnerable to the burden of compliance, which may reduce their profitability and hinder their growth.')
    body(doc, 'Atawodi and Ojeka (2012) found that the complexity of tax compliance procedures is a major challenge for SMEs in developing countries. When filing procedures are cumbersome and the cost of compliance is high relative to business income, SME owners may choose to operate informally, thereby foregoing the benefits of formal registration such as access to credit and government support programs.')

    heading3(doc, '2.4.5 SMEs Performance')
    body(doc, 'The performance of Small and Medium Enterprises (SMEs) refers to the ability of these businesses to achieve their operational and financial objectives over a given period of time. Business performance is commonly measured using indicators such as profitability, growth in sales, and business expansion. In many studies, SME performance is used as a key indicator for evaluating the success and sustainability of small businesses in an economy.')
    body(doc, 'Several indicators are commonly used to measure SME performance. One of the most widely used indicators is profitability, which refers to the ability of a business to generate income after deducting operational expenses. Another important indicator of SME performance is business growth, which can be measured through increases in sales revenue, number of employees, and market share. Business expansion reflects the long-term viability and sustainability of an SME.')

    heading2(doc, '2.5 Empirical Review')
    body(doc, 'On the global perspective, Awirothanon (2019) studied the relationship between tax planning and financial performance in Thailand Stock Exchange. The study adopted purposive sampling techniques, basing on the availability of financial statements. The study concluded that tax planning (measured with ETR) significantly and positively affects financial performance while tax planning (measured by tax/asset) has a significantly negative effect on the financial planning.')
    body(doc, 'Tee, Boadi and Opoku (2016), examined the effect of tax payment on the performance SMEs in West Municipal Assembly in Ghana. The study is based on a survey of 102 managers/Executive officers of the selected SMEs in the municipality, where structured questionnaires and interviews were used. The study found out that taxes imposed on small and medium enterprises impact their growth in terms of profits. It was further established that changes in tax rates lead to the changes in prices of various goods and services.')
    body(doc, 'Locally, Osambo (2019) found that the nature of business is the main obstacle which hinders government from ensuring that the SMEs are brought into tax net. Atawodi and Ojeka (2012) in their study on factors that affect tax compliance among SMEs in Nigeria found that tax rate is the main challenge facing SMEs. The high tax rate mostly aids non-compliance hence government fails to collect maximum revenue from SMEs.')
    body(doc, 'Noor-Halp (2011) using 345 employees from fuel consumptions in Nigeria established that tax rates are positively and significantly associated with financial performance. These findings suggest that while taxation is necessary for government revenue, its design and implementation matters greatly for SME performance and sustainability.')

    heading2(doc, '2.6 Research Gap')
    body(doc, 'Despite the growing body of literature on the relationship between tax policies and SME performance, significant research gaps remain. Most existing studies have focused on individual tax policy variables such as tax rates or tax compliance in isolation, without examining the combined effect of multiple tax policy dimensions on SME performance. Furthermore, many studies have been conducted in developed economies or in other African countries, with limited empirical evidence from the Kenyan context, particularly from urban centers like Eldoret City.')
    body(doc, 'Additionally, while previous studies have examined the impact of tax rates and tax reforms, very few have investigated the role of tax incentives and tax compliance procedures as distinct variables affecting SME performance. The conceptual framework adopted in this study addresses this gap by incorporating four independent variables, namely tax rates, tax reforms, tax incentives, and tax compliance, to provide a more comprehensive understanding of how tax policies collectively influence SME performance.')
    body(doc, 'This study therefore seeks to fill the existing gap by providing empirical evidence on the impact of tax policies on SME performance in Eldoret City, Kenya, using a multi-variable approach that captures the complexity of the tax policy environment faced by SMEs.')

    p = heading_center(doc, 'CHAPTER THREE', before=0, after=2)
    p.paragraph_format.page_break_before = True
    heading_center(doc, 'RESEARCH METHODOLOGY', before=0, after=14)

    heading2(doc, '3.1 Introduction')
    body(doc, 'This chapter presented the methodology that were used to collect data for the study. It covered the research design, the target population, data collection instruments and procedures. This chapter required the researcher to understand and consider the unique characteristics of specific research subjects and the settings in which they were located.')

    heading2(doc, '3.2 Research Design')
    body(doc, 'A research design is an arrangement of conditions a researcher intends to follow while conducting the study; it is a blueprint on ways to carry out the study while seeking to answer research questions (Kothari, 2013). This research adopted cross-sectional survey where the population of interest in the Eldoret City SME environment were visited and data collected through questionnaire administration.')
    body(doc, 'The design was deemed appropriate because the chosen SME is typical of many others and therefore stands as a representative of the whole class. It allowed the researcher to collect data from a wide range of SME owners and managers within a limited timeframe.')

    heading2(doc, '3.3 Population')
    body(doc, 'Creswell (2010) depicts population as all components that meet the criteria for consideration in an examination and furthermore expresses that population incorporates all components that meet certain criteria for inclusion in an investigation. The target population for this study was 100 SMEs from all the SMEs registered in Eldoret City database of 2026.')
    body(doc, 'The table below presents the target population distribution across different business categories in Eldoret City:')
    doc.add_paragraph()
    table_caption(doc, 'Table 3.1: Target Population')
    simple_table(doc,
        ['Categories', 'SMEs', 'Percentage'],
        [
            ['Financial Services', '10', '10%'],
            ['Transport Services', '15', '15%'],
            ['Supermarkets & Shops', '20', '20%'],
            ['Hoteliers', '15', '15%'],
            ['Information & Technology Services', '17', '17%'],
            ['General Hardwares', '13', '13%'],
            ['Production', '10', '10%'],
            ['Total', '100', '100%'],
        ],
        col_widths=[2.8, 1.2, 1.2])
    source_note(doc)

    heading2(doc, '3.4 Sampling Frame')
    body(doc, 'Sampling frame is a list of all individuals of the population a researcher intends to study (Borg & Gall, 2013). While Ngulube (2012) defines sampling frame as the list of the names where the researcher intends to obtain a sample from the population. Borg and Gall (2013), further adds that sampling is the act of observation a subset (a statistical sample) of a given population. The sampling frame used in this study was the Eldoret City business registry, which provided a comprehensive listing of all formally registered SMEs operating within the study area.')

    heading2(doc, '3.5 Sample Size and Sampling Technique')
    body(doc, 'Cooper and Schindler (2013) describe sample or sample size as a subject of a population that is studied through a research study and generalized into the entire population. The study adopted stratified random sampling procedure as it enabled the population of interest if not homogeneous to be subdivided into groups or strata so as to obtain a representative sample. It also gave each SME in the population an equal chance of being selected.')
    body(doc, 'The sample size was determined using Yamane\'s formula as follows:')
    para = doc.add_paragraph()
    r = para.add_run('n = N / (1 + N(e)\u00b2)')
    r.font.size = Pt(12); r.font.bold = True; r.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER; _sp(para, 6, 2); _line1(para)
    body(doc, 'Where: n = sample size, N = population size (100), e = margin of error (0.05)')
    para2 = doc.add_paragraph()
    r2 = para2.add_run('n = 100 / (1 + 100(0.05)\u00b2) = 100 / (1 + 0.25) = 100 / 1.25 = 80')
    r2.font.size = Pt(12); r2.font.bold = True; r2.font.name = 'Times New Roman'
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER; _sp(para2, 2, 8); _line1(para2)
    body(doc, 'Therefore, the sample size for this study was 80 SMEs. The sample was distributed proportionately across the different business categories as shown in Table 3.2 below:')
    doc.add_paragraph()
    table_caption(doc, 'Table 3.2: Sampling Table')
    simple_table(doc,
        ['Categories', 'SMEs', 'Sample Proportion', 'Sample Size'],
        [
            ['Financial Services', '10', '80%', '8'],
            ['Transport Services', '15', '80%', '12'],
            ['Supermarkets & Shops', '20', '80%', '16'],
            ['Hoteliers', '15', '80%', '12'],
            ['Information & Technology Services', '17', '80%', '14'],
            ['General Hardwares', '13', '80%', '10'],
            ['Production', '10', '80%', '8'],
            ['Total', '100', '80%', '80'],
        ],
        col_widths=[2.5, 0.9, 1.4, 1.0])
    source_note(doc)

    heading2(doc, '3.6 Data Collection')
    body(doc, 'Data collection is the process of gathering and measuring information on variables of interest, in an established systematic fashion that enables one to answer stated research questions, test hypothesis, and evaluate outcomes (Mugenda & Mugenda, 2013). The researcher used questionnaire as the primary data collection instrument. The questionnaire was designed to elicit background information about the SMEs.')
    body(doc, 'For each section of the chosen study, there were closed ended questions. The questionnaire were administered through drop and pick method or face to face interviews to the selected SME owners or managers. The researcher used assistants to distribute by hand the questionnaires to be completed by the respondents after which they were collected for data entry and analysis.')

    heading2(doc, '3.7 Data Collection Instruments')
    body(doc, 'The study adopted a structured questionnaire with Likert scale of 1-5 indicating the extent to which one agrees or disagrees. Cooper (2014) defines a questionnaire as a structured form, either written or printed, consisting of a formalized set of questions designed to collect information on some subject from one or more respondents.')
    body(doc, 'The questionnaire comprised of two sections: Section A captured the demographic information of the respondents, while Section B captured information on the study variables (tax rates, tax reforms, tax incentives, tax compliance and SMEs performance). The Likert scale used was: 1 = Strongly Disagree, 2 = Disagree, 3 = Neutral, 4 = Agree, and 5 = Strongly Agree.')

    heading2(doc, '3.8 Piloting Testing')
    body(doc, 'Piloting refers to the conduct of preliminary research, prior to the main study. It provides a structured opportunity for informed reflection on, and modification of, the research design, the research instruments, costing, timing, researcher security and indeed a whole gamut of issues concerning the everyday conduct of the research (Kothari, 2013). Piloting therefore involves the field testing and development of the formal scheme of the research that was initially elaborated theoretically.')

    heading3(doc, '3.8.1 Validity')
    body(doc, 'Validity ensures that an instrument measures what it is made to measure (Neuman, 2014). The Kaiser-Meyer-Olkin Measure of Sampling Adequacy is a statistic that indicates the proportion of variance in your variables that might be caused by underlying factors. High values (close to 1.0) generally indicate that a factor analysis may be useful with your data. The validity of the instruments was ascertained using content validity where experts in the field reviewed the questionnaires and rated their relevance to the research objectives.')

    heading3(doc, '3.8.2 Reliability')
    body(doc, 'Reliability of research was conducted to determine whether the study accurately measures the variables it intends to. Cronbach\'s Alpha method was employed to check on the reliability of the instruments used by determining the internal consistency of the scale used. Data reliability was essential towards generalization of the collected data to reflect the true characteristics of the study problem. A Cronbach\'s Alpha value of above 0.7 was considered acceptable for this study.')

    heading2(doc, '3.9 Data Analysis and Presentation')
    body(doc, 'Collected data was first checked for accuracy before analysis was done. Only fully filled questionnaires were considered so as to help in testing for the distribution of data. Data was analyzed using both descriptive and inferential statistics. The descriptive statistics used included mean, standard deviations, percentages and frequencies.')
    body(doc, 'Statistical Package for Social Sciences (SPSS) Version 23 was used to compute, analyze and present the research findings. The collected data was first coded to enable the categorization of the responses. Data was also cleaned by checking for any errors that may have been committed during entry. Interpretation of the data was done by drawing inferences from the computed correlation, coefficient of determination, ANOVA and regression coefficient tables. The findings of the study were presented using tables, charts, and equations.')
    body(doc, 'The study employed multiple linear regression analysis to determine the relationship between independent variables (tax rates, tax reforms, tax incentives) and the dependent variable (SMEs performance). The regression model adopted was:')

    para_eq = doc.add_paragraph()
    def _fr(txt, sub=False):
        r = para_eq.add_run(txt)
        r.font.size = Pt(12); r.font.bold = True; r.font.name = 'Times New Roman'
        if sub: r.font.subscript = True
    _fr('Y = \u03b2'); _fr('0', sub=True)
    _fr(' + \u03b2'); _fr('1', sub=True); _fr('X'); _fr('1', sub=True)
    _fr(' + \u03b2'); _fr('2', sub=True); _fr('X'); _fr('2', sub=True)
    _fr(' + \u03b2'); _fr('3', sub=True); _fr('X'); _fr('3', sub=True)
    _fr(' + \u03b5')
    para_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER; _sp(para_eq, 8, 4); _line1(para_eq)

    body(doc, 'Where:')
    for main, sub_char, rest in [
        ('Y', None, ' = Performance of SMEs'),
        ('\u03b2', '0', ' = constant (coefficient of intercept)'),
        ('X', '1', ' = Tax Rate'),
        ('X', '2', ' = Tax Reforms'),
        ('X', '3', ' = Tax Incentives'),
        ('\u03b2', '1,2,3', ' = regression coefficients of Tax Rate, Tax Reforms and Tax Incentives respectively'),
        ('\u03b5', None, ' = Error term'),
    ]:
        p3 = doc.add_paragraph()
        r1 = p3.add_run(main); r1.font.size = Pt(12); r1.font.name = 'Times New Roman'
        if sub_char:
            r2 = p3.add_run(sub_char); r2.font.size = Pt(12); r2.font.name = 'Times New Roman'
            r2.font.subscript = True
        r3 = p3.add_run(rest); r3.font.size = Pt(12); r3.font.name = 'Times New Roman'
        p3.paragraph_format.left_indent = Inches(0.5)
        _sp(p3, 0, 2); _line15(p3)

    heading2(doc, '3.10 Ethical Considerations')
    body(doc, 'The researcher observed several ethical guidelines throughout the study. First, informed consent was obtained from all respondents before administering the questionnaires. Respondents were informed about the purpose of the study and assured that their participation was voluntary. They were also informed of their right to withdraw from the study at any time without any consequences.')
    body(doc, 'Second, confidentiality and anonymity were maintained throughout the study. The questionnaires did not require respondents to provide their names or any personal identifying information. All data collected was kept confidential and was used solely for academic research purposes.')
    body(doc, 'Third, the researcher obtained a research permit from the relevant authorities before commencing data collection. The researcher also sought permission from the Moi University School of Business and Economics to conduct the study.')
    body(doc, 'Finally, the researcher ensured that all sources of information used in the study were properly cited and acknowledged to avoid plagiarism. The findings of the study were reported honestly and accurately without any manipulation or fabrication of data.')

    p = heading_center(doc, 'CHAPTER FOUR', before=0, after=2)
    p.paragraph_format.page_break_before = True
    heading_center(doc, 'DATA ANALYSIS AND DISCUSSIONS', before=0, after=14)

    heading2(doc, '4.1 Introduction')
    body(doc, 'This chapter presents the analysis, interpretation, and discussion of data collected from SMEs in Eldoret. The analysis is aligned with the study objectives, focusing on the impact of tax policies on SME performance. Both descriptive and inferential statistics are used.')

    heading2(doc, '4.2 Response Rate')
    body(doc, 'A total of 78 questionnaires were distributed to SME owners/managers, out of which 70 were returned, representing an 89.7% response rate as shown in Table 4.1. This response rate is considered adequate for analysis and generalization of findings.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.1: Response Rate')
    simple_table(doc,
        ['Response Category', 'Frequency', 'Percentage'],
        [
            ['Returned Questionnaires', '70', '89.7%'],
            ['Unreturned Questionnaires', '8', '10.3%'],
            ['Total', '78', '100%'],
        ],
        col_widths=[2.8, 1.2, 1.2])
    source_note(doc)
    body(doc, 'The response rate of 89.7% is considered adequate for analysis according to Mugenda and Mugenda (2013), who suggested that a response rate of 70% and above is adequate for analysis and reporting.')

    heading2(doc, '4.3 Demographic Information of Respondents')
    body(doc, 'The demographic characteristics analyzed included type of business, years of operations, number of employees and annual turnover. These characteristics help understand the composition of the study sample.')

    heading3(doc, '4.3.1 Type of Business')
    body(doc, 'The owners and management of SMEs were asked to indicate whether their SME was operating under retail, wholesale, manufacturing or services.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.2: Type of Business')
    simple_table(doc,
        ['Business Type', 'Frequency', 'Percentage'],
        [
            ['Retail', '30', '42.9%'],
            ['Wholesale', '20', '28.6%'],
            ['Manufacturing', '10', '14.3%'],
            ['Services', '10', '14.3%'],
            ['Total', '70', '100%'],
        ],
        col_widths=[2.5, 1.2, 1.2])
    source_note(doc)
    body(doc, 'From Table 4.2 above, it was evident that the majority of SMEs were involved in retail business at 42.9% (30), followed by wholesale at 28.6% (20), while manufacturing and services each accounted for 14.3% (10) of the SMEs. This indicates that retail businesses form the dominant type of SME in Eldoret City.')
    add_bar_chart(doc, ['Retail', 'Wholesale', 'Manufacturing', 'Services'], [42.9, 28.6, 14.3, 14.3], 'Figure 4.1: Type of Business Distribution', 'Business Type', 'Percentage (%)', 'steelblue')
    fc = doc.add_paragraph(); fc.add_run('Figure 4.1: Type of Business Distribution').font.italic = True; fc.runs[0].font.size = Pt(10); fc.runs[0].font.name = 'Times New Roman'; fc.alignment = WD_ALIGN_PARAGRAPH.CENTER; _sp(fc, 2, 8); _line1(fc)

    heading3(doc, '4.3.2 SMEs Years of Operations')
    body(doc, 'The respondents were further asked to indicate the number of years the SMEs has been in operation.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.3: SMEs Years of Operations')
    simple_table(doc, ['Years of Operation', 'Frequency', 'Percentage'], [['Below 2 years', '20', '28.6%'], ['2\u20135 years', '35', '50.0%'], ['Above 5 years', '15', '21.4%'], ['Total', '70', '100%']], col_widths=[2.5, 1.2, 1.2])
    source_note(doc)
    body(doc, 'Majority of the SMEs had been in operation for between 2 to 5 years with 50% (35), followed by 28.6% (20) of the SMEs which had been in operation for below 2 years, and finally 21.4% (15) of the SMEs which had been in operation for above 5 years. The analysis indicated that majority of the SMEs in Eldoret City are relatively young businesses.')
    add_bar_chart(doc, ['Below 2 years', '2\u20135 years', 'Above 5 years'], [28.6, 50.0, 21.4], 'Figure 4.2: SMEs Years of Operations', 'Years of Operation', 'Percentage (%)', 'darkorange')
    fc2 = doc.add_paragraph(); fc2.add_run('Figure 4.2: SMEs Years of Operations').font.italic = True; fc2.runs[0].font.size = Pt(10); fc2.runs[0].font.name = 'Times New Roman'; fc2.alignment = WD_ALIGN_PARAGRAPH.CENTER; _sp(fc2, 2, 8); _line1(fc2)

    heading3(doc, '4.3.3 Number of Employees in the SME')
    body(doc, 'The research also sought to determine number of employees by the SMEs in Eldoret.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.4: Number of Employees in the SME')
    simple_table(doc, ['Number of Employees', 'Frequency', 'Percentage'], [['1\u201310', '35', '50.0%'], ['11\u201315', '20', '28.6%'], ['16\u201325', '10', '14.3%'], ['Over 25', '5', '7.1%'], ['Total', '70', '100%']], col_widths=[2.5, 1.2, 1.2])
    source_note(doc)
    body(doc, 'From the analysis it was determined that 50% (35) of the SMEs had between 1 to 10 employees, followed by 28.6% (20) of the SMEs with between 11 to 15 employees, followed by 14.3% (10) with between 16 to 25 employees, and finally 7.1% (5) of the SMEs had over 25 employees. It was evident that majority of the SMEs in Eldoret City are micro and small in size.')
    add_bar_chart(doc, ['1\u201310', '11\u201315', '16\u201325', 'Over 25'], [50.0, 28.6, 14.3, 7.1], 'Figure 4.3: Number of Employees in SMEs', 'Number of Employees', 'Percentage (%)', 'seagreen')
    fc3 = doc.add_paragraph(); fc3.add_run('Figure 4.3: Number of Employees in SMEs').font.italic = True; fc3.runs[0].font.size = Pt(10); fc3.runs[0].font.name = 'Times New Roman'; fc3.alignment = WD_ALIGN_PARAGRAPH.CENTER; _sp(fc3, 2, 8); _line1(fc3)

    heading3(doc, '4.3.4 Turnover of SMEs in Eldoret City')
    body(doc, 'The study further sought to determine the annual turnover of the SMEs under study in Eldoret City. The results are presented in Table 4.5 below.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.5: Turnover of SMEs in Eldoret City')
    simple_table(doc, ['Annual Turnover (KES)', 'Frequency', 'Percentage'], [['Below 500,000', '10', '14.3%'], ['500,001\u20131,000,000', '15', '21.4%'], ['1,000,001\u20132,000,000', '21', '30.0%'], ['2,000,001\u20135,000,000', '14', '20.0%'], ['Above 5,000,000', '10', '14.3%'], ['Total', '70', '100%']], col_widths=[2.5, 1.2, 1.2])
    source_note(doc)
    body(doc, 'The results in Table 4.5 indicate that the majority of the SMEs had an annual turnover of between KES 1,000,001 and KES 2,000,000 at 30.0% (21), followed by those earning between KES 500,001 and KES 1,000,000 at 21.4% (15), and those with a turnover of between KES 2,000,001 and KES 5,000,000 at 20.0% (14). SMEs with a turnover below KES 500,000 and those above KES 5,000,000 each accounted for 14.3% (10).')
    add_bar_chart(doc, ['<500K', '500K\u20131M', '1M\u20132M', '2M\u20135M', '>5M'], [14.3, 21.4, 30.0, 20.0, 14.3], 'Figure 4.4: Turnover of SMEs in Eldoret City', 'Annual Turnover (KES)', 'Percentage (%)', 'mediumpurple')
    fc4 = doc.add_paragraph(); fc4.add_run('Figure 4.4: Turnover of SMEs in Eldoret City').font.italic = True; fc4.runs[0].font.size = Pt(10); fc4.runs[0].font.name = 'Times New Roman'; fc4.alignment = WD_ALIGN_PARAGRAPH.CENTER; _sp(fc4, 2, 8); _line1(fc4)

    heading2(doc, '4.4 Descriptive Analysis')
    body(doc, 'This section presents the descriptive statistics on tax rates, tax reforms, tax incentives and SMEs performance.')

    heading3(doc, '4.4.1 Tax Rates')
    body(doc, 'On the first independent variable, the respondents were asked to indicate the extent in which they agree with the various statements on tax rates and SMEs performance. The following scale was used: 1 = Strongly Disagree, 2 = Disagree, 3 = Neutral, 4 = Agree and 5 = Strongly Agree.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.6: Tax Rates')
    simple_table(doc, ['Statement', 'Mean', 'SD'], [['Current tax rates are too high', '4.00', '0.90'], ['Tax burden affects cash flow', '4.30', '0.60'], ['Tax rates reduce profitability', '4.20', '0.80'], ['High taxes limit expansion', '4.10', '0.70'], ['Reducing taxes would improve performance', '3.80', '0.95']], col_widths=[3.5, 0.8, 0.8])
    source_note(doc)
    body(doc, 'The analysis showed that the respondents strongly agreed that tax burden affects cashflow with (M=4.3; SD=0.6) and they strongly agreed that tax rates reduce SME profitability with (M=4.2; SD=0.8). The respondents further agreed that high taxes limit business expansion with (M=4.1; SD=0.7). These findings indicate that tax rates are perceived as a significant barrier to SME performance in Eldoret City.')

    heading3(doc, '4.4.2 Tax Reforms')
    body(doc, 'On the second independent variable, the respondents were further asked to indicate the extent in which they agree with the various statements on tax reforms and SMEs performance. The following scale was used: 1 = Strongly Disagree, 2 = Disagree, 3 = Neutral, 4 = Agree and 5 = Strongly Agree.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.7: Tax Reforms')
    simple_table(doc, ['Statement', 'Mean', 'SD'], [['Frequent reforms create uncertainty', '4.12', '0.89'], ['Reform processes are transparent', '2.96', '1.00'], ['Reforms increased compliance costs', '4.05', '0.89'], ['Reforms simplified filing', '2.80', '1.05'], ['Reforms positively impact performance', '2.70', '1.10']], col_widths=[3.5, 0.8, 0.8])
    source_note(doc)
    body(doc, 'The findings indicate that frequent tax reforms create uncertainty in business planning (Mean=4.12; SD=0.89). This suggests that most SMEs agree that frequent policy changes disrupt long-term financial planning. The statement that tax reforms increase compliance costs recorded a high mean score (Mean=4.05; SD=0.89). However, respondents were neutral on whether reform processes are transparent (Mean=2.96; SD=1.00).')
    body(doc, 'Overall, the results suggest that tax reforms significantly influence SME performance, particularly through increased uncertainty and compliance costs.')

    heading3(doc, '4.4.3 Tax Incentives')
    body(doc, 'On the third independent variable, the respondents were also asked to indicate their views on tax incentives and their influence on business growth. The following scale was used: 1 = Strongly Disagree, 2 = Disagree, 3 = Neutral, 4 = Agree and 5 = Strongly Agree.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.8: Tax Incentives')
    simple_table(doc, ['Statement', 'Mean', 'SD'], [['Aware of tax incentives', '3.02', '1.08'], ['Incentives reduced tax burden', '3.45', '0.95'], ['Incentives encouraged investment', '3.80', '0.85'], ['Incentives improved profitability', '4.31', '0.74'], ['Incentives encouraged growth', '4.25', '0.78']], col_widths=[3.5, 0.8, 0.8])
    source_note(doc)
    body(doc, 'The findings show strong agreement that tax incentives positively influence SME performance. Reduced tax rates improving profitability recorded the highest mean (Mean=4.31, SD=0.74), indicating strong consensus among respondents. Similarly, tax incentives encouraging SME growth recorded a high mean (Mean=4.25, SD=0.78). However, awareness of available tax incentives was relatively low (Mean=3.02, SD=1.08).')

    heading3(doc, '4.4.4 Tax Compliance')
    body(doc, 'On the fourth independent variable, the respondents were asked to indicate the extent in which they agree with the various statements on tax compliance procedures and SMEs operational efficiency. The following scale was used: 1 = Strongly Disagree, 2 = Disagree, 3 = Neutral, 4 = Agree and 5 = Strongly Agree.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.9: Tax Compliance')
    simple_table(doc, ['Statement', 'Mean', 'SD'], [['Registration procedures were straightforward', '2.85', '1.10'], ['Filing returns is time-consuming', '4.15', '0.82'], ['Cost of compliance is high', '4.08', '0.88'], ['Easy to understand filing procedures', '2.70', '1.05'], ['Simplified compliance would improve efficiency', '4.35', '0.72']], col_widths=[3.5, 0.8, 0.8])
    source_note(doc)
    body(doc, 'The findings reveal that respondents strongly agreed that simplified compliance procedures would improve business efficiency (M=4.35; SD=0.72). Filing tax returns was perceived as time-consuming and complex (M=4.15; SD=0.82), and the cost of complying with tax requirements was considered high relative to business income (M=4.08; SD=0.88). However, respondents disagreed that tax registration procedures were straightforward (M=2.85; SD=1.10) and that filing procedures were easy to understand (M=2.70; SD=1.05). These findings indicate that tax compliance procedures pose a significant burden on SME operations in Eldoret City.')

    heading3(doc, '4.4.5 SMEs Performance')
    body(doc, 'On the dependent variable, the respondents were asked to indicate the extent in which they agree with the various statements on the SMEs performance. The following scale was used: 1 = Strongly Disagree, 2 = Disagree, 3 = Neutral, 4 = Agree and 5 = Strongly Agree.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.10: SMEs Performance')
    simple_table(doc, ['Statement', 'Mean', 'SD'], [['Revenue increased over 3 years', '3.50', '0.98'], ['Tax policies positively influenced growth', '2.85', '1.10'], ['Expanded workforce in recent years', '3.10', '1.05'], ['Profitability improved due to tax policies', '3.87', '0.91'], ['Overall performance improved', '3.74', '0.95']], col_widths=[3.5, 0.8, 0.8])
    source_note(doc)
    body(doc, 'Based on the analysis it was evident that the respondents agreed that tax paid by SMEs reduces their profitability with (M=3.87; SD=0.91) and they agreed that the amount of tax levied on the small-scale business was too much with (M=3.74; SD=0.95). The respondents were neutral on the statements that tax policies have positively influenced business growth (M=2.85; SD=1.10).')

    heading2(doc, '4.5 Inferential Statistics')
    body(doc, 'This section presents the inferential statistics used to establish the relationship between tax policies and SME performance. Both correlation and regression analyses were conducted.')

    heading3(doc, '4.5.1 Correlation Analysis')
    body(doc, 'Pearson correlation analysis was conducted to determine the strength and direction of the relationship between the independent variables (tax rates, tax reforms, tax incentives, and tax compliance) and the dependent variable (SME performance). The results are presented in Table 4.11 below.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.11: Pearson Correlation Analysis')
    simple_table(doc, ['Variable', 'Tax Rates', 'Tax Ref.', 'Tax Inc.', 'Tax Comp.', 'SME Perf.'],
        [['Tax Rates', '1.000', '0.345**', '-0.298*', '0.412**', '-0.512**'],
         ['Tax Reforms', '0.345**', '1.000', '-0.215', '0.385**', '-0.489**'],
         ['Tax Incentives', '-0.298*', '-0.215', '1.000', '-0.267*', '0.573**'],
         ['Tax Compliance', '0.412**', '0.385**', '-0.267*', '1.000', '-0.456**'],
         ['SME Perf.', '-0.512**', '-0.489**', '0.573**', '-0.456**', '1.000']],
        col_widths=[1.4, 0.85, 0.85, 0.85, 0.85, 0.85])
    source_note(doc, 'Note: ** indicates correlation is significant at the 0.01 level (2-tailed). * indicates correlation is significant at the 0.05 level (2-tailed).')
    body(doc, 'The correlation results in Table 4.11 show that tax rates had a significant negative relationship with SME performance (r=-0.512, p<0.01), indicating that higher tax rates are associated with lower SME performance. Tax reforms also had a significant negative relationship with SME performance (r=-0.489, p<0.01), suggesting that frequent and complex tax reforms negatively affect SMEs. Tax incentives had a significant positive relationship with SME performance (r=0.573, p<0.01), indicating that tax incentives enhance SME performance. Tax compliance had a significant negative relationship with SME performance (r=-0.456, p<0.01), suggesting that burdensome compliance procedures reduce SME operational efficiency.')

    heading3(doc, '4.5.2 Regression Analysis')
    body(doc, 'Multiple linear regression analysis was conducted to determine the joint effect of tax rates, tax reforms, tax incentives, and tax compliance on SME performance. The results are presented in Tables 4.12, 4.13, and 4.14 below.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.12: Model Summary')
    simple_table(doc, ['Model', 'R', 'R Square', 'Adjusted R\u00b2', 'Std. Error'], [['1', '0.716', '0.512', '0.482', '0.515']], col_widths=[0.8, 0.9, 1.0, 1.2, 1.2])
    source_note(doc, 'Predictors: (Constant), Tax Rates, Tax Reforms, Tax Incentives, Tax Compliance')
    body(doc, 'From Table 4.12, the coefficient of determination (R\u00b2=0.512) indicates that tax rates, tax reforms, tax incentives, and tax compliance jointly explain 51.2% of the variance in SME performance in Eldoret City. The remaining 48.8% is explained by other factors not captured in this study. The adjusted R\u00b2 of 0.482 confirms this relationship after accounting for the number of predictors in the model.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.13: ANOVA')
    simple_table(doc, ['Model', 'Sum of Sq.', 'df', 'Mean Sq.', 'F', 'Sig.'],
        [['Regression', '18.113', '4', '4.528', '17.044', '0.000'], ['Residual', '17.263', '65', '0.266', '', ''], ['Total', '35.376', '69', '', '', '']],
        col_widths=[1.2, 1.0, 0.5, 1.0, 0.8, 0.7])
    source_note(doc, 'Dependent Variable: SME Performance. Predictors: Tax Rates, Tax Reforms, Tax Incentives, Tax Compliance')
    body(doc, 'The ANOVA results in Table 4.13 indicate that the regression model was statistically significant (F=17.044, p=0.000<0.05). This means that tax rates, tax reforms, tax incentives, and tax compliance are significant joint predictors of SME performance in Eldoret City, Kenya.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.14: Regression Coefficients')
    simple_table(doc, ['Variable', 'B', 'Std. Error', 'Beta', 't', 'Sig.'],
        [['(Constant)', '2.524', '0.298', '', '8.470', '0.000'], ['Tax Rates', '-0.285', '0.084', '-0.308', '-3.393', '0.001'], ['Tax Reforms', '-0.196', '0.072', '-0.248', '-2.722', '0.008'], ['Tax Incentives', '0.361', '0.088', '0.396', '4.102', '0.000'], ['Tax Compliance', '-0.195', '0.082', '-0.228', '-2.378', '0.020']],
        col_widths=[1.8, 0.7, 1.0, 0.8, 0.8, 0.7])
    source_note(doc, 'Dependent Variable: SME Performance')
    body(doc, 'From Table 4.14, the regression equation is:')
    para_reg = doc.add_paragraph()
    def _fr2(txt, sub=False):
        r = para_reg.add_run(txt)
        r.font.size = Pt(12); r.font.bold = True; r.font.name = 'Times New Roman'
        if sub: r.font.subscript = True
    _fr2('Y = 2.524 - 0.285X'); _fr2('1', sub=True)
    _fr2(' - 0.196X'); _fr2('2', sub=True)
    _fr2(' + 0.361X'); _fr2('3', sub=True)
    _fr2(' - 0.195X'); _fr2('4', sub=True)
    para_reg.alignment = WD_ALIGN_PARAGRAPH.CENTER; _sp(para_reg, 6, 6); _line1(para_reg)
    body(doc, 'Where X\u2081 = Tax Rates, X\u2082 = Tax Reforms, X\u2083 = Tax Incentives, X\u2084 = Tax Compliance')
    body(doc, 'The regression results show that tax rates had a significant negative effect on SME performance (\u03b2=-0.285, t=-3.393, p=0.001<0.05). This implies that a unit increase in the tax rate burden reduces SME performance by 0.285 units, holding other variables constant. Tax reforms also had a significant negative effect on SME performance (\u03b2=-0.196, t=-2.722, p=0.008<0.05), indicating that complex and frequent tax reforms negatively affect SME performance.')
    body(doc, 'Tax incentives had a significant positive effect on SME performance (\u03b2=0.361, t=4.102, p=0.000<0.05), indicating that a unit increase in tax incentives leads to a 0.361 unit increase in SME performance. Tax incentives had the largest positive beta coefficient (0.396), making them the strongest positive predictor of SME performance among the four variables studied.')
    body(doc, 'Tax compliance had a significant negative effect on SME performance (\u03b2=-0.195, t=-2.378, p=0.020<0.05), indicating that burdensome tax compliance procedures reduce SME operational efficiency by 0.195 units for every unit increase in compliance burden. This confirms that complex registration requirements, time-consuming filing processes, and high compliance costs significantly impair SME performance.')

    p = heading_center(doc, 'CHAPTER FIVE', before=0, after=2)
    p.paragraph_format.page_break_before = True
    heading_center(doc, 'SUMMARY, CONCLUSIONS AND RECOMMENDATIONS', before=0, after=14)

    heading2(doc, '5.1 Introduction')
    body(doc, 'This chapter presents the summary of the key findings from the study, conclusions drawn from the findings, and recommendations made based on the conclusions. The chapter is structured to address all the research objectives of the study, which sought to evaluate the influence of tax compliance procedures on SME operational efficiency, the effect of tax rates on performance of SMEs, the effect of tax reforms on the performance of SMEs, and the effect of tax incentives on the growth and sustainability of SMEs in Eldoret City, Kenya.')

    heading2(doc, '5.2 Summary of Findings')
    body(doc, 'The study sought to examine the impact of tax policies on the performance of small and medium enterprises in Eldoret City, Kenya. A cross-sectional survey design was adopted and data collected from 70 SME owners and managers who returned questionnaires out of 78 distributed, giving a response rate of 89.7%.')

    heading3(doc, '5.2.1 Theoretical Framework')
    body(doc, 'The study was anchored on three theoretical frameworks: the Ability-to-Pay Theory, Economic Based Theories, and Optimal Tax Theory. The Ability-to-Pay Theory provided the basis for understanding why SMEs with lower incomes should bear a proportionately lower tax burden. Economic Based Theories explained the rational decision-making process of SME owners regarding tax compliance, while the Optimal Tax Theory provided a framework for evaluating the efficiency of the current tax system in Kenya. All three theories were applied across the four independent variables (tax rates, tax reforms, tax incentives, and tax compliance) to provide a comprehensive theoretical foundation for the study.')

    heading3(doc, '5.2.2 Tax Rates and SME Performance')
    body(doc, 'The findings revealed that tax rates had a significant negative impact on SME performance in Eldoret City. The majority of respondents strongly agreed that tax burden affects cashflow (M=4.3; SD=0.6), tax rates reduce SME profitability (M=4.2; SD=0.8), and high taxes limit business expansion (M=4.1; SD=0.7). The regression analysis confirmed that tax rates had a significant negative effect on SME performance (\u03b2=-0.285, p=0.001).')

    heading3(doc, '5.2.3 Tax Reforms and SME Performance')
    body(doc, 'The study found that tax reforms had a significant negative effect on SME performance. Respondents agreed that frequent tax reforms create uncertainty in business planning (M=4.12; SD=0.89) and that recent tax reforms have increased compliance costs (M=4.05; SD=0.89). However, respondents expressed that SMEs are not adequately informed about tax reforms (M=2.96; SD=1.00), highlighting a communication gap between KRA and SME operators. The regression results confirmed that tax reforms had a significant negative effect on SME performance (\u03b2=-0.196, p=0.008).')

    heading3(doc, '5.2.4 Tax Incentives and SME Performance')
    body(doc, 'The study established that tax incentives had a significant positive effect on SME performance. Respondents strongly agreed that reduced tax rates improve SME profitability (M=4.31; SD=0.74) and that tax incentives encourage SME growth (M=4.25; SD=0.78). The regression analysis confirmed that tax incentives had the strongest positive effect on SME performance (\u03b2=0.361, p=0.000), making it the most influential positive predictor among the four variables.')

    heading3(doc, '5.2.5 Tax Compliance and SME Performance')
    body(doc, 'The study found that tax compliance procedures had a significant negative effect on SME operational efficiency. Respondents strongly agreed that simplified compliance procedures would improve business efficiency (M=4.35; SD=0.72) and that filing tax returns is time-consuming and complex (M=4.15; SD=0.82). The cost of compliance was considered high relative to business income (M=4.08; SD=0.88). The regression analysis confirmed that tax compliance had a significant negative effect on SME performance (\u03b2=-0.195, p=0.020), indicating that burdensome compliance procedures impair SME operations.')

    heading3(doc, '5.2.6 Overall Model Findings')
    body(doc, 'The multiple regression analysis showed that tax rates, tax reforms, tax incentives, and tax compliance jointly explain 51.2% of the variance in SME performance (R\u00b2=0.512, F=17.044, p=0.000). The regression model was found to be statistically significant, confirming that tax policies as a whole have a significant effect on SME performance in Eldoret City, Kenya.')

    heading2(doc, '5.3 Conclusions')
    body(doc, 'Based on the findings of the study, the following conclusions were made:')
    body(doc, 'First, tax rates have a significant negative impact on the performance of SMEs in Eldoret City, Kenya. High tax rates reduce the profitability of SMEs, limit their ability to expand operations, and impair their cashflow. This conclusion is consistent with existing literature that has identified high tax rates as one of the major challenges facing SMEs in developing countries.')
    body(doc, 'Second, tax reforms negatively affect SME performance by creating uncertainty in business planning and increasing compliance costs. The frequent changes in tax regulations make it difficult for SME owners to plan their finances effectively. The lack of adequate communication of tax reforms to SME owners exacerbates this problem.')
    body(doc, 'Third, tax incentives have a significant positive effect on SME performance. When tax incentives such as reduced tax rates, tax holidays, and VAT exemptions are accessible and effectively implemented, they improve SME profitability, encourage business growth, and support business expansion. However, the study found that many SME owners in Eldoret City find it difficult to access government tax incentives, limiting their positive impact.')
    body(doc, 'Fourth, tax compliance procedures significantly influence the operational efficiency of SMEs. Complex registration requirements, burdensome filing procedures, and high compliance costs divert resources from productive business activities. Simplifying compliance processes is essential for improving SME performance.')
    body(doc, 'Overall, the study concludes that tax policies significantly influence the performance of SMEs in Eldoret City, Kenya. The design and implementation of tax policies must therefore take into account the unique characteristics and challenges of SMEs to promote their growth and sustainability.')

    heading2(doc, '5.4 Recommendations')
    body(doc, 'Based on the conclusions of the study, the following recommendations are made:')
    body(doc, 'First, the government through the Kenya Revenue Authority (KRA) should consider reducing the tax burden on SMEs, particularly through lower turnover tax rates and simplified VAT procedures. A tax rate structure that is sensitive to the size and revenue of SMEs would ensure that taxes do not stifle their growth and profitability.')
    body(doc, 'Second, the government should minimize the frequency of tax policy changes and, where reforms are necessary, ensure that adequate and timely communication is provided to SME owners. This can be achieved through targeted outreach programs, workshops, and digital communication channels that reach SME owners at the grassroots level.')
    body(doc, 'Third, the government should make tax incentives more accessible to SMEs by simplifying the application processes and removing bureaucratic barriers. Tax authorities should conduct awareness campaigns to educate SME owners on the available tax incentives and the procedures for accessing them.')
    body(doc, 'Fourth, KRA should simplify tax compliance procedures for SMEs, including streamlining tax registration, reducing the frequency and complexity of filing requirements, and investing in user-friendly digital tax platforms that reduce the cost and time of compliance.')
    body(doc, 'Fifth, the county government of Uasin Gishu should work in collaboration with KRA to create a conducive tax environment for SMEs in Eldoret City. This includes providing tax education programs and establishing SME support centers where business owners can seek assistance on tax compliance matters.')
    body(doc, 'Sixth, SME owners are encouraged to keep proper financial records to enable accurate tax assessment and to take advantage of available tax incentives. They should also engage with tax authorities proactively to understand their tax obligations and rights.')

    heading2(doc, '5.5 Suggestions for Further Research')
    body(doc, 'This study was limited to SMEs in Eldoret City and focused on tax policy variables (tax rates, tax reforms, tax incentives, and tax compliance). Future research should consider:')
    for sugg in [
        'i. Conducting a similar study in other urban centers in Kenya to enable comparison of findings across different regions.',
        'ii. Examining additional tax policy variables such as tax administration efficiency, taxpayer education, and digital tax systems and their effect on SME performance.',
        'iii. Conducting a longitudinal study to track changes in SME performance over time in response to specific tax policy changes.',
        'iv. Exploring the mediating role of tax compliance behavior in the relationship between tax policies and SME performance.',
    ]:
        body(doc, sugg)

    heading2(doc, '5.6 Limitations of the Study')
    body(doc, 'The findings of this study cannot be generalized to other towns since it was based on SMEs in Eldoret City. Also, the variables considered in the study, that is tax rates, tax reforms, tax incentives, and tax compliance, can only explain 51.2% of SMEs performance in Eldoret City. Finally, the study was only limited on the small and medium enterprises, leaving out other sectors which contribute to tax collection in the country.')

    p = heading_center(doc, 'REFERENCES', before=0, after=14)
    p.paragraph_format.page_break_before = True
    references = [
        'Ali, M., Sjursen, I. H., & Michelsen, J. (2015). Factors affecting tax compliance attitude in Africa: Evidence from Kenya, Tanzania, Uganda and South Africa. Working Paper Series.',
        'Atawodi, O. W., & Ojeka, S. A. (2012). Factors that affect tax compliance among small and medium enterprises (SMEs) in North Central Nigeria. International Journal of Business and Management, 7(12), 87-96.',
        'Awirothanon, K. (2019). Relationship between tax planning and financial performance in Thailand Stock Exchange. International Business and Global Economy, 38, 209-218.',
        'Bird, R. M., & Zolt, E. M. (2008). Technology and taxation in developing countries: From hand to mouse. National Tax Journal, 61(4), 791-821.',
        'Bjork, G. (2013). Tax reforms and tax compliance: The elusive quest for fiscal legitimacy. Journal of Tax Research, 11(1), 77-100.',
        'Bolboros, D. (2016). Impact of tax rates on financial performance of small enterprises. Annals of the University of Craiova, 13(1), 145-158.',
        'Borg, W. R., & Gall, M. D. (2013). Educational research: An introduction (8th ed.). Longman Publishers.',
        'Cobham, A. (2012). Tax havens and illicit flows. Global Governance Program, European University Institute.',
        'Cooper, D. R., & Schindler, P. S. (2013). Business research methods (12th ed.). McGraw-Hill Education.',
        'Creswell, J. W. (2010). Research design: Qualitative, quantitative, and mixed methods approaches (3rd ed.). SAGE Publications.',
        'Holban, O. I. (2017). The taxation of small and medium enterprises: Between priorities and options. Journal of Business and Economics, 8(3), 212-225.',
        'Hong, Q. N., & Pluye, P. (2018). Conceptual frameworks in mixed methods research. Journal of Mixed Methods Research, 12(2), 151-173.',
        'IEA. (2012). Kenya economic report 2012: Creating an enabling environment for stimulating investment for competitive and sustainable counties. Institute of Economic Affairs.',
        'Kothari, C. R. (2012). Research methodology: Methods and techniques (2nd ed.). New Age International Publishers.',
        'Kothari, C. R. (2013). Research methodology: Methods and techniques (3rd ed.). New Age International Publishers.',
        'Mukras, M. S. (2003). Poverty reduction through strengthening small and medium enterprises. African Development, 28(1-2), 69-89.',
        'Mugenda, O. M., & Mugenda, A. G. (2013). Research methods: Quantitative and qualitative approaches. African Centre for Technology Studies.',
        'Neuman, W. L. (2014). Social research methods: Qualitative and quantitative approaches (7th ed.). Pearson Education.',
        'Ngulube, P. (2012). Using mixed methods research design in educational research. The Journal of Educational Research in Africa, 12(4), 22-39.',
        'Noor-Halp, M. (2011). Tax rates and financial performance: Evidence from Nigeria. International Journal of Finance and Accounting, 1(2), 89-102.',
        'Ocheni, S. I. (2015). Effect of multiple taxation on the performance of small and medium scale business enterprises in Lokoja, Kogi State. Mediterranean Journal of Social Sciences, 6(1), 86-96.',
        'Osambo, G. N. (2019). Effect of tax compliance on performance of small and medium enterprises in Nairobi, Kenya. International Journal of Research and Innovation in Social Science, 3(9), 1-15.',
        'Pope, J., & Abdul-Jabbar, H. (2018). Tax compliance costs of small and medium enterprises in Malaysia: Policy implications. International Journal of Business Research, 18(3), 65-83.',
        'Tee, E., Boadi, L. A., & Opoku, R. T. (2016). The effect of tax payment on the performance of SMEs: The case of selected SMEs in Ga West Municipal Assembly. European Journal of Business and Management, 8(20), 119-125.',
        'Toader, C., & Dragoti, E. (2014). The impact of taxation and net profit of the firm in Australia. Annals of the University of Craiova, Economic Sciences Series, 12(1), 52-64.',
        'Wagacha, M. (2019). Tax reform and economic development in Kenya. Nairobi: Kenya Institute for Public Policy Research and Analysis.',
        'World Bank. (2015). Small and medium enterprises (SMEs): Finance and development. World Bank Group Report.',
    ]
    for ref in references:
        p = doc.add_paragraph()
        r = p.add_run(ref)
        r.font.size = Pt(11); r.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _sp(p, 0, 6); _line15(p)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        p.paragraph_format.left_indent = Inches(0.4)

    p = heading_center(doc, 'APPENDICES', before=0, after=14)
    p.paragraph_format.page_break_before = True
    heading_center(doc, 'Appendix I: Research Questionnaire', size=12, before=0, after=10)

    def bold_body(doc, text, before=0, after=6):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _sp(para, before, after)
        _line15(para)
        return para

    def bold_label_body(doc, label, rest, before=0, after=6):
        para = doc.add_paragraph()
        r1 = para.add_run(label)
        r1.font.size = Pt(12); r1.font.name = 'Times New Roman'; r1.font.bold = True
        r2 = para.add_run(rest)
        r2.font.size = Pt(12); r2.font.name = 'Times New Roman'
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _sp(para, before, after); _line15(para)
        return para

    def section_title(doc, text, before=8, after=6):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _sp(para, before, after)
        _line15(para)
        return para

    CB = '\u2610'

    bold_body(doc, 'MOI UNIVERSITY', before=0, after=2)
    bold_body(doc, 'SCHOOL OF BUSINESS AND ECONOMICS', before=0, after=2)
    bold_body(doc, 'DEPARTMENT OF ACCOUNTING AND FINANCE', before=0, after=8)
    body(doc, '')
    bold_body(doc, 'RESEARCH QUESTIONNAIRE', before=0, after=8)

    bold_label_body(doc, 'Study Title: ', 'The Impact of Tax Policies on the Performance of Small and Medium Enterprises in Eldoret City, Kenya')
    bold_label_body(doc, 'Instructions: ', 'This questionnaire is for academic research purposes only. Your responses are completely confidential and anonymous. Please answer all questions honestly. Do not write your name anywhere on this questionnaire.')

    body(doc, '')
    section_title(doc, 'SECTION A: Demographic Information')

    body(doc, f'1. Type of Business:')
    body(doc, f'   {CB} Retail    {CB} Manufacturing    {CB} Services    {CB} Wholesale', indent=True)
    body(doc, f'   {CB} Food and Hospitality    {CB} Agro-processing    {CB} Other (specify): ________', indent=True)
    body(doc, f'2. Years of Operation:')
    body(doc, f'   {CB} Below 2 years    {CB} 2-5 years    {CB} Above 5 years', indent=True)
    body(doc, f'3. Number of Employees:')
    body(doc, f'   {CB} 1-10    {CB} 11-15    {CB} 16-25    {CB} Over 25', indent=True)
    body(doc, f'4. Annual Turnover of the Business:')
    body(doc, f'   {CB} Below KES 500,000    {CB} KES 500,001 - 1,000,000', indent=True)
    body(doc, f'   {CB} KES 1,000,001 - 2,000,000    {CB} KES 2,000,001 - 5,000,000', indent=True)
    body(doc, f'   {CB} Above KES 5,000,000', indent=True)

    body(doc, '')
    bold_label_body(doc, 'Rating Scale: ', '1 = Strongly Disagree, 2 = Disagree, 3 = Neutral, 4 = Agree, 5 = Strongly Agree')

    body(doc, '')
    section_title(doc, 'SECTION B: Tax Rates')
    for i, stmt in enumerate([
        'The current tax rates imposed on SMEs are too high.',
        'Tax burden significantly affects the cash flow of my business.',
        'High tax rates reduce the profitability of my business.',
        'Tax rates discourage me from expanding my business.',
        'Reducing tax rates would improve my business performance.',
    ], 5):
        body(doc, f'{i}. {stmt}', indent=True)

    body(doc, '')
    section_title(doc, 'SECTION C: Tax Reforms')
    for i, stmt in enumerate([
        'Frequent changes in tax laws create uncertainty in my business planning.',
        'Tax reform processes are transparent and fair to SMEs.',
        'New tax reforms have increased the cost of compliance for my business.',
        'Tax reforms have simplified tax filing processes for SMEs.',
        'Tax reforms have had a positive impact on my business performance.',
    ], 10):
        body(doc, f'{i}. {stmt}', indent=True)

    body(doc, '')
    section_title(doc, 'SECTION D: Tax Incentives')
    for i, stmt in enumerate([
        'I am aware of tax incentives available to SMEs in Kenya.',
        'Tax incentives have reduced my overall tax burden.',
        'Tax incentives have encouraged me to invest more in my business.',
        'Tax incentives have improved the profitability of my business.',
        'I believe tax incentives have encouraged business growth in Eldoret City.',
    ], 15):
        body(doc, f'{i}. {stmt}', indent=True)

    body(doc, '')
    section_title(doc, 'SECTION E: Tax Compliance')
    for i, stmt in enumerate([
        'Tax registration procedures for my business were straightforward and easy.',
        'Filing tax returns is time-consuming and complex for my business.',
        'The cost of complying with tax requirements is high relative to my business income.',
        'I find it easy to understand and follow the tax filing procedures.',
        'Simplified compliance procedures would improve my business efficiency.',
    ], 20):
        body(doc, f'{i}. {stmt}', indent=True)

    body(doc, '')
    section_title(doc, 'SECTION F: SME Performance')
    for i, stmt in enumerate([
        'My business revenue has increased over the past three years.',
        'Tax policies have positively influenced the growth of my business.',
        'My business has been able to expand its workforce in recent years.',
        'The profitability of my business has improved due to favorable tax policies.',
        'Overall, my business performance has improved in the last three years.',
    ], 25):
        body(doc, f'{i}. {stmt}', indent=True)

    body(doc, '')
    bold_body(doc, 'Thank you for your participation.')

    s = doc.sections[-1]
    sectPr = s._sectPr
    for existing in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(existing)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), 'decimal')
    pgNumType.set(qn('w:start'), '1')
    sectPr.append(pgNumType)

    add_page_numbers(doc)

    output_path = 'Calvince_Odhiambo_Research_Project.docx'
    doc.save(output_path)
    print(f'DOCX saved: {output_path}')
    return output_path


if __name__ == '__main__':
    create_docx()
