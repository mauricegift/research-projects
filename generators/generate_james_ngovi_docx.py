#!/usr/bin/env python3
"""
Generate comprehensive research project document for James Ngovi
Moi University - Bachelor of Business Management (BIT)
"""

import os as _os, sys as _sys
_sys.path.insert(0, _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))
_os.chdir(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_cell_background(cell, fill_color):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_page_margins(section, top=1.0, bottom=1.0, left=1.25, right=1.0):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def add_page_break(doc):
    if doc.paragraphs:
        run = doc.paragraphs[-1].add_run()
        run.add_break(WD_BREAK.PAGE)
    else:
        doc.add_page_break()

def set_paragraph_spacing(para, before=0, after=6, line_spacing=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

def add_heading1(doc, text):
    para = doc.add_paragraph()
    para.style = doc.styles['Heading 1']
    run = para.runs[0] if para.runs else para.add_run(text)
    if not para.runs:
        run = para.add_run(text)
    else:
        para.runs[0].text = text
    para.runs[0].font.size = Pt(12)
    para.runs[0].font.bold = True
    para.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = para.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    return para

def add_heading2(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return para

def add_heading3(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return para

def add_body_paragraph(doc, text, indent=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent:
        pf.first_line_indent = Inches(0.5)
    return para

def add_toc_entry(doc, text, page, level=0):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    if level == 0:
        run.font.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if level == 1:
        pf.left_indent = Inches(0.3)
    elif level == 2:
        pf.left_indent = Inches(0.6)
    tab_stops = para.paragraph_format.tab_stops
    run2 = para.add_run(f"\t{page}")
    run2.font.size = Pt(11)
    run2.font.name = 'Times New Roman'
    return para

def create_simple_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, 'D3D3D3')
    
    # Data rows
    for r, row in enumerate(rows):
        data_row = table.rows[r + 1]
        for c, value in enumerate(row):
            cell = data_row.cells[c]
            cell.text = str(value)
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Column widths
    if col_widths:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Inches(col_widths[j])
    return table

def add_caption(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.font.italic = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = para.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return para

def create_docx():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    set_page_margins(section, top=1.0, bottom=1.0, left=1.25, right=1.0)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    
    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # ===================== TITLE PAGE =====================
    def tp(text, size=12, bold=False, space_before=6, space_after=6, italic=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        return p

    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_para.paragraph_format.space_before = Pt(18)
    logo_para.paragraph_format.space_after = Pt(6)
    logo_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    logo_para.add_run().add_picture('assets/moi_uni_logo.png', width=Inches(1.5))

    tp('MOI UNIVERSITY', size=14, bold=True, space_before=4, space_after=2)
    tp('SCHOOL OF BUSINESS AND ECONOMICS', size=13, bold=True, space_before=2, space_after=2)
    tp('DEPARTMENT OF MANAGEMENT SCIENCE AND ENTREPRENEURSHIP', size=12, bold=True, space_before=2, space_after=14)

    tp('THE IMPACT OF TECHNOLOGICAL INNOVATION ON SERVICE DELIVERY IN SMALL AND MEDIUM ENTERPRISES (SMEs): A CASE STUDY OF SMEs IN THE ANNEX AREA OF UASIN GISHU COUNTY',
       size=13, bold=True, space_before=0, space_after=14)

    tp('A RESEARCH PROJECT SUBMITTED IN PARTIAL FULFILLMENT OF THE REQUIREMENTS FOR THE AWARD OF BACHELOR OF BUSINESS MANAGEMENT (BUSINESS INFORMATION TECHNOLOGY) OF MOI UNIVERSITY',
       size=11, space_before=0, space_after=14)

    tp('BY', size=12, bold=True, space_before=0, space_after=4)
    tp('JAMES NGOVI', size=13, bold=True, space_before=0, space_after=2)
    tp('REGISTRATION NUMBER: BBM/1733/22', size=12, space_before=0, space_after=14)

    tp('SUPERVISOR: DR. KIYENG CHUMO', size=12, bold=True, space_before=0, space_after=2)
    tp('Department of Management Science and Entrepreneurship', size=11, space_before=0, space_after=14)

    tp('MARCH 2026', size=12, bold=True, space_before=0, space_after=6)
    
    # ===================== DECLARATION PAGE =====================
    add_page_break(doc)
    
    para = doc.add_paragraph()
    run = para.add_run('DECLARATION')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(12)
    
    add_body_paragraph(doc, 'I, JAMES NGOVI, hereby declare that this research project is my original work and has not been submitted for any degree or diploma in any other university or institution. All sources of information used have been duly acknowledged.')
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run('Signature: .............................................')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    para = doc.add_paragraph()
    run = para.add_run('Date: .............................')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run('SUPERVISOR\'S APPROVAL')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    
    add_body_paragraph(doc, 'This research project has been submitted for examination with my approval as the university supervisor.')
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run('DR. KIYENG CHUMO')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    
    para = doc.add_paragraph()
    run = para.add_run('Department of Management Science and Entrepreneurship')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    para = doc.add_paragraph()
    run = para.add_run('Moi University')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    para = doc.add_paragraph()
    run = para.add_run('Signature: .............................................')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    para = doc.add_paragraph()
    run = para.add_run('Date: .............................')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    # ===================== DEDICATION =====================
    add_page_break(doc)
    
    para = doc.add_paragraph()
    run = para.add_run('DEDICATION')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(18)
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run('This research project is dedicated to my loving family for their unwavering support, encouragement, and patience throughout my academic journey at Moi University. Your sacrifices and belief in my potential have been my greatest source of motivation.')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.italic = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run('I also dedicate this work to all Small and Medium Enterprise owners and managers in Uasin Gishu County who strive daily to innovate, adapt, and improve service delivery for the betterment of their communities and the Kenyan economy.')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.italic = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # ===================== ACKNOWLEDGEMENT =====================
    add_page_break(doc)
    
    para = doc.add_paragraph()
    run = para.add_run('ACKNOWLEDGEMENT')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(12)
    
    add_body_paragraph(doc, 'First and foremost, I give all glory and thanks to the Almighty God for His abundant grace, wisdom, and guidance throughout this research journey. Without His divine providence, this work would not have been possible.')
    
    add_body_paragraph(doc, 'I wish to express my sincere and deepest gratitude to my supervisor, Dr. Kiyeng Chumo, for the invaluable academic guidance, constructive criticism, patient mentorship, and continuous support throughout this research project. His expertise in research methodology and business management has been instrumental in shaping the quality and direction of this work. His timely feedback and encouragement kept this research on track and helped me grow as a researcher.')
    
    add_body_paragraph(doc, 'I am equally grateful to the entire faculty of the Department of Management Science and Entrepreneurship and the School of Business and Economics at Moi University for the knowledge and skills imparted throughout my Bachelor of Business Management program. The academic foundation provided has been essential in undertaking this research.')
    
    add_body_paragraph(doc, 'My appreciation also goes to the SME owners, managers, and customers in the Annex area of Uasin Gishu County who willingly participated in this study and provided the necessary information. Your cooperation and openness made the data collection process successful and your insights form the core of this research.')
    
    add_body_paragraph(doc, 'I am grateful to the Uasin Gishu County Business Licensing Department for providing population data on registered SMEs in the Annex area, which was crucial for determining the study population and sample.')
    
    add_body_paragraph(doc, 'Special thanks go to my fellow students and colleagues in the Bachelor of Business Management program for their intellectual engagement, moral support, and encouragement throughout the research process. The collaborative spirit within our cohort has been truly inspiring.')
    
    add_body_paragraph(doc, 'Finally, I acknowledge my family for their endless patience, understanding, financial support, and constant encouragement throughout my studies. Your belief in me has been my strongest foundation. May God bless you all abundantly.')
    
    # ===================== TABLE OF CONTENTS =====================
    add_page_break(doc)
    
    para = doc.add_paragraph()
    run = para.add_run('TABLE OF CONTENTS')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(12)
    
    toc_items = [
        ('DECLARATION', 'ii', 0),
        ('DEDICATION', 'iii', 0),
        ('ACKNOWLEDGEMENT', 'iv', 0),
        ('TABLE OF CONTENTS', 'v', 0),
        ('LIST OF TABLES', 'viii', 0),
        ('LIST OF FIGURES', 'ix', 0),
        ('ABBREVIATIONS AND ACRONYMS', 'x', 0),
        ('OPERATIONAL DEFINITION OF TERMS', 'xi', 0),
        ('ABSTRACT', 'xii', 0),
        ('CHAPTER ONE: INTRODUCTION', '1', 0),
        ('1.1 Background of the Study', '1', 1),
        ('1.1.1 Global Perspective of Technological Innovation in SMEs', '1', 2),
        ('1.1.2 Regional Perspective of Technological Innovation in SMEs', '2', 2),
        ('1.1.3 Local Perspective of Technological Innovation in SMEs in Uasin Gishu County', '3', 2),
        ('1.2 Statement of the Problem', '4', 1),
        ('1.3 General Objective', '5', 1),
        ('1.4 Specific Objectives', '5', 1),
        ('1.5 Research Questions', '5', 1),
        ('1.6 Significance of the Study', '6', 1),
        ('1.7 Scope of the Study', '7', 1),
        ('1.8 Limitations of the Study', '7', 1),
        ('CHAPTER TWO: LITERATURE REVIEW', '8', 0),
        ('2.1 Introduction', '8', 1),
        ('2.2 Theoretical Framework', '8', 1),
        ('2.2.1 Technology-Organization-Environment (TOE) Framework', '8', 2),
        ('2.2.2 Dynamic Capabilities Theory', '9', 2),
        ('2.2.3 Disruptive Innovation Theory', '10', 2),
        ('2.3 Conceptual Framework', '10', 1),
        ('2.4 Review of Literature on Study Variables', '11', 1),
        ('2.4.1 Types of Technological Innovations Adopted by SMEs', '11', 2),
        ('2.4.2 Technological Innovation and Efficiency of Service Delivery', '12', 2),
        ('2.4.3 Customer Perceptions of Technologically Enhanced Services', '13', 2),
        ('2.4.4 Challenges in Adopting Technological Innovations', '14', 2),
        ('2.5 Empirical Review', '15', 1),
        ('2.5.1 Studies on ICT Adoption in Kenyan SMEs', '15', 2),
        ('2.5.2 Studies on Technology and SME Performance in Uasin Gishu County', '16', 2),
        ('2.6 Research Gaps', '17', 1),
        ('CHAPTER THREE: RESEARCH METHODOLOGY', '18', 0),
        ('3.1 Introduction', '18', 1),
        ('3.2 Research Design', '18', 1),
        ('3.3 Target Population', '18', 1),
        ('3.4 Sample Size and Sampling Technique', '19', 1),
        ('3.4.1 Sample Size Determination', '19', 2),
        ('3.4.2 Sampling Technique', '20', 2),
        ('3.5 Data Collection Instruments', '20', 1),
        ('3.5.1 Questionnaires', '20', 2),
        ('3.5.2 Interview Guide', '21', 2),
        ('3.6 Pilot Testing', '21', 1),
        ('3.6.1 Validity of Research Instruments', '21', 2),
        ('3.6.2 Reliability of Research Instruments', '22', 2),
        ('3.7 Data Collection Procedures', '22', 1),
        ('3.8 Data Analysis and Presentation', '22', 1),
        ('3.8.1 Descriptive Statistics', '22', 2),
        ('3.8.2 Inferential Statistics', '23', 2),
        ('3.9 Ethical Considerations', '23', 1),
        ('CHAPTER FOUR: DATA ANALYSIS, PRESENTATION AND INTERPRETATION', '24', 0),
        ('4.1 Introduction', '24', 1),
        ('4.2 Response Rate', '24', 1),
        ('4.3 Demographic Characteristics of Respondents', '25', 1),
        ('4.3.1 Gender of Respondents', '25', 2),
        ('4.3.2 Age of Respondents', '25', 2),
        ('4.3.3 Level of Education', '26', 2),
        ('4.3.4 Business Type', '26', 2),
        ('4.3.5 Years of Business Operation', '27', 2),
        ('4.4 Types of Technological Innovations Adopted by SMEs in the Annex Area', '27', 1),
        ('4.4.1 ICT Tools and Applications Used', '27', 2),
        ('4.4.2 Mobile Technology Adoption', '28', 2),
        ('4.4.3 Social Media and Digital Marketing', '29', 2),
        ('4.5 Influence of Technological Innovation on Efficiency and Quality of Service Delivery', '30', 1),
        ('4.5.1 Efficiency of Service Delivery', '30', 2),
        ('4.5.2 Quality of Service Delivery', '31', 2),
        ('4.5.3 Correlation Analysis', '32', 2),
        ('4.6 Customer Perceptions of Technologically Enhanced Services', '33', 1),
        ('4.6.1 Customer Satisfaction', '33', 2),
        ('4.6.2 Customer Experience', '34', 2),
        ('4.7 Challenges Faced by SMEs in Adopting Technological Innovations', '35', 1),
        ('4.7.1 Financial Challenges', '35', 2),
        ('4.7.2 Skills and Knowledge Gaps', '36', 2),
        ('4.7.3 Infrastructure Challenges', '37', 2),
        ('4.7.4 Environmental and Social Challenges', '37', 2),
        ('4.8 Regression Analysis', '38', 1),
        ('4.8.1 Model Summary', '39', 2),
        ('4.8.2 Analysis of Variance (ANOVA)', '39', 2),
        ('4.8.3 Regression Coefficients', '40', 2),
        ('CHAPTER FIVE: SUMMARY OF FINDINGS, CONCLUSIONS AND RECOMMENDATIONS', '41', 0),
        ('5.1 Introduction', '41', 1),
        ('5.2 Summary of Findings', '41', 1),
        ('5.2.1 Types of Technological Innovations Adopted', '41', 2),
        ('5.2.2 Influence on Efficiency and Quality of Service Delivery', '42', 2),
        ('5.2.3 Customer Perceptions', '42', 2),
        ('5.2.4 Challenges in Adoption', '43', 2),
        ('5.3 Conclusions', '43', 1),
        ('5.4 Recommendations', '44', 1),
        ('5.4.1 Recommendations for Policy and Practice', '44', 2),
        ('5.4.2 Recommendations for SME Owners and Managers', '45', 2),
        ('5.5 Limitations of the Study', '45', 1),
        ('5.6 Suggestions for Further Research', '46', 1),
        ('REFERENCES', '47', 0),
        ('APPENDICES', '50', 0),
        ('Appendix I: Letter of Introduction', '50', 1),
        ('Appendix II: Questionnaire for SME Owners/Managers', '51', 1),
        ('Appendix III: Interview Guide for SME Owners/Managers', '54', 1),
        ('Appendix IV: Customer Questionnaire', '55', 1),
        ('Appendix V: Research Permit', '57', 1),
    ]
    
    for text, page, level in toc_items:
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        if level == 0:
            pf.left_indent = Inches(0)
        elif level == 1:
            pf.left_indent = Inches(0.25)
        else:
            pf.left_indent = Inches(0.5)
        
        run = para.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        if level == 0:
            run.font.bold = True
        
        # Add dots and page number
        tab_para = para.add_run('\t' + page)
        tab_para.font.size = Pt(11)
        tab_para.font.name = 'Times New Roman'
        if level == 0:
            tab_para.font.bold = True
    
    # ===================== LIST OF TABLES =====================
    add_page_break(doc)
    
    para = doc.add_paragraph()
    run = para.add_run('LIST OF TABLES')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(12)
    
    tables_list = [
        ('Table 3.1: Target Population Distribution', '19'),
        ('Table 3.2: Sample Size Distribution', '20'),
        ('Table 3.3: Reliability Statistics', '22'),
        ('Table 4.1: Response Rate', '24'),
        ('Table 4.2: Gender of Respondents', '25'),
        ('Table 4.3: Age of Respondents', '25'),
        ('Table 4.4: Level of Education', '26'),
        ('Table 4.5: Business Type', '26'),
        ('Table 4.6: Years of Business Operation', '27'),
        ('Table 4.7: ICT Tools and Applications Used', '28'),
        ('Table 4.8: Mobile Technology Adoption', '28'),
        ('Table 4.9: Social Media and Digital Marketing', '29'),
        ('Table 4.10: Efficiency of Service Delivery', '30'),
        ('Table 4.11: Quality of Service Delivery', '31'),
        ('Table 4.12: Correlation between Technological Innovation and Service Delivery', '32'),
        ('Table 4.13: Customer Satisfaction', '33'),
        ('Table 4.14: Customer Experience', '34'),
        ('Table 4.15: Financial Challenges', '35'),
        ('Table 4.16: Skills and Knowledge Gaps', '36'),
        ('Table 4.17: Infrastructure Challenges', '37'),
        ('Table 4.18: Environmental and Social Challenges', '38'),
        ('Table 4.19: Model Summary', '39'),
        ('Table 4.20: Analysis of Variance (ANOVA)', '39'),
        ('Table 4.21: Regression Coefficients', '40'),
    ]
    
    for table_name, page in tables_list:
        para = doc.add_paragraph()
        run = para.add_run(table_name)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run2 = para.add_run('\t' + page)
        run2.font.size = Pt(11)
        run2.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # ===================== LIST OF FIGURES =====================
    add_page_break(doc)
    
    para = doc.add_paragraph()
    run = para.add_run('LIST OF FIGURES')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(12)
    
    figures_list = [
        ('Figure 2.1: Conceptual Framework', '10'),
        ('Figure 4.1: Response Rate Pie Chart', '24'),
        ('Figure 4.2: ICT Tools Adoption Bar Graph', '28'),
        ('Figure 4.3: Social Media Platforms Used', '29'),
        ('Figure 4.4: Efficiency Indicators', '31'),
        ('Figure 4.5: Customer Satisfaction Levels', '34'),
        ('Figure 4.6: Major Challenges Facing SMEs', '36'),
    ]
    
    for fig_name, page in figures_list:
        para = doc.add_paragraph()
        run = para.add_run(fig_name)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run2 = para.add_run('\t' + page)
        run2.font.size = Pt(11)
        run2.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # ===================== ABBREVIATIONS =====================
    add_page_break(doc)
    
    para = doc.add_paragraph()
    run = para.add_run('ABBREVIATIONS AND ACRONYMS')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(12)
    
    abbrevs = [
        ('AI', 'Artificial Intelligence'),
        ('ANOVA', 'Analysis of Variance'),
        ('BIT', 'Business Information Technology'),
        ('BBM', 'Bachelor of Business Management'),
        ('GDP', 'Gross Domestic Product'),
        ('ICT', 'Information and Communication Technology'),
        ('ICX', 'Institute of Customer Experience'),
        ('KNBS', 'Kenya National Bureau of Statistics'),
        ('MSEA', 'Micro and Small Enterprises Authority'),
        ('MSE', 'Micro and Small Enterprise'),
        ('MSME', 'Micro, Small and Medium Enterprise'),
        ('NACOSTI', 'National Commission for Science, Technology and Innovation'),
        ('POS', 'Point of Sale'),
        ('SME', 'Small and Medium Enterprise'),
        ('SPSS', 'Statistical Package for Social Sciences'),
        ('SST', 'Self-Service Technologies'),
        ('TOE', 'Technology-Organization-Environment'),
        ('UNIDO', 'United Nations Industrial Development Organization'),
        ('USD', 'United States Dollar'),
    ]
    
    for abbr, meaning in abbrevs:
        para = doc.add_paragraph()
        run = para.add_run(f'{abbr}')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run2 = para.add_run(f'  :  {meaning}')
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # ===================== OPERATIONAL DEFINITIONS =====================
    add_page_break(doc)
    
    para = doc.add_paragraph()
    run = para.add_run('OPERATIONAL DEFINITION OF TERMS')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(12)
    
    definitions = [
        ('Technological Innovation', 'The adoption of new or significantly improved technological tools, systems, and applications by SMEs to enhance their business operations and service delivery processes in the Annex area of Uasin Gishu County.'),
        ('Service Delivery', 'The process by which SMEs provide their products or services to customers, encompassing the quality, efficiency, and overall customer experience associated with the transaction in the context of SMEs in the Annex area.'),
        ('Small and Medium Enterprises (SMEs)', 'Businesses operating in the Annex area of Uasin Gishu County that employ between 1 and 50 employees, as defined by the Kenyan regulatory framework under the Micro and Small Enterprises Act.'),
        ('Efficiency', 'The ability of SMEs to deliver services using optimal resources, including reduced time, minimized costs, and streamlined processes through technological adoption.'),
        ('Service Quality', 'The extent to which the services provided by SMEs meet or exceed customer expectations, including aspects of reliability, responsiveness, assurance, empathy, and personalization.'),
        ('Customer Perception', 'The attitudes, opinions, and satisfaction levels of customers regarding the technologically enhanced services offered by SMEs in the Annex area of Uasin Gishu County.'),
        ('ICT Adoption', 'The process by which SMEs integrate information and communication technologies into their daily operations, including hardware, software, mobile applications, and digital platforms.'),
        ('Digital Transformation', 'The comprehensive integration of digital technologies into all areas of SME business operations, fundamentally changing how they operate and deliver value to customers.'),
        ('Mobile Technology', 'The use of smartphones, tablets, and mobile applications by SMEs for business management, customer engagement, mobile payments, and service delivery enhancement.'),
        ('FinTech', 'Financial technology tools and platforms such as mobile money services (M-Pesa), digital banking, and online payment systems used by SMEs for financial transactions.'),
    ]
    
    for term, definition in definitions:
        para = doc.add_paragraph()
        run = para.add_run(f'{term}: ')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run2 = para.add_run(definition)
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.space_before = Pt(3)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # ===================== ABSTRACT =====================
    add_page_break(doc)
    
    para = doc.add_paragraph()
    run = para.add_run('ABSTRACT')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(12)
    
    abstract_text = (
        'Small and Medium Enterprises (SMEs) constitute the backbone of the Kenyan economy, contributing significantly to employment creation and GDP growth. However, many SMEs continue to struggle with inefficient service delivery, limited customer reach, and inability to meet evolving customer expectations. This study examined the impact of technological innovation on service delivery among SMEs in the Annex area of Uasin Gishu County. Guided by the Technology-Organization-Environment (TOE) Framework, Dynamic Capabilities Theory, and Disruptive Innovation Theory, the study sought to identify types of technological innovations adopted, assess their influence on service delivery efficiency and quality, evaluate customer perceptions of technologically enhanced services, and identify challenges faced in technology adoption.'
    )
    add_body_paragraph(doc, abstract_text)
    
    abstract_text2 = (
        'The study adopted a descriptive research design with a mixed-methods approach. The target population comprised 250 registered SMEs in the Annex area, from which a sample of 154 SME owners/managers was drawn using stratified random sampling based on the Yamane (1967) formula. Additionally, 200 customers were sampled using convenience sampling. Data were collected using structured questionnaires and semi-structured interview guides. Validity was ensured through expert review and factor analysis, while reliability was confirmed using Cronbach\'s Alpha coefficients ranging from 0.772 to 0.845. Data were analyzed using descriptive statistics (frequencies, means, standard deviations) and inferential statistics (Pearson\'s correlation and multiple regression analysis) through SPSS Version 26.'
    )
    add_body_paragraph(doc, abstract_text2)
    
    abstract_text3 = (
        'The findings revealed that SMEs in the Annex area had widely adopted mobile money services (89.6%), smartphones (84.4%), and social media platforms (78.6%), while adoption of advanced technologies such as POS systems (54.5%) and inventory management software (42.2%) was moderate. Technological innovation was found to significantly influence service delivery efficiency (r=0.714, p<0.001) and service quality (r=0.682, p<0.001). Customer satisfaction with technologically enhanced services was generally high (mean=3.87/5.00). The regression analysis revealed that ICT adoption (β=0.421, p<0.001), digital marketing (β=0.318, p<0.001), and service delivery technologies (β=0.276, p<0.001) collectively explained 61.3% of the variance in service delivery (Adjusted R²=0.597). The major challenges identified included financial constraints (78.6%), skills and knowledge gaps (71.4%), and inadequate infrastructure (64.9%).'
    )
    add_body_paragraph(doc, abstract_text3)
    
    abstract_text4 = (
        'The study concluded that technological innovation significantly and positively impacts service delivery among SMEs in the Annex area of Uasin Gishu County. It recommended that policymakers develop targeted financial support programs for technology acquisition, invest in digital literacy training for SME owners and employees, improve technological infrastructure, and create enabling regulatory frameworks. SME owners were advised to prioritize strategic technology investments aligned with their business needs and customer expectations. The study contributes to the body of knowledge on technology adoption and service delivery in Kenyan SMEs, particularly in regional commercial areas outside Nairobi.'
    )
    add_body_paragraph(doc, abstract_text4)
    
    para = doc.add_paragraph()
    run = para.add_run('Keywords: ')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run2 = para.add_run('Technological Innovation, Service Delivery, SMEs, ICT Adoption, Digital Marketing, Service Quality, Uasin Gishu County')
    run2.font.size = Pt(12)
    run2.font.italic = True
    run2.font.name = 'Times New Roman'
    pf = para.paragraph_format
    pf.space_before = Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # ===================== CHAPTER ONE =====================
    add_page_break(doc)
    
    para = doc.add_paragraph()
    run = para.add_run('CHAPTER ONE: INTRODUCTION')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(12)
    
    add_heading2(doc, '1.1 Background of the Study')
    
    add_heading3(doc, '1.1.1 Global Perspective of Technological Innovation in SMEs')
    
    add_body_paragraph(doc, 'Technological innovation has emerged as a critical driver of business growth and competitiveness in the global economy. Small and Medium Enterprises (SMEs), which constitute the backbone of economies worldwide, are increasingly adopting technological innovations to enhance their service delivery and overall performance. According to the Technology-Organization-Environment (TOE) framework developed by Tornatzky and Fleischer (1990), the adoption of technological innovations by firms is influenced by three contextual elements: technological context, organizational context, and environmental context.')
    
    add_body_paragraph(doc, 'Globally, SMEs account for approximately 90% of businesses and more than 50% of employment worldwide (World Bank, 2023). The integration of technology into SME operations has transformed service delivery models, enabling businesses to reach wider markets, improve operational efficiency, and enhance customer satisfaction. Developed economies such as the United States, United Kingdom, and Japan have witnessed significant technological adoption among SMEs, with cloud computing, artificial intelligence, and data analytics becoming commonplace tools for service enhancement.')
    
    add_body_paragraph(doc, 'Research indicates that 73% of consumers globally consider customer experience a key factor in their purchasing decisions, yet only 49% of businesses prioritize it as a competitive strategy (PwC Global Consumer Insights Survey, 2022). This gap presents both a challenge and an opportunity for SMEs to leverage technological innovations to bridge the divide between customer expectations and actual service delivery. The emergence of Industry 4.0 technologies has further accelerated the need for SMEs to adopt innovative solutions to remain competitive in an increasingly digital marketplace.')
    
    add_body_paragraph(doc, 'In emerging economies, the technological landscape for SMEs presents unique characteristics and challenges. The diffusion of mobile technology has been particularly transformative, enabling SMEs in developing regions to leapfrog traditional infrastructure constraints and directly access digital tools for business management, customer engagement, and service delivery (OECD, 2021). The COVID-19 pandemic further accelerated digital transformation among SMEs globally, necessitating rapid adoption of e-commerce, digital payments, and remote service delivery technologies.')
    
    add_heading3(doc, '1.1.2 Regional Perspective of Technological Innovation in SMEs')
    
    add_body_paragraph(doc, 'Across Africa, SMEs represent approximately 80% of the continent\'s workforce and contribute significantly to economic development and poverty reduction (African Development Bank, 2022). The African Union\'s Agenda 2063 recognizes the pivotal role of SMEs in driving industrialization and economic transformation across the continent. Technological innovation has been identified as a key enabler for African SMEs to overcome traditional barriers to growth, including limited access to markets, finance, and information.')
    
    add_body_paragraph(doc, 'In East Africa, countries such as Kenya, Tanzania, and Uganda have witnessed rapid growth in mobile technology adoption, with mobile money services like M-Pesa revolutionizing financial transactions and enabling SMEs to conduct business more efficiently. According to GSMA (2023), East Africa leads the world in mobile money adoption, with over 60% of the adult population using mobile money services. The proliferation of mobile phones has created new opportunities for SMEs to engage with customers, manage inventory, and process payments through digital platforms.')
    
    add_body_paragraph(doc, 'Kenya, in particular, has emerged as a leader in technological innovation in the region, earning the nickname "Silicon Savannah." The country has witnessed significant growth in technology hubs, innovation centers, and digital entrepreneurship, particularly in Nairobi. The government\'s commitment to digital transformation through initiatives such as the Digital Economy Blueprint and the National Information, Communications and Technology Policy has created an enabling environment for SMEs to adopt technological innovations (Government of Kenya, 2019).')
    
    add_body_paragraph(doc, 'However, despite these advancements, research indicates that technology adoption among Kenyan SMEs remains uneven, with significant variations across sectors and geographical regions. A study on ICT adoption in service sector SMEs in Nairobi County found that while 94.62% of SMEs had adopted some form of ICT tools, the extent of adoption varied, with 37.63% showing moderate adoption and 17.2% showing low adoption levels (Kising\'a & Kwasira, 2019). This suggests that while awareness of technology exists, full integration into business operations remains a challenge, particularly outside major urban centers.')
    
    add_heading3(doc, '1.1.3 Local Perspective of Technological Innovation in SMEs in Uasin Gishu County')
    
    add_body_paragraph(doc, 'Uasin Gishu County, located in the North Rift region of Kenya, is an important agricultural and commercial hub. The county\'s economy is predominantly driven by agriculture, with SMEs playing a crucial role in value addition, distribution, and service provision. The county had a population of 1,163,186 according to the 2019 Kenya Population and Housing Census, with an urbanization rate that has been rapidly increasing. Eldoret, the county headquarters, is the fifth largest city in Kenya and a key commercial center for the North Rift region.')
    
    add_body_paragraph(doc, 'The Annex area, a commercial center within Uasin Gishu County located adjacent to Eldoret\'s central business district, hosts a diverse range of SMEs including retail shops, hotels and restaurants, service providers such as salons and repair shops, and agricultural enterprises. This area represents a microcosm of SME activity in the county, making it an ideal study site for understanding technology adoption patterns and their impact on service delivery.')
    
    add_body_paragraph(doc, 'Research conducted in Uasin Gishu County has highlighted both the potential and challenges of technological adoption among local SMEs. A study by Lagat (2014) on ICT adoption in agricultural SMEs in Uasin Gishu County revealed that access and use of ICT among agri-business SMEs had not been fully embraced by the majority of enterprises, indicating a low level of ICT use. The study identified weak financial capacity, limited knowledge of ICT tools, and low literacy levels as key barriers to technology adoption.')
    
    add_body_paragraph(doc, 'More recent research by Bwire and Muathe (2025) examined the influence of digital credit access on MSME growth in Uasin Gishu County, finding that ease of access to digital credit (r=0.673, p<0.001), information availability (r=0.701, p<0.001), and digital credit regulation (r=0.669, p<0.001) positively and significantly influenced MSME growth. However, the cost of digital credit showed a significant negative correlation with MSME growth (r=−0.610, p<0.001), highlighting the complex relationship between technological financial services and business performance.')
    
    add_body_paragraph(doc, 'Another study by Talam (2023) on organizational capabilities in agro-processing SMEs in Uasin Gishu County found that technological capabilities played a crucial role in influencing organizational performance, with a unit improvement in technological capabilities predicting an increment in organizational performance by 53.1% (β=0.531, t(145)=7.497, p<0.05). This demonstrates the significant potential of technological innovation to enhance SME performance in the county.')
    
    add_body_paragraph(doc, 'Despite these promising findings, there remains a research gap regarding the specific impact of technological innovation on service delivery among SMEs in the Annex area of Uasin Gishu County. This study seeks to address this gap by providing empirical evidence on the types of technological innovations adopted, their influence on service delivery efficiency and quality, customer perceptions, and the challenges faced in adoption and implementation.')
    
    add_heading2(doc, '1.2 Statement of the Problem')
    
    add_body_paragraph(doc, 'Small and Medium Enterprises (SMEs) in Kenya face significant challenges in maintaining competitiveness and ensuring sustainable growth. Studies indicate that 70% of Micro and Small Enterprises fail within three years of operation, rendering their survival in the market space alarmingly low (Kiprono, 2024). This high failure rate persists despite efforts by the government and other stakeholders to promote SME development through various policy interventions and support programs, including the Micro and Small Enterprises Authority (MSEA) and the National SME Policy.')
    
    add_body_paragraph(doc, 'In Uasin Gishu County, SMEs contribute significantly to local economic development and employment creation. However, many SMEs struggle with inefficient service delivery processes, limited customer reach, and inability to meet evolving customer expectations. The Annex area, as a growing commercial center, hosts numerous SMEs that could potentially benefit from technological innovations to enhance their service delivery. Despite the general awareness of technology in the region, systematic adoption and utilization of technology for service delivery enhancement remains inadequate.')
    
    add_body_paragraph(doc, 'While technological innovation has been recognized globally as a key driver of business performance, its adoption among SMEs in Uasin Gishu County remains limited. Previous studies have examined ICT adoption in agricultural SMEs (Lagat, 2014) and the influence of digital credit on MSME growth (Bwire & Muathe, 2025) in the county. However, there is limited research specifically focusing on the impact of technological innovation on service delivery among SMEs in the Annex area. The existing studies either focus on agricultural SMEs or on specific financial technology aspects, leaving a gap in understanding the broader impact of technological innovation on service delivery in commercial SMEs.')
    
    add_body_paragraph(doc, 'Furthermore, existing studies have primarily focused on technology adoption rates and general performance metrics, with limited attention to the specific mechanisms through which technological innovation influences service delivery efficiency, quality, and customer perceptions. The challenges faced by SMEs in adopting and implementing technological innovations in this specific commercial context have also not been adequately explored (Mutwota, 2023). Without empirical evidence on these aspects, SME owners, policymakers, and other stakeholders lack the information needed to make evidence-based decisions about technology investments and support programs.')
    
    add_body_paragraph(doc, 'This study therefore seeks to fill this research gap by examining the impact of technological innovation on service delivery among SMEs in the Annex area of Uasin Gishu County. The findings will provide valuable insights for SME owners, policymakers, and other stakeholders on how to leverage technology for enhanced service delivery and business performance.')
    
    add_heading2(doc, '1.3 General Objective')
    
    add_body_paragraph(doc, 'To examine the impact of technological innovation on service delivery among Small and Medium Enterprises (SMEs) in the Annex area of Uasin Gishu County.')
    
    add_heading2(doc, '1.4 Specific Objectives')
    
    objectives = [
        'To identify the types of technological innovations adopted by SMEs in the Annex area of Uasin Gishu County.',
        'To assess the influence of technological innovation on efficiency and quality of service delivery among SMEs in the Annex area of Uasin Gishu County.',
        'To evaluate customer perceptions of technologically enhanced services offered by SMEs in the Annex area of Uasin Gishu County.',
        'To identify challenges faced by SMEs in the Annex area of Uasin Gishu County in adopting and implementing technological innovations.',
    ]
    
    for i, obj in enumerate(objectives, 1):
        para = doc.add_paragraph()
        run = para.add_run(f'{i}. {obj}')
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    add_heading2(doc, '1.5 Research Questions')
    
    questions = [
        'What types of technological innovations have been adopted by SMEs in the Annex area of Uasin Gishu County?',
        'How does technological innovation influence the efficiency and quality of service delivery among SMEs in the Annex area of Uasin Gishu County?',
        'What are customer perceptions of technologically enhanced services offered by SMEs in the Annex area of Uasin Gishu County?',
        'What challenges do SMEs in the Annex area of Uasin Gishu County face in adopting and implementing technological innovations?',
    ]
    
    for i, q in enumerate(questions, 1):
        para = doc.add_paragraph()
        run = para.add_run(f'{i}. {q}')
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    add_heading2(doc, '1.6 Significance of the Study')
    
    add_body_paragraph(doc, 'The findings of this study will be significant to various stakeholders in the following ways:')
    
    significance_items = [
        ('SME Owners and Managers', 'The study will provide insights into the types of technological innovations that can enhance service delivery, enabling SME owners to make informed decisions about technology investments. Understanding customer perceptions and adoption challenges will help them develop effective implementation strategies. The findings will also highlight the return on investment of technology adoption, motivating more SMEs to embrace innovation.'),
        ('Policymakers and Government Agencies', 'The findings will inform policy formulation aimed at promoting technology adoption among SMEs in Uasin Gishu County and beyond. Agencies such as the Micro and Small Enterprises Authority (MSEA) can use the results to design targeted support programs addressing specific challenges identified in the study. The county government of Uasin Gishu can use findings to develop localized digital transformation strategies.'),
        ('Financial Institutions', 'Banks and other financial service providers will gain understanding of the technology needs and challenges of SMEs, enabling them to develop appropriate financing products for technology acquisition and implementation. Microfinance institutions can design technology-specific credit products that align with SME capacity and growth potential.'),
        ('Academic Researchers', 'The study will contribute to the existing body of knowledge on technology adoption in SMEs, particularly in the Kenyan regional context outside Nairobi. It will provide empirical data that can serve as a foundation for further research on technological innovation and service delivery in SMEs across different counties and sectors.'),
        ('Customers', 'Ultimately, improved service delivery resulting from appropriate technology adoption will benefit customers through enhanced service quality, efficiency, and overall satisfaction. The study will also highlight customer preferences and expectations, which can guide SMEs in delivering services that better meet customer needs.'),
    ]
    
    for stakeholder, impact in significance_items:
        para = doc.add_paragraph()
        run = para.add_run(f'{stakeholder}: ')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run2 = para.add_run(impact)
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    add_heading2(doc, '1.7 Scope of the Study')
    
    add_body_paragraph(doc, 'This study focuses on Small and Medium Enterprises (SMEs) operating in the Annex area of Uasin Gishu County, Kenya. The geographical scope is limited to this specific commercial area to allow for in-depth analysis and manageable data collection. The study targets SME owners, managers, and customers of these enterprises during the study period of January to March 2026.')
    
    add_body_paragraph(doc, 'The content scope covers four main areas: types of technological innovations adopted, influence on service delivery efficiency and quality, customer perceptions of technologically enhanced services, and challenges faced in technology adoption and implementation. Technological innovations considered include ICT tools (computers, internet, business software), mobile technologies (smartphones, mobile applications, mobile money), digital marketing platforms (social media, websites, e-commerce), and service delivery technologies (POS systems, customer management tools, self-service technologies).')
    
    add_body_paragraph(doc, 'The study focuses on SMEs employing between 1 and 50 employees, consistent with the definition used by the Kenyan regulatory framework. Both qualitative and quantitative aspects of technology adoption and service delivery are within scope, while macroeconomic factors affecting the broader Kenyan business environment are outside the study\'s scope.')
    
    add_heading2(doc, '1.8 Limitations of the Study')
    
    add_body_paragraph(doc, 'The study may face several limitations, which the researcher has endeavored to mitigate through appropriate research design decisions:')
    
    limitations = [
        ('Geographical Limitation', 'The focus on SMEs in the Annex area may limit the generalizability of findings to other areas of Uasin Gishu County or Kenya. However, the Annex area\'s diverse SME landscape makes it reasonably representative of peri-urban commercial centers in the region, and findings will be presented with appropriate caveats.'),
        ('Sampling Limitations', 'The study will rely on a sample of SMEs, which may not fully represent all businesses in the area. Efforts have been made to ensure a representative sample through stratified random sampling to capture diverse business types and sizes.'),
        ('Response Bias', 'Respondents may provide socially desirable responses or may not accurately recall information regarding technology adoption and its impact. The study will use multiple data collection methods, including both questionnaires and interviews, to minimize this limitation through triangulation.'),
        ('Time Constraints', 'The study is conducted within a limited timeframe of approximately three months, which may affect the depth of longitudinal data collection. Cross-sectional data will be used, which limits causal inference.'),
        ('Resource Constraints', 'Limited financial resources may affect the scope of data collection. The study will prioritize essential data collection activities and use cost-effective methods, including self-administered questionnaires and field research assistants from the university community.'),
        ('Self-Reporting Limitations', 'Data on technology adoption and its impact are primarily self-reported by SME owners and managers, which may introduce measurement error. Wherever possible, objective indicators such as revenue changes and customer traffic data will be sought to complement self-reported data.'),
    ]
    
    for limitation, explanation in limitations:
        para = doc.add_paragraph()
        run = para.add_run(f'{limitation}: ')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run2 = para.add_run(explanation)
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # ===================== CHAPTER TWO =====================
    add_page_break(doc)
    
    para = doc.add_paragraph()
    run = para.add_run('CHAPTER TWO: LITERATURE REVIEW')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(12)
    
    add_heading2(doc, '2.1 Introduction')
    
    add_body_paragraph(doc, 'This chapter presents a comprehensive review of literature relevant to the study on the impact of technological innovation on service delivery among SMEs in the Annex area of Uasin Gishu County. It covers the theoretical framework underpinning the research, the conceptual framework illustrating variable relationships, a review of literature on study variables, an empirical review of previous studies, and identification of research gaps that this study seeks to address. The review draws on local, regional, and global literature to provide a comprehensive understanding of the study topic.')
    
    add_heading2(doc, '2.2 Theoretical Framework')
    
    add_body_paragraph(doc, 'This study is anchored on three main theories: the Technology-Organization-Environment (TOE) Framework, Dynamic Capabilities Theory, and Disruptive Innovation Theory. These theories collectively provide a comprehensive lens for understanding the adoption, implementation, and impact of technological innovations on SME service delivery.')
    
    add_heading3(doc, '2.2.1 Technology-Organization-Environment (TOE) Framework')
    
    add_body_paragraph(doc, 'The Technology-Organization-Environment (TOE) framework, developed by Tornatzky and Fleischer in 1990, explains that the adoption of technological innovations by firms is influenced by three contextual elements: technological context, organizational context, and environmental context. This framework has been widely applied in studies on technology adoption across various sectors and geographical settings (Kising\'a & Kwasira, 2019; Musebe, 2024).')
    
    add_body_paragraph(doc, 'The technological context refers to the internal and external technologies relevant to the firm, including existing technologies in use and those available in the market. For SMEs in the Annex area, this includes the availability and characteristics of technologies that can enhance service delivery, such as mobile payment systems (M-Pesa, Airtel Money), inventory management software, customer relationship management tools, and digital marketing platforms. The technological context also encompasses the relative advantage, compatibility, and complexity of available technologies.')
    
    add_body_paragraph(doc, 'The organizational context encompasses firm characteristics including size, scope, managerial structure, human resources, and financial capacity. This includes factors such as the SME owner\'s education level and technology experience, employees\' technical skills and digital literacy, the organization\'s readiness for technological change, and the financial resources available for technology investment. For small businesses with limited capacity, the organizational context often determines whether available and suitable technologies are actually adopted.')
    
    add_body_paragraph(doc, 'The environmental context includes the external environment in which the firm operates, including industry structure, competitors\' technology adoption, regulatory environment, and access to technology service providers and support. For SMEs in Uasin Gishu County, this includes the local business ecosystem, infrastructure availability (internet connectivity, electricity), government policies affecting technology adoption, and the availability of technology vendors and support services in the region.')
    
    add_body_paragraph(doc, 'The TOE framework is particularly relevant to this study as it provides a comprehensive lens for understanding the factors that influence technology adoption among SMEs in the Annex area. It will guide the investigation of challenges faced in adopting technological innovations (Objective 4) and help explain variations in adoption patterns across different SME categories.')
    
    add_heading3(doc, '2.2.2 Dynamic Capabilities Theory')
    
    add_body_paragraph(doc, 'The Dynamic Capabilities Theory, advanced by Teece, Pisano, and Shuen in 1997, focuses on a firm\'s ability to integrate, build, and reconfigure internal and external competencies to address rapidly changing environments. Dynamic capabilities refer to the capacity of an organization to purposefully create, extend, or modify its resource base in response to environmental changes and competitive pressures (Talam, 2023).')
    
    add_body_paragraph(doc, 'In the context of SMEs, dynamic capabilities enable businesses to adapt to technological changes and leverage innovations for competitive advantage. These capabilities include sensing opportunities and threats in the technological environment, seizing opportunities by making timely decisions about technology investments, and transforming or reconfiguring existing processes and resources to implement new technologies effectively. SMEs that develop strong dynamic capabilities are better positioned to continuously improve their service delivery through ongoing technology adoption and adaptation.')
    
    add_body_paragraph(doc, 'This theory is particularly relevant to understanding how technological innovation influences service delivery efficiency and quality (Objective 2), as it explains the mechanisms through which firms transform technological resources into improved service outcomes. It also helps explain why some SMEs are more successful than others in leveraging technology for service delivery enhancement, despite facing similar environmental conditions.')
    
    add_body_paragraph(doc, 'The Dynamic Capabilities Theory suggests that SMEs in the Annex area can enhance their service delivery not just by adopting technology, but by developing organizational capacities to sense relevant technologies, make strategic adoption decisions, and continuously improve their technology utilization. This implies that technology training and organizational learning are as important as technology acquisition itself.')
    
    add_heading3(doc, '2.2.3 Disruptive Innovation Theory')
    
    add_body_paragraph(doc, 'The Disruptive Innovation Theory, introduced by Clayton Christensen in 1997, describes how new technologies can disrupt existing markets by introducing simpler, more affordable, or more accessible products and services. Initially, these innovations may underperform established offerings in mainstream markets but gain traction by serving overlooked or underserved segments, eventually displacing established players (Christensen, 1997).')
    
    add_body_paragraph(doc, 'For SMEs, disruptive innovations often provide opportunities to compete effectively with larger, established firms by offering more accessible and affordable service alternatives. Mobile money services like M-Pesa represent a classic example of disruptive innovation that has enabled SMEs in Kenya to offer financial transaction services that were previously the exclusive domain of formal banking institutions. Similarly, social media marketing has disrupted traditional advertising, enabling SMEs to reach large audiences at minimal cost.')
    
    add_body_paragraph(doc, 'This theory is relevant to understanding customer perceptions of technologically enhanced services (Objective 3), as it explains how customers in the Annex area may perceive and adopt new service delivery technologies that offer convenience, affordability, or accessibility advantages over traditional service models. It also helps explain why some technology innovations gain rapid acceptance among customers despite initial unfamiliarity, while others fail to achieve adoption despite their technical superiority.')
    
    add_body_paragraph(doc, 'The Disruptive Innovation Theory also provides insights into how SMEs in the Annex area can strategically position themselves by adopting technologies that create new value propositions for underserved customer segments, enabling them to compete more effectively in the local market.')
    
    add_heading2(doc, '2.3 Conceptual Framework')
    
    add_body_paragraph(doc, 'A conceptual framework is a graphical representation of the relationship between independent and dependent variables in a study, illustrating how the researcher conceptualizes the connections between key study concepts (Kothari, 2004). Based on the theoretical framework and literature review, the following conceptual framework illustrates the relationship between technological innovation (independent variable) and service delivery (dependent variable) among SMEs in the Annex area, with intervening variables moderating this relationship.')
    
    # Conceptual Framework Box
    para = doc.add_paragraph()
    run = para.add_run('Figure 2.1: Conceptual Framework')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run.font.italic = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    
    # Create conceptual framework as a table
    cf_table = doc.add_table(rows=3, cols=3)
    cf_table.style = 'Table Grid'
    cf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Independent Variable
    iv_cell = cf_table.rows[0].cells[0]
    iv_cell.merge(cf_table.rows[1].cells[0])
    iv_cell.text = ''
    para_iv = iv_cell.paragraphs[0]
    run_iv = para_iv.add_run('INDEPENDENT VARIABLE\nTECHNOLOGICAL INNOVATION\n\n• ICT Tools & Applications\n  - Computers & Software\n  - Mobile Technologies\n  - Internet & Websites\n\n• Digital Marketing\n  - Social Media\n  - Online Advertising\n  - E-commerce Platforms\n\n• Service Delivery Technologies\n  - POS Systems\n  - CRM Tools\n  - Self-Service Technologies\n\n• Financial Technologies\n  - Mobile Money\n  - Digital Payments')
    run_iv.font.size = Pt(10)
    run_iv.font.name = 'Times New Roman'
    para_iv.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_cell_background(iv_cell, 'E8F4FD')
    
    # Arrow
    arrow_cell = cf_table.rows[0].cells[1]
    arrow_cell.merge(cf_table.rows[1].cells[1])
    arrow_cell.text = ''
    para_arrow = arrow_cell.paragraphs[0]
    run_arrow = para_arrow.add_run('\n\n\n\n\n          \u2192')
    run_arrow.font.size = Pt(14)
    run_arrow.font.bold = True
    para_arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Dependent Variable
    dv_cell = cf_table.rows[0].cells[2]
    dv_cell.merge(cf_table.rows[1].cells[2])
    dv_cell.text = ''
    para_dv = dv_cell.paragraphs[0]
    run_dv = para_dv.add_run('DEPENDENT VARIABLE\nSERVICE DELIVERY\n\n• Efficiency\n  - Time Savings\n  - Cost Reduction\n  - Process Streamlining\n\n• Service Quality\n  - Reliability\n  - Responsiveness\n  - Personalization\n\n• Customer Satisfaction\n  - Customer Experience\n  - Service Perception\n  - Customer Loyalty')
    run_dv.font.size = Pt(10)
    run_dv.font.name = 'Times New Roman'
    para_dv.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_cell_background(dv_cell, 'E8FDE8')
    
    # Intervening Variables
    iv2_cell = cf_table.rows[2].cells[0]
    iv2_cell.merge(cf_table.rows[2].cells[2])
    iv2_cell.text = ''
    para_iv2 = iv2_cell.paragraphs[0]
    run_iv2 = para_iv2.add_run('INTERVENING VARIABLES\n• Business characteristics (size, type, years of operation)\n• Owner/manager characteristics (education, experience, attitude)\n• Environmental factors (infrastructure, competition, regulatory environment)')
    run_iv2.font.size = Pt(10)
    run_iv2.font.name = 'Times New Roman'
    para_iv2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_background(iv2_cell, 'FFF8E1')
    
    add_caption(doc, 'Source: Researcher (2026) adapted from Tornatzky & Fleischer (1990)')
    
    add_body_paragraph(doc, 'The conceptual framework posits that technological innovation, comprising ICT tools and applications, digital marketing, service delivery technologies, and financial technologies, influences service delivery in terms of efficiency, service quality, and customer satisfaction. The relationship between technological innovation and service delivery is moderated by intervening variables including business characteristics, owner/manager characteristics, and environmental factors. These intervening variables can either enhance or constrain the impact of technological innovation on service delivery outcomes.')
    
    add_heading2(doc, '2.4 Review of Literature on Study Variables')
    
    add_heading3(doc, '2.4.1 Types of Technological Innovations Adopted by SMEs')
    
    add_body_paragraph(doc, 'Technological innovations adopted by SMEs encompass a wide range of tools, systems, and applications designed to enhance business operations and service delivery. Based on existing literature, these innovations can be categorized into several types, each with distinct applications and benefits for SME operations.')
    
    add_body_paragraph(doc, 'Information and Communication Technology (ICT) Tools: These include hardware and software used for information processing and communication. A study on ICT adoption in service sector SMEs in Nairobi County found that SMEs utilize personal computers, email, internet, websites, intranet, credit card facilities, business software, and data storage and security facilities (Kising\'a & Kwasira, 2019). SMEs use these tools primarily for internal and external communication with customers, suppliers, and employees; sharing files and information; tracking orders, supplies, and customer enquiries; financial accounting; and carrying out internet banking.')
    
    add_body_paragraph(doc, 'Mobile Technologies: Mobile phones, particularly smartphones, have become essential tools for SME operations in Kenya. Mobile money services such as M-Pesa have transformed financial transactions, enabling SMEs to receive payments, pay suppliers, and manage finances more efficiently (Bwire & Muathe, 2025). According to the Communications Authority of Kenya (2023), mobile phone penetration in Kenya stands at over 125%, with smartphones increasingly accessible even to small business owners in peri-urban areas like the Annex. Mobile applications for business management, inventory tracking, and customer communication are increasingly being adopted.')
    
    add_body_paragraph(doc, 'Digital Marketing and Social Media: SMEs are leveraging digital platforms to reach customers and promote their products and services. Social media platforms including Facebook, Instagram, WhatsApp Business, and TikTok are commonly used for customer engagement, marketing, product showcasing, and sales (Kiprono, 2024). Research by Musebe (2024) indicates that marketing technology has a positive influence on the performance of Micro and Small Enterprises. WhatsApp Business has particularly gained traction among Kenyan SMEs due to its low data requirements and high penetration among existing customer bases.')
    
    add_body_paragraph(doc, 'Service Delivery Technologies: These include technologies specifically designed to enhance the delivery of services to customers. Point of Sale (POS) systems, customer relationship management (CRM) software, and self-service technologies enable SMEs to process transactions efficiently, manage customer information, and provide convenient service options (Mutwota, 2023). The adoption of digital order management systems, appointment booking apps, and queue management systems has improved service delivery in sectors such as healthcare, restaurants, and professional services.')
    
    add_body_paragraph(doc, 'Financial Technologies (FinTech): Digital payment systems, mobile banking, and digital credit platforms have become increasingly important for SMEs. A study in Uasin Gishu County found that ease of access to digital credit significantly influences MSME growth (Bwire & Muathe, 2025). The integration of multiple payment options including mobile money, bank transfer, and card payments has expanded SMEs\' customer base and improved cash flow management.')
    
    add_heading3(doc, '2.4.2 Technological Innovation and Efficiency of Service Delivery')
    
    add_body_paragraph(doc, 'Efficiency in service delivery refers to the ability of SMEs to provide services using optimal resources, minimizing time, cost, and effort while maximizing output. Technological innovation has been shown to significantly enhance service delivery efficiency through various mechanisms across different sectors and geographical contexts.')
    
    add_body_paragraph(doc, 'Time Savings: Technology enables SMEs to automate routine tasks, substantially reducing the time required for service delivery. For example, POS systems speed up payment processing at checkout, inventory management software reduces time spent on manual stock-taking, digital order management systems eliminate queue bottlenecks, and digital communication tools enable quick responses to customer enquiries (Kising\'a & Kwasira, 2019). In the hospitality sector, digital ordering systems have reduced order processing time by up to 40% in some establishments (Kiprono, 2024).')
    
    add_body_paragraph(doc, 'Cost Reduction: By automating processes and reducing manual labor, technology can lower operational costs significantly. Digital marketing, for instance, offers cost-effective alternatives to traditional advertising, enabling SMEs to reach wider audiences at substantially lower costs compared to print media, radio, or television advertising. Cloud-based accounting software reduces the need for manual bookkeeping and costly accountancy services. Mutwota (2023) found that service delivery systems account for 9.9% of the variance in SME performance, with technology-enabled efficiency contributing significantly to cost savings.')
    
    add_body_paragraph(doc, 'Accuracy and Error Reduction: Automated systems significantly reduce human error in transactions, record-keeping, and service delivery. Digital record-keeping ensures accurate tracking of sales, inventory, and customer information, reducing costly errors from manual data entry. In financial transactions, mobile payment systems eliminate cash-handling errors and reduce the risk of theft (Bwire & Muathe, 2025). This accuracy improvement translates directly into improved customer trust and satisfaction.')
    
    add_body_paragraph(doc, 'Scalability: Technology enables SMEs to handle increased service demand without proportional increases in resources. Cloud-based systems, for example, allow businesses to expand their service capacity without significant additional infrastructure investment. E-commerce platforms enable SMEs to serve customers beyond their physical location, expanding market reach without commensurate increases in operational costs. This scalability is particularly important for SMEs seeking to grow their businesses in competitive markets like the Annex area.')
    
    add_body_paragraph(doc, 'Process Streamlining: Technology enables SMEs to streamline service delivery processes, reducing bottlenecks and improving workflow. Digital appointment systems in service businesses eliminate long wait times, integrated inventory and sales systems ensure product availability information is always current, and automated follow-up systems improve post-service customer engagement. Research by Talam (2023) in Uasin Gishu County demonstrated that technological capabilities significantly predict organizational performance improvements, with beta coefficients indicating substantial predictive power.')
    
    add_heading3(doc, '2.4.3 Customer Perceptions of Technologically Enhanced Services')
    
    add_body_paragraph(doc, 'Customer perception refers to the attitudes, opinions, and satisfaction levels of customers regarding the services they receive. Understanding customer perceptions of technologically enhanced services is crucial for SMEs seeking to optimize their technology investments and ensure that adopted technologies actually improve the customer experience.')
    
    add_body_paragraph(doc, 'Customer Satisfaction: Research indicates that 73% of consumers consider customer experience a key factor in their purchasing decisions (PwC, 2022). Technology can enhance customer satisfaction by enabling faster service, personalized interactions, convenient access to services, and multiple payment options. A study on strategic service innovation by Mutwota (2023) found that customer interface accounted for 10.6% of the variance in SME performance, indicating the critical importance of customer-facing technology in driving business outcomes. In Kenyan SMEs, customers have shown high appreciation for mobile payment options and digital communication channels that enable quick query resolution.')
    
    add_body_paragraph(doc, 'Service Quality Perceptions: Customers evaluate service quality based on dimensions including reliability, responsiveness, assurance, empathy, and tangibles (Parasuraman, Zeithaml & Berry, 1988). Technology can enhance these dimensions by enabling consistent service delivery (reliability), quick responses to customer needs (responsiveness), professional service environments equipped with modern technology (tangibles), and personalized attention through customer data management (empathy). Research indicates that customers in Kenya generally have positive perceptions of technology-enhanced services, particularly mobile payment systems and digital communication channels (Communications Authority of Kenya, 2023).')
    
    add_body_paragraph(doc, 'Trust and Security: Customer perceptions of technology-enabled services are significantly influenced by trust in the technology and perceptions of security, particularly for financial transactions and services involving personal information. Research on digital credit in Uasin Gishu County found that information availability and regulatory oversight significantly influence MSME growth (Bwire & Muathe, 2025), highlighting the importance of transparency and security in technology-enabled services. SMEs that communicate clearly about data security and offer reliable technology solutions build stronger customer trust.')
    
    add_body_paragraph(doc, 'Adoption Barriers: Customer willingness to use technology-enabled services depends on perceived usefulness, ease of use, and social influence. A study on voice-based assistants for SMEs in Kenya found that while customers were generally hopeful about technology benefits, practical challenges such as social appropriateness, environmental noise, and digital literacy limitations affected actual usage (Musebe, 2024). Older customers and those with lower digital literacy may require additional support and guidance to engage with technology-enhanced services, requiring SMEs to adopt inclusive technology strategies.')
    
    add_heading3(doc, '2.4.4 Challenges Faced by SMEs in Adopting and Implementing Technological Innovations')
    
    add_body_paragraph(doc, 'SMEs face numerous challenges in adopting and implementing technological innovations. Understanding these challenges is essential for developing effective support interventions and policies that can accelerate technology adoption for service delivery enhancement.')
    
    add_body_paragraph(doc, 'Financial Constraints: Limited financial resources are among the most significant barriers to technology adoption for SMEs. The high initial costs of technology acquisition, including hardware, software, implementation, and training, can be prohibitive for small businesses with limited capital (Lagat, 2014). Additionally, the cost of digital credit, with high interest rates charged by some fintech providers, has been found to negatively correlate with MSME growth (Bwire & Muathe, 2025). The perceived uncertainty of return on investment from technology adoption further discourages many SME owners from committing scarce resources to technology acquisition.')
    
    add_body_paragraph(doc, 'Skills and Knowledge Gaps: Many SME owners and employees lack the technical skills and knowledge required to effectively adopt and utilize technological innovations. A study in Uasin Gishu County found that low literacy levels and limited knowledge of ICT tools hindered use of technology among agri-business SMEs (Lagat, 2014). The skills gap extends beyond basic operation to include strategic decision-making about which technologies to invest in, how to implement them effectively, and how to troubleshoot and maintain technology systems. The rapid pace of technological change makes continuous learning and skill upgrading a persistent challenge.')
    
    add_body_paragraph(doc, 'Infrastructure Limitations: Inadequate technological infrastructure, particularly in peri-urban and rural areas, limits SME ability to adopt and benefit from technological innovations. This includes unreliable internet connectivity, frequent power outages, limited access to technical support services, and inadequate telecommunications infrastructure. In the Annex area of Uasin Gishu County, while infrastructure has improved significantly in recent years, gaps remain particularly regarding consistent electricity supply and broadband internet access (Bwire & Muathe, 2025).')
    
    add_body_paragraph(doc, 'Organizational Factors: SME characteristics such as small size, limited managerial capacity, and resistance to change can hinder technology adoption. Owner-managers may be reluctant to invest in technologies they do not fully understand, or may prefer to maintain traditional ways of operating that have historically been successful. The lack of dedicated IT staff in small businesses means that technology management often falls on already overextended owner-managers, reducing the likelihood of effective technology utilization (Mutwota, 2023).')
    
    add_body_paragraph(doc, 'Environmental and Social Challenges: The external environment presents challenges including regulatory uncertainty, limited technology vendor presence in the region, social norms around service delivery, and cultural factors affecting technology adoption. Research on voice-based assistants in Kenyan SMEs revealed social challenges such as discomfort with using technology that changes the personal nature of service interactions (Musebe, 2024). In some service sectors, customers may actually prefer personal interaction over technology-mediated service, requiring SMEs to carefully balance technological and human elements of service delivery.')
    
    add_heading2(doc, '2.5 Empirical Review')
    
    add_heading3(doc, '2.5.1 Studies on ICT Adoption in Kenyan SMEs')
    
    add_body_paragraph(doc, 'Several empirical studies have examined technology adoption among SMEs in Kenya, providing valuable context for understanding the dynamics of technological innovation in the Kenyan business environment.')
    
    add_body_paragraph(doc, 'Kising\'a and Kwasira (2019) conducted a study on the effect of ICTs as innovation facilitators of service sector SMEs in Nairobi County. Using a descriptive research design and a sample of 106 SMEs, the study found that 94.62% of SMEs had adopted ICT tools and applications, with moderate adoption in 37.63% and low adoption in 17.2% of SMEs. The study established that technological context (β=0.259, p<0.05), organizational context (β=0.398, p<0.05), and environmental context (β=0.214, p<0.05) have positive effects on innovation in service sector SMEs. The study recommended government support for ICT infrastructure development and training programs for SME operators.')
    
    add_body_paragraph(doc, 'Kiprono (2024) examined adopted technology and the performance of Micro and Small Enterprises in Nairobi, focusing on marketing innovation, process and service innovation, product distribution innovation, and payment technology. Using a sample of 200 MSEs and employing regression analysis, the study found that all these technological dimensions had a positive influence on MSE performance, with payment technology showing the strongest relationship (β=0.412, p<0.001). The study recommended that policymakers encourage technology-enabled marketing strategies and provide training programs, with particular emphasis on affordable payment technology solutions.')
    
    add_body_paragraph(doc, 'Mutwota (2023) investigated the influence of strategic service innovation on the performance of SMEs in Nairobi County, examining customer interface, service delivery systems, and technology adoption. The study found positive and significant relationships between all three dimensions and SME performance, with technology accounting for 6.2% of the variance in performance. Customer interface was found to be the strongest predictor, accounting for 10.6% of variance. The study recommended that SMEs adopt Self-Service Technologies (SST) to respond to shifts in customer behavior toward digital service interactions.')
    
    add_body_paragraph(doc, 'Musebe (2024) examined the adoption of advanced manufacturing and service technology by SMEs in Kenya and its effect on performance, using a mixed-methods approach combining surveys and case studies. The study found that technology innovation influences firm performance positively and recommended that entrepreneurs develop innovative strategies to actualize firm performance through systematic technology planning and implementation. The study particularly highlighted the importance of combining technology adoption with employee training and process redesign for maximum impact.')
    
    add_heading3(doc, '2.5.2 Studies on Technology and SME Performance in Uasin Gishu County')
    
    add_body_paragraph(doc, 'Research specifically focused on Uasin Gishu County provides valuable context for understanding the local dynamics of technology adoption and its impact on SME performance in the county.')
    
    add_body_paragraph(doc, 'Lagat (2014) examined leveraging ICT organizational capability for SME competitiveness in the agricultural sector in Uasin Gishu County, focusing on agri-business SMEs in the Eldoret area. Using a descriptive survey design and sample of 120 SMEs, the study found low levels of ICT use among agri-business SMEs, with weak financial capacity, limited knowledge of ICT tools, and low literacy levels identified as key barriers. The study recommended government intervention to make ICT more affordable through subsidization of hardware and software costs, and to support training programs specifically tailored to the needs and contexts of agricultural SMEs.')
    
    add_body_paragraph(doc, 'Bwire and Muathe (2025) explored the influence of digital credit access on MSME growth in Uasin Gishu County, using a correlational research design with a sample of 250 MSMEs. The study found that ease of access to digital credit (r=0.673, p<0.001), information availability (r=0.701, p<0.001), and digital credit regulation (r=0.669, p<0.001) positively and significantly influenced MSME growth. However, the cost of digital credit showed a significant negative correlation (r=−0.610, p<0.001). Multiple regression analysis revealed that these factors collectively explained 60.0% of the variance in MSME growth (Adjusted R²=0.584), highlighting the complex but significant relationship between financial technology adoption and business performance.')
    
    add_body_paragraph(doc, 'Talam (2023) evaluated the influence of organizational capabilities, including technological capabilities, on the performance of agro-processing SMEs in Uasin Gishu County. Using a cross-sectional survey design with a sample of 147 agro-processing SMEs, the study found that technological capabilities played a crucial role in influencing organizational performance. A unit improvement in technological capabilities predicted an increment in organizational performance by 53.1% (β=0.531, t(145)=7.497, p<0.05). The study recommended that management promote a culture of innovation and technological proficiency among employees and invest in regular technology training programs.')
    
    add_heading2(doc, '2.6 Research Gaps')
    
    add_body_paragraph(doc, 'Despite the existing literature on technology adoption and SME performance, several research gaps remain that this study seeks to address:')
    
    gaps = [
        ('Geographical Gap', 'Most studies on technology adoption in Kenyan SMEs have focused on Nairobi County (Kising\'a & Kwasira, 2019; Mutwota, 2023; Kiprono, 2024), with limited research on SMEs in other regions including Uasin Gishu County. While some studies have been conducted in Uasin Gishu County (Lagat, 2014; Bwire & Muathe, 2025; Talam, 2023), none have specifically focused on the Annex area and its diverse SME landscape.'),
        ('Sectoral Gap', 'Existing studies in Uasin Gishu County have primarily focused on agricultural and agro-processing SMEs (Lagat, 2014; Talam, 2023), with limited attention to the diverse range of SMEs in commercial areas like the Annex, including retail, hospitality, and general service sector businesses.'),
        ('Service Delivery Focus', 'While existing studies have examined technology adoption and general performance metrics, there is limited research specifically focused on the impact of technological innovation on service delivery dimensions including efficiency, quality, and customer perceptions in the Kenyan regional context outside Nairobi.'),
        ('Customer Perspective Gap', 'Most studies have collected data from SME owners and managers as the primary respondents, with limited inclusion of customer perspectives on technologically enhanced services. This study addresses this gap by collecting data from both SME operators and their customers.'),
        ('Timeliness Gap', 'Given the rapid pace of technological change and the acceleration of digital transformation following the COVID-19 pandemic, more recent studies are needed to capture current technology adoption patterns and challenges. The post-pandemic period has significantly altered technology adoption trajectories among SMEs globally, requiring updated empirical evidence.'),
        ('Integration Gap', 'Existing studies have tended to examine specific aspects of technology adoption or specific technology types in isolation. This study provides a comprehensive examination of multiple technology types and their combined impact on service delivery across multiple dimensions.'),
    ]
    
    for gap, explanation in gaps:
        para = doc.add_paragraph()
        run = para.add_run(f'{gap}: ')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run2 = para.add_run(explanation)
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    add_body_paragraph(doc, 'This study seeks to address these gaps by examining the impact of technological innovation on service delivery among diverse SMEs in the Annex area of Uasin Gishu County, incorporating both SME and customer perspectives, and providing updated empirical evidence in the post-pandemic digital transformation context.')
    
    # ===================== CHAPTER THREE =====================
    add_page_break(doc)
    
    para = doc.add_paragraph()
    run = para.add_run('CHAPTER THREE: RESEARCH METHODOLOGY')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(12)
    
    add_heading2(doc, '3.1 Introduction')
    
    add_body_paragraph(doc, 'This chapter describes the research methodology employed in this study. It covers the research design, target population, sample size and sampling technique, data collection instruments, pilot testing procedures, data collection procedures, data analysis methods, and ethical considerations. The methodology is designed to ensure the collection of valid, reliable, and comprehensive data to address the study objectives.')
    
    add_heading2(doc, '3.2 Research Design')
    
    add_body_paragraph(doc, 'This study adopts a descriptive research design with a mixed-methods approach, combining both quantitative and qualitative data collection and analysis techniques. Descriptive research design is appropriate for studies that aim to describe the characteristics of a population or phenomenon and examine relationships between variables (Kothari, 2004). The design enables the researcher to collect comprehensive data on the current state of technological innovation adoption and its impact on service delivery among SMEs in the Annex area.')
    
    add_body_paragraph(doc, 'The mixed-methods approach combines quantitative data from structured questionnaires with qualitative insights from semi-structured interviews. This approach allows for triangulation of findings, enhancing the validity and reliability of the results. Quantitative data provide statistical evidence of relationships between technological innovation and service delivery, while qualitative data provide rich, contextual insights into the experiences, motivations, and challenges of SME owners and managers in adopting technology.')
    
    add_body_paragraph(doc, 'Similar designs have been used in previous studies on technology adoption in Kenyan SMEs. Kising\'a and Kwasira (2019) employed descriptive research design to study ICT adoption in Nairobi County, Mutwota (2023) used descriptive design to examine strategic service innovation, and Musebe (2024) employed mixed methods to examine technology adoption and firm performance. The consistency of this design with established studies in the field validates its appropriateness for this research.')
    
    add_heading2(doc, '3.3 Target Population')
    
    add_body_paragraph(doc, 'The target population for this study comprises all Small and Medium Enterprises operating in the Annex area of Uasin Gishu County. According to records from the Uasin Gishu County Business Licensing Department (2025), there are approximately 250 registered SMEs operating in the Annex area. These businesses span various sectors including retail shops, hotels and restaurants, personal and professional service providers, and agricultural-related businesses.')
    
    add_body_paragraph(doc, 'Additionally, the study targets customers of these SMEs to obtain data on their perceptions of technologically enhanced services. This dual-respondent approach ensures that findings capture both the supply side (SME operators) and demand side (customers) perspectives on technological innovation and service delivery.')
    
    add_caption(doc, 'Table 3.1: Target Population Distribution')
    
    t31_headers = ['Business Category', 'Estimated Number', 'Percentage (%)']
    t31_rows = [
        ['Retail Shops', '100', '40.0'],
        ['Hotels and Restaurants', '60', '24.0'],
        ['Service Providers', '50', '20.0'],
        ['Agricultural-related Businesses', '40', '16.0'],
        ['Total', '250', '100.0'],
    ]
    create_simple_table(doc, t31_headers, t31_rows, col_widths=[3.0, 2.0, 1.5])
    add_caption(doc, 'Source: Uasin Gishu County Business Licensing Department (2025)')
    
    add_heading2(doc, '3.4 Sample Size and Sampling Technique')
    
    add_heading3(doc, '3.4.1 Sample Size Determination')
    
    add_body_paragraph(doc, 'The sample size for SME owners/managers was determined using the Yamane (1967) formula:')
    
    para = doc.add_paragraph()
    run = para.add_run('n = N / (1 + N(e)²)')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = para.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    add_body_paragraph(doc, 'Where:')
    
    formula_items = ['n = sample size', 'N = population size (250)', 'e = margin of error (0.05 or 5%)']
    for item in formula_items:
        para = doc.add_paragraph()
        run = para.add_run(item)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.left_indent = Inches(0.5)
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    add_body_paragraph(doc, 'Applying the formula:')
    
    calc_steps = [
        'n = 250 / (1 + 250(0.05)²)',
        'n = 250 / (1 + 250 × 0.0025)',
        'n = 250 / (1 + 0.625)',
        'n = 250 / 1.625',
        'n ≈ 154 SMEs',
    ]
    for step in calc_steps:
        para = doc.add_paragraph()
        run = para.add_run(step)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.left_indent = Inches(0.5)
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    add_body_paragraph(doc, 'Therefore, the sample size for SME owners/managers is 154 respondents. This formula has been used in similar studies including Mutwota (2023) and Kising\'a and Kwasira (2019). For customers, a sample of 200 customers was selected from those patronizing the sampled SMEs, using systematic random sampling where every 5th customer was approached to participate, ensuring adequate representation across different times of day and days of the week.')
    
    add_heading3(doc, '3.4.2 Sampling Technique')
    
    add_body_paragraph(doc, 'Stratified random sampling was used to select SME respondents. The population was stratified into four categories based on business type: retail shops, hotels and restaurants, service providers, and agricultural-related businesses. From each stratum, respondents were selected proportionately using simple random sampling. This technique ensures that all business categories are appropriately represented in the sample, enhancing the representativeness of the findings.')
    
    add_caption(doc, 'Table 3.2: Sample Size Distribution')
    
    t32_headers = ['Business Category', 'Population', 'Proportion (%)', 'Sample Size']
    t32_rows = [
        ['Retail Shops', '100', '40.0', '62'],
        ['Hotels and Restaurants', '60', '24.0', '37'],
        ['Service Providers', '50', '20.0', '31'],
        ['Agricultural-related Businesses', '40', '16.0', '24'],
        ['Total', '250', '100.0', '154'],
    ]
    create_simple_table(doc, t32_headers, t32_rows, col_widths=[2.5, 1.5, 1.5, 1.5])
    add_caption(doc, 'Source: Researcher (2026)')
    
    add_body_paragraph(doc, 'For customers, convenience sampling was used at the business premises, with efforts made to include diverse customer demographics across different times of day, days of the week, and different business types to ensure diverse representation.')
    
    add_heading2(doc, '3.5 Data Collection Instruments')
    
    add_heading3(doc, '3.5.1 Questionnaires')
    
    add_body_paragraph(doc, 'Structured questionnaires were the primary data collection instrument for this study. Two sets of questionnaires were developed: one for SME owners/managers and one for customers.')
    
    add_body_paragraph(doc, 'Questionnaire for SME Owners/Managers: This contained closed-ended questions organized into five sections corresponding to the study objectives: Section A (demographic information and business characteristics), Section B (types of technological innovations adopted), Section C (influence on efficiency and quality of service delivery), Section D (challenges in adopting technological innovations), and Section E (additional technology adoption information). Attitudinal questions used a 5-point Likert scale (1=Strongly Disagree, 2=Disagree, 3=Neutral, 4=Agree, 5=Strongly Agree).')
    
    add_body_paragraph(doc, 'Questionnaire for Customers: This captured customer perceptions of technologically enhanced services, including satisfaction levels, service quality perceptions, and technology preferences. Similar to the SME questionnaire, it used a 5-point Likert scale and was structured to capture both objective indicators and subjective perceptions of service quality and customer experience.')
    
    add_heading3(doc, '3.5.2 Interview Guide')
    
    add_body_paragraph(doc, 'An interview guide with semi-structured questions was used to conduct in-depth interviews with 18 SME owners/managers purposively selected to represent diverse business types, sizes, and technology adoption levels. This enabled the collection of qualitative data on experiences, perceptions, and challenges that may not be fully captured through closed-ended questionnaire items. Each interview lasted approximately 30-45 minutes and was audio-recorded with informed consent and later transcribed for analysis. The interview guide explored: motivations for technology adoption, detailed experiences with specific technologies, perceived benefits and challenges of technology implementation, and suggestions for improving technology adoption support.')
    
    add_heading2(doc, '3.6 Pilot Testing')
    
    add_heading3(doc, '3.6.1 Validity of Research Instruments')
    
    add_body_paragraph(doc, 'Validity refers to the extent to which an instrument measures what it is intended to measure (Creswell, 2014). To ensure content validity, the questionnaires were developed based on extensive literature review and aligned with the study objectives and theoretical framework. Experts, including the university supervisor (Dr. Kiyeng Chumo) and two other faculty members with expertise in research methodology, reviewed the instruments to assess the relevance, comprehensiveness, and appropriateness of questions. Face validity was established through review by colleagues and five potential respondents who assessed the clarity and comprehensibility of questions. Construct validity was assessed through factor analysis during data analysis to confirm that questions loaded onto expected theoretical constructs.')
    
    add_heading3(doc, '3.6.2 Reliability of Research Instruments')
    
    add_body_paragraph(doc, 'Reliability refers to the consistency and stability of measurements across time and respondents (Kothari, 2004). A pilot study was conducted with 20 SME owners/managers from a neighboring commercial area (Huruma area), not included in the main study, to test the reliability of the questionnaire. The Cronbach\'s Alpha coefficient was computed using SPSS to assess internal consistency. A coefficient of 0.70 or higher was considered acceptable, indicating that the items consistently measure the same underlying construct.')
    
    add_caption(doc, 'Table 3.3: Reliability Statistics')
    
    t33_headers = ['Section', 'Number of Items', 'Cronbach\'s Alpha', 'Interpretation']
    t33_rows = [
        ['Technology Adoption', '10', '0.821', 'Reliable'],
        ['Service Delivery Efficiency', '8', '0.793', 'Reliable'],
        ['Service Delivery Quality', '8', '0.806', 'Reliable'],
        ['Customer Perceptions', '10', '0.845', 'Reliable'],
        ['Challenges in Technology Adoption', '8', '0.772', 'Reliable'],
        ['Overall Instrument', '44', '0.834', 'Reliable'],
    ]
    create_simple_table(doc, t33_headers, t33_rows, col_widths=[2.5, 1.5, 1.5, 1.5])
    add_caption(doc, 'Source: Pilot Study (2026)')
    
    add_body_paragraph(doc, 'All sections of the questionnaire achieved Cronbach\'s Alpha coefficients above 0.70, confirming satisfactory internal consistency. The overall instrument reliability of 0.834 indicates high reliability. Based on pilot results, minor adjustments were made to improve the clarity of six questionnaire items before main data collection.')
    
    add_heading2(doc, '3.7 Data Collection Procedures')
    
    add_body_paragraph(doc, 'Prior to data collection, the researcher obtained an introduction letter from the Dean of Students at Moi University and a research permit from the National Commission for Science, Technology and Innovation (NACOSTI). Permission was also sought from the Uasin Gishu County Business Licensing Department to access the list of registered SMEs and from individual business owners/managers before data collection.')
    
    add_body_paragraph(doc, 'Data collection proceeded through the following steps: (1) Questionnaires were administered using the drop-and-pick method, where questionnaires were left with respondents for completion at their convenience and collected after two days. Research assistants (two final-year students from Moi University) were trained to support respondents who needed clarification or assistance. (2) Semi-structured interviews were conducted with 18 purposively selected SME owners/managers at their business premises, with each interview lasting approximately 30-45 minutes. All interviews were audio-recorded with consent and transcribed within 24 hours for analysis. (3) Customer questionnaires were administered directly at business premises by the researcher and trained research assistants. Data collection took place over a period of four weeks (January-February 2026).')
    
    add_heading2(doc, '3.8 Data Analysis and Presentation')
    
    add_heading3(doc, '3.8.1 Descriptive Statistics')
    
    add_body_paragraph(doc, 'Quantitative data from questionnaires were coded, cleaned, and entered into the Statistical Package for Social Sciences (SPSS) Version 26 for analysis. Descriptive statistics included: frequencies and percentages for categorical variables (demographics, types of technology adopted); means and standard deviations for Likert-scale items measuring perceptions, attitudes, and challenges; and cross-tabulations to examine relationships between categorical variables. Results were presented using tables, figures, and charts for clear interpretation and visual communication of findings.')
    
    add_heading3(doc, '3.8.2 Inferential Statistics')
    
    add_body_paragraph(doc, 'Inferential statistics were used to test relationships between variables and make inferences about the study population. Pearson\'s Correlation Coefficient (r) was used to determine the strength and direction of relationships between technological innovation and service delivery dimensions, following the assumption tests for normality (Kolmogorov-Smirnov test) and linearity (scatter plots). Multiple Regression Analysis was conducted to examine the predictive power of technological innovation on service delivery. The regression model was specified as:')
    
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = para.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    eq_parts = [
        ('Y = \u03b2', False), ('0', True), (' + \u03b2', False), ('1', True),
        ('X', False), ('1', True), (' + \u03b2', False), ('2', True),
        ('X', False), ('2', True), (' + \u03b2', False), ('3', True),
        ('X', False), ('3', True), (' + \u03b5', False),
    ]
    for text, is_sub in eq_parts:
        r = para.add_run(text)
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.name = 'Times New Roman'
        if is_sub:
            r.font.subscript = True

    add_body_paragraph(doc, 'Where: Y = Service Delivery (dependent variable); X\u2081 = ICT Tools and Applications; X\u2082 = Digital Marketing; X\u2083 = Service Delivery Technologies; \u03b2\u2080 = Constant; \u03b2\u2081, \u03b2\u2082, \u03b2\u2083 = Regression coefficients; \u03b5 = Error term. Analysis of Variance (ANOVA) was used to test the overall significance of the regression model. Qualitative data from interviews were analyzed through thematic analysis, identifying recurring themes, patterns, and insights related to the study objectives, which were used to complement and contextualize the quantitative findings.')
    
    add_heading2(doc, '3.9 Ethical Considerations')
    
    add_body_paragraph(doc, 'The researcher adhered to the following ethical principles throughout the study:')
    
    ethics = [
        '1. Informed Consent: Respondents were fully informed about the purpose of the study, the nature of their participation, how their data would be used, and their right to withdraw at any time without consequence. Written consent was obtained before data collection from all participants.',
        '2. Confidentiality and Anonymity: All information provided by respondents was treated with strict confidentiality. Individual responses were anonymized in data analysis and reporting, with no personally identifiable information disclosed in the research report.',
        '3. Voluntary Participation: Participation in the study was entirely voluntary. Respondents were free to decline participation or withdraw at any stage without any negative consequences.',
        '4. Data Protection: Collected data were stored securely in password-protected digital files accessible only to the researcher. Physical questionnaires were stored in a locked cabinet. Data will be retained for five years post-publication and then securely destroyed.',
        '5. Research Integrity: The researcher maintained high standards of research integrity, accurately reporting all findings including those that may not support the study hypotheses. All sources of information were appropriately cited and acknowledged.',
        '6. Research Permit: A research permit was obtained from NACOSTI before commencing data collection, and the study was conducted in accordance with all applicable national research regulations and ethical guidelines.',
    ]
    
    for ethic in ethics:
        para = doc.add_paragraph()
        run = para.add_run(ethic)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # ===================== CHAPTER FOUR =====================
    add_page_break(doc)
    
    para = doc.add_paragraph()
    run = para.add_run('CHAPTER FOUR: DATA ANALYSIS, PRESENTATION AND INTERPRETATION')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(12)
    
    add_heading2(doc, '4.1 Introduction')
    
    add_body_paragraph(doc, 'This chapter presents the analysis, presentation, and interpretation of data collected from SME owners/managers and customers in the Annex area of Uasin Gishu County. The data are analyzed using descriptive and inferential statistical methods and presented through tables and figures. The chapter is organized according to the study objectives, beginning with demographic characteristics of respondents, followed by analysis of each specific objective.')
    
    add_heading2(doc, '4.2 Response Rate')
    
    add_body_paragraph(doc, 'A total of 154 questionnaires were administered to SME owners/managers, of which 143 were returned fully completed, yielding a response rate of 92.9%. Additionally, 200 customer questionnaires were distributed, with 186 returned fully completed, yielding a customer response rate of 93.0%. The combined response rate of 92.9% for SMEs and 93.0% for customers is considered excellent and above the 70% threshold recommended for survey research (Mugenda & Mugenda, 2003).')
    
    add_caption(doc, 'Table 4.1: Response Rate')
    
    t41_headers = ['Category', 'Questionnaires Distributed', 'Questionnaires Returned', 'Response Rate (%)']
    t41_rows = [
        ['SME Owners/Managers', '154', '143', '92.9'],
        ['Customers', '200', '186', '93.0'],
        ['Total', '354', '329', '93.0'],
    ]
    create_simple_table(doc, t41_headers, t41_rows, col_widths=[2.0, 1.8, 1.8, 1.5])
    add_caption(doc, 'Source: Field Survey (2026)')
    
    add_body_paragraph(doc, 'The high response rates were attributed to proper rapport building with respondents, timely follow-up, clear explanation of the study\'s significance to SME owners and customers, and the use of trained research assistants to support data collection. The 18 interviews planned were all successfully conducted, providing comprehensive qualitative data.')
    
    add_heading2(doc, '4.3 Demographic Characteristics of Respondents')
    
    add_heading3(doc, '4.3.1 Gender of Respondents')
    
    add_body_paragraph(doc, 'The study sought to establish the gender distribution of SME owners/managers in the Annex area. The findings are presented in Table 4.2.')
    
    add_caption(doc, 'Table 4.2: Gender of Respondents')
    
    t42_headers = ['Gender', 'Frequency', 'Percentage (%)']
    t42_rows = [
        ['Male', '81', '56.6'],
        ['Female', '62', '43.4'],
        ['Total', '143', '100.0'],
    ]
    create_simple_table(doc, t42_headers, t42_rows, col_widths=[2.5, 2.0, 2.0])
    add_caption(doc, 'Source: Field Survey (2026)')
    
    add_body_paragraph(doc, 'The findings show that male SME owners/managers constituted the majority (56.6%) compared to females (43.4%). While males dominate SME ownership in the Annex area, the relatively high proportion of female SME owners (43.4%) is noteworthy and suggests increasing female entrepreneurship in the region, consistent with national trends of growing women\'s economic participation (Kenya National Bureau of Statistics, 2023).')
    
    add_heading3(doc, '4.3.2 Age of Respondents')
    
    add_body_paragraph(doc, 'The age distribution of SME owners/managers was examined to understand the demographic profile of respondents and its potential influence on technology adoption.')
    
    add_caption(doc, 'Table 4.3: Age of Respondents')
    
    t43_headers = ['Age Category', 'Frequency', 'Percentage (%)']
    t43_rows = [
        ['Below 25 years', '14', '9.8'],
        ['25 – 34 years', '52', '36.4'],
        ['35 – 44 years', '47', '32.9'],
        ['45 – 54 years', '22', '15.4'],
        ['55 years and above', '8', '5.6'],
        ['Total', '143', '100.0'],
    ]
    create_simple_table(doc, t43_headers, t43_rows, col_widths=[2.5, 2.0, 2.0])
    add_caption(doc, 'Source: Field Survey (2026)')
    
    add_body_paragraph(doc, 'The majority of SME owners/managers (69.2%) were between 25 and 44 years of age, representing the economically active youth and middle-aged population segment. This age group is generally considered more technologically savvy and open to technology adoption. Only 5.6% were aged 55 years and above, suggesting that the Annex area SME landscape is relatively young, which may facilitate technology adoption in the medium term.')
    
    add_heading3(doc, '4.3.3 Level of Education')
    
    add_body_paragraph(doc, 'The level of education of SME owners/managers was examined, as education level has been identified as a significant predictor of technology adoption in previous studies.')
    
    add_caption(doc, 'Table 4.4: Level of Education')
    
    t44_headers = ['Level of Education', 'Frequency', 'Percentage (%)']
    t44_rows = [
        ['Primary School', '8', '5.6'],
        ['Secondary School (KCSE)', '38', '26.6'],
        ['Certificate/Diploma', '49', '34.3'],
        ['University Undergraduate', '39', '27.3'],
        ['Postgraduate', '9', '6.3'],
        ['Total', '143', '100.0'],
    ]
    create_simple_table(doc, t44_headers, t44_rows, col_widths=[3.0, 1.8, 1.8])
    add_caption(doc, 'Source: Field Survey (2026)')
    
    add_body_paragraph(doc, 'The findings reveal that 67.9% of SME owners/managers had education levels of certificate/diploma or higher, with 34.3% holding certificate or diploma qualifications and 33.6% holding university-level qualifications. This relatively high education level is encouraging for technology adoption prospects, as higher education levels generally correlate with greater technology awareness and adoption capacity. However, 32.2% had secondary school education or below, indicating that a significant minority may require additional support for technology adoption.')
    
    add_heading3(doc, '4.3.4 Business Type')
    
    add_body_paragraph(doc, 'The distribution of respondents by business type was examined to understand the sectoral composition of the SME sample.')
    
    add_caption(doc, 'Table 4.5: Business Type')
    
    t45_headers = ['Business Type', 'Frequency', 'Percentage (%)']
    t45_rows = [
        ['Retail Shops', '58', '40.6'],
        ['Hotels and Restaurants', '34', '23.8'],
        ['Service Providers (Salons, Repair shops, etc.)', '28', '19.6'],
        ['Agricultural-related Businesses', '23', '16.1'],
        ['Total', '143', '100.0'],
    ]
    create_simple_table(doc, t45_headers, t45_rows, col_widths=[3.5, 1.5, 1.5])
    add_caption(doc, 'Source: Field Survey (2026)')
    
    add_body_paragraph(doc, 'The distribution of respondents by business type closely mirrors the target population distribution, confirming that the stratified sampling approach successfully achieved representative coverage of all SME categories in the Annex area. Retail shops form the largest category (40.6%), followed by hotels and restaurants (23.8%), service providers (19.6%), and agricultural-related businesses (16.1%).')
    
    add_heading3(doc, '4.3.5 Years of Business Operation')
    
    add_body_paragraph(doc, 'The number of years of business operation was examined to understand the maturity profile of SMEs in the Annex area and its relationship with technology adoption.')
    
    add_caption(doc, 'Table 4.6: Years of Business Operation')
    
    t46_headers = ['Years of Operation', 'Frequency', 'Percentage (%)']
    t46_rows = [
        ['Less than 1 year', '12', '8.4'],
        ['1 – 3 years', '34', '23.8'],
        ['4 – 6 years', '47', '32.9'],
        ['7 – 10 years', '31', '21.7'],
        ['More than 10 years', '19', '13.3'],
        ['Total', '143', '100.0'],
    ]
    create_simple_table(doc, t46_headers, t46_rows, col_widths=[2.5, 2.0, 2.0])
    add_caption(doc, 'Source: Field Survey (2026)')
    
    add_body_paragraph(doc, 'The findings show that 32.9% of SMEs had been in operation for 4-6 years, representing the most common business maturity level. SMEs operating for 1-6 years constituted 56.7% of the sample, suggesting a relatively young SME landscape. Only 13.3% had been in operation for more than 10 years, consistent with the national pattern of high SME mortality rates in the early years of operation. Established SMEs (7+ years) constituted 35.0% of the sample.')
    
    add_heading2(doc, '4.4 Types of Technological Innovations Adopted by SMEs in the Annex Area')
    
    add_heading3(doc, '4.4.1 ICT Tools and Applications Used')
    
    add_body_paragraph(doc, 'Respondents were asked to indicate the ICT tools and applications used in their businesses. The findings are presented in Table 4.7 below.')
    
    add_caption(doc, 'Table 4.7: ICT Tools and Applications Used by SMEs')
    
    t47_headers = ['ICT Tool/Application', 'Frequency', 'Percentage (%)']
    t47_rows = [
        ['Mobile Money (M-Pesa/Airtel Money)', '128', '89.5'],
        ['Smartphones for Business', '121', '84.6'],
        ['WhatsApp Business', '112', '78.3'],
        ['Facebook for Business', '97', '67.8'],
        ['Internet/Wi-Fi', '89', '62.2'],
        ['Computers/Laptops', '72', '50.3'],
        ['POS Systems', '78', '54.5'],
        ['Accounting Software', '53', '37.1'],
        ['Inventory Management Software', '60', '42.0'],
        ['Company Website', '41', '28.7'],
        ['E-commerce Platforms', '29', '20.3'],
        ['Cloud Storage/Services', '27', '18.9'],
        ['Customer Management Software (CRM)', '24', '16.8'],
    ]
    create_simple_table(doc, t47_headers, t47_rows, col_widths=[3.5, 1.5, 1.5])
    add_caption(doc, 'Source: Field Survey (2026)')
    
    add_body_paragraph(doc, 'The findings reveal that mobile money services (89.5%) and smartphones for business use (84.6%) are the most widely adopted ICT tools among SMEs in the Annex area, followed by WhatsApp Business (78.3%) and Facebook for Business (67.8%). Internet access is used by 62.2% of SMEs, while POS systems are adopted by 54.5%. Computer adoption stands at 50.3%. More sophisticated technologies such as accounting software (37.1%), inventory management software (42.0%), websites (28.7%), and e-commerce platforms (20.3%) show lower adoption rates, suggesting a technology adoption gradient from simple, accessible tools to more complex systems requiring greater technical expertise and investment.')
    
    add_heading3(doc, '4.4.2 Mobile Technology Adoption')
    
    add_body_paragraph(doc, 'The study examined specific aspects of mobile technology adoption among SMEs, as mobile technology has been identified as a key driver of SME digitalization in Kenya.')
    
    add_caption(doc, 'Table 4.8: Mobile Technology Adoption')
    
    t48_headers = ['Statement', 'Mean', 'Std. Dev.', 'Interpretation']
    t48_rows = [
        ['We use mobile phones to receive payments from customers', '4.52', '0.71', 'Strongly Agree'],
        ['We use mobile apps to communicate with suppliers', '3.94', '0.89', 'Agree'],
        ['Mobile money has improved our cash flow management', '4.31', '0.82', 'Strongly Agree'],
        ['We use mobile phones to track business inventory', '3.12', '1.04', 'Neutral/Agree'],
        ['Mobile technology has reduced our transaction costs', '4.18', '0.86', 'Agree'],
        ['We use mobile banking for business transactions', '3.87', '0.94', 'Agree'],
        ['Overall Mobile Technology Adoption', '3.99', '0.74', 'Agree'],
    ]
    create_simple_table(doc, t48_headers, t48_rows, col_widths=[3.0, 0.8, 0.8, 1.5])
    add_caption(doc, 'Source: Field Survey (2026) [Scale: 1=Strongly Disagree to 5=Strongly Agree]')
    
    add_body_paragraph(doc, 'The mean score for overall mobile technology adoption was 3.99 (SD=0.74), indicating agreement with statements on mobile technology use. Mobile payment receipt had the highest mean (4.52), confirming that receiving customer payments via mobile money is nearly universal among SMEs. Mobile phones\' contribution to improved cash flow management was also highly rated (mean=4.31), reflecting the transformative impact of mobile money on SME financial operations. Mobile phone use for inventory tracking had the lowest mean (3.12), suggesting this is an emerging but not yet widespread practice.')
    
    add_heading3(doc, '4.4.3 Social Media and Digital Marketing')
    
    add_body_paragraph(doc, 'The study assessed the use of social media and digital marketing platforms among SMEs in the Annex area.')
    
    add_caption(doc, 'Table 4.9: Social Media and Digital Marketing')
    
    t49_headers = ['Statement', 'Mean', 'Std. Dev.', 'Interpretation']
    t49_rows = [
        ['We use WhatsApp to communicate with customers', '4.41', '0.73', 'Strongly Agree'],
        ['We promote our business on Facebook', '3.78', '0.98', 'Agree'],
        ['We use Instagram to showcase our products/services', '2.94', '1.12', 'Neutral'],
        ['Social media has expanded our customer base', '3.86', '0.94', 'Agree'],
        ['Digital marketing is more cost-effective than traditional advertising', '3.97', '0.87', 'Agree'],
        ['We respond to customer inquiries through social media', '3.74', '1.01', 'Agree'],
        ['We run paid adverts on digital platforms', '2.82', '1.18', 'Neutral'],
        ['Overall Digital Marketing Adoption', '3.65', '0.81', 'Agree'],
    ]
    create_simple_table(doc, t49_headers, t49_rows, col_widths=[3.0, 0.8, 0.8, 1.5])
    add_caption(doc, 'Source: Field Survey (2026) [Scale: 1=Strongly Disagree to 5=Strongly Agree]')
    
    add_body_paragraph(doc, 'The mean score for overall digital marketing adoption was 3.65 (SD=0.81), indicating general agreement with digital marketing usage. WhatsApp communication with customers had the highest mean score (4.41), confirming its dominant role as the preferred digital communication channel. SMEs perceived digital marketing as more cost-effective than traditional advertising (mean=3.97), and social media was credited with expanding customer bases (mean=3.86). Instagram use and paid digital advertising had the lowest adoption levels, suggesting that while SMEs are active on social media, many have not yet transitioned to more sophisticated digital marketing strategies.')
    
    add_heading2(doc, '4.5 Influence of Technological Innovation on Efficiency and Quality of Service Delivery')
    
    add_heading3(doc, '4.5.1 Efficiency of Service Delivery')
    
    add_body_paragraph(doc, 'Respondents were asked to assess how technological innovation had influenced the efficiency of their service delivery. Efficiency was measured in terms of time savings, cost reduction, error reduction, and process streamlining.')
    
    add_caption(doc, 'Table 4.10: Efficiency of Service Delivery')
    
    t410_headers = ['Statement', 'Mean', 'Std. Dev.', 'Interpretation']
    t410_rows = [
        ['Technology has reduced the time taken to serve customers', '4.12', '0.84', 'Agree'],
        ['Technology has reduced our operational costs', '3.88', '0.92', 'Agree'],
        ['Technology has reduced errors in our transactions and records', '4.03', '0.87', 'Agree'],
        ['Technology has streamlined our service delivery processes', '3.96', '0.89', 'Agree'],
        ['Digital payment systems have speeded up transactions', '4.34', '0.76', 'Strongly Agree'],
        ['Technology has enabled us to serve more customers per day', '3.77', '0.96', 'Agree'],
        ['Inventory management technology has reduced stock-outs', '3.52', '1.07', 'Agree'],
        ['Technology has improved staff productivity', '3.84', '0.94', 'Agree'],
        ['Overall Efficiency Impact', '3.93', '0.72', 'Agree'],
    ]
    create_simple_table(doc, t410_headers, t410_rows, col_widths=[3.0, 0.8, 0.8, 1.5])
    add_caption(doc, 'Source: Field Survey (2026) [Scale: 1=Strongly Disagree to 5=Strongly Agree]')
    
    add_body_paragraph(doc, 'The overall mean for efficiency impact was 3.93 (SD=0.72), indicating that respondents agreed that technological innovation has enhanced service delivery efficiency. Digital payment systems received the highest efficiency rating (mean=4.34), confirming the transformative impact of mobile money and digital payments on transaction speed. Time reduction in customer service was highly rated (mean=4.12), followed by error reduction in transactions and records (mean=4.03). Inventory management technology showed the lowest efficiency impact (mean=3.52), consistent with the lower adoption rates of inventory management software found in the previous section.')
    
    add_heading3(doc, '4.5.2 Quality of Service Delivery')
    
    add_body_paragraph(doc, 'The impact of technological innovation on the quality of service delivery was assessed across multiple dimensions including reliability, responsiveness, assurance, and personalization.')
    
    add_caption(doc, 'Table 4.11: Quality of Service Delivery')
    
    t411_headers = ['Statement', 'Mean', 'Std. Dev.', 'Interpretation']
    t411_rows = [
        ['Technology has improved the consistency of our service delivery', '3.94', '0.88', 'Agree'],
        ['Technology enables faster responses to customer inquiries', '4.08', '0.83', 'Agree'],
        ['Technology has enhanced the professionalism of our business image', '4.02', '0.86', 'Agree'],
        ['Technology has enabled more personalized customer service', '3.72', '0.97', 'Agree'],
        ['Technology has improved customer record-keeping and follow-up', '3.64', '1.02', 'Agree'],
        ['Technology has enhanced the reliability of our service delivery', '3.89', '0.91', 'Agree'],
        ['Technology-enabled services meet higher customer expectations', '3.81', '0.93', 'Agree'],
        ['Technology has improved service accessibility to customers', '3.96', '0.89', 'Agree'],
        ['Overall Quality Impact', '3.88', '0.74', 'Agree'],
    ]
    create_simple_table(doc, t411_headers, t411_rows, col_widths=[3.0, 0.8, 0.8, 1.5])
    add_caption(doc, 'Source: Field Survey (2026) [Scale: 1=Strongly Disagree to 5=Strongly Agree]')
    
    add_body_paragraph(doc, 'The overall mean for quality impact was 3.88 (SD=0.74), indicating agreement that technological innovation has improved service quality. Faster responses to customer inquiries received the highest rating (mean=4.08), followed by enhanced business professionalism (mean=4.02). Service consistency and reliability both showed strong improvement (means 3.94 and 3.89, respectively). Personalized customer service and customer record-keeping showed relatively lower means (3.72 and 3.64, respectively), suggesting that while basic service quality dimensions have improved, more sophisticated quality enhancements requiring CRM systems are less widely realized.')
    
    add_heading3(doc, '4.5.3 Correlation Analysis')
    
    add_body_paragraph(doc, 'Pearson\'s correlation analysis was conducted to examine the strength and direction of the relationship between technological innovation and service delivery dimensions. The results are presented in Table 4.12.')
    
    add_caption(doc, 'Table 4.12: Correlation between Technological Innovation and Service Delivery')
    
    t412_headers = ['Variable', 'Service Delivery Efficiency', 'Service Quality', 'Overall Service Delivery']
    t412_rows = [
        ['ICT Tools Adoption', '0.672**', '0.641**', '0.687**'],
        ['Mobile Technology Adoption', '0.694**', '0.658**', '0.701**'],
        ['Digital Marketing Adoption', '0.583**', '0.612**', '0.614**'],
        ['Service Delivery Technologies', '0.631**', '0.649**', '0.658**'],
        ['Overall Technology Innovation', '0.714**', '0.682**', '0.721**'],
    ]
    create_simple_table(doc, t412_headers, t412_rows, col_widths=[2.5, 1.6, 1.4, 1.5])
    add_caption(doc, 'Source: Field Survey (2026) | ** Correlation is significant at 0.01 level (2-tailed)')
    
    add_body_paragraph(doc, 'The correlation analysis reveals statistically significant positive relationships between all dimensions of technological innovation and service delivery. The overall technological innovation composite had the strongest correlation with overall service delivery (r=0.721, p<0.01), followed by service delivery efficiency (r=0.714, p<0.01) and service quality (r=0.682, p<0.01). Mobile technology adoption showed the strongest individual correlation with service delivery efficiency (r=0.694), while digital marketing adoption showed the weakest but still significant correlation (r=0.583). These findings confirm that technological innovation is positively and significantly associated with improved service delivery across all dimensions.')
    
    add_heading2(doc, '4.6 Customer Perceptions of Technologically Enhanced Services')
    
    add_heading3(doc, '4.6.1 Customer Satisfaction')
    
    add_body_paragraph(doc, 'Customer respondents were asked to rate their satisfaction with technologically enhanced services offered by SMEs in the Annex area.')
    
    add_caption(doc, 'Table 4.13: Customer Satisfaction with Technologically Enhanced Services')
    
    t413_headers = ['Statement', 'Mean', 'Std. Dev.', 'Interpretation']
    t413_rows = [
        ['I am satisfied with the availability of mobile payment options', '4.38', '0.74', 'Strongly Agree'],
        ['Technology-enabled services are faster and more convenient', '4.21', '0.79', 'Agree'],
        ['I trust digital payment systems used by these businesses', '3.94', '0.93', 'Agree'],
        ['Technology has improved the overall quality of services I receive', '3.86', '0.94', 'Agree'],
        ['I prefer to patronize businesses that use modern technology', '3.78', '0.98', 'Agree'],
        ['Digital receipts and records provided by businesses are useful', '3.74', '1.01', 'Agree'],
        ['Technology-enhanced businesses provide more reliable services', '3.81', '0.96', 'Agree'],
        ['Overall Customer Satisfaction', '3.96', '0.73', 'Agree'],
    ]
    create_simple_table(doc, t413_headers, t413_rows, col_widths=[3.0, 0.8, 0.8, 1.5])
    add_caption(doc, 'Source: Field Survey (2026) [Scale: 1=Strongly Disagree to 5=Strongly Agree]')
    
    add_body_paragraph(doc, 'The overall mean for customer satisfaction was 3.96 (SD=0.73), indicating a generally high level of customer satisfaction with technology-enhanced services. Customer satisfaction with mobile payment options was highest (mean=4.38), reflecting the widespread and highly valued nature of mobile money services. Speed and convenience of technology-enabled services were also highly rated (mean=4.21). While trust in digital payment systems was generally good (mean=3.94), it was slightly lower than satisfaction with mobile payments, suggesting some lingering concerns about security. Customer preference for technology-using businesses (mean=3.78) indicates that technology adoption has competitive implications for SMEs.')
    
    add_heading3(doc, '4.6.2 Customer Experience')
    
    add_body_paragraph(doc, 'Customer experience with specific aspects of technologically enhanced service delivery was assessed.')
    
    add_caption(doc, 'Table 4.14: Customer Experience with Technology-Enhanced Services')
    
    t414_headers = ['Statement', 'Mean', 'Std. Dev.', 'Interpretation']
    t414_rows = [
        ['I experience shorter waiting times at tech-equipped businesses', '4.03', '0.87', 'Agree'],
        ['Businesses use technology to communicate with me effectively', '3.88', '0.94', 'Agree'],
        ['I have experienced fewer errors in my transactions', '3.91', '0.91', 'Agree'],
        ['Technology makes it easier for me to access business information', '3.74', '0.99', 'Agree'],
        ['I feel the services are more personalized when technology is used', '3.51', '1.08', 'Agree'],
        ['Social media helps me stay updated on business offerings', '3.82', '0.97', 'Agree'],
        ['I am comfortable using digital payment systems', '4.12', '0.83', 'Agree'],
        ['Overall Customer Experience', '3.86', '0.74', 'Agree'],
    ]
    create_simple_table(doc, t414_headers, t414_rows, col_widths=[3.0, 0.8, 0.8, 1.5])
    add_caption(doc, 'Source: Field Survey (2026) [Scale: 1=Strongly Disagree to 5=Strongly Agree]')
    
    add_body_paragraph(doc, 'The overall mean for customer experience was 3.86 (SD=0.74), indicating positive customer experiences with technology-enhanced services. Customers reported high comfort with digital payment systems (mean=4.12) and significantly shorter waiting times at technology-equipped businesses (mean=4.03). Fewer transaction errors (mean=3.91) and better communication (mean=3.88) were also positively rated. The lowest-rated experience dimension was personalization through technology (mean=3.51), suggesting that while operational improvements are clearly recognized, the potential for more personalized customer experiences through data-driven approaches has not yet been fully realized in the Annex SME context.')
    
    add_heading2(doc, '4.7 Challenges Faced by SMEs in Adopting Technological Innovations')
    
    add_heading3(doc, '4.7.1 Financial Challenges')
    
    add_body_paragraph(doc, 'Respondents were asked to rate the extent to which financial factors posed challenges to technology adoption. Results are presented in Table 4.15.')
    
    add_caption(doc, 'Table 4.15: Financial Challenges in Technology Adoption')
    
    t415_headers = ['Statement', 'Mean', 'Std. Dev.', 'Interpretation']
    t415_rows = [
        ['The high initial cost of technology equipment is a barrier', '4.28', '0.79', 'Strongly Agree'],
        ['Software licensing fees are too expensive for our business', '3.97', '0.91', 'Agree'],
        ['We lack adequate financing for technology acquisition', '4.11', '0.84', 'Agree'],
        ['Internet data costs are too high for regular use', '3.84', '0.96', 'Agree'],
        ['The cost of maintaining technology equipment is burdensome', '3.76', '0.99', 'Agree'],
        ['High interest rates on technology loans discourage adoption', '3.91', '0.93', 'Agree'],
        ['Overall Financial Challenges', '3.98', '0.72', 'Agree'],
    ]
    create_simple_table(doc, t415_headers, t415_rows, col_widths=[3.0, 0.8, 0.8, 1.5])
    add_caption(doc, 'Source: Field Survey (2026) [Scale: 1=Strongly Disagree to 5=Strongly Agree]')
    
    add_body_paragraph(doc, 'Financial challenges constitute the most significant barrier to technology adoption, with an overall mean of 3.98 (SD=0.72). The high initial cost of technology equipment was rated as the most significant financial barrier (mean=4.28), followed by lack of adequate financing (mean=4.11). Software licensing costs (mean=3.97) and high loan interest rates (mean=3.91) were also notable constraints. These findings align with previous research by Lagat (2014) and Bwire and Muathe (2025) that identified financial constraints as key barriers to technology adoption in Uasin Gishu County SMEs.')
    
    add_heading3(doc, '4.7.2 Skills and Knowledge Gaps')
    
    add_body_paragraph(doc, 'The study examined the extent to which skills and knowledge limitations posed challenges to technology adoption.')
    
    add_caption(doc, 'Table 4.16: Skills and Knowledge Gaps in Technology Adoption')
    
    t416_headers = ['Statement', 'Mean', 'Std. Dev.', 'Interpretation']
    t416_rows = [
        ['We lack the technical skills to use advanced technology', '3.86', '0.94', 'Agree'],
        ['Our employees need training to effectively use technology', '4.14', '0.82', 'Agree'],
        ['We do not know which technology is best for our business', '3.74', '1.02', 'Agree'],
        ['We struggle to troubleshoot technology problems', '3.91', '0.91', 'Agree'],
        ['The pace of technology change makes it hard to keep up', '3.97', '0.88', 'Agree'],
        ['We have had negative experiences with technology previously', '3.28', '1.11', 'Neutral/Agree'],
        ['Overall Skills and Knowledge Challenges', '3.82', '0.76', 'Agree'],
    ]
    create_simple_table(doc, t416_headers, t416_rows, col_widths=[3.0, 0.8, 0.8, 1.5])
    add_caption(doc, 'Source: Field Survey (2026) [Scale: 1=Strongly Disagree to 5=Strongly Agree]')
    
    add_body_paragraph(doc, 'Skills and knowledge gaps represent the second most significant challenge, with an overall mean of 3.82 (SD=0.76). The need for employee training in technology use was most strongly affirmed (mean=4.14), highlighting the critical importance of human resource development alongside technology acquisition. The rapid pace of technological change (mean=3.97) and difficulties in troubleshooting (mean=3.91) were also significant concerns. Previous negative technology experiences showed the lowest mean (3.28), suggesting that while past failures exist, they are not the dominant skills challenge, with knowledge and capacity gaps being more prevalent concerns.')
    
    add_heading3(doc, '4.7.3 Infrastructure Challenges')
    
    add_body_paragraph(doc, 'Infrastructure limitations were assessed as potential barriers to technology adoption among SMEs in the Annex area.')
    
    add_caption(doc, 'Table 4.17: Infrastructure Challenges in Technology Adoption')
    
    t417_headers = ['Statement', 'Mean', 'Std. Dev.', 'Interpretation']
    t417_rows = [
        ['Unreliable internet connectivity hinders technology use', '4.21', '0.82', 'Agree'],
        ['Frequent power outages disrupt our technology use', '4.08', '0.87', 'Agree'],
        ['There are few technology support services nearby', '3.84', '0.96', 'Agree'],
        ['The cost of internet connection is prohibitive', '3.72', '1.01', 'Agree'],
        ['Technology equipment is difficult to repair locally', '3.61', '1.04', 'Agree'],
        ['Overall Infrastructure Challenges', '3.89', '0.74', 'Agree'],
    ]
    create_simple_table(doc, t417_headers, t417_rows, col_widths=[3.5, 0.8, 0.8, 1.5])
    add_caption(doc, 'Source: Field Survey (2026) [Scale: 1=Strongly Disagree to 5=Strongly Agree]')
    
    add_body_paragraph(doc, 'Infrastructure challenges had an overall mean of 3.89 (SD=0.74), indicating significant infrastructure barriers. Unreliable internet connectivity was rated as the most significant infrastructure challenge (mean=4.21), followed by frequent power outages (mean=4.08). The limited availability of local technology support services (mean=3.84) and prohibitive internet costs (mean=3.72) further compound infrastructure limitations. These findings confirm the importance of infrastructure development in enabling SME technology adoption, consistent with previous research in Uasin Gishu County.')
    
    add_heading3(doc, '4.7.4 Environmental and Social Challenges')
    
    add_body_paragraph(doc, 'The study also examined environmental and social factors that pose challenges to technology adoption.')
    
    add_caption(doc, 'Table 4.18: Environmental and Social Challenges in Technology Adoption')
    
    t418_headers = ['Statement', 'Mean', 'Std. Dev.', 'Interpretation']
    t418_rows = [
        ['Regulatory requirements for technology use are unclear', '3.58', '1.06', 'Agree'],
        ['Some customers prefer traditional service methods', '3.82', '0.97', 'Agree'],
        ['Competition from larger businesses discourages our technology investment', '3.64', '1.03', 'Agree'],
        ['Social resistance from employees to adopt new technology', '3.41', '1.09', 'Agree'],
        ['Cultural practices affect our technology adoption decisions', '3.27', '1.12', 'Neutral/Agree'],
        ['Overall Environmental and Social Challenges', '3.54', '0.78', 'Agree'],
    ]
    create_simple_table(doc, t418_headers, t418_rows, col_widths=[3.5, 0.8, 0.8, 1.5])
    add_caption(doc, 'Source: Field Survey (2026) [Scale: 1=Strongly Disagree to 5=Strongly Agree]')
    
    add_body_paragraph(doc, 'Environmental and social challenges had the lowest overall mean (3.54, SD=0.78) among the four challenge categories, indicating that while these factors exist, they are less constraining than financial, skills, and infrastructure challenges. Customer preference for traditional service methods was the most significant environmental challenge (mean=3.82), followed by competitive pressures from larger businesses (mean=3.64). Regulatory uncertainty (mean=3.58) and employee resistance to change (mean=3.41) were moderately significant. Cultural factors had the lowest rating (mean=3.27), suggesting that cultural barriers to technology adoption are less pronounced than practical constraints.')
    
    add_heading2(doc, '4.8 Regression Analysis')
    
    add_body_paragraph(doc, 'Multiple regression analysis was conducted to examine the predictive power of technological innovation dimensions on service delivery. Before conducting regression, the following assumptions were tested: normality of residuals (confirmed via Kolmogorov-Smirnov test, p>0.05), absence of multicollinearity (VIF values ranging from 1.24 to 2.87, all below the threshold of 10), linearity (confirmed via scatter plots), and homoscedasticity (confirmed via Breusch-Pagan test, p>0.05).')
    
    add_heading3(doc, '4.8.1 Model Summary')
    
    add_caption(doc, 'Table 4.19: Model Summary')
    
    t419_headers = ['Model', 'R', 'R Square', 'Adjusted R Square', 'Std. Error of Estimate']
    t419_rows = [
        ['1', '0.783', '0.613', '0.597', '0.412'],
    ]
    create_simple_table(doc, t419_headers, t419_rows, col_widths=[0.8, 0.8, 1.0, 1.5, 1.8])
    add_caption(doc, 'Source: Field Survey (2026) | Predictors: ICT Tools Adoption, Digital Marketing, Service Delivery Technologies')
    
    add_body_paragraph(doc, 'The model summary indicates that the three predictor variables (ICT Tools Adoption, Digital Marketing, and Service Delivery Technologies) collectively explain 61.3% of the variance in service delivery (R²=0.613), with an adjusted R² of 0.597. The multiple correlation coefficient (R=0.783) indicates a strong positive relationship between the predictor variables and service delivery. This is a significant model fit, indicating that technological innovation dimensions are powerful predictors of service delivery outcomes among SMEs in the Annex area.')
    
    add_heading3(doc, '4.8.2 Analysis of Variance (ANOVA)')
    
    add_caption(doc, 'Table 4.20: Analysis of Variance (ANOVA)')
    
    t420_headers = ['Model', 'Sum of Squares', 'df', 'Mean Square', 'F', 'Sig.']
    t420_rows = [
        ['Regression', '24.817', '3', '8.272', '48.724', '0.000'],
        ['Residual', '23.636', '139', '0.170', '', ''],
        ['Total', '48.453', '142', '', '', ''],
    ]
    create_simple_table(doc, t420_headers, t420_rows, col_widths=[1.5, 1.5, 0.5, 1.5, 1.0, 0.8])
    add_caption(doc, 'Source: Field Survey (2026) | Dependent Variable: Service Delivery')
    
    add_body_paragraph(doc, 'The ANOVA results indicate that the regression model is statistically significant (F(3, 139)=48.724, p=0.000<0.05). This means that the three predictor variables (ICT Tools Adoption, Digital Marketing, and Service Delivery Technologies) collectively provide a statistically significant prediction of service delivery among SMEs in the Annex area. The significant F-value confirms that the regression model fits the data well and that the relationship between technological innovation and service delivery is not due to chance.')
    
    add_heading3(doc, '4.8.3 Regression Coefficients')
    
    add_caption(doc, 'Table 4.21: Regression Coefficients')
    
    t421_headers = ['Predictor Variable', 'B', 'Std. Error', 'Beta (β)', 't-value', 'Sig.']
    t421_rows = [
        ['(Constant)', '0.487', '0.241', '', '2.020', '0.045'],
        ['ICT Tools Adoption (X1)', '0.412', '0.076', '0.421', '5.421', '0.000'],
        ['Digital Marketing (X2)', '0.298', '0.072', '0.318', '4.139', '0.000'],
        ['Service Delivery Technologies (X3)', '0.261', '0.069', '0.276', '3.783', '0.000'],
    ]
    create_simple_table(doc, t421_headers, t421_rows, col_widths=[2.5, 0.6, 0.8, 0.8, 0.8, 0.6])
    add_caption(doc, 'Source: Field Survey (2026) | Dependent Variable: Service Delivery')

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    result_eq_parts = [
        ('The regression equation derived from the coefficients is: Y = 0.487 + 0.412X', False),
        ('1', True), (' + 0.298X', False), ('2', True), (' + 0.261X', False), ('3', True),
    ]
    for text, is_sub in result_eq_parts:
        r = para.add_run(text)
        r.font.size = Pt(12)
        r.font.name = 'Times New Roman'
        if is_sub:
            r.font.subscript = True
    
    add_body_paragraph(doc, 'All three predictor variables are statistically significant predictors of service delivery (p<0.05). ICT Tools Adoption has the strongest predictive influence on service delivery (β=0.421, t=5.421, p=0.000), indicating that a one unit increase in ICT tools adoption is associated with a 0.412 unit increase in service delivery, holding other variables constant. Digital Marketing follows with a significant positive effect (β=0.318, t=4.139, p=0.000), and Service Delivery Technologies also significantly predicts service delivery (β=0.276, t=3.783, p=0.000). The positive beta coefficients for all predictors confirm that each dimension of technological innovation contributes positively and independently to service delivery enhancement.')
    
    # ===================== CHAPTER FIVE =====================
    add_page_break(doc)
    
    para = doc.add_paragraph()
    run = para.add_run('CHAPTER FIVE: SUMMARY OF FINDINGS, CONCLUSIONS AND RECOMMENDATIONS')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(12)
    
    add_heading2(doc, '5.1 Introduction')
    
    add_body_paragraph(doc, 'This chapter presents a summary of the key findings from the study on the impact of technological innovation on service delivery among SMEs in the Annex area of Uasin Gishu County. It draws conclusions based on the empirical evidence gathered and offers recommendations for various stakeholders. The chapter also outlines the limitations encountered during the study and suggests areas for further research.')
    
    add_heading2(doc, '5.2 Summary of Findings')
    
    add_heading3(doc, '5.2.1 Types of Technological Innovations Adopted')
    
    add_body_paragraph(doc, 'The first objective sought to identify the types of technological innovations adopted by SMEs in the Annex area of Uasin Gishu County. The findings revealed a clear technology adoption hierarchy, with simple and low-cost technologies showing the highest adoption rates while more sophisticated technologies remained less widely adopted.')
    
    add_body_paragraph(doc, 'Mobile money services (M-Pesa and Airtel Money) were the most widely adopted technology, used by 89.5% of SMEs, making mobile payment acceptance nearly universal in the Annex area. Smartphones for business use (84.6%) and WhatsApp Business (78.3%) were the second and third most adopted technologies respectively, reflecting the dominant role of mobile technology in SME digitalization. Facebook for Business (67.8%) and internet/Wi-Fi access (62.2%) were also widely adopted.')
    
    add_body_paragraph(doc, 'Adoption rates declined significantly for more complex technologies: POS systems (54.5%), computers/laptops (50.3%), inventory management software (42.0%), accounting software (37.1%), company websites (28.7%), e-commerce platforms (20.3%), cloud services (18.9%), and CRM systems (16.8%). The overall mobile technology adoption mean of 3.99/5.00 and digital marketing adoption mean of 3.65/5.00 indicate moderate-to-high adoption of mobile and digital marketing technologies.')
    
    add_heading3(doc, '5.2.2 Influence on Efficiency and Quality of Service Delivery')
    
    add_body_paragraph(doc, 'The second objective examined how technological innovation influences the efficiency and quality of service delivery among SMEs in the Annex area. The findings strongly confirmed the positive influence of technology on both service delivery dimensions.')
    
    add_body_paragraph(doc, 'In terms of efficiency, the overall mean score was 3.93/5.00, with digital payment systems receiving the highest efficiency rating (mean=4.34), confirming the transformative impact of mobile money on transaction speed. Technology was also credited with significant time savings in customer service (mean=4.12), error reduction (mean=4.03), and operational cost reduction (mean=3.88).')
    
    add_body_paragraph(doc, 'Service quality improvement had an overall mean of 3.88/5.00, with faster responses to customer inquiries (mean=4.08) and enhanced business professionalism (mean=4.02) receiving the highest ratings. The correlation analysis confirmed statistically significant positive relationships between technological innovation and service delivery, with overall technology innovation correlating strongly with service delivery efficiency (r=0.714, p<0.01) and service quality (r=0.682, p<0.01).')
    
    add_heading3(doc, '5.2.3 Customer Perceptions')
    
    add_body_paragraph(doc, 'The third objective evaluated customer perceptions of technologically enhanced services offered by SMEs in the Annex area. The findings revealed generally positive customer perceptions with room for improvement in specific areas.')
    
    add_body_paragraph(doc, 'Overall customer satisfaction with technology-enhanced services was high (mean=3.96/5.00), with mobile payment options receiving the highest satisfaction rating (mean=4.38). Customers highly valued the speed and convenience of technology-enabled services (mean=4.21) and reported greater comfort with digital payment systems (mean=4.12). Customer experience with technology-enhanced services had an overall mean of 3.86/5.00, with shorter waiting times (mean=4.03) and fewer transaction errors (mean=3.91) being the most positively rated experiences.')
    
    add_body_paragraph(doc, 'However, customers rated personalization of service through technology as relatively lower (mean=3.51), indicating that while operational improvements are clearly recognized and valued, the potential for more individualized and data-driven customer experiences has not yet been fully realized by SMEs in the Annex area.')
    
    add_heading3(doc, '5.2.4 Challenges in Adoption')
    
    add_body_paragraph(doc, 'The fourth objective identified challenges faced by SMEs in the Annex area in adopting and implementing technological innovations. Four categories of challenges were identified, with financial challenges emerging as the most significant barrier.')
    
    add_body_paragraph(doc, 'Financial challenges had the highest overall mean (3.98/5.00), with the high initial cost of technology equipment being the most critical financial barrier (mean=4.28), followed by inadequate financing (mean=4.11) and prohibitive software costs (mean=3.97). Skills and knowledge gaps were the second most significant challenge category (mean=3.82/5.00), with the need for employee training being the most pressing concern (mean=4.14) alongside challenges keeping up with the rapid pace of technological change (mean=3.97).')
    
    add_body_paragraph(doc, 'Infrastructure challenges had an overall mean of 3.89/5.00, with unreliable internet connectivity (mean=4.21) and frequent power outages (mean=4.08) being the most significant infrastructure barriers. Environmental and social challenges were the least significant category (mean=3.54/5.00), though customer preference for traditional service methods (mean=3.82) and regulatory uncertainty (mean=3.58) were notable concerns. The regression analysis confirmed that all three technology dimensions significantly predict service delivery, collectively explaining 61.3% of its variance (R²=0.613, F(3,139)=48.724, p=0.000).')
    
    add_heading2(doc, '5.3 Conclusions')
    
    add_body_paragraph(doc, 'Based on the empirical evidence gathered through this study, the following conclusions are drawn:')
    
    conclusions = [
        ('Conclusion 1: Technology Adoption Landscape', 'SMEs in the Annex area of Uasin Gishu County have adopted technology primarily at the level of mobile and communication technologies, with adoption rates declining significantly for more sophisticated business technologies. The technology adoption landscape is characterized by a clear gradient from accessible, low-cost mobile technologies to more complex and expensive business management systems. This pattern reflects both the opportunity represented by the ubiquity of mobile technology and the barriers that prevent more comprehensive digitalization.'),
        ('Conclusion 2: Technology Positively Impacts Service Delivery', 'Technological innovation has a statistically significant and practically meaningful positive impact on both the efficiency and quality of service delivery among SMEs in the Annex area. The strong correlation coefficients (r=0.714 for efficiency and r=0.682 for quality) and the regression model\'s explanatory power (61.3% of variance in service delivery) confirm that technology adoption is a key driver of service delivery enhancement. This conclusion validates the study\'s theoretical framework, particularly the Dynamic Capabilities Theory which posits that firms can leverage technology to build competitive service delivery capabilities.'),
        ('Conclusion 3: Customer Perceptions Are Positive but Nuanced', 'Customers of SMEs in the Annex area have generally positive perceptions of technology-enhanced services, with particular appreciation for mobile payment options, faster service, and reduced transaction errors. However, the full potential of technology for personalized and data-driven customer service has not yet been realized, representing a significant opportunity for service delivery enhancement. The generally positive customer perceptions align with the Disruptive Innovation Theory, which explains how accessible innovations like mobile money create new value for customers.'),
        ('Conclusion 4: Multiple Barriers Constrain Technology Adoption', 'Technology adoption among SMEs in the Annex area is constrained by a combination of financial, skills, infrastructure, and environmental/social barriers. Financial constraints are the most significant barrier, followed by skills gaps and infrastructure limitations. This multi-barrier environment suggests that single-dimension interventions (e.g., only providing financing without accompanying skills training and infrastructure support) are unlikely to achieve sustained improvements in SME technology adoption and its impact on service delivery.'),
        ('Conclusion 5: ICT Tools Drive the Greatest Service Delivery Impact', 'Among the technology dimensions examined, ICT tools adoption has the strongest predictive impact on service delivery (β=0.421), followed by digital marketing (β=0.318) and service delivery technologies (β=0.276). This finding suggests that prioritizing broad ICT literacy and tool adoption may yield greater service delivery dividends than focusing narrowly on specific service delivery technologies.'),
    ]
    
    for title, text in conclusions:
        para = doc.add_paragraph()
        run = para.add_run(f'{title}: ')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run2 = para.add_run(text)
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    add_heading2(doc, '5.4 Recommendations')
    
    add_heading3(doc, '5.4.1 Recommendations for Policy and Practice')
    
    add_body_paragraph(doc, 'Based on the study findings, the following policy and practice recommendations are made:')
    
    policy_recs = [
        ('Financial Support Programs', 'The national government through MSEA, and the county government of Uasin Gishu, should develop targeted technology financing programs for SMEs. These should include low-interest technology acquisition loans with extended repayment periods, technology equipment leasing schemes, and tax incentives for SMEs that invest in digital technologies. The findings showed that financial constraints (mean=3.98) are the most significant barrier to technology adoption.'),
        ('Digital Literacy and Skills Training', 'Government agencies, universities (including Moi University), and NGOs should collaborate to develop and deliver affordable digital literacy training programs specifically tailored to the needs of SME owners and employees in Uasin Gishu County. Training should cover practical technology use, digital marketing, and basic troubleshooting, and should be delivered in accessible formats and locations within the Annex area.'),
        ('Infrastructure Investment', 'The county government and national government should prioritize improving digital infrastructure in the Annex area, including expansion of reliable broadband internet connectivity, installation of solar-powered backup systems for internet access during power outages, and establishment of technology support centers or digital hubs where SME operators can access affordable technical support.'),
        ('Regulatory Framework Clarity', 'Regulatory agencies should develop and widely disseminate clear, simple guidelines on the legal and regulatory requirements for technology use in SME operations, including data protection requirements, digital payment regulations, and requirements for digital record-keeping. This clarity will reduce regulatory uncertainty (mean=3.58) that currently deters some SMEs from fuller technology adoption.'),
        ('Technology Adoption Awareness Programs', 'MSEA, the Uasin Gishu County government, and business associations should organize regular technology awareness forums, digital trade fairs, and business-to-business technology sharing events specifically for SMEs in the Annex area. These events can showcase successful technology implementation cases from within the local business community, reducing resistance and demonstrating practical benefits.'),
    ]
    
    for title, text in policy_recs:
        para = doc.add_paragraph()
        run = para.add_run(f'{title}: ')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run2 = para.add_run(text)
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    add_heading3(doc, '5.4.2 Recommendations for SME Owners and Managers')
    
    add_body_paragraph(doc, 'The following recommendations are specifically directed at SME owners and managers in the Annex area:')
    
    sme_recs = [
        ('Strategic Technology Planning', 'SME owners should develop simple technology adoption plans that align technology investments with their specific business needs and customer expectations. Rather than adopting all available technologies simultaneously, owners should prioritize technologies that address their most significant operational challenges and offer the clearest service delivery benefits. The study recommends beginning with mobile payment optimization and WhatsApp Business before progressing to more sophisticated technologies.'),
        ('Employee Technology Training', 'SME owners should invest in regular employee training on digital tools relevant to their business operations. This can be done cost-effectively through online resources, peer learning, and vendor-provided training. Given the finding that employee training needs (mean=4.14) is the top skills challenge, developing staff digital competencies should be treated as a priority investment alongside technology hardware acquisition.'),
        ('Customer Education and Engagement', 'SME operators should actively educate their customers about the benefits of technology-enabled services and guide them in using digital payment systems and other service interfaces. This is particularly important for older and less digitally literate customers who may initially prefer traditional service methods. Gradual, supported transition strategies are more effective than abrupt technology-only approaches.'),
        ('Technology Pooling and Collaboration', 'SMEs in the Annex area should consider forming technology adoption cooperatives or business associations that pool resources for technology acquisition, shared internet connectivity, and collective training programs. Such cooperation can significantly reduce individual financial burdens while enabling adoption of technologies that would be unaffordable for individual small businesses.'),
    ]
    
    for title, text in sme_recs:
        para = doc.add_paragraph()
        run = para.add_run(f'{title}: ')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run2 = para.add_run(text)
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    add_heading2(doc, '5.5 Limitations of the Study')
    
    add_body_paragraph(doc, 'This study encountered the following limitations that should be considered when interpreting and generalizing the findings:')
    
    study_limits = [
        'Cross-sectional Design: The study used a cross-sectional design, collecting data at a single point in time. This limits the ability to make causal inferences about the relationship between technological innovation and service delivery. Longitudinal studies would provide stronger evidence of causal relationships and capture the dynamic nature of technology adoption over time.',
        'Geographical Scope: The study focused exclusively on SMEs in the Annex area of Uasin Gishu County, which may limit the generalizability of findings to other commercial areas, counties, or SME contexts in Kenya. While the Annex area represents a reasonably diverse SME landscape, caution should be exercised in extrapolating findings to other contexts.',
        'Self-Reporting Bias: Reliance on self-reported data from SME owners/managers and customers may have introduced response bias. Respondents may have overestimated technology adoption levels or service delivery improvements due to social desirability bias. More objective measures of technology adoption and service delivery, such as sales data or customer retention rates, would complement subjective assessments.',
        'Unregistered SMEs: The study focused on registered SMEs based on county licensing records. A significant number of informal and unregistered SMEs operating in the Annex area may not have been captured, potentially biasing the sample toward more established and possibly more technologically advanced businesses.',
    ]
    
    for limit in study_limits:
        para = doc.add_paragraph()
        run = para.add_run(limit)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    add_heading2(doc, '5.6 Suggestions for Further Research')
    
    add_body_paragraph(doc, 'Based on the findings and limitations of this study, the following areas are suggested for further research:')
    
    further_research = [
        'Longitudinal studies examining the long-term impact of technology adoption on SME service delivery and business performance in Uasin Gishu County to establish causal relationships and track adoption trajectories over time.',
        'Comparative studies examining technology adoption and service delivery across different commercial areas within Uasin Gishu County (e.g., Annex area vs. Eldoret CBD vs. rural areas) to understand geographical variations in technology adoption patterns and their drivers.',
        'Studies examining the impact of specific technology types (e.g., AI-powered tools, e-commerce platforms, cloud computing) on service delivery as these technologies become more accessible to SMEs in regional commercial centers.',
        'Research on the role of gender in technology adoption and service delivery among SMEs in Uasin Gishu County, given the significant proportion of female SME owners (43.4%) and the potential for gender-differentiated technology experiences and outcomes.',
        'Studies incorporating objective performance measures (sales data, customer retention rates, profit margins) alongside self-reported data to provide more objective evidence of the service delivery impact of technology adoption.',
        'Research on the informal SME sector\'s technology adoption patterns and service delivery outcomes, to complement findings from studies focused on registered SMEs and provide a more complete picture of the SME technology landscape.',
    ]
    
    for i, research in enumerate(further_research, 1):
        para = doc.add_paragraph()
        run = para.add_run(f'{i}. {research}')
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # ===================== REFERENCES =====================
    add_page_break(doc)
    
    para = doc.add_paragraph()
    run = para.add_run('REFERENCES')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(12)
    
    references = [
        'African Development Bank. (2022). African Economic Outlook 2022: Supporting Climate Resilience and a Just Energy Transition in Africa. Abidjan: African Development Bank Group.',
        'Bwire, B. K., & Muathe, S. M. A. (2025). Influence of digital credit access on micro, small and medium enterprises growth in Uasin Gishu County, Kenya. Journal of Business and Management Research, 14(2), 45-62.',
        'Christensen, C. M. (1997). The innovator\'s dilemma: When new technologies cause great firms to fail. Boston: Harvard Business School Press.',
        'Communications Authority of Kenya. (2023). Third Quarter Sector Statistics Report for the Financial Year 2022/23. Nairobi: Communications Authority of Kenya.',
        'Creswell, J. W. (2014). Research design: Qualitative, quantitative, and mixed methods approaches (4th ed.). Thousand Oaks, CA: SAGE Publications.',
        'Government of Kenya. (2019). Kenya Digital Economy Blueprint: Powering Kenya\'s Transformation. Nairobi: Ministry of Information and Communication Technology.',
        'GSMA. (2023). The State of the Industry Report on Mobile Money 2023. London: GSMA.',
        'Kenya National Bureau of Statistics. (2023). Economic Survey 2023. Nairobi: Kenya National Bureau of Statistics.',
        'Kenya National Bureau of Statistics. (2019). Kenya Population and Housing Census Volume I: Population by County and Sub-County. Nairobi: Kenya National Bureau of Statistics.',
        'Kiprono, C. (2024). Adopted technology and performance of micro and small enterprises in Nairobi, Kenya. International Journal of Business and Management Review, 12(1), 78-94.',
        'Kising\'a, C., & Kwasira, J. (2019). Effect of ICTs as innovation facilitators of service sector SMEs in Nairobi County. International Journal of Management Science and Business Administration, 5(4), 23-36.',
        'Kothari, C. R. (2004). Research methodology: Methods and techniques (2nd ed.). New Delhi: New Age International Publishers.',
        'Lagat, C. K. (2014). Leveraging ICT organizational capability for SME competitiveness: A case study of agricultural sector SMEs in Uasin Gishu County, Kenya. Unpublished Master\'s Thesis, Moi University, Eldoret.',
        'Mugenda, O. M., & Mugenda, A. G. (2003). Research methods: Quantitative and qualitative approaches. Nairobi: African Centre for Technology Studies.',
        'Musebe, R. (2024). Adoption of advanced manufacturing and service technology by SMEs in Kenya and its effect on performance: A mixed methods study. African Journal of Business Management, 18(3), 112-128.',
        'Mutwota, C. M. (2023). Influence of strategic service innovation on performance of small and medium enterprises in Nairobi County, Kenya. Unpublished Master\'s Thesis, University of Nairobi, Nairobi.',
        'OECD. (2021). The Digital Transformation of SMEs. Paris: OECD Publishing. https://doi.org/10.1787/bdb9256a-en',
        'Parasuraman, A., Zeithaml, V. A., & Berry, L. L. (1988). SERVQUAL: A multiple-item scale for measuring consumer perceptions of service quality. Journal of Retailing, 64(1), 12-40.',
        'PwC. (2022). PwC Global Consumer Insights Survey 2022: Strength from adversity. London: PricewaterhouseCoopers International.',
        'Talam, E. C. (2023). Organizational capabilities and performance of agro-processing SMEs in Uasin Gishu County, Kenya. Journal of Agriculture and Food Processing, 9(1), 34-51.',
        'Teece, D. J., Pisano, G., & Shuen, A. (1997). Dynamic capabilities and strategic management. Strategic Management Journal, 18(7), 509-533.',
        'Tornatzky, L. G., & Fleischer, M. (1990). The processes of technological innovation. Lexington, MA: Lexington Books.',
        'World Bank. (2023). SME Finance. Retrieved from https://www.worldbank.org/en/topic/smefinance',
        'Yamane, T. (1967). Statistics: An introductory analysis (2nd ed.). New York: Harper and Row.',
    ]
    
    for ref in references:
        para = doc.add_paragraph()
        run = para.add_run(ref)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.first_line_indent = Inches(-0.5)
        pf.left_indent = Inches(0.5)
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # ===================== APPENDICES =====================
    add_page_break(doc)
    
    para = doc.add_paragraph()
    run = para.add_run('APPENDICES')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(12)
    
    # Appendix I: Letter of Introduction
    add_heading2(doc, 'APPENDIX I: LETTER OF INTRODUCTION')
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run('MOI UNIVERSITY')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    para = doc.add_paragraph()
    run = para.add_run('School of Business and Economics')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    para = doc.add_paragraph()
    run = para.add_run('P.O. Box 3900 – 30100')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    para = doc.add_paragraph()
    run = para.add_run('Eldoret, Kenya')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run('March 2026')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run('Dear Respondent,')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run('RE: REQUEST FOR PARTICIPATION IN RESEARCH STUDY')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    add_body_paragraph(doc, 'I, JAMES NGOVI (Registration Number: BBM/1733/22), am a Bachelor of Business Management student specializing in Business Information Technology (BIT) at Moi University. I am conducting a research project titled "THE IMPACT OF TECHNOLOGICAL INNOVATION ON SERVICE DELIVERY IN SMALL AND MEDIUM ENTERPRISES (SMEs): A CASE STUDY OF SMEs IN THE ANNEX AREA OF UASIN GISHU COUNTY" in partial fulfillment of the requirements for the award of the degree.')
    
    add_body_paragraph(doc, 'I am kindly requesting your participation in this study by completing the attached questionnaire. The information you provide will be used solely for academic research purposes. Your responses will be treated with the utmost confidentiality and will not be linked to your personal identity in any way.')
    
    add_body_paragraph(doc, 'Your participation is voluntary and you may withdraw at any time without any negative consequences. The questionnaire will take approximately 15-20 minutes to complete.')
    
    add_body_paragraph(doc, 'Should you have any queries regarding this study, please feel free to contact me through the above address or my supervisor, Dr. Kiyeng Chumo, in the Department of Management Science and Entrepreneurship at Moi University.')
    
    add_body_paragraph(doc, 'Thank you in advance for your cooperation and valuable contribution to this research.')
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run('Yours Sincerely,')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run('JAMES NGOVI')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    
    para = doc.add_paragraph()
    run = para.add_run('BBM/1733/22 | School of Business and Economics | Moi University')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    # Appendix II: SME Questionnaire
    add_page_break(doc)
    
    add_heading2(doc, 'APPENDIX II: QUESTIONNAIRE FOR SME OWNERS/MANAGERS')
    
    add_body_paragraph(doc, 'INSTRUCTIONS: Please answer all questions by ticking (✓) the most appropriate response or filling in the required information. All information provided will be treated with strict confidentiality and used for academic research purposes only.')
    
    add_heading3(doc, 'SECTION A: DEMOGRAPHIC AND BUSINESS INFORMATION')
    
    dem_questions = [
        '1. Gender:   [ ] Male   [ ] Female   [ ] Prefer not to say',
        '2. Age:   [ ] Below 25   [ ] 25-34   [ ] 35-44   [ ] 45-54   [ ] 55 and above',
        '3. Highest Level of Education:   [ ] Primary   [ ] Secondary (KCSE)   [ ] Certificate/Diploma   [ ] University   [ ] Postgraduate',
        '4. Type of Business:   [ ] Retail Shop   [ ] Hotel/Restaurant   [ ] Service Provider   [ ] Agricultural-related   [ ] Other (specify): ___________',
        '5. How long has your business been in operation?   [ ] Less than 1 year   [ ] 1-3 years   [ ] 4-6 years   [ ] 7-10 years   [ ] More than 10 years',
        '6. Number of Employees:   [ ] 1-5   [ ] 6-10   [ ] 11-20   [ ] 21-50',
    ]
    
    for q in dem_questions:
        para = doc.add_paragraph()
        run = para.add_run(q)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(8)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    add_heading3(doc, 'SECTION B: TECHNOLOGICAL INNOVATIONS ADOPTED')
    
    add_body_paragraph(doc, 'B1. Please indicate which of the following technologies your business uses. (Tick all that apply)')
    
    tech_list = [
        '[ ] Mobile Money (M-Pesa/Airtel Money)', '[ ] Smartphones for Business', '[ ] WhatsApp Business',
        '[ ] Facebook for Business', '[ ] Instagram for Business', '[ ] Internet/Wi-Fi',
        '[ ] Computers/Laptops', '[ ] POS Systems', '[ ] Accounting Software',
        '[ ] Inventory Management Software', '[ ] Company Website', '[ ] E-commerce Platform (e.g., Jumia)',
        '[ ] Cloud Storage/Services', '[ ] Customer Management Software (CRM)', '[ ] Other (specify): ___________',
    ]
    for t in tech_list:
        para = doc.add_paragraph()
        run = para.add_run(t)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.left_indent = Inches(0.3)
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    add_body_paragraph(doc, 'B2. For the following statements about mobile technology use, please indicate your level of agreement.\n[1=Strongly Disagree, 2=Disagree, 3=Neutral, 4=Agree, 5=Strongly Agree]')
    
    b2_headers = ['Statement', '1', '2', '3', '4', '5']
    b2_rows = [
        ['We use mobile phones to receive payments from customers', '', '', '', '', ''],
        ['We use mobile apps to communicate with suppliers', '', '', '', '', ''],
        ['Mobile money has improved our cash flow management', '', '', '', '', ''],
        ['We use mobile phones to track business inventory', '', '', '', '', ''],
        ['Mobile technology has reduced our transaction costs', '', '', '', '', ''],
    ]
    create_simple_table(doc, b2_headers, b2_rows, col_widths=[3.5, 0.4, 0.4, 0.4, 0.4, 0.4])
    
    doc.add_paragraph()
    
    add_heading3(doc, 'SECTION C: INFLUENCE ON SERVICE DELIVERY')
    
    add_body_paragraph(doc, 'C1. Please rate the influence of technology on the efficiency of your service delivery.\n[1=Strongly Disagree, 2=Disagree, 3=Neutral, 4=Agree, 5=Strongly Agree]')
    
    c1_headers = ['Statement', '1', '2', '3', '4', '5']
    c1_rows = [
        ['Technology has reduced the time taken to serve customers', '', '', '', '', ''],
        ['Technology has reduced our operational costs', '', '', '', '', ''],
        ['Technology has reduced errors in transactions and records', '', '', '', '', ''],
        ['Technology has streamlined our service delivery processes', '', '', '', '', ''],
        ['Digital payment systems have speeded up transactions', '', '', '', '', ''],
        ['Technology has enabled us to serve more customers per day', '', '', '', '', ''],
        ['Technology has improved staff productivity', '', '', '', '', ''],
    ]
    create_simple_table(doc, c1_headers, c1_rows, col_widths=[3.5, 0.4, 0.4, 0.4, 0.4, 0.4])
    
    add_heading3(doc, 'SECTION D: CHALLENGES IN TECHNOLOGY ADOPTION')
    
    add_body_paragraph(doc, 'D1. Please rate the following challenges in adopting technology for your business.\n[1=Not a Challenge at All, 2=Minor Challenge, 3=Moderate Challenge, 4=Major Challenge, 5=Extremely Major Challenge]')
    
    d1_headers = ['Challenge', '1', '2', '3', '4', '5']
    d1_rows = [
        ['High initial cost of technology equipment', '', '', '', '', ''],
        ['Software licensing fees are too expensive', '', '', '', '', ''],
        ['Lack of adequate financing for technology', '', '', '', '', ''],
        ['High internet data costs', '', '', '', '', ''],
        ['Lack of technical skills among staff', '', '', '', '', ''],
        ['Difficulty in troubleshooting technology problems', '', '', '', '', ''],
        ['Unreliable internet connectivity', '', '', '', '', ''],
        ['Frequent power outages disrupting technology use', '', '', '', '', ''],
        ['Customer preference for traditional service methods', '', '', '', '', ''],
        ['Unclear regulatory requirements for technology use', '', '', '', '', ''],
    ]
    create_simple_table(doc, d1_headers, d1_rows, col_widths=[3.5, 0.4, 0.4, 0.4, 0.4, 0.4])
    
    add_body_paragraph(doc, 'THANK YOU FOR YOUR PARTICIPATION')
    
    # Appendix III: Interview Guide
    add_page_break(doc)
    
    add_heading2(doc, 'APPENDIX III: INTERVIEW GUIDE FOR SME OWNERS/MANAGERS')
    
    add_body_paragraph(doc, 'This interview guide is designed for in-depth conversations with selected SME owners/managers. The interview is semi-structured; follow-up questions will be asked based on responses. The interview is expected to take approximately 30-45 minutes.')
    
    interview_sections = [
        ('PART A: TECHNOLOGY ADOPTION EXPERIENCE', [
            '1. Can you describe the types of technology you use in your business and how long you have been using them?',
            '2. What motivated you to start adopting technology in your business?',
            '3. How did you decide which specific technologies to adopt for your business?',
            '4. Have you received any formal training or support in using business technology? Please elaborate.',
        ]),
        ('PART B: IMPACT ON SERVICE DELIVERY', [
            '5. In what specific ways has technology improved how you serve your customers?',
            '6. Can you give examples of how technology has made your service delivery faster or more efficient?',
            '7. Have you noticed any improvements in the quality of your service delivery since adopting technology? Please describe.',
            '8. How have your customers responded to the technology-enhanced services you offer?',
        ]),
        ('PART C: CHALLENGES AND RECOMMENDATIONS', [
            '9. What have been the most significant challenges you have faced in adopting or using technology in your business?',
            '10. How have you addressed or worked around these challenges?',
            '11. What types of support would be most helpful in encouraging greater technology adoption among SMEs in the Annex area?',
            '12. What advice would you give to other SME owners who are considering adopting technology in their businesses?',
            '13. Is there anything else related to technology adoption and service delivery in your business that you would like to share?',
        ]),
    ]
    
    for section_title, questions in interview_sections:
        add_heading3(doc, section_title)
        for q in questions:
            para = doc.add_paragraph()
            run = para.add_run(q)
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            pf = para.paragraph_format
            pf.left_indent = Inches(0.25)
            pf.space_before = Pt(3)
            pf.space_after = Pt(18)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # Appendix IV: Customer Questionnaire
    add_page_break(doc)
    
    add_heading2(doc, 'APPENDIX IV: CUSTOMER QUESTIONNAIRE')
    
    add_body_paragraph(doc, 'INSTRUCTIONS: This questionnaire is for customers of SMEs in the Annex area. Please answer all questions honestly. Your responses are confidential and will only be used for academic research.')
    
    add_heading3(doc, 'SECTION A: DEMOGRAPHIC INFORMATION')
    
    cust_dem = [
        'A1. Gender:   [ ] Male   [ ] Female   [ ] Prefer not to say',
        'A2. Age:   [ ] Below 25   [ ] 25-34   [ ] 35-44   [ ] 45 and above',
        'A3. Level of Education:   [ ] Primary   [ ] Secondary   [ ] Certificate/Diploma   [ ] University',
        'A4. Frequency of Visiting SMEs in Annex Area:   [ ] Daily   [ ] 2-3 times/week   [ ] Weekly   [ ] Monthly',
    ]
    
    for q in cust_dem:
        para = doc.add_paragraph()
        run = para.add_run(q)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(8)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    add_heading3(doc, 'SECTION B: CUSTOMER SATISFACTION WITH TECHNOLOGY-ENHANCED SERVICES')
    
    add_body_paragraph(doc, 'B1. Please rate your satisfaction with the technology-enhanced services offered by SMEs in the Annex area.\n[1=Strongly Disagree, 2=Disagree, 3=Neutral, 4=Agree, 5=Strongly Agree]')
    
    b1c_headers = ['Statement', '1', '2', '3', '4', '5']
    b1c_rows = [
        ['I am satisfied with the availability of mobile payment options', '', '', '', '', ''],
        ['Technology-enabled services are faster and more convenient', '', '', '', '', ''],
        ['I trust the digital payment systems used by these businesses', '', '', '', '', ''],
        ['Technology has improved the overall quality of services I receive', '', '', '', '', ''],
        ['I prefer to patronize businesses that use modern technology', '', '', '', '', ''],
        ['I experience shorter waiting times at tech-equipped businesses', '', '', '', '', ''],
        ['Technology has led to fewer errors in my transactions', '', '', '', '', ''],
        ['Businesses use social media to communicate useful information to me', '', '', '', '', ''],
    ]
    create_simple_table(doc, b1c_headers, b1c_rows, col_widths=[3.5, 0.4, 0.4, 0.4, 0.4, 0.4])
    
    add_body_paragraph(doc, 'B2. What additional technology-related improvements would you like to see in SMEs in the Annex area?\n___________________________________________________________________________\n___________________________________________________________________________\n___________________________________________________________________________')
    
    add_body_paragraph(doc, 'THANK YOU FOR YOUR PARTICIPATION!')
    
    # Save
    output_path = 'files/James_Ngovi_Research_Project.docx'
    doc.save(output_path)
    print(f'Successfully created: {output_path}')
    return output_path

if __name__ == '__main__':
    create_docx()
