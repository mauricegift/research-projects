#!/usr/bin/env python3
"""
Generate research project DOCX for Agnetta Opisa
Moi University - Bachelor of Business Management (Banking and Finance)
Title: THE EFFECT OF FINANCIAL RISK MANAGEMENT ON FINANCIAL PERFORMANCE OF
       SMALL AND MEDIUM SCALE ENTERPRISES IN ELDORET TOWN, KENYA

Adapted from source documents with:
  - 3 IVs: Liquidity Risk, Credit Risk, Equity Risk -> Financial Performance
  - Theories: Prospect Theory, Modern Portfolio Theory, Moral Hazard Theory
  - Target population 230, sample size 146 (Yamane formula)
  - Descriptive statistics (Likert-scale tables with means and SDs)
  - Moi University formatting identical to Calvince's project
"""

import os as _os, sys as _sys
_sys.path.insert(0, _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))
import shutil as _shutil
_SOFFICE = (_shutil.which("libreoffice") or _shutil.which("soffice") or "/nix/store/0pa3zy5lid4paiw9miafpvjkjvlmxfgz-libreoffice-25.2.3.2-wrapped/bin/libreoffice")
_os.chdir(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import io

LOGO_PATH = 'assets/moi_uni_logo.png'
FONT_NAME = 'Liberation Serif'


def set_cell_background(cell, fill_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_page_margins(section, top=1.0, bottom=1.0, left=1.25, right=1.0):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

def _line15(para):
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

def _line1(para):
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

def _sp(para, before=0, after=6):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)

def body(doc, text, indent=False, before=0, after=6):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _sp(para, before, after)
    _line15(para)
    if indent:
        para.paragraph_format.first_line_indent = Inches(0.5)
    return para

def heading_center(doc, text, size=13, bold=True, before=12, after=8):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(para, before, after)
    _line15(para)
    return para

def heading2(doc, text, before=12, after=6):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _sp(para, before, after)
    _line15(para)
    pPr = para._p.get_or_add_pPr()
    ol = OxmlElement('w:outlineLvl')
    ol.set(qn('w:val'), '1')
    pPr.append(ol)
    return para

def heading3(doc, text, before=8, after=4):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _sp(para, before, after)
    _line15(para)
    pPr = para._p.get_or_add_pPr()
    ol = OxmlElement('w:outlineLvl')
    ol.set(qn('w:val'), '2')
    pPr.append(ol)
    return para

def table_caption(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.name = FONT_NAME
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _sp(para, 8, 4)
    _line1(para)
    return para

def source_note(doc, text="Source: Field Survey (2026)"):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.name = FONT_NAME
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _sp(para, 2, 8)
    _line1(para)
    return para

def simple_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].runs[0].font.name = FONT_NAME
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, 'D3D3D3')
    for r, row in enumerate(rows):
        dr = table.rows[r + 1]
        for c, val in enumerate(row):
            cell = dr.cells[c]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].runs[0].font.name = FONT_NAME
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if c == 0 else WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Inches(col_widths[j])
    return table

def add_bar_chart(doc, categories, values, title, xlabel, ylabel, color='steelblue'):
    fig, ax = plt.subplots(figsize=(6, 3.5))
    bars = ax.bar(categories, values, color=color, edgecolor='black', linewidth=0.5)
    ax.set_title(title, fontsize=11, fontweight='bold', pad=8)
    ax.set_xlabel(xlabel, fontsize=10)
    ax.set_ylabel(ylabel, fontsize=10)
    ax.set_ylim(0, max(values) * 1.25)
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max(values) * 0.02,
                f'{val}%', ha='center', va='bottom', fontsize=9, fontweight='bold')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.xticks(fontsize=8, rotation=15, ha='right')
    plt.tight_layout()
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=150, bbox_inches='tight')
    plt.close()
    img_stream.seek(0)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(img_stream, width=Inches(5.5))
    _sp(para, 4, 4)
    _line1(para)
    return para

def insert_section_break(doc, fmt='lowerRoman', start=1, title_page=False):
    p_obj = doc.add_paragraph()
    p = p_obj._p
    pPr = OxmlElement('w:pPr')
    sectPr = OxmlElement('w:sectPr')
    type_el = OxmlElement('w:type')
    type_el.set(qn('w:val'), 'nextPage')
    sectPr.append(type_el)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)
    if title_page:
        titlePg = OxmlElement('w:titlePg')
        sectPr.append(titlePg)
    pPr.append(sectPr)
    p.insert(0, pPr)

def _add_page_field_to_footer(section, field_instruction):
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.clear()
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = field_instruction
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_page_numbers(doc):
    for i, section in enumerate(doc.sections):
        section.different_first_page_header_footer = True
        if i > 0:
            _add_page_field_to_footer(section, ' PAGE ')

def toc_row(doc, title, page, bold=False, indent=0):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.name = FONT_NAME
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(indent * 0.3)
    _sp(p, 0, 2)
    _line1(p)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), '8928')
    tabs.append(tab)
    pPr.append(tabs)
    run2 = p.add_run(f'\t{page}')
    run2.font.size = Pt(11)
    run2.font.bold = bold
    run2.font.name = FONT_NAME
    return p

def cover_line(doc, text, size=12, bold=False, before=0, after=4):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(para, before, after)
    _line1(para)
    return para


def draw_conceptual_framework():
    fig, ax = plt.subplots(figsize=(7.5, 5.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis('off')

    iv_x = 0.3
    iv_w = 3.2
    dv_x = 6.5
    dv_w = 3.2

    iv_boxes = [
        ('Liquidity Risk Management', ['Cash Flow Management', 'Working Capital', 'Short-term Obligations'], 7.8),
        ('Credit Risk Management', ['Loan Default Rates', 'Credit Analysis', 'Interest Rate Computation'], 5.0),
        ('Equity Risk Management', ['Shareholder Interest', 'Debt-Equity Balance', 'Profit Retention'], 2.2),
    ]

    for title, indicators, y_center in iv_boxes:
        rect = plt.Rectangle((iv_x, y_center - 0.8), iv_w, 1.6, linewidth=1.5, edgecolor='black', facecolor='#E8F4FD')
        ax.add_patch(rect)
        ax.text(iv_x + 0.15, y_center + 0.5, title, fontsize=8.5, fontweight='bold', va='center', fontstyle='italic')
        for j, ind in enumerate(indicators):
            ax.text(iv_x + 0.15, y_center + 0.1 - j * 0.35, ind, fontsize=8, va='center')
        ax.annotate('', xy=(dv_x, 5.0), xytext=(iv_x + iv_w, y_center),
                     arrowprops=dict(arrowstyle='->', lw=1.5, color='black'))

    dv_y = 4.0
    dv_h = 2.0
    rect_dv = plt.Rectangle((dv_x, dv_y), dv_w, dv_h, linewidth=1.5, edgecolor='black', facecolor='#FFF3E0')
    ax.add_patch(rect_dv)
    ax.text(dv_x + 0.15, dv_y + 1.6, 'Financial Performance', fontsize=9, fontweight='bold', va='center', fontstyle='italic')
    for j, ind in enumerate(['Profitability', 'Return on Assets', 'Business Growth']):
        ax.text(dv_x + 0.15, dv_y + 1.15 - j * 0.4, ind, fontsize=8, va='center')

    ax.text(iv_x + iv_w / 2, 0.8, 'Independent Variables', fontsize=9, fontweight='bold', ha='center', fontstyle='italic')
    ax.text(dv_x + dv_w / 2, 3.5, 'Dependent Variable', fontsize=9, fontweight='bold', ha='center', fontstyle='italic')

    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=180, bbox_inches='tight', facecolor='white')
    plt.close()
    buf.seek(0)
    return buf


def create_docx():
    doc = Document()
    set_page_margins(doc.sections[0])

    doc.styles['Normal'].font.name = FONT_NAME
    doc.styles['Normal'].font.size = Pt(12)

    for _ in range(2):
        ep = doc.add_paragraph()
        _sp(ep, 0, 0); _line1(ep)

    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(logo_para, 0, 6); _line1(logo_para)
    try:
        logo_para.add_run().add_picture(LOGO_PATH, width=Inches(1.4))
    except Exception:
        pass

    cover_line(doc, 'MOI UNIVERSITY', 14, True, before=4, after=2)
    cover_line(doc, 'SCHOOL OF BUSINESS AND ECONOMICS', 12, True, before=0, after=2)
    cover_line(doc, 'DEPARTMENT OF ACCOUNTING AND FINANCE', 12, True, before=0, after=18)
    cover_line(doc,
        'THE EFFECT OF FINANCIAL RISK MANAGEMENT ON FINANCIAL\n'
        'PERFORMANCE OF SMALL AND MEDIUM SCALE ENTERPRISES\n'
        'IN ELDORET TOWN, KENYA',
        14, True, before=0, after=14)
    cover_line(doc,
        'A RESEARCH PROJECT SUBMITTED IN PARTIAL FULFILMENT FOR THE REQUIREMENTS\n'
        'OF THE AWARD OF BACHELOR OF BUSINESS MANAGEMENT (BANKING AND FINANCE)\n'
        'OF MOI UNIVERSITY',
        12, False, before=0, after=12)
    cover_line(doc, 'BY', 12, False, before=0, after=10)
    cover_line(doc, 'OPISA AGNETTA', 12, True, before=0, after=2)
    cover_line(doc, 'BBM/0038/23', 12, False, before=0, after=12)
    cover_line(doc, 'SUPERVISOR: MR. KIPLAGAT KOIMUR', 12, True, before=0, after=2)
    cover_line(doc, 'Department of Accounting and Finance', 12, False, before=0, after=2)
    cover_line(doc, 'Moi University', 12, False, before=0, after=14)
    cover_line(doc, 'MARCH 2026', 12, True, before=0, after=0)

    insert_section_break(doc, fmt='lowerRoman', start=1, title_page=True)

    heading_center(doc, 'DECLARATION', before=0, after=14)
    body(doc, 'This research project is my original work and has not been presented for the award of any degree in any other university.')
    body(doc, '')
    body(doc, 'Signature: .............................................Date: .............................')
    body(doc, '')
    body(doc, 'OPISA AGNETTA')
    body(doc, 'BBM/0038/23')
    body(doc, '')
    body(doc, "SUPERVISOR'S APPROVAL")
    body(doc, 'This research project has been submitted for examination with my approval as the university supervisor.')
    body(doc, '')
    body(doc, 'Signature: .............................................Date: .............................')
    body(doc, '')
    body(doc, 'MR. KIPLAGAT KOIMUR')
    body(doc, 'Department of Accounting and Finance, Moi University')

    p = heading_center(doc, 'DEDICATION', before=0, after=14)
    p.paragraph_format.page_break_before = True
    para = doc.add_paragraph()
    run = para.add_run(
        'This project is dedicated to my loving parent Madam Lilian Phoebe, Mary Matende, '
        'my sister Rose, Mr. Dhyke Odira, and Sharon who actively supported me both financially '
        'and morally throughout the year. Your constant encouragement and belief in my potential '
        'have been the foundation of my academic journey.')
    run.font.size = Pt(12)
    run.font.name = FONT_NAME
    run.font.italic = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(para, 30, 6)
    _line15(para)

    p = heading_center(doc, 'ACKNOWLEDGEMENT', before=0, after=14)
    p.paragraph_format.page_break_before = True
    body(doc, 'It has been an exciting and instructive study period at Moi University. I feel privileged to have had the opportunity to carry out this study as a demonstration of knowledge gained during the period of studying for my degree. With this, it would be impossible not to remember those who in one way or another, directly or indirectly, have played a role in the realization of this research project. Let me, therefore, thank them all equally.')
    body(doc, 'First, I am indebted to the all-powerful God for all the blessings He showered on me and for being with me throughout the study. I am deeply obliged to my supervisor, Mr. Kiplagat Koimur, for his exemplary guidance and support throughout the process, without whose help this project would not have been a success.')
    body(doc, 'I also appreciate my lecturers in the Department of Accounting and Finance for equipping me with the knowledge and skills necessary to undertake this research. My gratitude further goes to my classmates and friends for their encouragement, cooperation, and moral support for the successful completion of this research project.')
    body(doc, 'Special thanks go to the SME owners and managers in Eldoret Town who willingly participated in this study and provided the necessary information. Your cooperation and openness made the data collection process successful.')
    body(doc, 'Finally, yet importantly, I take this opportunity to express my deep gratitude to my loving mother and Mr. Odira who were a constant source of motivation and for their never-ending support and encouragement. May God bless you all abundantly.')

    p = heading_center(doc, 'ABSTRACT', before=0, after=14)
    p.paragraph_format.page_break_before = True
    body(doc, 'The study sought to establish the effect of financial risk management on financial performance of small and medium scale enterprises in Eldoret Town, Uasin Gishu County, Kenya. The study adopted a descriptive research design with a target population of 230 SMEs operating in Eldoret Town. A sample size of 146 respondents was determined using Yamane\u2019s formula and stratified random sampling was employed. Data was collected through primary instruments such as structured questionnaires and interviews.')
    body(doc, 'Data collected was analyzed using descriptive statistics which involved the use of frequency tables, percentages, means, and standard deviations. The study established that liquidity risk management is positively and significantly related to financial performance of SMEs in Eldoret Town, with insufficient cash flow being the most critical factor (M=4.57, SD=0.91). Credit risk management was also found to be positively and significantly related to financial performance, with high loan default rates being a major concern (M=1.68, SD=1.30). Similarly, equity risk management was positively and significantly related to financial performance of SMEs, with lack of knowledge on debt-equity balance being the most significant factor (M=4.82, SD=0.23).')
    body(doc, 'The study recommends that SME owners seek more training on credit management from financial and educational institutions, consult business experts on resource utilization, and scan the business environment to identify opportunities for attracting investors. The findings contribute to the existing body of knowledge on financial risk management practices among SMEs in Kenya.')
    para = doc.add_paragraph()
    r1 = para.add_run('Keywords: ')
    r1.font.bold = True; r1.font.size = Pt(12); r1.font.name = FONT_NAME
    r2 = para.add_run('Financial Risk Management, Liquidity Risk, Credit Risk, Equity Risk, Financial Performance, SMEs, Eldoret Town')
    r2.font.size = Pt(12); r2.font.name = FONT_NAME
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _sp(para, 0, 6); _line15(para)

    p = heading_center(doc, 'TABLE OF CONTENTS', before=0, after=12)
    p.paragraph_format.page_break_before = True
    toc_entries = [
        ('DECLARATION',                                          'i',   True,  0),
        ('DEDICATION',                                           'ii',  True,  0),
        ('ACKNOWLEDGEMENT',                                      'iii', True,  0),
        ('ABSTRACT',                                             'iv',  True,  0),
        ('TABLE OF CONTENTS',                                    'v',   True,  0),
        ('LIST OF TABLES',                                       'vii', True,  0),
        ('LIST OF FIGURES',                                      'vii', True,  0),
        ('ABBREVIATIONS AND ACRONYMS',                           'viii',True,  0),
        ('OPERATIONAL DEFINITION OF TERMS',                      'ix',  True,  0),
        ('CHAPTER ONE: INTRODUCTION',                            '1',   True,  0),
        ('1.1  Background of the Study',                         '1',   False, 1),
        ('1.1.1  Financial Risk Management',                     '2',   False, 2),
        ('1.1.2  Methods of Dealing with Financial Risk',        '3',   False, 2),
        ('1.1.3  Performance of SMEs',                           '4',   False, 2),
        ('1.2  Statement of the Problem',                        '4',   False, 1),
        ('1.3  Objectives of the Study',                         '6',   False, 1),
        ('1.3.1  General Objective',                             '6',   False, 2),
        ('1.3.2  Specific Objectives',                           '6',   False, 2),
        ('1.4  Research Questions',                              '6',   False, 1),
        ('1.5  Significance of the Study',                       '6',   False, 1),
        ('1.6  Scope of the Study',                              '7',   False, 1),
        ('1.7  Limitation of the Study',                         '7',   False, 1),
        ('CHAPTER TWO: LITERATURE REVIEW',                       '8',   True,  0),
        ('2.1  Introduction',                                    '8',   False, 1),
        ('2.2  Theoretical Framework',                           '8',   False, 1),
        ('2.2.1  Prospect Theory',                               '8',   False, 2),
        ('2.2.2  Modern Portfolio Theory',                       '9',   False, 2),
        ('2.2.3  Moral Hazard Theory',                           '9',   False, 2),
        ('2.3  Conceptual Framework',                            '10',  False, 1),
        ('2.4  Empirical Review',                                '11',  False, 1),
        ('2.4.1  Credit Risk Management',                        '12',  False, 2),
        ('2.4.2  Equity Risk Management',                        '12',  False, 2),
        ('2.4.3  Liquidity Risk Management',                     '13',  False, 2),
        ('2.5  Critique of Existing Literature',                 '13',  False, 1),
        ('2.6  Research Gaps',                                   '14',  False, 1),
        ('2.7  Summary of Literature',                           '14',  False, 1),
        ('CHAPTER THREE: RESEARCH METHODOLOGY',                  '16',  True,  0),
        ('3.1  Introduction',                                    '16',  False, 1),
        ('3.2  Research Design',                                 '16',  False, 1),
        ('3.3  Target Population',                               '16',  False, 1),
        ('3.4  Sample Design',                                   '17',  False, 1),
        ('3.5  Research Instruments',                            '17',  False, 1),
        ('3.6  Data Collection Procedure',                       '18',  False, 1),
        ('3.7  Data Analysis',                                   '18',  False, 1),
        ('3.8  Ethical Considerations',                          '18',  False, 1),
        ('CHAPTER FOUR: DATA ANALYSIS AND PRESENTATION',         '20',  True,  0),
        ('4.1  Introduction',                                    '20',  False, 1),
        ('4.2  Demographic Findings',                            '20',  False, 1),
        ('4.2.1  Response Rate',                                 '20',  False, 2),
        ('4.2.2  Age of Respondents',                            '20',  False, 2),
        ('4.2.3  Length of Business Operation',                   '21',  False, 2),
        ('4.2.4  Level of Education',                            '22',  False, 2),
        ('4.2.5  Sources of Capital',                            '23',  False, 2),
        ('4.3  Descriptive Findings',                            '24',  False, 1),
        ('4.3.1  Liquidity Risk Management',                     '24',  False, 2),
        ('4.3.2  Credit Risk Management',                        '25',  False, 2),
        ('4.3.3  Equity Risk Management',                        '27',  False, 2),
        ('4.3.4  Financial Performance',                         '28',  False, 2),
        ('CHAPTER FIVE: SUMMARY, CONCLUSIONS AND RECOMMENDATIONS','30', True,  0),
        ('5.1  Introduction',                                    '30',  False, 1),
        ('5.2  Summary of Findings',                             '30',  False, 1),
        ('5.2.1  Liquidity Risk Management',                     '30',  False, 2),
        ('5.2.2  Equity Risk Management',                        '30',  False, 2),
        ('5.2.3  Credit Risk Management',                        '31',  False, 2),
        ('5.3  Conclusions',                                     '31',  False, 1),
        ('5.4  Recommendations',                                 '32',  False, 1),
        ('5.5  Suggestions for Further Research',                '32',  False, 1),
        ('REFERENCES',                                           '34',  True,  0),
        ('APPENDICES',                                           '36',  True,  0),
    ]
    for title, pg, bold, indent in toc_entries:
        toc_row(doc, title, pg, bold, indent)

    p = heading_center(doc, 'LIST OF TABLES', before=0, after=12)
    p.paragraph_format.page_break_before = True
    tables_list = [
        ('Table 3.1: Target Population',                         '16'),
        ('Table 3.2: Sample Size',                               '17'),
        ('Table 4.1: Response Rate',                             '20'),
        ('Table 4.2: Age of Respondents',                        '21'),
        ('Table 4.3: Length of Business Operation',              '21'),
        ('Table 4.4: Level of Education',                        '22'),
        ('Table 4.5: Sources of Starting Capital',               '23'),
        ('Table 4.6: Liquidity Risk Management',                 '24'),
        ('Table 4.7: Credit Risk Management',                    '26'),
        ('Table 4.8: Equity Risk Management',                    '27'),
        ('Table 4.9: Financial Performance',                     '28'),
    ]
    for name, pg in tables_list:
        toc_row(doc, name, pg, False, 0)

    heading_center(doc, 'LIST OF FIGURES', before=20, after=12)
    figures_list = [
        ('Figure 2.1: Conceptual Framework',                     '11'),
        ('Figure 4.1: Age Distribution of Respondents',          '21'),
        ('Figure 4.2: Length of Business Operation',             '22'),
        ('Figure 4.3: Level of Education',                       '23'),
        ('Figure 4.4: Sources of Starting Capital',              '24'),
    ]
    for name, pg in figures_list:
        toc_row(doc, name, pg, False, 0)

    p = heading_center(doc, 'ABBREVIATIONS AND ACRONYMS', before=0, after=14)
    p.paragraph_format.page_break_before = True
    abbrevs = [
        ('BBM',   'Bachelor of Business Management'),
        ('ERM',   'Enterprise Risk Management'),
        ('FRM',   'Financial Risk Manager'),
        ('GDP',   'Gross Domestic Product'),
        ('MPT',   'Modern Portfolio Theory'),
        ('MSEA',  'Micro and Small Enterprise Authority'),
        ('PT',    'Prospect Theory'),
        ('SD',    'Standard Deviation'),
        ('SME',   'Small and Medium Enterprise'),
        ('SPSS',  'Statistical Package for Social Sciences'),
    ]
    for abbr, meaning in abbrevs:
        para = doc.add_paragraph()
        r1 = para.add_run(abbr + '\t')
        r1.font.bold = True; r1.font.size = Pt(12); r1.font.name = FONT_NAME
        r2 = para.add_run(meaning)
        r2.font.size = Pt(12); r2.font.name = FONT_NAME
        _sp(para, 1, 2); _line15(para)

    p = heading_center(doc, 'OPERATIONAL DEFINITION OF TERMS', before=0, after=14)
    p.paragraph_format.page_break_before = True
    terms = [
        ('Financial Risk Management', 'A process concerned with the management of uncertainties emanating from the financial market to minimize the risk of making losses from unforeseen events (Namatovu, 2010).'),
        ('Risk Management', 'An ongoing procedure which could help improve operations, gain overall performance objectives, enhance economic balance and in the long run, prevent loss or harm to the entity (Jayathilake, 2012).'),
        ('Financial Performance', 'The ability of the firm to produce new resources out of its daily production procedures over a certain period (Shuying & Mar, 2014).'),
        ('Small and Medium Enterprises', 'An enterprise that has fewer than 250 personnel and has either an annual turnover not exceeding \u20ac50 million or an annual balance-sheet total not exceeding \u20ac43 million (Gao, Sang and Zhang, 2013).'),
        ('Liquidity Risk', 'The possibility that over a specific time period, an organization becomes unable to settle obligations with immediacy (Diebold et al., 2010).'),
        ('Credit Risk', 'The risk of loss arising from a borrower who does not make payments as promised, measured by the five C\u2019s: credit history, capacity to repay, conditions, capital, and collateral.'),
        ('Equity Risk', 'The tactical allocation strategy that seeks to adjust a portfolio\u2019s equity exposure to potentially manage downside risk as volatility changes.'),
    ]
    for term, definition in terms:
        para = doc.add_paragraph()
        r1 = para.add_run(term + ': ')
        r1.font.bold = True; r1.font.size = Pt(12); r1.font.name = FONT_NAME
        r2 = para.add_run(definition)
        r2.font.size = Pt(12); r2.font.name = FONT_NAME
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _sp(para, 2, 4); _line15(para)

    insert_section_break(doc, fmt='decimal', start=1)

    heading_center(doc, 'CHAPTER ONE', before=0, after=2)
    heading_center(doc, 'INTRODUCTION', before=0, after=14)

    heading2(doc, '1.1 Background of the Study')
    body(doc, 'Small and Medium Scale enterprises in Kenya play a major role in economic development through output diversification, employment generation, improvement of local technology, and development of indigenous entrepreneurial capacity. This chapter provides the background of the study, statement of the problem, research questions, objectives of the study, significance of the study, and finally the scope of the study.')
    body(doc, 'Small and medium-sized businesses are increasingly being recognized as essential tools for economic and social development in many nations. In the present era, entrepreneurs face significant demand to generate novel business concepts, products, and improved services that meet the needs of the targeted consumers within particular market coverage. If entrepreneurs fail to upgrade their products and services, their businesses may become unattractive, leading to closure, which is not the primary goal of establishing a company (Rajnoha and Lorincava, 2016). The global economy creates both risks and opportunities for everyone engaged in business, which forces entrepreneurs to provide the best to withstand competition among themselves and the government to survive (Hraskova and Bartosova, 2016).')
    body(doc, 'With the rapid development of a market-oriented economy and the growth of the country\u2019s economy, small businesses face more competition in all sectors. Since financial risks exist everywhere and have a certain influence on enterprise performance, small businesses need to fully understand the actual characteristics of the present situation and the causes of financial risk to survive in market competition. They need to put forward effective prevention and control measures to manage risk and any other uncertainty by reducing the costs associated with cash flow volatility, which will enhance the performance of the enterprises and the development of business (European Association of Craft, 2017).')
    body(doc, 'According to CFI Team (2022), small-scale enterprises are businesses in both formal and informal sectors that employ 1 to 50 workers. They cut across all sectors of employment and provide one of the most prolific sources of employment creation, income generation, and poverty reduction. Within Kenya, SMEs cut across the economy and sustain the majority of households economically. They form a base upon which businesses are established, growing to maturity in size, where local companies operating today expand their operations and start employing more than 100 workers.')
    body(doc, 'The ability of small and medium-sized enterprises to appropriately identify various forms of risk and make sound decisions in tackling those risks will undoubtedly increase SMEs\u2019 profitability and economic growth in any country. By definition, SMEs are non-subsidiary, independent firms that employ less than a given number of employees. This number varies across countries, with the most frequent upper limit designating 250 employees, some at 200, while the United States considers SMEs to include firms with fewer than 500 employees (European Union, 2005; OECD, 2005).')
    body(doc, 'Hazards are uncertainties that are inevitably present in all business establishments that are established with the primary objective of generating profit. Financial institutions, in particular, are exposed to different types of risks, including credit risks, interest risks, liquidity risks, market risks, foreign risks, and operational risks.')

    heading3(doc, '1.1.1 Financial Risk Management')
    body(doc, 'Non-Business Risks are those risks that are not related to the business operations and are usually beyond the control of the company, such as natural disasters or political instability. Financial Risk, on the other hand, is the risk associated with the financial structure of the company, including its investments, financing, and liquidity.')
    body(doc, 'To mitigate Financial Risk, companies hire Financial Risk Managers (FRMs) who assess, monitor, and manage financial risks. The role of an FRM is to identify potential risks and develop strategies to manage them effectively. This involves analyzing financial data, monitoring market trends, and assessing the impact of external factors on the company\u2019s financial position. By implementing effective risk management strategies, companies can reduce their exposure to financial risks and protect their financial stability.')
    body(doc, 'In summary, financial risk management is a critical aspect of business management that requires careful attention and proactive measures. By understanding the different types of risks and hiring qualified professionals to manage them, companies can minimize their financial exposure and ensure long-term success.')
    body(doc, 'Business risks are taken by business enterprises in order to maximize shareholder value and profits. Companies undertake high costs in marketing to launch a new product in order to gain higher sales. The manager or the entrepreneur has the ability to focus on all aspects of the change together with taking all effective measures to control the risk within the firm (Shuying and Mei, 2014). Financial risk management has received increased attention over the past years which helps in the risk identification, assessment, and management process which is part of entrepreneurial strategic development in performance of its business which must be planned and integrated to control risk and monitor all kinds of risk together with dependence which businesses are exposed to. Small business enterprises are exposed to the following risks: operation risk, credit risk, contingent risk, inflation risk, liquidity risk, reputation risk, and business risk.')
    body(doc, 'Globally, financial risk management has been identified by major firms for profitability and sustainability. Berrell (2017) described an approach to Enterprise Risk Management (ERM) in China. A literature search identified areas of generic business risk in China. They used an Expert Panel as a Nominal Group Technique and Linear Rating Scales to identify and prioritize business risks in the Chinese economy. The approach demonstrated how an expedient and quick assessment of risk can occur using intellectual capital as a principal throughout.')
    body(doc, 'Financial risk appears in various forms, as external risk that is related to its external financial environment in which business operates, and lastly internal risk where the business itself is a source of risk. In this study, the following types of risk which are associated within SMEs are examined: Credit Risk, Liquidity Risk, and Equity Risk.')

    heading3(doc, '1.1.2 Methods of Dealing with Financial Risk')
    body(doc, 'In order to effectively oversee the financial hazards that a business may face, it is essential for small business proprietors to establish strategies that can safeguard them from the negative impact of uncertainties in the business environment (Rejda, 2013). Failing to proactively plan and underestimating the detrimental consequences of such risks can prove to be fatal. Small business owners must do everything in their power to minimize losses and maximize profits (Mwangi, 2015).')
    body(doc, 'To better manage risks, Fatai (2010) suggests that business owners should consider obtaining insurance coverage, implementing safety measures against fire and personal accidents that can be costly to the business, and investing in more than one venture to diversify risk and avoid issues with liquidity, cash flow, and bad debts. Investing in multiple ventures can provide an additional source of income and offer a better understanding of the most profitable sectors of the economy, particularly for novice business owners (Rejda, 2012).')

    heading3(doc, '1.1.3 Performance of SMEs')
    body(doc, 'SMEs are acknowledged as a crucial driver of economic growth in both developed and developing nations, and they represent a fundamental component of a country\u2019s economy (Abeywardhana, 2017). Small business owners play a pivotal role in boosting employment, augmenting Gross Domestic Product (GDP), promoting innovation, and fostering other economic endeavors. Additionally, they generate substantial wealth and support the livelihoods of individuals who conceive of ideas and transform them into business ventures.')
    body(doc, 'The correlation between financial risk and performance of SMEs arises when one is interested in understanding how operational risk, credit risk, inflation risk, volatility risk, liquidity risk, reputation risk, and business risk impact the performance of SMEs and how to mitigate these risks. To minimize uncertainty, companies must have risk management strategies in place. Effective financial risk management practices have played a crucial role in promoting the growth of SMEs.')
    body(doc, 'SMEs encounter challenges due to their intricate operations. In contrast to larger corporations, SMEs frequently lack the necessary resources, such as personnel, databases, and specialized knowledge, to execute standardized and structured risk management. This may be due to the fact that many SMEs do not conduct adequate analyses to identify their risks and lack mechanisms to mitigate them, leading to financial distress and eventual bankruptcy. In Kenya, SMEs are experiencing various financial risks in their business activities, causing them to incur losses and become unsustainable. Issues in the SME sector can create uncertainties and negatively impact the performance of SMEs.')

    heading2(doc, '1.2 Statement of the Problem')
    body(doc, 'Notwithstanding the endeavors of the government to assist Small and Medium-sized Enterprises (SMEs) in mitigating the factors that lower their performance, there exist additional elements that hinder their progress. In the present-day dynamic business environment, risks and hazards are on the rise and effective risk management has become indispensable for any enterprise, particularly for SMEs, which, due to their limited resources, intense competition, absence of support, and frail infrastructure, are more susceptible to the harmful effects.')
    body(doc, 'SMEs encounter obstacles due to a lack of supportive governance policies. They face issues such as harassment from local authorities, an unsupportive tax system, and exposure to corruption, as there is no legal framework to protect their interests. In recent years, the Kenyan government has initiated programs to promote SMEs, such as the formation of the Micro and Small Enterprise Authority (MSEA) under the Micro and Small Enterprise Act No. 55 of 2012. The MSEA is responsible for formulating and reviewing policies and programs for Micro, Small, and Medium Enterprises.')
    body(doc, 'Intense competition reduces the market share of SMEs by limiting the number of customers available. As SMEs depend on buyers who are customers, a reduction in customers leads to a decrease in profits and an increase in expenses. Another factor is infrastructure, as Kenya has not invested sufficiently in infrastructure in the past five years, resulting in poor road facilities, low electricity and water supply, and other challenges that have forced many SMEs to close.')
    body(doc, 'To address these challenges, the government has developed strategies to expedite the process of starting a business for small firms. These include enforcing legislation on local products for public projects, establishing \u2018Buy Kenya, Build Kenya\u2019 policies in public procurement, supporting research and development, and increasing funding for small business entrepreneurs. The government should also encourage SMEs to participate in the Entrepreneurial Development Program to improve their skills. Additionally, the government should support SMEs to minimize the effects of financial risk and enhance their operations not only in Eldoret Town but also throughout Kenya.')
    body(doc, 'The primary objective of this research was to identify solutions to the factors that affect the financial risk of small and medium-scale enterprises. The government should train entrepreneurs in management and workforce areas to help them address the challenges that SMEs face while running their businesses.')
    body(doc, 'The purpose of this study is to find out the effects of financial risk management on performance of small and medium scale enterprises in Eldoret Town, Kenya.')

    heading2(doc, '1.3 Objectives of the Study')
    heading3(doc, '1.3.1 General Objective')
    body(doc, 'The main objective of this study was to determine the effects of financial risk management on financial performance of small and medium scale enterprises in Eldoret Town, Kenya.')
    heading3(doc, '1.3.2 Specific Objectives')
    body(doc, 'i) To determine the extent to which credit risk management affects the financial performance of SMEs in Eldoret Town, Kenya.')
    body(doc, 'ii) To determine the effect of equity risk management on the financial performance of SMEs in Eldoret Town, Kenya.')
    body(doc, 'iii) To determine the effect of liquidity risk management on the financial performance of SMEs in Eldoret Town, Kenya.')

    heading2(doc, '1.4 Research Questions')
    body(doc, 'The study sought to answer the following research questions:')
    body(doc, 'i. What is the effect of credit risk management on the financial performance of SMEs in Eldoret Town, Kenya?')
    body(doc, 'ii. What is the effect of equity risk management on the financial performance of SMEs in Eldoret Town, Kenya?')
    body(doc, 'iii. What is the effect of liquidity risk management on the financial performance of SMEs in Eldoret Town, Kenya?')

    heading2(doc, '1.5 Significance of the Study')
    body(doc, 'The research holds immense significance for a diverse group of individuals. It will equip aspiring small and medium-sized entrepreneurs with the ability to evaluate the financial hazards linked with small-scale enterprises. Additionally, financial institutions seeking to target small and medium-sized enterprises can benefit from the study by gaining insights on how to incorporate all the indispensable factors essential for the sustainable growth of SMEs.')
    body(doc, 'Ultimately, the results can be utilized by policymakers in the SME sector, not just in Kenya but across the globe. It will empower the government to devise better approaches to facilitate the growth of SMEs. Future researchers and scholars will benefit from this study as it adds to the existing body of knowledge on financial risk management and SME performance.')

    heading2(doc, '1.6 Scope of the Study')
    body(doc, 'The aim of this study is to investigate the impact of financial risk management on the financial performance of Small and Medium Enterprises (SMEs) in Eldoret Town. The study focused on three independent variables: Liquidity Risk, Credit Risk, and Equity Risk. The dependent variable was financial performance, measured through income performance and return on assets. The scope of the study was based on SMEs in Kenya, with a specific focus on Eldoret Town, which is a growing town that relies on SMEs for job creation. Due to budget and time constraints, the study is assumed to reflect the current state of Kenyan SMEs. The relationship between the study variables was supported by Prospect Theory and Modern Portfolio Theory, which constituted the theoretical and temporal scope of the study.')

    heading2(doc, '1.7 Limitation of the Study')
    body(doc, 'The results of this study largely depend on primary information analyses. Therefore the study results are subjected to the limitations of the SMEs\u2019 financial statements as reported to the general public which were under custody of the Risk Management Department. The data available was only for the period 2021 to 2025. The study had the limitation of not having access to all targeted data. Additionally, some respondents were reluctant to provide information due to privacy concerns, which may have affected the comprehensiveness of the findings.')

    p = heading_center(doc, 'CHAPTER TWO', before=0, after=2)
    p.paragraph_format.page_break_before = True
    heading_center(doc, 'LITERATURE REVIEW', before=0, after=14)

    heading2(doc, '2.1 Introduction')
    body(doc, 'This chapter covers the theoretical review, empirical review, conceptual framework, study variables, the research summary, and the research gap. The chapter examines the relevant theoretical and empirical literature related to financial risk management and its effect on the financial performance of small and medium scale enterprises.')

    heading2(doc, '2.2 Theoretical Framework')
    body(doc, 'This study is anchored on three theories: Prospect Theory, Modern Portfolio Theory, and Moral Hazard Theory. The assumptions of these theories are used as a foundation in explaining the understanding of the study.')

    heading3(doc, '2.2.1 Prospect Theory')
    body(doc, 'This theory is based on the belief that individuals are in a dilemma on the approaches they can use when they are faced with uncertainty (Tversky & Kahneman, 2009). According to James Chen (2022), the theory assumes that losses and gains are valued differently, and thus individuals make decisions based on perceived gains instead of losses. Also known as loss-aversion theory, the general concept is that if two choices are put before an individual, both equal in terms of possible losses, the option presented in terms of potential gains will be chosen.')
    body(doc, 'It belongs to the behavioral economic subgroup describing how individuals make choices between probabilistic alternatives where risk is involved and the probability of different outcomes is unknown. This theory was formulated in 1979 and further developed in 1992 by Amos Tversky and Daniel Kahneman, deeming it more psychologically accurate of how decisions are made when compared to the expected utility theory.')
    body(doc, 'Based on the theoretical beliefs, it is deduced that a person\u2019s course of action in times of uncertainty can be classified into two (Rogerson, 2013). The immediate action a person needs to do is to understand the problem at hand, to be able to explain the problem, identify the potential source of the problem, and solutions that are applicable to solve the problem (Rogerson, 2013). The secondary action that a person must take is considered by determining the significance of the choices decided to solve the problem (Malz, 2011).')
    body(doc, 'This theory is relevant to this study because SME owners in Eldoret Town face financial uncertainties daily and must make decisions about how to manage liquidity risk, credit risk, and equity risk. The theory helps explain why some entrepreneurs may be risk-averse in their financial management approaches while others take calculated risks to improve their financial performance.')

    heading3(doc, '2.2.2 Modern Portfolio Theory')
    body(doc, 'This is a theory accredited to Markowitz (1959) which places a value on financial risk and gains from investment by an organization. MPT theory asks the investors to consider how much the risk of one investment can impact their entire portfolio. Investors are of two types: risk-takers who can invest in conditions of uncertainty, and risk-averse investors who are always careful about where to invest their money in portfolios where they are sure they will get a good return (Panigrahi, 2012).')
    body(doc, 'Through the assumptions of this theory, it is advisable to invest in more than one portfolio to spread the chances of avoiding risk in the market (Tarantino, 2010). It is expected that not all sectors of the economy can make losses at the same prevailing situations. As one sector suffers, the business might be booming in another sector for their environment is different, hence a chance of making an income in one or more avenues (Omisore, Munirat & Nwufo, 2012).')
    body(doc, 'The theory works well for risk-takers who can invest in more than one portfolio even at a time when the economy is not doing very well, with the hope of maximizing profits as the economy becomes stable. The theory assumes that every investor wants to achieve the highest possible long-term returns while minimizing market risk. The theory promotes the spread of money across different asset classes and investments.')
    body(doc, 'This theory is relevant to this study as it explains how SME owners can diversify their business operations and financial portfolios to manage equity risk and liquidity risk effectively, thereby improving their overall financial performance.')

    heading3(doc, '2.2.3 Moral Hazard Theory')
    body(doc, 'A moral hazard is where one party is responsible for the interests of another but has an incentive to put his or her own interests first. For example, one might take risks that someone else will have to bear. Moral hazards such as these are a pervasive and inevitable feature of the financial system and of the economy more generally (Wanjohi, Wanjohi & Ndambiri, 2017).')
    body(doc, 'Krugman (2009) described moral hazard as \u201cany situation in which one person makes the decision about how much risk to take, while someone else bears the cost if things go badly.\u201d According to Busato and Coletta (2017), a corporation\u2019s owners, namely the shareholders, do not have control over the day-to-day operations of their company. This separation is the consequence of asymmetric information within corporations.')
    body(doc, 'They further state that the contract between the shareholders and the managers leaves the latter too much discretion on how to run the business, because the managers, unlike shareholders, have the skills and knowledge to do so. This leads, in more than one case, to a moral hazard problem, when managers may act in their own interests instead of shareholders\u2019 best interests. This is a widespread phenomenon, occurring not only within corporations but in different realities of the modern economic structure.')
    body(doc, 'The study benefited from this theory through identification of various stakeholder and owner-related issues that bring about the financial risks experienced in a firm, particularly in the context of equity risk management where shareholders and managers may have conflicting interests affecting SME financial performance.')

    heading2(doc, '2.3 Conceptual Framework')
    body(doc, 'The conceptual framework illustrates the association existing between the independent variables and the dependent variable. The independent variables consist of liquidity risk management, credit risk management, and equity risk management. The dependent variable of the study is the financial performance of SMEs in Eldoret Town.')
    cf_buf = draw_conceptual_framework()
    cf_para = doc.add_paragraph()
    cf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cf_run = cf_para.add_run()
    cf_run.add_picture(cf_buf, width=Inches(5.5))
    _sp(cf_para, 4, 4); _line1(cf_para)
    fc = doc.add_paragraph()
    fc.add_run('Figure 2.1: Conceptual Framework').font.italic = True
    fc.runs[0].font.size = Pt(10); fc.runs[0].font.name = FONT_NAME
    fc.alignment = WD_ALIGN_PARAGRAPH.CENTER; _sp(fc, 2, 8); _line1(fc)

    heading2(doc, '2.4 Empirical Review')
    body(doc, 'Van den Boom (2019) analyzed financial risk management practices applied by Dutch Small and Medium-sized businesses. Specifically, the study targeted the techniques utilized by SMEs to manage credit risks, exchange rate risks, liquidity risks, and interest rate risks. The study adopted a descriptive research design and data were collected from a population of 97 Dutch SMEs via questionnaires for the period 2013-2017. Findings showed that SMEs focus on the risk process rather than establishing an organizational structure to embed the process.')
    body(doc, 'Belas, Kljucnikov, Vojtovic and Sobekova-Majkova (2015) performed a survey of entrepreneurs\u2019 opinions in 2015 within the Czech Republic. The study adopted a survey method and involved 1,000 entrepreneurs in the capital city. Data were collected through the use of questionnaires and interviews for evaluation. The study found that three quarters of entrepreneurs within the SME segment in the Czech Republic understand the extreme impact of financial risk, but at the same time their potential to properly manage these risks remains limited.')
    body(doc, 'Rotich and Wahome (2018) examined the influence of financial management practices on the financial performance of youth groups funded by the government in Kuresoi South, Kenya. The target population was 96 top officials from 32 government-funded youth organizations working in Kuresoi. A descriptive study was adopted in the form of a survey while a census survey was employed. Data was collected using questionnaires and analyzed using descriptive statistics. The findings suggest a statistically significant positive association between liquidity management and financial performance.')

    heading3(doc, '2.4.1 Credit Risk Management')
    body(doc, 'Customer credit risk can be measured by the five C\u2019s: credit history, capacity to repay the loan, conditions, capital, and associated collateral. Consumers posing high credit risk usually end up paying high interest rates on loans. Credit risk management usually enables the business to maximize sales while managing its risk exposure.')
    body(doc, 'Ashley Muteso (2021) carried out a study revealing the role of SMEs in the economy. She developed a novel model based on the original KMS model with improved parameters to measure credit risk. Patrick Behr (2007) estimated a logit scoring model for prediction of the probability of default by SMEs using a unique data set on SME loans. According to the research, the model helps entrepreneurs to maintain their banks\u2019 pricing behavior and reduce information asymmetry between lenders and borrowers.')
    body(doc, 'Jaroslav Belas (2018) conducted a survey on the business environment of SMEs in 2015 with a sample of 1,141 respondents. They confirmed that attitudes towards financial risks are determined by factors including credit risk management and sufficient funds for SMEs. They also confirmed that companies that run their business with cash were more aware of the impact of the crisis on their business finances.')

    heading3(doc, '2.4.2 Equity Risk Management')
    body(doc, 'Equity risk is the tactical allocation strategy that seeks to adjust a portfolio\u2019s equity exposure to potentially manage downside risk as volatility changes. Pieter Vijn (2010) investigated the association between brand equity measures and business unit financial performance. The study provided important insights for both managers and designers of performance measurement systems.')
    body(doc, 'Irene K. Njugi studied the capital structure as a mix of debt and equity. According to her findings, managers used various combinations of debt and equity that increase the net worth of the business while at the same time reducing the cost of obtaining finance. The study adopted a descriptive survey research design with a target population of 300 SMEs from which a sample size of 60 SMEs was drawn. The study revealed that SMEs had a greater preference for contributions from friends and ploughing back profit as a source of equity finance.')

    heading3(doc, '2.4.3 Liquidity Risk Management')
    body(doc, 'Esther W. Waweru (2017) explained that financial management is very important to ensure that small and medium enterprises remain solvent. Meeting financial obligations shows that SMEs have a right to continuity. She suggested that financial management challenges are consistent with record keeping and borrowing regulations and organizing financial analysis. In this regard, the study sought to identify the challenges faced by SMEs operating in Kenya.')
    body(doc, 'Jagongo (2021) sought to investigate the effect of liquidity risk management on financial performance of state-owned enterprises in Kenya. He adopted a desktop methodology using secondary data where he relied on already published studies, reports, and statistics which were easily accessed through online journals. The study found a positive and significant relationship between liquidity risk management and financial performance.')
    body(doc, 'Sunday (2020) focused on establishing the extent of concern of consumer goods companies regarding cash defensive incentives, long-term debts, and quick ratios for the purpose of turning around performance. Data was obtained from annual reports and accounts of studied companies. The study recommends that consumer goods companies should incorporate clear liquidity risk management appropriate to their strategic policy framework.')

    heading2(doc, '2.5 Critique of Existing Literature')
    body(doc, 'The evaluation of the literature means that most of the previous studies have analyzed the effects of financial indicators based on other indicators. This study is based on the financial risks and the financial results of the SMEs. Many studies have looked at managing financial risks but have not comprehensively addressed the relationship between financial risk management and financial performance indicators. The results of previous research done in Kenya have largely focused on financial institutions. Most project researchers in Kenya have shown that credit risks have a significant effect on SMEs. There is a lack of research on other types of risks, such as equity risks and liquidity risks, and their combined effect on SME financial performance.')

    heading2(doc, '2.6 Research Gaps')
    body(doc, 'It is evident that research on small and medium scale enterprises\u2019 financial risks has been done but not in a comprehensive approach. Most of the literature reviewed indicated that previous researchers only concentrated on credit risks, leaving out the components of equity risk and liquidity risk management.')
    body(doc, 'The current study has a wider scope by covering additional important variables of liquidity risk management and equity risk management that were omitted by previous studies. This makes the study more comprehensive. From the survey of relevant literature, it has been found that there are few studies specific to Kenya on the link between financial risk management and the performance of SMEs, particularly in Eldoret Town. This study therefore intended to fill these pertinent gaps in literature by establishing the effect of financial risk management on the financial performance of small and medium scale enterprises in Eldoret Town, Kenya.')

    heading2(doc, '2.7 Summary of Literature')
    body(doc, 'It is evident from the reviewed studies that the topic of financial risk management in SMEs has gained a lot of interest from local and international scholars. Ombworo (2014) sought to establish the relationship between liquidity and profitability of SMEs in Kenya and established that liquidity has a positive but insignificant effect on profitability. Van den Boom (2019) analyzed financial risk management practices carried out by Dutch SMEs and established that SMEs focus on the risk process rather than establishing an organizational structure to embed the process.')
    body(doc, 'The reviewed literature demonstrates that financial risk management\u2014encompassing liquidity risk, credit risk, and equity risk\u2014plays a critical role in determining the financial performance of SMEs. However, most studies have been conducted in different geographical contexts, necessitating a localized study in Eldoret Town, Kenya to understand the specific dynamics affecting SMEs in this region.')

    p = heading_center(doc, 'CHAPTER THREE', before=0, after=2)
    p.paragraph_format.page_break_before = True
    heading_center(doc, 'RESEARCH METHODOLOGY', before=0, after=14)

    heading2(doc, '3.1 Introduction')
    body(doc, 'Research refers to the process of answering unanswered questions to solve a given problem. Research methodology is the specific procedures used to identify, select, process, and analyze information of a given topic. This chapter presents the research design, target population, sample design, data collection instruments, data collection procedure, data analysis methods, and ethical considerations.')

    heading2(doc, '3.2 Research Design')
    body(doc, 'The study adopted a descriptive research design. A descriptive design is relevant in enabling the researcher to identify the relationship between variables, in this case the effects of financial risk management on the financial performance of small and medium scale enterprises in Eldoret Town, Kenya. The descriptive design was appropriate because it enabled the researcher to describe the characteristics of the population under study and establish the relationship between variables without manipulating them.')

    heading2(doc, '3.3 Target Population')
    body(doc, 'Target population refers to the population from which the information desired will be obtained (Ngechu, 2004). The study focused on the performance of SMEs in Eldoret Town. A total target population of 230 SME entrepreneurs was identified. The census technique was used which helps to improve the extent of accuracy and reliability (Mugenda, 2003). The target population was categorized into four strata as shown in Table 3.1 below.')
    doc.add_paragraph()
    table_caption(doc, 'Table 3.1: Target Population')
    simple_table(doc,
        ['Sector', 'Target', 'Percentage'],
        [
            ['Ethical and Professional Services', '30', '13.04%'],
            ['Education, Health, Private and Entertainment', '40', '17.39%'],
            ['Forestry, Natural Extracts and Agriculture', '50', '21.74%'],
            ['Storage, Communication and Transport', '110', '47.83%'],
            ['Total', '230', '100%'],
        ],
        col_widths=[3.0, 1.0, 1.2])
    source_note(doc, 'Source: Field Research (2026)')

    heading2(doc, '3.4 Sample Design')
    body(doc, 'According to Mugenda and Mugenda (2003), a sample is a small group chosen from a highly accessible population. Sampling is the process of selecting a specific number of individuals under research where the persons selected become a representative of the larger group. Stratified random sampling was used to accomplish the desired representation from diverse subgroups in the population.')
    body(doc, 'The sample size was calculated using Yamane\u2019s (1967) formula:')
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = formula_para.add_run('n = N \u00f7 (1 + Ne\u00b2)')
    fr.font.size = Pt(12); fr.font.bold = True; fr.font.name = FONT_NAME
    _sp(formula_para, 6, 4); _line1(formula_para)
    body(doc, 'Where: n = sample size, N = target population (230), e = margin of error (0.05)')
    body(doc, 'n = 230 \u00f7 (1 + 230 \u00d7 0.05\u00b2) = 230 \u00f7 (1 + 0.575) = 230 \u00f7 1.575 = 146')
    body(doc, 'The sample size for each stratum was proportionally allocated as shown in Table 3.2 below.')
    doc.add_paragraph()
    table_caption(doc, 'Table 3.2: Sample Size')
    simple_table(doc,
        ['Sector', 'Target', 'Sample'],
        [
            ['Ethical and Professional Services', '30', '19'],
            ['Education, Health, Private and Entertainment', '40', '25'],
            ['Forestry, Natural Extracts and Agriculture', '50', '32'],
            ['Storage, Communication and Transport', '110', '70'],
            ['Total', '230', '146'],
        ],
        col_widths=[3.0, 1.0, 1.0])
    source_note(doc, 'Source: Field Research (2026)')

    heading2(doc, '3.5 Research Instruments')
    body(doc, 'The research used both primary and secondary data. Primary data was collected using personally administered questionnaires and interviews to accumulate information from SME owners for analysis. The questionnaire consisted of two sections. The first section was designed to collect demographic information about the respondents. The second section was designed to gather data on the research objectives concerning financial risks.')
    body(doc, 'The questionnaire was designed to include both structured and unstructured questions. The structured questions were used in an effort to conserve time and money as well as to facilitate easier analysis as they were to be in immediately usable form. The unstructured questions were used to encourage the respondent to give an in-depth and felt response without feeling held back in revealing any information.')
    body(doc, 'The secondary data was received from annual reviews and financial statements of the sampled enterprises determined at their registered places of business and/or with their auditors. A data collection form was designed to record details of money spent, acquired, and profit generated.')

    heading2(doc, '3.6 Data Collection Procedure')
    body(doc, 'The data was collected in the month of February 2026. The researcher went to visit the firms in Eldoret Town. Data was collected using interviews and questionnaires whereby both small and medium enterprises were visited to give their views. This study collected quantitative data using a self-administered questionnaire.')
    body(doc, 'Nevertheless, where it proved difficult for the respondents to complete the questionnaires immediately, the questionnaire was left with the respondents and picked later. The questionnaires were hand-delivered and administered at the respondents\u2019 place of business to ensure objective response and reduce non-response rate. The respondents were assured of confidentiality of their names and responses and that the responses were to be used purely for academic purposes. Each questionnaire was coded and only the researcher had knowledge of which person responded.')

    heading2(doc, '3.7 Data Analysis')
    body(doc, 'The descriptive statistical tools were used to help the researcher describe the data and determine the extent of each variable. Analysis was done quantitatively and qualitatively by use of descriptive statistics. This included frequency distribution, tables, percentages, means, and standard deviations. The data was analyzed with the aid of SPSS software to generate the required statistical outputs.')

    heading2(doc, '3.8 Ethical Considerations')
    body(doc, 'A research letter from Moi University was employed to seek permission to conduct the study. The permit was sought from the relevant authorities in Eldoret Town. All questionnaires used in this study were properly referenced and cited. Respondents were informed about the purpose of the study and their participation was voluntary. Confidentiality of the data collected was maintained throughout the study, and the information was used solely for academic purposes.')

    p = heading_center(doc, 'CHAPTER FOUR', before=0, after=2)
    p.paragraph_format.page_break_before = True
    heading_center(doc, 'DATA ANALYSIS AND PRESENTATION', before=0, after=14)

    heading2(doc, '4.1 Introduction')
    body(doc, 'Chapter four contains data analyzed and presented in descriptive statistics. The findings are presented in tables, summaries, and charts to give a clear illustration of the relationships established between the independent variables (liquidity risk management, credit risk management, and equity risk management) and the dependent variable (financial performance of SMEs).')

    heading2(doc, '4.2 Demographic Findings')
    body(doc, 'The study sought to establish the general information of the study participants. Of particular interest were the respondents\u2019 age, length of business operation, level of education, and sources of starting capital.')

    heading3(doc, '4.2.1 Response Rate')
    body(doc, 'The study targeted 146 respondents based on the sample size calculation. Out of the 146 questionnaires distributed, 100 were filled and returned for analysis. The table below provides the frequency and percentage of the respondents.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.1: Response Rate')
    simple_table(doc,
        ['Response', 'Frequency', 'Percentage'],
        [
            ['Returned', '100', '68.5%'],
            ['Not Returned', '46', '31.5%'],
            ['Total', '146', '100%'],
        ],
        col_widths=[2.5, 1.2, 1.2])
    source_note(doc)
    body(doc, 'The response rate of 68.5% is considered adequate for analysis according to Mugenda and Mugenda (2003), who suggested that a response rate of 60% and above is adequate for analysis and reporting.')

    heading3(doc, '4.2.2 Age of Respondents')
    body(doc, 'The respondents were asked to indicate their age brackets. The results are presented in Table 4.2 below.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.2: Age of Respondents')
    simple_table(doc,
        ['Age Bracket', 'Frequency', 'Percentage'],
        [
            ['25\u201330 years', '15', '15%'],
            ['30\u201335 years', '12', '12%'],
            ['35\u201340 years', '65', '65%'],
            ['40\u201345 years', '5', '5%'],
            ['Above 50 years', '3', '3%'],
            ['Total', '100', '100%'],
        ],
        col_widths=[2.5, 1.2, 1.2])
    source_note(doc)
    body(doc, 'Table 4.2 above shows that the majority of the respondents at 65% were aged 35\u201340 years, followed by 15% aged 25\u201330 years, 12% aged 30\u201335 years, 5% aged 40\u201345 years, and lastly 3% aged above 50 years. According to the findings, participants involved in small and medium enterprises prefer self-employment to create a fortune and a stable income for their families.')
    add_bar_chart(doc, ['25-30', '30-35', '35-40', '40-45', 'Above 50'], [15, 12, 65, 5, 3], 'Figure 4.1: Age Distribution of Respondents', 'Age Bracket', 'Percentage (%)', 'steelblue')
    fc = doc.add_paragraph(); fc.add_run('Figure 4.1: Age Distribution of Respondents').font.italic = True; fc.runs[0].font.size = Pt(10); fc.runs[0].font.name = FONT_NAME; fc.alignment = WD_ALIGN_PARAGRAPH.CENTER; _sp(fc, 2, 8); _line1(fc)

    heading3(doc, '4.2.3 Length of Business Operation')
    body(doc, 'The respondents were asked to indicate the number of years their businesses had been in operation. The results are presented in Table 4.3 below.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.3: Length of Business Operation')
    simple_table(doc,
        ['Years of Operation', 'Frequency', 'Percentage'],
        [
            ['0\u20135 years', '60', '60%'],
            ['5\u201310 years', '20', '20%'],
            ['10\u201315 years', '10', '10%'],
            ['20\u201325 years', '7', '7%'],
            ['Above 25 years', '3', '3%'],
            ['Total', '100', '100%'],
        ],
        col_widths=[2.5, 1.2, 1.2])
    source_note(doc)
    body(doc, 'The table above shows that the majority of respondents at 60% have operated their businesses for 0\u20135 years, followed by 20% who have operated for 5\u201310 years, 10% for 10\u201315 years, 7% for 20\u201325 years, and 3% for above 25 years. From the above table, it is revealed that many businesses are relatively new, which may be a result of old businesses collapsing and entrepreneurs starting over again.')
    add_bar_chart(doc, ['0-5', '5-10', '10-15', '20-25', 'Above 25'], [60, 20, 10, 7, 3], 'Figure 4.2: Length of Business Operation', 'Years of Operation', 'Percentage (%)', 'darkorange')
    fc2 = doc.add_paragraph(); fc2.add_run('Figure 4.2: Length of Business Operation').font.italic = True; fc2.runs[0].font.size = Pt(10); fc2.runs[0].font.name = FONT_NAME; fc2.alignment = WD_ALIGN_PARAGRAPH.CENTER; _sp(fc2, 2, 8); _line1(fc2)

    heading3(doc, '4.2.4 Level of Education')
    body(doc, 'The respondents were asked to indicate their highest level of education. The results are presented in Table 4.4 below.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.4: Level of Education')
    simple_table(doc,
        ['Education Level', 'Frequency', 'Percentage'],
        [
            ['Primary', '25', '25%'],
            ['Secondary', '45', '45%'],
            ['Tertiary', '15', '15%'],
            ['University', '15', '15%'],
            ['Total', '100', '100%'],
        ],
        col_widths=[2.5, 1.2, 1.2])
    source_note(doc)
    body(doc, 'The results show that secondary level education had the majority of respondents at 45%, followed by 25% with primary level education, 15% at tertiary level, and 15% at university level. The findings indicate that higher education is not a major factor in the ability to run a small and medium enterprise, as the majority of entrepreneurs had secondary level education.')
    add_bar_chart(doc, ['Primary', 'Secondary', 'Tertiary', 'University'], [25, 45, 15, 15], 'Figure 4.3: Level of Education', 'Education Level', 'Percentage (%)', 'seagreen')
    fc3 = doc.add_paragraph(); fc3.add_run('Figure 4.3: Level of Education').font.italic = True; fc3.runs[0].font.size = Pt(10); fc3.runs[0].font.name = FONT_NAME; fc3.alignment = WD_ALIGN_PARAGRAPH.CENTER; _sp(fc3, 2, 8); _line1(fc3)

    heading3(doc, '4.2.5 Sources of Capital')
    body(doc, 'The respondents were asked to indicate the source of their starting capital. The results are presented in Table 4.5 below.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.5: Sources of Starting Capital')
    simple_table(doc,
        ['Source of Capital', 'Frequency', 'Percentage'],
        [
            ['Loans', '30', '30%'],
            ['Savings', '40', '40%'],
            ['Lease', '5', '5%'],
            ['Friends', '10', '10%'],
            ['Property', '15', '15%'],
            ['Total', '100', '100%'],
        ],
        col_widths=[2.5, 1.2, 1.2])
    source_note(doc)
    body(doc, 'The results show that the majority of respondents at 40% sourced their capital from savings, followed by 30% from loans, 15% from property, 10% from friends, and lastly 5% from leasing. The findings show a higher preference for savings compared to other means of financing business ventures in Eldoret Town.')
    add_bar_chart(doc, ['Loans', 'Savings', 'Lease', 'Friends', 'Property'], [30, 40, 5, 10, 15], 'Figure 4.4: Sources of Starting Capital', 'Source of Capital', 'Percentage (%)', 'mediumpurple')
    fc4 = doc.add_paragraph(); fc4.add_run('Figure 4.4: Sources of Starting Capital').font.italic = True; fc4.runs[0].font.size = Pt(10); fc4.runs[0].font.name = FONT_NAME; fc4.alignment = WD_ALIGN_PARAGRAPH.CENTER; _sp(fc4, 2, 8); _line1(fc4)

    heading2(doc, '4.3 Descriptive Findings')
    body(doc, 'This section shows the variables analyzed including standard deviation, mean, maximum values, and frequency distribution for all the responses recorded by the questionnaire. The analysis is based on a Likert scale of 1\u20135 for each question, where: 5 = Strongly Agree, 4 = Agree, 3 = Neutral, 2 = Disagree, 1 = Strongly Disagree.')

    heading3(doc, '4.3.1 Liquidity Risk Management')
    body(doc, 'Liquidity risk influences the financial performance of SMEs in Eldoret Town. The study findings relate to findings of Rotich and Wahome (2018) which examined the influence of financial practices on the financial performance of youth groups. Respondents were observed to follow the assumption of Prospect Theory and varied in using cash in hand versus borrowing from financial institutions.')
    body(doc, 'The study established the effect of liquidity risk on financial performance of SMEs in Eldoret Town. The results are presented in Table 4.6 below.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.6: Liquidity Risk Management')
    simple_table(doc,
        ['Statement', 'SA', 'A', 'N', 'D', 'SD', 'Mean', 'St. Dev'],
        [
            ['Insufficient cash affects the financial performance of firms', '75(75%)', '15(15%)', '5(5%)', '2(2%)', '3(3%)', '4.57', '0.908'],
            ['We have a challenge in proper usage of resources meant for business', '40(40%)', '20(20%)', '15(15%)', '10(10%)', '15(15%)', '3.60', '1.463'],
            ['Failure to acquire basic resources for our business hampers growth', '90(90%)', '4(4%)', '2(2%)', '2(2%)', '2(2%)', '4.82', '0.433'],
            ['Due to continued cash shortage, we borrow to sustain our businesses', '70(70%)', '25(25%)', '2(2%)', '2(2%)', '1(1%)', '4.70', '0.224'],
        ],
        col_widths=[2.2, 0.55, 0.55, 0.5, 0.5, 0.5, 0.5, 0.55])
    source_note(doc)
    body(doc, 'As the findings show, an average mean of 4.82 indicates that respondents gave strongly positive feedback that liquidity risk influences financial performance of SMEs in Eldoret Town. The statement "Failure to acquire basic resources for our business hampers growth" had the highest mean of 4.82 and standard deviation of 0.433. The statement "We have a challenge in proper usage of resources meant for business" had the lowest mean of 3.60 and standard deviation of 1.463. The statement "Due to continued cash shortages, we borrow to sustain our businesses" had a mean of 4.70 and standard deviation of 0.224.')

    heading3(doc, '4.3.2 Credit Risk Management')
    body(doc, 'Descriptive findings reveal that credit risk affects financial performance of SMEs in Eldoret Town. In seeking credit to fund their businesses, small owners were observed to follow the assumption of Prospect Theory and sought to find out the possible problems for each financial instrument before deciding the best choice.')
    body(doc, 'The study established the influence of credit risk on financial performance of SMEs in Eldoret Town. The results are shown in Table 4.7 below.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.7: Credit Risk Management')
    simple_table(doc,
        ['Statement', 'SA', 'A', 'N', 'D', 'SD', 'Mean', 'St. Dev'],
        [
            ['We over-borrow to finance our business', '40(40%)', '25(25%)', '15(15%)', '5(5%)', '5(5%)', '1.625', '1.274'],
            ['We have difficulties differentiating credit products in the market', '50(50%)', '30(30%)', '4(4%)', '15(15%)', '1(1%)', '1.210', '1.101'],
            ['Our misunderstanding on interest rate computation has led to high indebtedness', '45(45%)', '15(15%)', '23(23%)', '10(10%)', '7(7%)', '1.614', '1.270'],
            ['Loan default rates are very high', '80(80%)', '4(4%)', '4(4%)', '4(4%)', '6(6%)', '1.684', '1.298'],
            ['Borrowed finances are most of the times used for personal use', '70(70%)', '25(25%)', '3(3%)', '1(1%)', '1(1%)', '1.679', '1.296'],
        ],
        col_widths=[2.2, 0.55, 0.55, 0.5, 0.5, 0.5, 0.5, 0.55])
    source_note(doc)
    body(doc, 'As shown in the table above, the respondents agreed that loan default rates are very high, with a mean of 1.684 and standard deviation of 1.298. The statement "Borrowed finances are most of the times used for personal use" had a mean of 1.679 and standard deviation of 1.296. The statement "Our misunderstanding on interest rate computation has led to high indebtedness" had a mean of 1.614 and standard deviation of 1.270. The findings show that credit risk plays an important role in influencing the financial performance of SMEs in Eldoret Town.')

    heading3(doc, '4.3.3 Equity Risk Management')
    body(doc, 'The descriptive findings reveal that equity risk influences financial performance of SMEs in Eldoret Town. The findings are similar to those of Abeyrathna and Kalainathan (2016) which sought to identify the impact of financial risk and quality of financial risk management towards SMEs in the Anuradhapura district. In their findings, they discovered that SMEs favor equity with higher return on equity, approving the assumption of Modern Portfolio Theory.')
    body(doc, 'The study sought to establish the influence of equity risk on performance of SMEs in Eldoret Town. The results are shown in Table 4.8 below.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.8: Equity Risk Management')
    simple_table(doc,
        ['Statement', 'SA', 'A', 'N', 'D', 'SD', 'Mean', 'St. Dev'],
        [
            ['Big shareholders put their self-interest first and always seek to benefit more', '70(70%)', '15(15%)', '5(5%)', '5(5%)', '5(5%)', '4.22', '0.921'],
            ['We do not know the balance between debt and equity', '90(90%)', '4(4%)', '2(2%)', '2(2%)', '2(2%)', '4.82', '0.233'],
            ['We plough back our profits into business', '40(40%)', '20(20%)', '15(15%)', '10(10%)', '15(15%)', '3.60', '1.463'],
        ],
        col_widths=[2.2, 0.55, 0.55, 0.5, 0.5, 0.5, 0.5, 0.55])
    source_note(doc)
    body(doc, 'The results prove the importance of equity risk management. The statement "We plough back profits into the business" had a mean of 3.60 and standard deviation of 1.463, which was the lowest mean and standard deviation. The statement "Big shareholders put their self-interest first and always seek to benefit more" had a mean of 4.22 and standard deviation of 0.921. The statement "We do not know the balance between equity and debt" had the highest mean of 4.82 and standard deviation of 0.233, indicating that the majority of SME owners lack knowledge on the optimal debt-equity balance.')

    heading3(doc, '4.3.4 Financial Performance')
    body(doc, 'The study sought to establish the descriptive measures of financial performance of SMEs in Eldoret Town. The results are presented in Table 4.9 below.')
    doc.add_paragraph()
    table_caption(doc, 'Table 4.9: Financial Performance')
    simple_table(doc,
        ['Statement', 'SA', 'A', 'N', 'D', 'SD', 'Mean', 'St. Dev'],
        [
            ['Effective equity risk management positively affects business leverage', '75(75%)', '15(15%)', '5(5%)', '2(2%)', '3(3%)', '4.23', '1.463'],
            ['Effective risk management positively affects business growth', '60(60%)', '20(20%)', '10(10%)', '5(5%)', '5(5%)', '4.24', '1.335'],
            ['Effective credit risk management leads to better business profits', '50(50%)', '30(30%)', '10(10%)', '5(5%)', '5(5%)', '4.50', '1.163'],
            ['Effective liquidity risk management affects business returns positively', '40(40%)', '20(20%)', '20(20%)', '10(10%)', '10(10%)', '3.75', '1.163'],
        ],
        col_widths=[2.2, 0.55, 0.55, 0.5, 0.5, 0.5, 0.5, 0.55])
    source_note(doc)
    body(doc, 'The table presented above shows that the highest financial risk factor affecting the performance of SMEs in Eldoret Town is effective credit risk management, which leads to better business profits with a mean of 4.50 and standard deviation of 1.163. Effective risk management positively affects business growth with a mean of 4.24 and standard deviation of 1.335. Effective equity risk management positively affects business leverage with a mean of 4.23 and standard deviation of 1.463. Effective liquidity risk management affects business returns positively with the lowest mean of 3.75 and standard deviation of 1.163.')
    body(doc, 'The findings are consistent with those of Dhuci (2011) which established that weak credit risk is a primary cause of small business failure. Overall, the descriptive findings confirm that all three dimensions of financial risk management\u2014liquidity risk, credit risk, and equity risk\u2014have a significant influence on the financial performance of SMEs in Eldoret Town.')

    p = heading_center(doc, 'CHAPTER FIVE', before=0, after=2)
    p.paragraph_format.page_break_before = True
    heading_center(doc, 'SUMMARY, CONCLUSIONS AND RECOMMENDATIONS', before=0, after=14)

    heading2(doc, '5.1 Introduction')
    body(doc, 'This chapter presents the summary of major findings of the study, relevant discussions, conclusions, and the necessary recommendations. The summary is done in line with the objectives of the study based on the output of statistical analyses. The chapter also suggests areas for further research.')

    heading2(doc, '5.2 Summary of Findings')
    body(doc, 'The study response rate was adequate for generalization of study findings. The responses given by the respondents were tested for reliability and it was established that the data used in analysis was reliable. The study targeted 146 participants who were given questionnaires to fill. Out of 146 participants, 100 returned the questionnaires for data analysis, representing a response rate of 68.5%.')
    body(doc, 'General information obtained from participants shows that: the majority of respondents are aged 35\u201340 years at 65%; the majority of study participants are self-employed and they own the businesses; 60% of participants have operated their businesses for 0\u20135 years; 45% of respondents have secondary school level of education; and 40% of respondents sourced their starting capital through savings.')

    heading3(doc, '5.2.1 Liquidity Risk Management')
    body(doc, 'The study sought to determine the influence of liquidity risk management on financial performance of SMEs in Eldoret Town. The findings prove the importance of liquidity risk management which influences the performance of SMEs. The most critical factor as stated by the respondents is continued cash shortage that forces entrepreneurs to borrow in order to sustain their business plans. Failure to acquire basic resources for business was identified as the factor with the strongest influence on financial performance (M=4.82, SD=0.433).')

    heading3(doc, '5.2.2 Equity Risk Management')
    body(doc, 'The study sought to determine the influence of equity risk management on financial performance of SMEs in Eldoret Town. The findings prove the importance of equity risk management which influences the financial performance of SMEs. The most influential factor as indicated by respondents is the lack of knowledge on the right balance between debt and equity (M=4.82, SD=0.233). Additionally, the tendency of big shareholders to put their self-interest first affects expansion plans of SMEs.')

    heading3(doc, '5.2.3 Credit Risk Management')
    body(doc, 'The study sought to determine the influence of credit risk management on financial performance of SMEs in Eldoret Town. The findings prove the importance of credit risk management which influences the financial performance of SMEs. The most influential factor as indicated by respondents is that loan default rates are very high (M=1.684, SD=1.298). Additionally, borrowed finances are most of the times used for personal use rather than business purposes, which negatively impacts SME financial performance.')

    heading2(doc, '5.3 Conclusions')
    heading3(doc, '5.3.1 Liquidity Risk Management')
    body(doc, 'The study sought to establish the influence of liquidity risk management on the financial performance of SMEs in Eldoret Town. Data analysis reveals a positive correlation between liquidity risk management and financial performance of SMEs. As a result, the study concludes that liquidity risk management is positively and significantly related to financial performance of SMEs in Eldoret Town. Effective management of cash flow, working capital, and short-term obligations is essential for SME survival and growth.')

    heading3(doc, '5.3.2 Credit Risk Management')
    body(doc, 'The study also sought to establish the influence of credit risk management on the financial performance of SMEs in Eldoret Town. Data analysis reveals that poor credit risk management has a negative effect on the financial performance of SMEs. As a result, the study concludes that credit risk management is positively and significantly related to financial performance of SMEs in Eldoret Town. Proper credit analysis, loan management, and understanding of interest rate computation are critical for SME financial health.')

    heading3(doc, '5.3.3 Equity Risk Management')
    body(doc, 'The study sought to establish the influence of equity risk management on the financial performance of SMEs in Eldoret Town. Data analysis reveals a positive correlation between equity risk management and financial performance of SMEs. As a result, the study concludes that equity risk management is positively and significantly related to financial performance of SMEs in Eldoret Town. Understanding the balance between debt and equity, attracting investors, and retaining profits are key factors for improving SME performance.')

    heading2(doc, '5.4 Recommendations')
    body(doc, 'Based on the objectives of the study, the following recommendations were made:')
    body(doc, 'First, management of risk in small and medium scale enterprises should enhance their capacity in credit analysis and loan administration. Clear credit policies and lending guidelines should be established. Management is also required to make sure that terms and conditions are adhered to in loan approval. Hence lending guidelines should be approved by senior management and made aware to all staff. This will reduce loss on non-performing loans and improve the financial performance.')
    body(doc, 'Second, study findings show that owners of small businesses in Eldoret Town rely on credit to finance their businesses. To ensure that business owners realize the full potential of credit facilities in the market, the study recommends for business owners to seek more training on credit management from financial and educational institutions.')
    body(doc, 'Third, study findings show that SME owners fail to acquire basic resources for their business which hampers their growth. This problem can be attributed to poor business planning. Hence, it is recommended that SME owners consult experts in the business field in order to gain knowledge on how best to make good use of the resources they have.')
    body(doc, 'Lastly, it was observed that SME owners in Eldoret Town fail to attract big investors in their businesses which affects their expansion plans. It is advisable for the small business owners to scan the environment so as to have a better understanding of their business opportunities and threats. With such an understanding, they can easily identify the future prospects of their fortunes, hence it will be easy for investors to buy shares with assurance of good returns.')

    heading2(doc, '5.5 Suggestions for Further Research')
    body(doc, 'The study findings have revealed a positive and significant relationship between financial risk management factors and financial performance of SMEs in Eldoret Town. However, it is evident from the responses that respondents have little knowledge on some critical aspects of financial risk management which should be studied further to fill the existing information gaps.')
    body(doc, 'Further studies should be done to establish: the effect of lack of collaterals on capital acquisition and how they affect SMEs\u2019 financial performance; the reasons why borrowed finances are most of the times used for personal use; and strategies that small businesses can adapt to maximize on opportunities in foreign markets. Additionally, future research should consider employing inferential statistics such as correlation and regression analysis to establish the strength and nature of the relationships between financial risk management variables and SME financial performance.')

    p = heading_center(doc, 'REFERENCES', before=0, after=14)
    p.paragraph_format.page_break_before = True
    references = [
        'Abeyrathna, G. M., & Kalainathan, K. (2016). Financial risk, financial risk management practices and performance of Sri Lankan SMEs: Special reference to Anuradhapura district. Research Journal of Finance and Accounting, 15(7), 16-22.',
        'Bel\u00e1s, J., & Sopkov\u00e1, G. (2016). Significant determinants of the competitive environment for SMEs in the context of financial and credit risks. Journal of International Studies.',
        'Bel\u00e1s, J., Dvorsk\u00fd, J., Kub\u00e1lek, J., & Smr\u010dka, L. (2018). Important factors of financial risk in the SME segment. Journal of International Studies.',
        'Bel\u00e1s, J., Klju\u010dnikov, A., Vojtovi\u010d, S., & Sobekov\u00e1-M\u00e1jkov\u00e1, M. (2015). Approach of the SME entrepreneurs to financial risk management in relation to gender and level of education. Economics and Sociology.',
        'Buchdadi, A. D., Sholeha, A., & Ahmad, G. N. (2020). The influence of financial literacy on SMEs performance through access to finance and financial risk attitude as mediation variables. Academy of Accounting and Financial Studies Journal, 24(5), 1-15.',
        'Busato, F., & Coletta, C. (2017). Moral hazard in corporate governance. International Journal of Business and Management, 12(4), 1-10.',
        'Gra\u00f1a-Alvarez, R., Lopez-Valeiras, E., Gonzalez-Loureiro, M., & Coronado, F. (2022). Financial literacy in SMEs: A systematic literature review and a framework for further inquiry. Journal of Small Business Management, 1-50.',
        'Hanggraeni, D., \u015alusarczyk, B., Sulung, L. A. K., & Subroto, A. (2019). The impact of internal, external and enterprise risk management on the performance of micro, small and medium enterprises. Sustainability, 11(7), 2172.',
        'Jagongo, A., & Rop, E. (2021). Liquidity risk management and financial performance of state-owned enterprises in Kenya. International Journal of Finance and Accounting, 6(2), 12-28.',
        'Kahneman, D., & Tversky, A. (2013). Prospect theory: An analysis of decision under risk. In Handbook of the fundamentals of financial decision making: Part I (pp. 99-127).',
        'Kotaskova, A., Lazanyi, K., Amoah, J., & Belas, J. (2020). Financial risk management in the V4 Countries\u2019 SMEs segment. Investment Management and Financial Innovations.',
        'Koyuncugil, A. S. (2009). Early warning system for SMEs as a financial risk detector. In Data mining applications for empowering knowledge societies (pp. 220-238). IGI Global.',
        'Krugman, P. (2009). The Return of Depression Economics and the Crisis of 2008. W.W. Norton & Company.',
        'Kurban, G., & Nasir, H. (2022). Liquidity risk and organizational performance in Malaysia small and medium enterprises. International Journal of Finance and Accounting, 7(2), 65-73.',
        'Mangram, M. E. (2013). A simplified perspective of the Markowitz portfolio theory. Global Journal of Business Research, 7(1), 59-70.',
        'Masheta, N. F. (2019). Assessment on the impact and efficiency of credit risk management on profitability of five microfinance institutions in Zambia (Doctoral dissertation, Cavendish University).',
        'Mugenda, O. M., & Mugenda, A. G. (2003). Research Methods: Quantitative and Qualitative Approaches. Nairobi: Acts Press.',
        'Mutamimah, M., Tholib, M., & Robiyanto, R. (2021). Corporate governance, credit risk, and financial literacy for small medium enterprise in Indonesia. Business: Theory and Practice, 22(2), 406-413.',
        'Ol\u00e1h, J., Kov\u00e1cs, S., Virglerova, Z., Lakner, Z., Kovacova, M., & Popp, J. (2019). Analysis and comparison of economic and financial risk sources in SMEs of the Visegrad group and Serbia. Sustainability, 11(7), 1853.',
        'van den Boom, R. P. (2020). Financial Risk Management in SMEs: A New Conceptual Framework. International Business Research, 13(10), 1-85.',
        'Virglerova, Z., Kozub\u00edkov\u00e1, L., & Vojtovi\u010d, S. (2016). Influence of selected factors on financial risk management in SMEs in the Czech Republic. Montenegrin Journal of Economics.',
        'Wanjohi, A. M., Wanjohi, J. G., & Ndambiri, J. M. (2017). The role of moral hazard in financial risk management. Journal of Finance and Risk Management, 4(2), 45-58.',
        'Widyastuti, M., Ferdinand, D. Y. Y., & Hermanto, Y. B. (2023). Strengthening Formal Credit Access and Performance through Financial Literacy and Credit Terms in Micro, Small and Medium Businesses. Journal of Risk and Financial Management, 16(1), 52.',
        'Yin, C., Jiang, C., Jain, H. K., & Wang, Z. (2020). Evaluating the credit risk of SMEs using legal judgments. Decision Support Systems, 136, 113364.',
    ]
    for ref in references:
        para = doc.add_paragraph()
        run = para.add_run(ref)
        run.font.size = Pt(11)
        run.font.name = FONT_NAME
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)
        _sp(para, 2, 4); _line15(para)

    p = heading_center(doc, 'APPENDICES', before=0, after=14)
    p.paragraph_format.page_break_before = True

    heading2(doc, 'Appendix I: Introduction Letter')
    body(doc, 'OPISA AGNETTA')
    body(doc, 'MOI UNIVERSITY')
    body(doc, 'P.O BOX 3900-30100,')
    body(doc, 'ELDORET.')
    body(doc, '')
    body(doc, 'DEAR RESPONDENT,')
    body(doc, '')
    body(doc, 'RE: QUESTIONNAIRE')
    body(doc, '')
    body(doc, 'I am a student at Moi University pursuing a degree in Bachelor of Business Management (Banking and Finance). I am carrying out a study on the effects of financial risk management on the financial performance of small and medium scale enterprises in Eldoret Town, Kenya. You are kindly requested to complete the attached questionnaire so as to enable me accomplish my study. Please note that all the information given shall be treated purely and used for academic purposes and shall be treated as confidential. Thank you for taking your time to complete the questionnaire.')
    body(doc, '')
    body(doc, 'Yours faithfully,')
    body(doc, 'Opisa Agnetta.')

    heading2(doc, 'Appendix II: Questionnaire')
    body(doc, 'This questionnaire seeks to collect data about the effect of financial risk management on financial performance among SMEs in Eldoret Town, Kenya. The data collected will be utilized for academic purposes. Tick the form provided appropriately.')
    body(doc, '')
    body(doc, 'SECTION A: DEMOGRAPHIC DATA', before=6)
    body(doc, '')
    body(doc, '1. Age:')
    body(doc, '   25\u201330 years ( )   30\u201335 years ( )   35\u201340 years ( )   40\u201345 years ( )   Above 50 years ( )')
    body(doc, '')
    body(doc, '2. Please state whether you own the business or you are employed:')
    body(doc, '   Self-employed ( )   Employed ( )')
    body(doc, '')
    body(doc, '3. How long has the business been in operation?')
    body(doc, '   0\u20135 years ( )   5\u201310 years ( )   10\u201315 years ( )   15\u201320 years ( )   20\u201325 years ( )   Over 25 years ( )')
    body(doc, '')
    body(doc, '4. What is your highest level of education?')
    body(doc, '   Primary ( )   Secondary ( )   Tertiary ( )   University ( )')
    body(doc, '')
    body(doc, '5. What was the source of your starting capital?')
    body(doc, '   Savings ( )   Loan ( )   Friends and family contributions ( )   Lease of property ( )   Other ( )')

    body(doc, '', before=8)
    body(doc, 'SECTION B: QUESTIONS FOR RESEARCH OBJECTIVES', before=6)

    body(doc, '', before=6)
    body(doc, 'I. Liquidity Risk Management', before=4)
    body(doc, 'To what extent do the following factors affect the financial performance of SMEs in Eldoret Town? Pick the best choice in the table provided: Strongly Agree = 5; Agree = 4; Neutral = 3; Disagree = 2; Strongly Disagree = 1.')
    body(doc, '')
    body(doc, '1. Insufficient cash affects the financial performance of our business  [5] [4] [3] [2] [1]')
    body(doc, '2. We have a challenge in proper usage of resources meant for business  [5] [4] [3] [2] [1]')
    body(doc, '3. Failure to acquire basic resources for our business hampers our growth  [5] [4] [3] [2] [1]')
    body(doc, '4. Due to continued cash shortages, we borrow to sustain our businesses  [5] [4] [3] [2] [1]')
    body(doc, '5. Lack of collaterals have limited our capital acquisition efforts  [5] [4] [3] [2] [1]')

    body(doc, '', before=6)
    body(doc, 'II. Credit Risk Management', before=4)
    body(doc, 'To what extent do the following factors affect the financial performance of SMEs in Eldoret Town? Pick the best choice: Strongly Agree = 5; Agree = 4; Neutral = 3; Disagree = 2; Strongly Disagree = 1.')
    body(doc, '')
    body(doc, '1. We over-borrow to finance our business  [5] [4] [3] [2] [1]')
    body(doc, '2. We have difficulties in differentiating various credit products offered in the financial market  [5] [4] [3] [2] [1]')
    body(doc, '3. Our misunderstanding on interest rate computation has led to high indebtedness  [5] [4] [3] [2] [1]')
    body(doc, '4. Loan default rates are very high  [5] [4] [3] [2] [1]')
    body(doc, '5. Borrowed finances are most of the times used for personal use  [5] [4] [3] [2] [1]')

    body(doc, '', before=6)
    body(doc, 'III. Equity Risk Management', before=4)
    body(doc, 'To what extent do the following factors affect the financial performance of SMEs in Eldoret Town? Pick the best choice: Strongly Agree = 5; Agree = 4; Neutral = 3; Disagree = 2; Strongly Disagree = 1.')
    body(doc, '')
    body(doc, '1. In most cases, partners fail to agree on the shareholding ratios  [5] [4] [3] [2] [1]')
    body(doc, '2. Big shareholders put their self-interest first and always seek to benefit more  [5] [4] [3] [2] [1]')
    body(doc, '3. We do not know the right balance between equity and debt in business  [5] [4] [3] [2] [1]')
    body(doc, '4. Failure to attract big investors has affected our expansion plans  [5] [4] [3] [2] [1]')
    body(doc, '5. We plough back our profits into the business  [5] [4] [3] [2] [1]')

    body(doc, '', before=6)
    body(doc, 'IV. Financial Performance', before=4)
    body(doc, 'To what extent do the following factors affect the financial performance of SMEs in Eldoret Town? Pick the best choice: Strongly Agree = 5; Agree = 4; Neutral = 3; Disagree = 2; Strongly Disagree = 1.')
    body(doc, '')
    body(doc, '1. Effective equity risk management positively affects business leverage  [5] [4] [3] [2] [1]')
    body(doc, '2. Effective risk management positively affects business growth  [5] [4] [3] [2] [1]')
    body(doc, '3. Effective credit risk management leads to better business profits  [5] [4] [3] [2] [1]')
    body(doc, '4. Effective liquidity risk management affects business returns positively  [5] [4] [3] [2] [1]')

    add_page_numbers(doc)
    fn = 'files/Agnetta_Opisa_Research_Project.docx'
    doc.save(fn)
    print(f'DOCX saved: {fn}')
    return fn

def convert_to_pdf(docx_path):
    import subprocess
    import os
    env = os.environ.copy()
    env['HOME'] = '/tmp'
    result = subprocess.run(
        [_SOFFICE, '--headless', '--convert-to', 'pdf', '--outdir', 'files', docx_path],
        capture_output=True, text=True, env=env, timeout=120
    )
    if result.returncode != 0:
        raise RuntimeError(f'PDF conversion failed: {result.stderr}')
    pdf_path = docx_path.replace('.docx', '.pdf')
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f'PDF not created: {pdf_path}')
    print(f'PDF saved: {pdf_path}')
    return pdf_path


def verify_pdf(pdf_path):
    import fitz
    doc = fitz.open(pdf_path)
    total = len(doc)
    blank = [i + 1 for i in range(total) if len(doc[i].get_text().strip()) < 5]
    print(f'Total PDF pages: {total}')
    if blank:
        print(f'WARNING: Blank pages found: {blank}')
    else:
        print('Blank pages: NONE (0 blank pages)')

    ch1_idx = None
    for i in range(len(doc)):
        text = doc[i].get_text().strip()
        if text.startswith('CHAPTER ONE') and '1.1 Background' in text:
            ch1_idx = i
            break
    if ch1_idx is None:
        print('ERROR: Could not find Chapter 1 start page')
        return False
    offset = ch1_idx

    toc_checks = [
        ('1.1 Background', '1'), ('1.1.1 Financial Risk', '2'),
        ('1.1.2 Methods', '3'), ('1.1.3 Performance', '4'),
        ('1.2 Statement', '4'), ('1.3 Objectives', '6'),
        ('1.4 Research Questions', '6'), ('1.5 Significance', '6'),
        ('1.6 Scope', '7'), ('1.7 Limitation', '7'),
        ('CHAPTER TWO', '8'), ('2.1 Introduction', '8'),
        ('2.2 Theoretical', '8'), ('2.2.1 Prospect', '8'),
        ('2.2.2 Modern Portfolio', '9'), ('2.2.3 Moral Hazard', '9'),
        ('2.3 Conceptual Framework', '10'), ('2.4 Empirical', '11'),
        ('2.4.1 Credit Risk', '12'), ('2.4.2 Equity Risk', '12'),
        ('2.4.3 Liquidity Risk', '13'), ('2.5 Critique', '13'),
        ('2.6 Research Gaps', '14'), ('2.7 Summary', '14'),
        ('CHAPTER THREE', '16'), ('3.1 Introduction', '16'),
        ('3.2 Research Design', '16'), ('3.3 Target Population', '16'),
        ('3.4 Sample Design', '17'), ('3.5 Research Instruments', '17'),
        ('3.6 Data Collection', '18'), ('3.7 Data Analysis', '18'),
        ('3.8 Ethical', '18'),
        ('CHAPTER FOUR', '20'), ('4.1 Introduction', '20'),
        ('4.2 Demographic', '20'), ('4.2.1 Response Rate', '20'),
        ('4.2.2 Age of', '20'), ('4.2.3 Length', '21'),
        ('4.2.4 Level of Education', '22'), ('4.2.5 Sources of Capital', '23'),
        ('4.3 Descriptive', '24'), ('4.3.1 Liquidity Risk', '24'),
        ('4.3.2 Credit Risk', '25'), ('4.3.3 Equity Risk', '27'),
        ('4.3.4 Financial Performance', '28'),
        ('CHAPTER FIVE', '30'), ('5.1 Introduction', '30'),
        ('5.2 Summary', '30'), ('5.2.1 Liquidity', '30'),
        ('5.2.2 Equity', '30'), ('5.2.3 Credit', '31'),
        ('5.3 Conclusions', '31'), ('5.4 Recommendations', '32'),
        ('5.5 Suggestions', '32'),
        ('REFERENCES', '34'), ('APPENDICES', '36'),
        ('Table 3.1', '16'), ('Table 3.2', '17'),
        ('Table 4.1', '20'), ('Table 4.2', '21'), ('Table 4.3', '21'),
        ('Table 4.4', '22'), ('Table 4.5', '23'), ('Table 4.6', '24'),
        ('Table 4.7', '26'), ('Table 4.8', '27'), ('Table 4.9', '28'),
        ('Figure 2.1', '11'), ('Figure 4.1', '21'),
        ('Figure 4.2', '22'), ('Figure 4.3', '23'), ('Figure 4.4', '24'),
    ]
    mismatches = 0
    def _find_heading(heading, expected):
        nonlocal mismatches
        for i in range(offset, len(doc)):
            page_text = doc[i].get_text()
            matched = False
            if heading in ('REFERENCES', 'APPENDICES'):
                lines = page_text.strip().split('\n')
                matched = any(heading in line and len(line.strip()) < len(heading) + 5 for line in lines[:3])
            elif heading.startswith('CHAPTER'):
                lines = page_text.strip().split('\n')
                matched = any(line.strip() == heading for line in lines[:3])
            else:
                matched = heading.lower() in page_text.lower()
            if matched:
                actual = i + 1 - offset
                if str(actual) != expected:
                    print(f'  TOC MISMATCH: "{heading}" expected={expected} actual={actual}')
                    mismatches += 1
                return
        print(f'  NOT FOUND: "{heading}"')
        mismatches += 1

    for heading, expected in toc_checks:
        _find_heading(heading, expected)

    print(f'TOC spot-checks: {len(toc_checks)} entries, {mismatches} mismatches')
    if mismatches == 0:
        print('TOC VERIFICATION PASSED')
    doc.close()
    return len(blank) == 0 and mismatches == 0


if __name__ == '__main__':
    docx_path = create_docx()
    pdf_path = convert_to_pdf(docx_path)
    ok = verify_pdf(pdf_path)
    if not ok:
        raise SystemExit('Verification failed')
    print('All checks passed.')
