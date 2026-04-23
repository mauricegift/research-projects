#!/usr/bin/env python3
"""
Generate Mourice's elite, exam-style answers to BBM 433 past papers as a
single PDF: April 2024, April 2023 and July 2025. Each question is rewritten
verbatim, followed immediately by a detailed, well-structured answer.
"""

import os as _os
import sys as _sys

_sys.path.insert(0, _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))
_os.chdir(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))

from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

from generators.generate_mourice_bbm433_cats import (
    new_doc, add_cover, body, heading, subheading, page_break, cover_line,
    convert_to_pdf, set_run, set_spacing,
)


# ────────────────────────────────────────────────────────────────────────────
# Styled helpers for the Q&A document
# ────────────────────────────────────────────────────────────────────────────
def paper_title(doc, text):
    page_break(doc)
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)


def question(doc, qnum, qtext, marks=None):
    p = doc.add_paragraph()
    head = f'QUESTION {qnum}'
    if marks:
        head += f'   [{marks}]'
    r = p.add_run(head)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    if qtext:
        p2 = doc.add_paragraph()
        r2 = p2.add_run(qtext)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
        r2.font.italic = True
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_spacing(p2, before=0, after=6, spacing=1.5)


def sub_q(doc, label, qtext, marks):
    """A sub-question heading (e.g. '(a) ...  [5 marks]')."""
    p = doc.add_paragraph()
    r = p.add_run(f'{label} {qtext}')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.bold = True
    r2 = p.add_run(f'   ({marks})')
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    r2.font.bold = True
    r2.font.italic = True
    r2.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, before=8, after=4, spacing=1.5)


def ans_label(doc):
    p = doc.add_paragraph()
    r = p.add_run('Answer:')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.italic = True
    set_spacing(p, before=2, after=2, spacing=1.5)


def ans(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, before=0, after=6, spacing=1.5)


def ans_bullet(doc, label, text, indent=0.4):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r0 = p.add_run('• ')
    set_run(r0, size=12)
    if label:
        r1 = p.add_run(label)
        set_run(r1, bold=True, size=12)
    r2 = p.add_run(text)
    set_run(r2, size=12)


# ────────────────────────────────────────────────────────────────────────────
# Document
# ────────────────────────────────────────────────────────────────────────────
def build(output='files/Mourice_BBM_433_Past_Papers_Answers.docx'):
    doc = new_doc()
    add_cover(
        doc,
        course_code='BBM 433',
        course_title='RETAIL AND MERCHANDISE MANAGEMENT',
        assignment='PAST PAPERS — DETAILED ANSWERS',
        sub_date='APRIL 2026',
        marks='REVISION DOCUMENT',
    )

    # Intro
    cover_line(doc, 'INTRODUCTION', size=14, bold=True, after=8)
    body(doc,
         'This document presents detailed, exam-style model answers to '
         'three past BBM 433 (Retail and Merchandise Management) '
         'papers: April 2024, April 2023 and July 2025. Each question '
         'is reproduced verbatim and followed immediately by a '
         'structured, examiner-friendly answer that covers the key '
         'concepts, frameworks and Kenyan examples expected at degree '
         'level. Use this document for active revision: read the '
         'question, attempt your own answer, then compare with the '
         'model answer and refine your structure.', after=10)

    # ════════════════════════════════════════════════════════════════════
    # PAPER 1 — APRIL 2024
    # ════════════════════════════════════════════════════════════════════
    paper_title(doc, 'PAPER ONE — APRIL 2024')
    body(doc,
         'Instructions: Question One is compulsory. Answer any other '
         'three questions.', after=8)

    # Q1
    question(doc, 'ONE', None, marks='Compulsory')
    sub_q(doc, '(a)', 'Define the following terms: Retailing, '
          'Merchandise, Merchandise Management, Merchandise '
          'Assortment.', '2 marks each = 8 marks')
    ans_label(doc)
    ans_bullet(doc, 'Retailing: ',
               'The set of business activities involved in selling '
               'goods and services directly to the final consumer for '
               'personal, family or household use. It is the last '
               'link in the distribution chain (e.g. Naivas, Quickmart, '
               'Carrefour, Jumia, the local duka).')
    ans_bullet(doc, 'Merchandise: ',
               'The actual physical goods (or, in modern usage, the '
               'goods and services) that a retailer offers for sale '
               '— the inventory carried in stores or on a website '
               '(e.g. groceries on Naivas\' shelves, fashion on '
               'Vivo Activewear\'s racks).')
    ans_bullet(doc, 'Merchandise Management: ',
               'The strategic process of planning, buying, pricing, '
               'displaying and managing the goods a retailer sells '
               'so as to maximise sales and profit while satisfying '
               'customer needs. It covers the "5 Rights" — right '
               'product, right place, right time, right quantity, '
               'right price.')
    ans_bullet(doc, 'Merchandise Assortment: ',
               'The complete range of products offered by a retailer, '
               'described by two dimensions — breadth (number of '
               'different categories carried) and depth (number of '
               'SKUs/variants within each category). E.g. Carrefour '
               'has a wide-and-deep assortment; a kiosk has a '
               'narrow-and-shallow one.')

    sub_q(doc, '(b)', 'Discuss the wheel of retailing using examples '
          'here in Kenya.', '10 marks')
    ans_label(doc)
    ans(doc,
        'The Wheel of Retailing is a theory developed by Malcolm '
        'McNair in 1958 that explains how new retail formats enter '
        'the market and evolve over time. The theory likens the '
        'evolution of retail formats to a turning wheel that goes '
        'through three phases: entry, trading-up, and vulnerability.')
    ans_bullet(doc, '1. Entry Phase (Low-Status, Low-Price, Low-Margin): ',
               'New retailers enter the market as no-frills, low-cost '
               'operators with simple stores, narrow assortments and '
               'aggressive pricing. They target price-sensitive '
               'shoppers. In Kenya, Tuskys originally entered the '
               'market in this way, as did Naivas — basic stores, '
               'every-day-low-price proposition, focus on volume.')
    ans_bullet(doc, '2. Trading-Up Phase: ',
               'As they grow, these retailers add facilities, services, '
               'wider assortments, better store ambience, fresh-food '
               'sections, in-house bakeries, loyalty programs and '
               'home delivery. Margins and prices rise. Today\'s '
               'Naivas Hypermart in Westlands or Two Rivers is a '
               'completely different proposition from the original '
               'Naivas of the 1990s — smart layouts, butchery, '
               'sushi bar, Naivas Now app, BOPIS.')
    ans_bullet(doc, '3. Vulnerability Phase: ',
               'Having traded up, the retailer becomes vulnerable to a '
               'new wave of low-cost entrants who repeat the cycle. '
               'Quickmart, Chandarana FoodPlus discount tiers, and '
               'now Naivas Easy/Naivas Mtaa Wetu, Carrefour Express '
               'and the rapid expansion of mama-mboga delivery apps '
               '(Twiga, Greenspoon) are pulling price-sensitive '
               'shoppers back to no-frills formats — leaving the '
               'incumbents to defend with brand, experience and '
               'omnichannel.')
    ans(doc,
        'Limitations: the theory does not explain the entry of '
        'high-end formats (Two Rivers Mall luxury boutiques, the '
        'Sarit Centre re-launch) which entered at the top, not the '
        'bottom; nor does it fully account for digital-first '
        'entrants like Jumia or Copia. Nevertheless, it remains a '
        'powerful lens for understanding the long-run dynamics of '
        'Kenyan modern trade.')

    sub_q(doc, '(c)', 'Discuss the various ways a consumer can '
          'purchase goods from the retailer.', '7 marks')
    ans_label(doc)
    ans(doc,
        'In the modern omnichannel environment a consumer has many '
        'paths to purchase. The main ones are:')
    ans_bullet(doc, 'In-store (brick-and-mortar) purchase: ',
               'The traditional model — customer walks in, selects '
               'goods, pays at the till. Still ~80% of Kenyan retail.')
    ans_bullet(doc, 'Online purchase with home delivery: ',
               'Customer orders via website or app; goods are '
               'delivered (Jumia, Kilimall, Naivas Now, Carrefour '
               'app, Glovo Market).')
    ans_bullet(doc, 'BOPIS (Buy Online, Pick Up In Store): ',
               'Order online, collect from the nearest branch — saves '
               'delivery cost. Increasingly offered by Carrefour and '
               'Naivas.')
    ans_bullet(doc, 'Click-and-Collect / Curbside pickup: ',
               'Variation of BOPIS where staff bring the order to the '
               'customer\'s car.')
    ans_bullet(doc, 'Telephone / WhatsApp order: ',
               'Common in Kenyan SMEs — customer messages on '
               'WhatsApp Business, pays via M-PESA, receives delivery '
               'by boda-boda.')
    ans_bullet(doc, 'Social commerce: ',
               'Buying directly from Instagram, TikTok Shop or '
               'Facebook Marketplace.')
    ans_bullet(doc, 'Vending machines & kiosks: ',
               'Self-service points (e.g. PesaPoint, Naivas vending '
               'in CBD).')
    ans_bullet(doc, 'Subscription / auto-replenishment: ',
               'Recurring deliveries (e.g. milk subscriptions, '
               'Greenspoon weekly veg boxes).')
    ans_bullet(doc, 'Direct selling / door-to-door: ',
               'Avon, Tupperware, Tianshi — sales agents bring '
               'product to the customer.')
    ans_bullet(doc, 'Mobile commerce via M-PESA: ',
               'Cuts across most channels — STK push, Pay Bill, '
               'Buy Goods — making digital purchase frictionless '
               'for the unbanked.')

    # Q2
    question(doc, 'TWO', None, marks='15 marks')
    sub_q(doc, '(a)', 'Nowadays customers have the option of shopping '
          'while sitting at home. They can place their order through '
          'the internet, pay with debit/credit cards and have '
          'products delivered to their home. Discuss the advantages '
          'and disadvantages associated with this type of retailing '
          '(e-tailing).', '8 marks')
    ans_label(doc)
    subheading(doc, 'Advantages of E-tailing')
    ans_bullet(doc, 'Convenience: ',
               '24/7 shopping from anywhere; no traffic, queues or '
               'parking issues.')
    ans_bullet(doc, 'Wider assortment: ',
               'Endless aisle — far more SKUs than any physical store '
               'could carry.')
    ans_bullet(doc, 'Price comparison & transparency: ',
               'Shoppers can compare across Jumia, Kilimall and '
               'Carrefour in seconds.')
    ans_bullet(doc, 'Personalisation: ',
               'AI-driven recommendations based on browsing and '
               'purchase history.')
    ans_bullet(doc, 'Lower overheads for retailers: ',
               'No need for prime-location physical space, smaller '
               'staff footprint.')
    ans_bullet(doc, 'Geographic reach: ',
               'A single Nairobi-based store can serve the whole '
               'country (and beyond) via courier.')
    ans_bullet(doc, 'Rich data: ',
               'Every click, search and abandoned cart is captured '
               'for future targeting.')
    subheading(doc, 'Disadvantages of E-tailing')
    ans_bullet(doc, 'No tactile experience: ',
               'Cannot feel fabric, smell perfume or test-fit shoes '
               '— hence high return rates in fashion.')
    ans_bullet(doc, 'Delivery delays & last-mile costs: ',
               'Especially outside Nairobi/Mombasa where logistics '
               'remain expensive.')
    ans_bullet(doc, 'Cyber-security & fraud risk: ',
               'Card fraud, phishing, fake websites; shaken consumer '
               'trust in some markets.')
    ans_bullet(doc, 'Returns & reverse-logistics complexity: ',
               'Costly and operationally hard to handle.')
    ans_bullet(doc, 'Digital divide: ',
               'Excludes consumers without smartphones, data or '
               'banking access.')
    ans_bullet(doc, 'Lack of immediate gratification: ',
               'Customers must wait for delivery instead of walking '
               'out with the product.')
    ans_bullet(doc, 'Counterfeit risk: ',
               'Particularly on open marketplaces — fake electronics, '
               'cosmetics, supplements.')

    sub_q(doc, '(b)', 'Elaborate on the "Darwin\'s Theory of '
          'Retailing".', '7 marks')
    ans_label(doc)
    ans(doc,
        'Darwin\'s Theory of Retailing — also called the theory of '
        '"Natural Selection in Retailing" — applies Charles Darwin\'s '
        'biological concept of "survival of the fittest" to the '
        'retail industry. The theory was popularised by Davidson, '
        'Sweeney and Stampfl in the 1970s and argues that retail '
        'formats, like organisms, must continuously adapt to changes '
        'in their environment or face extinction.')
    subheading(doc, 'Key Tenets of the Theory')
    ans_bullet(doc, 'Environmental change is constant: ',
               'Consumer tastes, technology, regulation, competition '
               'and economic conditions are always shifting.')
    ans_bullet(doc, 'Adaptation is the key to survival: ',
               'Retailers that adjust their formats, assortments, '
               'prices, channels and service propositions in line '
               'with environmental change will thrive.')
    ans_bullet(doc, 'Failure to adapt = extinction: ',
               'Retailers who cling to outdated formats lose '
               'relevance and die. Globally, Sears, Toys "R" Us, '
               'Blockbuster and Mothercare are textbook examples; '
               'in Kenya, Nakumatt\'s collapse in 2018 — once the '
               'largest retailer — is the classic case of failing '
               'to adapt to changing competitive and financial '
               'conditions.')
    ans_bullet(doc, 'New formats emerge to fill gaps: ',
               'Just as new species evolve, new retail formats '
               '(supermarkets in the 1950s, hypermarkets in the '
               '1970s, e-commerce in the 1990s, social commerce and '
               'super-apps today) emerge to satisfy needs that '
               'incumbents cannot.')
    ans(doc,
        'Implications for Kenyan retailers: continuous environmental '
        'scanning (PESTEL), willingness to experiment with new '
        'formats (Naivas\' rapid roll-out of small-format Naivas '
        'Mtaa Wetu, Carrefour Express convenience stores), and '
        'investment in digital channels and customer data are no '
        'longer optional — they are the conditions for survival.')

    # Q3
    question(doc, 'THREE', None, marks='15 marks')
    sub_q(doc, '(a)', 'Explain the systematic process of retailing.',
          '8 marks')
    ans_label(doc)
    ans(doc,
        'The systematic process of retailing is the structured set '
        'of decisions and activities a retailer follows from '
        'identifying market opportunity through to delivering goods '
        'to the final customer. The eight key stages are:')
    ans_bullet(doc, '1. Environmental scanning: ',
               'PESTEL + Porter\'s 5-Forces analysis of the macro '
               'and micro environment to spot opportunities and '
               'threats.')
    ans_bullet(doc, '2. Market research & target-customer definition: ',
               'Segmentation by demographic, geographic, psychographic '
               'and behavioural variables; choice of target segment.')
    ans_bullet(doc, '3. Strategic planning & format choice: ',
               'Decide retail mix — store format, location, '
               'positioning, value proposition.')
    ans_bullet(doc, '4. Merchandise planning & sourcing: ',
               'Forecast demand, build the assortment, negotiate '
               'with suppliers, place purchase orders, manage Open-'
               'to-Buy.')
    ans_bullet(doc, '5. Logistics & inventory management: ',
               'Receive goods at the DC, distribute to stores, '
               'manage stock levels, replenish.')
    ans_bullet(doc, '6. Store operations & visual merchandising: ',
               'Layout, planogram, pricing, signage, staffing, '
               'service standards.')
    ans_bullet(doc, '7. Marketing, promotion & customer engagement: ',
               'Advertising, loyalty programs, digital marketing, '
               'CRM.')
    ans_bullet(doc, '8. Performance measurement & continuous '
                    'improvement: ',
               'Track KPIs (SSS, GMROI, AOV, NPS), gather feedback, '
               'refine the offer. Loop back to step 1.')

    sub_q(doc, '(b)', 'Expound on the duties of a retailer and that '
          'of a consumer.', '7 marks')
    ans_label(doc)
    subheading(doc, 'Duties of the Retailer')
    ans_bullet(doc, 'Stock the right merchandise: ',
               'Carry assortments that match the target market\'s '
               'needs.')
    ans_bullet(doc, 'Provide accurate product information: ',
               'Honest labelling, pricing, ingredients, expiry dates, '
               'usage instructions.')
    ans_bullet(doc, 'Fair pricing: ',
               'No price gouging, deceptive promotions or hidden '
               'charges; comply with the Consumer Protection Act '
               'and Kenya Bureau of Standards (KEBS) requirements.')
    ans_bullet(doc, 'Quality assurance & safety: ',
               'Ensure goods are genuine, safe and within shelf-life.')
    ans_bullet(doc, 'After-sales service: ',
               'Honour warranties, accept legitimate returns, handle '
               'complaints fairly.')
    ans_bullet(doc, 'Maintain a safe shopping environment: ',
               'Clean, well-lit stores; trained staff; fire-safety '
               'compliance.')
    ans_bullet(doc, 'Respect data privacy: ',
               'Protect customer information collected through '
               'loyalty cards, M-PESA receipts and apps.')
    subheading(doc, 'Duties of the Consumer')
    ans_bullet(doc, 'Pay the agreed price promptly: ',
               'Settle in cash, card or M-PESA without dispute.')
    ans_bullet(doc, 'Inspect goods before/at purchase: ',
               'Check expiry, packaging integrity, accuracy of '
               'change.')
    ans_bullet(doc, 'Handle merchandise with care: ',
               'Avoid damaging goods on display.')
    ans_bullet(doc, 'Honest behaviour: ',
               'No shoplifting, fraudulent returns or false claims.')
    ans_bullet(doc, 'Provide constructive feedback: ',
               'Reasonable complaints help retailers improve.')
    ans_bullet(doc, 'Respect store policies and staff: ',
               'Follow queueing, mask, and conduct rules.')
    ans_bullet(doc, 'Know and assert their rights: ',
               'Under the Consumer Protection Act 2012, consumers '
               'have rights to safety, information, choice, redress '
               'and a healthy environment.')

    # Q4
    question(doc, 'FOUR', None, marks='15 marks')
    sub_q(doc, '', 'Discuss eight (8) types of retail outlets and '
          'also the type of merchandise they stock for their '
          'customers.', '15 marks')
    ans_label(doc)
    ans(doc,
        'Retail outlets can be classified by the breadth and depth '
        'of their assortments, their service level and their '
        'pricing strategy. Eight major types are:')
    ans_bullet(doc, '1. Department Stores: ',
               'Large-format stores organised into departments; '
               'carry a wide and deep assortment of clothing, '
               'cosmetics, household goods, electronics and '
               'furniture. Examples: Deacons (now closed) and '
               'historic Mothercare and Woolworths in Kenya. '
               'Globally — Macy\'s, Selfridges.')
    ans_bullet(doc, '2. Supermarkets: ',
               'Self-service stores carrying a wide range of food, '
               'beverages, household and personal-care items at '
               'competitive prices. Examples: Naivas, Quickmart, '
               'Chandarana FoodPlus.')
    ans_bullet(doc, '3. Hypermarkets: ',
               'Very large stores combining supermarket + general '
               'merchandise (electronics, apparel, furniture) under '
               'one roof. Examples: Carrefour Two Rivers, Naivas '
               'Hypermart Westlands.')
    ans_bullet(doc, '4. Convenience Stores: ',
               'Small-format, long-hours stores stocking high-'
               'turnover daily essentials — bread, milk, snacks, '
               'soft drinks, top-up airtime. Examples: Naivas Mtaa '
               'Wetu, Quickmart Express, Shell Select, Total '
               'Bonjour.')
    ans_bullet(doc, '5. Specialty Stores: ',
               'Narrow but very deep assortment in a single product '
               'category. Examples: Bata (footwear), Vivo Activewear '
               '(women\'s fashion), Hotpoint (electronics), Toi '
               'Market traders specialising in one item.')
    ans_bullet(doc, '6. Discount Stores: ',
               'Carry a broad assortment at low prices and low '
               'service. Examples: Mwananchi Wholesale, Eastmatt, '
               'Naivas Easy.')
    ans_bullet(doc, '7. Warehouse / Cash-and-Carry Clubs: ',
               'Bulk-buying outlets serving small businesses and '
               'large households. Examples: Mahitaji Cash & Carry, '
               'Eastmatt Wholesale.')
    ans_bullet(doc, '8. Online / E-tailers: ',
               'Pure-play digital retailers carrying broad, deep '
               'assortments delivered to customers\' doorsteps. '
               'Examples: Jumia, Kilimall, Copia, Sky.Garden.')
    ans(doc,
        'Other notable formats include category killers (large '
        'specialty stores like Game stores in their hey-day), '
        'factory outlets (Bata Industrials), pop-up stores '
        '(Vivo seasonal pop-ups in Sarit Centre), and food/'
        'service retailers (KFC, Java House).')

    # Q5
    question(doc, 'FIVE', None, marks='15 marks')
    sub_q(doc, '', 'According to the concept of retailing, a retailer '
          'doesn\'t sell products in bulk; instead sells the '
          'merchandise in small units to the end-users. Explain the '
          'conditions available to Manufacturer Suggested Retail '
          'Price (MSRP) (also called List Price or Recommended '
          'Retail Price).', '15 marks')
    ans_label(doc)
    ans(doc,
        'The Manufacturer Suggested Retail Price (MSRP) — also '
        'called the List Price or Recommended Retail Price (RRP) — '
        'is the price at which the manufacturer recommends that the '
        'retailer sell the product to the end consumer. It is '
        'printed on the product, in the catalogue or in the '
        'manufacturer\'s price list. Whether and how a retailer can '
        'apply MSRP depends on a set of conditions:')
    subheading(doc, 'Key Conditions Surrounding MSRP')
    ans_bullet(doc, '1. Recommendation, not obligation: ',
               'In most jurisdictions (including Kenya under the '
               'Competition Act 2010) MSRP is a suggestion only; '
               'forcing a retailer to sell at MSRP — known as Resale '
               'Price Maintenance (RPM) — is generally illegal as it '
               'restricts competition.')
    ans_bullet(doc, '2. Branded and consumer-recognisable products: ',
               'MSRP works best on well-known branded merchandise '
               '(electronics, vehicles, books, pharmaceuticals) where '
               'consumers can compare prices across retailers.')
    ans_bullet(doc, '3. Adequate distribution margin: ',
               'The MSRP must include enough margin between the '
               'manufacturer\'s wholesale price and the MSRP to '
               'cover the retailer\'s operating costs and profit; '
               'otherwise retailers will refuse to stock.')
    ans_bullet(doc, '4. Stable supply & cost environment: ',
               'MSRP relies on predictable input costs. In high-'
               'inflation or currency-volatile markets like Kenya, '
               'MSRPs are revised frequently or replaced with '
               '"price on application".')
    ans_bullet(doc, '5. Anti-counterfeit & price-protection benefits: ',
               'MSRP signals a fair benchmark, helping consumers '
               'identify suspiciously low prices that may indicate '
               'counterfeits (common with electronics, cosmetics).')
    ans_bullet(doc, '6. Used as a reference for promotions: ',
               'Retailers commonly advertise discounts as "X% off '
               'RRP" or "Was Ksh 4,999, Now Ksh 3,499". Consumer '
               'protection regulators require the original RRP to '
               'be genuine and in recent use, otherwise the '
               'promotion is deceptive.')
    ans_bullet(doc, '7. Channel-conflict management: ',
               'Manufacturers use MSRP to discourage destructive '
               'price wars between their distributors and retailers.')
    ans_bullet(doc, '8. Brand-positioning safeguard: ',
               'Premium brands (Apple, Sony, Mercedes-Benz) enforce '
               'MSRP-aligned policies to protect brand equity from '
               'discounting.')
    ans_bullet(doc, '9. Discount tolerance: ',
               'MSRP allows retailers to discount downwards (sales, '
               'clearance) but not upwards on most goods; in price-'
               'controlled categories (e.g. petroleum in Kenya, '
               'regulated by EPRA) the cap is statutory.')
    ans_bullet(doc, '10. Information asymmetry reduction: ',
               'In categories where consumers cannot easily judge '
               'value (e.g. medicines), MSRP provides a trust '
               'anchor.')
    ans(doc,
        'In summary, MSRP is a guidance instrument that balances '
        'manufacturer interests (brand equity, channel discipline, '
        'consumer trust) with retailer freedom (the right to '
        'discount). It works best for branded, consumer-recognisable '
        'goods in stable cost environments, and must be applied '
        'within competition-law boundaries.')

    # Q6
    question(doc, 'SIX', None, marks='15 marks')
    sub_q(doc, '', 'What do you understand by the prestige pricing '
          'model? Under what circumstances can a retailer apply this '
          'model to their merchandise?', '15 marks')
    ans_label(doc)
    subheading(doc, 'Definition')
    ans(doc,
        'Prestige pricing — also called premium pricing or image '
        'pricing — is a pricing strategy in which a retailer '
        'deliberately sets a high price for its merchandise to '
        'signal superior quality, exclusivity, status and '
        'desirability. The high price is itself part of the value '
        'proposition: the consumer interprets the price as a '
        'quality cue and as a badge of social standing. It '
        'leverages Veblen-good behaviour, where demand actually '
        'rises as price rises, contradicting the standard demand '
        'curve.')
    subheading(doc, 'How Prestige Pricing Works')
    ans_bullet(doc, 'Quality signalling: ',
               'Customers use price as a heuristic for quality when '
               'they cannot easily evaluate the product technically.')
    ans_bullet(doc, 'Status & identity: ',
               'High-price products confer social status (Veblen '
               'effect) — owning a Rolex, Apple iPhone Pro Max, or '
               'a designer handbag signals success.')
    ans_bullet(doc, 'Exclusivity: ',
               'High prices restrict the customer base, preserving '
               'a sense of rarity (Snob effect).')
    ans_bullet(doc, 'High margin / lower volume: ',
               'The retailer earns a high margin per unit, '
               'compensating for lower transaction volumes.')
    subheading(doc, 'Circumstances Where a Retailer Can Apply Prestige Pricing')
    ans_bullet(doc, '1. Strong, prestigious brand equity: ',
               'The brand must have demonstrable heritage, '
               'craftsmanship, design or scarcity (e.g. Apple, '
               'Mercedes-Benz, Rolex, Louis Vuitton). In Kenya, '
               'high-end car dealers like DT Dobie, premium watch '
               'retailers in Two Rivers Mall, and Sarova/Fairmont '
               'hospitality outlets apply prestige pricing.')
    ans_bullet(doc, '2. Affluent target market: ',
               'A clearly identifiable upper-income segment willing '
               'and able to pay (Nairobi\'s upper-middle class, '
               'expatriate community, top-end SMEs).')
    ans_bullet(doc, '3. Differentiated, hard-to-compare product: ',
               'Goods where quality is genuinely superior or where '
               'comparison is difficult (luxury fashion, fine '
               'wine, jewellery, premium cosmetics).')
    ans_bullet(doc, '4. Limited supply / scarcity: ',
               'Hand-made, limited-edition or imported goods. '
               'Scarcity supports the high-price narrative.')
    ans_bullet(doc, '5. Premium retail environment: ',
               'Stores must reinforce the premium position — '
               'plush décor, attentive personal service, premium '
               'packaging, exclusive locations (Westgate, Sarit, '
               'Two Rivers, Village Market).')
    ans_bullet(doc, '6. Strong promotional & PR support: ',
               'Aspirational advertising, celebrity endorsement, '
               'high-end events. Mass-market promotion would '
               'dilute the prestige.')
    ans_bullet(doc, '7. Weak price-elasticity in the segment: ',
               'Demand is relatively insensitive to price — even '
               'inelastic, in Veblen-good categories.')
    ans_bullet(doc, '8. Service and after-sales excellence: ',
               'Premium clients expect concierge-level service, '
               'home delivery, loyalty privileges and seamless '
               'after-sales.')
    ans_bullet(doc, '9. Long-term brand-building objective: ',
               'The retailer is willing to forgo short-term volume '
               'for long-term brand equity.')
    ans_bullet(doc, '10. Stable competitive landscape: ',
               'Few or no aggressive low-price competitors in the '
               'same niche.')
    ans(doc,
        'Risks include alienating mass-market customers, '
        'vulnerability to economic downturns (luxury demand '
        'falls fastest in recessions), and reputational damage '
        'if quality fails to match price. When applied with '
        'discipline, however, prestige pricing builds enduring '
        'margin pools and brand defensibility — which is why '
        'the world\'s most valuable consumer brands rely on it.')

    # ════════════════════════════════════════════════════════════════════
    # PAPER 2 — APRIL 2023
    # ════════════════════════════════════════════════════════════════════
    paper_title(doc, 'PAPER TWO — APRIL 2023')
    body(doc,
         'Instructions: Answer Question One and any other three '
         'questions.', after=8)

    # Q1
    question(doc, 'ONE', None, marks='Compulsory — 25 marks')
    sub_q(doc, '(a)', 'Explain the functions of retailing.', '5 marks')
    ans_label(doc)
    ans_bullet(doc, 'Provision of an assortment: ',
               'Retailers gather products from many manufacturers '
               'into one convenient assortment so customers can '
               'one-stop-shop.')
    ans_bullet(doc, 'Breaking bulk: ',
               'Retailers buy in large quantities and resell in '
               'small consumer-friendly units.')
    ans_bullet(doc, 'Holding inventory: ',
               'Stocks are held so customers can buy at the moment '
               'of need rather than wait for a manufacturing run.')
    ans_bullet(doc, 'Providing services: ',
               'Information, advice, after-sales support, returns, '
               'credit, delivery, installation.')
    ans_bullet(doc, 'Place and time utility: ',
               'Goods are made available at the right place '
               '(near the consumer) and the right time (when '
               'needed).')
    ans_bullet(doc, 'Bridging the producer-consumer gap: ',
               'Retailers communicate consumer feedback back to '
               'manufacturers, supporting product improvement.')

    sub_q(doc, '(b)', 'Discuss store retailers and non-store '
          'retailers.', '10 marks')
    ans_label(doc)
    subheading(doc, 'Store Retailers')
    ans(doc,
        'Store retailers operate from a fixed physical location '
        '(shop, supermarket, mall stand) where customers visit to '
        'browse and buy. The store itself is the primary marketing '
        'asset: location, layout, ambience, staff and merchandise '
        'come together to create the brand experience.')
    ans_bullet(doc, 'Types include: ',
               'Department stores, supermarkets, hypermarkets, '
               'convenience stores, specialty stores, discount '
               'stores, warehouse clubs, category killers, factory '
               'outlets and food-service retailers.')
    ans_bullet(doc, 'Strengths: ',
               'Tactile experience, immediate possession, '
               'personal service, brand showcase, social '
               'experience.')
    ans_bullet(doc, 'Weaknesses: ',
               'High rent and labour costs, geographic '
               'limitations, fixed opening hours.')
    ans_bullet(doc, 'Kenyan examples: ',
               'Naivas, Quickmart, Carrefour, Bata, Hotpoint, '
               'KFC, Java House.')
    subheading(doc, 'Non-Store Retailers')
    ans(doc,
        'Non-store retailers reach customers without a fixed '
        'physical store. The transaction happens via direct '
        'channels — mail, phone, internet, vending, door-to-door, '
        'or automatic.')
    ans_bullet(doc, 'Types include: ',
               'Direct selling (Avon, Tianshi), direct marketing '
               '(catalogues, telemarketing), automated vending '
               'machines, online retailing/e-tailing (Jumia, '
               'Kilimall), social commerce, TV shopping, '
               'subscription boxes.')
    ans_bullet(doc, 'Strengths: ',
               'Wider geographic reach, lower overheads, 24/7 '
               'availability, rich data capture, scalability.')
    ans_bullet(doc, 'Weaknesses: ',
               'No tactile experience, delivery dependency, '
               'higher return rates, cyber-security and trust '
               'issues, last-mile cost.')
    ans_bullet(doc, 'Kenyan examples: ',
               'Jumia, Kilimall, Copia, Glovo Market, Naivas Now, '
               'WhatsApp/Instagram boutiques, Twiga Foods (B2B '
               'mobile app).')

    sub_q(doc, '(c)', 'Describe the five major bases for classifying '
          'retail outlets.', '5 marks')
    ans_label(doc)
    ans_bullet(doc, '1. Form of ownership: ',
               'Independent, chain stores, franchises, '
               'cooperatives, consumer cooperatives, leased '
               'departments.')
    ans_bullet(doc, '2. Merchandise carried (assortment): ',
               'Breadth and depth — convenience, supermarket, '
               'specialty, department store, hypermarket.')
    ans_bullet(doc, '3. Pricing strategy: ',
               'EDLP (every-day-low-pricing) vs High-Low; '
               'discount, off-price, full-price, prestige.')
    ans_bullet(doc, '4. Service level: ',
               'Self-service (Naivas), self-selection, limited '
               'service, full service (boutiques, jewellers).')
    ans_bullet(doc, '5. Location: ',
               'Central Business District (CBD), shopping malls, '
               'neighbourhood / strip locations, free-standing, '
               'online (no physical location).')

    sub_q(doc, '(d)', 'Explain the key financial records that a '
          'retail manager should keep at all times.', '5 marks')
    ans_label(doc)
    ans_bullet(doc, 'Sales records / Daily Sales Reports (DSR): ',
               'Track turnover by department, SKU, payment method '
               'and time-of-day.')
    ans_bullet(doc, 'Cash book / Petty cash book: ',
               'Record of cash received and cash paid out.')
    ans_bullet(doc, 'Purchases / Payables records: ',
               'Invoices from suppliers, GRNs (Goods Received '
               'Notes), and outstanding amounts owed.')
    ans_bullet(doc, 'Inventory / Stock records: ',
               'Stock-on-hand, GRN, stock-out, stock-take, '
               'shrinkage and aged stock.')
    ans_bullet(doc, 'Income Statement (P&L): ',
               'Periodic record of revenue, COGS, expenses and '
               'net profit.')
    ans_bullet(doc, 'Balance Sheet: ',
               'Statement of assets, liabilities and equity at a '
               'point in time.')
    ans_bullet(doc, 'Cash-Flow Statement: ',
               'Operating, investing and financing cash flows.')
    ans_bullet(doc, 'Payroll & statutory records: ',
               'Salaries, NSSF, NHIF/SHA, PAYE, NITA.')
    ans_bullet(doc, 'Tax records: ',
               'VAT (KRA iTax), corporation tax, withholding tax, '
               'turnover tax for SMEs.')

    # Q2
    question(doc, 'TWO', None, marks='15 marks')
    sub_q(doc, '(a)', 'List and explain the many types of retail '
          'outlets.', '10 marks')
    ans_label(doc)
    ans_bullet(doc, 'Department Store: ',
               'Wide-and-deep assortment organised in departments '
               '(historically: Selfridges, Macy\'s).')
    ans_bullet(doc, 'Supermarket: ',
               'Self-service food and household goods (Naivas, '
               'Quickmart).')
    ans_bullet(doc, 'Hypermarket: ',
               'Supermarket + general merchandise + apparel + '
               'electronics under one roof (Carrefour, Naivas '
               'Hypermart).')
    ans_bullet(doc, 'Convenience Store: ',
               'Small, long hours, top-up shopping (Naivas Mtaa '
               'Wetu, Shell Select).')
    ans_bullet(doc, 'Specialty Store: ',
               'Narrow but deep — single category (Bata, Vivo '
               'Activewear, Hotpoint).')
    ans_bullet(doc, 'Discount Store: ',
               'Broad assortment at low price (Mwananchi, '
               'Eastmatt).')
    ans_bullet(doc, 'Off-Price Retailer: ',
               'Brand-name goods at deep discounts (factory '
               'outlets, end-of-season clearance hubs).')
    ans_bullet(doc, 'Warehouse / Cash & Carry: ',
               'Bulk format for SMEs (Mahitaji, Eastmatt '
               'Wholesale).')
    ans_bullet(doc, 'Category Killer: ',
               'Large specialty store dominating one category '
               '(Game stores at peak; Mr Price Home for furniture).')
    ans_bullet(doc, 'Factory Outlet: ',
               'Manufacturer-owned discount store (Bata Industrials).')
    ans_bullet(doc, 'Online / E-Tailer: ',
               'Pure-play digital — Jumia, Kilimall, Copia.')
    ans_bullet(doc, 'Pop-Up Store: ',
               'Temporary, event-driven physical store (Vivo '
               'pop-ups, Black Friday container shops).')
    ans_bullet(doc, 'Food Service / QSR: ',
               'Restaurants, fast-food outlets — KFC, Java House, '
               'Pizza Inn.')

    sub_q(doc, '(b)', 'Apart from independent retail stores, list and '
          'discuss four (4) other types of retail store ownership.',
          '5 marks')
    ans_label(doc)
    ans_bullet(doc, '1. Chain Stores (Corporate Chains): ',
               'A single owner operates two or more outlets under '
               'a common name, central buying and uniform '
               'merchandising. Benefits include economies of '
               'scale, brand consistency and bargaining power '
               '(e.g. Naivas, Quickmart, Carrefour).')
    ans_bullet(doc, '2. Franchise Stores: ',
               'The franchisor licences its brand, systems and '
               'know-how to a franchisee in exchange for fees and '
               'royalties; the franchisee invests capital and '
               'operates locally (e.g. KFC, Java House, '
               'Subway, Total Bonjour, Pizza Inn).')
    ans_bullet(doc, '3. Cooperative Retail Stores: ',
               'Owned and managed by members (consumers or '
               'producers) who share profits as patronage '
               'dividends. The Kenya National Federation of '
               'Cooperatives represents many; SACCOs sometimes '
               'run consumer outlets.')
    ans_bullet(doc, '4. Leased Department / Concession Stores: ',
               'A retailer leases part of its store to a '
               'specialist operator who runs that department '
               '(e.g. cosmetics counters in department stores, '
               'restaurants and barber shops within Naivas '
               'hypermarkets, mobile-phone counters within '
               'Carrefour).')

    # Q3
    question(doc, 'THREE', None, marks='15 marks')
    sub_q(doc, '(a)', 'Based on the length and breadth of product '
          'assortment, list and describe six (6) types of retail '
          'outlets.', '10 marks')
    ans_label(doc)
    ans(doc,
        'Breadth = number of different product categories; '
        'Length/Depth = number of items within each category. '
        'Combining these two dimensions gives six characteristic '
        'retail formats:')
    ans_bullet(doc, '1. Convenience Store (Narrow & Shallow): ',
               'Few categories, few items per category — top-up '
               'essentials. Naivas Mtaa Wetu, Shell Select.')
    ans_bullet(doc, '2. Specialty Store (Narrow & Deep): ',
               'One category, many items/variants. Bata for '
               'shoes, Vivo Activewear for women\'s fashion, '
               'Hotpoint for electronics.')
    ans_bullet(doc, '3. Limited-line Store (Narrow & Moderate): ',
               'Focuses on a few related categories with moderate '
               'depth — e.g. baby-shop, pet-shop, photo-studio.')
    ans_bullet(doc, '4. Supermarket (Wide & Moderate): ',
               'Wide range of food, beverage, household and '
               'personal-care categories with moderate depth per '
               'category. Naivas, Quickmart, Chandarana FoodPlus.')
    ans_bullet(doc, '5. Department Store (Wide & Deep): ',
               'Many categories (apparel, cosmetics, household, '
               'electronics) with deep assortment in each — '
               'historically Macy\'s, Selfridges; in Kenya the '
               'closest current example is Carrefour\'s general-'
               'merchandise floor.')
    ans_bullet(doc, '6. Hypermarket / Superstore (Very Wide & Deep): ',
               'Combines supermarket + department-store + '
               'electronics + furniture under one roof — Carrefour '
               'Two Rivers, Naivas Hypermart Westlands, '
               'Quickmart Hypermart.')

    sub_q(doc, '(b)', 'Explain the role of visual merchandising.',
          '5 marks')
    ans_label(doc)
    ans(doc,
        'Visual Merchandising (VM) is the strategic presentation '
        'of products in a retail environment to attract attention, '
        'communicate brand and drive sales. Its key roles are:')
    ans_bullet(doc, 'Capture attention: ',
               'Eye-catching window displays and end-caps stop '
               'passers-by and pull them into the store.')
    ans_bullet(doc, 'Communicate brand identity: ',
               'Colour palette, signage, props and lighting tell '
               'the brand story (e.g. Apple\'s minimalism, '
               'Lululemon\'s wellness vibe, Vivo\'s vibrant '
               'African prints).')
    ans_bullet(doc, 'Direct customer flow: ',
               'Layout, sight-lines and focal points guide '
               'shoppers around the store and past key '
               'categories.')
    ans_bullet(doc, 'Stimulate impulse purchase: ',
               'Cross-merchandising (e.g. wine next to cheese, '
               'sun-cream next to swimwear) and check-out displays '
               'lift basket size.')
    ans_bullet(doc, 'Educate and inform: ',
               'POP signage explains features, benefits and prices '
               'so customers self-serve confidently.')
    ans_bullet(doc, 'Support promotions: ',
               'Themed displays for Christmas, Eid, Back-to-School, '
               'Valentines drive seasonal lift.')
    ans_bullet(doc, 'Maximise space productivity: ',
               'Sales-per-square-foot rises when fixtures, '
               'planograms and displays are optimised.')

    # Q4
    question(doc, 'FOUR', None, marks='15 marks')
    sub_q(doc, '(a)', 'Discuss factors to be considered when choosing '
          'a retail location.', '8 marks')
    ans_label(doc)
    ans_bullet(doc, '1. Customer demographics & catchment: ',
               'Population size, income, age, family-life-cycle '
               'and lifestyle of the surrounding area must match '
               'the retailer\'s target market.')
    ans_bullet(doc, '2. Footfall and traffic: ',
               'High pedestrian and vehicular traffic — CBDs, '
               'busy malls, transport hubs (Globe Roundabout, '
               'OTC, Westlands, Eastleigh).')
    ans_bullet(doc, '3. Visibility and accessibility: ',
               'Easy to see from the road, easy to enter; '
               'adequate parking and matatu/bus access.')
    ans_bullet(doc, '4. Competition: ',
               'Proximity of competitors (cluster benefit for '
               'comparison shopping vs cannibalisation risk).')
    ans_bullet(doc, '5. Co-tenancy / anchor stores: ',
               'In malls, presence of complementary anchor '
               'tenants (Naivas, Carrefour, KFC) drives traffic '
               'to satellite stores.')
    ans_bullet(doc, '6. Cost factors: ',
               'Rent, service charge, set-up, fit-out and '
               'utilities; must be sustainable relative to '
               'expected sales/sq ft.')
    ans_bullet(doc, '7. Legal & zoning compliance: ',
               'County licences, change-of-use permits, fire-'
               'safety, NEMA approvals.')
    ans_bullet(doc, '8. Infrastructure: ',
               'Reliable electricity, water, internet, security, '
               'waste disposal.')
    ans_bullet(doc, '9. Future development of the area: ',
               'Roads, BRT, real-estate growth (e.g. Tatu City, '
               'Konza, Northlands).')
    ans_bullet(doc, '10. Brand fit: ',
               'Premium brands need premium locations '
               '(Two Rivers, Sarit, Village Market); discount '
               'brands need high-density, lower-rent areas.')

    sub_q(doc, '(b)', 'With relevant examples, enumerate dynamics in '
          'retail business.', '7 marks')
    ans_label(doc)
    ans(doc,
        'The retail business is highly dynamic — constantly '
        'shaped by shifts in technology, consumer behaviour and '
        'the competitive landscape. The key dynamics are:')
    ans_bullet(doc, '1. Digital and mobile transformation: ',
               'Move from purely physical to omnichannel; e.g. '
               'Naivas Now app, Carrefour app, Jumia.')
    ans_bullet(doc, '2. Rise of mobile money: ',
               'M-PESA STK push, Lipa Na M-PESA and pay-bill '
               'have made digital payments universal in Kenya.')
    ans_bullet(doc, '3. Changing consumer expectations: ',
               'Today\'s shopper expects speed, convenience, '
               'personalisation and value.')
    ans_bullet(doc, '4. Format innovation: ',
               'Small-format convenience stores, dark stores, '
               'pop-ups and 15-minute delivery (Glovo, Jumia '
               'Express).')
    ans_bullet(doc, '5. Data-driven merchandising: ',
               'Loyalty data, AI recommendations and dynamic '
               'pricing.')
    ans_bullet(doc, '6. Supply-chain disruption: ',
               'COVID-19, Russia-Ukraine war, drought and '
               'currency volatility have forced agility.')
    ans_bullet(doc, '7. Sustainability and ethics: ',
               'Rising consumer demand for green packaging, '
               'fair trade and local sourcing.')
    ans_bullet(doc, '8. Consolidation and exits: ',
               'Nakumatt, Tuskys and Uchumi exited; Naivas, '
               'Quickmart and Carrefour expanded — illustrating '
               'Darwinian retail dynamics.')

    # Q5
    question(doc, 'FIVE', None, marks='15 marks')
    sub_q(doc, '(a)', 'Discuss the concept of retail promotion mix.',
          '9 marks')
    ans_label(doc)
    ans(doc,
        'The retail promotion mix is the integrated set of '
        'communication tools a retailer uses to inform, persuade '
        'and remind customers about its merchandise, services '
        'and brand. It comprises the following elements:')
    ans_bullet(doc, '1. Advertising: ',
               'Paid mass-media communication — TV, radio, print, '
               'outdoor billboards, digital display ads (e.g. '
               'Naivas TV adverts, Quickmart radio jingles, '
               'Jumia bus-stop billboards).')
    ans_bullet(doc, '2. Sales promotion: ',
               'Short-term incentives — discounts, coupons, '
               'BOGOF (buy-one-get-one-free), bundle deals, '
               'loyalty points (Naivas Linda Card).')
    ans_bullet(doc, '3. Personal selling: ',
               'Face-to-face interaction with customers by '
               'sales associates — particularly important in '
               'fashion, electronics, cars, jewellery.')
    ans_bullet(doc, '4. Public relations & publicity: ',
               'CSR activities, sponsorships, media coverage, '
               'community events (Safaricom Foundation, Naivas '
               '"Mtaa Wetu" community drives).')
    ans_bullet(doc, '5. Direct marketing: ',
               'Personalised communication — SMS, email, '
               'WhatsApp Business, app push notifications.')
    ans_bullet(doc, '6. Digital & social-media marketing: ',
               'Facebook, Instagram, TikTok, YouTube, search-'
               'engine marketing, influencer collaborations.')
    ans_bullet(doc, '7. Visual merchandising & in-store displays: ',
               'Window dressings, end-caps, POP signage — silent '
               'salesmen.')
    ans_bullet(doc, '8. Loyalty / CRM programs: ',
               'Long-term engagement through points, tiers, '
               'birthdays and personalised offers.')

    sub_q(doc, '(b)', 'Explain three factors affecting retailers.',
          '6 marks')
    ans_label(doc)
    ans_bullet(doc, '1. Economic environment: ',
               'Inflation, currency depreciation, interest rates, '
               'unemployment and disposable income directly '
               'affect consumer spending power and the cost of '
               'imported merchandise. The 2022-23 inflation in '
               'Kenya forced retailers to revise prices weekly.')
    ans_bullet(doc, '2. Technology: ',
               'E-commerce, mobile money, AI, cloud POS and '
               'automated logistics have lowered the cost of '
               'going omnichannel and raised customer '
               'expectations. Retailers that lag (Nakumatt, '
               'Uchumi) struggle.')
    ans_bullet(doc, '3. Competition: ',
               'New entrants (Carrefour\'s 2016 entry into '
               'Kenya, super-apps like Glovo) and existing '
               'rivals continually pressure prices, '
               'assortments and service. Quickmart\'s sub-100-'
               'KES "value packs" forced rivals to respond.')
    ans(doc,
        'Other factors worth knowing for full marks include: '
        'regulation (KEBS, Competition Authority, county '
        'licensing); social and cultural factors (changing '
        'lifestyles, religious considerations during Ramadan, '
        'urbanisation); demographics (youth bulge, growing '
        'middle class); and the supplier landscape '
        '(consolidation among FMCG manufacturers).')

    # Q6
    question(doc, 'SIX', None, marks='15 marks')
    sub_q(doc, '(a)', 'Explain the specific functions of a retail '
          'supervisor.', '10 marks')
    ans_label(doc)
    ans(doc,
        'A retail supervisor (or floor / department / store '
        'supervisor) is the front-line manager responsible for '
        'translating strategy into day-to-day execution on the '
        'shop floor. Specific functions include:')
    ans_bullet(doc, '1. Staff supervision and rostering: ',
               'Allocating shifts, coaching cashiers and floor '
               'attendants, monitoring punctuality and grooming '
               'standards.')
    ans_bullet(doc, '2. Customer service & complaint handling: ',
               'Being the first escalation point for customer '
               'queries; ensuring resolutions are fair and fast.')
    ans_bullet(doc, '3. Visual merchandising & shelf management: ',
               'Implementing planograms, refreshing displays, '
               'ensuring price tags and POP material are correct.')
    ans_bullet(doc, '4. Stock & inventory management: ',
               'Monitoring stock levels, raising replenishment '
               'requests, conducting cycle counts and shrinkage '
               'investigations.')
    ans_bullet(doc, '5. Sales target tracking: ',
               'Reviewing daily/weekly KPIs (sales, AOV, UPT, '
               'conversion), motivating the team to hit them.')
    ans_bullet(doc, '6. Cash & POS management: ',
               'Opening and closing tills, reconciling cash, '
               'managing M-PESA pay-bills, banking float.')
    ans_bullet(doc, '7. Loss prevention: ',
               'Watching for shoplifting, internal pilferage, '
               'damage and waste; implementing controls.')
    ans_bullet(doc, '8. Health, safety & hygiene: ',
               'Ensuring fire-safety compliance, cleanliness, '
               'PPE and food-safety standards (cold-chain, '
               'expiry).')
    ans_bullet(doc, '9. Promotions implementation: ',
               'Setting up advertised offers correctly on the '
               'floor and at the till; communicating to staff.')
    ans_bullet(doc, '10. Reporting upwards: ',
               'Daily/weekly reports to the store manager — '
               'sales, stock issues, staff issues, customer '
               'feedback.')

    sub_q(doc, '(b)', 'Discuss the emerging trends in the retail '
          'sector.', '5 marks')
    ans_label(doc)
    ans_bullet(doc, '1. Omnichannel and phygital retail: ',
               'Seamless integration of physical and digital — '
               'BOPIS, ship-from-store, app-and-store loyalty.')
    ans_bullet(doc, '2. AI & hyper-personalisation: ',
               'Recommendations, chatbots, dynamic pricing, '
               'generative-AI product copy and imagery.')
    ans_bullet(doc, '3. Mobile money & mobile commerce: ',
               'M-PESA, Airtel Money, super-apps; >70% of '
               'Kenyan retail transactions touch mobile money.')
    ans_bullet(doc, '4. Social commerce & live shopping: ',
               'TikTok Shop, Instagram, WhatsApp catalogues, '
               'live-stream selling.')
    ans_bullet(doc, '5. Sustainability and circular retail: ',
               'Green packaging, reusable bags, refill stations, '
               'second-hand and resale markets.')
    ans_bullet(doc, '6. Quick commerce / 15-minute delivery: ',
               'Dark stores and rider networks (Glovo, Bolt '
               'Food, Jumia Express).')
    ans_bullet(doc, '7. Cashier-less and self-checkout stores: ',
               'Self-checkout kiosks (Naivas, Carrefour pilot '
               'stores), and globally Amazon Go.')
    ans_bullet(doc, '8. Subscription & replenishment models: ',
               'Recurring orders for groceries, beauty, pet '
               'food.')
    ans_bullet(doc, '9. AR / VR shopping experiences: ',
               'Virtual try-on for fashion, furniture and '
               'cosmetics.')
    ans_bullet(doc, '10. Data-driven loyalty 2.0: ',
               'Tiered, gamified, personalised loyalty (Naivas '
               'Linda, Carrefour MyClub, Total Club).')

    # ════════════════════════════════════════════════════════════════════
    # PAPER 3 — JULY 2025
    # ════════════════════════════════════════════════════════════════════
    paper_title(doc, 'PAPER THREE — JULY 2025')
    body(doc,
         'Instructions: Answer Question One and any other three '
         'questions.', after=8)

    # Q1
    question(doc, 'ONE', None, marks='Compulsory — 25 marks')
    sub_q(doc, '(a)', 'Use the PESTEL Analysis of a retail business '
          'and cite notable examples here in Kenya.', '5 marks')
    ans_label(doc)
    ans(doc,
        'PESTEL analyses the macro-environment in which a '
        'retailer operates across six dimensions:')
    ans_bullet(doc, 'P – Political: ',
               'Government stability, taxation policy, county '
               'licensing, trade agreements (EAC, AfCFTA). E.g. '
               'higher fuel levies and VAT changes affect retail '
               'pricing decisions.')
    ans_bullet(doc, 'E – Economic: ',
               'Inflation, exchange rates (KES vs USD), interest '
               'rates and disposable income. The 2022-23 '
               'inflation forced supermarkets like Naivas to '
               'introduce value packs.')
    ans_bullet(doc, 'S – Social/Cultural: ',
               'Urbanisation, youth bulge, rising middle class, '
               'religious considerations (halal aisles for '
               'Eastleigh shoppers, Ramadan promotions), changing '
               'family structures.')
    ans_bullet(doc, 'T – Technological: ',
               'M-PESA, smartphones, fibre internet, AI, e-'
               'commerce platforms. Jumia, Glovo and Naivas Now '
               'are direct beneficiaries.')
    ans_bullet(doc, 'E – Environmental/Ecological: ',
               'Plastic-bag ban (NEMA 2017), drought affecting '
               'fresh produce supply, sustainability '
               'expectations.')
    ans_bullet(doc, 'L – Legal: ',
               'Consumer Protection Act 2012, Competition Act, '
               'Data Protection Act 2019, KEBS standards, food-'
               'safety regulations, employment law.')

    sub_q(doc, '(b)', 'Discuss the role of social-media marketing in '
          'creating brand awareness. How do Kenyan retailers like '
          'Jumia use social media to boost sales?', '5 marks')
    ans_label(doc)
    ans(doc,
        'Social-media marketing uses platforms such as Facebook, '
        'Instagram, TikTok, X (Twitter), YouTube and WhatsApp to '
        'build brand awareness, engage customers and drive '
        'sales. Roles in brand awareness include:')
    ans_bullet(doc, 'Reach & visibility: ',
               'Paid and organic posts put the brand in front '
               'of millions cheaply.')
    ans_bullet(doc, 'Two-way engagement: ',
               'Comments, DMs and reviews build relationships.')
    ans_bullet(doc, 'Influencer marketing: ',
               'Trusted creators amplify the message to '
               'targeted audiences.')
    ans_bullet(doc, 'User-generated content: ',
               'Customer photos and reviews create authentic '
               'social proof.')
    ans_bullet(doc, 'Targeted advertising: ',
               'Demographic, interest and look-alike targeting '
               'minimises waste.')
    ans_bullet(doc, 'Community building: ',
               'Branded hashtags, groups and challenges build '
               'tribe-like loyalty.')
    ans(doc,
        'Jumia in Kenya specifically uses: (i) flash-sale '
        'campaigns ("Jumia Black Friday", "Mobile Week", '
        '"Jumia Anniversary") amplified across Instagram, '
        'TikTok and Facebook; (ii) influencer partnerships '
        'with local celebrities and micro-influencers to '
        'demo products; (iii) targeted Facebook/Instagram '
        'ads with retargeting on abandoned carts; (iv) live '
        'shopping streams; (v) WhatsApp Business catalogues '
        'and customer-care; and (vi) JumiaPay promotions '
        'driving conversion.')

    sub_q(doc, '(c)', 'Using the ABC analysis method, how might a '
          'retailer categorise products like sugar, electronics or '
          'stationery? Why is this classification useful in '
          'inventory planning?', '6 marks')
    ans_label(doc)
    ans(doc,
        'ABC analysis (also known as the Pareto or 80/20 rule) '
        'classifies inventory items into three groups based on '
        'their contribution to sales value (or profit):')
    ans_bullet(doc, 'A-items (high value, ~10-20% of SKUs, ~70-80% '
                    'of sales): ',
               'Top-selling, high-margin items requiring tight '
               'control. E.g. fast-moving sugar (Mumias, Kabras) '
               'and best-selling smartphones (Tecno, Samsung A-'
               'series, iPhone) in an electronics store.')
    ans_bullet(doc, 'B-items (moderate value, ~30% of SKUs, ~15-20% '
                    'of sales): ',
               'Mid-tier items requiring routine control. E.g. '
               'mid-range stationery brands, mid-range '
               'televisions and home-appliances.')
    ans_bullet(doc, 'C-items (low value, ~50% of SKUs, ~5-10% of '
                    'sales): ',
               'Slow-moving long-tail items requiring relaxed '
               'controls. E.g. specialty stationery (calligraphy '
               'pens, rare colours), niche electronics '
               'accessories, premium-brand sugar.')
    subheading(doc, 'Usefulness in Inventory Planning')
    ans_bullet(doc, 'Focus management attention: ',
               'A-items get daily monitoring, accurate forecasts '
               'and frequent replenishment.')
    ans_bullet(doc, 'Optimise working capital: ',
               'B and C items are reviewed less often; C items '
               'may be moved to vendor-managed inventory or '
               'delisted.')
    ans_bullet(doc, 'Reduce stock-outs on critical items: ',
               'A-items receive safety-stock buffers because '
               'their stock-out is most damaging.')
    ans_bullet(doc, 'Avoid over-stocking the long tail: ',
               'C-items get tighter caps and more aggressive '
               'markdown discipline.')
    ans_bullet(doc, 'Negotiation leverage: ',
               'A-item suppliers receive priority negotiation '
               'and volume contracts.')
    ans_bullet(doc, 'Supports SKU rationalisation: ',
               'Persistent C-items that earn no margin are '
               'candidates for delisting.')

    sub_q(doc, '(d)', 'Formulate the Strategic Planning Process in '
          'a retail unit of a business.', '7 marks')
    ans_label(doc)
    ans(doc,
        'The strategic planning process in a retail unit '
        'follows a structured cycle:')
    ans_bullet(doc, '1. Define mission, vision and values: ',
               'Why we exist, what we aspire to, how we '
               'behave (e.g. Naivas: "Serving you better").')
    ans_bullet(doc, '2. Environmental scanning: ',
               'Internal (SWOT) and external (PESTEL + Porter\'s '
               '5-Forces) analysis.')
    ans_bullet(doc, '3. Set strategic objectives: ',
               'SMART goals — sales growth, market share, NPS, '
               'GMROI, store-count.')
    ans_bullet(doc, '4. Identify target market & positioning: ',
               'Segmentation, targeting, differentiation, '
               'positioning statement.')
    ans_bullet(doc, '5. Formulate the retail mix strategy: ',
               'Merchandise, pricing, store/format, location, '
               'communications, service.')
    ans_bullet(doc, '6. Resource allocation: ',
               'Capital, people, technology, inventory budget '
               '(Open-to-Buy).')
    ans_bullet(doc, '7. Implementation & execution: ',
               'Roll out plans across stores, channels and '
               'functions.')
    ans_bullet(doc, '8. Performance monitoring & control: ',
               'KPI dashboards, balanced scorecards, periodic '
               'reviews.')
    ans_bullet(doc, '9. Feedback and continuous improvement: ',
               'Refine the plan based on results — loop back '
               'to step 2.')

    sub_q(doc, '(e)', 'What do you understand by the term "Visual '
          'Merchandising"?', '2 marks')
    ans_label(doc)
    ans(doc,
        'Visual Merchandising (VM) is the strategic art and '
        'practice of presenting products in a retail '
        'environment — through window displays, store layout, '
        'fixtures, lighting, colour, signage and props — to '
        'attract attention, communicate the brand identity, '
        'guide customer flow and drive sales. In short, it is '
        '"the silent salesperson" of the store.')

    # Q2
    question(doc, 'TWO', None, marks='15 marks')
    sub_q(doc, '(a)', 'How can retail businesses use email marketing '
          'to enhance customer engagement and loyalty? What types of '
          'offers or content should they focus on?', '6 marks')
    ans_label(doc)
    ans(doc,
        'Email marketing is one of the highest-ROI channels in '
        'retail because it is owned (no platform fee), '
        'permission-based and easily personalised. Retailers can '
        'use it to enhance engagement and loyalty by:')
    ans_bullet(doc, 'Welcome / onboarding series: ',
               'Set expectations, share brand story, offer a '
               'first-purchase discount.')
    ans_bullet(doc, 'Personalised product recommendations: ',
               'Based on browsing, purchase and loyalty history.')
    ans_bullet(doc, 'Abandoned-cart and browse-recovery emails: ',
               'Recover lost sales — typically 10-15% recovery '
               'rate.')
    ans_bullet(doc, 'Loyalty updates: ',
               'Points balance, tier upgrades, expiry reminders, '
               'birthday rewards.')
    ans_bullet(doc, 'Post-purchase follow-up: ',
               'Order confirmation, shipping updates, review '
               'request, replenishment reminders.')
    ans_bullet(doc, 'Educational & lifestyle content: ',
               'Recipes, styling tips, how-to guides — keeps '
               'the brand top-of-mind.')
    ans_bullet(doc, 'VIP/early-access offers: ',
               'Reward best customers with first look at sales, '
               'limited editions, exclusive events.')
    subheading(doc, 'Types of Offers / Content to Focus On')
    ans_bullet(doc, '', 'Time-bound discounts and flash sales (24-'
               '72 hours).')
    ans_bullet(doc, '', 'Free-shipping or free-delivery thresholds.')
    ans_bullet(doc, '', 'Bundle deals and "Buy More, Save More".')
    ans_bullet(doc, '', 'Birthday & anniversary offers.')
    ans_bullet(doc, '', 'Seasonal/holiday campaigns (Easter, '
               'Madaraka, Eid, Christmas).')
    ans_bullet(doc, '', 'Refer-a-friend incentives.')
    ans_bullet(doc, '', 'New-arrival announcements with imagery '
               'and quick-shop links.')

    sub_q(doc, '(b)', 'With the rise of the internet and mobile '
          'technology, digital marketing has become essential in '
          'retail. List and explain various online channels '
          'retailers use to market and engage customers.',
          '6 marks')
    ans_label(doc)
    ans_bullet(doc, '1. Brand website / e-commerce store: ',
               'Owned digital storefront — the hub for content, '
               'product, search and checkout.')
    ans_bullet(doc, '2. Mobile app: ',
               'Personalised, push-notification capable, often '
               'integrated with loyalty (Naivas Now, Carrefour '
               'app).')
    ans_bullet(doc, '3. Search-engine marketing (SEM/SEO): ',
               'Google Ads + organic SEO capture purchase '
               'intent at the moment customers search.')
    ans_bullet(doc, '4. Social-media platforms: ',
               'Facebook, Instagram, TikTok, X, LinkedIn, '
               'Pinterest, YouTube — for awareness, '
               'engagement, social commerce.')
    ans_bullet(doc, '5. Email marketing: ',
               'Newsletters, promotions, transactional, '
               'lifecycle campaigns.')
    ans_bullet(doc, '6. SMS & WhatsApp Business: ',
               'High open rates; ideal for order confirmations, '
               'promos, customer service in Kenya.')
    ans_bullet(doc, '7. Influencer marketing: ',
               'Partnering with macro and micro influencers for '
               'reach and credibility.')
    ans_bullet(doc, '8. Online marketplaces: ',
               'Selling on Jumia, Kilimall, Sky.Garden, '
               'Glovo Market.')
    ans_bullet(doc, '9. Content marketing & blogs: ',
               'Long-form content (recipes, buying guides) for '
               'SEO and brand authority.')
    ans_bullet(doc, '10. Programmatic display & retargeting: ',
               'Banner ads served across the web to past '
               'visitors, lifting conversion.')
    ans_bullet(doc, '11. Affiliate marketing: ',
               'Pay partners a commission for sales they refer.')

    sub_q(doc, '(c)', 'What are the challenges of implementing CRM '
          'systems in small retail shops in Kenya?', '3 marks')
    ans_label(doc)
    ans_bullet(doc, 'High upfront cost: ',
               'Software licences, hardware, training and '
               'integration are unaffordable for many SMEs.')
    ans_bullet(doc, 'Limited ICT skills & literacy: ',
               'Owners and staff often lack the technical '
               'expertise to set up and maintain CRM systems.')
    ans_bullet(doc, 'Poor and inconsistent data quality: ',
               'Customer details captured manually are '
               'incomplete or inaccurate, undermining the CRM\'s '
               'value.')
    ans_bullet(doc, 'Unreliable infrastructure: ',
               'Frequent power outages, costly data and '
               'unstable internet hamper cloud-based CRMs.')
    ans_bullet(doc, 'Cultural and trust barriers: ',
               'Customers are reluctant to share personal '
               'data; concerns about privacy and SMS spam.')
    ans_bullet(doc, 'Integration with M-PESA, POS and stock: ',
               'Many off-the-shelf CRMs are not built for the '
               'Kenyan payment ecosystem.')
    ans_bullet(doc, 'Owner mindset: ',
               'Many small retailers focus on transactions, '
               'not relationships — CRM is seen as overhead, '
               'not investment.')

    # Q3
    question(doc, 'THREE', None, marks='15 marks')
    sub_q(doc, '(a)', 'How has the growth of mobile-payment systems '
          'like M-PESA impacted e-commerce in Kenya? Provide '
          'examples from local platforms.', '7 marks')
    ans_label(doc)
    ans(doc,
        'Since its launch by Safaricom in 2007, M-PESA has '
        'transformed Kenya into one of the most mobile-money-'
        'penetrated economies in the world (>96% of adults). '
        'Its impact on e-commerce has been profound:')
    ans_bullet(doc, '1. Closed the payment gap: ',
               'Before M-PESA, low credit-card penetration '
               'crippled online shopping. M-PESA enabled '
               'instant, trusted digital payment for the '
               'unbanked majority.')
    ans_bullet(doc, '2. Reduced cash-on-delivery (COD) reliance: ',
               'Jumia, Kilimall and others have steadily '
               'shifted from costly COD to upfront M-PESA '
               'payment via Lipa Na M-PESA / STK push.')
    ans_bullet(doc, '3. Boosted SME e-commerce: ',
               'WhatsApp/Instagram boutiques can now collect '
               'payment via Pay Bill or Buy Goods seconds '
               'after agreeing on a sale — no need for a '
               'merchant account.')
    ans_bullet(doc, '4. Enabled fintech and BNPL ecosystems: ',
               'M-PESA Overdraft (Fuliza), KCB-M-PESA, M-Shwari, '
               'and BNPL apps like Lipa Later, Aspira and '
               'Flexpay let consumers pay in instalments.')
    ans_bullet(doc, '5. Fuelled gig-economy logistics: ',
               'Glovo, Bolt Food, Little Riders, Sendy and '
               'Jumia Logistics riders are paid and tipped via '
               'M-PESA — making last-mile delivery viable.')
    ans_bullet(doc, '6. Powered super-apps: ',
               'M-PESA Super App now bundles bill-pay, '
               'shopping, transport and savings — pulling '
               'commerce inside the wallet.')
    ans_bullet(doc, '7. Cross-border and B2B reach: ',
               'M-PESA Global enables remittances; PesaLink '
               'and M-PESA Pay Bill enable B2B settlement '
               'between retailers and suppliers (e.g. Twiga '
               'Foods\' B2B platform).')
    ans(doc,
        'Local platform examples: Jumia (STK push checkout), '
        'Naivas Now, Carrefour app, Glovo, Bolt Food, Twiga '
        'Foods, Copia, Lipa Later, KCB Vooma, and thousands '
        'of WhatsApp/Instagram-based micro-retailers across '
        'Nairobi, Mombasa, Kisumu and Eldoret.')

    sub_q(doc, '(b)', 'Elaborate on inventory-control techniques '
          'available to retailers.', '8 marks')
    ans_label(doc)
    ans_bullet(doc, '1. Economic Order Quantity (EOQ): ',
               'A formula that calculates the order quantity '
               'minimising total ordering + holding cost. Best '
               'for stable-demand SKUs.')
    ans_bullet(doc, '2. Just-in-Time (JIT): ',
               'Inventory is ordered only as needed — minimises '
               'holding cost but requires reliable suppliers '
               '(e.g. fresh produce in Carrefour).')
    ans_bullet(doc, '3. ABC Analysis: ',
               'Pareto-based classification of SKUs by sales '
               'value to focus tight controls on the most '
               'important items.')
    ans_bullet(doc, '4. Reorder-Point (ROP) System: ',
               'Triggers a new order when stock falls to a '
               'predefined level (lead-time × demand + safety '
               'stock).')
    ans_bullet(doc, '5. Periodic Review System: ',
               'Stock is reviewed at fixed intervals (weekly, '
               'monthly) and topped up to a target level.')
    ans_bullet(doc, '6. Continuous / Perpetual Inventory: ',
               'Real-time tracking via POS/barcode/RFID; gives '
               'live stock visibility.')
    ans_bullet(doc, '7. First-In-First-Out (FIFO) and First-Expired-'
                    'First-Out (FEFO): ',
               'Critical for perishables (dairy, fresh produce, '
               'pharmaceuticals) to minimise waste.')
    ans_bullet(doc, '8. Vendor-Managed Inventory (VMI): ',
               'The supplier monitors and replenishes stock at '
               'the retailer\'s shelf — common with Coca-Cola, '
               'EABL, Bidco at Kenyan supermarkets.')
    ans_bullet(doc, '9. Cycle Counting: ',
               'Continuous spot stock-takes throughout the year '
               'instead of one big annual count — improves '
               'accuracy.')
    ans_bullet(doc, '10. Safety-Stock & Buffer Inventory: ',
               'Extra stock held to cushion against demand '
               'spikes or supply delays.')
    ans_bullet(doc, '11. Drop-shipping: ',
               'Retailer accepts the order; supplier ships '
               'directly to the customer — zero retailer '
               'inventory.')
    ans_bullet(doc, '12. RFID & barcode automation: ',
               'Technology-enabled accuracy and shrinkage '
               'control.')

    # Q4
    question(doc, 'FOUR', None, marks='15 marks')
    sub_q(doc, '(a)', 'What cultural considerations should retailers '
          'in Kenya keep in mind when designing signage and store '
          'displays?', '6 marks')
    ans_label(doc)
    ans_bullet(doc, '1. Language and literacy: ',
               'Use both English and Kiswahili (and sometimes '
               'local languages — Kikuyu, Luo, Luhya, Kalenjin, '
               'Somali) in signage to be inclusive.')
    ans_bullet(doc, '2. Religious sensitivity: ',
               'Respect Muslim, Christian and Hindu sensibilities '
               '— halal labelling for the Eastleigh and Coast '
               'markets, modest mannequin styling, separate '
               'displays for alcohol where appropriate.')
    ans_bullet(doc, '3. Colour symbolism: ',
               'Black is used in mourning in some communities; '
               'red is celebratory; green is associated with '
               'agriculture and Islam. Avoid colour clashes that '
               'send unintended signals.')
    ans_bullet(doc, '4. Imagery and representation: ',
               'Use diverse, locally-relatable models (skin '
               'tones, age, gender, ability). Avoid imported '
               'imagery that does not resonate.')
    ans_bullet(doc, '5. Festivals and seasonality: ',
               'Plan for Christmas, Easter, Madaraka, Jamhuri, '
               'Eid, Diwali, Ramadan, Back-to-School, Valentines '
               '— with appropriate themes.')
    ans_bullet(doc, '6. Modesty and decency: ',
               'In family-oriented Kenyan culture, very revealing '
               'imagery may offend; aim for tasteful, family-'
               'friendly visuals.')
    ans_bullet(doc, '7. Local pride: ',
               '"Made in Kenya", "Buy Kenya, Build Kenya" and '
               '"Pesa ya Wakenya" messaging resonates strongly.')
    ans_bullet(doc, '8. Currency display: ',
               'Always display prices in KES, with VAT-inclusive '
               'figures, in clear, large fonts.')

    sub_q(doc, '(b)', 'What factors should a retailer consider when '
          'selecting a local supplier in Kenya?', '5 marks')
    ans_label(doc)
    ans_bullet(doc, '1. Quality & consistency: ',
               'Ability to meet KEBS standards and the retailer\'s '
               'specifications consistently.')
    ans_bullet(doc, '2. Price & cost competitiveness: ',
               'Total landed cost (price + transport + duties + '
               'losses) compared to alternatives.')
    ans_bullet(doc, '3. Reliability & lead time: ',
               'On-time-in-full (OTIF) record; capacity to scale '
               'with demand spikes (festive season).')
    ans_bullet(doc, '4. Capacity & financial stability: ',
               'Audited accounts; ability to honour large '
               'orders without collapse.')
    ans_bullet(doc, '5. Compliance & ethics: ',
               'Tax compliance (KRA PIN/iTax), labour standards, '
               'environmental compliance (NEMA), no child '
               'labour.')
    ans_bullet(doc, '6. Geographic location & logistics: ',
               'Proximity to retailer\'s DC reduces transport '
               'cost and lead time.')
    ans_bullet(doc, '7. Payment terms & flexibility: ',
               'Credit days, M-PESA acceptance, ability to '
               'invoice on consignment.')
    ans_bullet(doc, '8. Innovation & product range: ',
               'Ability to introduce new SKUs and respond to '
               'market trends.')
    ans_bullet(doc, '9. Reputation & references: ',
               'Track record with other retailers, consumer '
               'reviews.')
    ans_bullet(doc, '10. After-sales & service support: ',
               'Returns, warranty handling, marketing support, '
               'merchandising assistance.')

    sub_q(doc, '(c)', 'Why is it important for retailers to monitor '
          'technological trends?', '4 marks')
    ans_label(doc)
    ans_bullet(doc, '1. Stay competitive: ',
               'Failure to adopt e-commerce, mobile payments or '
               'modern POS leads to loss of market share — '
               'Nakumatt being the cautionary tale.')
    ans_bullet(doc, '2. Meet evolving customer expectations: ',
               'Customers expect mobile checkout, free Wi-Fi, '
               'self-checkout, BOPIS and personalised offers.')
    ans_bullet(doc, '3. Improve operational efficiency: ',
               'Cloud POS, automated inventory, AI demand-'
               'forecasting and RFID cut costs and shrinkage.')
    ans_bullet(doc, '4. Better customer insights: ',
               'Loyalty programs, CDPs and analytics turn data '
               'into personalised marketing and merchandising '
               'decisions.')
    ans_bullet(doc, '5. Spot disruptive threats early: ',
               'Live shopping, dark stores and 15-minute '
               'delivery can erode share quickly if missed.')
    ans_bullet(doc, '6. Enhance security & loss prevention: ',
               'Biometrics, CCTV-AI analytics, fraud-detection '
               'in payments protect margin.')
    ans_bullet(doc, '7. Compliance: ',
               'New laws (Data Protection Act 2019, e-invoicing '
               'TIMS / eTIMS by KRA) require technological '
               'adaptation.')

    # Q5
    question(doc, 'FIVE', None, marks='15 marks')
    sub_q(doc, '(a)', 'Should Kenyan retailers invest in custom '
          'inventory software or use off-the-shelf solutions? Why?',
          '6 marks')
    ans_label(doc)
    ans(doc,
        'The choice depends on the size, complexity and '
        'strategic ambitions of the retailer.')
    subheading(doc, 'Off-the-Shelf Solutions (Recommended for SMEs)')
    ans_bullet(doc, 'Examples: ',
               'QuickBooks POS, Vend (Lightspeed), Loyverse, '
               'Zoho Inventory, Odoo, Sage, ERPNext.')
    ans_bullet(doc, 'Advantages: ',
               'Lower upfront cost, faster deployment '
               '(days/weeks), proven reliability, vendor '
               'support, regular updates, good integration '
               'with M-PESA, KRA eTIMS, Excel and accounting '
               'systems.')
    ans_bullet(doc, 'Disadvantages: ',
               'Limited customisation, you adapt your '
               'processes to the software, recurring '
               'subscription cost.')
    subheading(doc, 'Custom Software (Justified for Large or '
                    'Complex Chains)')
    ans_bullet(doc, 'Advantages: ',
               'Tailored to unique processes (e.g. Naivas\' '
               'multi-format network, Carrefour\'s import '
               'workflows), competitive differentiation, '
               'tighter integration with proprietary loyalty/'
               'CRM, full ownership of the IP and data.')
    ans_bullet(doc, 'Disadvantages: ',
               'Very high cost (millions of shillings), long '
               'development time (12-24 months), risk of '
               'over-engineering, ongoing maintenance and '
               'developer dependency.')
    ans(doc,
        'Recommendation: most Kenyan retailers — especially '
        'SMEs and growing mid-market chains — should start '
        'with a proven off-the-shelf solution and graduate '
        'to a hybrid (off-the-shelf core + custom modules '
        'via APIs) once volume and process complexity '
        'justify the investment. Only the very largest '
        'chains (Naivas, Quickmart, Carrefour) can justify '
        'fully custom systems.')

    sub_q(doc, '(c)', 'Imagine you are a buyer for a mid-sized '
          'retailer in Kenya — what criteria would you use to select '
          'suppliers for fresh vegetables?', '5 marks')
    ans_label(doc)
    ans_bullet(doc, '1. Freshness and quality consistency: ',
               'Daily-harvest produce, properly graded; meets '
               'KEBS and HACCP standards.')
    ans_bullet(doc, '2. Reliability and on-time delivery: ',
               'Daily or twice-daily deliveries — fresh produce '
               'has a 24-72 hour window.')
    ans_bullet(doc, '3. Cold-chain capability: ',
               'Refrigerated transport from Mwea/Limuru/Naivasha '
               'farms to the store; reduces wastage.')
    ans_bullet(doc, '4. Price competitiveness: ',
               'Wholesale price within market range; volume '
               'discounts available.')
    ans_bullet(doc, '5. Geographic proximity: ',
               'Close to Nairobi (Limuru, Kiambu, Kajiado, '
               'Mai Mahiu) reduces transport cost and CO\u2082 '
               'footprint.')
    ans_bullet(doc, '6. Capacity & seasonal flexibility: ',
               'Can scale up for festive periods; aggregator '
               'model (e.g. Twiga Foods) helps smooth supply.')
    ans_bullet(doc, '7. Compliance: ',
               'Ministry of Agriculture certification, food-'
               'safety records, KEBS labelling for packaged '
               'lines.')
    ans_bullet(doc, '8. Sustainability: ',
               'Use of safe pesticides, water-efficient '
               'irrigation, fair-trade for smallholder '
               'farmers.')
    ans_bullet(doc, '9. Payment & credit terms: ',
               'M-PESA acceptance, reasonable credit days '
               '(7-14 days for fresh produce).')
    ans_bullet(doc, '10. Returns / wastage policy: ',
               'Willingness to take back damaged or unsold '
               'stock at agreed terms.')

    sub_q(doc, '(d)', 'What are some challenges Kenyan retailers '
          'face when forecasting demand during inflation or '
          'political uncertainty?', '4 marks')
    ans_label(doc)
    ans_bullet(doc, '1. Volatile consumer spending: ',
               'Inflation erodes disposable income; consumers '
               'down-trade and switch brands unpredictably.')
    ans_bullet(doc, '2. Currency volatility: ',
               'KES depreciation against USD increases the '
               'cost of imported goods overnight, distorting '
               'price-elasticity assumptions.')
    ans_bullet(doc, '3. Supply-chain disruption: ',
               'Fuel shortages, fuel-price hikes, port '
               'congestion (Mombasa) and political demonstrations '
               '(e.g. 2024 Gen-Z protests) interrupt deliveries.')
    ans_bullet(doc, '4. Panic buying & hoarding: ',
               'Around elections or fuel-price announcements, '
               'demand spikes for staples (sugar, cooking oil, '
               'maize flour) — historical data fails to '
               'predict.')
    ans_bullet(doc, '5. Frequent price changes: ',
               'Suppliers raise prices weekly; planogram '
               'pricing and promotions become hard to lock '
               'down.')
    ans_bullet(doc, '6. Reduced data reliability: ',
               'Historical sales patterns are no longer '
               'predictive; AI/ML models need frequent re-'
               'training.')
    ans_bullet(doc, '7. Working-capital strain: ',
               'Higher inflation increases the cost of carrying '
               'inventory; over-forecasting locks up cash.')

    # Q6
    question(doc, 'SIX', None, marks='15 marks')
    sub_q(doc, '(a)', 'In what ways can retailers influence the '
          'purchasing decisions of consumers?', '3 marks')
    ans_label(doc)
    ans_bullet(doc, 'Visual merchandising & store layout: ',
               'Placement of high-margin items at eye-level, '
               'end-cap displays and check-out impulse zones.')
    ans_bullet(doc, 'Pricing tactics: ',
               'Charm pricing (Ksh 999 vs 1,000), bundle '
               'discounts, BOGOF, anchor pricing.')
    ans_bullet(doc, 'Promotions and advertising: ',
               'Sales, flash deals, coupons, seasonal '
               'campaigns.')
    ans_bullet(doc, 'In-store sampling and demonstrations: ',
               'Tasting stations, beauty consultants, tech '
               'demos.')
    ans_bullet(doc, 'Personal selling and recommendations: ',
               'Trained staff who up-sell and cross-sell.')
    ans_bullet(doc, 'Loyalty programs and rewards: ',
               'Points, tiers, birthday offers (Naivas Linda).')
    ans_bullet(doc, 'Sensory marketing: ',
               'Music, scent, lighting and colour to influence '
               'mood and dwell time.')
    ans_bullet(doc, 'Digital personalisation: ',
               'AI-driven recommendations, retargeting, '
               'personalised emails.')
    ans_bullet(doc, 'Social proof: ',
               'Reviews, testimonials, user-generated content, '
               'influencer endorsements.')

    sub_q(doc, '(b)', 'Expound on the Retailing Mix — why does it '
          'matter to a retail business\' operations?', '6 marks')
    ans_label(doc)
    ans(doc,
        'The Retailing Mix (also called the "6 Ps" of retail) '
        'is the combination of controllable variables a '
        'retailer uses to satisfy its target market and '
        'achieve its objectives. The classic elements are:')
    ans_bullet(doc, '1. Product (Merchandise assortment): ',
               'Width, depth, quality, brand mix and '
               'private-label strategy.')
    ans_bullet(doc, '2. Price: ',
               'Pricing strategy (EDLP, High-Low, prestige), '
               'discounts, loyalty pricing, dynamic pricing.')
    ans_bullet(doc, '3. Place (Location & channels): ',
               'Store location, format, omnichannel reach '
               '(physical + digital + delivery).')
    ans_bullet(doc, '4. Promotion: ',
               'Advertising, sales promotion, PR, digital '
               'marketing, in-store communication.')
    ans_bullet(doc, '5. People (Staff and service): ',
               'Recruitment, training, motivation, customer-'
               'service standards.')
    ans_bullet(doc, '6. Presentation (Store atmospherics & '
                    'visual merchandising): ',
               'Layout, fixtures, lighting, signage, scent, '
               'music — the in-store experience.')
    subheading(doc, 'Why It Matters')
    ans_bullet(doc, 'Strategic alignment: ',
               'Each element must reinforce the brand position '
               '(e.g. Vivo Activewear\'s premium pricing must '
               'be matched by premium presentation and '
               'service).')
    ans_bullet(doc, 'Differentiation: ',
               'A unique mix is hard for competitors to '
               'replicate, building defensible advantage.')
    ans_bullet(doc, 'Customer satisfaction & loyalty: ',
               'A coherent mix delivers on the promise the '
               'customer expects.')
    ans_bullet(doc, 'Profitability: ',
               'Right product at right price, in right place, '
               'with right people drives sales and margin.')
    ans_bullet(doc, 'Adaptability: ',
               'The mix can be adjusted as the market shifts '
               '— e.g. adding e-commerce post-COVID, expanding '
               'private-label during inflation.')

    sub_q(doc, '(c)', 'Imagine you are setting up a loyalty program '
          'for a new grocery store in Kisumu. What features would '
          'you include to make it effective?', '6 marks')
    ans_label(doc)
    ans(doc,
        'A successful Kisumu grocery loyalty program must be '
        'simple, mobile-first and locally relevant. Core '
        'features:')
    ans_bullet(doc, '1. Easy sign-up via phone number / M-PESA: ',
               'No need for plastic cards or e-mail; one-time '
               'verification.')
    ans_bullet(doc, '2. Earn-and-burn points system: ',
               'E.g. 1 point per Ksh 100 spent; 100 points = '
               'Ksh 50 off next purchase.')
    ans_bullet(doc, '3. Tiered membership: ',
               'Bronze, Silver, Gold based on annual spend, '
               'with escalating benefits.')
    ans_bullet(doc, '4. Personalised offers: ',
               'Targeted discounts on items members buy '
               'often (e.g. Omena, Sukuma Wiki, Ugali flour).')
    ans_bullet(doc, '5. Birthday and anniversary rewards: ',
               'Free gift or double points around the '
               'member\'s special day.')
    ans_bullet(doc, '6. M-PESA integration: ',
               'Pay and earn in one STK push; cash-back '
               'directly to M-PESA.')
    ans_bullet(doc, '7. Family / household sharing: ',
               'Allow points pooling among related phone '
               'numbers — popular in Kenyan family shopping.')
    ans_bullet(doc, '8. Seasonal & community campaigns: ',
               'Bonus points during back-to-school, '
               'Christmas, Sondu farmers\' market season.')
    ans_bullet(doc, '9. Referral rewards: ',
               'Refer a friend → both get bonus points.')
    ans_bullet(doc, '10. WhatsApp and SMS communication: ',
               'High open-rate channels in Kenya; in '
               'Kiswahili and English.')
    ans_bullet(doc, '11. Easy redemption at till and online: ',
               'No friction — cashier scans the phone number; '
               'app shows balance.')
    ans_bullet(doc, '12. Data-driven personalisation: ',
               'Use loyalty data to refine assortment, '
               'pricing and marketing for the Kisumu '
               'shopper.')
    ans_bullet(doc, '13. Charitable giving option: ',
               'Allow members to donate points to a local '
               'cause (e.g. school feeding, Lake Victoria '
               'clean-up) — builds emotional brand bond.')

    # Closing
    page_break(doc)
    cover_line(doc, 'EXAM TECHNIQUE — FINAL POINTERS', size=14,
               bold=True, after=10)
    body(doc,
         'When tackling the actual paper, allocate time strictly '
         'in proportion to marks (≈1.8 minutes per mark in a 3-'
         'hour, 100-mark paper). Always: (1) underline the '
         'command word — define, discuss, explain, evaluate, '
         'illustrate, compare; (2) structure with sub-headings '
         'and bullets; (3) anchor with a Kenyan example for '
         'every major concept; (4) end "discuss" answers with a '
         'short critical comment or limitation; (5) keep your '
         'handwriting legible and your answer-booklet tidy. '
         'These habits add up to 5-10 marks across a paper.')
    body(doc,
         'Good luck Mourice — revise actively, write structured '
         'answers and walk into the exam confident.', before=8)

    doc.save(output)
    print(f'DOCX saved: {output}')
    convert_to_pdf(output)
    try:
        _os.remove(output)
        print(f'Removed intermediate DOCX: {output}')
    except Exception:
        pass


def main():
    build()


if __name__ == '__main__':
    main()
