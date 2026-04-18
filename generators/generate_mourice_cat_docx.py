#!/usr/bin/env python3
"""
Generate formatted DOCX for Mourice Onyango BBM 453 CAT
Distributed Systems - Critical Analysis
"""

import os as _os, sys as _sys
_sys.path.insert(0, _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))
_os.chdir(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_run(run, bold=False, italic=False, size=12, color=(0, 0, 0)):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(*color)


def set_spacing(para, before=0, after=6, spacing=1.5):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if spacing == 1.5:
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    elif spacing == 1.0:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    else:
        pf.line_spacing = Pt(spacing * 12)


def cover_line(doc, text, size=12, bold=False, after=6,
               align=WD_ALIGN_PARAGRAPH.CENTER):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run(run, bold=bold, size=size)
    para.alignment = align
    set_spacing(para, before=0, after=after, spacing=1.0)
    return para


def body(doc, text, bold=False, italic=False, before=0, after=6,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run(run, bold=bold, italic=italic, size=12)
    para.alignment = align
    set_spacing(para, before=before, after=after, spacing=1.5)
    if indent:
        para.paragraph_format.left_indent = Inches(indent)
    return para


def heading(doc, text, before=14, after=6):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run(run, bold=True, size=12)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(para, before=before, after=after, spacing=1.5)
    return para


def subheading(doc, text, before=10, after=4):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run(run, bold=True, italic=True, size=12)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(para, before=before, after=after, spacing=1.5)
    return para


def bullet(doc, label, text, before=0, after=4):
    para = doc.add_paragraph()
    r1 = para.add_run(label)
    set_run(r1, bold=True, size=12)
    r2 = para.add_run(text)
    set_run(r2, size=12)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(para, before=before, after=after, spacing=1.5)
    para.paragraph_format.left_indent = Inches(0.35)
    return para


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def generate(output='files/Mourice_BBM_453_CAT.docx'):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.0)

    # ─── COVER PAGE ───────────────────────────────────────────────
    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(logo, before=6, after=10, spacing=1.0)
    try:
        logo.add_run().add_picture(
            'assets/moi_uni_logo.png', width=Inches(1.4))
    except Exception:
        pass

    cover_line(doc, 'MOI UNIVERSITY', size=16, bold=True, after=4)
    cover_line(doc, 'ANNEX CAMPUS',   size=14, bold=True, after=4)
    cover_line(doc, 'SCHOOL OF BUSINESS & ECONOMICS', size=13, bold=True, after=4)
    cover_line(doc, 'DEPARTMENT OF MANAGEMENT SCIENCE & ENTREPRENEURSHIP',
               size=12, bold=True, after=18)

    details = [
        ('PROGRAMME',        'BACHELOR OF BUSINESS MANAGEMENT'),
        ('ACADEMIC YEAR',    'YEAR 4'),
        ('COURSE CODE',      'BBM 453'),
        ('COURSE TITLE',     'DISTRIBUTED SYSTEMS'),
        ('SEMESTER',         '2025/26: SEM II'),
        ('ASSIGNMENT',       'RESEARCH (CAT)'),
        ('SUBMISSION DATE',  '10TH APRIL 2026'),
        ('MARKS',            '20 MARKS'),
        ('NAME',             'MOURICE ONYANGO'),
        ('REG NUMBER',       'BBM/1891/22'),
    ]
    for label, value in details:
        para = doc.add_paragraph()
        r1 = para.add_run(f'{label:<18}: ')
        set_run(r1, bold=True, size=12)
        r2 = para.add_run(value)
        set_run(r2, size=12)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(para, before=0, after=7, spacing=1.0)

    page_break(doc)

    # ─── INTRODUCTION ─────────────────────────────────────────────
    heading(doc, '1. INTRODUCTION', before=0)
    for t in [
        'Distributed systems represent one of the most transformative paradigms in modern '
        'computing. A distributed system is a collection of autonomous computing elements '
        'that appear to users as a single, coherent system (Tanenbaum & Van Steen, 2017). '
        'Unlike centralized systems, distributed systems spread computation, data storage, '
        'and processing across multiple nodes connected via a network, enabling fault '
        'tolerance, scalability, and geographic flexibility.',
        'The current technological landscape is characterised by a convergence of several '
        'distributed computing trends: cloud computing, edge and fog computing, blockchain '
        'technology, microservices architecture, the Internet of Things (IoT), and real-time '
        'data streaming platforms. Together, these technologies are reshaping how institutions '
        'and enterprises deliver services, manage data, and respond to operational demands.',
        'In Kenya, the adoption of distributed systems has accelerated markedly over the past '
        'decade, driven by the growth of mobile internet penetration, government digitalization '
        'initiatives such as the Digital Superhighway program, and the emergence of a vibrant '
        'technology ecosystem centred around Nairobi\'s Silicon Savannah. The COVID-19 '
        'pandemic further accelerated digital transformation across all sectors, making '
        'distributed systems a critical enabler of continuity in healthcare, education, '
        'finance, security, and commerce.',
        'This paper critically analyses the current trends of distributed systems and '
        'evaluates their impact on five key sectors in Kenya: healthcare, finance and banking, '
        'education and academia, national security, and small and medium enterprises (SMEs). '
        'For each sector, a practical Kenyan example is provided to illustrate the tangible '
        'effects of distributed systems on service delivery and performance.',
    ]:
        body(doc, t)

    page_break(doc)

    # ─── SECTION A: HEALTHCARE ────────────────────────────────────
    heading(doc, 'a) DISTRIBUTED SYSTEMS IN HEALTHCARE', before=0)

    subheading(doc, 'Current Trends')
    for t in [
        'The healthcare sector globally is undergoing a digital revolution anchored on '
        'distributed systems. Key trends include cloud-based Electronic Health Records (EHR) '
        'systems, telemedicine platforms, IoT-connected medical devices, and distributed '
        'data analytics for epidemiological surveillance. Edge computing is increasingly '
        'deployed to process medical sensor data in real time at the point of care, reducing '
        'latency and improving clinical decision-making.',
        'Interoperability between disparate health information systems — achieved through '
        'distributed middleware and Application Programming Interfaces (APIs) — is enabling '
        'the seamless exchange of patient information across hospitals, clinics, laboratories, '
        'and pharmacies. This interoperability is particularly significant in low-resource '
        'settings where fragmented data has historically hampered public health management.',
    ]:
        body(doc, t)

    subheading(doc, 'Impact in Kenya')
    for t in [
        'Kenya\'s Ministry of Health adopted the District Health Information System 2 (DHIS2), '
        'a distributed, open-source health information platform, as the national health '
        'management information system. DHIS2 aggregates health data from over 10,000 '
        'health facilities across the country, enabling the Ministry to track disease '
        'surveillance indicators, immunization coverage, maternal health outcomes, and '
        'resource utilization in near real time.',
        'Additionally, M-TIBA, a mobile-based healthcare financing platform developed by '
        'CarePay International and Safaricom, leverages distributed cloud architecture to '
        'connect patients, healthcare providers, and insurance funds. The platform enables '
        'Kenyans — including low-income households — to save for healthcare, receive health '
        'funds, and pay for medical services directly from their mobile phones.',
    ]:
        body(doc, t)

    subheading(doc, 'Practical Example: DHIS2 and M-TIBA')
    for t in [
        'During the COVID-19 pandemic, Kenya\'s Ministry of Health utilized DHIS2 as a '
        'real-time distributed data platform for tracking infections, hospitalisations, and '
        'vaccine administration. County health departments uploaded daily case data from '
        'tablets and smartphones across all 47 counties. The distributed architecture '
        'ensured that even in counties with intermittent internet connectivity, data could '
        'be entered offline and synchronized when connectivity was restored — a critical '
        'feature for remote areas.',
        'M-TIBA similarly demonstrated the power of distributed systems during the pandemic '
        'by facilitating cashless payment for COVID-19 testing and treatment services, '
        'reducing physical contact at health facilities. By March 2021, M-TIBA had processed '
        'over 4 million health fund transactions and onboarded more than 1,500 healthcare '
        'providers across Kenya.',
    ]:
        body(doc, t)

    subheading(doc, 'Impact on Service Delivery and Performance')
    for label, text in [
        ('Improved Data Quality:  ',
         'Real-time, distributed data collection reduced reporting delays from weeks to '
         'hours, enabling faster policy responses to disease outbreaks.'),
        ('Enhanced Access:  ',
         'Telemedicine platforms powered by distributed cloud infrastructure extended '
         'specialist consultations to rural areas previously underserved by qualified doctors.'),
        ('Cost Reduction:  ',
         'M-TIBA\'s digital payment model reduced administrative costs for healthcare '
         'providers by eliminating paper-based billing and cash handling.'),
        ('Accountability:  ',
         'Distributed audit trails in DHIS2 improved accountability in health commodity '
         'supply chains, reducing stock-outs of essential medicines at facility level.'),
    ]:
        bullet(doc, label, text)

    page_break(doc)

    # ─── SECTION B: FINANCE & BANKING ────────────────────────────
    heading(doc, 'b) DISTRIBUTED SYSTEMS IN FINANCE AND BANKING', before=0)

    subheading(doc, 'Current Trends')
    for t in [
        'The financial sector has been among the earliest and most enthusiastic adopters of '
        'distributed systems technologies. Current trends include microservices-based core '
        'banking architectures, real-time payment processing networks, blockchain and '
        'distributed ledger technology (DLT) for transaction transparency, cloud-native '
        'banking platforms, and AI-driven distributed fraud detection systems.',
        'Open Banking — enabled by distributed APIs — is allowing third-party developers '
        'to build financial applications on top of bank infrastructure, democratizing access '
        'to financial services. Central Bank Digital Currencies (CBDCs), being explored by '
        'several African central banks, are also premised on distributed ledger architectures '
        'that ensure transparency and immutability of monetary transactions.',
    ]:
        body(doc, t)

    subheading(doc, 'Impact in Kenya')
    for t in [
        'Kenya\'s financial sector is globally recognised for its pioneering use of '
        'distributed mobile money systems. M-Pesa, launched by Safaricom in 2007 and '
        'built on a distributed transaction processing architecture, revolutionized '
        'financial inclusion in Kenya. The platform processes millions of transactions '
        'daily through a distributed network of agents, servers, and mobile nodes '
        'spread across the country.',
        'Equity Bank\'s Equitel platform and the PesaLink interbank payment system, '
        'operated by the Kenya Bankers Association (KBA), further exemplify the power '
        'of distributed systems in enabling real-time, interoperable financial transactions '
        'among competing financial institutions.',
    ]:
        body(doc, t)

    subheading(doc, 'Practical Example: M-Pesa Distributed Architecture')
    for t in [
        'M-Pesa\'s technical infrastructure is a distributed system comprising Safaricom\'s '
        'core transaction servers, a vast network of distributed agent terminals (over '
        '250,000 agents nationwide), mobile subscriber endpoints, and integration gateways '
        'to banks, utilities, and government services. The system uses geographically '
        'redundant data centres to ensure 99.99% uptime, with automatic failover mechanisms '
        'that switch to backup nodes in the event of a primary node failure.',
        'The M-Pesa Global platform extended this distributed architecture across borders, '
        'enabling diaspora remittances from 14 countries to reach recipients in Kenya within '
        'seconds. PesaLink, operating on a distributed real-time gross settlement (RTGS) '
        'infrastructure, enables instant bank-to-bank transfers 24/7, including on weekends '
        'and public holidays — a capability unavailable under earlier centralized batch '
        'processing systems.',
    ]:
        body(doc, t)

    subheading(doc, 'Impact on Service Delivery and Performance')
    for label, text in [
        ('Financial Inclusion:  ',
         'M-Pesa has brought over 30 million Kenyans into the formal financial system, '
         'including unbanked populations in rural and peri-urban areas, dramatically '
         'improving their access to savings, credit, and insurance products.'),
        ('Transaction Speed:  ',
         'Real-time distributed payment systems reduced transaction settlement time from '
         'T+3 business days under traditional banking to near-instantaneous processing.'),
        ('Resilience:  ',
         'Geographically distributed data centres ensure continuity of financial services '
         'even during localized infrastructure failures or cyber attacks.'),
        ('Fraud Detection:  ',
         'Distributed machine learning models that analyse transaction patterns across '
         'millions of nodes in real time have significantly reduced fraudulent transactions '
         'in mobile banking, protecting both financial institutions and consumers.'),
    ]:
        bullet(doc, label, text)

    page_break(doc)

    # ─── SECTION C: EDUCATION ─────────────────────────────────────
    heading(doc, 'c) DISTRIBUTED SYSTEMS IN EDUCATION AND ACADEMIA', before=0)

    subheading(doc, 'Current Trends')
    for t in [
        'Education is experiencing a profound transformation through distributed learning '
        'management systems (LMS), cloud-hosted academic repositories, virtual classrooms, '
        'distributed video conferencing infrastructure, and peer-to-peer (P2P) collaborative '
        'platforms. Massive Open Online Courses (MOOCs), delivered via globally distributed '
        'content delivery networks (CDNs), have made quality education accessible to learners '
        'in remote and underserved regions.',
        'Edge computing is beginning to be deployed in educational institutions to support '
        'bandwidth-intensive applications such as virtual reality (VR) labs and real-time '
        'collaborative coding environments. Blockchain-based academic credential systems '
        'are emerging as a solution to certificate fraud by enabling verifiable, tamper-proof '
        'digital certificates on distributed ledgers.',
    ]:
        body(doc, t)

    subheading(doc, 'Impact in Kenya')
    for t in [
        'Kenya\'s transition to a competency-based curriculum and the government\'s Digital '
        'Literacy Programme (DLP) have been underpinned by distributed computing '
        'infrastructure. The Kenya Education Cloud (edCloud), managed by the Kenya '
        'Education Network (KENET), provides a distributed cloud hosting platform for '
        'universities and research institutions, supporting e-learning portals, digital '
        'libraries, and research data repositories.',
        'Moi University\'s online learning portal and the Jomo Kenyatta University of '
        'Agriculture and Technology (JKUAT) virtual campus are examples of distributed '
        'academic platforms that enabled continuity of learning during the COVID-19 school '
        'closures of 2020-2021, when all physical campuses were shut down by government '
        'directive.',
    ]:
        body(doc, t)

    subheading(doc, 'Practical Example: KENET and BBM Annex')
    for t in [
        'KENET operates a distributed research and education network connecting over 60 '
        'institutions of higher learning across Kenya through a high-speed fibre backbone '
        'and peering arrangements with regional and international networks. This distributed '
        'infrastructure supports video conferencing, remote access to academic journals via '
        'the Kenya Universities and Colleges Central Placement Service (KUCCPS) portal, and '
        'high-performance computing resources for scientific research.',
        'At Moi University\'s Annex Campus, the student-developed BBM Annex platform '
        '(https://bbm.giftedtech.co.ke) exemplifies a grassroots distributed academic '
        'resource-sharing system. The platform aggregates lecture notes, past examination '
        'papers, and study guides uploaded by students across different academic years, '
        'creating a distributed knowledge base accessible to all BBM students via mobile '
        'devices. Research by Onyango (2026) found that 71.8% of BBM students reported '
        'improved examination preparedness after accessing resources through the platform.',
    ]:
        body(doc, t)

    subheading(doc, 'Impact on Service Delivery and Performance')
    for label, text in [
        ('Continuity of Learning:  ',
         'Distributed e-learning platforms ensured that over 1 million university students '
         'in Kenya continued their studies during the COVID-19 pandemic, preventing an entire '
         'academic year from being lost.'),
        ('Resource Accessibility:  ',
         'KENET\'s distributed network gave students in remote campuses access to the same '
         'academic databases and digital libraries as students in Nairobi, reducing '
         'information inequality between urban and rural institutions.'),
        ('Research Collaboration:  ',
         'Distributed cloud repositories enabled Kenyan researchers to collaborate with '
         'international counterparts in real time, contributing to a measurable increase '
         'in research output and co-authored publications.'),
        ('Reduced Costs:  ',
         'Cloud-hosted LMS platforms eliminated the need for expensive on-premise server '
         'infrastructure in individual institutions, significantly reducing the cost of '
         'deploying and maintaining e-learning environments.'),
    ]:
        bullet(doc, label, text)

    page_break(doc)

    # ─── SECTION D: NATIONAL SECURITY ────────────────────────────
    heading(doc, 'd) DISTRIBUTED SYSTEMS IN NATIONAL SECURITY', before=0)

    subheading(doc, 'Current Trends')
    for t in [
        'Modern national security architectures are increasingly premised on distributed '
        'computing paradigms. Key trends include distributed surveillance networks '
        'integrating CCTV cameras, drones, and biometric terminals; distributed intelligence '
        'sharing platforms connecting multiple security agencies; cloud-based command and '
        'control systems; and distributed cybersecurity infrastructure for real-time threat '
        'detection and response.',
        'Blockchain technology is being explored for securing sensitive government records '
        'and identity documents, as its distributed and immutable ledger prevents '
        'unauthorized modification of data. The Zero Trust security model — which assumes '
        'no inherently trusted nodes in a distributed network — is gaining traction in '
        'government cybersecurity frameworks, requiring continuous verification of every '
        'user and device.',
    ]:
        body(doc, t)

    subheading(doc, 'Impact in Kenya')
    for t in [
        'The Government of Kenya has invested significantly in distributed security '
        'infrastructure through several flagship programs. The National Integrated Identity '
        'Management System (NIIMS), known commercially as Huduma Namba, is a distributed '
        'biometric identification system that consolidates citizen identity data from '
        'multiple government registries — including the National Registration Bureau, '
        'Kenya Revenue Authority, and Registrar of Persons — into a single interoperable '
        'platform.',
        'The Nairobi City Surveillance Project, implemented by the National Police Service '
        'in partnership with Huawei, deployed over 1,800 networked CCTV cameras across '
        'the city, connected to a central command center via a distributed fiber and '
        'wireless network. This system has been expanded to other major cities including '
        'Mombasa, Kisumu, and Eldoret as part of Kenya\'s Safe City initiative.',
    ]:
        body(doc, t)

    subheading(doc, 'Practical Example: Huduma Namba and the Safe City CCTV Network')
    for t in [
        'The Huduma Namba system exemplifies a large-scale distributed identity management '
        'deployment. Data is collected at distributed registration centers across all '
        '47 counties using biometric enrollment devices (fingerprint scanners, iris cameras, '
        'and facial recognition terminals) and synchronized with a central cloud-hosted '
        'database managed by the State Department for Immigration. The distributed '
        'architecture allows for real-time identity verification at border checkpoints, '
        'airport immigration counters, and government service kiosks nationwide.',
        'The Nairobi Safe City CCTV network operates as a distributed surveillance system '
        'where video feeds from thousands of cameras are processed at edge nodes for '
        'preliminary analysis (e.g., number plate recognition, crowd density monitoring) '
        'before being aggregated at the central command center. This edge processing '
        'reduces bandwidth requirements and enables faster incident response compared to '
        'a purely centralized architecture.',
    ]:
        body(doc, t)

    subheading(doc, 'Impact on Service Delivery and Performance')
    for label, text in [
        ('Crime Reduction:  ',
         'The Nairobi Safe City surveillance network contributed to a reported 46% '
         'reduction in crime in covered areas according to the National Police Service, '
         'demonstrating the deterrence effect of distributed video surveillance.'),
        ('Identity Verification:  ',
         'The Huduma Namba system has streamlined access to over 15 government services, '
         'reducing duplication of identity documents and enabling cross-agency verification '
         'in real time, thereby curbing identity fraud.'),
        ('Inter-Agency Coordination:  ',
         'Distributed intelligence-sharing platforms connecting the National Intelligence '
         'Service, Kenya Police Service, and Kenya Defence Forces have improved '
         'coordination in counter-terrorism operations.'),
        ('Border Security:  ',
         'Distributed biometric verification at border points has strengthened immigration '
         'control, enabling the flagging of wanted persons and stolen documents in real '
         'time against a centrally maintained but distributed watchlist database.'),
    ]:
        bullet(doc, label, text)

    page_break(doc)

    # ─── SECTION E: SMEs ──────────────────────────────────────────
    heading(doc, 'e) DISTRIBUTED SYSTEMS IN SMALL AND MEDIUM ENTERPRISE (SME) BUSINESS',
            before=0)

    subheading(doc, 'Current Trends')
    for t in [
        'Small and medium enterprises worldwide are leveraging distributed systems to '
        'compete effectively in digital markets that were previously accessible only to '
        'large corporations with significant IT infrastructure investments. The democratization '
        'of cloud computing has been particularly transformative: SMEs can now access '
        'enterprise-grade distributed computing resources — databases, data analytics '
        'platforms, AI tools, and global content delivery networks — on a pay-as-you-use '
        'basis with minimal upfront capital expenditure.',
        'Current trends driving SME adoption of distributed systems include cloud-based '
        'Enterprise Resource Planning (ERP) systems, mobile Point-of-Sale (POS) terminals '
        'integrated with cloud inventory management, e-commerce platforms operating on '
        'globally distributed infrastructure, distributed supply chain management systems, '
        'and mobile payment APIs that connect SME businesses to millions of mobile money '
        'users instantaneously.',
    ]:
        body(doc, t)

    subheading(doc, 'Impact in Kenya')
    for t in [
        'Kenya\'s SME sector, which accounts for approximately 98% of all businesses and '
        'contributes around 30% of GDP and 80% of employment (KNBS, 2021), has been '
        'significantly transformed by distributed systems. The integration of M-Pesa\'s '
        'Lipa Na M-Pesa Till and Paybill distributed payment APIs into SME operations '
        'has enabled even small roadside kiosks and market stalls to accept cashless '
        'payments and reconcile sales data through cloud-connected terminals.',
        'Platforms such as Jumia Kenya, Copia Global, and Twiga Foods have deployed '
        'distributed supply chain and e-commerce systems that connect small-scale '
        'retailers, farmers, and manufacturers across Kenya\'s diverse geographic landscape, '
        'enabling them to access markets, source inputs, and receive payments through a '
        'unified digital infrastructure.',
    ]:
        body(doc, t)

    subheading(doc, 'Practical Example: Twiga Foods Distributed Supply Chain Platform')
    for t in [
        'Twiga Foods is a Nairobi-based agri-tech company that operates a distributed '
        'supply chain management platform connecting smallholder farmers upcountry with '
        'informal food vendors (mama mbogas) in urban areas. The platform uses a '
        'distributed mobile ordering system where vendors place orders via USSD or a '
        'smartphone app; orders are aggregated by a cloud-based platform, matched to '
        'the nearest fulfillment center, and dispatched via a distributed network of '
        'contracted delivery vehicles tracked in real time through GPS-connected tablets.',
        'Copia Global similarly operates a distributed last-mile e-commerce platform '
        'targeting low-income consumers in peri-urban and rural Kenya. The platform uses '
        'a network of distributed agents (small shop owners who act as order collection '
        'points) connected to a central cloud platform, enabling customers without '
        'smartphones or reliable internet to order products digitally through their '
        'local agent.',
    ]:
        body(doc, t)

    subheading(doc, 'Impact on Service Delivery and Performance')
    for label, text in [
        ('Market Reach:  ',
         'Distributed e-commerce platforms have enabled Kenyan SMEs to reach customers '
         'beyond their immediate geographic areas, with some micro-enterprises reporting '
         'doubling of their customer base within 12 months of joining digital marketplaces.'),
        ('Supply Chain Efficiency:  ',
         'Twiga Foods\' distributed platform reduced post-harvest food losses by up to 30% '
         'by optimising the matching of supply with demand and eliminating inefficient '
         'middlemen in the food value chain.'),
        ('Financial Access:  ',
         'SMEs transacting through distributed mobile payment platforms build a digital '
         'credit history that enables them to access micro-loans from fintech lenders such '
         'as M-Shwari, KCB M-Pesa, and Tala, facilitating business expansion.'),
        ('Operational Efficiency:  ',
         'Cloud-based distributed ERP and POS systems have reduced SME administrative '
         'overhead by automating inventory management, sales reconciliation, and tax '
         'reporting, freeing business owners to focus on growth activities.'),
    ]:
        bullet(doc, label, text)

    page_break(doc)

    # ─── CONCLUSION ───────────────────────────────────────────────
    heading(doc, '2. CONCLUSION', before=0)
    for t in [
        'This paper has critically examined the current trends of distributed systems and '
        'their impact across five key sectors in Kenya. The analysis demonstrates that '
        'distributed computing is not merely a technical advancement but a fundamental '
        'enabler of socio-economic transformation. From DHIS2\'s distributed health '
        'data platform improving pandemic response, to M-Pesa\'s distributed transaction '
        'network revolutionising financial inclusion, to Twiga Foods\' distributed supply '
        'chain reducing food insecurity — distributed systems are at the heart of Kenya\'s '
        'digital development story.',
        'Across all five sectors examined, a consistent pattern emerges: distributed '
        'systems improve service delivery by enhancing scalability, resilience, and '
        'real-time responsiveness; they improve performance by enabling data-driven '
        'decision-making, automating repetitive processes, and connecting previously '
        'isolated stakeholders into collaborative digital ecosystems. The geographic '
        'flexibility of distributed architectures is particularly significant in the '
        'Kenyan context, where infrastructure quality varies enormously between urban '
        'centres and rural areas.',
        'However, the benefits of distributed systems are not without challenges. '
        'Cybersecurity risks increase as the attack surface expands across multiple '
        'distributed nodes. Data privacy concerns arise when personal health, financial, '
        'and identity data is processed across geographically dispersed servers. '
        'Connectivity inequalities risk creating a two-tier digital economy where '
        'the gains of distributed systems are concentrated in well-connected urban '
        'areas while rural communities remain excluded.',
        'To maximise the transformative potential of distributed systems in Kenya, '
        'stakeholders — government, private sector, academia, and civil society — must '
        'collaborate to invest in digital infrastructure, strengthen data governance '
        'frameworks, build local technical capacity, and design inclusive distributed '
        'systems that serve all Kenyans regardless of location or income level.',
    ]:
        body(doc, t)

    page_break(doc)

    # ─── REFERENCES ───────────────────────────────────────────────
    heading(doc, 'REFERENCES', before=0)
    ref_style_para = body  # reuse body function
    refs = [
        'Bird, R. M., & Zolt, E. M. (2008). Technology and taxation in developing countries: '
        'From hand to mouse. National Tax Journal, 61(4), 791-821.',
        'Buyya, R., Yeo, C. S., Venugopal, S., Broberg, J., & Brandic, I. (2009). Cloud '
        'computing and emerging IT platforms: Vision, hype, and reality for delivering '
        'computing as the 5th utility. Future Generation Computer Systems, 25(6), 599-616.',
        'Chetty, K., Qigui, L., Gcora, N., Josie, J., Sheng, L., & Fang, C. (2018). '
        'Bridging the digital divide: Measuring digital literacy. Economics, 12(1), 1-20.',
        'Government of Kenya. (2022). Kenya National Digital Master Plan 2022-2032. '
        'Ministry of Information, Communications and Technology, Innovation and Youth '
        'Affairs. Nairobi: Government Printer.',
        'Hashem, I. A. T., Yaqoob, I., Anuar, N. B., Mokhtar, S., Gani, A., & Khan, S. U. '
        '(2015). The rise of "big data" on cloud computing: Review and open research '
        'issues. Information Systems, 47, 98-115.',
        'Kenya National Bureau of Statistics. (2021). Economic Survey 2021. Nairobi: KNBS.',
        'Lomas, E. (2017). Cyber security: Ensuring confidentiality, integrity and '
        'availability of information. Records Management Journal, 27(3), 264-278.',
        'Maitland, C., & Alamgir, H. (2019). Distributed systems and humanitarian '
        'information management: Implications for design. The Electronic Journal of '
        'Information Systems in Developing Countries, 85(5), e12093.',
        'Onyango, M. (2026). Effectiveness of software development on Moi University '
        'students\' learning behaviour: A case study of BBM Annex. Bachelor of Business '
        'Management Research Project. Moi University.',
        'Safaricom PLC. (2023). Integrated Report and Financial Statements 2023. '
        'Nairobi: Safaricom PLC.',
        'Tanenbaum, A. S., & Van Steen, M. (2017). Distributed systems: Principles '
        'and paradigms (3rd ed.). Prentice Hall.',
        'World Bank. (2022). Kenya Digital Economy Assessment. Washington DC: '
        'World Bank Group.',
        'Zuboff, S. (2019). The age of surveillance capitalism: The fight for a human '
        'future at the new frontier of power. Public Affairs.',
    ]
    for r in refs:
        para = doc.add_paragraph()
        run = para.add_run(r)
        set_run(run, size=12)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_spacing(para, before=0, after=6, spacing=1.5)
        para.paragraph_format.first_line_indent = Inches(-0.35)
        para.paragraph_format.left_indent = Inches(0.35)

    doc.save(output)
    print(f'DOCX saved: {output}')


if __name__ == '__main__':
    generate('files/Mourice_BBM_453_CAT.docx')
