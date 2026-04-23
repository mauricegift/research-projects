#!/usr/bin/env python3
"""
Generate formatted DOCX + PDF for Mourice Onyango BBM 433 CAT 1 and CAT 2.
Course: Retail and Merchandise / Digital Retailing.
Cover-page pattern matches the BBM 453 CAT.
"""

import os as _os
import sys as _sys
import shutil as _shutil
import subprocess

_sys.path.insert(0, _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))
_os.chdir(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))

_SOFFICE = (
    _shutil.which('libreoffice')
    or _shutil.which('soffice')
    or '/nix/store/0pa3zy5lid4paiw9miafpvjkjvlmxfgz-libreoffice-25.2.3.2-wrapped/bin/libreoffice'
)

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK


# ────────────────────────────────────────────────────────────────────────────
# Helpers
# ────────────────────────────────────────────────────────────────────────────
def set_run(run, bold=False, italic=False, size=12, color=(0, 0, 0)):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(*color)


def set_spacing(para, before=0, after=6, spacing=1.5):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if spacing == 1.5:
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    elif spacing == 1.0:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    else:
        pf.line_spacing = Pt(spacing * 12)


def cover_line(doc, text, size=12, bold=False, after=6,
               align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, bold=bold, size=size)
    p.alignment = align
    set_spacing(p, before=0, after=after, spacing=1.0)
    return p


def body(doc, text, bold=False, italic=False, before=0, after=6,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, bold=bold, italic=italic, size=12)
    p.alignment = align
    set_spacing(p, before=before, after=after, spacing=1.5)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    return p


def heading(doc, text, before=14, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, bold=True, size=12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=before, after=after, spacing=1.5)
    return p


def subheading(doc, text, before=10, after=4):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, bold=True, italic=True, size=12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=before, after=after, spacing=1.5)
    return p


def labelled(doc, label, text, before=0, after=4, indent=0.35):
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    set_run(r1, bold=True, size=12)
    r2 = p.add_run(text)
    set_run(r2, size=12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, before=before, after=after, spacing=1.5)
    p.paragraph_format.left_indent = Inches(indent)
    return p


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_cover(doc, course_code, course_title, assignment, sub_date, marks):
    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(logo, before=6, after=10, spacing=1.0)
    try:
        logo.add_run().add_picture('assets/moi_uni_logo.png', width=Inches(1.4))
    except Exception:
        pass

    cover_line(doc, 'MOI UNIVERSITY', size=16, bold=True, after=4)
    cover_line(doc, 'ANNEX CAMPUS', size=14, bold=True, after=4)
    cover_line(doc, 'SCHOOL OF BUSINESS & ECONOMICS', size=13, bold=True, after=4)
    cover_line(doc, 'DEPARTMENT OF MANAGEMENT SCIENCE & ENTREPRENEURSHIP',
               size=12, bold=True, after=18)

    details = [
        ('PROGRAMME',       'BACHELOR OF BUSINESS MANAGEMENT'),
        ('ACADEMIC YEAR',   'YEAR 4'),
        ('COURSE CODE',     course_code),
        ('COURSE TITLE',    course_title),
        ('SEMESTER',        '2025/26: SEM II'),
        ('ASSIGNMENT',      assignment),
        ('SUBMISSION DATE', sub_date),
        ('MARKS',           marks),
        ('NAME',            'MOURICE ONYANGO'),
        ('REG NUMBER',      'BBM/1891/22'),
    ]
    for label, value in details:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{label:<18}: ')
        set_run(r1, bold=True, size=12)
        r2 = p.add_run(value)
        set_run(r2, size=12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p, before=0, after=7, spacing=1.0)

    page_break(doc)


def new_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.0)
    return doc


def convert_to_pdf(docx_path):
    out_dir = _os.path.dirname(docx_path) or '.'
    env = _os.environ.copy()
    env['HOME'] = '/tmp'
    res = subprocess.run(
        [_SOFFICE, '--headless', '--convert-to', 'pdf', '--outdir', out_dir, docx_path],
        capture_output=True, text=True, env=env, timeout=180,
    )
    if res.returncode != 0:
        raise RuntimeError(f'PDF conversion failed: {res.stderr}')
    pdf_path = _os.path.splitext(docx_path)[0] + '.pdf'
    print(f'PDF saved: {pdf_path}')


# ────────────────────────────────────────────────────────────────────────────
# CAT 1
# ────────────────────────────────────────────────────────────────────────────
def build_cat1(output='files/Mourice_BBM_433_CAT_1.docx'):
    doc = new_doc()
    add_cover(
        doc,
        course_code='BBM 433',
        course_title='RETAIL AND MERCHANDISE',
        assignment='CAT 1',
        sub_date='24TH APRIL 2026',
        marks='30 MARKS',
    )

    # ─── QUESTION ONE — Naivas ────────────────────────────────────
    heading(doc, 'QUESTION ONE: NAIVAS DIGITAL MERCHANDISING (15 MARKS)', before=0)

    body(doc,
         'Naivas Supermarkets launched naivas.online to defend market share '
         'against Carrefour\'s online channel and emerging delivery apps such '
         'as Glovo, Jumia Food and Yum. The platform attracts strong traffic '
         'but converts only 1.2 % of visitors against an industry benchmark '
         'of 2.5 – 3 %, with 68 % of users dropping off at the product page '
         'or checkout. The diagnosis below applies the 7C Framework for '
         'Digital Retailing (Rayport & Jaworski, 2003) and Conversion Rate '
         'Optimization (CRO) principles to identify the failure points and '
         'propose a structured action plan.')

    subheading(doc, '(a) Three Critical 7Cs for Naivas — Weaknesses and Improvements (9 marks)')

    body(doc,
         'For a grocery e-commerce business serving Kenyan households, the '
         'three most critical Cs are Context, Content and Commerce. Grocery '
         'is a high-frequency, low-margin, fresh-perishable category in '
         'which trust, speed and payment friction make or break each '
         'purchase. The remaining Cs (Community, Connection, Customization, '
         'Communication) are important but secondary to the three below.')

    subheading(doc, '1. Context — Site Layout and Navigation')
    labelled(doc, 'Weakness 1: ',
             'Cluttered top-level navigation. The current menu lists more '
             'than twelve departments side-by-side, forcing shoppers into '
             'three or four taps before they can locate everyday staples '
             'such as bread, milk, unga and sukari. This violates the '
             '"three-tap rule" used by best-in-class grocery sites.')
    labelled(doc, 'Weakness 2: ',
             'Slow mobile load times (often above five seconds on a 3G '
             'connection) caused by un-optimised banner imagery and heavy '
             'JavaScript carousels. Google research shows bounce probability '
             'rises 32 % when load time goes from one to three seconds.')
    labelled(doc, 'Improvement: ',
             'Re-architect the homepage around a "frequently bought" rail '
             'driven by user purchase history; lazy-load all below-the-fold '
             'images; replace the JavaScript carousel with a static hero '
             'banner; and enforce a maximum-three-tap rule from landing → '
             'category → product → cart.')

    subheading(doc, '2. Content — Product Information Quality')
    labelled(doc, 'Weakness 1: ',
             'Generic manufacturer-supplied product photos with no scale '
             'reference, ingredient close-ups or "in-basket" lifestyle shots. '
             'Fresh-food categories (meat, vegetables, dairy) particularly '
             'suffer because shoppers cannot judge freshness or portion size.')
    labelled(doc, 'Weakness 2: ',
             'Inconsistent weight and quantity descriptions — for example, '
             'identical thumbnails for "Sukari 1 kg" and "Sukari 2 kg" packs '
             'lead to wrong-item purchases and negative reviews that further '
             'erode trust.')
    labelled(doc, 'Improvement: ',
             'Standardise every product detail page (PDP) with a high-quality '
             'primary image, a secondary scale shot, the gross weight in bold, '
             'an explicit expiry-buffer guarantee (e.g. "minimum 7 days to '
             'expiry on delivery"), and a structured "what\'s inside" '
             'specification table.')

    subheading(doc, '3. Commerce — Checkout and Payment Infrastructure')
    labelled(doc, 'Weakness 1: ',
             'Mandatory account creation before checkout — a documented '
             'friction point that Baymard Institute (2024) measures as '
             'responsible for 24 % of global cart abandonment. First-time '
             'shoppers in Kenya prefer to test the experience before '
             'committing personal data.')
    labelled(doc, 'Weakness 2: ',
             'Limited delivery-slot transparency. Available slots are only '
             'revealed after the customer has paid, breaching expectation. '
             'Customers in Nairobi who need same-day delivery rage-quit '
             'when they discover the earliest slot is two days later.')
    labelled(doc, 'Improvement: ',
             'Enable guest checkout backed by an M-PESA Express STK push, '
             'and surface a colour-coded delivery-slot calendar before the '
             'cart-review screen so the customer commits to a slot up front.')

    subheading(doc, '(b) Two CRO Tactics for the Cart-Abandonment Problem (6 marks)')

    subheading(doc, 'Tactic 1: One-Tap M-PESA Express Checkout with Persistent Cart')
    body(doc,
         'Replace the multi-step billing → shipping → payment funnel with a '
         'single STK push triggered from the cart page. The cart is keyed '
         'to the shopper\'s phone number and persists across sessions, so a '
         'customer interrupted by a phone call or weak signal can resume '
         'from any device without losing items.')
    labelled(doc, '7C Justification: ',
             'This tactic operates on Commerce (transaction infrastructure) '
             'and Connection (linking the website to the customer\'s '
             'preferred payment instrument). It removes the highest-friction '
             'step in Kenyan e-commerce — manual paybill entry — and '
             'directly attacks the 68 % checkout drop-off identified in '
             'the case data.')

    subheading(doc, 'Tactic 2: Exit-Intent "Save Your Basket" via WhatsApp')
    body(doc,
         'When the user moves the cursor toward the close-tab button or '
         'puts the app to background, present a modal offering to send the '
         'basket contents to their WhatsApp. The message contains a '
         'one-click deep-link that returns them to a pre-filled checkout '
         'with the same items, prices and delivery slot held for 30 minutes.')
    labelled(doc, '7C Justification: ',
             'This tactic operates on Communication (message channel) and '
             'Customization (personalised recovery content). WhatsApp '
             'delivers near-100 % open rates among Kenyan shoppers and the '
             'recovery message is personal rather than generic, so it '
             'recovers a meaningful share of the 68 % product-page and '
             'checkout drop-offs that other channels (email, push) fail '
             'to reach.')

    page_break(doc)

    # ─── QUESTION TWO — Vivo Fashion ──────────────────────────────
    heading(doc, 'QUESTION TWO: VIVO FASHION APP CONVERSION (15 MARKS)', before=0)

    body(doc,
         'Vivo Fashion has built a feature-rich mobile app, yet 55 % of '
         'users abandon their cart, the product detail page bounce rate is '
         '40 %, and only 12 % of users engage with the personalised '
         'recommendation engine. Acting as a digital merchandising '
         'consultant, the analysis below applies CRO principles and the '
         '7C Framework to identify the root causes and propose corrective '
         'experiments.')

    subheading(doc, '(a) Three Reasons for the High Abandonment and Bounce Rates with A/B Tests (9 marks)')

    subheading(doc, 'Reason 1: Weak Visual Hierarchy on the Product Detail Page')
    body(doc,
         'On the current PDP, the "Add to Cart" call-to-action competes '
         'visually with three equally-weighted secondary buttons — wishlist, '
         'share and size-guide. This breaches the F-pattern eye-tracking '
         'principle (Nielsen Norman Group) and dilutes the primary action.')
    labelled(doc, 'A/B Test: ',
             'Variant A (control) keeps the existing layout. Variant B '
             'promotes the "Add to Cart" CTA to a sticky bottom-of-screen '
             'pill in Vivo brand-orange, while greying the secondary '
             'actions and moving them into a "more options" overflow menu. '
             'Run for 14 days with at least 5,000 sessions per variant; '
             'primary metric: PDP-to-cart conversion rate.')

    subheading(doc, 'Reason 2: Absence of Urgency and Scarcity Cues')
    body(doc,
         'Fashion is an emotionally driven, identity-led purchase. The '
         'current app shows no stock indicators or time-bound offers, so '
         'shoppers feel no cost to delaying the decision. This is a classic '
         'CRO scarcity-and-urgency gap.')
    labelled(doc, 'A/B Test: ',
             'Variant A unchanged. Variant B introduces a "Only 3 left in '
             'your size" badge on PDPs and a 30-minute "Hold this item" '
             'reservation timer that activates when the user adds to cart. '
             'Primary metric: cart-to-checkout conversion rate; secondary: '
             'average time from cart-add to purchase.')

    subheading(doc, 'Reason 3: Excessive Checkout Friction')
    body(doc,
         'The current funnel forces shoppers through five separate screens '
         'between cart and confirmation, including a mandatory delivery-'
         'address typing step even for repeat customers whose addresses are '
         'already on file. Every additional screen compounds the abandonment '
         'risk.')
    labelled(doc, 'A/B Test: ',
             'Variant A (5-step current flow). Variant B condenses the flow '
             'into a single accordion screen with the saved delivery address '
             'auto-filled and biometric (fingerprint or Face ID) M-PESA '
             'confirmation. Primary metric: checkout completion rate; '
             'secondary: median time-to-purchase.')

    subheading(doc, '(b) Two 7C Elements Redesigned for Kenyan Omnichannel Fashion (6 marks)')

    subheading(doc, '1. Customization — Behaviour-Aware Recommendations Bridging In-App and In-Store Data')
    body(doc,
         'Replace the static "Recommended for you" carousel (currently '
         'engaged with by only 12 % of users) with a behaviour-aware '
         'recommendation engine that fuses in-app browsing history with '
         'in-store fitting-room data captured via RFID tags on every '
         'garment. A customer who tried on three dresses at the Vivo Sarit '
         'branch sees those exact items at the top of her app feed when she '
         'opens it within 24 hours, paired with matching shoes and bags. '
         'This addresses the dominant Kenyan fashion-shopper behaviour of '
         '"browse online, fit in-store, decide later" and turns physical '
         'visits into a personalisation signal rather than a lost touch '
         'point.')

    subheading(doc, '2. Communication — Event-Driven WhatsApp Recovery Threads')
    body(doc,
         'Replace generic broadcast push notifications (largely muted by '
         'users) with event-driven WhatsApp Business conversations. When a '
         'cart is abandoned for more than one hour, a polite WhatsApp '
         'message is sent containing the abandoned items, the nearest Vivo '
         'branch holding stock in the customer\'s size, and a "reserve for '
         'in-store collection" call-to-action. This deliberately monetises '
         'the omnichannel browse-online-buy-in-store pattern by converting '
         'the cart-abandonment leak into store footfall, rather than '
         'treating it as a lost sale. WhatsApp\'s near-universal reach in '
         'Kenya guarantees the message is seen, and the personalised, '
         'conversational tone outperforms transactional email and push by '
         'a wide margin.')

    body(doc,
         'In conclusion, Vivo\'s app does not need more features — it needs '
         'sharper visual hierarchy, stronger urgency triggers, a one-screen '
         'checkout and a customization-and-communication redesign that '
         'embraces, rather than ignores, the omnichannel reality of the '
         'Kenyan fashion shopper.', before=10)

    doc.save(output)
    print(f'DOCX saved: {output}')
    convert_to_pdf(output)


# ────────────────────────────────────────────────────────────────────────────
# CAT 2
# ────────────────────────────────────────────────────────────────────────────
def build_cat2(output='files/Mourice_BBM_433_CAT_2.docx'):
    doc = new_doc()
    add_cover(
        doc,
        course_code='BBM 433',
        course_title='RETAIL AND MERCHANDISE',
        assignment='CAT 2',
        sub_date='24TH APRIL 2026',
        marks='30 MARKS',
    )

    # ─── QUESTION ONE — Hotpoint ──────────────────────────────────
    heading(doc, 'QUESTION ONE: HOTPOINT PHYGITAL SHOWROOM (15 MARKS)', before=0)

    body(doc,
         'Hotpoint Appliances faces the classic showrooming dilemma: '
         'customers visit the Sarit Centre flagship to inspect appliances '
         'physically and then complete the transaction with a competitor '
         'offering a sharper online price or faster delivery. Remodelling '
         'the flagship into a "phygital" showroom requires a tightly '
         'integrated visual-merchandising and digital-integration plan that '
         'captures the in-store visit as a sale on any channel.')

    subheading(doc, '(a) Three Visual Merchandising Elements for the Sarit Showroom (6 marks)')

    subheading(doc, '1. Zoned Lifestyle Vignettes')
    body(doc,
         'Break the floor into four staged living spaces — a modern Kenyan '
         'kitchen, a family living room, a young-professional bedsit and an '
         'executive home office — in which appliances are merchandised in '
         'use rather than lined up on shelves. Kenyan electronics buyers '
         'place high cultural value on the home as a status symbol, so '
         'visualising a Samsung Bespoke fridge inside a complete kitchen '
         'vignette triggers the endowment effect, lengthens dwell time and '
         'lifts perceived value of the appliance against a sterile '
         'comparison-shopping experience online.')

    subheading(doc, '2. Strategic Focal Lighting and Premium Signage')
    body(doc,
         'Place category leaders — the new Samsung Bespoke fridge, the LG '
         'OLED TV, the Bosch dishwasher — directly under dropped pendant '
         'spotlights at 3,000 K colour temperature, with backlit brand-name '
         'signage above each piece. This guides the eye through the floor '
         'in a deliberate sequence and triggers the halo effect: shoppers '
         'who entered planning to buy a mid-range model frequently trade '
         'up when premium models are presented dramatically. The lighting '
         'budget is small relative to the average-basket-size lift this '
         'technique reliably produces in showroom retail.')

    subheading(doc, '3. Test-and-Touch Interactive Bays')
    body(doc,
         'Convert passive shelving into interactive bays: washing machines '
         'wired to demonstrate a live drum spin, blenders running at '
         'half-speed with crushed-ice samples on offer, and sound-bars '
         'looped on a familiar Sauti Sol track with a "press to play" '
         'button. Kenyan buyers culturally insist on physical verification '
         '("nione kwanza, nigusane") before committing to high-ticket '
         'purchases. Interactive bays satisfy this need at the showroom '
         'and remove the rationale for showrooming to a competitor — once '
         'the customer has touched, played with and trusted the product, '
         'the Sarit floor becomes the natural place to close the sale.')

    subheading(doc, '(b) Three Digital Integrations (6 marks)')

    subheading(doc, '1. QR-Coded "Scan to Compare" Shelf Labels')
    body(doc,
         'Every shelf-edge label carries a QR code that loads, on the '
         'customer\'s own phone, a live comparison sheet showing the price '
         'across Hotpoint.co.ke, Jumia, Killimall and the local Hotpoint '
         'price (with a price-match guarantee). This neutralises the '
         'showrooming threat directly: the customer no longer needs to '
         'leave the app or open a browser to compare prices.')
    labelled(doc, 'Local Challenge: ',
             'Unreliable in-mall Wi-Fi and cellular dead-spots near the '
             'service core.')
    labelled(doc, 'Mitigation: ',
             'Provision free Hotpoint guest Wi-Fi with a captive-portal '
             'sign-in (which doubles as a marketing data-capture channel), '
             'and ensure all QR-target pages are lightweight (< 300 KB) '
             'AMP pages cached on Cloudflare so they load even on weak 3G.')

    subheading(doc, '2. Endless-Aisle Touchscreen Kiosks')
    body(doc,
         'Install one 32-inch touchscreen kiosk per category corner where '
         'shoppers (or a staff member assisting them) can browse Hotpoint\'s '
         'full national inventory — including stock at the Mombasa and '
         'Kisumu warehouses — and order any out-of-stock SKU for free '
         'home delivery. This unifies physical and online inventory and '
         'turns "we don\'t have your colour in store" from a lost sale '
         'into a fulfilment opportunity.')
    labelled(doc, 'Local Challenge: ',
             'Varying digital literacy among older or rural visitors who '
             'are intimidated by touch interfaces.')
    labelled(doc, 'Mitigation: ',
             'Default the kiosk UI to Swahili with large icons, generous '
             'touch-target sizes, and a single prominent red "Itisha '
             'Msaidizi" (call assistant) button that pages a floor staff '
             'member through the in-store PA within 30 seconds.')

    subheading(doc, '3. WhatsApp Store-Associate Chat')
    body(doc,
         'A static QR code at the entrance and at the customer-care desk '
         'opens WhatsApp pre-loaded with the Sarit branch number. A '
         'numbered store associate (whose badge matches the WhatsApp '
         'profile) responds via tablet within 60 seconds for product '
         'questions, stock checks, or follow-up after the visit. The same '
         'thread continues post-visit, giving Hotpoint a warm channel for '
         'price-match offers, delivery confirmation and post-purchase '
         'support — all on the platform Kenyan shoppers already use daily.')
    labelled(doc, 'Local Challenge: ',
             'Staff training and 24/7 response coverage when associates '
             'go off shift.')
    labelled(doc, 'Mitigation: ',
             'Operate a rota of two trained associates per shift working '
             'from a 30-question playbook, with a chatbot answering '
             'after-hours messages and flagging high-intent conversations '
             '(price questions, "ready to buy") for follow-up first thing '
             'the next morning.')

    subheading(doc, '(c) Two KPIs for the Phygital Transformation (3 marks)')
    labelled(doc, 'KPI 1 — Phygital Conversion Rate. ',
             'The proportion of unique Sarit-store visitors (counted at the '
             'door by an automatic counter) who complete a purchase on any '
             'Hotpoint channel — in-store, on Hotpoint.co.ke, or via the '
             'WhatsApp associate — within 7 days of the visit, attributed '
             'through the Hotpoint loyalty card or a QR-scan customer ID. '
             'Target: lift from a baseline of approximately 8 % to 18 % '
             'within twelve months.')
    labelled(doc, 'KPI 2 — Cross-Channel Revenue per Visitor (CCRPV). ',
             'Total in-store + online revenue tagged to a customer ID, '
             'divided by the total number of Sarit visits in the period. '
             'This KPI captures whether the showroom is generating online '
             'sales rather than cannibalising them. Target: KSh 1,500 or '
             'higher per visit, with online attribution making up at least '
             '25 % of that figure.')

    page_break(doc)

    # ─── QUESTION TWO — Quickmart ─────────────────────────────────
    heading(doc, 'QUESTION TWO: QUICKMART OMNICHANNEL STRATEGY (15 MARKS)', before=0)

    body(doc,
         'Quickmart operates more than 60 outlets but its third-party '
         'online channel contributes only 4 % of revenue, against more '
         'than 15 % at Carrefour, whose physical and digital channels are '
         'tightly fused. Quickmart\'s 60-store network is currently an '
         'underutilised asset. The strategy below redefines the role of '
         'the physical store from a simple point of sale to an active hub '
         'in an integrated omnichannel model, supported by digital in-store '
         'touchpoints designed for the Kenyan operating context.')

    subheading(doc, '(a) Three Roles Physical Stores Can Play Beyond Points of Sale (6 marks)')

    subheading(doc, '1. Last-Mile Micro-Fulfilment Hubs')
    body(doc,
         'Instead of warehousing online orders centrally and shipping from '
         'a single distribution centre, Quickmart can convert each '
         'high-volume outlet (Lavington, Westlands, Eastleigh, Ruaka, '
         'Embakasi) into a micro-fulfilment hub serving online orders '
         'within a 10 km radius. Implementation: install picker stations '
         'and short-term holding shelves in the back-of-store, equip '
         'pickers with handheld scanners synced to the online order queue, '
         'and partner with Glovo, Bolt or Sendy for the final mile. The '
         'same shelf stock then serves walk-ins and online buyers, '
         'reducing inventory costs and enabling 2-hour same-day delivery '
         'without building any new warehouses.')

    subheading(doc, '2. Click-and-Collect (BOPIS) Pickup Points')
    body(doc,
         'Ring-fence the first two square metres at the entrance of every '
         'store as a "Quick Collect" desk, staffed by one dedicated '
         'employee. Online customers pay through the app and collect their '
         'order at a fixed slot without queueing through tills. '
         'Implementation: equip the desk with a tablet that pulls up the '
         'order on customer ID, give the customer a 60-second pickup '
         'experience, and route returns through the same counter. This '
         'opens a high-margin, low-cost fulfilment channel particularly '
         'attractive to matatu-bound office workers who pass a Quickmart '
         'on their commute.')

    subheading(doc, '3. Returns, Refunds and Customer-Acquisition Centres')
    body(doc,
         'The physical store handles online returns instantly — refund or '
         'exchange in the same visit — and serves as an in-person '
         'sign-up point for the Quickmart loyalty programme. '
         'Implementation: train the Quick Collect desk staff to issue '
         'instant refunds via M-PESA reversal up to a defined ceiling, '
         'and run weekly loyalty-card sign-up drives at the door staffed '
         'by a brand ambassador. Friction-free returns build online-'
         'purchase confidence (industry studies show free in-store returns '
         'lift online conversion 9 – 15 %), and on-the-ground sign-ups '
         'capture the meaningful share of Kenyan customers who are not '
         'comfortable signing up through an app.')

    subheading(doc, '(b) Two Digital In-Store Touchpoints (6 marks)')

    subheading(doc, '1. Self-Checkout Kiosks (M-PESA and Card)')
    body(doc,
         'Install two to three self-checkout kiosks at every flagship '
         'store. Customers scan their items, bag them, and pay by M-PESA '
         'STK push or card without queueing for a cashier. The kiosks '
         'should accept a maximum of fifteen items per transaction so '
         'that they remain genuinely "express" lanes and complement '
         'rather than replace conventional tills.')
    labelled(doc, 'Operational Challenge: ',
             'Traffic congestion at the kiosks during peak commuter hours '
             '(5:30 – 7:30 pm at urban outlets) and elevated theft risk '
             'when customers are unsupervised.')
    labelled(doc, 'Mitigation: ',
             'Pair every two kiosks with one supervising attendant on a '
             'walking patrol, install overhead cameras with weight-sensor '
             'verification on the bagging area to flag scan-and-bag '
             'mismatches, and limit each kiosk lane to fifteen items as '
             'an enforced soft safeguard. Add a "call assistant" button '
             'on every screen for customers who hit a problem, so the '
             'queue keeps moving.')

    subheading(doc, '2. "Scan-as-You-Shop" Mobile App with Express Pay Lane')
    body(doc,
         'Customers scan barcodes with the Quickmart app while filling '
         'their trolley; final payment is one M-PESA STK push at a '
         'dedicated express-pay lane where staff perform a brief audit '
         'check (e.g. random 10 % verification). This converts shopping '
         'time into checkout time, eliminates the till queue entirely '
         'for app users, and feeds rich basket data into the loyalty '
         'engine.')
    labelled(doc, 'Operational Challenge: ',
             'Staff training and customer adoption — both the floor team '
             'and the customer must trust the technology, and a scan-as-'
             'you-shop flow is unfamiliar to most Kenyan supermarket '
             'shoppers.')
    labelled(doc, 'Mitigation: ',
             'Roll out in two flagship stores first (e.g. Lavington and '
             'Thika Road Mall), staff each store with a "scan-and-shop '
             'ambassador" for the first month to coach customers, offer '
             'a 50-point sign-up bonus and a 10 % discount on the first '
             'scan-and-shop basket, and refresh staff training every six '
             'weeks tied to the loyalty data dashboard so any drop in '
             'adoption is identified early and addressed.')

    body(doc,
         'In conclusion, Quickmart\'s physical network is its most '
         'powerful — and currently most under-leveraged — omnichannel '
         'asset. Reframing the 60 stores as fulfilment hubs, BOPIS '
         'pickup points and customer-acquisition centres, and equipping '
         'them with self-checkout kiosks and a scan-as-you-shop app, '
         'will close the digital gap on Carrefour while making the '
         'physical visit faster, smarter and more profitable.', before=10)

    doc.save(output)
    print(f'DOCX saved: {output}')
    convert_to_pdf(output)


def main():
    build_cat1()
    build_cat2()


if __name__ == '__main__':
    main()
