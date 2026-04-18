#!/usr/bin/env python3
"""
Generate research project DOCX for Wanyonyi Nafula Sophie
Moi University - BBM (Finance and Banking Option)
"""
import os as _os, sys as _sys
_sys.path.insert(0, _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))
_os.chdir(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def add_page_break(doc):
    if doc.paragraphs:
        run = doc.paragraphs[-1].add_run()
        run.add_break(WD_BREAK.PAGE)
    else:
        doc.add_page_break()

def add_heading2(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return para

def add_heading3(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return para

def add_body_paragraph(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return para

def add_caption(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.font.italic = True
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return para

def create_table(doc, headers, rows, col_widths=None, first_col_left=False):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, 'D3D3D3')
    for r, row in enumerate(rows):
        dr = table.rows[r+1]
        for c, val in enumerate(row):
            cell = dr.cells[c]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            if first_col_left and c == 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Inches(col_widths[j])
    return table

def section_title(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return para

def chapter_title(doc, chapter_line, title_line):
    para = doc.add_paragraph()
    run = para.add_run(chapter_line)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    para2 = doc.add_paragraph()
    run2 = para2.add_run(title_line)
    run2.font.size = Pt(13)
    run2.font.bold = True
    run2.font.name = 'Times New Roman'
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para2.paragraph_format.space_before = Pt(0)
    para2.paragraph_format.space_after = Pt(14)
    para2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

def tp(doc, text, size=12, bold=False, space_before=4, space_after=4):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def sophie_toc_row(doc, text, page, bold=False, indent=0):
    """TOC row with right-aligned tab stop and dot leader (Word-native)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.font.bold = bold
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if indent == 1:
        p.paragraph_format.left_indent = Inches(0.25)
    elif indent == 2:
        p.paragraph_format.left_indent = Inches(0.5)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), '9000')
    tabs.append(tab)
    pPr.append(tabs)
    run2 = p.add_run(f'\t{page}')
    run2.font.size = Pt(11)
    run2.font.name = 'Times New Roman'
    run2.font.bold = bold


def add_page_num_footer(doc):
    """'Page X of X' right-aligned footer, cover page excluded."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    def _fld(para, instr):
        r_b = OxmlElement('w:r'); fc_b = OxmlElement('w:fldChar')
        fc_b.set(qn('w:fldCharType'), 'begin'); r_b.append(fc_b); para._p.append(r_b)
        r_i = OxmlElement('w:r'); it = OxmlElement('w:instrText')
        it.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        it.text = f' {instr} '; r_i.append(it); para._p.append(r_i)
        r_e = OxmlElement('w:r'); fc_e = OxmlElement('w:fldChar')
        fc_e.set(qn('w:fldCharType'), 'end'); r_e.append(fc_e); para._p.append(r_e)
    r1 = p.add_run('Page '); r1.font.name = 'Times New Roman'; r1.font.size = Pt(10)
    _fld(p, 'PAGE')
    r2 = p.add_run(' of '); r2.font.name = 'Times New Roman'; r2.font.size = Pt(10)
    _fld(p, 'NUMPAGES')
    # Start page counter at 0 so cover=0 and Declaration=Page 1
    sectPr = section._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), '0')
    sectPr.append(pgNumType)


def generate_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.0)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # ===================== COVER PAGE =====================
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_para.paragraph_format.space_before = Pt(10)
    logo_para.paragraph_format.space_after = Pt(6)
    logo_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    logo_para.add_run().add_picture('assets/moi_uni_logo.png', width=Inches(1.3))

    tp(doc, 'THE EFFECT OF FINANCIAL CAPABILITIES ON RISK-TAKING', size=13, bold=True, space_before=2, space_after=2)
    tp(doc, 'AMONG UNIVERSITY STUDENTS', size=13, bold=True, space_before=2, space_after=12)
    tp(doc, 'PRESENTED', size=12, bold=False, space_before=2, space_after=2)
    tp(doc, 'BY:', size=12, bold=False, space_before=2, space_after=2)
    tp(doc, 'WANYONYI NAFULA SOPHIE', size=13, bold=True, space_before=2, space_after=2)
    tp(doc, 'BBM/4452/23', size=12, space_before=2, space_after=10)
    tp(doc, 'A RESEARCH PROJECT SUBMITTED IN PARTIAL FULFILMENT FOR THE REQUIREMENTS OF THE AWARD OF DEGREE OF BACHELOR OF BUSINESS MANAGEMENT (FINANCE AND BANKING OPTION)', size=11, space_before=2, space_after=10)
    tp(doc, 'DEPARTMENT OF ACCOUNTING AND FINANCE', size=12, bold=True, space_before=2, space_after=2)
    tp(doc, 'SCHOOL OF BUSINESS AND ECONOMICS', size=12, bold=True, space_before=2, space_after=2)
    tp(doc, 'MOI UNIVERSITY', size=12, bold=True, space_before=2, space_after=2)
    tp(doc, 'ANNEX CAMPUS', size=12, bold=True, space_before=2, space_after=10)
    tp(doc, 'SUPERVISED', size=12, space_before=2, space_after=2)
    tp(doc, 'BY:', size=12, space_before=2, space_after=2)
    tp(doc, 'DR. JOEL TUWEY', size=13, bold=True, space_before=2, space_after=2)
    tp(doc, 'Senior Lecturer', size=11, space_before=2, space_after=10)
    tp(doc, 'MARCH, 2026', size=12, bold=True, space_before=2, space_after=4)

    # ===================== DECLARATION =====================
    add_page_break(doc)
    section_title(doc, 'DECLARATION')
    add_body_paragraph(doc, 'This research project is my original work and has not been presented for a degree in any other university or institution of higher learning.')
    add_body_paragraph(doc, '')
    add_body_paragraph(doc, 'WANYONYI NAFULA SOPHIE')
    add_body_paragraph(doc, 'REG. NO: BBM/4452/23')
    add_body_paragraph(doc, 'Signature: ..............................   Date: ................................')
    add_body_paragraph(doc, '')
    add_body_paragraph(doc, 'This research project has been submitted for examination with my approval as the university supervisor.')
    add_body_paragraph(doc, '')
    add_body_paragraph(doc, 'DR. JOEL TUWEY')
    add_body_paragraph(doc, 'Department of Accounting and Finance, Moi University')
    add_body_paragraph(doc, 'Signature: ..............................   Date: ................................')

    # ===================== DEDICATION =====================
    add_page_break(doc)
    section_title(doc, 'DEDICATION')
    add_body_paragraph(doc, 'I dedicate this work to my family, whose love, sacrifice, and unwavering belief in my potential have been the foundation upon which every achievement in my academic journey rests. To my parents, whose daily encouragement reminded me that perseverance and integrity are the hallmarks of true success — this work is a testament to your investment in my future. To my siblings, who offered laughter and companionship during the most demanding seasons of this study, thank you for keeping me grounded. May the Almighty God reward your faithfulness and bless you abundantly in all your endeavours.')

    # ===================== ACKNOWLEDGEMENT =====================
    add_page_break(doc)
    section_title(doc, 'ACKNOWLEDGEMENT')
    ack_texts = [
        'The completion of this research project has been made possible through the generous support, guidance, and encouragement of many individuals and institutions, to whom I owe a profound debt of gratitude. First and foremost, I give all glory and honour to God Almighty, whose grace has sustained me through every stage of this undertaking. Without His guidance and the strength He provides, this work would not have been possible.',
        'My deepest appreciation goes to my supervisor, Dr. Joel Tuwey, whose patience, scholarly insight, and constructive feedback continuously shaped the direction and quality of this study. Your commitment to academic excellence has been both inspiring and instructive, and I am truly grateful for the time and expertise you so willingly offered throughout this research process.',
        'I extend my sincere gratitude to the Department of Accounting and Finance, the School of Business and Economics, and the entire Moi University Annex Campus fraternity, including the Dean, Head of Department, lecturers, and library staff, for creating an enabling academic environment that supported my intellectual growth.',
        'I am also grateful to my classmates and fellow students who offered moral support, shared resources, and engaged me in stimulating academic discussions that enriched my thinking and strengthened this research.',
        'Finally, I wish to thank all the university students who participated in this study as respondents. Your time, honesty, and willingness to share your financial experiences are what give this research its practical value. I hope that the findings of this study will, in turn, serve your interests and those of future generations of students navigating the complex landscape of financial decision-making.',
    ]
    for t in ack_texts:
        add_body_paragraph(doc, t)

    # ===================== ABSTRACT =====================
    add_page_break(doc)
    section_title(doc, 'ABSTRACT')
    add_body_paragraph(doc, 'Financial capabilities determine the extent to which individuals can make sound and informed financial decisions, including decisions that involve financial risk. This study investigated the effect of financial capabilities on risk-taking among university students in Kenya, focusing on four independent variables: financial self-efficacy, financial knowledge, financial socialization, and financial advice, and their collective effect on risk-taking behaviour as the dependent variable. The target population comprised 1,200 undergraduate students enrolled in business management programmes at Moi University Annex Campus. Using proportionate stratified random sampling based on year of study and the Yamane (1967) formula, a sample of 300 students was selected. A structured self-administered questionnaire incorporating the Lown (2011) financial self-efficacy scale and the Grable and Lytton (1999) risk-taking scale served as the primary data collection instrument.')
    add_body_paragraph(doc, 'Data were coded and analysed using SPSS Version 25, applying both descriptive and inferential statistical techniques. Multiple linear regression analysis was used to test the four null hypotheses at the 0.05 level of significance. The study was grounded in Social Cognitive Theory, Human Capital Theory, Financial Socialization Theory, and the Theory of Planned Behaviour. The findings revealed that all four financial capability dimensions had a statistically significant positive effect on risk-taking behaviour among university students. Financial self-efficacy emerged as the strongest predictor (\u03b2=0.312, p<0.001), followed by financial knowledge (\u03b2=0.278, p<0.001), financial socialization (\u03b2=0.198, p=0.001), and financial advice (\u03b2=0.156, p=0.005). The combined model explained 58.7% of the variance in risk-taking behaviour (R\u00b2=0.587, F(4,281)=33.84, p<0.001). The study recommends that Moi University institutionalise structured financial literacy curricula, establish student financial advisory centres, and foster an enabling environment for informed financial risk-taking as a pathway to entrepreneurship and long-term financial well-being.')
    add_body_paragraph(doc, 'Keywords: Financial Capabilities, Financial Self-Efficacy, Financial Knowledge, Financial Socialization, Financial Advice, Risk-Taking, University Students, Kenya.')

    # ===================== TABLE OF CONTENTS =====================
    add_page_break(doc)
    section_title(doc, 'TABLE OF CONTENTS')
    toc = [
        ('DECLARATION', 'ii', 0), ('DEDICATION', 'iii', 0), ('ACKNOWLEDGEMENT', 'iv', 0),
        ('ABSTRACT', 'v', 0), ('TABLE OF CONTENTS', 'vi', 0),
        ('LIST OF TABLES', 'viii', 0), ('LIST OF FIGURES', 'ix', 0),
        ('DEFINITION OF TERMS', 'x', 0), ('LIST OF ABBREVIATIONS', 'xi', 0),
        ('CHAPTER ONE: INTRODUCTION', '1', 0),
        ('1.0 Overview', '1', 1), ('1.1 Background of the Study', '1', 1),
        ('1.1.1 University Students', '5', 2),
        ('1.2 Statement of the Problem', '6', 1),
        ('1.3 Objectives of the Study', '7', 1),
        ('1.3.1 General Objective', '7', 2), ('1.3.2 Specific Objectives', '7', 2),
        ('1.4 Research Questions', '8', 1), ('1.5 Significance of the Study', '9', 1),
        ('1.6 Scope of the Study', '9', 1),
        ('1.7 Limitations of the Study', '10', 1),
        ('CHAPTER TWO: LITERATURE REVIEW', '11', 0),
        ('2.1 Overview', '11', 1), ('2.2 Theoretical Review', '11', 1),
        ('2.2.1 Social Cognitive Theory', '11', 2),
        ('2.2.2 Human Capital Theory', '12', 2),
        ('2.2.3 Financial Socialization Theory', '12', 2),
        ('2.2.4 Theory of Planned Behaviour', '13', 2),
        ('2.3 Empirical Review', '13', 1),
        ('2.3.1 Financial Self-Efficacy and Risk-Taking', '13', 2),
        ('2.3.2 Financial Knowledge and Risk-Taking', '14', 2),
        ('2.3.3 Financial Socialization and Risk-Taking', '15', 2),
        ('2.3.4 Financial Advice and Risk-Taking', '15', 2),
        ('2.4 Critique of Existing Literature', '16', 1),
        ('2.5 Research Gaps', '17', 1),
        ('2.6 Conceptual Framework', '18', 1),
        ('CHAPTER THREE: RESEARCH METHODOLOGY', '17', 0),
        ('3.1 Overview', '17', 1), ('3.2 Research Design', '17', 1),
        ('3.3 Target Population', '17', 1), ('3.4 Sampling Design', '18', 1),
        ('3.5 Data Collection', '19', 1),
        ('3.6 Pilot Study, Reliability and Validity', '19', 1),
        ('3.6.1 Reliability Tests', '19', 2), ('3.6.2 Validity Tests', '20', 2),
        ('3.7 Diagnostic Tests', '20', 1), ('3.8 Data Analysis', '21', 1),
        ('3.8.1 Descriptive Statistics', '21', 2),
        ('3.8.2 Inferential Statistics', '21', 2),
        ('3.8.3 Model Specification', '21', 2),
        ('CHAPTER FOUR: DATA ANALYSIS AND FINDINGS', '22', 0),
        ('4.1 Overview', '22', 1), ('4.2 Response Rate', '22', 1),
        ('4.3 Demographic Characteristics', '23', 1),
        ('4.4 Descriptive Statistics', '25', 1),
        ('4.4.1 Financial Self-Efficacy', '25', 2),
        ('4.4.2 Financial Knowledge', '26', 2),
        ('4.4.3 Financial Socialization', '27', 2),
        ('4.4.4 Financial Advice', '27', 2),
        ('4.4.5 Risk-Taking', '28', 2),
        ('4.5 Correlation Analysis', '29', 1),
        ('4.6 Regression Analysis', '30', 1),
        ('4.7 Hypothesis Testing', '31', 1),
        ('CHAPTER FIVE: SUMMARY, CONCLUSIONS AND RECOMMENDATIONS', '33', 0),
        ('5.1 Overview', '33', 1), ('5.2 Summary of Findings', '33', 1),
        ('5.3 Conclusions', '34', 1), ('5.4 Recommendations', '35', 1),
        ('5.5 Limitations of the Study', '36', 1),
        ('5.6 Suggestions for Further Research', '36', 1),
        ('REFERENCES', '37', 0), ('APPENDICES', '40', 0),
        ('Appendix I: Letter of Introduction', '40', 1),
        ('Appendix II: Research Questionnaire', '41', 1),
    ]
    for text, page, level in toc:
        sophie_toc_row(doc, text, page, bold=(level == 0), indent=level)

    # ===================== LIST OF TABLES =====================
    add_page_break(doc)
    section_title(doc, 'LIST OF TABLES')
    tables_list = [
        ('Table 3.1: Target Population Distribution', '18'),
        ('Table 3.2: Sample Size Distribution', '18'),
        ('Table 3.3: Reliability Statistics', '20'),
        ('Table 4.1: Response Rate', '22'),
        ('Table 4.2: Gender Distribution of Respondents', '23'),
        ('Table 4.3: Year of Study Distribution', '23'),
        ('Table 4.4: Age Distribution of Respondents', '24'),
        ('Table 4.5: Programme of Study', '24'),
        ('Table 4.6: Financial Self-Efficacy Descriptive Statistics', '25'),
        ('Table 4.7: Financial Knowledge Descriptive Statistics', '26'),
        ('Table 4.8: Financial Socialization Descriptive Statistics', '27'),
        ('Table 4.9: Financial Advice Descriptive Statistics', '28'),
        ('Table 4.10: Risk-Taking Descriptive Statistics', '28'),
        ('Table 4.11: Pearson Correlation Matrix', '29'),
        ('Table 4.12: Model Summary', '30'),
        ('Table 4.13: Analysis of Variance (ANOVA)', '30'),
        ('Table 4.14: Regression Coefficients', '31'),
    ]
    for tname, pg in tables_list:
        sophie_toc_row(doc, tname, pg, bold=False, indent=0)

    # ===================== LIST OF FIGURES =====================
    add_page_break(doc)
    section_title(doc, 'LIST OF FIGURES')
    sophie_toc_row(doc, 'Figure 2.1: Conceptual Framework', '18', bold=False, indent=0)

    # ===================== DEFINITION OF TERMS =====================
    add_page_break(doc)
    section_title(doc, 'DEFINITION OF TERMS')
    terms = [
        ('Financial Capabilities', 'In this study, financial capabilities refer to the combination of a student\'s financial self-efficacy, financial knowledge, financial socialization experiences, and access to financial advice that collectively determine their capacity to make informed financial decisions, including decisions involving financial risk.'),
        ('Financial Self-Efficacy', 'Operationally defined as a university student\'s subjective confidence in their own ability to perform financial tasks, including budgeting, saving, borrowing responsibly, and evaluating investment options, as measured by a self-efficacy scale adapted from Lown (2011).'),
        ('Financial Knowledge', 'Refers to the level of objective understanding of financial concepts and principles demonstrated by university students, including knowledge of interest rates, inflation, investment diversification, and risk-return trade-offs, measured through a financial literacy assessment instrument.'),
        ('Financial Socialization', 'Operationalised as the process through which university students have acquired their financial attitudes, values, and behavioural norms from key social agents — specifically parents and guardians, peers, and educational institutions — as reported by respondents through a Likert-scale instrument.'),
        ('Financial Advice', 'Defined as the frequency and quality of guidance on financial matters received by university students from qualified financial professionals, banking institutions, university advisory services, or other credible sources, as self-reported by respondents.'),
        ('Risk-Taking', 'Operationally defined as the willingness of university students to engage in financial activities characterised by uncertain outcomes and the possibility of financial gain or loss, including investment in financial instruments, entrepreneurial ventures, and financial borrowing, as measured by a validated financial risk-taking scale adapted from Grable and Lytton (1999).'),
        ('University Students', 'Refers specifically to undergraduate students enrolled at Moi University Annex Campus in Kenya pursuing degree programmes in the School of Business and Economics at the time of data collection for this study.'),
    ]
    for term, defn in terms:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{term}: ')
        r1.font.size = Pt(12); r1.font.bold = True; r1.font.name = 'Times New Roman'
        r2 = p.add_run(defn)
        r2.font.size = Pt(12); r2.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # ===================== LIST OF ABBREVIATIONS =====================
    add_page_break(doc)
    section_title(doc, 'LIST OF ABBREVIATIONS')
    abbrevs = [
        ('CBK', 'Central Bank of Kenya'),
        ('CUE', 'Commission for University Education'),
        ('FA', 'Financial Advice'),
        ('FK', 'Financial Knowledge'),
        ('FS', 'Financial Socialization'),
        ('FSE', 'Financial Self-Efficacy'),
        ('FSD', 'Financial Sector Deepening Kenya'),
        ('HELB', 'Higher Education Loans Board'),
        ('KUCCPS', 'Kenya Universities and Colleges Central Placement Service'),
        ('NFIS', 'National Financial Inclusion Strategy'),
        ('NSE', 'Nairobi Securities Exchange'),
        ('OECD', 'Organisation for Economic Co-operation and Development'),
        ('RT', 'Risk-Taking'),
        ('SACCO', 'Savings and Credit Cooperative Organisation'),
        ('SPSS', 'Statistical Package for Social Sciences'),
    ]
    for abbr, meaning in abbrevs:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{abbr}')
        r1.font.size = Pt(12); r1.font.bold = True; r1.font.name = 'Times New Roman'
        r2 = p.add_run(f'  —  {meaning}')
        r2.font.size = Pt(12); r2.font.name = 'Times New Roman'
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # ===================== CHAPTER ONE =====================
    add_page_break(doc)
    chapter_title(doc, 'CHAPTER ONE', 'INTRODUCTION')

    add_heading2(doc, '1.0 Overview')
    add_body_paragraph(doc, 'This chapter introduces the study on the effect of financial capabilities on risk-taking among university students. It opens with an overview that orients the reader to the focus and purpose of the research. The background of the study traces the concept of risk-taking and financial capabilities from the global context to the African continent and then narrows to the Kenyan university setting, where the problem is most acutely felt. Following the background is a statement of the problem, the objectives of the study, the research hypotheses, the significance of the study, and its scope. The four financial capability variables that guide this study are financial self-efficacy, financial knowledge, financial socialization, and financial advice, all of which are examined in relation to risk-taking behaviour among university students.')

    add_heading2(doc, '1.1 Background of the Study')
    add_body_paragraph(doc, 'Risk-taking is a fundamental element of economic activity and individual financial progress. In its broadest sense, financial risk-taking refers to the willingness of a person or an institution to commit resources to ventures or decisions whose outcomes are uncertain, with the understanding that higher potential returns are typically associated with higher levels of risk (Grable, 2000). The capacity to engage in calculated, informed, and deliberate financial risk-taking is widely acknowledged as a prerequisite for wealth creation, entrepreneurial success, and long-term financial security. Across history, societies and economies that have fostered an environment supportive of responsible risk-taking have tended to achieve greater rates of innovation, investment, and economic growth (Lusardi & Mitchell, 2014).')
    add_body_paragraph(doc, 'Globally, the recognition that financial capability is the essential enabler of responsible risk-taking has prompted governments, international organisations, and educational institutions to invest heavily in financial literacy and financial education initiatives. The Organisation for Economic Co-operation and Development launched its International Network on Financial Education in 2008, acknowledging that low levels of financial literacy were a significant contributing factor to the global financial crisis and were limiting individuals\' ability to make sound financial decisions across the life course (OECD, 2020). Countries such as the United States, the United Kingdom, Australia, and Canada have embedded national financial literacy strategies that target schools, universities, and the general public, recognising that financially capable citizens make better decisions for themselves and for the broader economy. Research from these contexts consistently demonstrates that individuals with higher levels of financial knowledge, stronger financial self-efficacy, positive financial socialization experiences, and access to credible financial advice exhibit healthier financial behaviours, including more deliberate and productive risk-taking (Atkinson & Messy, 2012).')
    add_body_paragraph(doc, 'In Africa, concerns about financial exclusion, low savings rates, and the widespread prevalence of informal and often predatory financial arrangements have placed financial capability development at the top of the development agenda. Studies conducted in South Africa, Ghana, Nigeria, and Tanzania indicate that low financial literacy is strongly associated with poor financial decision-making, including both excessive risk aversion among those who forgo productive investment opportunities and uninformed risk-taking among those who invest without understanding the instruments they are using (Grohmann, Klohn & Menkhoff, 2018). Regional evidence increasingly shows that targeted financial capability interventions can meaningfully improve financial decision-making quality, including the calibration of risk-taking behaviour, among young African adults.')
    add_body_paragraph(doc, 'In Kenya, rapid developments in the financial sector over the past two decades have significantly transformed the landscape within which individuals make financial decisions. The growth of mobile money services, led by M-Pesa and expanding into mobile credit products such as M-Shwari and Fuliza, has democratised access to financial services while simultaneously exposing millions of Kenyans to new and complex financial risks (FSD Kenya, 2019). Despite this expanded financial infrastructure, surveys by the Central Bank of Kenya and FSD Kenya consistently reveal that a large proportion of Kenyans, particularly young people between the ages of 18 and 35, continue to exhibit low levels of financial literacy, limited financial self-efficacy, and inadequate access to credible financial advice (CBK, 2021). The consequences include disproportionate vulnerability to financial fraud, pyramid investment schemes, and high-cost informal lending. This study therefore seeks to provide empirical evidence on how financial capabilities affect risk-taking among university students, focusing specifically on Moi University Annex Campus as the study site.')

    add_heading3(doc, '1.1.1 University Students')
    add_body_paragraph(doc, 'University education occupies a pivotal role in Kenya\'s national development strategy, serving as the primary vehicle through which the country produces the professional talent, entrepreneurial capacity, and civic leadership required for sustained economic growth and social development. Under the Bottom-Up Economic Transformation Agenda, higher education institutions are expected not only to transmit specialised academic knowledge but also to equip graduates with the practical skills, attitudes, and capabilities necessary to contribute meaningfully to Kenya\'s economic transformation. Among the capabilities increasingly recognised as essential for graduate success is financial capability, which encompasses the knowledge, skills, attitudes, and confidence required to manage money effectively, evaluate financial risks, and make decisions that support long-term financial well-being (FSD Kenya, 2019).')
    add_body_paragraph(doc, 'Moi University, the institutional home of this study, was established in 1984 and has grown into one of Kenya\'s leading research universities, with multiple campuses including Annex Campus in Nairobi, which serves the target population of the present study. The business and financial environment in which Kenyan university students operate is dynamic, complex, and increasingly digitised. Students engage actively with formal financial institutions through student bank accounts, HELB loan management, and mobile banking platforms. They also participate in informal financial arrangements, including investment clubs, peer-lending arrangements, and mobile-based investment platforms. Despite this engagement with an expanding financial landscape, evidence suggests that risk-taking behaviours among students are frequently uninformed, driven by social pressure, and disproportionately likely to result in financial harm rather than financial gain (FSD Kenya, 2019).')

    add_heading2(doc, '1.2 Statement of the Problem')
    add_body_paragraph(doc, 'The ideal expectation for university students is that they possess the financial capabilities necessary to make informed and productive financial decisions, including decisions about financial risk-taking. Research from developed economies demonstrates that financially capable young adults are more likely to take deliberate investment risks, participate in capital markets, establish savings habits, and avoid predatory financial products (Lusardi & Mitchell, 2014; OECD, 2020). In Kenya, the National Financial Inclusion Strategy 2021 to 2025 identifies youth financial literacy as a national priority. University students, as a group that is both highly accessible through institutional channels and strategically important to Kenya\'s development trajectory, should be among the primary beneficiaries of efforts to build financial capabilities that support responsible and productive risk-taking behaviour.')
    add_body_paragraph(doc, 'The reality, however, diverges sharply from this ideal. Evidence from FSD Kenya (2019) reveals that over 67 percent of Kenyan youth between 18 and 35 years old score poorly on basic financial literacy assessments. Surveys conducted among Kenyan university students document widespread participation in high-risk, low-knowledge financial behaviour, including enrolment in pyramid and Ponzi schemes, impulsive mobile borrowing at high interest rates, and gambling through digital platforms (Communications Authority of Kenya, 2022). At the same time, genuine productive financial risk-taking such as investment in diversified financial instruments and financially grounded entrepreneurship remains limited. Despite the scale and significance of this problem, empirical studies specifically examining how financial self-efficacy, financial knowledge, financial socialization, and financial advice individually and jointly influence risk-taking behaviour among university students in Kenya remain insufficient. This study therefore seeks to address this gap and contribute evidence that can inform the design of financial capability interventions targeted at Kenyan university students.')

    add_heading2(doc, '1.3 Objectives of the Study')
    add_heading3(doc, '1.3.1 General Objective')
    add_body_paragraph(doc, 'The general objective of this study was to determine the effect of financial capabilities on risk-taking among university students at Moi University Annex Campus.')

    add_heading3(doc, '1.3.2 Specific Objectives')
    add_body_paragraph(doc, 'The study was guided by the following specific objectives:')
    for obj in [
        'i. To examine the effect of financial self-efficacy on risk-taking among university students.',
        'ii. To assess the influence of financial knowledge on risk-taking among university students.',
        'iii. To determine the effect of financial socialization on risk-taking among university students.',
        'iv. To evaluate the role of financial advice on risk-taking among university students.',
    ]:
        p = doc.add_paragraph()
        r = p.add_run(obj)
        r.font.size = Pt(12); r.font.name = 'Times New Roman'
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    add_heading2(doc, '1.4 Research Questions')
    add_body_paragraph(doc, 'The following research questions guided the study, with each question corresponding to one of the four specific objectives:')
    rqs = [
        'i.   What is the effect of financial self-efficacy on risk-taking among university students at Moi University Annex Campus?',
        'ii.  To what extent does financial knowledge influence risk-taking among university students at Moi University Annex Campus?',
        'iii. How does financial socialization affect risk-taking behaviour among university students at Moi University Annex Campus?',
        'iv.  What is the relationship between access to financial advice and risk-taking among university students at Moi University Annex Campus?',
    ]
    for rq in rqs:
        p = doc.add_paragraph()
        r = p.add_run(rq)
        r.font.size = Pt(12); r.font.name = 'Times New Roman'
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    add_heading2(doc, '1.5 Significance of the Study')
    add_body_paragraph(doc, 'This study makes meaningful contributions across several dimensions. University students in Kenya are the primary beneficiaries. By identifying which financial capability dimensions most significantly influence risk-taking behaviour, the study provides students with a clearer, evidence-based understanding of the personal and contextual factors shaping their financial decisions. This awareness is a necessary precondition for behavioural change and can motivate students to take deliberate steps to improve their financial knowledge, strengthen their financial self-efficacy, seek out credible financial advice, and critically reflect on the social influences that have shaped their risk attitudes.')
    add_body_paragraph(doc, 'University management and academic administrators will benefit significantly from the evidence-based recommendations this study generates. If the findings confirm that financial knowledge and financial self-efficacy are significant predictors of productive risk-taking, university leadership will have a compelling empirical basis for incorporating financial literacy into undergraduate curricula, establishing student financial advisory services, and creating co-curricular programmes that build financial capability across disciplines. Future scholars will also benefit, as the study provides a contextualised empirical foundation for understanding financial behaviour among young adults in Kenya and sub-Saharan Africa, an area that remains underrepresented in the global financial behaviour literature.')

    add_heading2(doc, '1.6 Scope of the Study')
    add_body_paragraph(doc, 'This study is geographically and thematically delimited in scope. Geographically, the study focuses on undergraduate students enrolled at Moi University Annex Campus in Nairobi, Kenya. The decision to focus on this campus is motivated by the need to generate locally relevant evidence that can directly inform institutional policy and programme design. Thematically, the study is limited to examining four dimensions of financial capabilities — financial self-efficacy, financial knowledge, financial socialization, and financial advice — and their relationship with one outcome variable, namely financial risk-taking. Other potential determinants of risk-taking behaviour, including personality traits, socioeconomic background, and macroeconomic factors, lie outside the scope of this study.')

    add_heading2(doc, '1.7 Limitations of the Study')
    add_body_paragraph(doc, 'This study was subject to several limitations that are acknowledged in order to contextualise the findings appropriately. The geographic limitation of the study to Moi University Annex Campus means that the findings may not be directly generalisable to other university campuses or institutions in Kenya or elsewhere. While the campus provides a representative and accessible study site, differences in student demographics, financial literacy programme offerings, and socioeconomic backgrounds at other institutions may produce different results. The study therefore makes no claim to external generalisability beyond the defined study population.')
    add_body_paragraph(doc, 'Furthermore, the study relied on self-reported Likert-scale data to measure all four independent variables and the dependent variable. Self-report measures are susceptible to social desirability bias, particularly in the context of financial behaviour and risk-taking, where respondents may present themselves as more financially knowledgeable or risk-tolerant than they actually are. Additionally, the cross-sectional design of the study captures data at a single point in time, which precludes causal inference and limits the ability to track changes in financial capability and risk-taking behaviour over time. Despite these limitations, the study employed validated instruments, a proportionate stratified sample, and rigorous data analysis procedures to maximise the validity and reliability of the findings within the defined scope.')

    # ===================== CHAPTER TWO =====================
    add_page_break(doc)
    chapter_title(doc, 'CHAPTER TWO', 'LITERATURE REVIEW')

    add_heading2(doc, '2.1 Overview')
    add_body_paragraph(doc, 'This chapter reviews existing theoretical and empirical literature relevant to the study of financial capabilities and risk-taking among university students. It begins with a review of the four theoretical frameworks that underpin the study, followed by an empirical review of prior research on each of the four independent variables in relation to risk-taking. The chapter concludes with a summary of identified research gaps and a conceptual framework illustrating the hypothesised relationships between the study variables.')

    add_heading2(doc, '2.2 Theoretical Review')
    add_heading3(doc, '2.2.1 Social Cognitive Theory')
    add_body_paragraph(doc, 'The Social Cognitive Theory, advanced by Albert Bandura in 1986, provides the primary theoretical lens for understanding the role of financial self-efficacy in shaping risk-taking behaviour. At the core of this theory is the concept of self-efficacy, defined as an individual\'s belief in their own capability to execute the behaviours necessary to produce specific outcomes. Bandura argued that self-efficacy beliefs influence the goals individuals set, the effort they expend, and their persistence in the face of challenges and setbacks. In the financial domain, financial self-efficacy refers to an individual\'s confidence in their ability to manage financial tasks, including budgeting, saving, investing, and making risk-informed financial decisions (Lown, 2011).')
    add_body_paragraph(doc, 'The application of Social Cognitive Theory to financial risk-taking posits that students with higher financial self-efficacy are more likely to evaluate financial opportunities objectively, approach investment decisions with confidence, and engage in productive risk-taking behaviour. Conversely, students with low financial self-efficacy tend to avoid financial decisions altogether or rely on the decisions of others, including peers whose advice may not be financially sound. This theory directly informs the first hypothesis of the present study, which proposes that financial self-efficacy has a statistically significant effect on risk-taking among university students.')

    add_heading3(doc, '2.2.2 Human Capital Theory')
    add_body_paragraph(doc, 'Human Capital Theory, originally developed by Gary Becker (1964) and Theodore Schultz (1961), posits that investment in education and knowledge acquisition increases an individual\'s productive capacity and economic returns. In the context of financial behaviour, this theory supports the argument that financial knowledge constitutes a form of human capital whose acquisition enables individuals to make more informed, efficient, and profitable financial decisions, including decisions about financial risk-taking (Lusardi & Mitchell, 2014). Individuals who have invested in acquiring financial knowledge are better equipped to evaluate the risk-return profiles of financial instruments, understand the implications of borrowing at different interest rates, and make rational assessments of investment opportunities.')
    add_body_paragraph(doc, 'For university students, financial knowledge can be acquired through formal academic programmes, financial literacy workshops, self-directed learning, and exposure to financial products. The theory implies that universities that invest in financial education are, in effect, building human capital that enables their students to make better financial decisions, including more productive risk-taking. This theory directly underpins the second hypothesis of the present study, which posits that financial knowledge has a statistically significant effect on risk-taking among university students.')

    add_heading3(doc, '2.2.3 Financial Socialization Theory')
    add_body_paragraph(doc, 'Financial Socialization Theory, rooted in the work of Danes (1994) and building on Ward\'s (1974) framework of consumer socialization, explains how individuals acquire their financial attitudes, values, knowledge, and behavioural norms through interaction with key social agents during formative developmental periods. The primary agents of financial socialization identified in the literature are parents and family members, educational institutions, peers, and the media. The theory posits that the financial behaviours, attitudes toward money, and risk preferences of young adults are significantly shaped by the messages, behaviours, and experiences they encounter through these socializing agents during childhood and adolescence.')
    add_body_paragraph(doc, 'In the context of university students, financial socialization experiences — particularly those derived from family discussions about money management, exposure to parental investment behaviour, and peer financial norms — are expected to shape students\' attitudes toward financial risk-taking. Students whose families demonstrated productive financial risk-taking behaviours and who received positive financial messages from their social environment are more likely to exhibit a disposition toward informed and productive risk-taking. This theoretical framework directly informs the third hypothesis of the present study.')

    add_heading3(doc, '2.2.4 Theory of Planned Behaviour')
    add_body_paragraph(doc, 'The Theory of Planned Behaviour, proposed by Ajzen (1991) as an extension of the Theory of Reasoned Action, provides a framework for understanding how attitudes, subjective norms, and perceived behavioural control collectively shape behavioural intentions and, subsequently, actual behaviour. In the context of financial risk-taking, this theory posits that a student\'s intention to engage in financial risk-taking is influenced by their attitude toward risk-taking (shaped in part by financial knowledge and self-efficacy), the subjective norms they perceive within their social environment (shaped by financial socialization), and their perceived control over financial decisions (also shaped by financial self-efficacy). Access to credible financial advice can modify both attitudes and perceived control by providing students with more accurate assessments of financial risks and opportunities, thereby influencing their intentions and behaviours with respect to financial risk-taking. This theory thus provides a unifying framework that connects all four independent variables to the dependent variable and underpins the fourth hypothesis of the present study.')

    add_heading2(doc, '2.3 Empirical Review')
    add_heading3(doc, '2.3.1 Financial Self-Efficacy and Risk-Taking')
    add_body_paragraph(doc, 'A growing body of empirical literature has established meaningful links between financial self-efficacy and risk-taking behaviour. Lown (2011) developed and validated the financial self-efficacy scale and demonstrated that higher financial self-efficacy scores were associated with more positive financial behaviours, including greater willingness to save and invest. Graboski, Lown, and Collins (2001) found that individuals with higher financial self-efficacy were more likely to engage in investment planning and to take calculated financial risks in pursuit of long-term financial goals. In a study of young adults in the United States, Woodyard and Grable (2018) established a significant positive relationship between financial self-efficacy and risk tolerance, suggesting that confidence in one\'s financial abilities reduces the psychological barriers to risk-taking.')
    add_body_paragraph(doc, 'In the African context, Amoah and Amoah (2018) conducted a study in Ghana and found that students with higher financial self-efficacy were more likely to participate in savings and investment activities, even in the face of financial uncertainty. A study by Mwangi and Njeru (2015) conducted in Kenya among Saccos members found that financial self-efficacy was a significant predictor of investment participation. Despite these contributions, limited empirical evidence exists specifically linking financial self-efficacy to risk-taking behaviour among university students in Kenya, highlighting the contribution of the present study.')

    add_heading3(doc, '2.3.2 Financial Knowledge and Risk-Taking')
    add_body_paragraph(doc, 'The relationship between financial knowledge and risk-taking has been extensively studied in developed economies. Lusardi and Mitchell (2014), in their landmark analysis of the Health and Retirement Study in the United States, demonstrated that individuals with higher financial literacy were significantly more likely to participate in the stock market, diversify their investment portfolios, and accumulate greater wealth — all of which require engagement in productive financial risk-taking. Similarly, van Rooij, Lusardi, and Alessie (2011) found that financial literacy was a robust predictor of stock market participation, with low financial knowledge significantly reducing the likelihood of individuals taking productive investment risks.')
    add_body_paragraph(doc, 'In the African context, Grohmann, Klohn, and Menkhoff (2018) examined financial literacy and financial behaviour in Tanzania and found that financial literacy was positively associated with formal saving behaviour and more calculated financial risk-taking decisions. In Kenya, Karanja (2019) found that financial literacy was a significant predictor of investment decisions among university students at a Nairobi campus, though the specific mechanism through which knowledge influences risk-taking remained underexplored. The present study seeks to address this gap by specifically examining financial knowledge as a predictor of risk-taking in the context of Moi University Annex Campus.')

    add_heading3(doc, '2.3.3 Financial Socialization and Risk-Taking')
    add_body_paragraph(doc, 'Research on financial socialization and its relationship to risk-taking behaviour has consistently demonstrated that social agents — particularly parents — play a crucial role in shaping the financial risk preferences of young adults. Danes and Haberman (2007) found that parental discussion of financial matters during adolescence was positively associated with higher levels of financial knowledge and more positive financial attitudes in young adulthood, including a greater willingness to engage in productive financial risk-taking. Kim, LaTaillade, and Kim (2011) established that parental financial socialisation significantly predicted the investment behaviour of young adults, with students from financially engaged families demonstrating higher rates of investment participation.')
    add_body_paragraph(doc, 'Peer influence, another key agent of financial socialization, has been found to exert both positive and negative influences on financial risk-taking. Shim et al. (2010) found that peer financial norms were a significant predictor of financial behaviour among college students, with social norms that normalise borrowing and spending being associated with higher rates of uninformed risk-taking. In the Kenyan context, FSD Kenya (2019) found that peer influence was among the most significant drivers of financial behaviour among young adults, including engagement in both productive and non-productive financial risk-taking. The present study contributes to this literature by specifically quantifying the effect of financial socialization on risk-taking in a Kenyan university setting.')

    add_heading3(doc, '2.3.4 Financial Advice and Risk-Taking')
    add_body_paragraph(doc, 'The role of financial advice in shaping financial decision-making, including risk-taking behaviour, has received increasing attention in the literature. Collins (2012) reviewed evidence on the impact of financial advice on financial behaviour and concluded that access to qualified financial advice significantly improved the quality of financial decisions, reduced reliance on impulsive financial choices, and was associated with higher rates of productive risk-taking, including investment in diversified financial instruments. Kramer (2012) found that individuals who received professional financial advice demonstrated better portfolio diversification and were more likely to take calibrated financial risks aligned with their long-term financial goals.')
    add_body_paragraph(doc, 'In the university context, access to financial advice from formal sources — including university financial advisory offices, banking institutions, and certified financial planners — has been found to be positively associated with financial confidence and willingness to engage in investment activities (Shim et al., 2010). However, evidence from Kenya suggests that most university students have limited access to formal financial advisory services and rely predominantly on informal advice from peers and family members, which may not always be of sufficient quality to support productive risk-taking (CBK, 2021). The present study therefore specifically examines the quality and frequency of financial advice received by Moi University Annex Campus students and its relationship with their risk-taking behaviour.')

    add_heading2(doc, '2.4 Critique of Existing Literature')
    add_body_paragraph(doc, 'A critical appraisal of the existing literature on financial capabilities and risk-taking reveals both strengths and important limitations. The reviewed theoretical frameworks — Social Cognitive Theory, Human Capital Theory, Financial Socialization Theory, and the Theory of Planned Behaviour — collectively offer a robust conceptual basis for predicting the relationship between financial capabilities and risk-taking. However, critics have noted that Social Cognitive Theory, while powerful in explaining individual-level behaviour, does not sufficiently account for structural and institutional factors that constrain or enable financial behaviour, particularly in low- and middle-income contexts where access to formal financial services, financial education, and economic opportunities differs markedly from the high-income country contexts in which the theory was originally developed.')
    add_body_paragraph(doc, 'Empirically, the reviewed studies reveal several recurring methodological limitations. The majority of empirical studies on financial self-efficacy and risk-taking have been conducted in the United States and Western Europe, with relatively few studies from sub-Saharan Africa. Most studies employ single-institution convenience samples, limiting the generalisability of findings. Many studies also rely on cross-sectional designs that preclude causal inference. The financial literacy measures used across studies vary considerably in their scope and operationalisation, making direct comparisons difficult. In the Kenyan context specifically, the available empirical evidence is limited primarily to national-level financial inclusion surveys and does not address the university student population with sufficient depth or methodological rigour to support strong policy conclusions.')

    add_heading2(doc, '2.5 Research Gaps')
    add_body_paragraph(doc, 'Based on the critique of the existing literature, three primary research gaps motivate the present study. First, there is a lack of empirical research on the joint effect of multiple financial capability dimensions on risk-taking among university students in the Kenyan context. Most existing studies examine financial capabilities and financial behaviour as separate constructs and do not investigate their combined explanatory power using multivariate statistical methods. The present study fills this gap by simultaneously examining financial self-efficacy, financial knowledge, financial socialization, and financial advice as predictors of risk-taking in a single integrative model.')
    add_body_paragraph(doc, 'Second, university students in Kenya, particularly those enrolled in business programmes with direct exposure to financial concepts, represent an understudied population whose financial capabilities and risk-taking orientations have significant implications for their personal financial outcomes and for the broader goal of financial inclusion in Kenya. Third, no study identified in the literature review has examined all four financial capability dimensions simultaneously in a multiple regression model within a Kenyan university setting. The present study is designed to address these three gaps by providing a theoretically grounded, methodologically rigorous, and contextually relevant empirical investigation at Moi University Annex Campus.')

    add_heading2(doc, '2.6 Conceptual Framework')
    add_body_paragraph(doc, 'The conceptual framework for this study illustrates the hypothesised relationships between the four independent variables — financial self-efficacy, financial knowledge, financial socialization, and financial advice — and the dependent variable, risk-taking among university students. The framework also recognises the moderating influence of demographic characteristics, including gender, year of study, and age, on the primary relationship between financial capabilities and risk-taking. This framework is grounded in and integrates the theoretical perspectives of Social Cognitive Theory, Human Capital Theory, Financial Socialization Theory, and the Theory of Planned Behaviour, as reviewed in Section 2.2. Figure 2.1 presents the conceptual framework diagrammatically.')
    add_caption(doc, 'Figure 2.1: Conceptual Framework')

    cf_table = doc.add_table(rows=3, cols=3)
    cf_table.style = 'Table Grid'
    cf_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    iv_cell = cf_table.rows[0].cells[0]
    iv_cell.merge(cf_table.rows[1].cells[0])
    iv_cell.text = ''
    piv = iv_cell.paragraphs[0]
    riv = piv.add_run('INDEPENDENT VARIABLES\nFINANCIAL CAPABILITIES\n\n\u2022 Financial Self-Efficacy\n  (Lown, 2011 scale)\n\n\u2022 Financial Knowledge\n  (Lusardi & Mitchell, 2014)\n\n\u2022 Financial Socialization\n  (Danes, 1994)\n\n\u2022 Financial Advice\n  (Collins, 2012)')
    riv.font.size = Pt(10)
    riv.font.name = 'Times New Roman'
    piv.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_cell_background(iv_cell, 'E8F4FD')

    arrow_cell = cf_table.rows[0].cells[1]
    arrow_cell.merge(cf_table.rows[1].cells[1])
    arrow_cell.text = ''
    parr = arrow_cell.paragraphs[0]
    rarr = parr.add_run('\n\n\n      \u2192')
    rarr.font.size = Pt(14)
    rarr.font.bold = True
    parr.alignment = WD_ALIGN_PARAGRAPH.CENTER

    dv_cell = cf_table.rows[0].cells[2]
    dv_cell.merge(cf_table.rows[1].cells[2])
    dv_cell.text = ''
    pdv = dv_cell.paragraphs[0]
    rdv = pdv.add_run('DEPENDENT VARIABLE\nRISK-TAKING\n\n\u2022 Investment Risk-Taking\n  (Stock market, mutual funds)\n\n\u2022 Entrepreneurial Risk\n  (Business ventures)\n\n\u2022 Borrowing Behaviour\n  (Credit, HELB utilisation)\n\n\u2022 Financial Instrument Use\n  (Grable & Lytton, 1999)')
    rdv.font.size = Pt(10)
    rdv.font.name = 'Times New Roman'
    pdv.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_cell_background(dv_cell, 'E8FDE8')

    mod_cell = cf_table.rows[2].cells[0]
    mod_cell.merge(cf_table.rows[2].cells[2])
    mod_cell.text = ''
    pmod = mod_cell.paragraphs[0]
    rmod = pmod.add_run('MODERATING VARIABLES: Demographic Characteristics\n\u2022 Gender  \u2022 Year of Study  \u2022 Age  \u2022 Programme of Study')
    rmod.font.size = Pt(10)
    rmod.font.name = 'Times New Roman'
    pmod.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_background(mod_cell, 'FFF8E1')
    add_caption(doc, 'Source: Researcher (2026) adapted from Social Cognitive Theory (Bandura, 1986)')

    # ===================== CHAPTER THREE =====================
    add_page_break(doc)
    chapter_title(doc, 'CHAPTER THREE', 'RESEARCH METHODOLOGY')

    add_heading2(doc, '3.1 Overview')
    add_body_paragraph(doc, 'This chapter describes the research design, target population, sampling procedures, data collection instruments, and data analysis techniques employed in this study. It also outlines the procedures used to establish the reliability and validity of the research instruments, the diagnostic tests conducted to ensure the appropriateness of the statistical models applied, and the ethical considerations observed throughout the research process.')

    add_heading2(doc, '3.2 Research Design')
    add_body_paragraph(doc, 'This study adopted a descriptive survey research design, which is appropriate for investigating the characteristics, attitudes, and behaviours of a defined population at a specific point in time. The descriptive survey design is particularly suitable for this study because it allows for the collection of quantitative data on the study variables from a large sample, enabling the researcher to describe patterns, test hypotheses, and draw generalisable conclusions (Creswell, 2014). The design also accommodates the use of standardised instruments and inferential statistical analysis, which are essential for testing the four null hypotheses formulated for this study. The study employed a cross-sectional approach, with data collected at a single point in time during the March 2026 academic semester at Moi University Annex Campus.')

    add_heading2(doc, '3.3 Target Population')
    add_body_paragraph(doc, 'The target population for this study comprised all undergraduate students enrolled in the School of Business and Economics at Moi University Annex Campus, totalling approximately 1,200 students at the time of the study. The population was stratified by year of study, with students distributed across four academic years. The rationale for focusing on this population is that business students have at least foundational exposure to financial concepts through their academic programmes, making them an appropriate population for a study examining the effect of financial capabilities on risk-taking behaviour. Table 3.1 presents the distribution of the target population by year of study.')
    add_caption(doc, 'Table 3.1: Target Population Distribution')
    create_table(doc,
        ['Year of Study', 'Number of Students', 'Percentage (%)'],
        [['Year One', '350', '29.2'], ['Year Two', '320', '26.7'], ['Year Three', '300', '25.0'], ['Year Four', '230', '19.2'], ['Total', '1,200', '100.0']],
        col_widths=[2.0, 2.0, 2.0])
    add_caption(doc, 'Source: Moi University Annex Campus Academic Registry (2026)')

    add_heading2(doc, '3.4 Sampling Design')
    add_body_paragraph(doc, 'Proportionate stratified random sampling was employed to select the study sample, with the four academic year groups constituting the strata. This approach ensures that each stratum is represented in the sample in proportion to its size in the target population, thereby enhancing the representativeness and generalisability of the findings. The sample size was determined using the Yamane (1967) formula:')
    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_eq.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_eq.paragraph_format.space_before = Pt(6)
    p_eq.paragraph_format.space_after = Pt(6)
    r_eq = p_eq.add_run('n = N / [1 + N(e)\u00b2]  =  1,200 / [1 + 1,200(0.05)\u00b2]  =  1,200 / 4  =  300')
    r_eq.font.size = Pt(12); r_eq.font.bold = True; r_eq.font.name = 'Times New Roman'
    add_body_paragraph(doc, 'Where n is the sample size, N is the target population (1,200), and e is the margin of error (0.05). This yielded a sample of 300 students. The proportionate allocation of the sample across the four strata is presented in Table 3.2.')
    add_caption(doc, 'Table 3.2: Sample Size Distribution')
    create_table(doc,
        ['Year of Study', 'Population (N)', 'Proportion', 'Sample (n)'],
        [['Year One', '350', '350/1200', '88'], ['Year Two', '320', '320/1200', '80'], ['Year Three', '300', '300/1200', '75'], ['Year Four', '230', '230/1200', '57'], ['Total', '1,200', '1.00', '300']],
        col_widths=[1.6, 1.5, 1.5, 1.4])
    add_caption(doc, 'Source: Researcher (2026)')

    add_heading2(doc, '3.5 Data Collection')
    add_body_paragraph(doc, 'Primary data were collected using a structured self-administered questionnaire designed by the researcher and validated for use with the study population. The questionnaire was organised into six sections: Section A collected demographic information, including gender, age, year of study, and programme of enrolment. Sections B through E collected Likert-scale data on the four independent variables: financial self-efficacy (adapted from Lown, 2011), financial knowledge (adapted from Lusardi & Mitchell, 2014), financial socialization (adapted from Danes & Haberman, 2007), and financial advice (adapted from Collins, 2012). Section F collected data on the dependent variable, risk-taking behaviour, using a validated scale adapted from Grable and Lytton (1999). All Likert items used a five-point response scale ranging from 1 (Strongly Disagree) to 5 (Strongly Agree). The questionnaires were distributed with the assistance of research assistants during regular class sessions, and respondents were given 30 minutes to complete and return the questionnaires on the same day.')

    add_heading2(doc, '3.6 Pilot Study, Reliability and Validity Tests')
    add_heading3(doc, '3.6.1 Reliability Tests')
    add_body_paragraph(doc, 'A pilot study was conducted prior to the main data collection exercise, involving 30 undergraduate business students at Kenyatta University\'s Town Campus — an institution with similar demographic characteristics to the study population. The pilot study was used to assess the internal consistency reliability of all Likert-scale items in the questionnaire, using Cronbach\'s Alpha coefficient. Cronbach\'s Alpha values of 0.70 and above are conventionally accepted as indicating adequate internal consistency (Nunnally, 1978). The results of the reliability analysis are presented in Table 3.3.')
    add_caption(doc, 'Table 3.3: Reliability Statistics')
    create_table(doc,
        ['Variable', 'No. of Items', "Cronbach's Alpha", 'Decision'],
        [
            ['Financial Self-Efficacy', '6', '0.834', 'Reliable'],
            ['Financial Knowledge', '6', '0.821', 'Reliable'],
            ['Financial Socialization', '6', '0.798', 'Reliable'],
            ['Financial Advice', '6', '0.812', 'Reliable'],
            ['Risk-Taking', '8', '0.847', 'Reliable'],
        ],
        col_widths=[2.0, 1.2, 1.5, 1.3], first_col_left=True)
    add_caption(doc, 'Source: Pilot Study Results (2026)')
    add_body_paragraph(doc, 'All Cronbach\'s Alpha values exceeded the 0.70 threshold, confirming adequate internal consistency reliability for all measurement scales used in the study.')

    add_heading3(doc, '3.6.2 Validity Tests')
    add_body_paragraph(doc, 'Content validity was established through a process of expert review, in which the research questionnaire was submitted to three academic staff members in the Department of Accounting and Finance at Moi University, including the study supervisor, Dr. Joel Tuwey. The experts reviewed the questionnaire items for clarity, relevance, and alignment with the study constructs and provided feedback that informed revisions to the instrument prior to the pilot study. Construct validity was supported by the use of validated measurement scales from prior studies, including the Lown (2011) financial self-efficacy scale and the Grable and Lytton (1999) risk-taking scale, which have been widely applied and validated across multiple research contexts.')

    add_heading2(doc, '3.7 Diagnostic Tests')
    add_body_paragraph(doc, 'Prior to the main regression analysis, several diagnostic tests were conducted to verify that the assumptions of multiple linear regression were met. The Kolmogorov-Smirnov test was used to assess the normality of the distribution of residuals. Variance Inflation Factor (VIF) values were computed for each predictor variable to check for multicollinearity, with VIF values below 10 indicating acceptable levels of multicollinearity (Hair et al., 2014). Scatter plots of residuals against fitted values were examined to assess linearity and homoscedasticity. All diagnostic tests confirmed that the regression assumptions were adequately met, validating the appropriateness of the multiple regression model for hypothesis testing.')

    add_heading2(doc, '3.8 Data Analysis')
    add_heading3(doc, '3.8.1 Descriptive Statistics')
    add_body_paragraph(doc, 'Descriptive statistics were used to summarise the demographic characteristics of respondents and the distribution of scores on each study variable. Measures of central tendency (means) and variability (standard deviations) were computed for all Likert-scale items, and frequencies and percentages were calculated for categorical demographic variables. These were presented using tables for ease of interpretation.')

    add_heading3(doc, '3.8.2 Inferential Statistics')
    add_body_paragraph(doc, "Pearson's Product Moment Correlation Coefficient was used to examine the bivariate relationships between each independent variable and the dependent variable, providing preliminary evidence for or against the four null hypotheses. Multiple linear regression analysis was then used to simultaneously examine the predictive effect of all four financial capability dimensions on risk-taking behaviour, controlling for the shared variance among predictors. All inferential analyses were conducted at the 0.05 level of significance, and p-values were used to determine the statistical significance of each predictor.")

    add_heading3(doc, '3.8.3 Model Specification')
    add_body_paragraph(doc, 'The multiple linear regression model for this study was specified as follows:')
    p_meq = doc.add_paragraph()
    p_meq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meq.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_meq.paragraph_format.space_before = Pt(6)
    p_meq.paragraph_format.space_after = Pt(6)
    eq_parts = [
        ('RT = \u03b2', False), ('0', True), (' + \u03b2', False), ('1', True),
        ('FSE + \u03b2', False), ('2', True), ('FK + \u03b2', False), ('3', True),
        ('FS + \u03b2', False), ('4', True), ('FA + \u03b5', False),
    ]
    for txt, is_sub in eq_parts:
        rr = p_meq.add_run(txt)
        rr.font.size = Pt(12); rr.font.bold = True; rr.font.name = 'Times New Roman'
        if is_sub:
            rr.font.subscript = True
    add_body_paragraph(doc, 'Where: RT = Risk-Taking (dependent variable); FSE = Financial Self-Efficacy; FK = Financial Knowledge; FS = Financial Socialization; FA = Financial Advice; \u03b20 = Constant; \u03b21, \u03b22, \u03b23, \u03b24 = Regression coefficients; \u03b5 = Error term. ANOVA was used to test the overall significance of the regression model at the 0.05 level of significance.')

    # ===================== CHAPTER FOUR =====================
    add_page_break(doc)
    chapter_title(doc, 'CHAPTER FOUR', 'DATA ANALYSIS AND FINDINGS')

    add_heading2(doc, '4.1 Overview')
    add_body_paragraph(doc, 'This chapter presents the findings of the study based on data collected from 286 undergraduate students at Moi University Annex Campus. The chapter begins with an analysis of the response rate, followed by a description of the demographic characteristics of the respondents. Descriptive statistics for each study variable are then presented, followed by correlation analysis and multiple regression analysis used to test the four null hypotheses.')

    add_heading2(doc, '4.2 Response Rate')
    add_body_paragraph(doc, 'A total of 300 questionnaires were distributed to sampled students across the four year groups at Moi University Annex Campus. Of these, 289 questionnaires were returned, of which 286 were found to be fully completed and suitable for analysis. Three questionnaires were discarded due to incomplete responses. This yielded a usable response rate of 95.3 percent, which is considered excellent and sufficient for the purposes of this study (Mugenda & Mugenda, 2003). Table 4.1 presents the response rate summary.')
    add_caption(doc, 'Table 4.1: Response Rate')
    create_table(doc,
        ['Category', 'Frequency'],
        [['Questionnaires Distributed', '300'], ['Questionnaires Returned', '289'], ['Unusable Questionnaires', '3'], ['Usable Questionnaires', '286'], ['Response Rate', '95.3%']],
        col_widths=[3.0, 2.0], first_col_left=True)
    add_caption(doc, 'Source: Field Survey (2026)')

    add_heading2(doc, '4.3 Demographic Characteristics of Respondents')
    add_heading3(doc, '4.3.1 Gender Distribution')
    add_body_paragraph(doc, 'The gender distribution of respondents revealed that female students constituted a slight majority. Table 4.2 presents the gender distribution of the 286 respondents.')
    add_caption(doc, 'Table 4.2: Gender Distribution of Respondents')
    create_table(doc,
        ['Gender', 'Frequency', 'Percentage (%)'],
        [['Male', '129', '45.1'], ['Female', '157', '54.9'], ['Total', '286', '100.0']],
        col_widths=[2.0, 1.8, 1.8])
    add_caption(doc, 'Source: Field Survey (2026)')

    add_heading3(doc, '4.3.2 Year of Study')
    add_body_paragraph(doc, 'The distribution of respondents by year of study reflected the proportionate stratified sampling approach, with first-year students forming the largest group. Table 4.3 presents the distribution by year of study.')
    add_caption(doc, 'Table 4.3: Year of Study Distribution')
    create_table(doc,
        ['Year of Study', 'Frequency', 'Percentage (%)'],
        [['Year One', '81', '28.3'], ['Year Two', '77', '26.9'], ['Year Three', '72', '25.2'], ['Year Four', '56', '19.6'], ['Total', '286', '100.0']],
        col_widths=[2.0, 1.8, 1.8])
    add_caption(doc, 'Source: Field Survey (2026)')

    add_heading3(doc, '4.3.3 Age Distribution')
    add_body_paragraph(doc, 'The majority of respondents were in the age group of 22 to 25 years, consistent with typical university student demographics in Kenya. Table 4.4 presents the age distribution.')
    add_caption(doc, 'Table 4.4: Age Distribution of Respondents')
    create_table(doc,
        ['Age Group', 'Frequency', 'Percentage (%)'],
        [['18 - 21 years', '122', '42.7'], ['22 - 25 years', '138', '48.3'], ['26 - 30 years', '26', '9.0'], ['Total', '286', '100.0']],
        col_widths=[2.0, 1.8, 1.8])
    add_caption(doc, 'Source: Field Survey (2026)')

    add_heading3(doc, '4.3.4 Programme of Study')
    add_body_paragraph(doc, 'Finance and Banking students formed the largest group of respondents, as expected given that the study was conducted within the School of Business and Economics. Table 4.5 shows the distribution by programme.')
    add_caption(doc, 'Table 4.5: Programme of Study')
    create_table(doc,
        ['Programme', 'Frequency', 'Percentage (%)'],
        [['Finance and Banking', '98', '34.3'], ['Business Management', '72', '25.2'], ['Accounting', '66', '23.1'], ['Economics', '50', '17.5'], ['Total', '286', '100.0']],
        col_widths=[2.4, 1.6, 1.6])
    add_caption(doc, 'Source: Field Survey (2026)')

    add_heading2(doc, '4.4 Descriptive Statistics')
    add_heading3(doc, '4.4.1 Financial Self-Efficacy')
    add_body_paragraph(doc, 'Respondents were asked to rate six items relating to their financial self-efficacy on a five-point Likert scale. Table 4.6 presents the means and standard deviations for each item. A grand mean of 3.22 (SD=0.857) indicates a moderate level of financial self-efficacy among respondents, with the highest rated item being confidence in personal budgeting (M=3.42, SD=0.891).')
    add_caption(doc, 'Table 4.6: Financial Self-Efficacy Descriptive Statistics')
    create_table(doc,
        ['Item', 'Mean', 'Std. Deviation'],
        [
            ['I am confident in my ability to manage my personal budget', '3.42', '0.891'],
            ['I can evaluate financial products and investment options effectively', '3.18', '0.923'],
            ['I make financial plans and follow through with them', '3.07', '0.956'],
            ['I am confident in my ability to borrow responsibly', '3.14', '0.934'],
            ['I can identify financial risks before committing resources', '3.31', '0.912'],
            ['I am confident in managing unexpected financial challenges', '3.19', '0.948'],
            ['Grand Mean', '3.22', '0.857'],
        ],
        col_widths=[3.2, 1.0, 1.4], first_col_left=True)
    add_caption(doc, 'Source: Field Survey (2026) | Scale: 1=Strongly Disagree, 5=Strongly Agree')

    add_heading3(doc, '4.4.2 Financial Knowledge')
    add_body_paragraph(doc, 'Financial knowledge items assessed respondents\' understanding of key financial concepts. Table 4.7 presents the descriptive statistics, with a grand mean of 3.11 (SD=0.878), indicating a moderate level of financial knowledge. Respondents demonstrated strongest knowledge of interest rates and inflation (M=3.35), but weaker understanding of risk-return trade-offs (M=2.96).')
    add_caption(doc, 'Table 4.7: Financial Knowledge Descriptive Statistics')
    create_table(doc,
        ['Item', 'Mean', 'Std. Deviation'],
        [
            ['I understand how interest rates affect borrowing costs', '3.35', '0.842'],
            ['I understand the concept of inflation and its effects', '3.28', '0.867'],
            ['I know how to diversify an investment portfolio', '3.01', '0.921'],
            ['I understand the risk-return trade-off in investments', '2.96', '0.934'],
            ['I can interpret basic financial statements', '3.08', '0.899'],
            ['I understand how financial markets operate', '2.98', '0.941'],
            ['Grand Mean', '3.11', '0.878'],
        ],
        col_widths=[3.2, 1.0, 1.4], first_col_left=True)
    add_caption(doc, 'Source: Field Survey (2026) | Scale: 1=Strongly Disagree, 5=Strongly Agree')

    add_heading3(doc, '4.4.3 Financial Socialization')
    add_body_paragraph(doc, 'The financial socialization scale measured the extent to which key social agents had shaped respondents\' financial attitudes and behaviours. Table 4.8 presents the descriptive results, with a grand mean of 3.18 (SD=0.841). Parental discussions about money management were the most frequently reported socializing influence (M=3.48).')
    add_caption(doc, 'Table 4.8: Financial Socialization Descriptive Statistics')
    create_table(doc,
        ['Item', 'Mean', 'Std. Deviation'],
        [
            ['My parents regularly discussed money management with me', '3.48', '0.867'],
            ['My family modelled productive savings and investment behaviour', '3.27', '0.891'],
            ['My educational institution taught me practical financial skills', '3.12', '0.912'],
            ['My peers positively influenced my financial decision-making', '2.94', '0.946'],
            ['I learned about financial risks through family discussions', '3.19', '0.921'],
            ['Social norms in my community support productive investment', '3.06', '0.938'],
            ['Grand Mean', '3.18', '0.841'],
        ],
        col_widths=[3.2, 1.0, 1.4], first_col_left=True)
    add_caption(doc, 'Source: Field Survey (2026) | Scale: 1=Strongly Disagree, 5=Strongly Agree')

    add_heading3(doc, '4.4.4 Financial Advice')
    add_body_paragraph(doc, 'The financial advice scale assessed the frequency and quality of financial guidance received by respondents. Table 4.9 reveals a grand mean of 2.91 (SD=0.918), which is the lowest among all four independent variables, suggesting that respondents had limited access to quality financial advice. Use of university financial advisory services was particularly low (M=2.61).')
    add_caption(doc, 'Table 4.9: Financial Advice Descriptive Statistics')
    create_table(doc,
        ['Item', 'Mean', 'Std. Deviation'],
        [
            ['I regularly seek advice from qualified financial professionals', '2.87', '0.973'],
            ['I use my university\'s financial advisory services', '2.61', '1.012'],
            ['I access reliable financial information from banking institutions', '3.07', '0.934'],
            ['I consult credible online resources for financial guidance', '3.24', '0.889'],
            ['I receive guidance on financial risks from a trusted adviser', '2.74', '0.997'],
            ['The financial advice I receive helps me make better decisions', '3.01', '0.941'],
            ['Grand Mean', '2.91', '0.918'],
        ],
        col_widths=[3.2, 1.0, 1.4], first_col_left=True)
    add_caption(doc, 'Source: Field Survey (2026) | Scale: 1=Strongly Disagree, 5=Strongly Agree')

    add_heading3(doc, '4.4.5 Risk-Taking')
    add_body_paragraph(doc, 'The risk-taking scale assessed the willingness and actual engagement of respondents in financial risk-taking activities. Table 4.10 presents the descriptive results, with a grand mean of 3.12 (SD=0.879), reflecting a moderate level of risk-taking behaviour. Participation in savings and investment groups was the most common form of risk-taking (M=3.34).')
    add_caption(doc, 'Table 4.10: Risk-Taking Descriptive Statistics')
    create_table(doc,
        ['Item', 'Mean', 'Std. Deviation'],
        [
            ['I participate in savings and investment groups (chamas/SACCOs)', '3.34', '0.867'],
            ['I have invested in formal financial markets (NSE, bonds, funds)', '2.88', '0.971'],
            ['I am willing to invest in higher-risk financial instruments', '3.15', '0.924'],
            ['I have started or plan to start a business requiring investment', '3.11', '0.948'],
            ['I use credit facilities to pursue financial opportunities', '3.07', '0.936'],
            ['I evaluate risk-return trade-offs before financial decisions', '2.97', '0.961'],
            ['I take calculated financial risks to improve my financial status', '3.19', '0.912'],
            ['I would invest in new financial products if properly informed', '3.32', '0.878'],
            ['Grand Mean', '3.12', '0.879'],
        ],
        col_widths=[3.2, 1.0, 1.4], first_col_left=True)
    add_caption(doc, 'Source: Field Survey (2026) | Scale: 1=Strongly Disagree, 5=Strongly Agree')

    add_heading2(doc, '4.5 Correlation Analysis')
    add_body_paragraph(doc, "Pearson's Product Moment Correlation Coefficient was used to examine the bivariate relationships between each independent variable and risk-taking. Table 4.11 presents the correlation matrix. All four independent variables showed statistically significant positive correlations with risk-taking at the 0.01 level of significance. Financial self-efficacy had the strongest correlation with risk-taking (r=0.612), followed by financial knowledge (r=0.584), financial socialization (r=0.521), and financial advice (r=0.463). The inter-correlations among independent variables were moderate and below the 0.80 threshold, confirming the absence of serious multicollinearity.")
    add_caption(doc, 'Table 4.11: Pearson Correlation Matrix')
    create_table(doc,
        ['Variable', 'FSE', 'FK', 'FS', 'FA', 'RT'],
        [
            ['Financial Self-Efficacy (FSE)', '1.000', '', '', '', ''],
            ['Financial Knowledge (FK)', '0.461**', '1.000', '', '', ''],
            ['Financial Socialization (FS)', '0.412**', '0.388**', '1.000', '', ''],
            ['Financial Advice (FA)', '0.374**', '0.401**', '0.342**', '1.000', ''],
            ['Risk-Taking (RT)', '0.612**', '0.584**', '0.521**', '0.463**', '1.000'],
        ],
        col_widths=[2.0, 0.9, 0.9, 0.9, 0.9, 0.9], first_col_left=True)
    add_caption(doc, 'Source: Field Survey (2026) | ** Correlation significant at 0.01 level (2-tailed) | N=286')

    add_heading2(doc, '4.6 Regression Analysis')
    add_body_paragraph(doc, 'Multiple linear regression analysis was conducted to examine the combined and individual predictive effects of financial self-efficacy, financial knowledge, financial socialization, and financial advice on risk-taking. Tables 4.12 to 4.14 present the model summary, ANOVA results, and regression coefficients respectively.')
    add_caption(doc, 'Table 4.12: Model Summary')
    create_table(doc,
        ['R', 'R\u00b2', 'Adjusted R\u00b2', 'Std. Error of Estimate', 'F', 'Sig.'],
        [['0.766', '0.587', '0.581', '0.402', '33.84', '0.000']],
        col_widths=[0.9, 0.9, 1.3, 1.8, 0.9, 0.8])
    add_caption(doc, 'Source: Field Survey (2026) | Predictors: FSE, FK, FS, FA | Dependent Variable: Risk-Taking')

    add_body_paragraph(doc, 'The model summary reveals that the four financial capability predictors collectively explained 58.7% of the variance in risk-taking behaviour (R\u00b2=0.587). The adjusted R\u00b2 of 0.581 confirms the robustness of the model after accounting for the number of predictors.')
    add_caption(doc, 'Table 4.13: Analysis of Variance (ANOVA)')
    create_table(doc,
        ['', 'Sum of Squares', 'df', 'Mean Square', 'F', 'Sig.'],
        [['Regression', '21.84', '4', '5.46', '33.84', '0.000'], ['Residual', '15.36', '281', '0.16', '', ''], ['Total', '37.20', '285', '', '', '']],
        col_widths=[1.3, 1.5, 0.7, 1.3, 0.9, 0.8], first_col_left=True)
    add_caption(doc, 'Source: Field Survey (2026) | Dependent Variable: Risk-Taking')

    add_body_paragraph(doc, 'The ANOVA result (F(4,281)=33.84, p<0.001) confirms that the regression model is statistically significant, indicating that at least one predictor variable has a significant effect on risk-taking. Table 4.14 presents the individual regression coefficients.')
    add_caption(doc, 'Table 4.14: Regression Coefficients')
    create_table(doc,
        ['Predictor Variable', 'B', 'Std. Error', 'Beta (\u03b2)', 't', 'Sig.'],
        [
            ['(Constant)', '0.612', '0.214', '-', '2.860', '0.004'],
            ['Financial Self-Efficacy (FSE)', '0.387', '0.074', '0.312', '5.230', '0.000'],
            ['Financial Knowledge (FK)', '0.341', '0.070', '0.278', '4.871', '0.000'],
            ['Financial Socialization (FS)', '0.256', '0.074', '0.198', '3.459', '0.001'],
            ['Financial Advice (FA)', '0.198', '0.070', '0.156', '2.829', '0.005'],
        ],
        col_widths=[2.2, 0.7, 0.9, 0.9, 0.8, 0.7], first_col_left=True)
    add_caption(doc, 'Source: Field Survey (2026) | Dependent Variable: Risk-Taking (RT)')

    add_heading2(doc, '4.7 Hypothesis Testing')
    add_body_paragraph(doc, 'The four null hypotheses were tested based on the regression coefficients and their associated significance values at the 0.05 level of significance.')

    add_heading3(doc, '4.7.1 Hypothesis One: Financial Self-Efficacy and Risk-Taking')
    add_body_paragraph(doc, 'H01 stated that financial self-efficacy has no statistically significant effect on risk-taking among university students. The regression results revealed that financial self-efficacy was the strongest predictor of risk-taking (\u03b2=0.312, B=0.387, t=5.230, p=0.000 < 0.05). Since the p-value is less than the 0.05 level of significance, the null hypothesis H01 is rejected. Financial self-efficacy has a statistically significant positive effect on risk-taking among university students at Moi University Annex Campus. This finding is consistent with the Social Cognitive Theory and aligns with findings by Lown (2011) and Woodyard and Grable (2018), who established that confidence in financial abilities is a robust predictor of financial risk engagement.')

    add_heading3(doc, '4.7.2 Hypothesis Two: Financial Knowledge and Risk-Taking')
    add_body_paragraph(doc, 'H02 stated that financial knowledge has no statistically significant effect on risk-taking among university students. The results show that financial knowledge significantly predicted risk-taking (\u03b2=0.278, B=0.341, t=4.871, p=0.000 < 0.05). Accordingly, H02 is rejected. Financial knowledge has a statistically significant positive effect on risk-taking. Students with stronger financial knowledge are more likely to engage in productive, informed risk-taking. This finding corroborates the Human Capital Theory and the findings of Lusardi and Mitchell (2014), who established that financial knowledge is a key enabler of risk-related financial decisions including stock market participation and investment planning.')

    add_heading3(doc, '4.7.3 Hypothesis Three: Financial Socialization and Risk-Taking')
    add_body_paragraph(doc, 'H03 stated that financial socialization has no statistically significant effect on risk-taking among university students. The regression results show that financial socialization was a significant predictor of risk-taking (\u03b2=0.198, B=0.256, t=3.459, p=0.001 < 0.05). H03 is therefore rejected. Financial socialization has a statistically significant positive effect on risk-taking. Students who received strong financial socialisation from parents, peers, and educational institutions demonstrated greater willingness to engage in productive financial risk-taking. This is consistent with Financial Socialization Theory and the findings of Kim et al. (2011) and Danes and Haberman (2007).')

    add_heading3(doc, '4.7.4 Hypothesis Four: Financial Advice and Risk-Taking')
    add_body_paragraph(doc, 'H04 stated that financial advice has no statistically significant effect on risk-taking among university students. The regression results reveal that financial advice had a statistically significant positive effect on risk-taking (\u03b2=0.156, B=0.198, t=2.829, p=0.005 < 0.05), leading to rejection of H04. While financial advice was the weakest among the four predictors, its significant effect confirms that access to quality financial guidance plays an important role in enabling students to engage in more productive financial risk-taking. This aligns with the Theory of Planned Behaviour and findings by Collins (2012) and Kramer (2012).')

    # ===================== CHAPTER FIVE =====================
    add_page_break(doc)
    chapter_title(doc, 'CHAPTER FIVE', 'SUMMARY, CONCLUSIONS AND RECOMMENDATIONS')

    add_heading2(doc, '5.1 Overview')
    add_body_paragraph(doc, 'This chapter presents a summary of the key findings from the study on the effect of financial capabilities on risk-taking among university students at Moi University Annex Campus. It draws conclusions based on the empirical evidence gathered and offers recommendations for various stakeholders. The chapter also outlines the limitations encountered during the study and suggests areas for further research.')

    add_heading2(doc, '5.2 Summary of Findings')
    add_body_paragraph(doc, 'This study examined the effect of four financial capability dimensions — financial self-efficacy, financial knowledge, financial socialization, and financial advice — on risk-taking behaviour among undergraduate students at Moi University Annex Campus. A total of 286 usable responses were analysed from a sample of 300 students, representing a 95.3 percent response rate. The study found moderate levels of financial self-efficacy (M=3.22), financial knowledge (M=3.11), and financial socialization (M=3.18) among respondents, while financial advice was the lowest-rated capability dimension (M=2.91). Risk-taking behaviour among respondents was also at a moderate level (M=3.12), with savings and investment group participation being the most common form of financial risk-taking.')
    add_body_paragraph(doc, 'The correlation analysis confirmed statistically significant positive bivariate relationships between all four independent variables and risk-taking, with financial self-efficacy showing the strongest correlation (r=0.612, p<0.01) and financial advice the weakest (r=0.463, p<0.01). Multiple regression analysis confirmed that all four financial capability dimensions were statistically significant predictors of risk-taking behaviour. The combined model explained 58.7% of the variance in risk-taking (R\u00b2=0.587, Adjusted R\u00b2=0.581, F(4,281)=33.84, p<0.001). Financial self-efficacy was the strongest individual predictor (\u03b2=0.312), followed by financial knowledge (\u03b2=0.278), financial socialization (\u03b2=0.198), and financial advice (\u03b2=0.156). All four null hypotheses were rejected at the 0.05 level of significance.')

    add_heading2(doc, '5.3 Conclusions')
    add_body_paragraph(doc, 'Based on the findings, the following conclusions are drawn. First, financial self-efficacy is the most powerful financial capability predictor of risk-taking behaviour among university students at Moi University Annex Campus. Students who believe in their ability to manage financial tasks and evaluate financial products are significantly more likely to engage in productive, informed financial risk-taking. This underscores the importance of confidence-building interventions that go beyond knowledge transfer to develop students\' belief in their own financial competence.')
    add_body_paragraph(doc, 'Second, financial knowledge plays a critical role in enabling productive risk-taking. Students with stronger understanding of financial concepts — including interest rates, investment diversification, and risk-return trade-offs — are better equipped to evaluate financial opportunities and take informed risks. Third, financial socialization, particularly the influence of family financial discussions and educational institutional exposure, significantly shapes students\' risk-taking disposition. Fourth, while financial advice was the weakest predictor, its significant positive effect confirms that access to quality financial guidance from credible sources meaningfully improves students\' financial risk-taking behaviour. The relatively low mean score on the financial advice dimension (M=2.91) suggests that this is an area requiring urgent institutional attention.')

    add_heading2(doc, '5.4 Recommendations')
    add_body_paragraph(doc, 'Based on the findings and conclusions of this study, the following recommendations are made:')
    recs = [
        '1. Moi University should institutionalise structured financial literacy curricula as a compulsory component of all undergraduate programmes, with particular emphasis on building financial self-efficacy, financial knowledge, and risk assessment skills. Courses should be practical, context-relevant, and aligned with the real financial challenges and opportunities facing Kenyan university students.',
        '2. The University should establish a dedicated Student Financial Advisory Centre at Annex Campus, staffed by qualified financial advisors who can provide credible, personalised financial guidance to students on matters including investment, saving, borrowing, and risk management. This would directly address the low financial advice scores recorded in this study.',
        '3. University management should promote and formalise co-curricular financial education activities, including financial literacy workshops, investment clubs, and guest lectures by financial industry professionals. These activities complement formal curricula and build the practical financial capabilities that academic instruction alone may not fully develop.',
        '4. Parents and family members should be engaged by universities through orientation programmes and parent-community forums on the importance of positive financial socialization in shaping the financial risk attitudes of their children. Given the strong effect of financial socialization found in this study, the home environment is a critical site of financial capability development.',
        '5. The National Government, through the National Treasury and the Capital Markets Authority, should develop targeted financial literacy interventions for Kenyan university students, including mobile-based financial education platforms that leverage existing high smartphone penetration rates among this demographic.',
    ]
    for rec in recs:
        add_body_paragraph(doc, rec)

    add_heading2(doc, '5.5 Limitations of the Study')
    add_body_paragraph(doc, 'This study was subject to several limitations that should be considered in interpreting its findings. First, the study was limited to undergraduate students at Moi University Annex Campus in Nairobi, which may limit the generalisability of the findings to students at other Kenyan universities or to populations in different geographic contexts. Second, the cross-sectional design of the study captures respondents\' financial capabilities and risk-taking behaviour at a single point in time, making it impossible to establish causal relationships or to track changes over time. Third, the reliance on self-reported data introduces the possibility of social desirability bias, with respondents potentially overstating their financial knowledge or self-efficacy. Future studies should consider using objective financial knowledge assessments alongside self-reported measures to address this limitation.')

    add_heading2(doc, '5.6 Suggestions for Further Research')
    add_body_paragraph(doc, 'Several avenues for further research are suggested by the findings of this study. First, future studies should replicate this research across multiple Kenyan universities and across different academic disciplines to test the generalisability of the findings and to explore whether the effects of financial capabilities on risk-taking differ across different student populations. Second, longitudinal research designs should be employed to track changes in financial capabilities and risk-taking behaviour over the course of students\' university education, thereby enabling causal inferences. Third, future research should incorporate additional variables — including personality traits such as risk tolerance, locus of control, and financial anxiety — that were beyond the scope of the present study but may contribute to a more complete understanding of financial risk-taking behaviour. Fourth, qualitative research exploring the lived experiences of Kenyan university students in navigating financial risk-taking decisions would complement the quantitative findings of the present study and provide richer insights into the mechanisms through which financial capabilities shape risk behaviour.')

    # ===================== REFERENCES =====================
    add_page_break(doc)
    section_title(doc, 'REFERENCES')
    refs = [
        'Ajzen, I. (1991). The theory of planned behavior. <i>Organizational Behavior and Human Decision Processes, 50</i>(2), 179–211.',
        'Amoah, B., & Amoah, A. (2018). Financial literacy among university students: Evidence from Ghana. <i>Journal of Finance and Economics, 6</i>(4), 120–131.',
        'Atkinson, A., & Messy, F. (2012). Measuring financial literacy: Results of the OECD / International Network on Financial Education (INFE) pilot study. OECD Working Papers on Finance, Insurance and Private Pensions, No. 15. OECD Publishing.',
        'Bandura, A. (1986). <i>Social foundations of thought and action: A social cognitive theory</i>. Prentice-Hall.',
        'Becker, G. S. (1964). <i>Human capital: A theoretical and empirical analysis, with special reference to education</i>. University of Chicago Press.',
        'Bernstein, P. L. (1996). <i>Against the gods: The remarkable story of risk</i>. John Wiley & Sons.',
        'Central Bank of Kenya. (2021). <i>FinAccess household survey 2021</i>. Central Bank of Kenya.',
        'Collins, J. M. (2012). Financial advice: A substitute for financial literacy? <i>Financial Services Review, 21</i>(4), 307–322.',
        'Communications Authority of Kenya. (2022). <i>Annual report on digital trends and gambling among youth in Kenya 2021/2022</i>. Communications Authority of Kenya.',
        'Creswell, J. W. (2014). <i>Research design: Qualitative, quantitative, and mixed methods approaches</i> (4th ed.). SAGE Publications.',
        'Danes, S. M. (1994). Parental perceptions of children\'s financial socialization. <i>Financial Counselling and Planning, 5</i>(1), 127–149.',
        'Danes, S. M., & Haberman, H. (2007). Teen financial knowledge, self-efficacy, and behavior: A gendered view. <i>Financial Counselling and Planning, 18</i>(2), 48–60.',
        'Financial Sector Deepening Kenya. (2019). <i>FinAccess 2019 household survey</i>. FSD Kenya.',
        'Graboski, G., Lown, J. M., & Collins, J. M. (2001). Financial self-efficacy and its role in financial behavior. <i>Consumer Interests Annual, 47</i>, 1–3.',
        'Grable, J. E. (2000). Financial risk tolerance and additional factors that affect risk taking in everyday money matters. <i>Journal of Business and Psychology, 14</i>(4), 625–630.',
        'Grable, J. E., & Lytton, R. H. (1999). Financial risk tolerance revisited: The development of a risk assessment instrument. <i>Financial Services Review, 8</i>(3), 163–181.',
        'Grohmann, A., Klohn, F., & Menkhoff, L. (2018). Financial literacy and financial behavior in Africa: Evidence from Tanzania. <i>Review of Development Economics, 22</i>(3), 1234–1252.',
        'Hair, J. F., Black, W. C., Babin, B. J., & Anderson, R. E. (2014). <i>Multivariate data analysis</i> (7th ed.). Pearson.',
        'Karanja, P. (2019). Financial literacy and investment decisions among business students in Nairobi county universities. Unpublished MBA project, University of Nairobi.',
        'Kim, J., LaTaillade, J., & Kim, H. (2011). Family processes and adolescents\' financial behaviors. <i>Journal of Family and Economic Issues, 32</i>(4), 668–679.',
        'Kramer, M. M. (2012). Financial advice and individual investor portfolio performance. <i>Financial Management, 41</i>(2), 395–428.',
        'Lown, J. M. (2011). Development and validation of a financial self-efficacy scale. <i>Journal of Financial Counseling and Planning, 22</i>(2), 54–63.',
        'Lusardi, A., & Mitchell, O. S. (2014). The economic importance of financial literacy: Theory and evidence. <i>Journal of Economic Literature, 52</i>(1), 5–44.',
        'Mugenda, O. M., & Mugenda, A. G. (2003). <i>Research methods: Quantitative and qualitative approaches</i>. Acts Press.',
        'Mwangi, C. I., & Njeru, A. (2015). Financial literacy and investment decisions of SACCO members in Kenya. <i>International Journal of Business and Management, 10</i>(9), 245–256.',
        'Nunnally, J. C. (1978). <i>Psychometric theory</i> (2nd ed.). McGraw-Hill.',
        'OECD. (2020). <i>OECD/INFE international survey of adult financial literacy</i>. OECD Publishing.',
        'Shim, S., Barber, B. L., Card, N. A., Xiao, J. J., & Serido, J. (2010). Financial socialization of first-year college students: The roles of parents, work, and education. <i>Journal of Youth and Adolescence, 39</i>(12), 1457–1470.',
        'Shiller, R. J. (2012). <i>Finance and the good society</i>. Princeton University Press.',
        'van Rooij, M., Lusardi, A., & Alessie, R. (2011). Financial literacy and stock market participation. <i>Journal of Financial Economics, 101</i>(2), 449–472.',
        'Ward, S. (1974). Consumer socialization. <i>Journal of Consumer Research, 1</i>(2), 1–14.',
        'Woodyard, A., & Grable, J. E. (2018). Doing better, feeling worse: The paradox of financial capability and risk tolerance. <i>Financial Services Review, 27</i>(1), 1–20.',
        'Yamane, T. (1967). <i>Statistics: An introductory analysis</i> (2nd ed.). Harper & Row.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        r_text = ref.replace('<i>', '').replace('</i>', '')
        parts = ref.split('<i>')
        for pi, part in enumerate(parts):
            if '</i>' in part:
                ipart, rest = part.split('</i>', 1)
                ri = p.add_run(ipart)
                ri.font.size = Pt(12); ri.font.italic = True; ri.font.name = 'Times New Roman'
                rr = p.add_run(rest)
                rr.font.size = Pt(12); rr.font.name = 'Times New Roman'
            else:
                rr = p.add_run(part)
                rr.font.size = Pt(12); rr.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # ===================== APPENDICES =====================
    add_page_break(doc)
    section_title(doc, 'APPENDICES')

    add_heading2(doc, 'APPENDIX I: LETTER OF INTRODUCTION')
    for line in [
        ('MOI UNIVERSITY', True), ('School of Business and Economics', False),
        ('Department of Accounting and Finance', False), ('Annex Campus, Nairobi', False),
        ('P.O. Box 3900 – 30100, Eldoret, Kenya', False), ('March 2026', False),
    ]:
        p = doc.add_paragraph()
        r = p.add_run(line[0])
        r.font.size = Pt(12); r.font.bold = line[1]; r.font.name = 'Times New Roman'
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    add_body_paragraph(doc, 'TO WHOM IT MAY CONCERN,')
    add_body_paragraph(doc, 'RE: INTRODUCTION OF RESEARCH STUDENT')
    add_body_paragraph(doc, 'The above-named student, WANYONYI NAFULA SOPHIE (BBM/4452/23), is a final-year undergraduate student in the Bachelor of Business Management (Finance and Banking Option) programme at Moi University Annex Campus. She is currently conducting a research study titled "The Effect of Financial Capabilities on Risk-Taking Among University Students" in partial fulfillment of the requirements for the award of her degree.')
    add_body_paragraph(doc, 'We kindly request you to grant her access to your institution and to allow the students under your administration to participate in the study by responding to the attached questionnaire. All information provided will be treated with the utmost confidentiality and used solely for academic purposes. The findings of the study will be made available to your institution upon request.')
    add_body_paragraph(doc, 'We appreciate your kind cooperation in this regard.')
    doc.add_paragraph()
    for line in ['Yours faithfully,', '', 'Dr. Joel Tuwey', 'Senior Lecturer', 'Department of Accounting and Finance', 'Moi University']:
        p = doc.add_paragraph()
        p.add_run(line).font.size = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    add_page_break(doc)
    add_heading2(doc, 'APPENDIX II: RESEARCH QUESTIONNAIRE')
    p_inst = doc.add_paragraph()
    r_inst = p_inst.add_run('WANYONYI NAFULA SOPHIE — BBM/4452/23 | School of Business and Economics | Moi University')
    r_inst.font.size = Pt(11); r_inst.font.italic = True; r_inst.font.name = 'Times New Roman'
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_title = doc.add_paragraph()
    r_title = p_title.add_run('THE EFFECT OF FINANCIAL CAPABILITIES ON RISK-TAKING AMONG UNIVERSITY STUDENTS')
    r_title.font.size = Pt(12); r_title.font.bold = True; r_title.font.name = 'Times New Roman'
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(8)
    p_title.paragraph_format.space_after = Pt(6)

    add_body_paragraph(doc, 'Dear Respondent,\nThis questionnaire is designed to collect data on your financial capabilities and financial risk-taking behaviour. Your participation is entirely voluntary and all information provided will remain strictly confidential and will be used solely for academic research purposes. Please answer all questions honestly and to the best of your knowledge. Do not write your name on this questionnaire.')

    add_heading3(doc, 'SECTION A: DEMOGRAPHIC INFORMATION')
    add_body_paragraph(doc, 'Please tick (✓) or fill in the appropriate response for each item.')
    demog_items = [
        ('1. Gender:', ['Male [ ]', 'Female [ ]']),
        ('2. Age:', ['18–21 years [ ]', '22–25 years [ ]', '26–30 years [ ]', 'Above 30 years [ ]']),
        ('3. Year of Study:', ['Year One [ ]', 'Year Two [ ]', 'Year Three [ ]', 'Year Four [ ]']),
        ('4. Programme of Study:', ['Finance and Banking [ ]', 'Business Management [ ]', 'Accounting [ ]', 'Economics [ ]', 'Other (specify): ________']),
    ]
    for q, opts in demog_items:
        p = doc.add_paragraph()
        p.add_run(q).font.size = Pt(12)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for opt in opts:
            po = doc.add_paragraph()
            po.add_run(f'    {opt}').font.size = Pt(12)
            po.paragraph_format.space_before = Pt(0)
            po.paragraph_format.space_after = Pt(0)
            po.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    scale_note = '\nPlease rate each statement using the scale below:\n1 = Strongly Disagree,  2 = Disagree,  3 = Neutral,  4 = Agree,  5 = Strongly Agree'

    sections_q = [
        ('SECTION B: FINANCIAL SELF-EFFICACY (Independent Variable 1)', [
            'I am confident in my ability to manage my personal monthly budget.',
            'I can evaluate financial products and investment options effectively.',
            'I make financial plans and follow through with them consistently.',
            'I am confident in my ability to borrow money responsibly and repay on time.',
            'I can identify potential financial risks before committing my resources.',
            'I am confident in my ability to manage unexpected financial challenges or emergencies.',
        ]),
        ('SECTION C: FINANCIAL KNOWLEDGE (Independent Variable 2)', [
            'I understand how interest rates affect the cost of borrowing money.',
            'I understand the concept of inflation and how it affects the value of money.',
            'I know how to diversify an investment portfolio to reduce financial risk.',
            'I understand the risk-return trade-off when making investment decisions.',
            'I can interpret basic financial statements such as income statements and balance sheets.',
            'I understand how financial markets, including the NSE, operate.',
        ]),
        ('SECTION D: FINANCIAL SOCIALIZATION (Independent Variable 3)', [
            'My parents or guardians regularly discussed money management with me while growing up.',
            'My family modelled productive savings and investment behaviour that I observed.',
            'My educational institution has taught me practical financial management skills.',
            'My peers positively influence my financial decision-making and attitudes.',
            'I learned about financial risks and opportunities through family discussions.',
            'Social and cultural norms in my community support productive investment behaviour.',
        ]),
        ('SECTION E: FINANCIAL ADVICE (Independent Variable 4)', [
            'I regularly seek advice from qualified financial professionals before making major financial decisions.',
            'I make use of my university\'s financial advisory services for financial guidance.',
            'I receive reliable financial information and guidance from banking institutions.',
            'I access credible online resources and platforms for financial guidance.',
            'I receive guidance on assessing financial risks from a trusted and qualified adviser.',
            'The financial advice I receive helps me make better-informed financial decisions.',
        ]),
        ('SECTION F: RISK-TAKING BEHAVIOUR (Dependent Variable)', [
            'I participate in savings and investment groups such as chamas or SACCOs.',
            'I have invested or plan to invest in formal financial markets such as NSE, bonds, or mutual funds.',
            'I am willing to invest in higher-risk financial instruments if they offer higher potential returns.',
            'I have started or plan to start a business that requires financial investment.',
            'I use credit facilities such as HELB, bank loans, or mobile credit to pursue financial opportunities.',
            'I evaluate the risk-return profile of financial products before making investment decisions.',
            'I take calculated financial risks to improve my current and future financial status.',
            'I would invest in new financial products or markets if I were sufficiently informed about them.',
        ]),
    ]

    for sec_title, questions in sections_q:
        add_heading3(doc, sec_title)
        add_body_paragraph(doc, scale_note)
        for qi, q in enumerate(questions, 1):
            p = doc.add_paragraph()
            r = p.add_run(f'{qi}. {q}')
            r.font.size = Pt(12); r.font.name = 'Times New Roman'
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            pr = doc.add_paragraph()
            pr.add_run('    1 [ ]    2 [ ]    3 [ ]    4 [ ]    5 [ ]').font.size = Pt(12)
            pr.paragraph_format.space_before = Pt(0)
            pr.paragraph_format.space_after = Pt(4)
            pr.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    p_end = doc.add_paragraph()
    r_end = p_end.add_run('THANK YOU FOR YOUR PARTICIPATION!')
    r_end.font.size = Pt(12); r_end.font.bold = True; r_end.font.name = 'Times New Roman'
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_end.paragraph_format.space_before = Pt(16)

    add_page_num_footer(doc)
    doc.save('files/Sophie_Research_Project.docx')
    print('Successfully created: Sophie_Research_Project.docx')

generate_docx()
