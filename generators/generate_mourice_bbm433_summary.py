#!/usr/bin/env python3
"""
Generate Mourice's elite, exam-ready BBM 433 summary as a single PDF.
Covers all nine course-outline topics in detail with Kenyan examples.
Reuses the cover-page + formatting helpers from the BBM 433 CATs generator.
"""

import os as _os
import sys as _sys

_sys.path.insert(0, _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))
_os.chdir(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))

from docx.enum.text import WD_ALIGN_PARAGRAPH

from generators.generate_mourice_bbm433_cats import (
    new_doc, add_cover, body, heading, subheading, labelled, page_break,
    cover_line, convert_to_pdf,
)


# ────────────────────────────────────────────────────────────────────────────
# Light wrappers
# ────────────────────────────────────────────────────────────────────────────
def topic_title(doc, num, title):
    page_break(doc)
    p = doc.add_paragraph()
    r = p.add_run(f'TOPIC {num}: {title.upper()}')
    from docx.shared import Pt, RGBColor
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)


def bullet(doc, text, indent=0.35):
    p = doc.add_paragraph(style='List Bullet')
    p.runs.clear() if p.runs else None
    # python-docx List Bullet style adds a bullet glyph; we just set text + format
    if p.runs:
        for r in p.runs:
            r.text = ''
    r = p.add_run(text)
    from docx.shared import Pt, Inches
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(3)
    from docx.enum.text import WD_LINE_SPACING
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def b_label(doc, label, text, indent=0.4):
    """Bullet-style row with bold label and body text."""
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_LINE_SPACING
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r0 = p.add_run('• ')
    r0.font.name = 'Times New Roman'
    r0.font.size = Pt(12)
    r1 = p.add_run(label)
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r1.font.bold = True
    r2 = p.add_run(text)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    return p


# ────────────────────────────────────────────────────────────────────────────
# Document
# ────────────────────────────────────────────────────────────────────────────
def build(output='files/Mourice_BBM_433_Notes_Summary.docx'):
    doc = new_doc()
    add_cover(
        doc,
        course_code='BBM 433',
        course_title='RETAIL AND MERCHANDISE MANAGEMENT',
        assignment='COMPREHENSIVE NOTES SUMMARY',
        sub_date='APRIL 2026',
        marks='REVISION DOCUMENT',
    )

    # ── Table of Contents ───────────────────────────────────────────────
    cover_line(doc, 'TABLE OF CONTENTS', size=14, bold=True, after=10)
    toc_items = [
        ('Course Overview', 'i'),
        ('Topic 1: The New Retail Landscape', '1'),
        ('Topic 2: The Omnichannel Customer Journey & Experience Management', '4'),
        ('Topic 3: Retail Strategy — Brand, Value Creation & Ethics', '8'),
        ('Topic 4: Retail Analytics & Financial Performance', '12'),
        ('Topic 5: Merchandise Planning & Assortment Strategy', '16'),
        ('Topic 6: Digital Merchandising & the 7C Framework', '20'),
        ('Topic 7: Visual Merchandising & Store Experience', '24'),
        ('Topic 8: Supply Chain Management & Technological Innovation', '28'),
        ('Topic 9: Emerging Technologies & the Future of Retail', '32'),
        ('Conclusion & Exam-Style Pointers', '36'),
    ]
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    for title, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r = p.add_run(title)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r2 = p.add_run('\t' + page)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)

    # ── Course Overview ─────────────────────────────────────────────────
    page_break(doc)
    cover_line(doc, 'COURSE OVERVIEW', size=14, bold=True, after=10,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    body(doc,
         'BBM 433 — Retail and Merchandise Management — explores how '
         'retailers in the 21st century design, deliver, measure and '
         'continually re-imagine the shopping experience across '
         'physical, digital and social channels. The course is built '
         'around nine interlocking topics that move from the macro '
         'forces reshaping retail (Topic 1) through the customer '
         'journey, brand strategy, analytics, merchandise planning, '
         'digital and visual merchandising, supply chain, and finally '
         'into the emerging technologies that will define the next '
         'decade of retail (Topic 9). The unifying theme is the shift '
         'from channel-centric thinking to customer-centric, data-'
         'driven, omnichannel retailing.')
    body(doc,
         'This document is a single, exam-ready summary of the official '
         'BBM 433 notes (2025/26 academic year) prepared by Mourice '
         'Onyango (BBM/1891/22). Each topic is condensed into the key '
         'definitions, frameworks, KPIs, examples and practical '
         'implications you need to recall under exam pressure. Kenyan '
         'illustrations (Naivas, Quickmart, Carrefour, Jumia, Safaricom '
         'M-PESA, Vivo Activewear, Hotpoint, Twiga Foods) are woven in '
         'throughout to anchor each concept in the local market.',
         after=12)

    # ── Topic 1 ─────────────────────────────────────────────────────────
    topic_title(doc, '1', 'The New Retail Landscape')
    subheading(doc, '1.1 Defining 21st-Century Retailing')
    body(doc,
         'Retailing is the set of business activities that add value to '
         'goods and services sold to consumers for their personal, '
         'family or household use. In the 21st century, however, the '
         'retailer is no longer just the final link in the supply chain '
         '— they are the orchestrator of an end-to-end customer '
         'experience that spans physical stores, mobile apps, social '
         'commerce, marketplaces, voice assistants and same-day '
         'fulfilment networks.')
    subheading(doc, '1.2 Key Forces of Change')
    b_label(doc, 'Technology disruption: ',
            'Mobile commerce, AI personalisation, AR/VR try-ons, '
            'cashier-less checkout, IoT shelf sensors and cloud-native '
            'commerce platforms have collapsed the cost and complexity '
            'of going omnichannel.')
    b_label(doc, 'Empowered consumers: ',
            'Shoppers are always-on, price-transparent and review-'
            'driven. Showrooming (research in-store, buy online) and '
            'web-rooming (research online, buy in-store) are now the '
            'norm. Loyalty must be earned every transaction.')
    b_label(doc, 'Channel proliferation: ',
            'TikTok Shop, WhatsApp Business, Instagram Reels, Jumia, '
            'Kilimall and same-app super-apps have multiplied the '
            'paths to purchase. The "store" is wherever the customer '
            'is paying attention.')
    b_label(doc, 'Sustainability & ethics: ',
            'Conscious consumers reward retailers that demonstrate '
            'fair sourcing, low-carbon logistics, circular packaging '
            'and inclusive employment.')
    b_label(doc, 'Macro-economic volatility: ',
            'Inflation, currency depreciation, fuel costs and '
            'political uncertainty (especially in emerging markets '
            'like Kenya) compress margins and force agile pricing, '
            'sourcing and assortment decisions.')
    b_label(doc, 'Data as a strategic asset: ',
            'First-party data (loyalty cards, app behaviour, BOPIS '
            'pickups) is now the most defensible competitive moat — '
            'especially as third-party cookies and ad targeting fade.')
    subheading(doc, '1.3 The Strategic Decision-Making Model')
    body(doc,
         'Modern retail strategy follows a continuous loop: '
         '(1) Scan the environment using PESTEL + Porter\'s Five '
         'Forces; (2) Define the target customer and value '
         'proposition; (3) Choose the retail format and channel mix; '
         '(4) Plan merchandise, pricing and store experience; '
         '(5) Execute and measure through KPIs (sales/sq ft, GMROI, '
         'CLV, NPS, conversion); (6) Learn and re-deploy. Retailers '
         'that loop fastest — Zara, Amazon, Naivas — win.')
    subheading(doc, '1.4 Kenyan Snapshot')
    body(doc,
         'Kenya\'s retail sector is growing at ~6% per year and is '
         'split between modern trade (Naivas with ~110 stores, '
         'Quickmart with ~60, Carrefour, Chandarana) and a vibrant '
         'informal/dukawalla segment that still controls roughly 70% '
         'of FMCG spend. M-PESA, fibre roll-out, and rising urban '
         'middle-class incomes are pulling the formal share upwards '
         'every year, while pure-play e-commerce (Jumia, Kilimall, '
         'Copia) and social commerce on WhatsApp/Instagram continue '
         'to fragment shopper attention.')

    # ── Topic 2 ─────────────────────────────────────────────────────────
    topic_title(doc, '2', 'The Omnichannel Customer Journey & Experience Management')
    subheading(doc, '2.1 Multichannel vs Cross-channel vs Omnichannel')
    b_label(doc, 'Multichannel: ',
            'Multiple, siloed channels (separate store ops and e-com). '
            'Inventory, pricing and customer data are not shared.')
    b_label(doc, 'Cross-channel: ',
            'Some integration (e.g. buy online, return in store) but '
            'the customer view is still partial.')
    b_label(doc, 'Omnichannel: ',
            'A single, unified customer view, unified inventory, '
            'consistent pricing, persistent cart and seamless '
            'transitions between every touchpoint.')
    body(doc,
         'Why it matters: omnichannel customers spend 30% more and '
         'have higher lifetime value (CLV) than single-channel '
         'shoppers. In Kenya, a Naivas customer who uses Naivas Now '
         'delivery + in-store + the loyalty card is 2-3× more valuable '
         'than a single-channel shopper.')
    subheading(doc, '2.2 The Modern (Non-Linear) Customer Journey')
    body(doc,
         'The traditional AIDA funnel (Awareness → Interest → Desire '
         '→ Action) has been replaced by a non-linear, channel-hopping '
         'loop: Awareness → Consideration & Research → Purchase → '
         'Experience & Usage → Post-Purchase Service → Loyalty & '
         'Advocacy → (back to Awareness for the next category). At '
         'every stage, the customer can jump channels — see a TikTok, '
         'check Jumia reviews, walk into Quickmart to feel the '
         'product, then order via the app for home delivery.')
    subheading(doc, '2.3 Customer Journey Mapping')
    body(doc,
         'A customer journey map visualises the end-to-end experience '
         'across personas, stages, touchpoints, customer actions, '
         'emotions, pain points and opportunities. It is the single '
         'most powerful tool for diagnosing where customers drop off '
         'and where to invest. Outputs typically take the form of a '
         'swim-lane diagram or a journey-stage table.')
    subheading(doc, '2.4 Customer Experience Management (CXM)')
    body(doc,
         'CXM is the discipline of designing, measuring and optimising '
         'customer interactions to build loyalty. The five core '
         'principles are: customer-centricity, consistency, continuity '
         '(persistent context across channels), contextual relevance '
         '(personalisation by location, device, history) and empathy/'
         'responsiveness.')
    body(doc,
         'Key CXM capabilities: (1) Unified customer profile via a '
         'Customer Data Platform (CDP); (2) Personalisation engine '
         '(AI-driven recommendations); (3) Journey orchestration '
         '(triggered cross-channel messages); (4) Unified commerce '
         '(one OMS, one inventory, one cart); (5) Unified customer '
         'service (chat → phone → email with full context).')
    subheading(doc, '2.5 Omnichannel Capabilities Customers Now Expect')
    b_label(doc, 'BOPIS (Buy Online, Pick Up In Store): ',
            'Customer orders on the app, collects from the nearest '
            'branch — saves shipping cost and fuels footfall.')
    b_label(doc, 'Ship-from-Store: ',
            'Stores act as fulfilment nodes for nearby online orders, '
            'cutting last-mile cost.')
    b_label(doc, 'Endless Aisle: ',
            'In-store kiosks let staff order out-of-stock items for '
            'home delivery, salvaging the sale.')
    b_label(doc, 'Return Anywhere: ',
            'Online purchase can be returned in any branch.')
    b_label(doc, 'Real-time Inventory Visibility: ',
            'Customers see live stock-availability per store before '
            'they travel.')
    subheading(doc, '2.6 Omnichannel KPIs')
    body(doc,
         'Move beyond channel-only metrics to: Customer Lifetime '
         'Value (CLV), Omnichannel Share of Wallet, Cross-Channel '
         'Conversion Rate, BOPIS Adoption Rate, Return-to-Store Rate, '
         'Customer Effort Score (CES) and Net Promoter Score (NPS).')

    # ── Topic 3 ─────────────────────────────────────────────────────────
    topic_title(doc, '3', 'Retail Strategy — Brand, Value Creation & Ethics')
    subheading(doc, '3.1 What is a Retail Strategy?')
    body(doc,
         'A retail strategy is the comprehensive plan that answers: '
         'who is our target customer, what value will we deliver, '
         'through which channels and formats, how will we differentiate '
         'and how will we sustain a profitable advantage. It rests on '
         'three pillars — Brand Positioning, Value Creation and '
         'Ethics — that must be designed together, not in isolation.')
    subheading(doc, '3.2 Brand Positioning')
    body(doc,
         'Positioning is the act of designing the company\'s offering '
         'and image to occupy a distinctive place in the target '
         'customer\'s mind. The process is: Segmentation → Targeting '
         '→ Differentiation → Positioning Statement (STP-D).')
    b_label(doc, 'Segmentation bases: ',
            'Demographic (age, income), Geographic (urban vs rural), '
            'Psychographic (lifestyle, values), Behavioural (usage, '
            'loyalty, benefits sought).')
    b_label(doc, 'Differentiation variables: ',
            'Assortment (breadth vs depth), service level, price/'
            'value (EDLP vs High/Low), experience, access/'
            'convenience.')
    body(doc,
         'A positioning statement template: "To [target market], '
         '[retailer] is the brand of [frame of reference] that '
         'provides [point of difference] because [reason to believe]." '
         'Example — Naivas: "To value-conscious Kenyan families, '
         'Naivas is the supermarket that delivers fresh produce, '
         'wide assortment and proudly local brands at fair prices, '
         'because we are 100% Kenyan-owned with the country\'s '
         'largest store network."')
    subheading(doc, '3.3 Value Creation')
    body(doc,
         'Customer Value = (Functional Benefits + Emotional Benefits) '
         '÷ (Monetary Price + Time + Effort). Retailers create value '
         'through their value chain — primary activities (merchandise '
         'management, store/digital experience, marketing & sales) '
         'plus support activities (supply chain, technology/data, '
         'human resources).')
    b_label(doc, 'Utilitarian value: ',
            'Functional, task-oriented (e.g. Mwananchi Wholesale — '
            'lowest prices in bulk).')
    b_label(doc, 'Hedonic value: ',
            'Emotional, experiential (e.g. The Junction Mall, Sephora, '
            'Vivo Activewear pop-up runways).')
    b_label(doc, 'Social value: ',
            'Identity and belonging (e.g. Patagonia, Kikoy Co.\'s '
            '"proudly Kenyan, ethically made" community).')
    subheading(doc, '3.4 Ethics in Retail')
    body(doc,
         'Ethics in retail goes beyond legal compliance. The five '
         'most common ethical dilemmas are: (1) Supply-chain & '
         'sourcing (labour, child labour, environmental damage); '
         '(2) Pricing practices (gouging, deceptive promotions); '
         '(3) Data privacy & security; (4) Environmental '
         'sustainability (packaging, waste, carbon); (5) Employee '
         'relations (wages, scheduling, benefits).')
    body(doc,
         'The business case for ethics is strong: it manages risk, '
         'builds brand equity, attracts talent and underpins long-'
         'term profitability. Carroll\'s CSR pyramid (Economic → '
         'Legal → Ethical → Philanthropic) is the dominant framework. '
         'Patagonia, IKEA and Safaricom Foundation are textbook '
         'examples; locally, Bidco and Naivas\' "Mtaa Wetu" '
         'community programmes earn similar trust dividends.')

    # ── Topic 4 ─────────────────────────────────────────────────────────
    topic_title(doc, '4', 'Retail Analytics & Financial Performance')
    subheading(doc, '4.1 Why Analytics is Now Core to Retail')
    body(doc,
         'Retail margins are thin (often 2-4% net) and inventory-'
         'intensive. Analytics turns the avalanche of POS, e-commerce, '
         'loyalty and supply-chain data into decisions on pricing, '
         'assortment, replenishment, staffing and personalisation. '
         'The retailers that win are not those with the most data — '
         'they are those that act on it fastest.')
    subheading(doc, '4.2 The Analytics Maturity Model')
    b_label(doc, 'Descriptive: ',
            'What happened? (last week\'s sales by branch).')
    b_label(doc, 'Diagnostic: ',
            'Why did it happen? (drop in Kisumu was caused by a '
            'stockout of cooking oil).')
    b_label(doc, 'Predictive: ',
            'What will happen? (forecast demand for back-to-school '
            'stationery in January).')
    b_label(doc, 'Prescriptive: ',
            'What should we do? (recommend optimal markdown for '
            'slow-moving Easter chocolate).')
    subheading(doc, '4.3 Key Retail Financial KPIs')
    b_label(doc, 'Same-Store Sales (SSS): ',
            'Sales from stores open ≥ 12 months — isolates organic '
            'growth from new openings.')
    b_label(doc, 'Sales per Square Foot: ',
            'Net Sales ÷ selling area — the headline productivity '
            'measure for physical retail.')
    b_label(doc, 'Average Transaction Value (ATV): ',
            'Total Sales ÷ Number of Transactions — pushed up by '
            'cross-sell, upsell and bundling.')
    b_label(doc, 'Units per Transaction (UPT): ',
            'Effectiveness of merchandising and cross-sell.')
    b_label(doc, 'Gross Margin: ',
            '(Net Sales − COGS) ÷ Net Sales — the pool that pays for '
            'rent, salaries, marketing and profit.')
    b_label(doc, 'GMROI (Gross Margin Return on Investment): ',
            'Gross Margin ÷ Average Inventory Cost — single most '
            'important productivity metric in retail. Below 1.0 means '
            'the inventory is losing money.')
    b_label(doc, 'Inventory Turnover (Stock Turn): ',
            'COGS ÷ Average Inventory — how many times stock cycles '
            'through in a period.')
    b_label(doc, 'Sell-Through Rate: ',
            'Units Sold ÷ Units Received — short-horizon merchandise '
            'health, especially for fashion/seasonal.')
    b_label(doc, 'Customer Acquisition Cost (CAC) & CLV: ',
            'Marketing spend per new customer; lifetime profit per '
            'customer. CLV ÷ CAC ratio guides marketing investment.')
    subheading(doc, '4.4 Analytical Frameworks Every Retailer Uses')
    b_label(doc, 'DuPont / Strategic Profit Model: ',
            'ROA = Net Profit Margin × Asset Turnover; decomposes '
            'profitability into pricing, cost and asset-efficiency '
            'levers.')
    b_label(doc, 'ABC Analysis (Pareto 80/20): ',
            'A-items = top 20% of SKUs that drive 80% of sales — '
            'tight stock control. C-items = long tail — relax '
            'controls, consider delisting.')
    b_label(doc, 'Market Basket Analysis: ',
            'Association rules ("customers who buy nappies also buy '
            'wipes") — drives planogram design, bundles and '
            'recommendations.')
    subheading(doc, '4.5 The Three Core Financial Statements')
    body(doc,
         'Income Statement (P&L) reports performance over a period — '
         'revenue, COGS, gross profit, operating expenses, EBIT, net '
         'profit. Balance Sheet reports the financial position at a '
         'point in time — assets (mainly inventory + receivables + '
         'fixed assets), liabilities (payables, debt), equity. Cash '
         'Flow Statement reports cash in/out across operating, '
         'investing and financing activities — critical because '
         'retail can be profitable on paper but cash-poor.')

    # ── Topic 5 ─────────────────────────────────────────────────────────
    topic_title(doc, '5', 'Merchandise Planning & Assortment Strategy')
    subheading(doc, '5.1 What is Merchandise Planning?')
    body(doc,
         'Merchandise planning is the disciplined process of forecasting, '
         'budgeting and managing the right products, in the right '
         'quantities, at the right time, in the right place and at the '
         'right price (the "5 Rights"). It balances four trade-offs: '
         'in-stock vs over-stock, breadth vs depth, full-price sell-'
         'through vs markdown clearance, and centralisation vs '
         'localisation.')
    subheading(doc, '5.2 The Planning Hierarchy & OTB')
    body(doc,
         'Plans cascade from Company → Division → Department → Class → '
         'Sub-class → SKU. Each level owns a sales plan, a margin plan '
         'and an inventory plan. The Open-to-Buy (OTB) system controls '
         'how much new inventory the buyer can commit to in a given '
         'month: OTB = Planned EOM Inventory + Planned Sales + Planned '
         'Markdowns − BOM Inventory − On-Order. OTB prevents over-'
         'buying and protects margin.')
    subheading(doc, '5.3 Assortment Strategy')
    b_label(doc, 'Breadth (variety): ',
            'Number of different product categories carried.')
    b_label(doc, 'Depth (assortment within a category): ',
            'Number of SKUs per category — sizes, colours, brands.')
    b_label(doc, 'Wide & Shallow: ',
            'Many categories, few SKUs each (e.g. Naivas convenience '
            'store).')
    b_label(doc, 'Narrow & Deep: ',
            'Few categories, many SKUs (e.g. Bata shoes, Vivo '
            'Activewear).')
    b_label(doc, 'Wide & Deep: ',
            'Hypermarket model (Carrefour, Naivas Hypermart).')
    subheading(doc, '5.4 Category Management (8 Steps)')
    body(doc,
         'Define category → assess role (destination, routine, '
         'convenience, occasional/seasonal) → set objectives → set '
         'strategies → set tactics (assortment, price, promotion, '
         'shelf, supply) → implement → review → refine. Category '
         'roles drive shelf space, price aggressiveness and '
         'promotional intensity.')
    subheading(doc, '5.5 SKU Rationalisation, Localisation & Replenishment')
    body(doc,
         'SKU rationalisation removes duplicate or under-performing '
         'SKUs (Pareto: kill the bottom 20% that contribute < 5% of '
         'sales). Localisation tailors the range to each store '
         '(Eldoret may carry more agricultural inputs; Mombasa more '
         'seafood). Replenishment models include reorder point, '
         'periodic review, and continuous (auto-replenishment) — '
         'JIT for fast-movers, EOQ for stable demand, vendor-managed '
         'inventory (VMI) for collaborative supply.')
    subheading(doc, '5.6 Managing Seasonal & Fashion Goods')
    body(doc,
         'Seasonal/fashion goods follow a planned mark-down curve: '
         'full price → first markdown (15-25%) → second markdown '
         '(40-50%) → clearance (60%+). The aim is to maximise total '
         'gross margin dollars, not to avoid markdowns altogether. '
         'Buy-and-react (Zara) beats buy-and-hope.')

    # ── Topic 6 ─────────────────────────────────────────────────────────
    topic_title(doc, '6', 'Digital Merchandising & the 7C Framework')
    subheading(doc, '6.1 What is Digital Merchandising?')
    body(doc,
         'Digital merchandising is the practice of curating, presenting '
         'and optimising products on digital storefronts (website, '
         'app, marketplace listings, social commerce) to maximise '
         'discovery, engagement and conversion. The pixel is the new '
         'shelf; site search, product imagery, recommendations and '
         'PDP (product-detail-page) design are the new visual '
         'merchandising tools.')
    subheading(doc, '6.2 The 7C Framework for Digital Retailing')
    b_label(doc, 'Context: ',
            'Site/app aesthetic, layout, navigation, brand feel and '
            'mobile-first design. First impression in 3 seconds.')
    b_label(doc, 'Content: ',
            'Product copy, photography, video, 360° views, size '
            'guides, user-generated content, expert reviews.')
    b_label(doc, 'Community: ',
            'Reviews, Q&A, forums, social proof, ambassador programs '
            '— builds trust and social validation.')
    b_label(doc, 'Customisation: ',
            'Personalised home page, recommendations, dynamic '
            'merchandising by segment / behaviour / location.')
    b_label(doc, 'Communication: ',
            'Two-way channels — chatbots, WhatsApp, push, email, in-'
            'app messages.')
    b_label(doc, 'Connection: ',
            'Integration with social platforms, marketplaces, '
            'partner ecosystems and offline channels (BOPIS, ship-'
            'from-store).')
    b_label(doc, 'Commerce: ',
            'Frictionless checkout — saved cards, M-PESA STK push, '
            'one-click, buy-now-pay-later, multiple delivery '
            'options.')
    subheading(doc, '6.3 Conversion Rate Optimisation (CRO)')
    body(doc,
         'CRO is the systematic A/B testing of every conversion lever '
         '— PDP layout, button colour and copy, photo type, social '
         'proof position, shipping-cost transparency, checkout step '
         'count. A 1% lift in conversion on a Ksh 100 M e-commerce '
         'business is Ksh 1 M of pure margin.')
    subheading(doc, '6.4 Digital-Merchandising Tactics')
    b_label(doc, 'Search & navigation: ',
            'Faceted filters, synonym handling, typo-tolerance, '
            'auto-suggest. ~30% of revenue comes from on-site search.')
    b_label(doc, 'Visual merchandising online: ',
            'Hero banners, "shop the look", lifestyle photography, '
            'video. Curated landing pages for events (Black Friday, '
            'Back-to-School, Eid).')
    b_label(doc, 'Recommendations: ',
            '"You may also like", "Frequently bought together", '
            '"Recently viewed" — typically lift AOV by 10-30%.')
    b_label(doc, 'Pricing & promotions online: ',
            'Dynamic pricing, flash sales, bundle pricing, free-'
            'shipping thresholds.')

    # ── Topic 7 ─────────────────────────────────────────────────────────
    topic_title(doc, '7', 'Visual Merchandising & Store Experience')
    subheading(doc, '7.1 Why the Physical Store Still Matters')
    body(doc,
         'Even after a decade of e-commerce growth, ~80% of global '
         'retail sales still happen in physical stores. The store is '
         'now a brand stage, a fulfilment hub (BOPIS, returns), a '
         'service touchpoint and the most powerful media channel a '
         'retailer owns. Visual Merchandising (VM) is the strategic '
         'art of presenting product to drive sales, communicate brand '
         'and shape behaviour.')
    subheading(doc, '7.2 The Atmospherics Framework')
    b_label(doc, 'Exterior: ',
            'Façade, signage, window display, entrance — the first '
            'impression. Aim: stop the passer-by within 3 seconds.')
    b_label(doc, 'Layout & circulation: ',
            'Grid (supermarkets — Naivas), Free-flow (apparel — Vivo), '
            'Loop/Racetrack (department stores — Carrefour), Boutique '
            '(small specialty). Decompression zone at entrance, then '
            'guided path past key categories.')
    b_label(doc, 'Fixtures & displays: ',
            'Gondolas, end-caps, dump bins, mannequins, focal walls. '
            'End-caps generate 30%+ of category sales in supermarkets.')
    b_label(doc, 'Lighting: ',
            'Ambient, accent, task and decorative — sets mood and '
            'directs the eye.')
    b_label(doc, 'Colour: ',
            'Warm colours (red, orange) attract; cool colours (blue, '
            'green) calm. Brand colour palette must be consistent.')
    b_label(doc, 'Sound: ',
            'Slow tempo → longer dwell, higher spend; up-tempo → '
            'faster turnover (good for QSR & convenience).')
    b_label(doc, 'Scent (olfactory marketing): ',
            'Fresh-baked bread at the supermarket entrance, vanilla in '
            'apparel — proven to lift dwell time and basket size.')
    subheading(doc, '7.3 Visual Merchandising Techniques')
    body(doc,
         'Display types include window, focal, point-of-purchase '
         '(POP), end-cap and lifestyle vignettes. Effective VM uses '
         'colour blocking, the rule-of-three (focal + flanking '
         'support), the AIDA principle (Attention-Interest-Desire-'
         'Action) and clear visual hierarchy. Mannequin styling, '
         'cross-merchandising (display the dress with the bag and '
         'shoes), and storytelling props turn product into '
         'aspiration.')
    subheading(doc, '7.4 Service & Experiential Retail')
    body(doc,
         'Store experience extends beyond visuals: friendly, '
         'knowledgeable staff; comfortable fitting rooms; fast '
         'checkout (cash-wrap is the most under-rated touchpoint); '
         'in-store events, classes, sampling and "retail-tainment" '
         '(e.g. Apple Today at Apple, Lululemon yoga classes, '
         'Carrefour cooking demos).')
    subheading(doc, '7.5 Omnichannel Integration In-Store')
    body(doc,
         'Modern stores blend digital seamlessly: QR codes on shelves, '
         'mobile self-checkout (Naivas Now), endless-aisle kiosks, '
         'BOPIS pickup desks, smart fitting rooms (RFID + tablet), '
         'digital signage tied to live inventory, and store-staff '
         'tablets that show the customer\'s online history.')

    # ── Topic 8 ─────────────────────────────────────────────────────────
    topic_title(doc, '8', 'Supply Chain Management & Technological Innovation')
    subheading(doc, '8.1 Why SCM Matters in Retail')
    body(doc,
         'In retail, the supply chain is the business. It determines '
         'in-stock rates, working-capital efficiency, gross margin '
         '(through landed cost), speed-to-market for new ranges, and '
         'fulfilment cost-to-serve in e-commerce. A 1% improvement in '
         'in-stock rate typically lifts sales 1-2%; a 1-day reduction '
         'in lead time can free millions in working capital.')
    subheading(doc, '8.2 Core SCM Components')
    b_label(doc, 'Demand planning & forecasting: ',
            'Statistical + ML models that fuse history, weather, '
            'promotions, holidays and macro signals.')
    b_label(doc, 'Sourcing & procurement: ',
            'Supplier selection, contracts, ethical sourcing, total-'
            'cost-of-ownership thinking (not just unit price).')
    b_label(doc, 'Inbound logistics: ',
            'Freight (sea, air, road), customs clearance, container '
            'consolidation.')
    b_label(doc, 'Warehousing & inventory: ',
            'Distribution centres, cross-docking, slotting, cycle '
            'counting, RFID/IoT.')
    b_label(doc, 'Outbound & last-mile: ',
            'Store deliveries, e-com fulfilment, BOPIS, ship-from-'
            'store, third-party riders (Glovo, Bolt Food).')
    b_label(doc, 'Reverse logistics: ',
            'Returns, refurbishment, recycling — increasingly '
            'critical in e-commerce.')
    subheading(doc, '8.3 Key SCM Metrics')
    body(doc,
         'On-Time-In-Full (OTIF), Perfect Order Rate, In-Stock %, '
         'Inventory Turns, Days of Supply, Fill Rate, Order-Cycle '
         'Time, Cost-to-Serve, and Total Landed Cost.')
    subheading(doc, '8.4 Technologies Reshaping the Supply Chain')
    b_label(doc, 'AI & ML: ',
            'Demand forecasting, dynamic routing, anomaly detection, '
            'auto-replenishment.')
    b_label(doc, 'IoT: ',
            'Smart shelves, temperature sensors for cold chain, GPS '
            'on trucks.')
    b_label(doc, 'Robotics & automation: ',
            'Goods-to-person picking, autonomous mobile robots in '
            'DCs (e.g. Amazon Kiva, Ocado).')
    b_label(doc, 'Blockchain: ',
            'Provenance tracking (e.g. Walmart\'s leafy-greens '
            'traceability) and smart contracts for payments.')
    b_label(doc, 'Digital twins: ',
            'Virtual replica of the supply chain to simulate '
            'disruptions before they happen.')
    b_label(doc, 'Cloud & APIs: ',
            'Composable, real-time data exchange between retailer, '
            'supplier and 3PL.')
    subheading(doc, '8.5 Agile, Responsive & Sustainable Supply Chains')
    body(doc,
         'The classic trade-off between efficiency (lean, low-cost, '
         'predictable) and responsiveness (fast, flexible, customer-'
         'driven) is now resolved through hybrid networks: lean '
         'core, agile periphery. Sustainability — fair labour, '
         'lower-carbon transport, recyclable packaging, circular '
         'returns — is becoming a hard requirement, not a nice-to-'
         'have.')

    # ── Topic 9 ─────────────────────────────────────────────────────────
    topic_title(doc, '9', 'Emerging Technologies & the Future of Retail')
    subheading(doc, '9.1 The Drivers of the Retail Revolution')
    body(doc,
         'Five forces are reshaping retail simultaneously: ubiquitous '
         'connectivity (5G, fibre), generative AI, the rise of social '
         'commerce, sustainability mandates, and a new generation of '
         'shoppers (Gen Z, Gen Alpha) who are mobile-native, values-'
         'driven and creator-influenced.')
    subheading(doc, '9.2 Key Emerging Technologies')
    b_label(doc, 'Artificial Intelligence & Machine Learning: ',
            'Hyper-personalisation, demand forecasting, dynamic '
            'pricing, chat-based shopping assistants, auto-generated '
            'product copy and imagery (generative AI).')
    b_label(doc, 'Augmented Reality (AR) & Virtual Reality (VR): ',
            'Virtual try-on (eyewear, make-up, apparel), AR room-'
            'planning (IKEA Place), virtual stores in the metaverse.')
    b_label(doc, 'Internet of Things (IoT) & Smart Stores: ',
            'Smart shelves, electronic shelf labels (ESL), connected '
            'fitting rooms, environmental sensors.')
    b_label(doc, 'Robotics & Automation: ',
            'Warehouse robots, in-store cleaning/scanning robots, '
            'autonomous delivery (drones, sidewalk bots).')
    b_label(doc, 'Blockchain: ',
            'Provenance, anti-counterfeit, loyalty interoperability.')
    b_label(doc, '5G & Edge Computing: ',
            'Real-time inventory, low-latency AR experiences, '
            'computer-vision checkout.')
    b_label(doc, 'Biometrics & Computer Vision: ',
            'Pay-by-palm/face (Amazon One), cashier-less stores '
            '(Amazon Go), shrink/loss detection, footfall analytics.')
    subheading(doc, '9.3 Future Retail Transformations')
    b_label(doc, 'Phygital integration: ',
            'Physical and digital merge into a single, frictionless '
            'experience.')
    b_label(doc, 'Autonomous & frictionless stores: ',
            'Walk-in, walk-out checkout-free retail.')
    b_label(doc, 'Hyper-personalisation at scale: ',
            'A store of one — every shopper sees a unique assortment, '
            'price and message.')
    b_label(doc, 'Immersive / experiential retail: ',
            'Stores designed for content creation, community and '
            'discovery — not just transaction.')
    b_label(doc, 'Sustainable & circular retail: ',
            'Resale, rental, repair, refill — built into the business '
            'model (e.g. H&M garment collecting, Patagonia Worn Wear).')
    b_label(doc, 'New commerce models: ',
            'Live shopping, social commerce, voice commerce, '
            'subscription, super-apps (Safaricom\'s M-PESA Super App, '
            'Jumia One).')
    subheading(doc, '9.4 Strategic Implications & Ethical Considerations')
    body(doc,
         'The winning retailers of the next decade will be those that '
         'combine: a sharp, value-led brand position; a unified, data-'
         'rich technology stack; a culture of experimentation and '
         'rapid learning; and a credible sustainability and ethics '
         'posture. Ethical considerations — algorithmic bias, '
         'surveillance, job displacement, data privacy — must be '
         'built in by design, not bolted on after a scandal.')

    # ── Closing ─────────────────────────────────────────────────────────
    page_break(doc)
    cover_line(doc, 'CONCLUSION & EXAM-STYLE POINTERS', size=14, bold=True,
               after=10)
    body(doc,
         'BBM 433 is ultimately about one big shift: from selling '
         'product through channels to orchestrating customer '
         'experience across an integrated, data-driven, omnichannel '
         'ecosystem. Every topic in the course feeds into that core '
         'idea — Topics 1-3 set the strategic context; Topic 4 '
         'measures it; Topics 5-7 execute it across merchandise, '
         'digital and physical; Topic 8 powers it from the supply '
         'side; and Topic 9 looks ahead.')
    subheading(doc, 'Quick Exam Tips')
    bullet(doc, 'Always anchor definitions in a Kenyan example '
                '(Naivas, Quickmart, Carrefour, Jumia, M-PESA, Vivo, '
                'Hotpoint, Twiga) — examiners reward local '
                'application.')
    bullet(doc, 'For "discuss" questions, use the structure: '
                'Definition → Why it matters → Components/Types → '
                'Example → Limitations.')
    bullet(doc, 'Memorise the headline KPIs (GMROI, SSS, Sales/sq ft, '
                'AOV, UPT, CLV, NPS, OTIF) and the formulas — they '
                'crop up as quick-marks questions every year.')
    bullet(doc, 'Memorise frameworks by acronym: STP, AIDA, PESTEL, '
                'ABC, 7C, OTB, BOPIS, CSR pyramid, DuPont.')
    bullet(doc, 'Use sub-headings, bullets and short paragraphs in '
                'your script — examiners mark on structure as much '
                'as content.')
    body(doc,
         'Good luck Mourice — revise actively, write structured '
         'answers and you will walk out with a strong A.', before=10)

    doc.save(output)
    print(f'DOCX saved: {output}')
    convert_to_pdf(output)
    # Remove intermediate DOCX (PDF only as requested)
    try:
        _os.remove(output)
        print(f'Removed intermediate DOCX: {output}')
    except Exception:
        pass


def main():
    build()


if __name__ == '__main__':
    main()
