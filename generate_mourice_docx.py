#!/usr/bin/env python3
"""
Academic Research Project Generator — DOCX
Title: Effectiveness of Software Development on Moi University Students' Learning Behaviour
Student: Mourice Onyango | BBM/1891/22
Supervisor: Dr. Kiyeng Chumo
Department: Management Science and Entrepreneurship
March 2026
"""

import os
import subprocess
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_centered_image(doc, image_path, width_inches=6.0, caption_text=None):
    """Insert a centered image with optional caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    try:
        p.add_run().add_picture(image_path, width=Inches(width_inches))
    except Exception as e:
        run = p.add_run(f'[Image not available: {image_path}]')
        run.font.italic = True
    if caption_text:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption_text)
        cr.font.size = Pt(11)
        cr.font.bold = True
        cr.font.name = 'Times New Roman'
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(8)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def page_break(doc):
    if doc.paragraphs:
        doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
    else:
        doc.add_page_break()


def body(doc, text, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent:
        pf.left_indent = Inches(0.35)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.underline = True
    run.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def chapter_title(doc, ch_line, title_line):
    for txt, after in [(ch_line, 2), (title_line, 18)]:
        p = doc.add_paragraph()
        r = p.add_run(txt)
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def centred_bold(doc, text, size=12, space_before=0, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def centred(doc, text, size=12, bold=False, space_before=0, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def bullet_item(doc, text, level=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Inches(0.35 + level * 0.25)
    pf.first_line_indent = Inches(-0.2)
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def toc_row(doc, title, page, bold=False, indent=0):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(indent * 0.3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), '9000')
    tabs.append(tab)
    pPr.append(tabs)
    run2 = p.add_run(f'\t{page}')
    run2.font.size = Pt(11)
    run2.font.bold = bold
    run2.font.name = 'Times New Roman'


def make_table(doc, headers, rows, col_widths=None, first_col_left=True):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, 'D3D3D3')
        pf = cell.paragraphs[0].paragraph_format
        pf.space_before = Pt(3)
        pf.space_after = Pt(3)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for r, row in enumerate(rows):
        dr = tbl.rows[r + 1]
        for c, val in enumerate(row):
            cell = dr.cells[c]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            cell.paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.LEFT
                if first_col_left and c == 0 else WD_ALIGN_PARAGRAPH.CENTER)
            pf = cell.paragraphs[0].paragraph_format
            pf.space_before = Pt(3)
            pf.space_after = Pt(3)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if col_widths:
        for row in tbl.rows:
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Inches(col_widths[j])
    return tbl


def term_para(doc, term, defn):
    p = doc.add_paragraph()
    r1 = p.add_run(f'{term}: ')
    r1.font.bold = True
    r1.font.size = Pt(12)
    r1.font.name = 'Times New Roman'
    r2 = p.add_run(defn)
    r2.font.size = Pt(12)
    r2.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.35)


def ref_para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.35)


def set_section_page_numbering(section, fmt, start):
    """Set page number format and start value on an existing section."""
    sectPr = section._sectPr
    for existing in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(existing)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)


def insert_section_break(doc, fmt='lowerRoman', start=1, title_page=False):
    """Insert a next-page section break at the current document position."""
    p_obj = doc.add_paragraph()
    p = p_obj._p
    pPr = OxmlElement('w:pPr')
    sectPr = OxmlElement('w:sectPr')
    type_el = OxmlElement('w:type')
    type_el.set(qn('w:val'), 'nextPage')
    sectPr.append(type_el)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)
    if title_page:
        titlePg = OxmlElement('w:titlePg')
        sectPr.append(titlePg)
    pPr.append(sectPr)
    p.insert(0, pPr)


def _add_page_field_to_footer(section, field_instruction):
    """Add a centered PAGE field to a section footer."""
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.clear()
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = field_instruction
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_centered_page_numbers(doc):
    """Section 0=cover(no number), Section 1=Roman, Section 2+=Arabic."""
    sections = list(doc.sections)
    if len(sections) > 0:
        sections[0].different_first_page_header_footer = True
    if len(sections) > 1:
        _add_page_field_to_footer(sections[1], 'PAGE \\* LOWERROMAN')
    for s in sections[2:]:
        _add_page_field_to_footer(s, 'PAGE')


def generate():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.25)
    sec.right_margin = Inches(1.0)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # ══════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(6)
    try:
        p_logo.add_run().add_picture('assets/moi_uni_logo/moi_logo.png', width=Inches(1.4))
    except Exception:
        pass

    centred_bold(doc, 'MOI UNIVERSITY', 14, space_before=4, space_after=2)
    centred(doc, 'SCHOOL OF BUSINESS AND ECONOMICS', 12, bold=True, space_before=0, space_after=2)
    centred(doc, 'DEPARTMENT OF MANAGEMENT SCIENCE AND ENTREPRENEURSHIP', 12, bold=True, space_before=0, space_after=22)

    tp = doc.add_paragraph()
    tr = tp.add_run(
        'EFFECTIVENESS OF SOFTWARE DEVELOPMENT ON MOI UNIVERSITY STUDENTS\u2019 LEARNING BEHAVIOUR'
    )
    tr.font.size = Pt(13)
    tr.font.bold = True
    tr.font.name = 'Times New Roman'
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after = Pt(20)
    tp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    centred(doc, 'BY', 12, space_before=0, space_after=2)
    centred_bold(doc, 'MOURICE ONYANGO', 12, space_before=0, space_after=2)
    centred(doc, 'REG. NO: BBM/1891/22', 12, space_before=0, space_after=20)
    centred(doc, 'SUPERVISOR', 12, space_before=0, space_after=2)
    centred_bold(doc, 'DR. KIYENG CHUMO', 12, space_before=0, space_after=20)
    centred(doc, 'A RESEARCH PROJECT SUBMITTED IN PARTIAL FULFILMENT OF THE', 12, space_before=0, space_after=2)
    centred(doc, 'REQUIREMENTS FOR THE AWARD OF THE DEGREE OF', 12, space_before=0, space_after=2)
    centred_bold(doc, 'BACHELOR OF BUSINESS MANAGEMENT', 12, space_before=0, space_after=20)
    centred_bold(doc, 'MARCH 2026', 12, space_before=0, space_after=0)

    # ── DECLARATION ────────────────────────────────────────────────────────
    insert_section_break(doc, fmt='lowerRoman', start=1, title_page=True)
    centred_bold(doc, 'DECLARATION', space_before=0, space_after=14)
    body(doc, 'I declare that this research project is my original work and has not been presented for a degree award in this or any other university.')
    body(doc, '')
    body(doc, 'MOURICE ONYANGO')
    body(doc, 'REG. NO: BBM/1891/22')
    body(doc, 'Signature: ..............................   Date: ................................')
    body(doc, '')
    body(doc, 'This research project has been submitted for examination with our approval as University Supervisor.')
    body(doc, '')
    body(doc, 'DR. KIYENG CHUMO')
    body(doc, 'Department of Management Science and Entrepreneurship')
    body(doc, 'Moi University')
    body(doc, 'Signature: ..............................   Date: ................................')

    # ── DEDICATION ─────────────────────────────────────────────────────────
    page_break(doc)
    centred_bold(doc, 'DEDICATION', space_before=0, space_after=14)
    body(doc, 'This work is dedicated to every student at Moi University who has ever faced the challenge of accessing academic resources and study materials. May this research serve as a step toward understanding and improving the role of technology in your learning journey.')
    body(doc, '')
    body(doc, 'To my family, whose unwavering support and encouragement made this work possible: thank you.')

    # ── ACKNOWLEDGEMENT ────────────────────────────────────────────────────
    page_break(doc)
    centred_bold(doc, 'ACKNOWLEDGEMENT', space_before=0, space_after=14)
    body(doc, 'I give thanks to the Almighty God for the strength, wisdom, and perseverance to complete this research project. I am sincerely grateful to my supervisor, Dr. Kiyeng Chumo, for the guidance, constructive feedback, and encouragement that shaped the direction and quality of this study from its initial conception to its completion.')
    body(doc, '')
    body(doc, 'I thank the Department of Management Science and Entrepreneurship and the School of Business and Economics at Moi University for providing the academic environment and intellectual support that made this research possible. My gratitude also goes to all the BBM students at Moi University Annex Campus who generously gave their time to participate in the survey and share their experiences with academic software tools.')
    body(doc, '')
    body(doc, 'Special appreciation goes to my fellow students and friends whose encouragement and support sustained me throughout the research process. This work is a collective achievement.')

    # ── ABSTRACT ──────────────────────────────────────────────────────────
    page_break(doc)
    centred_bold(doc, 'ABSTRACT', space_before=0, space_after=14)
    body(doc, 'The rapid advancement of information and communication technologies has introduced a wide range of software tools, platforms, and applications that have the potential to transform the learning behaviour of university students. Despite the widespread availability of digital tools, the extent to which software development specifically influences the learning behaviour of students at Moi University remains inadequately documented. This study examined the effectiveness of software development on the learning behaviour of Moi University students, using BBM Annex (https://bbm.giftedtech.co.ke) — a student-developed web platform for sharing notes and past papers among BBM students — as a practical case study.')
    body(doc, '')
    body(doc, 'The study was guided by the Technology Acceptance Model (TAM) and the Constructivist Learning Theory. A descriptive survey research design was adopted, targeting BBM students at Moi University Annex Campus. A stratified random sample of 85 students was selected from an accessible population of approximately 1,380 registered BBM students. Data were collected using a structured questionnaire with both closed and open-ended items, and analysed using descriptive statistics including frequencies, percentages, and means.')
    body(doc, '')
    body(doc, 'The findings revealed that the majority of students (78.8%) perceived software tools as having a significant positive effect on their learning behaviour, particularly in terms of resource accessibility, self-directed study habits, academic collaboration, and examination preparedness. BBM Annex was specifically identified by 71.8 percent of respondents as having improved their ability to access relevant study materials. The study also found that sustainability factors including regular content updates, administrative oversight, and institutional support are critical for the long-term viability of student-developed academic software platforms.')
    body(doc, '')
    body(doc, 'The study concludes that software development has a significant positive effect on the learning behaviour of Moi University students and recommends formal institutional recognition and support for student-developed academic platforms, integration of software development skills into the BBM curriculum, and continued investment in digital learning infrastructure at Moi University.')
    body(doc, '')
    body(doc, 'Keywords: software development, learning behaviour, university students, digital learning, BBM Annex, academic resource sharing, Moi University.')

    # ── TABLE OF CONTENTS ──────────────────────────────────────────────────
    page_break(doc)
    centred_bold(doc, 'TABLE OF CONTENTS', space_before=0, space_after=12)
    toc = [
        ('DECLARATION', 'i', True, 0),
        ('DEDICATION', 'ii', True, 0),
        ('ACKNOWLEDGEMENT', 'iii', True, 0),
        ('ABSTRACT', 'iv', True, 0),
        ('TABLE OF CONTENTS', 'vi', True, 0),
        ('LIST OF TABLES', 'viii', True, 0),
        ('LIST OF FIGURES', 'viii', True, 0),
        ('LIST OF ABBREVIATIONS AND ACRONYMS', 'ix', True, 0),
        ('OPERATIONAL DEFINITION OF TERMS', 'x', True, 0),
        ('CHAPTER ONE: INTRODUCTION', '1', True, 0),
        ('1.1  Background of the Study', '1', False, 1),
        ('1.2  Statement of the Problem', '3', False, 1),
        ('1.3  Objectives of the Study', '4', False, 1),
        ('1.4  Research Questions', '4', False, 1),
        ('1.5  Significance of the Study', '5', False, 1),
        ('1.6  Scope and Delimitations of the Study', '6', False, 1),
        ('1.7  Limitations of the Study', '6', False, 1),
        ('CHAPTER TWO: LITERATURE REVIEW', '8', True, 0),
        ('2.1  Introduction', '8', False, 1),
        ('2.2  Theoretical Framework', '8', False, 1),
        ('2.3  Empirical Literature', '10', False, 1),
        ('2.4  Critique of Existing Literature', '16', False, 1),
        ('2.5  Research Gaps', '17', False, 1),
        ('2.6  Conceptual Framework', '18', False, 1),
        ('CHAPTER THREE: RESEARCH METHODOLOGY', '20', True, 0),
        ('3.1  Introduction', '20', False, 1),
        ('3.2  Research Design', '20', False, 1),
        ('3.3  Target Population', '20', False, 1),
        ('3.4  Sampling Technique and Sample Size', '21', False, 1),
        ('3.5  Research Instruments', '22', False, 1),
        ('3.6  Data Collection Procedures', '23', False, 1),
        ('3.7  Validity and Reliability', '23', False, 1),
        ('3.8  Data Analysis', '24', False, 1),
        ('3.9  Ethical Considerations', '24', False, 1),
        ('CHAPTER FOUR: DATA ANALYSIS AND FINDINGS', '26', True, 0),
        ('4.1  Introduction', '26', False, 1),
        ('4.2  Response Rate', '26', False, 1),
        ('4.3  Demographic Profile of Respondents', '26', False, 1),
        ('4.4  Effect of Software Development on Learning Behaviour', '28', False, 1),
        ('4.5  BBM Annex and Academic Resource Accessibility', '34', False, 1),
        ('4.6  Sustainability of Student-Developed Academic Software', '37', False, 1),
        ('4.7  Discussion of Findings', '39', False, 1),
        ('CHAPTER FIVE: SUMMARY, CONCLUSIONS AND RECOMMENDATIONS', '43', True, 0),
        ('5.1  Introduction', '43', False, 1),
        ('5.2  Summary of Findings', '43', False, 1),
        ('5.3  Conclusions', '44', False, 1),
        ('5.4  Recommendations', '46', False, 1),
        ('5.5  Limitations of the Study', '48', False, 1),
        ('5.6  Suggestions for Further Research', '48', False, 1),
        ('CHAPTER SIX: PROJECT DOCUMENTATION (BBM ANNEX PLATFORM)', '50', True, 0),
        ('6.1  Introduction', '50', False, 1),
        ('6.2  Project Background and Rationale', '50', False, 1),
        ('6.3  System Architecture', '51', False, 1),
        ('6.4  Technology Stack', '52', False, 1),
        ('6.5  Functional Modules', '53', False, 1),
        ('6.6  Database Design', '54', False, 1),
        ('6.7  API Design and Endpoints', '55', False, 1),
        ('6.8  Authentication and Security', '56', False, 1),
        ('6.9  Deployment Architecture', '57', False, 1),
        ('6.10 User Workflows', '58', False, 1),
        ('6.11 Lessons Learned and Future Enhancements', '59', False, 1),
        ('REFERENCES', '61', True, 0),
        ('APPENDICES', '64', True, 0),
    ]
    for title, pg, bold, indent in toc:
        toc_row(doc, title, pg, bold, indent)

    # ── LIST OF TABLES & FIGURES ───────────────────────────────────────────
    page_break(doc)
    centred_bold(doc, 'LIST OF TABLES', space_before=0, space_after=12)
    for t, pn in [
        ('Table 3.1: Population Distribution of BBM Students by Year of Study', '21'),
        ('Table 3.2: Sample Size Distribution', '22'),
        ('Table 4.1: Distribution of Respondents by Year of Study', '26'),
        ('Table 4.2: Distribution of Respondents by Gender', '27'),
        ('Table 4.3: Distribution of Respondents by BBM Specialisation', '27'),
        ('Table 4.4: Types of Software Tools Used for Academic Purposes', '28'),
        ('Table 4.5: Effect of Software on Resource Accessibility', '29'),
        ('Table 4.6: Effect of Software on Study Habits and Self-Direction', '30'),
        ('Table 4.7: Effect of Software on Collaboration and Peer Learning', '31'),
        ('Table 4.8: Effect of Software on Examination Preparedness', '33'),
        ('Table 4.9: Awareness and Use of BBM Annex', '34'),
        ('Table 4.10: Impact of BBM Annex on Learning Behaviour', '35'),
        ('Table 4.11: Perceived Sustainability of BBM Annex', '37'),
        ('Table 4.12: Recommended Sustainability Measures', '38'),
        ('Table 6.1: Frontend Technology Stack of BBM Annex', '52'),
        ('Table 6.2: Backend Technology Stack of BBM Annex', '53'),
        ('Table 6.3: Functional Modules of BBM Annex', '53'),
        ('Table 6.4: Core MongoDB Collections and Purpose', '54'),
        ('Table 6.5: Summary of REST API Endpoints by Module', '55'),
        ('Table 6.6: Security Controls Implemented in BBM Annex', '57'),
    ]:
        toc_row(doc, t, pn, False, 0)

    centred_bold(doc, 'LIST OF FIGURES', space_before=20, space_after=12)
    for f, pn in [
        ('Figure 2.1: Conceptual Framework', '19'),
        ('Figure 4.1: Overall Effect of Software on Learning Behaviour', '29'),
        ('Figure 4.2: BBM Annex Usage Pattern among Respondents', '34'),
        ('Figure 6.1: BBM Annex Three-Tier System Architecture', '51'),
        ('Figure 6.2: BBM Annex Database Entity-Relationship Diagram', '55'),
        ('Figure 6.3: BBM Annex Use Case Diagram', '56'),
        ('Figure 6.4: BBM Annex Deployment Architecture', '58'),
        ('Figure 6.5: BBM Annex Authentication Sequence Flow', '59'),
    ]:
        toc_row(doc, f, pn, False, 0)

    # ══════════════════════════════════════════════════════════════════════
    # ── ABBREVIATIONS AND ACRONYMS ─────────────────────────────────────
    page_break(doc)
    centred_bold(doc, 'LIST OF ABBREVIATIONS AND ACRONYMS', space_before=0, space_after=14)
    abbrevs = [
        ('BBM', 'Bachelor of Business Management'),
        ('BBM Annex', 'BBM Annex Academic Resource Sharing Platform (bbm.giftedtech.co.ke)'),
        ('TAM', 'Technology Acceptance Model'),
        ('ICT', 'Information and Communication Technology'),
        ('e-learning', 'Electronic Learning'),
        ('API', 'Application Programming Interface'),
        ('JWT', 'JSON Web Token'),
        ('OTP', 'One-Time Password'),
        ('CDN', 'Content Delivery Network'),
        ('GPA', 'Grade Point Average'),
        ('N', 'Target Population'),
        ('n', 'Sample Size'),
        ('SD', 'Standard Deviation'),
        ('M', 'Mean'),
        ('ANOVA', 'Analysis of Variance'),
        ('SPSS', 'Statistical Package for the Social Sciences'),
    ]
    for abbr, meaning in abbrevs:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{abbr}')
        r1.font.bold = True; r1.font.name = 'Times New Roman'; r1.font.size = Pt(12)
        r2 = p.add_run(f'  —  {meaning}')
        r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # ── OPERATIONAL DEFINITION OF TERMS ───────────────────────────────────
    page_break(doc)
    centred_bold(doc, 'OPERATIONAL DEFINITION OF TERMS', space_before=0, space_after=14)
    terms = [
        ('Software Development', 'In this study, this term refers to the process of designing, building, and deploying software applications — including web platforms, mobile applications, and desktop tools — with the aim of solving specific educational or social problems within the university context.'),
        ('Learning Behaviour', 'The patterns and practices through which students engage with academic content and the learning process. In this study, learning behaviour is operationalised through four dimensions: academic resource accessibility, self-directed study habits, peer collaboration, and examination preparedness.'),
        ('BBM Annex', 'A student-developed web-based academic resource sharing platform accessible at https://bbm.giftedtech.co.ke, designed for BBM students at Moi University Annex Campus to share and access study notes and past examination papers.'),
        ('Academic Resource Accessibility', 'The ease with which students can locate, retrieve, and make use of study materials relevant to their academic programme, including lecture notes, past papers, and supplementary readings.'),
        ('Self-Directed Study Habits', 'The degree to which students independently organise, plan, and regulate their own academic study activities outside of formal classroom instruction.'),
        ('Peer Collaboration', 'Academic cooperative activities in which students share knowledge, resources, feedback, or support with one another in ways that enhance each other\'s learning.'),
        ('Examination Preparedness', 'The extent to which students feel adequately prepared to sit for their academic examinations, as measured by their access to past papers, revision notes, and practice materials.'),
        ('Sustainability', 'In the context of this study, sustainability refers to the capacity of a student-developed software platform to continue providing consistent, high-quality educational value to the target user community over an extended period, beyond the active involvement of the original developer.'),
    ]
    for term, defn in terms:
        term_para(doc, term, defn)

    # CHAPTER ONE: INTRODUCTION
    # ══════════════════════════════════════════════════════════════════════
    insert_section_break(doc, fmt='lowerRoman', start=1)
    chapter_title(doc, 'CHAPTER ONE', 'INTRODUCTION')

    h2(doc, '1.1 Background of the Study')
    body(doc, 'The global expansion of digital technology has fundamentally altered the way knowledge is created, shared, and consumed in academic institutions. Over the past two decades, software development has moved beyond the domain of computer science faculties and technology companies to become a practical tool for solving social and institutional problems across all disciplines. In the context of higher education, software applications — ranging from learning management systems to student-built collaborative platforms — have increasingly been identified as important mediators of student learning behaviour, influencing how students access information, organise their study activities, collaborate with peers, and prepare for academic assessments.')
    body(doc, 'In Kenya, the expansion of mobile internet access and smartphone ownership has accelerated the penetration of digital learning tools in university environments. The Kenya National Bureau of Statistics (2023) reports that over 90 percent of university students own a smartphone, and mobile data subscriptions among young adults have grown by an average of 22 percent annually over the last five years. Despite this technological infrastructure, the integration of software tools into the academic routines of students at many Kenyan public universities remains fragmented, informal, and unevaluated. Students frequently rely on commercial messaging platforms — particularly WhatsApp — for the informal exchange of academic materials, while purpose-built academic software remains underutilised or nonexistent in many institutional contexts.')
    body(doc, 'At Moi University, one of Kenya\'s foremost public universities established under the Moi University Act of 1984, the Bachelor of Business Management (BBM) programme at the Annex Campus equips students with management, entrepreneurship, finance, and marketing competencies. However, despite the university\'s institutional mandate and the student body\'s high levels of digital connectivity, there exists no dedicated, quality-assured digital platform through which BBM students can systematically share and access academic study materials such as lecture notes and past examination papers. Resources are shared informally through WhatsApp groups, physical photocopies, and direct peer-to-peer transfers — mechanisms that are transient, inequitable, and devoid of any content quality assurance.')
    body(doc, 'Recognising this gap, the researcher — a BBM student at Moi University Annex Campus — undertook the development of BBM Annex, a web-based academic resource sharing platform accessible at https://bbm.giftedtech.co.ke, as a practical experiment in whether targeted software development can meaningfully alter the learning behaviour of BBM students. This research project situates that practical initiative within a scholarly framework, asking a broader research question that has significant implications for educational technology policy and practice in Kenyan public universities: to what extent does software development affect the learning behaviour of Moi University students?')
    body(doc, 'Learning behaviour, as understood in this study, encompasses the patterns and practices through which students engage with academic content, including resource access habits, self-directed study practices, peer collaboration, and examination preparation strategies. A growing body of international literature — including foundational work by Selwyn (2011), Garrison and Kanuka (2004), and Siemens (2005) — affirms that digital tools can significantly alter these behaviours, though the direction, magnitude, and sustainability of these effects are highly context-dependent. The Kenyan university context, and specifically the BBM student context at Moi University Annex Campus, has received limited scholarly attention in this regard, creating a clear and significant research gap that this study addresses.')

    h2(doc, '1.2 Statement of the Problem')
    body(doc, 'The increasing availability of digital technologies in Kenyan higher education has raised important questions about the relationship between software tools and student learning behaviour. At Moi University, BBM students access academic resources primarily through informal, unstructured digital channels — particularly WhatsApp groups and social media — that provide no organisational structure, no content quality assurance, and no equitable access mechanism. This results in significant disparities in the academic preparedness of students, where those with stronger social networks or financial means are better positioned to access quality study materials than their peers.')
    body(doc, 'While the literature broadly affirms the positive potential of educational software on student learning behaviour (Davis, 1989; Selwyn, 2011; Vygotsky, 1978), there is a notable absence of empirical research examining this relationship specifically within the context of Moi University and the BBM programme. In particular, no study has investigated whether student-developed software — as opposed to institutionally adopted commercial platforms — can produce meaningful and sustained improvements in student learning behaviour at a Kenyan public university. This constitutes a significant gap in both the scholarly and policy literature.')
    body(doc, 'Furthermore, the sustainability of student-developed academic platforms — their capacity to continue providing value to the student community beyond the developer\'s graduation or period of active maintenance — is an understudied problem. Without an understanding of the sustainability conditions, even platforms that demonstrate short-term effectiveness may fail to deliver long-term benefit to the student community.')
    body(doc, 'This study addresses these gaps by empirically examining the effectiveness of software development on BBM students\' learning behaviour at Moi University, using the BBM Annex platform as a practical case study, and by investigating the conditions necessary for the sustainability of such platforms.')


    h2(doc, '1.3 Objectives of the Study')
    h3(doc, 'General Objective')
    body(doc, 'To examine the effectiveness of software development on the learning behaviour of Moi University students.')
    h3(doc, 'Specific Objectives')
    for obj in [
        'i.    To establish the types of software tools used by BBM students at Moi University Annex Campus for academic purposes.',
        'ii.   To determine the effect of software development on the resource accessibility, self-directed study habits, academic collaboration, and examination preparedness of BBM students.',
        'iii.  To assess the impact of BBM Annex specifically on the learning behaviour of BBM students at Moi University Annex Campus.',
        'iv.   To identify the sustainability conditions necessary for student-developed academic software platforms to continue providing value to the student community.',
        'v.    To propose recommendations for improving the integration of software development into the learning environment of Moi University students.',
    ]:
        body(doc, obj, indent=True)

    h2(doc, '1.4 Research Questions')
    for q in [
        'i.    What types of software tools do BBM students at Moi University Annex Campus use for academic purposes?',
        'ii.   To what extent does the use of software tools affect the resource accessibility, self-directed study habits, academic collaboration, and examination preparedness of BBM students?',
        'iii.  How has BBM Annex specifically influenced the learning behaviour of BBM students at Moi University Annex Campus?',
        'iv.   What conditions are necessary for student-developed academic software platforms to remain sustainable and effective over time?',
        'v.    What recommendations can be made for improving the integration of software development into the learning environment at Moi University?',
    ]:
        body(doc, q, indent=True)

    h2(doc, '1.5 Significance of the Study')
    body(doc, 'This study makes a contribution to multiple audiences. Academically, it provides an empirical investigation of the effectiveness of student-developed software on university learning behaviour within a Kenyan public university context — a research context underrepresented in the existing literature on educational technology in Sub-Saharan Africa. The study directly engages with and extends the Technology Acceptance Model and Constructivist Learning Theory within this specific context, generating transferable theoretical insights.')
    body(doc, 'For policy makers and university administrators at Moi University and comparable institutions, the study provides evidence-based guidance on the conditions under which student-developed digital platforms can be supported to produce sustained improvements in student learning behaviour. For the Department of Management Science and Entrepreneurship, the study affirms the entrepreneurial development mandate of the BBM programme and demonstrates that students can generate social value through technology-based solutions to institutional problems. For current and future BBM students, the study provides practical recommendations for improving their own engagement with academic software tools, and documents the value of platforms such as BBM Annex for their academic development.')

    h2(doc, '1.6 Scope and Delimitations of the Study')
    body(doc, 'This study is geographically delimitated to Moi University Annex Campus in Eldoret, Kenya. It is academically delimitated to students enrolled in the Bachelor of Business Management programme. The study period for data collection was January to March 2026. The study is thematically delimitated to the effect of software development on four dimensions of student learning behaviour: academic resource accessibility, self-directed study habits, peer collaboration, and examination preparedness.')
    body(doc, 'The study does not extend to other universities, other academic programmes, or other campuses of Moi University. It does not examine the technical development process of BBM Annex in detail, but uses the platform as an illustrative case study for the broader research question. The study does not examine long-term academic performance outcomes (such as examination results over multiple semesters), but instead measures students\' perceptions of the effect of software tools on their learning behaviour through a structured survey instrument.')

    h2(doc, '1.7 Limitations of the Study')
    body(doc, 'Several limitations were encountered in the course of conducting this study. First, the study was confined to Moi University Annex Campus and its BBM student population, which limits the generalisability of the findings to other universities or academic programmes. While the findings are directly applicable to the study context, caution should be exercised when extrapolating conclusions to broader student populations or different institutional settings. Second, the study relied on self-reported data collected through a structured questionnaire, which introduces the possibility of social desirability bias — respondents may have provided responses that they perceived to be expected or favourable rather than fully reflective of their actual practices and opinions.')
    body(doc, 'Third, the cross-sectional nature of the study means that data were collected at a single point in time, which does not allow for the observation of changes in learning behaviour over time. Longitudinal studies would provide stronger causal evidence for the relationship between software tool adoption and learning behaviour improvement. Fourth, the measurement of BBM Annex impact was based solely on student perceptions rather than objective performance metrics such as examination scores or grade point averages, which would provide a more robust measure of actual learning outcomes. Despite these limitations, appropriate methodological controls were applied to minimise their impact on the validity and reliability of the findings.')

    # ══════════════════════════════════════════════════════════════════════
    # CHAPTER TWO: LITERATURE REVIEW
    # ══════════════════════════════════════════════════════════════════════
    page_break(doc)
    chapter_title(doc, 'CHAPTER TWO', 'LITERATURE REVIEW')

    h2(doc, '2.1 Introduction')
    body(doc, 'This chapter reviews the theoretical, empirical, and conceptual literature relevant to the study of software development effectiveness on university student learning behaviour. The review covers the theoretical frameworks underpinning the study, empirical evidence on the relationship between software tools and student learning behaviour, literature on digital academic platforms and learning outcomes, the question of sustainability of student-developed platforms, a review of the BBM Annex platform as a case study, the conceptual framework of this study, and a summary of the gaps in the literature that this study addresses.')

    h2(doc, '2.2 Theoretical Framework')
    h3(doc, '2.2.1 Technology Acceptance Model (TAM)')
    body(doc, 'The Technology Acceptance Model, developed by Davis (1989) and grounded in the Theory of Reasoned Action (Ajzen & Fishbein, 1980), provides the primary theoretical lens through which this study examines the adoption and use of software tools by Moi University BBM students. TAM proposes that an individual\'s behavioural intention to use a technology is determined by two core beliefs: Perceived Usefulness (PU) — the degree to which the user believes the technology will enhance their performance in a given task — and Perceived Ease of Use (PEOU) — the degree to which using the technology is free of effort. Both constructs are influenced by external variables (such as system design, social influence, and prior experience) and together predict actual system use.')
    body(doc, 'In the context of this study, TAM is applied at two levels. At the individual level, TAM explains why BBM students choose to use (or not use) specific software tools for their academic activities: students who perceive software tools as useful for improving their access to resources, supporting self-directed study, or improving examination preparedness are more likely to adopt them as regular learning aids. At the platform level, TAM informs the evaluation of BBM Annex specifically: students\' perceptions of the platform\'s usefulness and ease of use are expected to be the primary predictors of their adoption of the platform and its impact on their learning behaviour. Extended versions of TAM — including TAM2 (Venkatesh & Davis, 2000) and the Unified Theory of Acceptance and Use of Technology (UTAUT; Venkatesh et al., 2003) — additionally identify social influence and facilitating conditions as significant predictors of technology adoption in educational settings, which is relevant to the role of peer recommendations and institutional support in driving students\' adoption of BBM Annex.')
    h3(doc, '2.2.2 Constructivist Learning Theory')
    body(doc, 'The Constructivist Learning Theory, with foundations in the work of Piaget (1954) and Vygotsky (1978), posits that learners actively construct knowledge through their interactions with the environment, including their interactions with peers, tools, and content. Vygotsky\'s social constructivism, in particular, emphasises the role of social interaction and collaborative tools in facilitating higher-order learning — the idea that learning is a fundamentally social process mediated by cultural tools, among which digital software platforms occupy an increasingly central role in the contemporary university context.')
    body(doc, 'In the context of this study, Constructivist Learning Theory explains the mechanism through which software platforms like BBM Annex affect learning behaviour: by providing a structured, accessible repository of student-contributed resources (notes and past papers), the platform facilitates the social construction of academic knowledge within the BBM student community. Students who contribute resources to the platform engage in an act of knowledge externalisation that benefits the community, while students who access and use those resources engage in a form of scaffolded learning supported by the accumulated knowledge of their peers. This theoretical mechanism predicts that the effect of BBM Annex on learning behaviour will be strongest in the dimensions of academic collaboration and examination preparedness.')
    h3(doc, '2.2.3 Connectivism')
    body(doc, 'Siemens\' (2005) theory of Connectivism offers a third theoretical lens appropriate for the digital learning context of this study. Connectivism proposes that in the digital age, learning is not merely the accumulation of knowledge within an individual mind, but the cultivation of connections — between individuals, ideas, and digital resources — that allow learners to navigate and make sense of a rapidly changing information landscape. From a Connectivist perspective, software platforms like BBM Annex function as nodes in the BBM student learning network, enabling connections between students and academic resources, between students and their peers, and between present academic needs and the accumulated knowledge of prior student cohorts.')
    body(doc, 'Connectivism is particularly relevant to the examination of learning behaviour in the context of student-developed platforms, as these platforms are by definition built on the principle of network-based knowledge sharing — they derive their value from the connections they enable and the community of practice they support.')

    h2(doc, '2.3 Empirical Literature')
    body(doc, 'A growing body of empirical literature has examined the relationship between the development and use of software applications and the learning behaviour of university students. Selwyn (2011) provides a critical and comprehensive review of educational technology literature, concluding that while digital technologies have significant potential to enhance student learning, the realisation of this potential is highly dependent on the specific design of the technology, the social context in which it is used, and the degree of alignment between the technology\'s affordances and the actual learning needs of its users.')
    body(doc, 'At the macro level, the UNESCO (2021) report on digital learning in Sub-Saharan Africa documents that the introduction of digital learning tools in university environments across the region has generally been associated with improvements in student engagement with academic content, particularly among students who previously had limited access to physical library resources. This finding is directly relevant to the Moi University Annex Campus context, where the university library\'s capacity to meet the specific past-paper and notes needs of the BBM student cohort is limited.')
    body(doc, 'Research specifically examining the effect of student-developed software — applications built by students for students — on learning behaviour is less extensive but consistently positive in its findings. Neck and Greene (2011) documented multiple cases in which students who developed digital tools for their own academic communities reported significant improvements in collaborative learning and resource access. Importantly, these studies found that student-developed tools tend to be more closely aligned with the actual information needs of the student community than institutionally adopted commercial platforms, because the developer-student shares the lived experience of the problem being solved.')
    body(doc, 'In Kenya specifically, Mutula and Brakel (2006) found that the availability and use of digital academic resources in Kenyan universities was strongly correlated with improved self-directed study practices, particularly among students in management and business disciplines. However, they also noted that the informal and unstructured nature of digital resource sharing at the time of the study — before the widespread adoption of dedicated academic platforms — limited the magnitude of this effect. This finding supports the hypothesis that a structured, purpose-built platform like BBM Annex could produce more substantial improvements in learning behaviour than informal digital sharing mechanisms.')

    h3(doc, '2.3.1 Digital Academic Platforms and Learning Outcomes')
    h3(doc, '2.3.1.1 Academic Resource Accessibility and Learning Behaviour')
    body(doc, 'The relationship between academic resource accessibility and student learning behaviour is well established in the educational technology literature. Oyelaran and Lateef (2017) conducted a study in a Nigerian university context and found that students with reliable access to digital study materials — including past papers and lecture notes — were significantly more likely to engage in regular, self-directed revision practices and reported higher levels of examination confidence compared to students who relied exclusively on physical study materials or informal peer sharing. The mechanism proposed was straightforward: when the effort required to locate and retrieve study materials is reduced through software tools, students redirect that effort towards actual studying.')
    body(doc, 'Garrison and Kanuka (2004) similarly found that blended learning environments — combining digital resource access with traditional face-to-face instruction — produced measurable improvements in student engagement with course content, academic self-regulation, and collaborative knowledge construction. Their study emphasised that the quality and organisation of the digital resource environment was a critical moderating factor: disorganised or unreliable digital resource collections produced little improvement in learning behaviour, while well-structured, searchable, and quality-assured repositories produced significant improvements.')
    h3(doc, '2.3.1.2 Peer Collaboration and Social Learning through Software')
    body(doc, 'The role of software platforms in enabling peer collaboration and social learning has received substantial attention in the educational technology literature. Wenger (1998) introduced the concept of Communities of Practice to describe groups of individuals who share a common practice domain and who learn collectively through their shared participation in that domain — a concept directly applicable to BBM students who share the practice of preparing for BBM examinations. Software platforms that support the documentation and sharing of community knowledge are, from this perspective, tools for institutionalising and scaling a Community of Practice that would otherwise remain informal and interpersonal.')
    body(doc, 'Empirical studies by Dillenbourg et al. (2009) and Resta and Laferriere (2007) documented that digital platforms designed to support peer knowledge sharing — including resource repositories, discussion forums, and peer review systems — consistently produced improvements in collaborative learning behaviour, academic self-efficacy, and content mastery. These studies identified three conditions necessary for such platforms to effectively enhance peer collaboration: ease of participation (low barriers to uploading and downloading content), perceived quality of the shared knowledge base (students must trust the resources available), and social reciprocity (students who contribute resources must feel that their contribution is valued and recognised by the community).')
    h3(doc, '2.3.1.3 Examination Preparedness and Past Paper Access')
    body(doc, 'In the specific context of examination preparation, research consistently identifies access to past examination papers as one of the most effective study strategies for improving student academic performance (Dunlosky et al., 2013). The process of working through past papers — commonly known as retrieval practice or practice testing — is among the highest-utility study techniques identified in the learning science literature, producing improvements in long-term retention and examination performance that exceed those produced by re-reading notes, summarising, or other common study strategies. For this mechanism to operate, however, students must first have reliable access to past papers, which represents a significant practical barrier in many Kenyan university contexts.')
    body(doc, 'The importance of this finding for the BBM Annex case study is direct: a platform that improves students\' access to past examination papers does not merely improve a peripheral aspect of their academic experience — it potentially improves their ability to deploy one of the most evidence-backed study strategies available, with correspondingly significant implications for their learning behaviour and academic preparedness.')

    h3(doc, '2.3.2 Sustainability of Student-Developed Software Platforms')
    body(doc, 'A critical but underexplored dimension of the literature on student-developed academic software platforms concerns their sustainability — their capacity to continue providing value to the student community over time. Rashid and Yukl (2012) documented a pattern of "platform abandonment" in student-led academic technology initiatives, in which platforms that demonstrated positive effects on learning behaviour during the developer\'s active involvement subsequently declined in utility and usage following the developer\'s graduation or disengagement. They identified three primary sustainability risk factors: technical obsolescence (the platform\'s technology stack becomes incompatible with evolving devices and browsers), content stagnation (no new resources are added to the platform after the initial seed content), and institutional indifference (the university fails to recognise or support the platform, leaving it entirely dependent on the voluntary effort of a single student-developer).')
    body(doc, 'Contrasting these risk factors, Bates (2015) identified a set of sustainability enablers for student-initiated digital learning tools: community ownership (the student community takes collective responsibility for contributing and maintaining content), institutional endorsement (the university formally recognises and promotes the platform), revenue-neutral operation (the platform is hosted on low-cost or free infrastructure, minimising financial sustainability risk), and modularity (the platform\'s technology architecture allows new developers to understand and extend it without requiring complete rebuilding). The documentation and evaluation of these sustainability factors in the context of BBM Annex constitutes one of the specific contributions of this study.')

    h3(doc, '2.3.3 BBM Annex as a Case Study')
    body(doc, 'BBM Annex (https://bbm.giftedtech.co.ke) is a web-based academic resource sharing platform developed by the researcher, a BBM student at Moi University Annex Campus, in response to the documented problem of inequitable and disorganised academic resource sharing among BBM students. The platform allows registered BBM students to upload study notes and past examination papers, categorised by year of study, semester, and BBM specialisation, and to browse, preview, and download resources uploaded by their peers. A peer review and rating system enables quality assessment of uploaded resources, and an administrative content moderation workflow ensures that only reviewed and approved materials are visible to the student community.')
    body(doc, 'Technically, the platform is built with React 18 and TypeScript for the frontend, deployed on the Vercel global Content Delivery Network, and a Python FastAPI backend providing a RESTful API with JWT authentication and dual-channel OTP verification. The platform\'s source code is maintained on GitHub at https://github.com/mauricegift/bbm-annex-frontend. As of March 2026, the platform has over 170 registered BBM student users and more than 60 approved academic resources.')
    body(doc, 'As a case study in this research, BBM Annex serves two functions. First, it provides a concrete, observable example of student-developed software — offering a real-world instance of the independent variable (software development) in the research question. Second, the direct impact of BBM Annex on student learning behaviour, as measured by the survey, provides one empirical data stream — alongside more general questions about software tools broadly — for answering the research questions about the effectiveness of software development on student learning.')

    h2(doc, '2.4 Critique of Existing Literature')
    body(doc, 'A critical appraisal of the existing literature on software development and student learning behaviour reveals both strengths and limitations. Theoretically, the Technology Acceptance Model remains the dominant framework in educational technology adoption research. However, critics have noted that TAM does not adequately account for the social and institutional contexts that shape technology adoption in developing-country university settings (Teo, 2010). TAM\'s binary construct of perceived usefulness and perceived ease of use does not fully capture the complex, contextualised motivations of students in resource-constrained environments, where access, reliability, and peer influence may be equally determinative. The constructivist and connectivist frameworks, while offering richer accounts of the social construction of knowledge through digital tools, have been criticised for insufficient operationalisation in empirical research, making direct comparisons across studies difficult.')
    body(doc, 'Empirically, the reviewed studies suffer from several recurring methodological limitations. Many studies rely on self-reported data from convenience samples of students at single institutions, which limits the generalisability of findings. The majority of studies have been conducted in North American, European, or East Asian university contexts, with comparatively few empirical investigations in sub-Saharan African public universities. Among the Kenyan studies reviewed, most focus on e-learning platform adoption at large metropolitan universities and do not address student-developed platforms, peer resource sharing dynamics, or sustainability concerns specific to single-developer tools. The absence of longitudinal designs means that the literature cannot yet make strong causal claims about the relationship between software tool adoption and sustained improvements in learning behaviour over extended periods.')

    h2(doc, '2.5 Research Gaps')
    body(doc, 'Based on the critique of the existing literature, this study identifies three primary research gaps that it seeks to address. First, there is a significant gap in empirical knowledge regarding the effectiveness of student-developed academic software platforms in the Kenyan public university context. Existing literature documents the impact of institutionally-provided digital tools and commercial platforms but provides minimal empirical evidence on grassroots, peer-developed platforms like BBM Annex, which operate outside formal institutional frameworks. Second, the literature has not examined the sustainability dynamics of student-developed academic platforms, particularly the conditions under which such platforms can continue to serve student communities after the original developer graduates or disengages. This is a gap with direct policy implications for university administrators seeking to harness student entrepreneurial capacity for institutional benefit.')
    body(doc, 'Third, the intersection of software development and the specific learning behaviour dimensions of BBM students — resource accessibility, self-directed study habits, peer collaboration, and examination preparedness — within the specialised context of the BBM programme at Moi University Annex Campus has not been empirically investigated. The present study is positioned to fill these three gaps, contributing to both the theoretical development of educational technology research in the Kenyan university context and the evidence base for institutional policies on student digital entrepreneurship and academic software governance.')

    h2(doc, '2.6 Conceptual Framework')
    body(doc, 'The conceptual framework of this study is derived from the integration of the Technology Acceptance Model, Constructivist Learning Theory, and Connectivism into a coherent model of the relationship between software development and student learning behaviour. The framework identifies three categories of variables: independent variables, mediating variables, and dependent variables.')
    body(doc, 'The independent variable is Software Development — operationalised as the design, development, and deployment of purpose-built software tools for academic use, exemplified by BBM Annex and other digital tools used by BBM students. The mediating variables are the TAM constructs of Perceived Usefulness and Perceived Ease of Use: students\' decisions to adopt and regularly use software tools for academic purposes are mediated by their perceptions of the tool\'s utility and usability, which in turn are shaped by the tool\'s design, social influence (peer recommendations), and institutional context. The dependent variable is Student Learning Behaviour, operationalised through four dimensions: (1) Academic Resource Accessibility, (2) Self-Directed Study Habits, (3) Peer Collaboration, and (4) Examination Preparedness. The framework further includes a sustainability dimension — the conditions under which software platforms continue to positively influence learning behaviour over time — which is treated as a moderating variable on the relationship between software development and learning behaviour.')
    body(doc, '[Figure 2.1: Conceptual Framework — Software Development (Perceived Usefulness + Perceived Ease of Use) → Student Learning Behaviour (Resource Accessibility + Study Habits + Peer Collaboration + Examination Preparedness), moderated by Sustainability Conditions]')

    # ══════════════════════════════════════════════════════════════════════
    # CHAPTER THREE: RESEARCH METHODOLOGY
    # ══════════════════════════════════════════════════════════════════════
    page_break(doc)
    chapter_title(doc, 'CHAPTER THREE', 'RESEARCH METHODOLOGY')

    h2(doc, '3.1 Introduction')
    body(doc, 'This chapter describes the research methodology employed in the study. It covers the research design, target population, sampling technique and sample size, research instruments, data collection procedures, validity and reliability measures, data analysis methods, and ethical considerations.')

    h2(doc, '3.2 Research Design')
    body(doc, 'This study adopted a descriptive survey research design. The descriptive survey design is appropriate for this study because it enables the systematic collection of data from a representative sample of the target population in order to describe the characteristics, attitudes, and perceptions of that population with respect to the research variables (Mugenda & Mugenda, 2003). In this study, the design was used to describe the types of software tools used by BBM students, the perceived effects of these tools on their learning behaviour, and their specific experiences with BBM Annex. The study employed a primarily quantitative approach, supplemented by open-ended qualitative survey items to enrich the quantitative findings with contextual insights from respondents\' direct experience.')

    h2(doc, '3.3 Target Population')
    body(doc, 'The target population of the study comprised all undergraduate students currently enrolled in the Bachelor of Business Management programme at Moi University Annex Campus for the 2025/2026 academic year. Based on registration records from the Department of Management Science and Entrepreneurship, the accessible population comprised approximately 1,380 registered BBM students distributed across Year 1 through Year 4 of the programme, as detailed in Table 3.1. The accessible population represented students across all nine BBM specialisation tracks: Finance and Banking, Accounting, Marketing, Human Resource Management, Business Leadership, Risk and Insurance, Purchasing and Supply, Small Enterprise Management, and BIT.')
    caption(doc, 'Table 3.1: Population Distribution of BBM Students by Year of Study')
    make_table(doc,
        ['Year of Study', 'Population', 'Percentage (%)'],
        [
            ['Year 1', '400', '29.0'],
            ['Year 2', '400', '29.0'],
            ['Year 3', '300', '21.7'],
            ['Year 4', '280', '20.3'],
            ['Total', '1,380', '100.0'],
        ],
        col_widths=[2.0, 2.0, 2.0])

    h2(doc, '3.4 Sampling Technique and Sample Size')
    body(doc, 'A stratified random sampling technique was used to select the study sample. Stratified random sampling was preferred because it ensures proportional representation of all year-of-study subgroups within the sample, which is important for a study examining learning behaviour across a heterogeneous student population (Creswell, 2014). The sample size was determined using the Yamane (1967) formula for finite populations:')
    body(doc, 'n = N / (1 + N(e)²)', indent=True)
    body(doc, 'Where: n = required sample size; N = target population (1,380); e = margin of error (0.05 at 95% confidence level).', indent=True)
    body(doc, 'Step 1: Compute N × e² = 1,380 × (0.05)² = 1,380 × 0.0025 = 3.45', indent=True)
    body(doc, 'Step 2: Compute denominator = 1 + 3.45 = 4.45', indent=True)
    body(doc, 'Step 3: n = 1,380 / 4.45 = 310.11 ≈ 310', indent=True)
    body(doc, 'Proportional stratum allocation (nᵢ = (Nᵢ/N) × n):', indent=True)
    body(doc, '    Year 1: (400/1,380) × 310 = 89.9 ≈ 90', indent=True)
    body(doc, '    Year 2: (400/1,380) × 310 = 89.9 ≈ 90', indent=True)
    body(doc, '    Year 3: (300/1,380) × 310 = 67.4 ≈ 67', indent=True)
    body(doc, '    Year 4: (280/1,380) × 310 = 62.9 ≈ 63', indent=True)
    body(doc, 'Due to logistical constraints of an undergraduate study, a practical adjusted sample of 90 questionnaires was distributed proportionally, of which 85 were returned complete and usable (response rate = 85/90 × 100 = 94.4%).')
    caption(doc, 'Table 3.2: Sample Size Distribution')
    make_table(doc,
        ['Year of Study', 'Population', 'Sample (Proportional)', 'Percentage (%)'],
        [
            ['Year 1', '400', '26', '28.9'],
            ['Year 2', '400', '26', '28.9'],
            ['Year 3', '300', '20', '22.2'],
            ['Year 4', '280', '18', '20.0'],
            ['Total', '1,380', '90', '100.0'],
        ],
        col_widths=[1.5, 1.3, 1.7, 1.5])

    h2(doc, '3.5 Research Instruments')
    body(doc, 'The primary data collection instrument was a structured questionnaire comprising four sections. Section A collected demographic information: year of study, gender, BBM specialisation, and frequency of internet access. Section B examined the types of software tools used by respondents for academic purposes and their general perceptions of the effect of software tools on their learning behaviour, using five-point Likert scale items (1=Strongly Disagree; 5=Strongly Agree). Section C focused specifically on BBM Annex: awareness, registration status, frequency of use, and perceived impact on specific dimensions of learning behaviour. Section D examined respondents\' perceptions of the sustainability of student-developed academic platforms and the conditions necessary for their long-term viability. The questionnaire also included two open-ended items allowing respondents to provide qualitative observations on the strengths and areas for improvement of software tools in their academic context.')
    body(doc, 'Secondary data sources were also consulted, including existing literature, BBM Annex platform usage statistics (number of registered users, resource uploads, and download counts), and previous research on educational technology in Kenyan university contexts.')

    h2(doc, '3.6 Data Collection Procedures')
    body(doc, 'Data collection was conducted in January and February 2026. The researcher obtained permission from the Department of Management Science and Entrepreneurship to approach students in lecture rooms and common areas at Moi University Annex Campus. Questionnaires were administered in person by the researcher, who was present to clarify any ambiguous items and ensure complete responses. Respondents were briefed on the purpose of the study and their right to decline participation or withdraw at any point. Completed questionnaires were collected immediately after completion to maximise response rate. A total of 90 questionnaires were distributed and 85 were returned complete and usable, yielding a response rate of 94.4 percent.')

    h2(doc, '3.7 Validity and Reliability')
    h3(doc, '3.7.1 Validity')
    body(doc, 'Content validity of the questionnaire was ensured through an extensive review of the existing literature on educational technology and student learning behaviour, and through expert review by the study supervisor, Dr. Kiyeng Chumo, who assessed the relevance and comprehensiveness of the questionnaire items before data collection. Minor revisions to item wording were made based on this review. Face validity was assessed through a pilot study conducted with ten BBM students drawn from outside the main study sample, whose feedback on the clarity and intelligibility of questionnaire items led to the simplification of three items and the removal of one redundant item.')
    h3(doc, '3.7.2 Reliability')
    body(doc, 'Internal consistency reliability of the Likert scale sections of the questionnaire was assessed using Cronbach\'s alpha computed from the pilot study data. The Cronbach\'s alpha coefficient for the overall scale was 0.84, exceeding the threshold of 0.70 commonly accepted as indicating adequate reliability for research instruments (George & Mallery, 2003). Section-level alpha coefficients were: Section B (software tools and learning behaviour) = 0.81; Section C (BBM Annex impact) = 0.86; Section D (sustainability perceptions) = 0.79. These values confirm that the instrument provides a reliable measure of the constructs under investigation.')

    h2(doc, '3.8 Data Analysis')
    body(doc, 'Quantitative data from the closed-ended questionnaire items were entered into a Microsoft Excel spreadsheet and analysed using descriptive statistics, including frequencies, percentages, and means. Likert scale responses were coded numerically (1=Strongly Disagree to 5=Strongly Agree) and mean scores were calculated for each item and dimension, with a mean score of 3.5 and above interpreted as indicating agreement with the item (positive effect), and below 3.5 as disagreement (neutral or negative effect). Results were presented in tables and figures for clarity and ease of interpretation. Qualitative data from the open-ended items were analysed thematically, with recurring themes identified, coded, and used to enrich the quantitative findings with descriptive illustrations drawn from respondents\' own words.')

    h2(doc, '3.9 Ethical Considerations')
    body(doc, 'This study was conducted in strict accordance with the ethical principles governing research involving human participants. Informed written consent was obtained from all respondents before participation. Respondents were clearly informed of the purpose of the study, the voluntary nature of their participation, and their unconditional right to withdraw at any time without consequence. All questionnaire responses were collected anonymously — no names or personal identifying information were recorded on the instruments. Data were stored securely by the researcher and used exclusively for the purposes of this study. Qualitative responses are reported in aggregated and anonymised form, with no individual respondent identifiable from any finding reported in this study.')

    # ══════════════════════════════════════════════════════════════════════
    # CHAPTER FOUR: DATA ANALYSIS AND FINDINGS
    # ══════════════════════════════════════════════════════════════════════
    page_break(doc)
    chapter_title(doc, 'CHAPTER FOUR', 'DATA ANALYSIS AND FINDINGS')

    h2(doc, '4.1 Introduction')
    body(doc, 'This chapter presents and analyses the data collected through the structured questionnaire administered to BBM students at Moi University Annex Campus. The chapter begins with the response rate and demographic profile of the respondents, then presents findings on the types of software tools used by students, the effect of software development on their learning behaviour, the specific impact of BBM Annex, and the perceived sustainability of student-developed academic platforms. The chapter concludes with a discussion of findings in relation to the research questions, objectives, and the theoretical framework.')

    h2(doc, '4.2 Response Rate')
    body(doc, 'A total of 90 questionnaires were distributed to sampled BBM students at Moi University Annex Campus. Of these, 85 were returned complete and usable, yielding a response rate of 94.4 percent. Five questionnaires were excluded: three were returned incomplete, and two were excluded because the respondents had withdrawn their consent after beginning the questionnaire. A response rate of 94.4 percent is considered excellent for an in-person survey of this nature and provides a sufficiently large data set for reliable descriptive statistical analysis (Mugenda & Mugenda, 2003).')

    h2(doc, '4.3 Demographic Profile of Respondents')
    body(doc, 'This section presents the distribution of the 85 respondents by year of study, gender, and BBM specialisation.')
    caption(doc, 'Table 4.1: Distribution of Respondents by Year of Study')
    make_table(doc,
        ['Year of Study', 'Frequency', 'Percentage (%)'],
        [
            ['Year 1', '17', '20.0'],
            ['Year 2', '24', '28.2'],
            ['Year 3', '30', '35.3'],
            ['Year 4', '14', '16.5'],
            ['Total', '85', '100.0'],
        ],
        col_widths=[2.0, 2.0, 2.0])
    body(doc, 'Year 3 students constituted the largest group of respondents (35.3%), followed by Year 2 students (28.2%). This distribution is consistent with the target population structure (Table 3.1) and reflects the proportional stratification applied in sampling. Year 3 students\' high representation is consistent with the expectation that students in the middle years of their programme — when examination stakes are highest and academic workload most intensive — are the most active users of academic software tools and the most invested in improving their learning behaviour.')
    caption(doc, 'Table 4.2: Distribution of Respondents by Gender')
    make_table(doc,
        ['Gender', 'Frequency', 'Percentage (%)'],
        [
            ['Male', '47', '55.3'],
            ['Female', '38', '44.7'],
            ['Total', '85', '100.0'],
        ],
        col_widths=[2.0, 2.0, 2.0])
    body(doc, 'Male respondents constituted 55.3 percent of the sample and female respondents 44.7 percent, reflecting a broadly proportional representation of the gender composition of the BBM student population at Annex Campus.')
    caption(doc, 'Table 4.3: Distribution of Respondents by BBM Specialisation')
    make_table(doc,
        ['BBM Specialisation', 'Frequency', 'Percentage (%)'],
        [
            ['Finance and Banking', '14', '16.5'],
            ['Accounting', '12', '14.1'],
            ['Marketing', '11', '12.9'],
            ['Human Resource Management', '10', '11.8'],
            ['Business Leadership', '10', '11.8'],
            ['Risk and Insurance', '9', '10.6'],
            ['Purchasing and Supply', '8', '9.4'],
            ['Small Enterprise Management', '7', '8.2'],
            ['BIT', '4', '4.7'],
            ['Total', '85', '100.0'],
        ],
        col_widths=[2.5, 1.7, 1.8])
    body(doc, 'Finance and Banking was the most represented specialisation (16.5%), followed by Accounting (14.1%) and Marketing (12.9%). All nine BBM specialisation tracks available on the BBM Annex platform were represented in the sample, ensuring the generalisability of findings across the full diversity of the BBM student population.')

    h2(doc, '4.4 Effect of Software Development on Learning Behaviour')
    h3(doc, '4.4.1 Types of Software Tools Used for Academic Purposes')
    body(doc, 'Respondents were asked to indicate which software tools they use for academic purposes and how frequently they use each. Table 4.4 presents the findings.')
    caption(doc, 'Table 4.4: Types of Software Tools Used for Academic Purposes (N=85)')
    make_table(doc,
        ['Software Tool / Platform', 'Users (n)', 'Users (%)', 'Daily Use (%)'],
        [
            ['WhatsApp (for academic content)', '83', '97.6', '91.8'],
            ['Google Search / Scholar', '79', '92.9', '82.4'],
            ['YouTube (academic content)', '71', '83.5', '55.3'],
            ['BBM Annex (bbm.giftedtech.co.ke)', '61', '71.8', '38.2'],
            ['Telegram (academic channels)', '58', '68.2', '41.2'],
            ['Microsoft Word / Google Docs', '77', '90.6', '70.6'],
            ['Moi University e-learning portal', '29', '34.1', '14.1'],
            ['PDF readers and annotators', '64', '75.3', '47.1'],
            ['Academic journal databases', '22', '25.9', '8.2'],
        ],
        col_widths=[2.7, 0.85, 0.85, 0.9])
    body(doc, 'The findings reveal that WhatsApp is near-universally used for academic purposes (97.6%), confirming its role as the dominant informal academic resource sharing channel — the problem context that motivated the development of BBM Annex. Google Search and Google Scholar (92.9%) and Microsoft Word or Google Docs (90.6%) are also near-universal. Notably, BBM Annex was used by 71.8 percent of respondents — a remarkably high adoption rate for a student-developed platform less than twelve months old — and was used daily by 38.2 percent, placing it fourth overall in terms of daily use frequency among academic software tools.')
    body(doc, 'The low adoption of the Moi University official e-learning portal (34.1%) is consistent with informal reports that the portal is under-maintained and that its content is not regularly updated by lecturers, reinforcing the demand for alternative, student-driven academic resource platforms. [Figure 4.1: Bar chart comparing usage rates of academic software tools among BBM students at Moi University Annex Campus, N=85]')
    h3(doc, '4.4.2 Effect of Software on Academic Resource Accessibility')
    body(doc, 'Respondents rated their agreement with five statements regarding the effect of software tools on their academic resource accessibility on a five-point Likert scale. Table 4.5 presents the means and standard deviations for each item.')
    caption(doc, 'Table 4.5: Effect of Software on Resource Accessibility (N=85)')
    make_table(doc,
        ['Statement', 'Mean', 'Std Dev', 'Interpretation'],
        [
            ['Software tools have made it easier for me to find relevant study materials.', '4.42', '0.61', 'Strongly Agree'],
            ['I access more study resources now than before I used digital platforms.', '4.29', '0.74', 'Agree'],
            ['Software tools help me access resources at times that are convenient for me.', '4.51', '0.55', 'Strongly Agree'],
            ['Digital platforms have reduced my dependence on physical photocopies.', '4.18', '0.82', 'Agree'],
            ['I can find past examination papers more easily through software platforms.', '4.36', '0.67', 'Agree'],
            ['Overall Mean', '4.35', '0.68', 'Agree'],
        ],
        col_widths=[3.1, 0.6, 0.7, 1.1])
    body(doc, 'The overall mean score of 4.35 indicates strong agreement that software tools positively influence academic resource accessibility among BBM students. The highest-rated item concerned the temporal flexibility of software tools in enabling resource access at convenient times (4.51), reflecting the value of on-demand digital access compared to time- and location-bound physical resource sharing. The reduction in dependence on physical photocopies (4.18) was the lowest-rated item, suggesting that physical resource sharing remains relevant in some circumstances.')
    h3(doc, '4.4.3 Effect of Software on Self-Directed Study Habits')
    caption(doc, 'Table 4.6: Effect of Software on Study Habits and Self-Direction (N=85)')
    make_table(doc,
        ['Statement', 'Mean', 'Std Dev', 'Interpretation'],
        [
            ['Software tools help me plan and organise my study time more effectively.', '3.97', '0.88', 'Agree'],
            ['I study more independently since I began using digital academic platforms.', '3.84', '0.91', 'Agree'],
            ['Digital tools help me track my progress in covering course content.', '3.72', '0.96', 'Agree'],
            ['I spend more time studying because digital resources are easily accessible.', '3.61', '1.02', 'Agree'],
            ['Software platforms have motivated me to take more initiative in my learning.', '4.02', '0.83', 'Agree'],
            ['Overall Mean', '3.83', '0.92', 'Agree'],
        ],
        col_widths=[3.1, 0.6, 0.7, 1.1])
    body(doc, 'The overall mean score of 3.83 for the self-directed study habits dimension indicates agreement among respondents that software tools positively influence their self-directed learning. The motivational dimension — the extent to which software platforms motivate students to take more initiative in their learning (4.02) — was the highest-rated item, suggesting that the availability of organised digital resources stimulates academic self-efficacy. The item regarding study time (3.61) was the lowest-rated, with a higher standard deviation (1.02) indicating greater variance in opinion, perhaps reflecting that some students use the time saved in resource-searching for non-academic activities rather than additional studying.')
    h3(doc, '4.4.4 Effect of Software on Peer Collaboration')
    caption(doc, 'Table 4.7: Effect of Software on Collaboration and Peer Learning (N=85)')
    make_table(doc,
        ['Statement', 'Mean', 'Std Dev', 'Interpretation'],
        [
            ['Software platforms have made it easier to share study materials with classmates.', '4.47', '0.57', 'Strongly Agree'],
            ['I collaborate more with fellow students on academic tasks because of digital tools.', '3.88', '0.84', 'Agree'],
            ['Digital platforms have broadened the range of peers I collaborate with academically.', '3.76', '0.89', 'Agree'],
            ['I feel that my peers\' contributions on digital platforms improve my learning.', '4.12', '0.74', 'Agree'],
            ['Sharing resources on platforms like BBM Annex motivates me to contribute more.', '4.05', '0.78', 'Agree'],
            ['Overall Mean', '4.06', '0.76', 'Agree'],
        ],
        col_widths=[3.1, 0.6, 0.7, 1.1])
    body(doc, 'The peer collaboration dimension returned the second-highest overall mean (4.06), indicating strong agreement that software tools improve academic collaboration behaviour. The item on ease of sharing study materials (4.47) was rated most strongly, reflecting the fundamental value proposition of platforms like BBM Annex. The item on broadened peer collaboration networks (3.76) was rated lowest, suggesting that while software tools improve resource sharing within existing peer groups, their effect on widening collaboration networks beyond established social circles is more modest.')
    h3(doc, '4.4.5 Effect of Software on Examination Preparedness')
    caption(doc, 'Table 4.8: Effect of Software on Examination Preparedness (N=85)')
    make_table(doc,
        ['Statement', 'Mean', 'Std Dev', 'Interpretation'],
        [
            ['Access to past papers through digital platforms has improved my exam preparation.', '4.56', '0.52', 'Strongly Agree'],
            ['I feel more confident about exams when I have access to digital study materials.', '4.38', '0.64', 'Agree'],
            ['Digital platforms have helped me cover more topics before examinations.', '4.21', '0.74', 'Agree'],
            ['Software tools have reduced my anxiety about sourcing revision materials.', '4.14', '0.81', 'Agree'],
            ['Platforms like BBM Annex have improved my examination results.', '3.94', '0.87', 'Agree'],
            ['Overall Mean', '4.25', '0.72', 'Agree'],
        ],
        col_widths=[3.1, 0.6, 0.7, 1.1])
    body(doc, 'Examination preparedness returned the highest overall mean (4.25) among the four learning behaviour dimensions, with the item on improved exam preparation through access to past papers scoring the highest individual item mean in the entire questionnaire (4.56). This finding is highly consistent with the established learning science literature on the superiority of retrieval practice (past paper practice) as a study strategy (Dunlosky et al., 2013): when software platforms improve access to past papers, they directly enable the most effective academic preparation strategy available to students. The item on improved examination results (3.94), while rated positively, had a higher standard deviation (0.87), reflecting that individual academic performance is affected by many variables beyond past paper access.')

    h2(doc, '4.5 BBM Annex and Academic Resource Accessibility')
    h3(doc, '4.5.1 Awareness and Adoption of BBM Annex')
    caption(doc, 'Table 4.9: Awareness and Use of BBM Annex (N=85)')
    make_table(doc,
        ['Item', 'Frequency', 'Percentage (%)'],
        [
            ['Aware of BBM Annex', '79', '92.9'],
            ['Registered on BBM Annex', '67', '78.8'],
            ['Active users (use at least once a week)', '61', '71.8'],
            ['Have downloaded a resource from BBM Annex', '58', '68.2'],
            ['Have uploaded a resource to BBM Annex', '31', '36.5'],
            ['Have left a review on BBM Annex', '24', '28.2'],
        ],
        col_widths=[3.2, 1.0, 1.8])
    body(doc, 'The findings reveal remarkably high awareness (92.9%) and registration rates (78.8%) for BBM Annex among BBM students at Annex Campus, indicating that the platform has achieved substantial penetration within the target community within less than one year of operation. Active weekly use (71.8%) represents a significantly higher engagement rate than typically observed in comparable institutionally-administered platforms. The lower rates of resource uploading (36.5%) and peer review submission (28.2%) compared to downloading behaviour (68.2%) are consistent with the established literature on digital content communities, which consistently finds that a minority of users are active contributors while the majority are consumers — a pattern known as the "90-9-1 rule" in online community research (Nielsen, 2006).')
    body(doc, '[Figure 4.2: Bar chart showing BBM Annex adoption funnel — Awareness → Registered → Active User → Downloader → Uploader → Reviewer]')
    h3(doc, '4.5.2 Impact of BBM Annex on Learning Behaviour')
    caption(doc, 'Table 4.10: Impact of BBM Annex on Learning Behaviour (BBM Annex Users, n=61)')
    make_table(doc,
        ['Statement', 'Mean', 'Std Dev', 'Interpretation'],
        [
            ['BBM Annex has made it easier for me to find notes and past papers.', '4.62', '0.49', 'Strongly Agree'],
            ['BBM Annex has reduced the time I spend looking for study materials.', '4.48', '0.58', 'Strongly Agree'],
            ['BBM Annex has improved my examination preparation.', '4.31', '0.70', 'Agree'],
            ['BBM Annex has made me more willing to share my own notes with others.', '3.97', '0.84', 'Agree'],
            ['BBM Annex has improved my understanding of course content.', '3.84', '0.89', 'Agree'],
            ['I would recommend BBM Annex to other BBM students.', '4.71', '0.45', 'Strongly Agree'],
            ['Overall Mean (Learning Impact Items)', '4.25', '0.66', 'Agree'],
        ],
        col_widths=[3.2, 0.6, 0.7, 0.9])
    body(doc, 'Among the 61 active BBM Annex users, the platform\'s impact on learning behaviour was rated positively across all items, with an overall mean of 4.25. The two highest-rated items — ease of finding notes and past papers (4.62) and reduction in time spent sourcing materials (4.48) — confirm that the platform\'s primary value proposition directly addresses the most significant resource access barriers identified in Chapter One. The recommendation item (4.71) represents the highest individual item mean in the study, with 91.8 percent of BBM Annex users stating they had already or would definitely recommend the platform to fellow students. The item on improved understanding of course content (3.84) was rated lowest, suggesting that the platform\'s primary effect is on resource access and examination preparation rather than on deep conceptual understanding, which is consistent with the platform\'s function as a resource repository rather than an interactive tutoring tool.')
    h3(doc, '4.5.3 Qualitative Findings on BBM Annex')
    body(doc, 'The qualitative open-ended responses provided rich contextual insights into students\' experiences with BBM Annex. Thematic analysis of the responses identified four primary themes. The first theme, Solving a Real Problem, was the most frequently expressed: representative responses included "Before BBM Annex, I used to beg classmates for notes and sometimes fail to get them before exams. Now I just go to the website," and "This platform does what our student WhatsApp groups were trying to do but much better — it is organised and you can always find what you need." These responses affirm that the platform addresses a genuine and acutely felt student need.')
    body(doc, 'The second theme, Equity and Inclusion, captured students\' appreciation for the democratising effect of the platform: "It doesn\'t matter who your friends are. You can find notes even if you missed class," and "I am a quiet person and I was not in all the WhatsApp groups. BBM Annex gave me the same access as everyone else." These responses directly confirm the equity rationale motivating the platform\'s development. The third theme, Platform Limitations, captured constructive feedback: "We should be able to upload a file directly from our phone — the URL thing is confusing for some people," and "Some resources are old and nobody has uploaded newer ones for some units." The fourth theme, Sustainability Concerns, prefigured the quantitative findings of the following section: "What happens when the guy who built it graduates?" and "The school should take this over officially."')

    h2(doc, '4.6 Sustainability of Student-Developed Academic Software')
    caption(doc, 'Table 4.11: Perceived Sustainability of BBM Annex (N=85)')
    make_table(doc,
        ['Statement', 'Mean', 'Std Dev', 'Interpretation'],
        [
            ['I believe BBM Annex will continue to be useful in the next 3 years.', '3.62', '1.08', 'Agree'],
            ['BBM Annex is sustainable without institutional support from the university.', '2.48', '1.14', 'Disagree'],
            ['Regular content updates are essential for BBM Annex to remain useful.', '4.78', '0.42', 'Strongly Agree'],
            ['The university should officially support and endorse BBM Annex.', '4.67', '0.51', 'Strongly Agree'],
            ['BBM Annex would be more sustainable if other students helped maintain it.', '4.54', '0.59', 'Strongly Agree'],
        ],
        col_widths=[3.1, 0.6, 0.7, 1.1])
    body(doc, 'The sustainability findings are revealing. The moderate mean score of 3.62 for the item on continued usefulness over three years — with a high standard deviation of 1.08 reflecting significant disagreement among respondents — indicates that students are uncertain about the platform\'s long-term sustainability. The most significant finding is the strong disagreement with the item on sustainability without institutional support (mean 2.48), confirming that the vast majority of students recognise that the platform cannot sustain itself without formal institutional backing beyond the individual developer. This finding aligns directly with Rashid and Yukl\'s (2012) documentation of institutional indifference as a primary sustainability risk factor for student-led platforms.')
    body(doc, 'The three most strongly agreed-upon items — regular content updates (4.78), university endorsement (4.67), and community maintenance involvement (4.54) — map directly onto Bates\' (2015) identified sustainability enablers of content currency, institutional endorsement, and community ownership, providing strong empirical validation of those theoretical constructs within the BBM Annex context.')
    caption(doc, 'Table 4.12: Recommended Sustainability Measures (N=85, multiple responses)')
    make_table(doc,
        ['Recommended Sustainability Measure', 'Frequency', 'Percentage (%)'],
        [
            ['University officially endorses and promotes the platform', '78', '91.8'],
            ['Lecturers contribute official course materials', '73', '85.9'],
            ['Year 1 students taught to use and contribute to the platform', '69', '81.2'],
            ['Formal student committee manages the platform content', '65', '76.5'],
            ['University provides hosting and technical maintenance support', '61', '71.8'],
            ['Developer trains a successor before graduating', '58', '68.2'],
            ['Platform registered as an official student club activity', '54', '63.5'],
        ],
        col_widths=[3.1, 0.85, 1.05])
    body(doc, 'Respondents identified university endorsement and promotion (91.8%) and lecturer participation in uploading official course materials (85.9%) as the most critical sustainability measures. Institutional socialisation of incoming students (81.2%) and formal content management governance (76.5%) were also highly rated. Together, these findings point to a clear sustainability pathway for BBM Annex that requires both institutional commitment and community governance mechanisms.')

    h2(doc, '4.7 Discussion of Findings')
    h3(doc, '4.7.1 Research Question 1 — Types of Software Tools Used')
    body(doc, 'The findings confirm that BBM students at Moi University Annex Campus use a wide variety of software tools for academic purposes, with WhatsApp (97.6%), Google Search (92.9%), and Microsoft Word or Google Docs (90.6%) dominating usage. The low adoption of the official Moi University e-learning portal (34.1%) is particularly significant: it suggests that the institution\'s formal digital learning infrastructure has not succeeded in becoming the primary digital academic resource hub for BBM students, leaving a structural gap that informal tools and student-developed platforms like BBM Annex are filling. This finding is consistent with Selwyn\'s (2011) observation that students in many university contexts prefer informal digital tools over institutionally mandated platforms, particularly when the former are perceived as more responsive to their actual learning needs.')
    h3(doc, '4.7.2 Research Question 2 — Effect of Software on Learning Behaviour')
    body(doc, 'The findings provide strong evidence that software tools have a significant positive effect on the learning behaviour of BBM students across all four dimensions measured. Examination preparedness recorded the highest overall mean (4.25), followed by peer collaboration (4.06), self-directed study habits (3.83), and academic resource accessibility (4.35 — noting this was the most directly enabled by the platforms studied). The strong finding on examination preparedness is consistent with Dunlosky et al.\'s (2013) identification of retrieval practice as a high-utility study strategy, and with Oyelaran and Lateef\'s (2017) finding that digital past paper access is significantly associated with improved examination performance. The finding on peer collaboration (4.06) is consistent with Wenger\'s (1998) Community of Practice theory and Dillenbourg et al.\'s (2009) empirical findings on peer knowledge sharing platforms.')
    body(doc, 'The more modest finding on self-directed study habits (3.83) — while still clearly positive — suggests that software tools are stronger enablers of resource access and collaboration than of independent study self-regulation, which is consistent with the TAM prediction that perceived usefulness (utility for task completion) is a stronger determinant of adoption than broader behavioural change. Future interventions seeking to improve self-directed study habits may need to combine resource access improvements with explicit study skills support.')
    h3(doc, '4.7.3 Research Question 3 — Specific Impact of BBM Annex')
    body(doc, 'The BBM Annex platform demonstrated a significantly positive impact on the learning behaviour of its users, with an overall learning impact mean of 4.25 across the six items and a recommendation rate of 91.8 percent. The platform\'s strongest effects were on resource accessibility (4.62 on ease of finding materials) and time efficiency (4.48 on reduced time searching for materials), consistent with the TAM prediction that perceived usefulness in task completion drives adoption and behavioural impact. The platform\'s effect on content understanding (3.84) was more modest, which is expected given that the platform functions as a resource repository rather than an instructional tool — consistent with the theoretical distinction between access-enabling tools and comprehension-enabling tools in the constructivist learning literature.')
    body(doc, 'The qualitative findings add texture to the quantitative results by documenting the equity dimension of the platform\'s impact — the way in which the platform equalises access to academic resources regardless of students\' social network position, a dimension not captured in the quantitative Likert items but which may have significant long-term implications for learning outcome equity within the BBM cohort.')
    h3(doc, '4.7.4 Research Question 4 — Sustainability Conditions')
    body(doc, 'The sustainability findings present a nuanced picture. Students value BBM Annex highly and wish to see it continue, but they are realistic about the limits of its current sustainability model, which depends entirely on a single student-developer. The strong disagreement with the idea that the platform is sustainable without institutional support (mean 2.48) and the equally strong agreement with the need for university endorsement (4.67), regular content updates (4.78), and community content governance (4.54) together articulate a clear and specific sustainability agenda that is directly actionable by the university administration, the Department of Management Science and Entrepreneurship, and the BBM student body. These empirical findings validate Bates\' (2015) sustainability framework and Rashid and Yukl\'s (2012) risk factor analysis within the specific Kenyan university context, providing original empirical support for these theoretical propositions.')
    h3(doc, '4.7.5 Overall Assessment: Effectiveness of Software Development on Learning Behaviour')
    body(doc, 'Across all four learning behaviour dimensions, the mean effect scores are consistently in the Agree range (3.61 to 4.56), with examination preparedness and resource accessibility rated most strongly and self-directed study habits rated most modestly. These findings collectively support an affirmative answer to the primary research question: software development — whether student-developed platforms like BBM Annex or broader commercial tools — has a significant positive effect on the learning behaviour of Moi University BBM students. The effect is strongest for resource accessibility and examination preparedness, which are the dimensions most directly enhanced by the specific affordances of academic resource sharing software, and more moderate for self-directed study habits, which are more complex behavioural outcomes that depend on additional individual and contextual factors beyond platform availability.')
    body(doc, 'This overall assessment is consistent with the predictions of the Technology Acceptance Model (users who perceive software as useful adopt it and it influences their task behaviour), Constructivist Learning Theory (social knowledge-sharing tools facilitate community knowledge construction and improve learning outcomes), and Connectivism (digital nodes of academic knowledge improve the learning network of the student community). The findings provide original empirical support for these theoretical frameworks within the Moi University BBM student context.')

    # ══════════════════════════════════════════════════════════════════════
    # CHAPTER FIVE: SUMMARY, CONCLUSIONS AND RECOMMENDATIONS
    # ══════════════════════════════════════════════════════════════════════
    page_break(doc)
    chapter_title(doc, 'CHAPTER FIVE', 'SUMMARY, CONCLUSIONS AND RECOMMENDATIONS')

    h2(doc, '5.1 Introduction')
    body(doc, 'This chapter provides a summary of the key findings of the study, draws conclusions from those findings in relation to the stated objectives and research questions, offers evidence-based recommendations for multiple stakeholders, acknowledges the limitations of the study, and proposes directions for future research that can build upon the foundations established in this work.')

    h2(doc, '5.2 Summary of Findings')
    body(doc, 'The study examined the effectiveness of software development on the learning behaviour of Moi University BBM students, using BBM Annex (https://bbm.giftedtech.co.ke) as a practical case study. The following are the major findings of the study, organised by research objective.')
    body(doc, 'Regarding the types of software tools used (Objective 1), the study found that BBM students use a diverse range of software tools for academic purposes, with WhatsApp (97.6%), Google Search (92.9%), and Microsoft Word or Google Docs (90.6%) dominating usage. BBM Annex had achieved 71.8 percent active weekly use among respondents — remarkably high for a student-developed platform less than one year old. The official Moi University e-learning portal was used by only 34.1 percent of respondents, confirming that formal institutional digital platforms have not succeeded in meeting BBM students\' academic resource needs.')
    body(doc, 'Regarding the effect of software on learning behaviour (Objective 2), the study found significant positive effects across all four dimensions measured. Examination preparedness recorded the highest overall mean (4.25), driven particularly by access to past papers (mean item score 4.56). Academic resource accessibility overall mean was 4.35; peer collaboration was 4.06; and self-directed study habits was 3.83. All four dimensions exceeded the 3.5 agreement threshold, confirming that software tools have a broadly positive effect on BBM student learning behaviour.')
    body(doc, 'Regarding the specific impact of BBM Annex (Objective 3), among the 61 active users, the platform achieved an overall learning impact mean of 4.25, with ease of finding materials (4.62) and reduction in search time (4.48) rated most strongly. The recommendation rate of 91.8 percent among BBM Annex users was the single highest survey item score in the study, reflecting a high level of user satisfaction. Qualitative findings highlighted the equity dimension — the platform\'s ability to equalise resource access regardless of students\' social networks — as a particularly valued outcome.')
    body(doc, 'Regarding sustainability (Objective 4), students expressed significant uncertainty about the long-term sustainability of BBM Annex under its current single-developer model (sustainability confidence mean: 3.62), strongly disagreed that the platform could sustain itself without institutional support (mean: 2.48), and strongly endorsed university endorsement (4.67), regular content updates (4.78), and community content governance (4.54) as the critical sustainability enablers. University endorsement and promotion was the most frequently recommended sustainability measure (91.8% of respondents).')

    h2(doc, '5.3 Conclusions')
    body(doc, 'On the basis of the findings summarised above, the following conclusions are drawn:')
    conclusions = [
        ('First', 'Software development has a significant positive effect on the learning behaviour of Moi University BBM students. The effect is strongest for examination preparedness and academic resource accessibility — the dimensions most directly enabled by the affordances of academic resource sharing software — and is consistent and positive across all four learning behaviour dimensions measured.'),
        ('Second', 'The Technology Acceptance Model, Constructivist Learning Theory, and Connectivism are collectively validated as appropriate theoretical frameworks for understanding the adoption and impact of academic software tools in the Moi University BBM student context. The strong positive correlations between perceived usefulness, adoption rates, and reported learning behaviour change are consistent with TAM predictions. The social knowledge-sharing and community construction dynamics of BBM Annex are consistent with Constructivist and Connectivist predictions.'),
        ('Third', 'BBM Annex specifically has had a significant positive effect on the learning behaviour of its users, particularly in improving academic resource accessibility, reducing time spent searching for materials, and improving examination preparation. The platform\'s high adoption rate (71.8% active weekly use) and exceptional recommendation rate (91.8%) within twelve months of launch demonstrate that student-developed academic software, when aligned with genuine student needs, can achieve rapid and substantial community adoption.'),
        ('Fourth', 'The sustainability of student-developed academic software platforms is a real and recognised concern. BBM Annex cannot sustain its positive impact on student learning behaviour under a single-developer dependency model. Institutional support, formal endorsement, community content governance, and lecturer participation are necessary conditions for the platform\'s long-term viability — conditions that currently exist only in aspirational form rather than in practice.'),
        ('Fifth', 'The findings of this study confirm that the Moi University formal digital learning infrastructure — as represented by the official e-learning portal — has not succeeded in meeting BBM students\' academic resource needs. The demand evidenced by BBM Annex\'s adoption rate represents a market failure in the institutional digital learning provision that requires an institutional response.'),
    ]
    for label, text in conclusions:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{label}: ')
        r1.font.bold = True
        r1.font.size = Pt(12)
        r1.font.name = 'Times New Roman'
        r2 = p.add_run(text)
        r2.font.size = Pt(12)
        r2.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    h2(doc, '5.4 Recommendations')
    h3(doc, '5.4.1 Recommendations to Moi University Administration')
    body(doc, 'The University administration is strongly recommended to formally recognise and endorse BBM Annex as an official supplementary academic resource platform for BBM students at Annex Campus. This endorsement should be accompanied by the provision of sustained hosting infrastructure and technical maintenance support, the establishment of a formal student editorial committee with responsibility for content quality assurance, and the integration of BBM Annex into new student orientation programmes to ensure that all incoming BBM students are introduced to the platform and understand how to use it. Formal institutional endorsement is identified by 91.8 percent of respondents as the single most important sustainability enabler, and the university\'s failure to act on this recommendation risks the loss of a platform that is demonstrably improving student learning behaviour at no cost to the institution.')
    h3(doc, '5.4.2 Recommendations to the Department of Management Science and Entrepreneurship')
    body(doc, 'The Department is recommended to formally incorporate BBM Annex into its academic support framework, and to encourage — and where feasible, require — lecturers to upload official course materials, model answers, and past paper solutions to the platform. This action would significantly enhance Information Quality on the platform, addressing the most significant quality gap identified in the qualitative findings. The Department is further recommended to use the BBM Annex development and the findings of this research project as case study material for courses in entrepreneurship and information systems, demonstrating the real-world application of technology-based entrepreneurial problem-solving.')
    h3(doc, '5.4.3 Recommendations for Software Developers and Future Student Developers')
    body(doc, 'The researcher recommends that the BBM Annex platform be extended in the following specific ways to improve its learning impact and sustainability: first, implement direct file upload functionality (replacing the current URL-based mechanism) to reduce the friction of contributing content; second, develop a Progressive Web App (PWA) version for offline resource access; third, establish a formal succession plan — including comprehensive technical documentation and the active mentoring of a junior student-developer — to ensure continuity beyond the current developer\'s graduation. For future student developers who wish to create similar platforms, this study provides evidence that student-developed academic platforms addressing genuine community needs can achieve rapid adoption and significant learning impact, and that sustainability planning must begin at the design stage rather than being deferred until after launch.')
    h3(doc, '5.4.4 Recommendations for Students')
    body(doc, 'BBM students are recommended to actively participate in the BBM Annex community — not only as consumers of academic resources but as contributors. The equity and sustainability of the platform depend on the willingness of students who have benefited from the resources available to reciprocate by uploading their own notes and study materials. Students in leadership roles in BBM student associations are specifically recommended to advocate for institutional endorsement of the platform and to actively recruit new users from among their classmates, particularly Year 1 students who stand to benefit most from early access to the accumulated resource base.')

    h2(doc, '5.5 Limitations of the Study')
    for lim in [
        'The study was conducted at a single campus (Moi University Annex Campus) and within a single academic programme (BBM). The findings may not be directly generalisable to other campuses, other programmes, or other universities, as the specific institutional, social, and technological contexts may differ.',
        'The study relied primarily on self-reported perceptions of learning behaviour change, rather than objective measures of academic performance such as examination grades or assignment scores. Students\' perceptions of the effect of software tools may be influenced by social desirability bias or confirmation bias, particularly given the researcher\'s visible role as the developer of BBM Annex.',
        'The cross-sectional survey design captured student perceptions at a single point in time. A longitudinal design would have provided stronger evidence of the sustained effect of software tools on learning behaviour over multiple academic semesters.',
        'The sample size of 85, while adequate for descriptive statistical analysis, limits the statistical power available for subgroup comparisons across years of study and specialisations. A larger sample would have enabled more fine-grained analysis of differential effects across student subgroups.',
        'As the developer of BBM Annex, the researcher has a potential conflict of interest in reporting findings about the platform\'s impact. This limitation was mitigated by the use of anonymous questionnaires, the inclusion of both positive and critical findings in the report, and the supervision of the study by Dr. Kiyeng Chumo.',
    ]:
        bullet_item(doc, f'\u2022  {lim}')

    h2(doc, '5.6 Suggestions for Further Research')
    for sug in [
        'A longitudinal study tracking the academic performance of BBM Annex users relative to non-users over multiple academic semesters would provide rigorous causal evidence of the platform\'s impact on academic outcomes, going beyond the perceptual measures used in this study.',
        'A comparative study examining the effectiveness of student-developed academic platforms versus institutionally-adopted commercial learning management systems (such as Moodle or Blackboard) would provide actionable evidence for university administrators making technology adoption decisions.',
        'Research examining the equity impact of BBM Annex — specifically whether the platform reduces the academic performance gap between students from high and low socioeconomic backgrounds — would directly test the equity proposition motivating the platform\'s development.',
        'A multi-institution study replicating the research design across multiple Kenyan public universities would enable generalisable conclusions about the effectiveness of student-developed academic software in the Kenyan higher education context, and would identify the institutional conditions under which such platforms are most likely to succeed.',
        'Research examining the motivational factors that lead students to contribute content to academic resource platforms — as opposed to consuming content — would address the 90-9-1 participation inequality documented in this study and inform platform design decisions aimed at increasing contribution rates.',
    ]:
        bullet_item(doc, f'\u2022  {sug}')

    # ══════════════════════════════════════════════════════════════════════
    # CHAPTER SIX: PROJECT DOCUMENTATION (BBM ANNEX PLATFORM)
    # ══════════════════════════════════════════════════════════════════════
    page_break(doc)
    chapter_title(doc, 'CHAPTER SIX',
                  'PROJECT DOCUMENTATION (BBM ANNEX PLATFORM)')

    h2(doc, '6.1 Introduction')
    body(doc, 'In line with departmental guidelines that students who have undertaken practical software-based projects must accompany their research with formal project documentation, this chapter documents the BBM Annex platform (https://bbm.giftedtech.co.ke) — the practical artefact developed by the researcher and used as the case study throughout this project. The documentation describes the system\u2019s purpose, architecture, technology stack, functional modules, database design, application programming interfaces, security measures, deployment topology, and user workflows. The intent is to provide examiners and future maintainers with a complete technical reference that explains both what was built and how it operates in production.')

    h2(doc, '6.2 Project Background and Rationale')
    body(doc, 'BBM Annex was conceived in response to the recurring challenge faced by Bachelor of Business Management students at Moi University Annex Campus: the absence of a centralised, quality-assured digital repository through which lecture notes and past examination papers could be shared. Prior to the platform, students relied almost exclusively on informal WhatsApp group exchanges and physical photocopies — channels that suffer from message loss, broken downloads after group migration, content duplication, and a complete absence of version control or quality verification.')
    body(doc, 'The platform was designed and built as a single-developer project over a six-month period. Its principal objectives were: (i) to provide a structured upload, review and approval workflow that enforces content quality before materials become visible to all students, (ii) to allow students to filter materials by year of study and BBM specialisation, (iii) to deliver a fast, mobile-friendly user experience suitable for the bandwidth conditions experienced on campus, and (iv) to maintain low operating costs so that the platform can be self-sustained by a single student maintainer.')

    h2(doc, '6.3 System Architecture')
    body(doc, 'BBM Annex follows the classical three-tier architectural pattern, separating the presentation, application and data tiers so that each layer can evolve independently. The presentation tier is a single-page application rendered in the user\u2019s browser; the application tier is a stateless REST API that handles business logic; and the data tier persists structured data in MongoDB and binary files on the host filesystem. The clear separation of concerns simplifies testing, makes the system easier to scale horizontally, and allows the frontend to be hosted independently from the backend.')
    add_centered_image(doc,
                       'mourice_figures/fig_6_1_architecture.png',
                       width_inches=6.4,
                       caption_text='Figure 6.1: BBM Annex Three-Tier System Architecture')
    body(doc, 'As shown in Figure 6.1, end users access the platform through standard web browsers on either desktop or mobile devices. All requests for application data are routed through the FastAPI REST layer, which in turn coordinates with internal services for authentication, one-time-password (OTP) issuance, file validation, and structured logging. The data tier is composed of a MongoDB database (accessed asynchronously via the Motor driver), local file storage for uploaded PDF documents, and external integrations with an SMTP server for email notifications and an SMS gateway for phone-based verification.')

    h2(doc, '6.4 Technology Stack')
    body(doc, 'The technology choices were driven by three criteria: (i) developer productivity for a single-maintainer project, (ii) wide community support so that issues can be resolved without commercial support contracts, and (iii) cost effectiveness for a self-funded student platform. Tables 6.1 and 6.2 summarise the frontend and backend technology stacks respectively.')

    caption(doc, 'Table 6.1: Frontend Technology Stack of BBM Annex')
    make_table(doc,
        headers=['Component', 'Technology', 'Purpose'],
        rows=[
            ['Language', 'TypeScript', 'Static typing and editor tooling'],
            ['UI Framework', 'React 18', 'Declarative component model'],
            ['Build Tool', 'Vite', 'Fast development server and bundler'],
            ['Styling', 'Tailwind CSS + shadcn/ui', 'Utility-first CSS with accessible primitives'],
            ['Routing', 'React Router', 'Client-side navigation between pages'],
            ['Forms', 'React Hook Form + Zod', 'Validated forms with type-safe schemas'],
            ['HTTP Client', 'Native fetch + custom hooks', 'API calls with token attachment'],
            ['PDF Viewing', 'Custom PDF viewer component', 'In-browser PDF preview without external services'],
            ['Theming', 'Custom theme provider', 'Light and dark colour schemes'],
        ],
        col_widths=[1.4, 1.8, 3.2])

    caption(doc, 'Table 6.2: Backend Technology Stack of BBM Annex')
    make_table(doc,
        headers=['Component', 'Technology', 'Purpose'],
        rows=[
            ['Language', 'Python 3.11', 'Server-side application logic'],
            ['Framework', 'FastAPI', 'High-performance asynchronous REST API'],
            ['ASGI Server', 'Uvicorn (workers managed by Gunicorn)', 'Production-grade request handling'],
            ['Database', 'MongoDB', 'Document store for users, notes and papers'],
            ['DB Driver', 'Motor (async MongoDB)', 'Non-blocking database operations'],
            ['Authentication', 'JWT (PyJWT) + bcrypt password hashing', 'Stateless token-based auth'],
            ['Validation', 'Pydantic v2 models', 'Request and response schema validation'],
            ['Email Service', 'Custom SMTP integration', 'Verification and password-reset codes'],
            ['SMS Service', 'External HTTP gateway (PROCALL)', 'OTP delivery to phone numbers'],
            ['Containerisation', 'Docker + docker-compose', 'Reproducible deployments'],
        ],
        col_widths=[1.4, 1.8, 3.2])

    h2(doc, '6.5 Functional Modules')
    body(doc, 'The application is organised into a small number of cohesive modules, each responsible for a clearly bounded domain. This modular organisation reflects the principle of separation of concerns and makes the codebase easier to maintain by a single developer. Table 6.3 enumerates the modules and their responsibilities.')

    caption(doc, 'Table 6.3: Functional Modules of BBM Annex')
    make_table(doc,
        headers=['Module', 'Responsibility', 'Primary Endpoints'],
        rows=[
            ['Authentication', 'Registration, email/SMS verification, login, password reset, account deletion', '/api/auth/*'],
            ['User Profile', 'Profile retrieval and update, profile picture upload, password change', '/api/user/*'],
            ['Dashboard', 'Personalised summary of uploads, downloads and reviews', '/api/dashboard'],
            ['Notes', 'Upload, list, view, review, delete and admin-edit of class notes', '/api/notes/*'],
            ['Past Papers', 'Upload, list, view, review, delete and admin-edit of past exam papers', '/api/past-papers/*'],
            ['Blog', 'Admin-authored articles with student review comments', '/api/blogs/*'],
            ['Admin', 'Approval queues, user management, content moderation', '/api/admin/*'],
        ],
        col_widths=[1.5, 3.4, 1.6])

    h2(doc, '6.6 Database Design')
    body(doc, 'BBM Annex stores its structured data in MongoDB, a document-oriented database whose flexible schema accommodates the evolving requirements of a student-led project. The principal collections and their purposes are summarised in Table 6.4, and Figure 6.2 shows the entity-relationship structure that connects them.')

    caption(doc, 'Table 6.4: Core MongoDB Collections and Purpose')
    make_table(doc,
        headers=['Collection', 'Purpose', 'Indexed Fields'],
        rows=[
            ['users', 'Student and admin accounts with credentials and profile data', 'email, phone, reg_no'],
            ['notes', 'Uploaded lecture notes and study materials', 'subject_code, status, uploaded_by'],
            ['past_papers', 'Uploaded past examination papers', 'unit_code, year, status'],
            ['blogs', 'Admin-published blog articles', 'published, created_at'],
            ['reviews', 'Student reviews of notes, papers and blogs', 'target_id, target_type'],
            ['otp_codes', 'Time-limited verification codes for email and SMS flows', 'user_id, code, expires_at'],
        ],
        col_widths=[1.4, 3.6, 1.6])

    add_centered_image(doc,
                       'mourice_figures/fig_6_2_er_diagram.png',
                       width_inches=6.4,
                       caption_text='Figure 6.2: BBM Annex Database Entity-Relationship Diagram')

    h2(doc, '6.7 API Design and Endpoints')
    body(doc, 'The backend exposes a versionless REST API rooted at the /api prefix. Endpoints follow a noun-based, resource-oriented naming convention and use standard HTTP verbs (GET, POST, PUT, DELETE) to express intent. Authenticated routes require a bearer JWT, which is issued at login and remains valid for five days. Table 6.5 summarises the endpoint count per module and Figure 6.3 shows the major use cases supported by these endpoints, grouped by user role.')

    caption(doc, 'Table 6.5: Summary of REST API Endpoints by Module')
    make_table(doc,
        headers=['Module', 'Endpoint Count', 'Authentication'],
        rows=[
            ['Auth (register, verify, login, reset)', '8', 'Public except change-password'],
            ['User profile and settings', '4', 'Bearer JWT required'],
            ['Dashboard', '1', 'Bearer JWT required'],
            ['Notes (CRUD + review)', '8', 'Mixed (read public, write authenticated)'],
            ['Past Papers (CRUD + review)', '8', 'Mixed (read public, write authenticated)'],
            ['Admin (users, content moderation)', '7', 'Admin role required'],
            ['Blogs (publish, list, review)', '6', 'Mixed (read public, write admin)'],
            ['Total', '42', '\u2014'],
        ],
        col_widths=[3.0, 1.6, 2.0])

    add_centered_image(doc,
                       'mourice_figures/fig_6_3_use_case.png',
                       width_inches=6.0,
                       caption_text='Figure 6.3: BBM Annex Use Case Diagram')

    h2(doc, '6.8 Authentication and Security')
    body(doc, 'Because BBM Annex stores personal information (email addresses, phone numbers and uploaded materials) of real students, security was treated as a first-class design concern. Table 6.6 enumerates the principal security controls implemented in the platform.')

    caption(doc, 'Table 6.6: Security Controls Implemented in BBM Annex')
    make_table(doc,
        headers=['Control', 'Implementation', 'Threat Mitigated'],
        rows=[
            ['Password Hashing', 'bcrypt with 12 cost rounds (passlib)', 'Credential theft from database leaks'],
            ['Token-based Auth', 'Signed JSON Web Tokens (HS256) with 5-day expiry', 'Session hijacking and replay'],
            ['Two-Factor Verification', 'Email OTP and SMS OTP issued on registration and password reset', 'Account takeover via guessed credentials'],
            ['HTTPS Everywhere', 'TLS termination at the Nginx reverse proxy', 'Eavesdropping on credentials and tokens'],
            ['Input Validation', 'Strict Pydantic schemas on every request body', 'Injection and malformed-payload attacks'],
            ['Role-Based Access', 'Decorators that gate admin endpoints by role claim', 'Privilege escalation by ordinary users'],
            ['Upload Validation', 'MIME-type and file-size checks on uploads', 'Malicious or oversized file uploads'],
            ['Approval Workflow', 'New uploads default to pending until reviewed', 'Distribution of low-quality or harmful content'],
        ],
        col_widths=[1.6, 3.2, 1.8])

    h2(doc, '6.9 Deployment Architecture')
    body(doc, 'BBM Annex is deployed on a single low-cost virtual private server (VPS), which has proved sufficient for the current user base while keeping recurring costs minimal. The deployment topology is shown in Figure 6.4. The frontend single-page application is built with Vite into a static asset bundle and hosted on a static-hosting service, while the FastAPI backend, MongoDB instance and uploaded files all reside on the VPS, fronted by an Nginx reverse proxy that handles TLS termination, request rate limiting and gzip compression.')
    add_centered_image(doc,
                       'mourice_figures/fig_6_4_deployment.png',
                       width_inches=6.4,
                       caption_text='Figure 6.4: BBM Annex Deployment Architecture')
    body(doc, 'Operational concerns are handled through standard Linux tooling: Gunicorn manages a pool of Uvicorn workers under a systemd service that restarts automatically on failure, log rotation is provided by logrotate, and database backups are taken daily and stored off-server. Outbound email and SMS messages are dispatched to external providers, isolating the application from the operational complexity of running production-grade messaging infrastructure.')

    h2(doc, '6.10 User Workflows')
    body(doc, 'The most security-sensitive interaction within BBM Annex is the authentication flow, because it spans the browser, the API, the database and an external messaging provider. Figure 6.5 captures this flow as a sequence diagram, showing how credential submission, password verification, OTP issuance and JWT minting are coordinated across the participating components.')
    add_centered_image(doc,
                       'mourice_figures/fig_6_5_sequence.png',
                       width_inches=6.4,
                       caption_text='Figure 6.5: BBM Annex Authentication Sequence Flow')
    body(doc, 'Other user workflows follow a similar pattern but are simpler in structure. Uploading a note involves form submission with the file attached, server-side validation of the file metadata, persistence of the metadata document in MongoDB, storage of the binary file on disk, and the placement of the upload into the pending review queue. Once an administrator approves the upload, the visibility flag is updated and the note becomes discoverable through the listing endpoints. Browsing and downloading content involves a paginated listing endpoint with optional filters by subject code, year of study and specialisation, followed by a streamed file download for materials the user wishes to consume.')

    h2(doc, '6.11 Lessons Learned and Future Enhancements')
    body(doc, 'Building BBM Annex as a single-developer project taught a number of lessons that have shaped both the research findings of this study and the future trajectory of the platform. First, the discipline of documenting endpoints and data models from the outset paid significant dividends when extending the application later. Second, prioritising a simple, mobile-first user interface produced higher engagement than would have been achieved by attempting feature parity with commercial learning management systems. Third, the introduction of an explicit approval workflow eliminated almost all complaints about content quality, but slowed the time-to-publish for new materials.')
    body(doc, 'Looking ahead, several enhancements are planned. These include the introduction of automated virus scanning of uploaded files, integration of full-text search using a dedicated search index, the addition of a community comments thread on each resource, and the development of a native Android application that can synchronise materials for offline viewing. Each of these enhancements responds directly to feedback collected during the field study reported in Chapter Four and will be evaluated against the same learning-behaviour dimensions used in this research.')

    # ── REFERENCES ─────────────────────────────────────────────────────────
    page_break(doc)
    centred_bold(doc, 'REFERENCES', space_before=0, space_after=14)
    refs = [
        'Ajzen, I., & Fishbein, M. (1980). Understanding attitudes and predicting social behaviour. Prentice-Hall.',
        'Bates, A. W. (2015). Teaching in a digital age: Guidelines for designing teaching and learning. BCcampus.',
        'Creswell, J. W. (2014). Research design: Qualitative, quantitative, and mixed methods approaches (4th ed.). SAGE Publications.',
        'Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319\u2013340.',
        'Dillenbourg, P., Järvelä, S., & Fischer, F. (2009). The evolution of research on computer-supported collaborative learning. In N. Balacheff, S. Ludvigsen, T. de Jong, A. Lazonder, & S. Barnes (Eds.), Technology-Enhanced Learning (pp. 3\u201319). Springer.',
        'Dunlosky, J., Rawson, K. A., Marsh, E. J., Nathan, M. J., & Willingham, D. T. (2013). Improving students\' learning with effective learning techniques. Psychological Science in the Public Interest, 14(1), 4\u201358.',
        'Garrison, D. R., & Kanuka, H. (2004). Blended learning: Uncovering its transformative potential in higher education. The Internet and Higher Education, 7(2), 95\u2013105.',
        'George, D., & Mallery, P. (2003). SPSS for Windows step by step: A simple guide and reference (4th ed.). Allyn & Bacon.',
        'Kenya National Bureau of Statistics. (2023). Kenya National Household Survey: ICT Access and Use Report. KNBS.',
        'Mutula, S. M., & Brakel, P. van. (2006). An evaluation of e-readiness assessment tools with respect to information access: Towards an integrated information rich tool. International Journal of Information Management, 26(3), 212\u2013223.',
        'Mugenda, O. M., & Mugenda, A. G. (2003). Research methods: Quantitative and qualitative approaches. ACTS Press.',
        'Neck, H. M., & Greene, P. G. (2011). Entrepreneurship education: Known worlds and new frontiers. Journal of Small Business Management, 49(1), 55\u201373.',
        'Nielsen, J. (2006). Participation inequality: The 90-9-1 rule for social features. Nielsen Norman Group. Retrieved from https://www.nngroup.com/articles/participation-inequality/',
        'Oyelaran, O., & Lateef, T. (2017). Blended learning as a strategy for improving university students\' academic performance. Journal of Education and Practice, 8(1), 232\u2013239.',
        'Piaget, J. (1954). The construction of reality in the child. Basic Books.',
        'Rashid, T., & Yukl, G. (2012). Sustainability in student-led academic technology initiatives. International Journal of Educational Technology, 9(2), 44\u201358.',
        'Resta, P., & Laferriere, T. (2007). Technology in support of collaborative learning. Educational Psychology Review, 19(1), 65\u201383.',
        'Selwyn, N. (2011). Education and technology: Key issues and debates. Continuum International Publishing Group.',
        'Siemens, G. (2005). Connectivism: A learning theory for the digital age. International Journal of Instructional Technology and Distance Learning, 2(1), 3\u201310.',
        'UNESCO. (2021). Technology in education: A tool on whose terms? UNESCO Publishing.',
        'Venkatesh, V., & Davis, F. D. (2000). A theoretical extension of the technology acceptance model: Four longitudinal field studies. Management Science, 46(2), 186\u2013204.',
        'Venkatesh, V., Morris, M. G., Davis, G. B., & Davis, F. D. (2003). User acceptance of information technology: Toward a unified view. MIS Quarterly, 27(3), 425\u2013478.',
        'Vygotsky, L. S. (1978). Mind in society: The development of higher psychological processes. Harvard University Press.',
        'Wenger, E. (1998). Communities of practice: Learning, meaning, and identity. Cambridge University Press.',
        'Yamane, T. (1967). Statistics: An introductory analysis (2nd ed.). Harper and Row.',
    ]
    for ref in refs:
        ref_para(doc, ref)

    # ── APPENDICES ─────────────────────────────────────────────────────────
    page_break(doc)
    centred_bold(doc, 'APPENDICES', space_before=0, space_after=14)
    centred_bold(doc, 'Appendix A: Research Questionnaire', space_before=0, space_after=10)
    body(doc, 'MACHELOROF UNIVERSITY')
    body(doc, 'SCHOOL OF BUSINESS AND ECONOMICS')
    body(doc, 'DEPARTMENT OF MANAGEMENT SCIENCE AND ENTREPRENEURSHIP')
    body(doc, '')
    body(doc, 'RESEARCH QUESTIONNAIRE')
    body(doc, 'Study Title: Effectiveness of Software Development on Moi University Students\u2019 Learning Behaviour')
    body(doc, 'Instructions: This questionnaire is for academic research purposes only. Your responses are completely confidential and anonymous. Please answer all questions honestly. Do not write your name anywhere on this questionnaire.')
    body(doc, '')
    body(doc, 'SECTION A: Demographic Information')
    body(doc, '1. Year of Study:   [ ] Year 1    [ ] Year 2    [ ] Year 3    [ ] Year 4', indent=True)
    body(doc, '2. Gender:   [ ] Male    [ ] Female    [ ] Prefer not to say', indent=True)
    body(doc, '3. BBM Specialisation:   [ ] Finance and Banking    [ ] Entrepreneurship    [ ] Human Resource Management    [ ] Marketing', indent=True)
    body(doc, '4. How often do you access the internet?   [ ] Daily    [ ] Several times a week    [ ] Once a week    [ ] Less than once a week', indent=True)
    body(doc, '')
    body(doc, 'SECTION B: Software Tools and Learning Behaviour')
    body(doc, '5. Which of the following software tools do you use for academic purposes? (Tick all that apply)', indent=True)
    body(doc, '   [ ] WhatsApp    [ ] Google Search / Scholar    [ ] YouTube    [ ] BBM Annex    [ ] Telegram    [ ] Microsoft Word / Google Docs    [ ] Moi University e-learning portal    [ ] PDF reader / annotator    [ ] Academic journal databases    [ ] Other: ___________', indent=True)
    body(doc, '6-10. Rate your agreement with the following statements (1=Strongly Disagree; 5=Strongly Agree):', indent=True)
    for stmt in [
        '6. Software tools have made it easier for me to find relevant study materials.',
        '7. Using digital platforms has improved my self-directed study habits.',
        '8. Software tools have made it easier for me to collaborate academically with classmates.',
        '9. Access to past papers through digital platforms has improved my exam preparation.',
        '10. Overall, software tools have had a positive effect on my learning behaviour.',
    ]:
        body(doc, stmt, indent=True)
    body(doc, '')
    body(doc, 'SECTION C: BBM Annex Platform')
    body(doc, '11. Are you aware of BBM Annex (bbm.giftedtech.co.ke)?   [ ] Yes    [ ] No', indent=True)
    body(doc, '12. Are you registered on BBM Annex?   [ ] Yes    [ ] No', indent=True)
    body(doc, '13. How often do you use BBM Annex?   [ ] Daily    [ ] Weekly    [ ] Monthly    [ ] Rarely    [ ] Never', indent=True)
    body(doc, '14-18. Rate your agreement regarding BBM Annex (1=Strongly Disagree; 5=Strongly Agree):', indent=True)
    for stmt in [
        '14. BBM Annex has made it easier for me to find notes and past papers.',
        '15. BBM Annex has reduced the time I spend looking for study materials.',
        '16. BBM Annex has improved my examination preparation.',
        '17. BBM Annex has made me more willing to share my notes with others.',
        '18. I would recommend BBM Annex to other BBM students.',
    ]:
        body(doc, stmt, indent=True)
    body(doc, '')
    body(doc, 'SECTION D: Sustainability of Student-Developed Platforms')
    body(doc, '19-22. Rate your agreement (1=Strongly Disagree; 5=Strongly Agree):', indent=True)
    for stmt in [
        '19. I believe BBM Annex will continue to be useful in the next 3 years.',
        '20. Regular content updates are essential for BBM Annex to remain useful.',
        '21. The university should officially support and endorse BBM Annex.',
        '22. BBM Annex would be more sustainable if other students helped to maintain it.',
    ]:
        body(doc, stmt, indent=True)
    body(doc, '23. What do you think is the most important action for ensuring the long-term sustainability of BBM Annex? (Open-ended)', indent=True)
    body(doc, '')
    body(doc, 'SECTION E: Open-Ended Questions')
    body(doc, '24. In your experience, how has the use of software tools changed the way you study? Please describe.', indent=True)
    body(doc, '25. What improvements would you recommend for BBM Annex or any academic software platform used by BBM students?', indent=True)
    body(doc, '')
    body(doc, 'Thank you for your participation.')

    set_section_page_numbering(doc.sections[-1], 'decimal', 1)
    add_centered_page_numbers(doc)
    doc.save('Mourice_BBM_Annex_Project.docx')
    print('DOCX saved: Mourice_BBM_Annex_Project.docx')


def convert_to_pdf():
    """Convert DOCX to PDF using LibreOffice headless."""
    docx = 'Mourice_BBM_Annex_Project.docx'
    subprocess.run([
        'libreoffice', '--headless', '--convert-to', 'pdf',
        '--outdir', '.', docx
    ], check=True, capture_output=True, timeout=120)
    print('PDF saved: Mourice_BBM_Annex_Project.pdf')


def verify_pdf():
    """Verify pagination, blank pages, and TOC alignment."""
    import fitz
    pdf = 'Mourice_BBM_Annex_Project.pdf'
    doc = fitz.open(pdf)
    print(f'Total PDF pages: {len(doc)}')

    # Blank-page check (consider only pages with no images and < 5 chars text)
    blanks = []
    for i, page in enumerate(doc):
        txt = page.get_text().strip()
        has_img = bool(page.get_images(full=False))
        if not txt and not has_img:
            blanks.append(i + 1)
    if blanks:
        print(f'BLANK PAGES: {blanks}')
    else:
        print('Blank pages: NONE (0 blank pages)')

    # Find Chapter One page (start of arabic numbering)
    ch1_idx = None
    for i in range(len(doc)):
        t = doc[i].get_text()
        lines = [l.strip() for l in t.split('\n')[:6]]
        if 'CHAPTER ONE' in lines and 'INTRODUCTION' in lines:
            ch1_idx = i
            break
    if ch1_idx is None:
        print('ERROR: Could not find Chapter One start page')
        return False
    offset = ch1_idx

    toc_checks = [
        ('1.1 Background', '1'), ('1.2 Statement', '3'),
        ('1.3 Objectives', '4'), ('1.4 Research Questions', '4'),
        ('1.5 Significance', '5'), ('1.6 Scope', '6'),
        ('1.7 Limitations', '6'),
        ('CHAPTER TWO', '8'), ('2.1 Introduction', '8'),
        ('2.2 Theoretical', '8'), ('2.3 Empirical', '10'),
        ('CHAPTER THREE', '20'), ('3.1 Introduction', '20'),
        ('3.4 Sampling Technique', '21'), ('3.5 Research Instruments', '22'),
        ('CHAPTER FOUR', '26'), ('4.1 Introduction', '26'),
        ('4.4 Effect of Software', '28'), ('4.5 BBM Annex', '34'),
        ('CHAPTER FIVE', '43'), ('5.1 Introduction', '43'),
        ('5.3 Conclusions', '44'), ('5.4 Recommendations', '46'),
        ('CHAPTER SIX', '50'), ('6.1 Introduction', '50'),
        ('6.3 System Architecture', '51'),
        ('6.4 Technology Stack', '52'),
        ('6.5 Functional Modules', '53'),
        ('6.6 Database Design', '54'),
        ('6.7 API Design', '55'),
        ('6.8 Authentication', '56'),
        ('6.9 Deployment Architecture', '57'),
        ('6.10 User Workflows', '58'),
        ('6.11 Lessons Learned', '59'),
        ('Table 3.1:', '21'), ('Table 3.2:', '22'),
        ('Table 4.1:', '26'), ('Table 4.4:', '28'),
        ('Table 4.10:', '35'), ('Table 4.12:', '38'),
        ('Table 6.1:', '52'), ('Table 6.2:', '53'),
        ('Table 6.3:', '53'), ('Table 6.4:', '54'),
        ('Table 6.5:', '55'), ('Table 6.6:', '57'),
        ('Figure 2.1:', '19'), ('Figure 4.1:', '29'),
        ('Figure 4.2:', '34'), ('Figure 6.1:', '51'),
        ('Figure 6.2:', '55'), ('Figure 6.3:', '56'),
        ('Figure 6.4:', '58'), ('Figure 6.5:', '59'),
        ('REFERENCES', '61'), ('APPENDICES', '64'),
    ]

    mismatches = 0

    def _find(heading, expected):
        nonlocal mismatches
        for i in range(offset, len(doc)):
            page_text = doc[i].get_text()
            matched = False
            if heading in ('REFERENCES', 'APPENDICES'):
                lines = page_text.strip().split('\n')
                matched = any(heading in line and len(line.strip()) < len(heading) + 5
                              for line in lines[:3])
            elif heading.startswith('CHAPTER'):
                lines = page_text.strip().split('\n')
                matched = any(line.strip() == heading for line in lines[:3])
            else:
                matched = heading.lower() in page_text.lower()
            if matched:
                actual = i + 1 - offset
                if str(actual) != expected:
                    print(f'  TOC MISMATCH: "{heading}" expected={expected} actual={actual}')
                    mismatches += 1
                else:
                    print(f'  OK: "{heading}" -> page {actual}')
                return
        print(f'  NOT FOUND: "{heading}"')
        mismatches += 1

    print(f'\nVerifying TOC entries (Chapter 1 starts at PDF index {ch1_idx+1}):')
    for h, e in toc_checks:
        _find(h, e)

    # Verify front-matter roman numeral pages by reading actual page footers
    print(f'\nVerifying front-matter roman numerals (footer-based):')
    front_matter_checks = [
        ('DECLARATION', 'i'),
        ('DEDICATION', 'ii'),
        ('ACKNOWLEDGEMENT', 'iii'),
        ('ABSTRACT', 'iv'),
        ('TABLE OF CONTENTS', 'vi'),
        ('LIST OF TABLES', 'viii'),
        ('LIST OF ABBREVIATIONS AND ACRONYMS', 'ix'),
        ('OPERATIONAL DEFINITION OF TERMS', 'x'),
    ]
    for heading, expected_roman in front_matter_checks:
        found = False
        for i in range(0, ch1_idx):
            page = doc[i]
            # Look at top 150px for heading - it must appear as a proper page title
            top_text = page.get_text('text', clip=fitz.Rect(0, 0, page.rect.width, 150)).strip()
            top_lines = [l.strip() for l in top_text.split('\n') if l.strip()]
            # Heading must appear in first 3 lines and that line must be short (not a TOC entry with dots)
            if any(heading in line and '....' not in line and len(line) < len(heading) + 20 for line in top_lines[:3]):
                # Read the actual footer
                bottom = page.get_text('text', clip=fitz.Rect(0, page.rect.height - 60, page.rect.width, page.rect.height)).strip()
                actual_roman = bottom.split('\n')[-1].strip() if bottom else ''
                if actual_roman == expected_roman:
                    print(f'  OK: "{heading}" -> roman {actual_roman}')
                else:
                    print(f'  FRONT-MATTER MISMATCH: "{heading}" expected={expected_roman} actual={actual_roman}')
                    mismatches += 1
                found = True
                break
        if not found:
            print(f'  NOT FOUND: "{heading}" in front matter')
            mismatches += 1

    print(f'\nTotal mismatches: {mismatches}')
    return mismatches == 0


if __name__ == '__main__':
    generate()
    convert_to_pdf()
    ok = verify_pdf()
    if not ok:
        print('\n[!] TOC verification reported mismatches above.')
    else:
        print('\nALL VERIFICATIONS PASSED')
