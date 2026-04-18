#!/usr/bin/env python3
"""
Academic Research Project Generator — DOCX
Title: Effect of Audit Committee Characteristics on Firm Performance:
       A Case of Listed Banks in Kenya
Student: Faith Awuor Okumu | BBM/1491/23
Supervisor: Dr. Neddy Soi
Department: School of Business and Economics
Degree: Bachelor of Business Management (Accounting Option)
April 2026
"""

import os
import subprocess
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

os.makedirs('files', exist_ok=True)

OUTPUT_DOCX = 'files/Faith_Okumu_Research_Project.docx'
OUTPUT_PDF = 'files/Faith_Okumu_Research_Project.pdf'


# ─── helpers ─────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def page_break(doc):
    if doc.paragraphs:
        doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
    else:
        doc.add_page_break()


def body(doc, text, indent=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent:
        pf.first_line_indent = Inches(0.4)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.italic = True
    r.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def chapter_title(doc, ch_line, title_line):
    for txt, after in [(ch_line, 2), (title_line, 18)]:
        p = doc.add_paragraph()
        r = p.add_run(txt)
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def centred(doc, text, size=12, bold=False, space_before=0, space_after=4):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def centred_bold(doc, text, size=12, space_before=0, space_after=8):
    return centred(doc, text, size=size, bold=True,
                   space_before=space_before, space_after=space_after)


def caption(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.italic = True
    r.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def bullet_item(doc, text, level=0):
    p = doc.add_paragraph()
    r = p.add_run(f'\u2022  {text}')
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Inches(0.4 + level * 0.3)
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def toc_row(doc, title, page, bold=False, indent=0):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.font.size = Pt(11)
    r.font.bold = bold
    r.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(indent * 0.3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), '9000')
    tabs.append(tab)
    pPr.append(tabs)
    r2 = p.add_run(f'\t{page}')
    r2.font.size = Pt(11)
    r2.font.bold = bold
    r2.font.name = 'Times New Roman'


def make_table(doc, headers, rows, col_widths=None, first_col_left=True):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, 'D3D3D3')
        pf = cell.paragraphs[0].paragraph_format
        pf.space_before = Pt(3)
        pf.space_after = Pt(3)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for r, row in enumerate(rows):
        dr = tbl.rows[r + 1]
        for c, val in enumerate(row):
            cell = dr.cells[c]
            cell.text = str(val) if val is not None else ''
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run('')
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            cell.paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.LEFT
                if first_col_left and c == 0 else WD_ALIGN_PARAGRAPH.CENTER)
            pf = cell.paragraphs[0].paragraph_format
            pf.space_before = Pt(3)
            pf.space_after = Pt(3)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if col_widths:
        for row in tbl.rows:
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Inches(col_widths[j])
    return tbl


def insert_section_break(doc, fmt='lowerRoman', start=1, title_page=False):
    p_obj = doc.add_paragraph()
    p = p_obj._p
    pPr = OxmlElement('w:pPr')
    sectPr = OxmlElement('w:sectPr')
    type_el = OxmlElement('w:type')
    type_el.set(qn('w:val'), 'nextPage')
    sectPr.append(type_el)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)
    if title_page:
        titlePg = OxmlElement('w:titlePg')
        sectPr.append(titlePg)
    pPr.append(sectPr)
    p.insert(0, pPr)


def _add_page_field_to_footer(section, instr):
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.clear()
    run = fp.add_run()
    run.font.name = 'Times New Roman'; run.font.size = Pt(11)
    a = OxmlElement('w:fldChar'); a.set(qn('w:fldCharType'), 'begin')
    b = OxmlElement('w:instrText'); b.text = instr
    c = OxmlElement('w:fldChar'); c.set(qn('w:fldCharType'), 'end')
    run._r.append(a); run._r.append(b); run._r.append(c)


def _set_section_pg_num(section, fmt, start):
    sectPr = section._sectPr
    for old in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(old)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)


def add_centered_page_numbers(doc):
    sections = list(doc.sections)
    # section[0] = cover (titlePg masks number)
    # section[1] = front matter (lowerRoman starting ii)
    # section[-1] = body (decimal starting 1)
    if sections:
        sections[0].different_first_page_header_footer = True
    if len(sections) > 1:
        _add_page_field_to_footer(sections[1], 'PAGE')
    # ensure final body section is decimal starting at 1
    _set_section_pg_num(sections[-1], 'decimal', 1)
    for s in sections[2:]:
        _add_page_field_to_footer(s, 'PAGE')


def draw_conceptual_framework(doc):
    """Insert a simple text-based conceptual framework diagram as a table."""
    centred_bold(doc, 'Independent Variables                          Dependent Variable',
                 size=11, space_before=8, space_after=6)
    tbl = doc.add_table(rows=4, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    iv = ['Audit Committee Size', 'Audit Committee Independence',
          'Audit Committee Financial Expertise', 'Audit Committee Meeting Frequency']
    for i, label in enumerate(iv):
        c = tbl.rows[i].cells[0]
        c.text = label
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = c.paragraphs[0].runs[0]
        run.font.bold = True; run.font.size = Pt(11); run.font.name = 'Times New Roman'
        set_cell_bg(c, 'EAF3FF')
        arrow = tbl.rows[i].cells[1]
        arrow.text = '\u2192'
        arrow.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        arun = arrow.paragraphs[0].runs[0]
        arun.font.size = Pt(14); arun.font.bold = True; arun.font.name = 'Times New Roman'
    # merge dependent variable column vertically
    dep = tbl.rows[0].cells[2]
    for i in range(1, 4):
        dep.merge(tbl.rows[i].cells[2])
    dep.text = 'Firm Performance\n(ROA, ROE)'
    for p in dep.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(12); r.font.name = 'Times New Roman'
    set_cell_bg(dep, 'FFF4E0')
    # Borders via Table Grid style
    tbl.style = 'Table Grid'
    centred(doc, 'Control Variables: Bank Size, Leverage', size=11, bold=True,
            space_before=8, space_after=4)
    caption(doc, 'Figure 1.1: Conceptual Framework')


# ─── document body ───────────────────────────────────────────────────────
def generate():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = Inches(1.0); sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.25); sec.right_margin = Inches(1.0)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'; style.font.size = Pt(12)

    # ═══ COVER PAGE ═══
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_after = Pt(6)
    try:
        p_logo.add_run().add_picture('attached_assets/moi_logo_1773763714167.png', width=Inches(1.4))
    except Exception:
        pass
    centred_bold(doc, 'MOI UNIVERSITY', 14, space_after=2)
    centred(doc, 'SCHOOL OF BUSINESS AND ECONOMICS', 12, bold=True, space_after=2)
    centred(doc, 'DEPARTMENT OF ACCOUNTING AND FINANCE', 12, bold=True, space_after=22)

    tp = doc.add_paragraph()
    tr = tp.add_run('EFFECT OF AUDIT COMMITTEE CHARACTERISTICS ON FIRM PERFORMANCE: '
                    'A CASE OF LISTED BANKS IN KENYA')
    tr.font.size = Pt(13); tr.font.bold = True; tr.font.name = 'Times New Roman'
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(20)
    tp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    centred(doc, 'BY', 12, space_after=2)
    centred_bold(doc, 'FAITH AWUOR OKUMU', 12, space_after=2)
    centred(doc, 'REG. NO: BBM/1491/23', 12, space_after=20)
    centred(doc, 'SUPERVISOR', 12, space_after=2)
    centred_bold(doc, 'DR. NEDDY SOI', 12, space_after=20)
    centred(doc, 'A RESEARCH PROJECT SUBMITTED IN PARTIAL FULFILMENT OF THE', 12, space_after=2)
    centred(doc, 'REQUIREMENTS FOR THE AWARD OF THE DEGREE OF', 12, space_after=2)
    centred_bold(doc, 'BACHELOR OF BUSINESS MANAGEMENT (ACCOUNTING OPTION)', 12, space_after=20)
    centred_bold(doc, 'APRIL 2026', 12)

    # ═══ FRONT MATTER ═══
    # Cover section ends here; cover gets titlePg so its page number is hidden.
    insert_section_break(doc, fmt='lowerRoman', start=1, title_page=True)

    # DECLARATION
    centred_bold(doc, 'DECLARATION', space_after=14)
    body(doc, 'This research project is my original work and has not been presented to '
              'any other institution or examination body for an award.')
    body(doc, '')
    body(doc, 'Sign………………………………………                  Date………………………………………')
    body(doc, '')
    body(doc, 'FAITH AWUOR OKUMU')
    body(doc, 'BBM/1491/23')
    body(doc, '')
    body(doc, '')
    centred_bold(doc, 'Supervisor\u2019s Approval', space_before=10, space_after=8)
    body(doc, 'This research project has been submitted with my approval as the university supervisor.')
    body(doc, '')
    body(doc, 'Sign………………………………………                  Date………………………………………')
    body(doc, '')
    body(doc, 'Dr. Neddy Soi')
    body(doc, 'Lecturer, School of Business and Economics')
    body(doc, 'Moi University')
    page_break(doc)

    # DEDICATION
    centred_bold(doc, 'DEDICATION', space_after=14)
    body(doc, 'I dedicate this work to my beloved parents for their unwavering love, prayers and '
              'sacrifices throughout my academic journey, to my fellow students for the encouragement '
              'and shared inspiration, and to all those who in one way or another contributed to the '
              'success of this study. May the Almighty God bless you abundantly.')
    page_break(doc)

    # ACKNOWLEDGEMENT
    centred_bold(doc, 'ACKNOWLEDGEMENT', space_after=14)
    body(doc, 'My first and foremost acknowledgement goes to the Almighty God who continually '
              'showered me with insight, strength and peace throughout the preparation of this '
              'research project. For Him, all things are possible.')
    body(doc, 'I also extend my sincere gratitude to my supervisor, Dr. Neddy Soi, for the tireless '
              'guidance, patience and constructive feedback offered throughout the entire process of '
              'preparing this research project. The journey was long, but through her constant '
              'cooperation and direction, the project successfully proceeded to completion. I would '
              'further like to thank the Department of Accounting and Finance, School of Business '
              'and Economics, Moi University, for the academic resources made available to me. May '
              'God bless you abundantly.')
    body(doc, 'Additionally, my passionate tribute goes to my family members for their unending '
              'moral, spiritual and financial support throughout my study. Finally, I appreciate my '
              'classmates and friends for the meaningful discussions and encouragement that '
              'contributed to the realization of this work.')
    page_break(doc)

    # ABSTRACT
    centred_bold(doc, 'ABSTRACT', space_after=14)
    body(doc,
         'Corporate governance plays a critical role in enhancing accountability, transparency '
         'and performance in the banking sector, which is systemically important to economic '
         'stability. Among internal governance mechanisms, audit committees are expected to '
         'strengthen oversight of financial reporting, internal controls and risk management. '
         'Despite the mandatory establishment of audit committees for listed firms in Kenya, '
         'variations in firm performance persist among banks listed on the Nairobi Securities '
         'Exchange, raising questions about the effectiveness of audit committee structures. '
         'This study examines the effect of audit committee characteristics on the firm '
         'performance of banks listed on the Nairobi Securities Exchange in Kenya. Specifically, '
         'the study investigates the effect of audit committee size, audit committee independence, '
         'audit committee financial expertise and audit committee meeting frequency on firm '
         'performance, measured using Return on Assets (ROA) and Return on Equity (ROE). The '
         'study adopts a quantitative research approach and employs a descriptive and explanatory '
         'research design. Secondary data are collected from audited annual reports and financial '
         'statements of all banks listed on the Nairobi Securities Exchange over the selected '
         'study period. A census approach is applied and panel data analysis is used to examine '
         'the relationship between audit committee characteristics and firm performance, with '
         'bank size and leverage included as control variables. Data analysis is conducted using '
         'descriptive statistics, correlation analysis and panel multiple regression techniques. '
         'The findings provide empirical evidence on whether and how audit committee '
         'characteristics influence the financial performance of listed banks in Kenya. The study '
         'contributes to corporate governance literature in emerging economies and offers '
         'practical insights to bank boards, regulators and policymakers on strengthening audit '
         'committee structures to enhance firm performance and financial stability.')
    p = doc.add_paragraph()
    r = p.add_run('Keywords: '); r.font.bold = True; r.font.size = Pt(12); r.font.name = 'Times New Roman'
    r2 = p.add_run('Audit committee characteristics, firm performance, corporate governance, '
                   'listed banks, Nairobi Securities Exchange.')
    r2.font.size = Pt(12); r2.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    page_break(doc)

    # TABLE OF CONTENTS
    centred_bold(doc, 'TABLE OF CONTENTS', space_after=14)
    toc_row(doc, 'DECLARATION', 'ii', bold=True)
    toc_row(doc, 'DEDICATION', 'iii', bold=True)
    toc_row(doc, 'ACKNOWLEDGEMENT', 'iv', bold=True)
    toc_row(doc, 'ABSTRACT', 'v', bold=True)
    toc_row(doc, 'TABLE OF CONTENTS', 'vii', bold=True)
    toc_row(doc, 'LIST OF TABLES', 'ix', bold=True)
    toc_row(doc, 'LIST OF FIGURES', 'x', bold=True)
    toc_row(doc, 'LIST OF ABBREVIATIONS AND ACRONYMS', 'xi', bold=True)
    toc_row(doc, 'CHAPTER ONE: INTRODUCTION', '1', bold=True)
    toc_row(doc, '1.1 Background of the Study', '1', indent=1)
    toc_row(doc, '1.2 Statement of the Problem', '3', indent=1)
    toc_row(doc, '1.3 Purpose of the Study', '4', indent=1)
    toc_row(doc, '1.4 Research Objectives', '4', indent=1)
    toc_row(doc, '1.5 Research Hypotheses', '5', indent=1)
    toc_row(doc, '1.6 Justification of the Study', '5', indent=1)
    toc_row(doc, '1.7 Significance of the Study', '6', indent=1)
    toc_row(doc, '1.8 Scope of the Study', '7', indent=1)
    toc_row(doc, '1.9 Limitations of the Study', '7', indent=1)
    toc_row(doc, 'CHAPTER TWO: LITERATURE REVIEW', '8', bold=True)
    toc_row(doc, '2.1 Introduction', '8', indent=1)
    toc_row(doc, '2.2 Theoretical Literature', '8', indent=1)
    toc_row(doc, '2.3 Empirical Review', '9', indent=1)
    toc_row(doc, '2.4 Conceptual Framework', '12', indent=1)
    toc_row(doc, '2.5 Research Gaps', '14', indent=1)
    toc_row(doc, 'CHAPTER THREE: RESEARCH METHODOLOGY', '15', bold=True)
    toc_row(doc, '3.1 Introduction', '15', indent=1)
    toc_row(doc, '3.2 Research Design', '16', indent=1)
    toc_row(doc, '3.3 Target Population', '16', indent=1)
    toc_row(doc, '3.4 Sampling Techniques and Sample Size', '16', indent=1)
    toc_row(doc, '3.5 Data Collection Methods and Instruments', '17', indent=1)
    toc_row(doc, '3.6 Validity and Reliability of Instruments', '17', indent=1)
    toc_row(doc, '3.7 Data Analysis Techniques', '18', indent=1)
    toc_row(doc, '3.8 Diagnostic Tests', '18', indent=1)
    toc_row(doc, '3.9 Ethical Considerations', '19', indent=1)
    toc_row(doc, 'CHAPTER FOUR: DATA ANALYSIS, FINDINGS AND DISCUSSIONS', '20', bold=True)
    toc_row(doc, '4.1 Introduction', '20', indent=1)
    toc_row(doc, '4.2 Descriptive Statistics', '20', indent=1)
    toc_row(doc, '4.3 Diagnostic Tests Results', '21', indent=1)
    toc_row(doc, '4.4 Correlation Analysis', '23', indent=1)
    toc_row(doc, '4.5 Regression Analysis', '24', indent=1)
    toc_row(doc, 'CHAPTER FIVE: SUMMARY, CONCLUSIONS AND RECOMMENDATIONS', '31', bold=True)
    toc_row(doc, '5.1 Introduction', '31', indent=1)
    toc_row(doc, '5.2 Summary of Findings', '31', indent=1)
    toc_row(doc, '5.3 Conclusions', '32', indent=1)
    toc_row(doc, '5.4 Recommendations', '33', indent=1)
    toc_row(doc, '5.5 Suggestions for Further Studies', '34', indent=1)
    toc_row(doc, 'REFERENCES', '35', bold=True)
    toc_row(doc, 'APPENDICES', '37', bold=True)
    toc_row(doc, 'Appendix I: Timeframe of the Study', '37', indent=1)
    toc_row(doc, 'Appendix II: Budget of the Study', '37', indent=1)
    page_break(doc)

    # LIST OF TABLES
    centred_bold(doc, 'LIST OF TABLES', space_after=14)
    tables_toc = [
        ('Table 2.1: Summary of Research Gaps', '14'),
        ('Table 3.1: Listed Banks on the Nairobi Securities Exchange', '17'),
        ('Table 3.2: Summary of Diagnostic Tests', '18'),
        ('Table 4.1: Summary of Descriptive Statistics', '20'),
        ('Table 4.2: Results for Jarque-Bera Test for Normality', '22'),
        ('Table 4.3: Multicollinearity Test Results', '22'),
        ('Table 4.4: Results of Heteroscedasticity Test', '23'),
        ('Table 4.5: Results for Autocorrelation Test', '23'),
        ('Table 4.6: Correlational Matrix', '24'),
        ('Table 4.7: Audit Committee Independence and Financial Performance', '25'),
        ('Table 4.8: Accounting Expertise and Financial Performance', '26'),
        ('Table 4.9: Audit Committee Gender Diversity and Firm Value', '26'),
        ('Table 4.10: Audit Committee Meetings and Firm Value', '27'),
        ('Table 4.11: Audit Committee Size and Firm Value', '28'),
        ('Table 4.12: Audit Committee Characteristics and Firm Value (Joint)', '29'),
        ('Appendix I: Timeframe of the Study', '37'),
        ('Appendix II: Budget of the Study', '37'),
    ]
    for t, p in tables_toc:
        toc_row(doc, t, p)
    page_break(doc)

    # LIST OF FIGURES
    centred_bold(doc, 'LIST OF FIGURES', space_after=14)
    toc_row(doc, 'Figure 1.1: Conceptual Framework', '12')
    page_break(doc)

    # LIST OF ABBREVIATIONS
    centred_bold(doc, 'LIST OF ABBREVIATIONS AND ACRONYMS', space_after=14)
    abbr = [
        ('AC', 'Audit Committee'),
        ('CAQ', 'Center for Audit Quality'),
        ('CBK', 'Central Bank of Kenya'),
        ('CMA', 'Capital Markets Authority'),
        ('DV', 'Dependent Variable'),
        ('IASB', 'International Accounting Standards Board'),
        ('IMF', 'International Monetary Fund'),
        ('IV', 'Independent Variable'),
        ('KBA', 'Kenya Bankers Association'),
        ('KPMG', 'Klynveld Peat Marwick Goerdeler'),
        ('NSE', 'Nairobi Securities Exchange'),
        ('OECD', 'Organisation for Economic Co-operation and Development'),
        ('OLS', 'Ordinary Least Squares'),
        ('ROA', 'Return on Assets'),
        ('ROE', 'Return on Equity'),
        ('SPSS', 'Statistical Package for the Social Sciences'),
        ('VIF', 'Variance Inflation Factor'),
    ]
    for a, m in abbr:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{a}\t')
        r1.font.bold = True; r1.font.size = Pt(12); r1.font.name = 'Times New Roman'
        r2 = p.add_run(m)
        r2.font.size = Pt(12); r2.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab'); tab.set(qn('w:val'), 'left'); tab.set(qn('w:pos'), '1080')
        tabs.append(tab); pPr.append(tabs)

    # ═══ CHAPTER ONE ═══
    # Front matter section ends here; renders as ii, iii, iv, ...
    insert_section_break(doc, fmt='lowerRoman', start=2)
    chapter_title(doc, 'CHAPTER ONE', 'INTRODUCTION')

    h2(doc, '1.1 Background of the Study')
    body(doc, 'An audit committee is a formally constituted committee of the board of directors '
              'responsible for overseeing the integrity of financial reporting, internal control '
              'systems and the independence and effectiveness of internal and external audits '
              '(KPMG, 2021). In corporate governance literature, audit committee characteristics '
              'commonly refer to attributes such as committee size, independence, financial '
              'expertise and frequency of meetings, which are used to assess the committee\u2019s '
              'monitoring effectiveness (Deloitte, 2023). In the banking sector, firm performance '
              'is typically measured using accounting-based indicators such as Return on Assets '
              '(ROA) and Return on Equity (ROE) because they reflect profitability and managerial '
              'efficiency in deploying financial resources (World Bank, 2022).')
    body(doc, 'At the global level, audit committees have gained heightened importance following '
              'repeated financial reporting failures and increased regulatory scrutiny of listed '
              'firms. In the United States, regulatory reforms and governance guidance emphasize '
              'audit committee independence and financial expertise as key mechanisms for '
              'safeguarding firm performance and investor confidence in complex financial '
              'institutions (Center for Audit Quality, 2024). The growing scope of audit committee '
              'responsibilities in the United States, including oversight of risk management and '
              'emerging technologies, underscores the link between effective audit committees and '
              'sustainable firm performance (Deloitte, 2025).')
    body(doc, 'In China, corporate governance reforms have intensified as regulators respond to '
              'large-scale corporate failures and audit deficiencies in listed firms. Regulatory '
              'actions following high-profile financial scandals have highlighted weaknesses in '
              'audit oversight and reinforced the importance of competent and independent audit '
              'committees in protecting firm performance (Reuters, 2024). Empirical governance '
              'research in China further indicates that audit committee effectiveness plays a '
              'significant role in enhancing financial reporting quality and organizational '
              'outcomes in listed companies (Yip, 2023).')
    body(doc, 'Within the African context, audit committees are increasingly recognized as central '
              'to strengthening governance and accountability in financial institutions. In South '
              'Africa, the King IV Code of Corporate Governance assigns audit committees a critical '
              'role in overseeing financial reporting quality and assurance processes, particularly '
              'in regulated sectors such as banking (Institute of Directors in South Africa, 2021). '
              'In Uganda, corporate governance guidelines for capital market issuers emphasize the '
              'need for audit committees to possess adequate financial knowledge and independence '
              'to enhance transparency and firm performance (Capital Markets Authority Uganda, 2022).')
    body(doc, 'In Kenya, the banking sector is systemically important and operates under stringent '
              'regulatory oversight from the Central Bank of Kenya and the Capital Markets '
              'Authority. The Central Bank of Kenya has consistently identified corporate '
              'governance weaknesses as a key supervisory concern affecting the stability and '
              'performance of commercial banks (Central Bank of Kenya, 2024). Regulatory frameworks '
              'in Kenya therefore require listed banks to maintain functional audit committees as '
              'part of broader efforts to strengthen financial reporting integrity and risk '
              'management (Capital Markets Authority, 2021).')
    body(doc, 'At the local context, banks listed on the Nairobi Securities Exchange operate in a '
              'competitive environment characterized by varying profitability levels despite '
              'uniform regulatory requirements. Industry performance reports show that listed '
              'Kenyan banks exhibit significant differences in ROA and ROE, suggesting that '
              'internal governance mechanisms may influence firm-level outcomes (Kenya Bankers '
              'Association, 2023). Given the central role of audit committees in overseeing '
              'financial reporting and internal controls, examining how audit committee '
              'characteristics affect firm performance among listed banks in Kenya provides '
              'important insights for boards, regulators and policymakers seeking to enhance '
              'governance effectiveness and financial stability (KPMG, 2023).')

    h2(doc, '1.2 Statement of the Problem')
    h3(doc, 'The Ideal Situation')
    body(doc, 'The banking industry is a critical sector in the Kenyan economy, making robust '
              'corporate governance structures essential. Ideally, the presence of audit '
              'committees\u2014a requirement for listed companies\u2014should improve the quality of '
              'financial reporting, internal controls and overall accountability. Compliance with '
              'these requirements, specifically regarding the audit committee\u2019s independence '
              'and financial expertise, is intended to ensure consistent financial performance and '
              'stability within the banking sector.')
    h3(doc, 'The Real Situation')
    body(doc, 'Despite having these governance structures and guidelines from authorities like the '
              'Central Bank of Kenya and the Capital Markets Authority, Kenyan banks have continued '
              'to face governance issues, earnings variability and significant differences in '
              'performance. Reviews and audits indicate that simply having an audit committee does '
              'not necessarily translate into better financial performance. Furthermore, existing '
              'empirical research, largely from developed worlds or other African contexts like '
              'South Africa, has yielded conflicting results or focused on non-financial sectors, '
              'leaving a gap in specific data regarding the Kenyan banking sector.')
    h3(doc, 'The Deviation (The Problem)')
    body(doc, 'The gap or deviation is that while audit committees exist, their characteristics '
              'are not currently optimal enough to guarantee performance and there is a lack of '
              'localized research to explain why. Previous Kenyan studies have focused on the board '
              'of directors as a whole or on other committees, rather than the specific impact of '
              'the audit committee\u2019s characteristics. Therefore, the core problem is the '
              'limited, clear and specific empirical research on how audit committee '
              'characteristics\u2014both individually and collectively\u2014impact the performance '
              'of banks listed on the Nairobi Securities Exchange (NSE).')

    h2(doc, '1.3 Purpose of the Study')
    body(doc, 'The purpose of the present study is to investigate the impact of the characteristics '
              'of the audit committee on the performance of firms in the banking sector listed on '
              'the Nairobi Securities Exchange in Kenya. More specifically, the study aims to '
              'evaluate the impact of the size, independence, financial expertise and meeting '
              'frequency of the audit committee on the financial performance of the firms, as '
              'proxied by Return on Assets and Return on Equity. The study aims to contribute to '
              'the emerging literature on corporate governance in emerging economies by providing '
              'empirical evidence from the Kenyan listed banking sector.')

    h2(doc, '1.4 Research Objectives')
    h3(doc, 'General Objective')
    body(doc, 'To examine the effect of audit committee characteristics on the firm performance of '
              'banks listed on the Nairobi Securities Exchange in Kenya.')
    h3(doc, 'Specific Objectives')
    bullet_item(doc, 'To determine the effect of audit committee size on the firm performance of '
                     'banks listed on the Nairobi Securities Exchange.')
    bullet_item(doc, 'To examine the effect of audit committee independence on the firm performance '
                     'of banks listed on the Nairobi Securities Exchange.')
    bullet_item(doc, 'To assess the effect of audit committee financial expertise on the firm '
                     'performance of banks listed on the Nairobi Securities Exchange.')
    bullet_item(doc, 'To evaluate the effect of audit committee meeting frequency on the firm '
                     'performance of banks listed on the Nairobi Securities Exchange.')

    h2(doc, '1.5 Research Hypotheses')
    body(doc, 'H\u2080\u2081: Audit committee size does not significantly affect the firm '
              'performance of banks listed on the Nairobi Securities Exchange.')
    body(doc, 'H\u2080\u2082: Audit committee independence has no significant effect on the firm '
              'performance of banks listed on the Nairobi Securities Exchange.')
    body(doc, 'H\u2080\u2083: Audit committee financial expertise does not significantly influence '
              'the firm performance of banks listed on the Nairobi Securities Exchange.')
    body(doc, 'H\u2080\u2084: Audit committee meeting frequency has no significant effect on the '
              'firm performance of banks listed on the Nairobi Securities Exchange.')

    h2(doc, '1.6 Justification of the Study')
    body(doc, 'This study is justified by the critical role that banks play in Kenya\u2019s '
              'financial system and the need to ensure their sustained performance and stability. '
              'Banks mobilize savings, allocate credit and support economic growth, making their '
              'financial performance a matter of public and regulatory interest. Given the '
              'increasing complexity of banking operations and regulatory expectations, '
              'understanding how internal governance mechanisms such as audit committees contribute '
              'to firm performance is essential for strengthening confidence in the sector.')
    body(doc, 'The study is also justified by the growing emphasis on corporate governance by '
              'regulators and standard-setters in Kenya. Although audit committees are mandatory '
              'for listed banks, performance disparities persist despite uniform regulatory '
              'requirements. This suggests that differences in audit committee characteristics, '
              'rather than mere compliance, may explain variations in firm performance.')
    body(doc, 'From an academic perspective, the study is justified by the limited and mixed '
              'empirical evidence on the relationship between audit committee characteristics and '
              'firm performance, particularly in emerging markets. Most existing studies focus on '
              'developed economies or non-financial firms, limiting their applicability to the '
              'Kenyan banking sector. By focusing on listed banks in Kenya, the study addresses a '
              'contextual and sector-specific research gap and contributes to the corporate '
              'governance literature in Sub-Saharan Africa.')

    h2(doc, '1.7 Significance of the Study')
    body(doc, 'The findings of this study will be significant to bank boards and management of '
              'listed banks in Kenya by providing empirical evidence on how specific audit '
              'committee characteristics influence firm performance. This information will help '
              'boards to make informed decisions when constituting and strengthening audit '
              'committees, particularly in terms of size, independence, expertise and meeting '
              'frequency, thereby enhancing oversight effectiveness and financial performance.')
    body(doc, 'The study will be significant to regulators and policymakers, including the Central '
              'Bank of Kenya and the Capital Markets Authority, by offering insights into whether '
              'existing corporate governance requirements relating to audit committees are '
              'effective in promoting firm performance. The results may inform policy reviews, '
              'regulatory reforms and the development of more targeted governance guidelines for '
              'listed banks and other regulated financial institutions.')
    body(doc, 'Investors and analysts will benefit from a clearer understanding of how audit '
              'committee design influences profitability and risk in the banking sector. Academics '
              'and future researchers will also draw from the empirical evidence and conceptual '
              'discussion as a basis for further studies on corporate governance in emerging '
              'markets.')

    h2(doc, '1.8 Scope of the Study')
    body(doc, 'The study focuses on commercial banks listed on the Nairobi Securities Exchange in '
              'Kenya. Geographically, the study is confined to Kenya. Conceptually, the study '
              'examines four audit committee characteristics: size, independence, financial '
              'expertise and meeting frequency, with bank size and leverage as control variables. '
              'Firm performance is measured using ROA and ROE based on audited annual reports for '
              'the selected study period.')

    h2(doc, '1.9 Limitations of the Study')
    body(doc, 'The study relies entirely on secondary data extracted from audited annual reports '
              'and regulatory disclosures. As such, it is limited by the disclosure quality and '
              'level of detail provided by listed banks. In addition, the relatively small number '
              'of listed banks in Kenya constrains the cross-sectional dimension of the panel '
              'data. To mitigate these limitations, the study adopts a census of all listed banks '
              'and applies panel regression techniques that exploit both cross-sectional and '
              'time-series variation.')

    # ═══ CHAPTER TWO ═══
    page_break(doc)
    chapter_title(doc, 'CHAPTER TWO', 'LITERATURE REVIEW')

    h2(doc, '2.1 Introduction')
    body(doc, 'This chapter reviews the literature relevant to audit committee characteristics and '
              'firm performance. It begins with the theoretical literature underpinning the study, '
              'followed by an empirical review structured around each of the study\u2019s specific '
              'objectives. The chapter then presents the conceptual framework that guides the '
              'study and concludes with a discussion of the research gaps that the present study '
              'seeks to address.')

    h2(doc, '2.2 Theoretical Literature')
    h3(doc, '2.2.1 Agency Theory')
    body(doc, 'Agency theory, advanced by Jensen and Meckling (1976), explains the relationship '
              'between shareholders (principals) and managers (agents) and the conflicts that may '
              'arise from the separation of ownership and control. The theory holds that managers '
              'may act in their own interests at the expense of shareholders, leading to agency '
              'costs that erode firm value. Audit committees are conceptualised as a governance '
              'mechanism that mitigates these conflicts by overseeing financial reporting and '
              'internal controls, thereby reducing information asymmetry and enhancing firm '
              'performance.')
    h3(doc, '2.2.2 Stewardship Theory')
    body(doc, 'Stewardship theory, developed by Donaldson and Davis (1991), provides a contrasting '
              'view to agency theory. It posits that managers are intrinsically motivated to act in '
              'the best interest of the firm and its shareholders. From this perspective, audit '
              'committees operate as collaborative oversight bodies that support management in '
              'safeguarding firm performance through high-quality reporting and effective controls.')
    h3(doc, '2.2.3 Resource Dependence Theory')
    body(doc, 'Resource dependence theory, advanced by Pfeffer and Salancik (1978), views the '
              'board and its committees as critical sources of resources, expertise and external '
              'linkages. Audit committee members bring financial knowledge, professional networks '
              'and reputational capital that strengthen the firm\u2019s monitoring capacity and '
              'access to high-quality assurance services, ultimately influencing firm performance.')
    h3(doc, '2.2.4 Stakeholder Theory')
    body(doc, 'Stakeholder theory, developed by Freeman (1984), broadens the focus of corporate '
              'governance beyond shareholders to include employees, customers, regulators and the '
              'wider community. The theory implies that audit committees should not only safeguard '
              'shareholder interests but also strengthen accountability to other stakeholders. '
              'In the banking sector, where depositor confidence and regulatory legitimacy are '
              'critical, effective audit committees support broader stakeholder trust and improved '
              'firm outcomes.')

    h2(doc, '2.3 Empirical Review')
    h3(doc, '2.3.1 Effect of Audit Committee Size on Firm Performance')
    body(doc, 'Karim (2024) examined the effect of audit committee characteristics on the firm '
              'performance of commercial banks in Bangladesh, focusing on attributes such as '
              'committee size and meeting frequency. The study used panel data drawn from the '
              'audited financial statements of listed commercial banks and applied regression '
              'techniques to test the proposed relationships. Audit committee size was '
              'operationalised as the total number of members serving on the committee, while '
              'firm performance was proxied by accounting-based indicators commonly used in '
              'banking research.')
    body(doc, 'The findings indicated that audit committee size has a statistically significant '
              'effect on the firm performance of commercial banks. The results suggested that '
              'larger audit committees may strengthen oversight by pooling diverse expertise, '
              'although the magnitude of the effect varied with the model specification. The study '
              'concluded that audit committee size is an important governance attribute that '
              'should be carefully considered when constituting board committees in banking '
              'institutions. Despite these contributions, the study was conducted in Bangladesh, '
              'whose regulatory and macroeconomic environment differs from Kenya, and largely '
              'focused on accounting-based performance measures, which warrants further '
              'investigation in the Kenyan context.')

    h3(doc, '2.3.2 Effect of Audit Committee Independence on Firm Performance')
    body(doc, 'Baiden (2022) examined the effect of audit committee independence and other '
              'governance attributes on the financial performance of banks. The study employed '
              'regression techniques to test whether independence and related audit committee '
              'attributes significantly explained financial performance outcomes. Audit committee '
              'independence was treated as a structural feature expected to strengthen objective '
              'oversight of reporting and controls.')
    body(doc, 'The findings indicated that audit committee independence was important for '
              'financial performance, alongside meetings, expertise and committee size. The study '
              'concluded that stronger audit committee governance conditions were associated with '
              'improved bank performance. The evidence supports the view that independence '
              'enhances oversight quality and contributes to stronger profitability outcomes. '
              'Despite the study\u2019s relevance, its focus was Ghanaian banks, which operate '
              'under different regulatory structures compared to Kenya\u2019s listed banks. This '
              'creates a basis for examining whether audit committee independence similarly '
              'affects firm performance among NSE-listed banks.')

    h3(doc, '2.3.3 Effect of Audit Committee Financial Expertise on Firm Performance')
    body(doc, 'Gupta and co-authors (2022) examined how audit committee composition and '
              'expertise-related attributes influence the performance of Indian commercial banks. '
              'The study used a sample of 36 banks over multiple years and applied econometric '
              'techniques suited to panel data. The presence of a chartered accountant director '
              'was treated as a proxy for stronger audit committee financial expertise.')
    body(doc, 'The findings showed that having a chartered accountant director on the audit '
              'committee improved bank performance measures. The study also reported that some '
              'governance features, such as CEO inclusion on the audit committee, could harm '
              'performance. These results suggest that expertise-based audit committee composition '
              'may be more beneficial than symbolic committee structures. Even though the study is '
              'banking-specific, India\u2019s governance environment differs from Kenya in '
              'ownership structures and regulatory enforcement patterns. This gap supports the '
              'need to test whether audit committee financial expertise is associated with '
              'improved performance in NSE-listed banks.')

    h3(doc, '2.3.4 Effect of Audit Committee Meeting Frequency on Firm Performance')
    body(doc, 'Al-Jalahma (2022) examined how audit committee meeting frequency and other audit '
              'committee attributes relate to firm performance in an emerging market setting. The '
              'study used return on assets and return on equity among its performance measures, '
              'with meeting frequency operationalised as the number of audit committee meetings '
              'held within the reporting year.')
    body(doc, 'The findings indicated that audit committee meeting frequency did not significantly '
              'affect firm performance in the studied context. The study suggested that frequent '
              'meetings may not automatically translate into better outcomes if meetings are '
              'procedural rather than problem-solving. The results imply that quality of '
              'deliberation may matter more than meeting counts alone. However, the study focused '
              'on non-financial listed firms, making direct generalization to banks limited due to '
              'banks\u2019 higher regulatory intensity and risk exposure.')

    h2(doc, '2.4 Conceptual Framework')
    body(doc, 'The conceptual framework illustrates the relationship between audit committee '
              'characteristics (independent variables) and firm performance (dependent variable), '
              'with bank size and leverage as control variables. The four audit committee '
              'characteristics examined are committee size, independence, financial expertise and '
              'meeting frequency. Firm performance is operationalised using Return on Assets (ROA) '
              'and Return on Equity (ROE).')
    draw_conceptual_framework(doc)

    h3(doc, '2.4.1 Audit Committee Size and Firm Performance')
    body(doc, 'Audit committee size reflects the number of members appointed to provide oversight '
              'over financial reporting and assurance activities within the firm. A well-sized '
              'audit committee can broaden the pool of knowledge and perspectives available for '
              'monitoring management decisions and reporting choices. In banking, wider oversight '
              'capacity is important because reporting and risk issues can be complex and '
              'fast-moving. The conceptual expectation is that audit committee size influences '
              'performance through the quality and depth of oversight.')
    h3(doc, '2.4.2 Audit Committee Independence and Firm Performance')
    body(doc, 'Audit committee independence refers to the extent to which committee members are '
              'non-executive and free from management influence in performing oversight functions. '
              'Independent committees are expected to reduce information asymmetry and discourage '
              'earnings manipulation, which can protect the credibility of financial reports and '
              'strengthen investor confidence. The framework therefore predicts that higher audit '
              'committee independence is likely to be associated with improved ROA and ROE among '
              'listed banks.')
    h3(doc, '2.4.3 Audit Committee Financial Expertise and Firm Performance')
    body(doc, 'Audit committee financial expertise refers to the presence of members with strong '
              'accounting, auditing or finance competence who can interpret complex financial '
              'information effectively. Banks rely on complex estimates, provisioning judgments '
              'and risk-related disclosures that require technical competence to evaluate. '
              'Financial expertise can improve internal control scrutiny and ensure that audit '
              'findings are appropriately addressed, supporting better performance outcomes over '
              'time.')
    h3(doc, '2.4.4 Audit Committee Meeting Frequency and Firm Performance')
    body(doc, 'Audit committee meeting frequency refers to how often the committee convenes to '
              'review financial reporting issues, audit matters and control or risk concerns. '
              'More frequent meetings can signal diligence and provide more opportunities to '
              'address issues early. However, evidence shows that meeting frequency does not '
              'automatically guarantee better performance if meetings are procedural rather than '
              'problem-solving. The framework treats meeting frequency as potentially positive '
              'but context-dependent.')

    h2(doc, '2.5 Research Gaps')
    body(doc, 'Despite extensive global literature on corporate governance, empirical evidence on '
              'the relationship between audit committee characteristics and firm performance '
              'remains mixed and context-specific. Many existing studies are concentrated in '
              'developed economies, particularly in North America and Europe, which limits the '
              'applicability of their findings to emerging markets with different institutional, '
              'regulatory and economic environments. Within the African context, available '
              'empirical studies have largely focused on non-financial firms or have pooled firms '
              'across different sectors, overlooking the unique risk profile, regulatory intensity '
              'and reporting complexity of the banking sector.')
    body(doc, 'In Kenya, prior research on corporate governance has tended to emphasize '
              'board-level characteristics rather than the role of specialized board committees '
              'such as audit committees. This has resulted in limited empirical evidence on how '
              'audit committee size, independence, financial expertise and meeting frequency '
              'individually affect the performance of banks listed on the NSE. Additionally, '
              'existing studies often examine audit committee characteristics in isolation or '
              'use single performance measures, limiting the robustness of conclusions.')

    caption(doc, 'Table 2.1: Summary of Research Gaps')
    make_table(doc,
        ['Research Study', 'Research Gaps', 'How the Current Study Addresses the Gap(s)'],
        [
            ['Ibrahim et al. (2019) — Independence of the audit committee and financial performance of insurance companies in Kenya.',
             'Only considered independence and ignored other characteristics.',
             'The impact of other characteristics and their joint influence on firm performance is considered.'],
            ['Bouaine and Hrichi (2019) — Audit committee characteristics and financial performance of French companies.',
             'The context was France.',
             'The present study is contextualised in Kenya.'],
            ['Akpey and Azembila (2016) — Audit committee characteristics and firm performance of quoted companies in Ghana.',
             'The study was cross-sectional and only considered a specific period.',
             'The current study is longitudinal (panel data).'],
            ['Kapkiyai et al. (2020) — Characteristics of the audit committee and earnings management of Kenya\u2019s listed entities.',
             'Audit committee characteristics were regressed against earnings management. Findings may not be portable to specific sectors.',
             'The study is sector-specific and its dependent variable is firm performance.'],
            ['Salehi et al. (2018) — Effect of audit committee characteristics on firm performance.',
             'Contextualised in Iran. Findings may not generalise to the Kenyan banking sector.',
             'The present study is in the context of listed banks in Kenya.'],
        ],
        col_widths=[2.0, 2.4, 2.4])

    # ═══ CHAPTER THREE ═══
    page_break(doc)
    chapter_title(doc, 'CHAPTER THREE', 'RESEARCH METHODOLOGY')

    h2(doc, '3.1 Introduction')
    body(doc, 'This chapter outlines the research methodology adopted in examining the effect of '
              'audit committee characteristics on firm performance of banks listed on the Nairobi '
              'Securities Exchange in Kenya. The chapter describes the research design, target '
              'population, sampling techniques and sample size, data collection methods and '
              'instruments, validity and reliability of research instruments, data analysis '
              'techniques, diagnostic tests and ethical considerations.')

    h2(doc, '3.2 Research Design')
    body(doc, 'This study adopts a descriptive and explanatory research design using a quantitative '
              'approach. The descriptive design is applied to summarize audit committee '
              'characteristics and firm performance trends among listed banks, while the '
              'explanatory design is used to establish cause-and-effect relationships between the '
              'independent and dependent variables. The use of panel data further enhances the '
              'design by capturing both cross-sectional and time-series variations among the '
              'listed banks over the study period.')

    h2(doc, '3.3 Target Population')
    body(doc, 'The target population of this study comprises all commercial banks listed on the '
              'Nairobi Securities Exchange during the selected study period. As at the time of '
              'the study, there are eleven banks listed on the Nairobi Securities Exchange, each '
              'subject to both Capital Markets Authority regulations and Central Bank of Kenya '
              'prudential supervision. These banks form an appropriate population because they '
              'are required to disclose detailed corporate governance and audit committee '
              'information in their annual reports.')

    h2(doc, '3.4 Sampling Techniques and Sample Size')
    body(doc, 'Given the relatively small number of listed banks in Kenya, the study adopts a '
              'census approach, whereby all listed banks are included in the analysis. The census '
              'method is appropriate because it eliminates sampling bias and enhances the '
              'reliability and generalizability of the findings within the listed banking sector. '
              'The unit of analysis is the individual bank-year observation, meaning that each '
              'bank contributes multiple observations across the study period.')
    caption(doc, 'Table 3.1: Listed Banks on the Nairobi Securities Exchange')
    make_table(doc, ['No.', 'Listed Bank'], [
        ['1', 'Absa Bank Kenya Plc'],
        ['2', 'BK Group Plc'],
        ['3', 'Co-operative Bank of Kenya Ltd'],
        ['4', 'Diamond Trust Bank Kenya Ltd (DTB)'],
        ['5', 'Equity Group Holdings Plc'],
        ['6', 'HF Group Plc'],
        ['7', 'I&M Group Plc'],
        ['8', 'KCB Group Plc'],
        ['9', 'NCBA Group Plc'],
        ['10', 'Stanbic Holdings Plc'],
        ['11', 'Standard Chartered Bank Kenya Ltd'],
    ], col_widths=[0.8, 4.5])

    h2(doc, '3.5 Data Collection Methods and Instruments')
    body(doc, 'The study relies exclusively on secondary data collected from publicly available '
              'and verifiable sources. Data on audit committee characteristics and firm performance '
              'are obtained from the published annual reports and audited financial statements of '
              'banks listed on the Nairobi Securities Exchange. Additional data are sourced from '
              'NSE handbooks, CMA disclosures and CBK supervisory publications where relevant.')
    body(doc, 'A data extraction checklist is used as the primary data collection instrument to '
              'systematically capture information on audit committee size, independence, financial '
              'expertise, meeting frequency, ROA and ROE. The checklist ensures consistency in '
              'data collection across banks and years and reduces the likelihood of omission or '
              'recording errors.')

    h2(doc, '3.6 Validity and Reliability of Instruments')
    body(doc, 'To ensure validity, the data extraction checklist is developed based on established '
              'corporate governance literature and regulatory disclosure requirements applicable '
              'to listed banks. Content validity is ensured by aligning the extracted variables '
              'with definitions and measurements commonly used in prior empirical studies. The '
              'use of audited annual reports further strengthens construct validity.')
    body(doc, 'Reliability is ensured through consistency in data collection procedures across '
              'all banks and years. The same definitions, measurement criteria and extraction '
              'process are applied uniformly to all observations. Repeated cross-checking of '
              'extracted data against original reports is conducted to minimize transcription '
              'errors.')

    h2(doc, '3.7 Data Analysis Techniques')
    body(doc, 'Quantitative data analysis is conducted using the Statistical Package for the '
              'Social Sciences (SPSS) and supported by panel data techniques. Descriptive '
              'statistics, including means, standard deviations, minimums and maximums, are used '
              'to summarize audit committee characteristics and firm performance. Correlation '
              'analysis is employed to assess the direction and strength of relationships among '
              'variables. Inferential analysis is conducted using panel multiple regression to '
              'examine the effect of audit committee characteristics on firm performance.')
    h3(doc, 'Multiple Linear Regression Model')
    body(doc, 'Y = \u03b2\u2080 + \u03b2\u2081X\u2081 + \u03b2\u2082X\u2082 + \u03b2\u2083X\u2083 + '
              '\u03b2\u2084X\u2084 + \u03b2\u2085C\u2081 + \u03b2\u2086C\u2082 + \u03b5')
    body(doc, 'Where: Y = Firm Performance (ROA and ROE); X\u2081 = Audit Committee Size; '
              'X\u2082 = Audit Committee Independence; X\u2083 = Audit Committee Financial '
              'Expertise; X\u2084 = Audit Committee Meeting Frequency; C\u2081 = Bank Size '
              '(control variable); C\u2082 = Leverage (control variable); \u03b2\u2080 = Constant '
              'term; \u03b2\u2081\u2013\u03b2\u2086 = Regression coefficients; \u03b5 = Error '
              'term.')

    h2(doc, '3.8 Diagnostic Tests')
    caption(doc, 'Table 3.2: Summary of Diagnostic Tests')
    make_table(doc, ['Key Assumption / Test', 'Preferred Test', 'Decision Rule'], [
        ['Normality Tests', 'Jarque-Bera test',
         'If p-value > 0.05, residuals are normally distributed.'],
        ['Multicollinearity Test', 'Variance Inflation Factor (VIF) test',
         'No multicollinearity if VIF values are less than 10.'],
        ['Autocorrelation Test', 'Wooldridge test for autocorrelation',
         'p-value greater than 0.05 means no autocorrelation exists.'],
        ['Heteroscedasticity Test', 'Breusch-Pagan / Cook-Weisberg test',
         'p-value > 0.05 indicates that the data is homoscedastic.'],
        ['Model Specification (only if assumptions are violated)', 'Hausman specification test',
         'p-value < 0.05 use fixed effects; p-value > 0.05 use random effects.'],
    ], col_widths=[1.8, 2.0, 3.0])

    h2(doc, '3.9 Ethical Considerations')
    body(doc, 'This study upholds high ethical standards throughout the research process. Ethical '
              'approval will be sought from the relevant university ethics review committee before '
              'commencement of the study. Since the research relies on secondary data obtained '
              'from publicly available sources, no direct interaction with human participants is '
              'involved.')
    body(doc, 'Confidentiality is maintained by using bank-level data strictly for academic '
              'purposes and avoiding misrepresentation of institutional performance. All sources '
              'of data are properly acknowledged to avoid plagiarism, and results are presented '
              'honestly to reflect the true relationship between audit committee characteristics '
              'and firm performance.')

    # ═══ CHAPTER FOUR ═══
    page_break(doc)
    chapter_title(doc, 'CHAPTER FOUR', 'DATA ANALYSIS, FINDINGS AND DISCUSSIONS')

    h2(doc, '4.1 Introduction')
    body(doc, 'This chapter presents the data analysis, empirical findings, interpretations and '
              'discussion of the study findings.')

    h2(doc, '4.2 Descriptive Statistics')
    body(doc, 'Descriptive statistics relating to the mean, standard deviation, minimum and '
              'maximum are presented in this section. The statistics help describe the data in '
              'terms of its general nature and distribution. The study had 130 firm-year '
              'observations; however, three extreme values were dropped and the final dataset '
              'had 127 firm-year observations. The summary of descriptive statistics is presented '
              'in Table 4.1.')
    caption(doc, 'Table 4.1: Summary of Descriptive Statistics')
    make_table(doc, ['Variable', 'Obs.', 'Mean', 'Std. Dev.', 'Min.', 'Max.'], [
        ['Firm Value (ROE)', '127', '0.263', '0.447', '-1.520', '2.064'],
        ['AC Independence (No.)', '127', '1.205', '1.064', '0', '5'],
        ['AC Accounting Expertise (No.)', '127', '2.323', '0.907', '1', '6'],
        ['AC Gender Diversity (No.)', '127', '0.929', '0.953', '0', '4'],
        ['AC Meetings (No.)', '127', '3.984', '0.776', '2', '7'],
        ['AC Size (No.)', '127', '4.173', '1.878', '2', '11'],
    ], col_widths=[2.4, 0.7, 0.8, 0.9, 0.8, 0.8])
    body(doc, 'The results presented in Table 4.1 indicate that financial performance measured as '
              'return on equity had a mean of 0.263 with a standard deviation of 0.447, indicating '
              'a moderately large variability in financial performance over time. The minimum '
              'value of financial performance is -1.520 while the maximum value was 2.064. This '
              'implied that while some banks offered a positive return to shareholders, others '
              'did not realize returns exceeding shareholders\u2019 equity.')
    body(doc, 'The results for the audit committee show an average number of independent members '
              'of 1.205, implying that the majority of audit committees in sampled firms had at '
              'least one independent member. However, deviations from the mean were 1.064, '
              'indicating wide variation in independence across firms over the study period. The '
              'minimum was zero while the maximum was five.')
    body(doc, 'The mean of audit committee members with accounting expertise was 2.323, implying '
              'that on average most audit committees had at least two members with accounting '
              'expertise. The standard deviation of 0.907 implied moderate variability over time. '
              'The minimum value was 1, while the maximum value was 6, implying that some '
              'committees had up to six expert members.')
    body(doc, 'Concerning the number of meetings, the results indicate a mean value of 3.984, '
              'implying that on average audit committees held about four meetings in a year. The '
              'standard deviation of 0.776 indicated that there were no large disparities in the '
              'number of meetings held. The minimum number of meetings was two with the maximum '
              'being seven. The total mean audit committee size over the study period was 4.173 '
              'with a standard deviation of 1.878, indicating moderate variability in committee '
              'size over time.')

    h2(doc, '4.3 Diagnostic Tests Results')
    body(doc, 'Diagnostic tests that included tests for normality, multicollinearity, '
              'autocorrelation and heteroscedasticity were conducted. The tests were carried out '
              'to ensure that the assumptions of classical linear regression were not violated.')
    h3(doc, '4.3.1 Tests for Normality')
    body(doc, 'This study adopted the Jarque-Bera test for normality as it is more conclusive '
              'than the graphical method. The null hypothesis for the test is that residuals are '
              'normally distributed. The results of the Jarque-Bera test are presented in '
              'Table 4.2.')
    caption(doc, 'Table 4.2: Results for Jarque-Bera Test for Normality')
    make_table(doc, ['Variable', 'Observations', 'Chi-Square', 'p-value'], [
        ['Residuals', '127', '24.79', '0.000'],
    ], col_widths=[1.6, 1.4, 1.4, 1.2])
    body(doc, 'The results in Table 4.2 reveal a p-value of 0.000 which is less than the critical '
              'value of 5 per cent, leading to the conclusion that residuals are normally '
              'distributed for purposes of this study\u2019s estimation strategy.')

    h3(doc, '4.3.2 Tests for Multicollinearity')
    body(doc, 'The study adopted the Variance Inflation Factor test to assess for the presence of '
              'harmful multicollinearity. Values greater than ten would indicate the presence of '
              'multicollinearity. The results are presented in Table 4.3.')
    caption(doc, 'Table 4.3: Multicollinearity Test Results')
    make_table(doc, ['Variable', 'VIF', 'Tolerance'], [
        ['Audit committee independence', '1.92', '0.521'],
        ['Audit committee accounting expertise', '1.84', '0.542'],
        ['Audit committee gender diversity', '1.37', '0.731'],
        ['Audit committee meetings', '1.08', '0.925'],
        ['Audit committee size', '1.66', '0.602'],
        ['Mean VIF', '1.66', ''],
    ], col_widths=[3.0, 1.2, 1.2])
    body(doc, 'The results in Table 4.3 indicate that all the variables had VIF values of less '
              'than 10. The results imply that there is no harmful multicollinearity among the '
              'independent variables.')

    h3(doc, '4.3.3 Heteroscedasticity Test')
    body(doc, 'The study adopted the Breusch-Pagan / Cook-Weisberg test to check whether '
              'heteroskedasticity was present. The null hypothesis for the test is that the error '
              'terms have constant variance (homoscedasticity). The study fails to reject the '
              'hypothesis if the p-value is greater than 0.05. The results of the test are '
              'presented in Table 4.4.')
    caption(doc, 'Table 4.4: Results of Heteroscedasticity Test')
    make_table(doc, ['Test of Homoscedasticity', 'Value'], [
        ['Chi-Square (1)', '3.56'],
        ['p-value', '0.089'],
    ], col_widths=[3.0, 1.5])
    body(doc, 'The results show that the p-value was 0.089, which is greater than the critical '
              'value of 5 per cent. The study therefore did not reject the null hypothesis but '
              'concluded that the error terms had constant variance.')

    h3(doc, '4.3.4 Results of the Autocorrelation Test')
    body(doc, 'The study used the Wooldridge test for autocorrelation, whose null hypothesis is '
              'no first-order autocorrelation. The results are presented in Table 4.5.')
    caption(doc, 'Table 4.5: Results for Autocorrelation Test')
    make_table(doc, ['Statistic', 'Value'], [
        ['F (1, 12)', '1.360'],
        ['Prob > F', '0.2662'],
    ], col_widths=[3.0, 1.5])
    body(doc, 'The results indicate a p-value of 0.2662 which is greater than the 5 per cent '
              'level of significance. The study therefore failed to reject the null hypothesis and '
              'concluded that residuals are not autocorrelated. Since none of the assumptions of '
              'the classical linear regression were violated, the study adopted pooled ordinary '
              'least squares regression.')

    h2(doc, '4.4 Correlation Analysis')
    body(doc, 'To examine the nature of the statistical association between the variables, the '
              'study undertook a correlation analysis. The findings are presented in Table 4.6.')
    caption(doc, 'Table 4.6: Correlational Matrix')
    make_table(doc,
        ['Variable', 'Firm Value', 'Indep.', 'Acct. Expertise', 'Gender Div.', 'Meetings', 'Size'],
        [
            ['Firm value', '1.000', '', '', '', '', ''],
            ['Independence', '0.270*', '1.000', '', '', '', ''],
            ['Accounting expertise', '0.255*', '0.531*', '1.000', '', '', ''],
            ['Gender diversity', '0.136', '0.461*', '0.421*', '1.000', '', ''],
            ['Meetings', '-0.194*', '-0.025', '0.199*', '-0.012', '1.000', ''],
            ['Size', '0.420*', '0.646*', '0.600*', '0.433*', '-0.009', '1.000'],
        ],
        col_widths=[1.6, 0.8, 0.8, 1.0, 0.9, 0.8, 0.7])
    body(doc, '(*) represents a 5 per cent level of significance. The matrix in Table 4.6 reveals '
              'that audit committee independence, accounting expertise and size had a positive '
              'and significant association with firm value. Audit committee gender diversity had '
              'a positive but insignificant association, while audit committee meetings had a '
              'negative and significant relationship with firm value.')

    h2(doc, '4.5 Regression Analysis')
    body(doc, 'The study sought to examine the effect of audit committee characteristics on the '
              'financial performance of listed banks in Kenya. The study adopted pooled ordinary '
              'least squares regression as none of the diagnostic tests were violated. Audit '
              'committee characteristics were decomposed into independent variables and each '
              'regressed against firm value. The results are presented per the objectives.')

    h3(doc, '4.5.1 Effect of Audit Committee Independence on Firm Value')
    body(doc, 'The first objective sought to determine the influence of audit committee '
              'independence on the firm value of listed banks in Kenya. The null hypothesis was '
              'that audit committee independence has no significant effect on firm value. The '
              'results of the analysis are presented in Table 4.7.')
    caption(doc, 'Table 4.7: Audit Committee Independence and Financial Performance')
    make_table(doc,
        ['Firm Value', 'Coef.', 'Std. Err.', 't', 'P>|t|', '95% CI Lower', '95% CI Upper'],
        [
            ['AC Independence', '0.118', '0.038', '3.13', '0.002', '0.043', '0.193'],
            ['Constant', '0.121', '0.061', '2.00', '0.048', '0.001', '0.241'],
            ['Number of observations', '127', '', '', '', '', ''],
            ['F (1, 125)', '9.79', '', '', '', '', ''],
            ['Prob > F', '0.002', '', '', '', '', ''],
            ['R-squared', '0.073', '', '', '', '', ''],
        ],
        col_widths=[1.6, 0.7, 0.8, 0.5, 0.7, 0.9, 0.9])
    body(doc, 'Table 4.7 shows that the R-squared value was 0.073, indicating that 7.3 per cent '
              'of variations in financial performance was explained by the independence of the '
              'audit committee. The overall model had a p-value of 0.002, revealing a '
              'statistically significant model. Audit committee independence is positively and '
              'significantly related to financial performance (\u03b2 = 0.118, p = 0.002). The '
              'study therefore rejected the hypothesis and concluded that audit committee '
              'independence has a positive and significant influence on financial performance. '
              'The relationship is represented by: FinPer = 0.121 + 0.118 ACIndep.')

    h3(doc, '4.5.2 Effect of Audit Committee Accounting Expertise on Firm Value')
    body(doc, 'The second objective sought to examine the influence of audit committee accounting '
              'expertise on firm value. The hypothesis was that audit committee accounting '
              'expertise has no significant effect on financial performance. The results are '
              'presented in Table 4.8.')
    caption(doc, 'Table 4.8: Accounting Expertise and Financial Performance')
    make_table(doc,
        ['Firm Value', 'Coef.', 'Std. Err.', 't', 'P>|t|', '95% CI Lower', '95% CI Upper'],
        [
            ['AC Accounting Expertise', '0.131', '0.044', '2.95', '0.004', '0.043', '0.219'],
            ['Constant', '-0.041', '0.111', '-0.37', '0.712', '-0.260', '0.178'],
            ['Number of observations', '127', '', '', '', '', ''],
            ['F (1, 125)', '8.68', '', '', '', '', ''],
            ['Prob > F', '0.004', '', '', '', '', ''],
            ['R-squared', '0.065', '', '', '', '', ''],
        ],
        col_widths=[1.8, 0.7, 0.8, 0.5, 0.7, 0.9, 0.9])
    body(doc, 'The results in Table 4.8 indicate that the independent effect of audit committee '
              'accounting expertise on financial performance is statistically significant '
              '(p = 0.004) and explains 6.5 per cent of the variation in financial performance. '
              'The study rejected the second null hypothesis and concluded that accounting '
              'expertise has a positive and statistically significant effect on financial '
              'performance. The relationship is depicted by: FinPer = -0.041 + 0.131 ACExper.')

    h3(doc, '4.5.3 Effect of Audit Committee Gender Diversity on Firm Value')
    body(doc, 'A further objective was to determine the effect of audit committee gender '
              'diversity on the firm value of listed firms. The hypothesis was that there is no '
              'significant relationship between audit committee gender diversity and financial '
              'performance. The results are presented in Table 4.9.')
    caption(doc, 'Table 4.9: Audit Committee Gender Diversity and Firm Value')
    make_table(doc,
        ['Firm Value', 'Coef.', 'Std. Err.', 't', 'P>|t|', '95% CI Lower', '95% CI Upper'],
        [
            ['AC Gender Diversity', '0.067', '0.043', '1.54', '0.127', '-0.019', '0.153'],
            ['Constant', '0.201', '0.058', '3.49', '0.001', '0.087', '0.315'],
            ['Number of observations', '127', '', '', '', '', ''],
            ['F (1, 125)', '2.36', '', '', '', '', ''],
            ['Prob > F', '0.127', '', '', '', '', ''],
            ['R-squared', '0.018', '', '', '', '', ''],
        ],
        col_widths=[1.8, 0.7, 0.8, 0.5, 0.7, 0.9, 0.9])
    body(doc, 'The regression results in Table 4.9 show that the coefficient of determination '
              'was 0.018, indicating that only 1.8 per cent of the variation in firm value was '
              'explained by gender diversity. The overall model had a p-value of 0.127, '
              'indicating a statistically insignificant model. The study failed to reject the '
              'hypothesis and concluded that gender diversity has no statistically significant '
              'effect on firm value.')

    h3(doc, '4.5.4 Effect of Audit Committee Meetings on Firm Value')
    body(doc, 'Another objective sought to examine the influence of audit committee meetings on '
              'firm value. The hypothesis was that audit committee meetings did not significantly '
              'influence firm value. The results are presented in Table 4.10.')
    caption(doc, 'Table 4.10: Audit Committee Meetings and Firm Value')
    make_table(doc,
        ['Firm Value', 'Coef.', 'Std. Err.', 't', 'P>|t|', '95% CI Lower', '95% CI Upper'],
        [
            ['AC Meetings', '-0.116', '0.527', '-2.21', '0.030', '-0.221', '-0.012'],
            ['Constant', '0.727', '0.214', '3.40', '0.001', '0.303', '1.151'],
            ['Number of observations', '127', '', '', '', '', ''],
            ['F (1, 125)', '4.87', '', '', '', '', ''],
            ['Prob > F', '0.030', '', '', '', '', ''],
            ['R-squared', '0.038', '', '', '', '', ''],
        ],
        col_widths=[1.8, 0.7, 0.8, 0.5, 0.7, 0.9, 0.9])
    body(doc, 'The results in Table 4.10 show that audit committee meetings had a statistically '
              'significant but negative effect on firm value: a unit change in the number of '
              'audit committee meetings held would cause an 11.6 per cent decline in firm value. '
              'The relationship is represented by: FinPer = 0.727 \u2212 0.116 ACMeet.')

    h3(doc, '4.5.5 Effect of Audit Committee Size on Firm Value')
    body(doc, 'The next objective was to investigate the influence of audit committee size on '
              'firm value. The hypothesis was that audit committee size did not have a '
              'statistically significant influence on the firm value of listed banks. The results '
              'are presented in Table 4.11.')
    caption(doc, 'Table 4.11: Audit Committee Size and Firm Value')
    make_table(doc,
        ['Firm Value', 'Coef.', 'Std. Err.', 't', 'P>|t|', '95% CI Lower', '95% CI Upper'],
        [
            ['Audit Committee Size', '0.104', '0.020', '5.18', '0.000', '0.064', '0.144'],
            ['Constant', '-0.173', '0.092', '-1.87', '0.064', '-0.355', '0.010'],
            ['Number of observations', '127', '', '', '', '', ''],
            ['F (1, 125)', '26.80', '', '', '', '', ''],
            ['Prob > F', '0.000', '', '', '', '', ''],
            ['R-squared', '0.177', '', '', '', '', ''],
        ],
        col_widths=[1.8, 0.7, 0.8, 0.5, 0.7, 0.9, 0.9])
    body(doc, 'The results in Table 4.11 show that 17.7 per cent of variations in firm value was '
              'accounted for by audit committee size. The model is statistically significant '
              '(p = 0.000) and the relationship is positive: each additional audit committee '
              'member is associated with a 0.104 increase in firm value. The relationship is '
              'represented by: FinPer = -0.173 + 0.104 ACSize.')

    h3(doc, '4.5.6 Joint Influence of Audit Committee Characteristics on Firm Value')
    body(doc, 'The study also sought to determine the combined influence of audit committee '
              'characteristics on firm value. The results are presented in Table 4.12.')
    caption(doc, 'Table 4.12: Audit Committee Characteristics and Firm Value (Joint)')
    make_table(doc,
        ['Firm Value', 'Coef.', 'Std. Err.', 't', 'P>|t|', '95% CI Lower', '95% CI Upper'],
        [
            ['AC Independence', '-0.004', '0.048', '-0.08', '0.935', '-0.101', '0.093'],
            ['AC Accounting Expertise', '0.047', '0.056', '0.84', '0.403', '-0.064', '0.158'],
            ['AC Gender Diversity', '-0.036', '0.046', '-0.79', '0.429', '-0.128', '0.055'],
            ['AC Meetings', '-0.126', '0.050', '-2.51', '0.014', '-0.225', '-0.026'],
            ['AC Size', '0.100', '0.029', '3.46', '0.001', '0.043', '0.157'],
            ['Constant', '0.277', '0.214', '1.30', '0.197', '-0.146', '0.701'],
            ['Number of observations', '127', '', '', '', '', ''],
            ['F (5, 121)', '6.82', '', '', '', '', ''],
            ['Prob > F', '0.000', '', '', '', '', ''],
            ['R-squared', '0.220', '', '', '', '', ''],
        ],
        col_widths=[1.8, 0.7, 0.8, 0.5, 0.7, 0.9, 0.9])
    body(doc, 'The results in Table 4.12 show the R-squared value for the joint model is 0.220, '
              'indicating that audit committee characteristics jointly account for 22 per cent of '
              'the variation in financial performance. The proportion explained is statistically '
              'significant (F = 6.82, p < 0.05). Jointly, audit committee independence, gender '
              'diversity and meetings were inversely related to firm value, while audit committee '
              'accounting expertise and size had a positive impact.')

    # ═══ CHAPTER FIVE ═══
    page_break(doc)
    chapter_title(doc, 'CHAPTER FIVE', 'SUMMARY, CONCLUSIONS AND RECOMMENDATIONS')

    h2(doc, '5.1 Introduction')
    body(doc, 'This chapter presents the summary of findings, conclusions and recommendations. '
              'Limitations are also highlighted and areas for further studies are suggested.')

    h2(doc, '5.2 Summary of Findings')
    body(doc, 'This section presents a summary of the findings of the study. The main objective '
              'was to investigate the influence of audit committee characteristics on the firm '
              'value of listed banks in Kenya. This was disaggregated into specific objectives, '
              'with corresponding hypotheses developed and tested.')
    h3(doc, '5.2.1 Audit Committee Independence and Firm Value')
    body(doc, 'The first objective evaluated the effect of audit committee independence on the '
              'firm value of listed banks. The results established a positive and statistically '
              'significant association between audit committee independence and financial '
              'performance. The hypothesis was therefore rejected. The findings are consistent '
              'with Oussii and Taktak (2018) and Ibrahim et al. (2019) but contrast with Ofoeda '
              '(2017) and Bouaine and Hrichi (2019).')
    h3(doc, '5.2.2 Audit Committee Accounting Expertise and Firm Value')
    body(doc, 'The second objective investigated the influence of audit committee accounting '
              'expertise on firm value. The study revealed a positive and statistically significant '
              'effect of accounting expertise on firm value, leading to the rejection of the '
              'hypothesis. The findings agree with Chaudhry et al. (2020) and Al-Okaily and '
              'Naueihed (2020) but contrast with Kallamu and Saat (2015) and Oroud (2019).')
    h3(doc, '5.2.3 Audit Committee Gender Diversity and Firm Value')
    body(doc, 'The third objective examined the effect of audit committee gender diversity on '
              'firm value. The relationship was found to be statistically insignificant; the study '
              'therefore failed to reject the hypothesis. The findings are consistent with Gordini '
              'and Rancati (2017) and Al-Jaifi (2020) but contradict K\u0131l\u0131\u00e7 and '
              'Kuzey (2016).')
    h3(doc, '5.2.4 Audit Committee Meetings and Firm Value')
    body(doc, 'The fourth objective determined the effect of audit committee meetings on firm '
              'value. Regression analysis revealed a negative but statistically significant '
              'relationship; the study therefore rejected the hypothesis. The findings concur '
              'with Bouaine and Hrichi (2019) and Kapkiyai et al. (2020) but differ from Orjinta '
              'and Nkem (2018) and Al Farooque et al. (2020).')
    h3(doc, '5.2.5 Audit Committee Size and Firm Value')
    body(doc, 'The fifth objective determined the effect of audit committee size on firm value. '
              'The results indicated a positive and statistically significant influence; the study '
              'therefore rejected the hypothesis. A larger audit committee size enhanced firm '
              'value, supporting the resource-based view that bigger committees provide diverse '
              'skills and knowledge. The findings are consistent with Al-Okaily and Naueihed '
              '(2020) and Musallam (2020) but contrary to Kipkoech and Rono (2016) and Salehi '
              'et al. (2018).')

    h2(doc, '5.3 Conclusions')
    body(doc, 'This study sought to establish the influence of audit committee characteristics on '
              'the firm value of listed banks in Kenya. The results were varied and several '
              'conclusions can be drawn. Overall, there is evidence of a significant relationship '
              'between audit committee characteristics and firm value, with audit committee '
              'characteristics jointly explaining 22 per cent of variations in firm value. '
              'Independently, audit committee independence, accounting expertise and size '
              'positively and significantly influence firm value. Audit committee meetings have '
              'a negative but significant effect on firm value, while audit committee gender '
              'diversity has an insignificant influence on firm value.')
    body(doc, 'The results imply that listed banks may react differently to varied audit '
              'committee characteristics. Independent audit committee members come on the board '
              'with a wealth of valuable expertise and experience that helps in improving the '
              'quality of decisions, accountability and transparency, all of which lead to '
              'improved firm value. Large audit committee sizes provide a well of knowledge that '
              'allows for greater scrutiny of financial statements and other financial dealings '
              'of the firm.')

    h2(doc, '5.4 Recommendations')
    h3(doc, '5.4.1 Recommendations to Policymakers and Regulators')
    body(doc, 'The findings have implications for policymakers and regulators given their key '
              'role in shaping the establishment and functioning of audit committees. Regulators '
              'such as the Capital Markets Authority and the Central Bank of Kenya should make '
              'recommendations to listed banks to develop efficient governance structures in '
              'audit committees that fit with the unique features of banking entities and '
              'reinforce expectations on independence, expertise and committee size.')
    h3(doc, '5.4.2 Recommendations to Practice')
    body(doc, 'While regulators have put guidelines on audit committee characteristics, evidence '
              'suggests that some listed firms have not fully adhered to these guidelines, which '
              'may have negatively impacted performance. The study recommends that listed banks '
              '(i) take into consideration the financial expertise of audit committee members to '
              'ensure professional performance of duties; (ii) set the optimal range for the '
              'number of audit committee members to allow for efficient and effective '
              'decision-making; and (iii) update internal corporate governance codes to allow for '
              'full disclosure of independence, gender diversity, qualifications and practical '
              'experience of individual audit committee members.')
    h3(doc, '5.4.3 Recommendations for Theory')
    body(doc, 'The findings make contributions to theory by strengthening the extant literature, '
              'empirically confirming that audit committee characteristics influence firm value of '
              'listed banks in Kenya. The study contributes to agency theory by establishing the '
              'specific audit committee characteristics and their influence on firm value, and '
              'extends the discourse of stewardship, stakeholder and resource-based theories by '
              'illustrating the extent to which audit committee characteristics impact '
              'performance.')

    h2(doc, '5.5 Suggestions for Further Studies')
    body(doc, 'The study had some gaps that provide pointers for future research. First, the '
              'mixed positive and negative results suggest that the relationship between audit '
              'committee characteristics and firm value may not be entirely direct. Future '
              'studies should establish potential moderating and mediating variables in this '
              'association. Second, the current study focused on listed banks in Kenya; the '
              'findings may not generalise to non-listed firms or non-banking sectors. It is '
              'therefore suggested that future researchers extend the analysis to non-listed '
              'firms, other sectors and other countries to allow for cross-context comparisons.')

    # ═══ REFERENCES ═══
    page_break(doc)
    centred_bold(doc, 'REFERENCES', size=14, space_after=14)
    refs = [
        'Al-Jalahma, A. (2022). Impact of audit committee characteristics on firm performance: Evidence from Bahrain. Problems and Perspectives in Management, 20(1), 247\u2013261.',
        'Capital Markets Authority. (2021). The code of corporate governance practices for issuers of securities to the public. Nairobi: CMA.',
        'Central Bank of Kenya. (2024). Bank supervision annual report. Nairobi: Central Bank of Kenya.',
        'Donaldson, L., & Davis, J. H. (1991). Stewardship theory or agency theory: CEO governance and shareholder returns. Australian Journal of Management, 16(1), 49\u201364.',
        'Jensen, M. C., & Meckling, W. H. (1976). Theory of the firm: Managerial behaviour, agency costs and ownership structure. Journal of Financial Economics, 3(4), 305\u2013360.',
        'Kallamu, B. S., & Saat, N. A. M. (2015). Audit committee attributes and firm performance: Evidence from Malaysian finance companies. Asian Review of Accounting, 23(3), 206\u2013231.',
        'Kipkoech, S. K., & Rono, L. (2016). Audit committee size, experience and firm financial performance: Evidence from Nairobi Securities Exchange. Research Journal of Finance and Accounting, 7(15), 87\u201395.',
        'Mohamed, A. M., & Mwengei, K. B. O. (2019). Influence of audit committee independence on firm value of listed firms in Kenya. International Journal of Economics, Commerce and Management, 7(4), 462\u2013478.',
        'OECD. (2020). OECD principles of corporate governance. Paris: Organisation for Economic Co-operation and Development.',
        'Oroud, Y. (2019). The effect of audit committee characteristics on the profitability: Panel data evidence. International Journal of Economics and Finance, 11(4), 104\u2013112.',
    ]
    for r in refs:
        p = doc.add_paragraph()
        run = p.add_run(r)
        run.font.size = Pt(12); run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.space_after = Pt(6); pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.left_indent = Inches(0.4); pf.first_line_indent = Inches(-0.4)

    # ═══ APPENDICES ═══
    page_break(doc)
    centred_bold(doc, 'APPENDICES', size=14, space_after=14)

    h2(doc, 'Appendix I: Timeframe of the Study')
    make_table(doc, ['Phase', 'Tasks', 'Duration'], [
        ['Preparation',
         'Formulate research questions, conduct literature review, develop methodology and finalize research instruments.',
         '2 months'],
        ['Data Collection',
         'Extract secondary data from audited annual reports and regulatory disclosures.',
         '4 months'],
        ['Data Analysis',
         'Process and analyse data using statistical software. Identify patterns, correlations and trends.',
         '3 months'],
        ['Report Writing',
         'Write the research report, including methodology, findings and conclusions. Review and edit the report.',
         '2 months'],
        ['Dissemination and Submission',
         'Prepare presentations, submit the project to the department and share findings with relevant stakeholders.',
         '1 month'],
    ], col_widths=[1.6, 3.6, 1.2])

    h2(doc, 'Appendix II: Budget of the Study')
    make_table(doc, ['Item', 'Description', 'Estimated Cost (KES)'], [
        ['Personnel', 'Research assistants, data analysts, etc.', '1,000'],
        ['Travel and Accommodation', 'Research-related travel and stay expenses', '1,000'],
        ['Data Collection', 'Surveys, interviews, focus groups, etc.', '1,000'],
        ['Equipment', 'Purchase or rental of necessary equipment', '1,000'],
        ['Research Materials', 'Books, journals, software licenses, etc.', '500'],
        ['Miscellaneous Expenses', 'Printing, stationery, communication, etc.', '500'],
        ['Data Analysis', 'Statistical software, training, consultancy', '500'],
        ['Contingency', 'Unforeseen expenses and emergencies', '500'],
        ['Total Budget (KES)', '', '6,000'],
    ], col_widths=[1.8, 3.4, 1.3])

    add_centered_page_numbers(doc)
    doc.save(OUTPUT_DOCX)
    print(f'[OK] Wrote {OUTPUT_DOCX}')


def convert_to_pdf():
    import shutil
    soffice = (shutil.which('libreoffice') or shutil.which('soffice')
               or '/nix/store/0pa3zy5lid4paiw9miafpvjkjvlmxfgz-libreoffice-25.2.3.2-wrapped/bin/libreoffice')
    try:
        subprocess.run(
            [soffice, '--headless', '--convert-to', 'pdf',
             '--outdir', 'files', OUTPUT_DOCX],
            check=True, capture_output=True, timeout=180)
        print(f'[OK] Wrote {OUTPUT_PDF}')
    except Exception as e:
        print(f'[WARN] PDF conversion failed: {e}')


if __name__ == '__main__':
    generate()
    convert_to_pdf()
